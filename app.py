import os
import re
import json
import shutil
import unicodedata
from datetime import date, datetime, timedelta
import platform
from pathlib import Path
from urllib.parse import quote

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageOps
try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
except Exception:
    pass

try:
    from streamlit_drawable_canvas import st_canvas
except Exception:
    st_canvas = None

# =========================================
# INTEGRAÇÃO GOOGLE SHEETS
# =========================================

SHEETS_ID = "1uvZ6qfYCYFl_feGGgvIIXMQlUWvx0MTzTuC8TwfPBlM"
DRIVE_FOTOS_FOLDER_ID = "1KNtPKvLl_NJw-Vm_26ABxc4LG3CiZqDR"


def conectar_drive():
    """Conecta ao Google Drive usando as mesmas credenciais do Sheets."""
    try:
        from googleapiclient.discovery import build
        from google.oauth2.service_account import Credentials

        SCOPES = [
            "https://www.googleapis.com/auth/drive",
            "https://www.googleapis.com/auth/spreadsheets",
        ]

        if "gcp_service_account" in st.secrets:
            creds_dict = dict(st.secrets["gcp_service_account"])
            creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
        else:
            creds_path = Path(__file__).parent / "aqua-gestao-rt-87316ebf5331.json"
            if not creds_path.exists():
                return None
            creds = Credentials.from_service_account_file(str(creds_path), scopes=SCOPES)

        service = build("drive", "v3", credentials=creds)
        return service
    except Exception as e:
        _log_sheets_erro("conectar_drive", e)
        return None


def drive_criar_pasta(nome_pasta: str, pasta_pai_id: str) -> str | None:
    """Cria pasta no Drive se não existir. Retorna o ID da pasta."""
    try:
        service = conectar_drive()
        if not service:
            return None

        # Verifica se já existe
        query = f"name='{nome_pasta}' and '{pasta_pai_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        resultado = service.files().list(q=query, fields="files(id,name)").execute()
        arquivos = resultado.get("files", [])
        if arquivos:
            return arquivos[0]["id"]

        # Cria nova
        meta = {
            "name": nome_pasta,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [pasta_pai_id],
        }
        pasta = service.files().create(body=meta, fields="id").execute()
        return pasta.get("id")
    except Exception as e:
        _log_sheets_erro("drive_criar_pasta", e)
        return None


def drive_upload_foto(arquivo_bytes: bytes, nome_arquivo: str, nome_condominio: str, mes_ano: str = None) -> str | None:
    """Faz upload de foto para o Drive. Retorna o ID do arquivo."""
    try:
        import io
        from googleapiclient.http import MediaIoBaseUpload

        service = conectar_drive()
        if not service:
            return None

        if not mes_ano:
            mes_ano = datetime.now().strftime("%Y-%m")

        # Estrutura: Aqua Gestão – Fotos / Condomínio / Ano-Mês
        pasta_cond = drive_criar_pasta(nome_condominio, DRIVE_FOTOS_FOLDER_ID)
        if not pasta_cond:
            return None
        pasta_mes = drive_criar_pasta(mes_ano, pasta_cond)
        if not pasta_mes:
            return None

        # Detecta tipo
        ext = nome_arquivo.lower().split(".")[-1]
        mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "webp": "image/webp", "heic": "image/heic"}
        mime = mime_map.get(ext, "image/jpeg")

        meta = {"name": nome_arquivo, "parents": [pasta_mes]}
        media = MediaIoBaseUpload(io.BytesIO(arquivo_bytes), mimetype=mime)
        arquivo = service.files().create(body=meta, media_body=media, fields="id,webViewLink").execute()
        return arquivo.get("id")
    except Exception as e:
        _log_sheets_erro("drive_upload_foto", e)
        return None


def drive_baixar_foto(file_id: str) -> bytes | None:
    """Baixa foto do Drive pelo ID. Retorna bytes da imagem."""
    try:
        service = conectar_drive()
        if not service:
            return None
        conteudo = service.files().get_media(fileId=file_id).execute()
        return conteudo
    except Exception as e:
        _log_sheets_erro("drive_baixar_foto", e)
        return None


# =========================================
# GESTÃO DE OPERADORES
# =========================================

def _normalizar_chave_acesso(texto: str) -> str:
    """Normaliza nomes para comparação exata de PINs, operadores e condomínios."""
    texto = re.sub(r"\s+", " ", str(texto or "").strip())
    return texto.casefold()


def normalizar_texto_busca(valor: str) -> str:
    """Normaliza texto para comparação robusta: remove acento, caixa, espaços e símbolos."""
    valor = str(valor or "").strip().lower()
    valor = unicodedata.normalize("NFKD", valor)
    valor = "".join(c for c in valor if not unicodedata.combining(c))
    valor = re.sub(r"[^a-z0-9]+", " ", valor)
    valor = re.sub(r"\s+", " ", valor).strip()
    return valor


def nomes_condominio_equivalentes(a: str, b: str) -> bool:
    """Compara nomes de condomínios tolerando acentos, espaços e pequenas variações."""
    na = normalizar_texto_busca(a)
    nb = normalizar_texto_busca(b)
    if not na or not nb:
        return False
    return na == nb or na in nb or nb in na


def normalizar_data_visita(valor) -> str:
    """Converte datas como 17/04/26, 170426, 2026-04-17 para dd/mm/aaaa."""
    texto = str(valor or "").strip()
    if not texto:
        return ""

    formatos = [
        "%d/%m/%Y", "%d/%m/%y",
        "%d-%m-%Y", "%d-%m-%y",
        "%Y-%m-%d",
        "%d%m%Y", "%d%m%y",
    ]
    for fmt in formatos:
        try:
            dt = datetime.strptime(texto, fmt)
            return dt.strftime("%d/%m/%Y")
        except Exception:
            pass

    digitos = re.sub(r"\D", "", texto)
    if len(digitos) == 6:
        try:
            dt = datetime.strptime(digitos, "%d%m%y")
            return dt.strftime("%d/%m/%Y")
        except Exception:
            pass
    if len(digitos) == 8:
        for fmt in ("%d%m%Y", "%Y%m%d"):
            try:
                dt = datetime.strptime(digitos, fmt)
                return dt.strftime("%d/%m/%Y")
            except Exception:
                pass

    return texto


def lancamento_pertence_mes_ano(data_lancamento: str, mes: str, ano: str) -> bool:
    """Confere se uma visita pertence ao mês/ano do relatório."""
    data_norm = normalizar_data_visita(data_lancamento)
    try:
        dt = datetime.strptime(data_norm, "%d/%m/%Y")
        mes_int = int(str(mes).zfill(2))
        ano_int = int(str(ano))
    except Exception:
        return False
    return dt.month == mes_int and dt.year == ano_int


def _montar_resumo_dosagens_lancamento(lancamento: dict) -> str:
    """Monta texto resumido das dosagens para gravar em planilha."""
    dosagens = lancamento.get("dosagens", []) or []
    partes = []
    for d in dosagens:
        if not isinstance(d, dict):
            continue
        produto = str(d.get("produto", "")).strip()
        qtd = str(d.get("quantidade", "")).strip()
        unid = str(d.get("unidade", "")).strip()
        finalidade = str(d.get("finalidade", "")).strip()
        if produto or qtd or finalidade:
            partes.append(f"{produto} {qtd}{unid} - {finalidade}".strip(" -"))
    return " | ".join(partes)


def _condominios_organizar(condominios: list[str] | None) -> list[str]:
    """Limpa, deduplica e preserva a ordem dos condomínios informados."""
    resultado = []
    vistos = set()
    for item in condominios or []:
        valor = re.sub(r"\s+", " ", str(item or "").strip())
        if not valor:
            continue
        chave = _normalizar_chave_acesso(valor)
        if chave in vistos:
            continue
        vistos.add(chave)
        resultado.append(valor)
    return resultado


def _resolver_condominios_permitidos_exatos(condominios_permitidos: list[str], todos_condominios: list[str]) -> list[str]:
    """Resolve permissões usando comparação tolerante a acentos e variações.

    Ex.: Triad, Tríad, Tríad Vertical e Condomínio Tríad podem casar com o
    nome oficial salvo no cadastro. Mantém o nome oficial disponível no sistema.
    """
    permitidos = []
    vistos = set()
    disponiveis = _condominios_organizar(todos_condominios or [])

    for nome_perm in _condominios_organizar(condominios_permitidos or []):
        if _normalizar_chave_acesso(nome_perm) == "todos":
            continue

        escolhido = None
        # 1) Tentativa exata normalizada
        chave_perm = _normalizar_chave_acesso(nome_perm)
        for nome_disp in disponiveis:
            if _normalizar_chave_acesso(nome_disp) == chave_perm:
                escolhido = nome_disp
                break

        # 2) Tentativa tolerante: remove acentos, aceita abreviações e substrings
        if escolhido is None:
            for nome_disp in disponiveis:
                if nomes_condominio_equivalentes(nome_perm, nome_disp):
                    escolhido = nome_disp
                    break

        if escolhido:
            chave_final = _normalizar_chave_acesso(escolhido)
            if chave_final not in vistos:
                vistos.add(chave_final)
                permitidos.append(escolhido)

    return permitidos


def _pin_operador_em_uso(pin: str, nome_ignorar: str = "") -> bool:
    """Verifica se o PIN já está em uso por outro operador."""
    pin_limpo = str(pin or "").strip()
    nome_ignorar_norm = _normalizar_chave_acesso(nome_ignorar)
    if not pin_limpo:
        return False

    for op in sheets_listar_operadores():
        if str(op.get("pin", "")).strip() == pin_limpo:
            if _normalizar_chave_acesso(op.get("nome", "")) != nome_ignorar_norm:
                return True

    for op in carregar_operadores():
        if str(op.get("pin", "")).strip() == pin_limpo:
            if _normalizar_chave_acesso(op.get("nome", "")) != nome_ignorar_norm:
                return True
    return False


@st.cache_data(ttl=45, show_spinner=False)
def sheets_listar_operadores() -> list[dict]:
    """Lista operadores da aba 👷 Operadores do Sheets."""
    try:
        sh = conectar_sheets()
        if sh is None:
            return []
        try:
            aba = obter_aba_sheets("👷 Operadores")
        except Exception:
            return []
        todos = aba.get_all_values()
        operadores = []
        for row in todos:
            if len(row) >= 4 and str(row[0]).strip() and str(row[0]).strip() != "Nome":
                nome = re.sub(r"\s+", " ", str(row[0]).strip())
                pin  = str(row[1]).strip()
                conds_raw = str(row[2]).strip()
                ativo = str(row[3]).strip().lower() in ("sim", "ativo", "1", "true", "yes")
                conds = _condominios_organizar(conds_raw.split("|")) if conds_raw else []
                acesso_total = any(_normalizar_chave_acesso(c) == "todos" for c in conds) or not conds
                operadores.append({
                    "nome": nome,
                    "pin": pin,
                    "condomínios": conds,
                    "ativo": ativo,
                    "acesso_total": acesso_total,
                })
        return operadores
    except Exception as e:
        _log_sheets_erro("sheets_listar_operadores", e)
        return []


def sheets_criar_aba_operadores():
    """Cria a aba 👷 Operadores no Sheets se não existir."""
    try:
        sh = conectar_sheets()
        if sh is None:
            return False
        try:
            obter_aba_sheets("👷 Operadores")
            return True  # já existe
        except Exception:
            pass
        aba = sh.add_worksheet(title="👷 Operadores", rows=100, cols=6)
        # Cabeçalho
        aba.update("A1:F1", [["Nome", "PIN", "Condomínios (separados por |)", "Ativo", "Cadastrado_em", "Obs"]])
        aba.format("A1:F1", {"textFormat": {"bold": True}, "backgroundColor": {"red": 0.07, "green": 0.16, "blue": 0.46}})
        return True
    except Exception as e:
        _log_sheets_erro("sheets_criar_aba_operadores", e)
        return False


def sheets_salvar_operador(nome: str, pin: str, condomínios: list, ativo: bool = True) -> bool:
    """Salva ou atualiza operador na aba 👷 Operadores."""
    try:
        nome_limpo = re.sub(r"\s+", " ", str(nome or "").strip())
        pin_limpo = str(pin or "").strip()
        conds_limpos = _condominios_organizar(condomínios)

        if not nome_limpo or not pin_limpo:
            st.session_state["_operadores_erro"] = "Nome e PIN são obrigatórios."
            return False

        if _pin_operador_em_uso(pin_limpo, nome_ignorar=nome_limpo):
            st.session_state["_operadores_erro"] = f"O PIN {pin_limpo} já está em uso por outro operador."
            return False

        sh = conectar_sheets()
        if sh is None:
            return False
        sheets_criar_aba_operadores()
        aba = obter_aba_sheets("👷 Operadores")
        todos = aba.get_all_values()
        conds_str = " | ".join(conds_limpos)
        ativo_str = "Sim" if ativo else "Não"
        nova_linha = [nome_limpo, pin_limpo, conds_str, ativo_str, datetime.now().strftime("%Y-%m-%d"), ""]
        # Verifica se já existe (pelo nome)
        for i, row in enumerate(todos):
            if len(row) > 0 and _normalizar_chave_acesso(row[0]) == _normalizar_chave_acesso(nome_limpo):
                aba.update(f"A{i+1}:F{i+1}", [nova_linha])
                st.cache_data.clear()
                st.session_state.pop("_operadores_erro", None)
                return True
        # Insere novo
        linha_destino = max(len(todos) + 1, 8)
        aba.update(
            f"A{linha_destino}:Z{linha_destino}",
            [nova_linha],
            value_input_option="RAW"
        )
        st.cache_data.clear()
        st.session_state.pop("_operadores_erro", None)
        return True
    except Exception as e:
        _log_sheets_erro("sheets_salvar_operador", e)
        return False


def sheets_deletar_operador(nome: str) -> bool:
    """Remove operador da aba 👷 Operadores."""
    try:
        sh = conectar_sheets()
        if sh is None:
            return False
        aba = obter_aba_sheets("👷 Operadores")
        todos = aba.get_all_values()
        for i, row in enumerate(todos):
            if len(row) > 0 and _normalizar_chave_acesso(row[0]) == _normalizar_chave_acesso(nome):
                aba.delete_rows(i + 1)
                st.cache_data.clear()
                return True
        return False
    except Exception as e:
        _log_sheets_erro("sheets_deletar_operador", e)
        return False


def verificar_pin_operador(pin_digitado: str) -> dict | None:
    """Verifica PIN e retorna dados do operador, ou None se inválido."""
    return validar_pin_operador(pin_digitado)

def _log_sheets_erro(contexto: str, erro: Exception):
    """Armazena o último erro do Google Sheets no session_state para diagnóstico."""
    import traceback
    msg = f"[{contexto}] {type(erro).__name__}: {erro}\n{traceback.format_exc()}"
    st.session_state["_sheets_ultimo_erro"] = msg


@st.cache_resource(ttl=3600, show_spinner=False)
def conectar_sheets():
    """Conecta ao Google Sheets usando as credenciais do st.secrets ou arquivo local."""
    try:
        import gspread
        from google.oauth2.service_account import Credentials
    except ImportError as e:
        _log_sheets_erro("conectar_sheets/import", e)
        st.session_state["_sheets_ultimo_erro"] = (
            "ERRO: gspread ou google-auth não está instalado no ambiente atual.\n"
            "Verifique o requirements.txt e force um redeploy no Streamlit Cloud."
        )
        return None

    try:
        SCOPES = [
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive",
        ]

        # Tenta carregar do st.secrets (Streamlit Cloud)
        if "gcp_service_account" in st.secrets:
            creds_dict = dict(st.secrets["gcp_service_account"])
            creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
        else:
            # Fallback: arquivo local (uso no computador)
            creds_path = Path(__file__).parent / "aqua-gestao-rt-87316ebf5331.json"
            if not creds_path.exists():
                st.session_state["_sheets_ultimo_erro"] = (
                    "ERRO: Nenhuma credencial encontrada.\n"
                    "No Streamlit Cloud: verifique st.secrets['gcp_service_account'].\n"
                    "Localmente: arquivo aqua-gestao-rt-87316ebf5331.json não encontrado."
                )
                return None
            creds = Credentials.from_service_account_file(str(creds_path), scopes=SCOPES)

        gc = gspread.authorize(creds)
        sh = gc.open_by_key(SHEETS_ID)
        # Limpa erro anterior se conexão ok
        st.session_state.pop("_sheets_ultimo_erro", None)
        return sh
    except Exception as e:
        _log_sheets_erro("conectar_sheets", e)
        return None


@st.cache_resource(ttl=1800, show_spinner=False)
def obter_aba_sheets(nome_aba: str):
    """Retorna uma worksheet do Google Sheets com cache de recurso.

    Correção v4: a versão anterior chamava obter_aba_sheets(nome_aba)
    dentro dela mesma, gerando recursão e falhas silenciosas nas leituras
    do Google Sheets. Aqui a função passa a buscar a aba real no objeto
    da planilha retornado por conectar_sheets().
    """
    try:
        sh = conectar_sheets()
        if sh is None:
            return None
        return sh.worksheet(nome_aba)
    except Exception as e:
        _log_sheets_erro(f"obter_aba_sheets/{nome_aba}", e)
        return None


def limpar_payload_para_sheets(dados: dict) -> dict:
    """Remove campos pesados antes de salvar payload no Google Sheets.

    O Google Sheets limita cada célula a 50.000 caracteres.
    Fotos/base64/assinaturas/canvas não devem ser salvos em célula.
    """
    if not isinstance(dados, dict):
        return {}

    proibidos = {
        "foto", "fotos", "imagem", "imagens", "image", "images",
        "base64", "bytes", "arquivo", "arquivos", "upload",
        "assinatura", "assinatura_base64", "signature",
        "canvas", "drawing", "pdf", "html"
    }

    limpo = {}

    for k, v in dados.items():
        chave = str(k).lower()

        if any(p in chave for p in proibidos):
            continue

        if isinstance(v, dict):
            limpo[k] = limpar_payload_para_sheets(v)

        elif isinstance(v, list):
            nova_lista = []
            for item in v:
                if isinstance(item, dict):
                    nova_lista.append(limpar_payload_para_sheets(item))
                else:
                    texto = str(item)
                    if len(texto) <= 1000:
                        nova_lista.append(item)
            limpo[k] = nova_lista

        else:
            texto = str(v)
            if len(texto) > 1000:
                limpo[k] = texto[:1000] + "..."
            else:
                limpo[k] = v

    return limpo


def sheets_salvar_lancamento_campo(lancamento: dict, nome_condominio: str):
    """Salva lançamento de campo na aba 🔬 Visitas do Google Sheets.

    Esta função é crítica para o relatório mensal: PDF de visita não alimenta relatório
    se a linha não for gravada no Sheets. Por isso, grava dados básicos e também um
    Payload JSON completo para preservar múltiplas piscinas, fotos e dosagens.
    """
    try:
        sh = conectar_sheets()
        if sh is None:
            return False

        try:
            aba = obter_aba_sheets("🔬 Visitas")
        except Exception:
            aba = sh.add_worksheet(title="🔬 Visitas", rows=1000, cols=26)
            aba.update(
                "A1:Z1",
                [[
                    "", "ID Visita", "Data", "ID Cliente", "Condomínio",
                    "pH", "CRL", "CT", "Alcalinidade", "Dureza", "CYA",
                    "Foto Antes", "Foto Depois", "Foto Casa Máquinas", "Observação",
                    "Dosagem Cloro", "Dosagem Bicarb", "Alerta pH", "Alerta Cloro",
                    "Status", "Operador", "Problemas", "Payload JSON", "Salvo em",
                    "Fonte", "Mês/Ano",
                ]]
            )

        todos = aba.get_all_values()
        visitas_existentes = [
            r for r in todos
            if len(r) > 1 and str(r[1]).strip().startswith("V")
        ]
        proximo_num = len(visitas_existentes) + 1
        id_visita = f"V{proximo_num:05d}"

        data_normalizada = normalizar_data_visita(lancamento.get("data", ""))
        nome_condominio = str(nome_condominio or lancamento.get("condominio", "")).strip()

        # Busca ID do cliente por nome normalizado
        id_cliente = ""
        try:
            aba_clientes = obter_aba_sheets("👥 Clientes")
            clientes = aba_clientes.get_all_values()
            for row in clientes:
                if len(row) > 2 and nomes_condominio_equivalentes(nome_condominio, row[2]):
                    id_cliente = row[1]
                    break
        except Exception:
            id_cliente = ""

        # Se houver múltiplas piscinas, grava os parâmetros principais com a primeira piscina,
        # mas conserva tudo no Payload JSON.
        piscinas = lancamento.get("piscinas", []) or []
        base_param = piscinas[0] if isinstance(piscinas, list) and piscinas and isinstance(piscinas[0], dict) else lancamento

        dosagem_txt = _montar_resumo_dosagens_lancamento(lancamento)

        payload = dict(lancamento)
        payload["data"] = data_normalizada
        payload["condominio"] = nome_condominio
        payload["id_visita"] = id_visita
        payload["status"] = payload.get("status", "Concluída")
        try:
            payload_json = json.dumps(payload, ensure_ascii=False)
        except Exception:
            payload_json = ""

        mes_ano = ""
        try:
            dt = datetime.strptime(data_normalizada, "%d/%m/%Y")
            mes_ano = dt.strftime("%m/%Y")
        except Exception:
            pass

        nova_linha = [
            "",                                      # A
            id_visita,                               # B
            data_normalizada,                        # C
            id_cliente,                              # D
            nome_condominio,                         # E
            base_param.get("ph", lancamento.get("ph", "")),
            base_param.get("cloro_livre", lancamento.get("cloro_livre", "")),
            base_param.get("cloro_total", lancamento.get("cloro_total", "")),
            base_param.get("alcalinidade", lancamento.get("alcalinidade", "")),
            base_param.get("dureza", lancamento.get("dureza", "")),
            base_param.get("cianurico", lancamento.get("cianurico", "")),
            "",                                      # L foto antes
            "",                                      # M foto depois
            "",                                      # N foto casa máquinas
            lancamento.get("observacao", ""),        # O
            dosagem_txt,                              # P
            "", "", "",                              # Q/R/S
            "Concluída",                             # T
            lancamento.get("operador", ""),          # U
            lancamento.get("problemas", ""),         # V
            payload_json,                             # W
            _agora_brasilia(),
            "Modo Campo",                             # Y
            mes_ano,                                  # Z
        ]

        linha_destino = max(len(todos) + 1, 8)
        aba.update(
            f"A{linha_destino}:Z{linha_destino}",
            [nova_linha],
            value_input_option="RAW"
        )
        st.cache_data.clear()
        return True

    except Exception as e:
        _log_sheets_erro("sheets_salvar_lancamento_campo", e)
        return False


def sheets_salvar_cliente(nome: str, cnpj: str, endereco: str, contato: str, telefone: str,
                           vol_adulto: float = 0, vol_infantil: float = 0, vol_family: float = 0,
                           empresa: str = "Aqua Gestão"):
    """Salva novo cliente na aba Clientes do Google Sheets.
    
    Insere sempre logo após o último cliente real (C001, C002...),
    mantendo a formatação da planilha com cabeçalho e linha de total intactos.
    Colunas J/K/L = Vol_Adulto_m3, Vol_Infantil_m3, Vol_Family_m3
    """
    try:
        import re as _re

        sh = conectar_sheets()
        if sh is None:
            return False

        aba = obter_aba_sheets("👥 Clientes")
        todos = aba.get_all_values()

        # Identifica linhas reais de clientes: col B começa com C + dígitos
        ultima_linha_cliente = 0  # índice 0-based
        nums = []
        nomes_existentes = []

        for i, row in enumerate(todos):
            if len(row) > 2:
                id_val = str(row[1]).strip()
                nome_val = str(row[2]).strip()
                if _re.match(r"^C[0-9]+$", id_val) and nome_val:
                    ultima_linha_cliente = i
                    nomes_existentes.append(nome_val.lower())
                    m = _re.match(r"^C([0-9]+)$", id_val)
                    if m:
                        nums.append(int(m.group(1)))

        # Verifica duplicata
        if nome.lower().strip() in nomes_existentes:
            return True  # Já existe, considera sucesso

        # Próximo ID baseado no maior existente
        proximo_num = (max(nums) + 1) if nums else 1
        id_cliente = f"C{proximo_num:03d}"

        vol_total = (vol_adulto or 0) + (vol_infantil or 0) + (vol_family or 0)
        nova_linha = [
            "",                                    # A - vazia
            id_cliente,                            # B - ID
            nome,                                  # C - Nome
            str(vol_total) if vol_total else "",   # D - Volume total m3
            formatar_telefone(telefone),           # E - Telefone
            contato,                               # F - Contato síndico
            endereco,                              # G - Endereço
            datetime.now().strftime("%Y-%m-%d"),   # H - Data cadastro
            "Ativo",                               # I - Status
            str(vol_adulto) if vol_adulto else "", # J - Vol Adulto m3
            str(vol_infantil) if vol_infantil else "", # K - Vol Infantil m3
            str(vol_family) if vol_family else "", # L - Vol Family m3
            empresa,                               # M - Empresa
            cnpj,                                  # N - CNPJ # _CNPJ_COLUNA_N_
        ]

        # Determina posicao alfabetica dentro do bloco de clientes
        nome_novo_lower = nome.lower().strip()
        linha_insercao = ultima_linha_cliente + 2  # fallback: final do bloco
        for i, row in enumerate(todos):
            if len(row) > 2:
                id_val2 = str(row[1]).strip()
                nome_val2 = str(row[2]).strip()
                if _re.match(r"^C[0-9]+$", id_val2) and nome_val2:
                    if nome_val2.lower().strip() > nome_novo_lower:
                        linha_insercao = i + 1
                        break
        aba.insert_row(nova_linha, linha_insercao, value_input_option="USER_ENTERED")
        st.cache_data.clear()
        return True
    except Exception as e:
        _log_sheets_erro("sheets_salvar_cliente", e)
        return False


@st.cache_data(ttl=45, show_spinner=False)
def sheets_listar_clientes() -> list[str]:
    """Retorna lista de nomes de clientes da aba Clientes."""
    try:
        sh = conectar_sheets()
        if sh is None:
            return []
        aba = obter_aba_sheets("👥 Clientes")
        todos = aba.get_all_values()
        nomes = []
        for row in todos:
            if len(row) > 2 and str(row[1]).startswith("C") and row[2].strip():
                nomes.append(row[2].strip())
        return nomes
    except Exception:
        return []


@st.cache_data(ttl=45, show_spinner=False)
def sheets_listar_clientes_completo() -> list[dict]:
    """Retorna lista de dicts com dados completos de cada cliente do Sheets.
    
    Mapeamento das colunas da planilha:
      B=ID, C=Nome, D=Volume_m3, E=Contato_Sindico/Telefone, 
      F=Email_Sindico, G=Endereco, H=Data_Cadastro, I=Status
    """
    import re as _re
    try:
        sh = conectar_sheets()
        if sh is None:
            return []
        aba = obter_aba_sheets("👥 Clientes")
        todos = aba.get_all_values()
        clientes = []
        for row in todos:
            if not row or len(row) < 3:
                continue
            id_val = str(row[1]).strip() if len(row) > 1 else ""
            if not _re.match(r"^C[0-9]+$", id_val):
                continue
            nome = str(row[2]).strip() if len(row) > 2 else ""
            if not nome:
                continue
            # Detecta se col E é telefone ou email
            col_e = str(row[4]).strip() if len(row) > 4 else ""
            col_f = str(row[5]).strip() if len(row) > 5 else ""
            # Se col_e tem @ é email do síndico; senão é telefone
            if "@" in col_e:
                telefone = ""
                contato  = ""
                email    = col_e
            else:
                telefone = formatar_telefone(col_e) if col_e else ""
                contato  = col_f if col_f else ""
                email    = ""
            # Volumes das piscinas (colunas J=9, K=10, L=11)
            def _vol(r, idx):
                try: return float(str(r[idx]).replace(",",".").strip() or 0) if len(r) > idx else 0.0
                except: return 0.0
            vol_adulto   = _vol(row, 9)
            vol_infantil = _vol(row, 10)
            vol_family   = _vol(row, 11)
            vol_total    = _vol(row, 3) or (vol_adulto + vol_infantil + vol_family)

            _empresa_cl = str(row[12]).strip() if len(row) > 12 else "Aqua Gestão"
            if not _empresa_cl:
                _empresa_cl = "Aqua Gestão"
            _cnpj_cl = str(row[13]).strip() if len(row) > 13 else ""  # _CNPJ_LER_COLUNA_N_
            cliente_base = {
                "id":           id_val,
                "nome":         nome,
                "cnpj":         _cnpj_cl,
                "telefone":     telefone,
                "contato":      contato,
                "email":        email,
                "endereco":     str(row[6]).strip() if len(row) > 6 else "",
                "status":       str(row[8]).strip() if len(row) > 8 else "Ativo",
                "vol_total":    vol_total,
                "vol_adulto":   vol_adulto,
                "vol_infantil": vol_infantil,
                "vol_family":   vol_family,
                "empresa":      _empresa_cl,
                "piscinas_extras": [],  # carregado do JSON local se disponível
            }
            clientes.append(_enriquecer_cliente_com_dados_locais(cliente_base))
        return clientes
    except Exception as e:
        _log_sheets_erro("sheets_listar_clientes_completo", e)
        return []


def sheets_editar_cliente(id_cliente: str, nome: str, cnpj: str, endereco: str,
                           contato: str, telefone: str,
                           vol_adulto: float = 0, vol_infantil: float = 0, vol_family: float = 0,
                           empresa: str = "") -> bool:
    """Edita cliente existente na aba Clientes pelo ID."""
    import re as _re
    try:
        sh = conectar_sheets()
        if sh is None:
            return False
        aba = obter_aba_sheets("👥 Clientes")
        todos = aba.get_all_values()
        vol_total = (vol_adulto or 0) + (vol_infantil or 0) + (vol_family or 0)
        for i, row in enumerate(todos):
            if len(row) > 1 and str(row[1]).strip() == id_cliente.strip():
                linha_sheets = i + 1
                # Preserva empresa existente se não informada
                _empresa_atual = str(row[12]).strip() if len(row) > 12 else ""
                _empresa_final = empresa if empresa else (_empresa_atual or "Aqua Gestão")
                nova = [
                    "",
                    id_cliente,
                    nome,
                    str(vol_total) if vol_total else "",
                    formatar_telefone(telefone),
                    contato,
                    endereco,
                    str(row[7]).strip() if len(row) > 7 else datetime.now().strftime("%Y-%m-%d"),
                    str(row[8]).strip() if len(row) > 8 else "Ativo",
                    str(vol_adulto) if vol_adulto else "",
                    str(vol_infantil) if vol_infantil else "",
                    str(vol_family) if vol_family else "",
                    _empresa_final,                # M - Empresa
                ]
                aba.update(f"A{linha_sheets}:M{linha_sheets}", [nova], value_input_option="USER_ENTERED")
                return True
        return False
    except Exception as e:
        _log_sheets_erro("sheets_editar_cliente", e)
        return False

def sheets_carregar_cliente_por_nome(nome: str) -> dict:
    """Retorna dict com dados do cliente pelo nome (busca parcial)."""
    clientes = sheets_listar_clientes_completo()
    nome_lower = nome.lower().strip()
    for c in clientes:
        if c["nome"].lower().strip() == nome_lower:
            return c
    # Busca parcial
    for c in clientes:
        if nome_lower in c["nome"].lower():
            return c
    return {}


# =========================================
# GESTÃO DE OPERADORES
# =========================================

OPERADORES_JSON = None  # será inicializado após BASE_DIR

def _get_operadores_path():
    return BASE_DIR / "_operadores.json"

def carregar_operadores() -> list[dict]:
    """Carrega lista de operadores do arquivo JSON local."""
    try:
        p = _get_operadores_path()
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
        return []
    except Exception:
        return []

def salvar_operadores(lista: list):
    """Salva lista de operadores no arquivo JSON local."""
    try:
        p = _get_operadores_path()
        p.write_text(json.dumps(lista, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

def validar_pin_operador(pin: str) -> dict | None:
    """Valida PIN do operador. Retorna dict do operador ou None se inválido.
    Também aceita PIN global 2940 (acesso total)."""
    pin_limpo = str(pin or "").strip()
    # PIN global continua funcionando — acesso total
    if pin_limpo == PIN_OPERADOR:
        return {"nome": "Operador", "pin": pin_limpo, "condomínios": ["TODOS"], "acesso_total": True}
    # Busca nos operadores do Sheets
    try:
        operadores = sheets_listar_operadores()
        for op in operadores:
            if op.get("pin", "").strip() == pin_limpo and op.get("ativo", True):
                op["condomínios"] = _condominios_organizar(op.get("condomínios", []))
                op["acesso_total"] = op.get("acesso_total", False) or any(
                    _normalizar_chave_acesso(c) == "todos" for c in op["condomínios"]
                ) or not op["condomínios"]
                return op
    except Exception:
        pass
    # Fallback: JSON local
    operadores_local = carregar_operadores()
    for op in operadores_local:
        if str(op.get("pin", "")).strip() == pin_limpo and op.get("ativo", True):
            op["condomínios"] = _condominios_organizar(op.get("condomínios", []))
            op["acesso_total"] = op.get("acesso_total", False) or any(
                _normalizar_chave_acesso(c) == "todos" for c in op["condomínios"]
            ) or not op["condomínios"]
            return op
    return None

@st.cache_data(ttl=45, show_spinner=False)
def sheets_listar_lancamentos(nome_condominio: str) -> list[dict]:
    """Retorna lançamentos de visitas de um condomínio a partir da aba 🔬 Visitas."""
    try:
        sh = conectar_sheets()
        if sh is None:
            return []

        aba = obter_aba_sheets("🔬 Visitas")
        todos = aba.get_all_values()
        lancamentos = []

        for row in todos:
            if len(row) < 5:
                continue

            id_visita = str(row[1]).strip() if len(row) > 1 else ""
            data_raw = str(row[2]).strip() if len(row) > 2 else ""
            cond_linha = str(row[4]).strip() if len(row) > 4 else ""

            # Ignora cabeçalho e linhas que não são visita real.
            if not id_visita.startswith("V"):
                continue

            if not nomes_condominio_equivalentes(nome_condominio, cond_linha):
                continue

            def _r(i):
                return row[i] if len(row) > i else ""

            data_norm = normalizar_data_visita(data_raw)
            payload = {}
            payload_raw = _r(22)  # Coluna W
            if payload_raw:
                try:
                    payload = json.loads(payload_raw)
                except Exception:
                    payload = {}

            if payload:
                payload["id_visita"] = payload.get("id_visita", id_visita)
                payload["data"] = normalizar_data_visita(payload.get("data", data_norm))
                payload["condominio"] = payload.get("condominio", cond_linha)
                payload["status"] = payload.get("status", _r(19))
                payload["operador"] = payload.get("operador", _r(20))
                lancamentos.append(payload)
                continue

            # Compatibilidade com linhas antigas sem Payload JSON.
            lancamentos.append({
                "id_visita": id_visita,
                "data": data_norm,
                "condominio": cond_linha,
                "ph": _r(5),
                "cloro_livre": _r(6),
                "cloro_total": _r(7),
                "alcalinidade": _r(8),
                "dureza": _r(9),
                "cianurico": _r(10),
                "observacao": _r(14),
                "dosagem_resumo": _r(15),
                "status": _r(19),
                "operador": _r(20),
                "problemas": _r(21),
            })

        return lancamentos

    except Exception as e:
        _log_sheets_erro("sheets_listar_lancamentos", e)
        return []


@st.cache_data(ttl=45, show_spinner=False)
def sheets_listar_todas_visitas() -> list[dict]:
    """Lê a aba 🔬 Visitas uma única vez e retorna todas as visitas.

    Use esta função no dashboard e em relatórios que precisam de totais,
    em vez de chamar sheets_listar_lancamentos() em loop por cliente.
    Evita erro 429 (Quota exceeded) do Google Sheets.
    """
    try:
        aba = obter_aba_sheets("🔬 Visitas")
        todos = aba.get_all_values()
        visitas = []

        for row in todos:
            if len(row) < 5:
                continue
            id_visita = str(row[1]).strip() if len(row) > 1 else ""
            if not id_visita.startswith("V"):
                continue

            def _r(i):
                return row[i] if len(row) > i else ""

            payload_raw = _r(22)
            payload = {}
            if payload_raw:
                try:
                    payload = json.loads(payload_raw)
                except Exception:
                    payload = {}

            visitas.append({
                "id_visita": id_visita,
                "data": normalizar_data_visita(_r(2)),
                "id_cliente": _r(3),
                "condominio": _r(4),
                "ph": _r(5),
                "crl": _r(6),
                "ct": _r(7),
                "alcalinidade": _r(8),
                "dureza": _r(9),
                "cya": _r(10),
                "observacao": _r(14),
                "dosagem": _r(15),
                "status": _r(19),
                "operador": _r(20),
                "problemas": _r(21),
                "payload_json": payload_raw,
                "salvo_em": _r(23),
                "fonte": _r(24),
                "mes_ano": _r(25),
                "_payload": payload,
            })

        return visitas

    except Exception as e:
        _log_sheets_erro("sheets_listar_todas_visitas", e)
        return []


def extrair_lancamento_de_pdf_visita(pdf_bytes: bytes, nome_condominio_padrao: str = "") -> dict:
    """Extrai dados básicos de um PDF de Relatório de Visita Aqua Gestão.

    Uso: permite recuperar visitas antigas que geraram PDF, mas não foram gravadas
    na aba 🔬 Visitas do Google Sheets. Não inventa dados: só preenche o que achar no PDF.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        texto = "\n".join(page.get_text("text") for page in doc)
    except Exception as e:
        _log_sheets_erro("extrair_lancamento_de_pdf_visita/fitz", e)
        return {}

    linhas = [re.sub(r"\s+", " ", l).strip() for l in texto.splitlines() if str(l).strip()]

    def _apos(rotulo: str) -> str:
        rot = normalizar_texto_busca(rotulo)
        for i, linha in enumerate(linhas):
            if normalizar_texto_busca(linha) == rot and i + 1 < len(linhas):
                return linhas[i + 1].strip()
        return ""

    nome_cond = _apos("CONDOMÍNIO / LOCAL") or nome_condominio_padrao
    data_vis = normalizar_data_visita(_apos("DATA DA VISITA"))
    operador = _apos("OPERADOR")
    resp = _apos("RESP. TÉCNICO")

    observacao = ""
    for i, linha in enumerate(linhas):
        if normalizar_texto_busca(linha) == "observacoes" and i + 1 < len(linhas):
            prox = linhas[i + 1].strip().strip('"')
            if normalizar_texto_busca(prox) not in ("registro fotografico", "antes do tratamento"):
                observacao = prox
            break

    def _extrair_bloco(nome_piscina: str, inicio_idx: int, fim_idx: int) -> dict:
        bloco = linhas[inicio_idx:fim_idx]

        def _valor(prefixos: list[str]) -> str:
            for ln in bloco:
                ln_limpa = ln.strip()
                ln_norm = normalizar_texto_busca(ln_limpa)
                for pref in prefixos:
                    pref_norm = normalizar_texto_busca(pref)
                    if ln_norm.startswith(pref_norm):
                        resto = ln_limpa
                        for alvo in ["pH", "CRL mg/L", "CRL", "CT mg/L", "CT", "Alc. mg/L (15d)", "Alc", "Dureza mg/L (15d)", "Dureza", "CYA mg/L (15d)", "CYA", "Cloraminas"]:
                            if normalizar_texto_busca(resto).startswith(normalizar_texto_busca(alvo)):
                                resto = re.sub(re.escape(alvo), "", resto, count=1, flags=re.IGNORECASE).strip()
                                break
                        m = re.search(r"(—|-|\d+(?:[,.]\d+)?)", resto)
                        return m.group(1) if m else ""
            return ""

        return {
            "nome": nome_piscina,
            "ph": _valor(["pH"]),
            "cloro_livre": _valor(["CRL mg/L", "CRL"]),
            "cloro_total": _valor(["CT mg/L", "CT"]),
            "alcalinidade": _valor(["Alc. mg/L (15d)", "Alc"]),
            "dureza": _valor(["Dureza mg/L (15d)", "Dureza"]),
            "cianurico": _valor(["CYA mg/L (15d)", "CYA"]),
            "cloraminas": _valor(["Cloraminas"]),
        }

    indices_piscinas = []
    for i, linha in enumerate(linhas):
        ln = normalizar_texto_busca(linha)
        if "piscina adulto" in ln:
            indices_piscinas.append(("Piscina Adulto", i))
        elif "piscina infantil" in ln:
            indices_piscinas.append(("Piscina Infantil", i))
        elif "piscina family" in ln or "piscina spa" in ln:
            indices_piscinas.append((linha.replace("■", "").strip(), i))

    piscinas = []
    for pos, (nome_pisc, idx) in enumerate(indices_piscinas):
        fim = indices_piscinas[pos + 1][1] if pos + 1 < len(indices_piscinas) else len(linhas)
        for j in range(idx + 1, fim):
            if normalizar_texto_busca(linhas[j]) in ("alertas tecnicos", "dosagens aplicadas", "observacoes", "registro fotografico"):
                fim = j
                break
        p = _extrair_bloco(nome_pisc, idx, fim)
        if any(p.get(k) for k in ["ph", "cloro_livre", "cloro_total", "alcalinidade", "dureza", "cianurico", "cloraminas"]):
            piscinas.append(p)

    dosagens = []
    for i, linha in enumerate(linhas):
        if normalizar_texto_busca(linha) == "dosagens aplicadas":
            for ln in linhas[i + 1:i + 8]:
                ln_norm = normalizar_texto_busca(ln)
                if ln_norm in ("produto quantidade finalidade", "observacoes", "registro fotografico"):
                    continue
                if ln_norm.startswith("observacoes") or ln_norm.startswith("registro fotografico"):
                    break
                m = re.match(r"(.+?)\s+(\d+(?:[,.]\d+)?\s*(?:kg|g|ml|l|litros?))\s+(.+)$", ln, flags=re.IGNORECASE)
                if m:
                    dosagens.append({
                        "produto": m.group(1).strip(),
                        "quantidade": m.group(2).strip(),
                        "unidade": "",
                        "finalidade": m.group(3).strip(),
                    })
                    break
            break

    problemas = []
    coletando = False
    for ln in linhas:
        ln_norm = normalizar_texto_busca(ln)
        if ln_norm == "alertas tecnicos":
            coletando = True
            continue
        if coletando and ln_norm in ("dosagens aplicadas", "observacoes", "registro fotografico"):
            break
        if coletando:
            problemas.append(ln.replace("■", "").strip())

    base = piscinas[0] if piscinas else {}
    lancamento = {
        "data": data_vis,
        "operador": operador,
        "responsavel_tecnico_pdf": resp,
        "condominio": nome_cond,
        "ph": base.get("ph", ""),
        "cloro_livre": base.get("cloro_livre", ""),
        "cloro_total": base.get("cloro_total", ""),
        "cloraminas": base.get("cloraminas", ""),
        "alcalinidade": base.get("alcalinidade", ""),
        "dureza": base.get("dureza", ""),
        "cianurico": base.get("cianurico", ""),
        "piscinas": piscinas,
        "problemas": " | ".join(problemas),
        "observacao": observacao,
        "dosagens": dosagens,
        "parecer": "Importado de PDF de visita",
        "salvo_em": _agora_brasilia(),
        "fonte": "PDF de Relatório de Visita",
    }
    if not lancamento.get("data") and not piscinas and not dosagens:
        return {}
    if not lancamento.get("condominio"):
        lancamento["condominio"] = nome_condominio_padrao
    return lancamento

# =========================================
# OPERADORES — CONTROLE DE ACESSO
# =========================================


def filtrar_condomínios_por_operador(nome_operador: str, todos_condomínios: list[str]) -> list[str]:
    """Retorna lista de condomínios que o operador tem permissão de ver."""
    if not nome_operador.strip():
        return todos_condomínios  # sem nome → mostra todos (modo antigo)
    operadores = sheets_listar_operadores()
    for op in operadores:
        if _normalizar_chave_acesso(op["nome"]) == _normalizar_chave_acesso(nome_operador):
            conds = _condominios_organizar(op.get("condomínios", []))
            if any(_normalizar_chave_acesso(c) == "todos" for c in conds) or not conds:
                return todos_condomínios
            return _resolver_condominios_permitidos_exatos(conds, todos_condomínios)
    # Operador não cadastrado → mostra todos (retrocompatibilidade)
    return todos_condomínios


# =========================================
# MOTOR DE SUGESTÕES DE DOSAGEM
# =========================================

# Doses padrão APSP/WHO
DOSE_HIPOCLORITO_65  = 13.0   # g/m³ por ppm de CRL
DOSE_DICLORO_56      = 15.0   # g/m³ por ppm de CRL
DOSE_ACIDO_MURIATICO = 18.0   # mL/m³ por 0,1 de pH
DOSE_BICARBONATO     = 15.0   # g/m³ por 10 ppm de Alc
DOSE_CLORETO_CALCIO  = 17.0   # g/m³ por 10 ppm de Dureza
FATOR_DEMANDA        = 1.20   # +20% por demanda orgânica/UV

# Metas ideais (centro da faixa)
META_CRL        = 3.0    # ppm
META_PH         = 7.35   # centro entre 7,2-7,5
META_ALC        = 100.0  # ppm
META_DUREZA     = 225.0  # ppm
META_CYA        = 40.0   # ppm (só monitoramento)

# Faixas ideais
FAIXA_PH_MIN    = 7.2;  FAIXA_PH_MAX    = 7.8
FAIXA_CRL_MIN   = 0.5;  FAIXA_CRL_MAX   = 5.0
FAIXA_ALC_MIN   = 80;   FAIXA_ALC_MAX   = 120
FAIXA_DC_MIN    = 150;  FAIXA_DC_MAX    = 300
FAIXA_CYA_MIN   = 30;   FAIXA_CYA_MAX   = 50


def calcular_sugestoes_dosagem(ph: float | None, crl: float | None,
                                alc: float | None, dc: float | None,
                                cya: float | None, volume_m3: float) -> list[dict]:
    """
    Calcula sugestões de produtos e doses baseado nos parâmetros medidos.
    Retorna lista de dicts com: produto, quantidade, unidade, prioridade, justificativa.
    """
    sugestoes = []

    if volume_m3 <= 0:
        return sugestoes

    # ── 1. pH — SEMPRE CORRIGIR ANTES DO CLORO ───────────────────────────────
    if ph is not None:
        if ph > 7.8:
            # pH muito alto — reduzir antes de clorar
            deficit_ph = round(ph - 7.5, 1)
            dose_ml = round(deficit_ph / 0.1 * DOSE_ACIDO_MURIATICO * volume_m3 * FATOR_DEMANDA)
            sugestoes.append({
                "prioridade": 1,
                "produto": "Ácido muriático 31%",
                "quantidade": dose_ml,
                "unidade": "mL",
                "acao": "Reduzir pH",
                "justificativa": f"pH {ph:.1f} acima de 7,8 — reduzir para 7,4–7,5 antes de clorar. "
                                 f"Cloro perde >50% eficiência com pH alto.",
                "norma": "APSP / ABNT NBR 10339",
            })
        elif ph < 7.2:
            # pH baixo — hipoclorito vai elevar naturalmente
            sugestoes.append({
                "prioridade": 1,
                "produto": "Observação",
                "quantidade": 0,
                "unidade": "",
                "acao": "pH baixo — Hipoclorito de cálcio vai elevar",
                "justificativa": f"pH {ph:.1f} abaixo de 7,2 — aplicar Hipoclorito de cálcio 65% "
                                 f"(pH ~11,5) vai elevar o pH naturalmente enquanto desinfeta.",
                "norma": "APSP / WHO",
            })

    # ── 2. CLORO — produto baseado no pH ─────────────────────────────────────
    if crl is not None and crl < META_CRL:
        deficit_crl = round(META_CRL - crl, 2)

        # Seleciona produto pelo pH
        if ph is None or ph <= 7.5:
            produto_cloro = "Hipoclorito de cálcio 65%"
            fator_cloro   = DOSE_HIPOCLORITO_65
            motivo_cloro  = f"pH {ph:.1f} ≤ 7,5 — produto ideal nesta faixa" if ph else "pH não medido — usando padrão"
        else:
            produto_cloro = "Dicloro 56%"
            fator_cloro   = DOSE_DICLORO_56
            motivo_cloro  = f"pH {ph:.1f} entre 7,5–7,8 — Dicloro é mais indicado (mais ácido)"

        dose_g = round(deficit_crl * volume_m3 * fator_cloro * FATOR_DEMANDA)

        if dose_g >= 1000:
            qtd_fmt = round(dose_g / 1000, 2)
            unid = "kg"
        else:
            qtd_fmt = dose_g
            unid = "g"

        sugestoes.append({
            "prioridade": 2,
            "produto": produto_cloro,
            "quantidade": qtd_fmt,
            "unidade": unid,
            "acao": f"Elevar CRL de {crl:.1f} → {META_CRL:.1f} ppm",
            "justificativa": f"CRL {crl:.1f} ppm abaixo da meta ({META_CRL} ppm). "
                             f"Déficit: {deficit_crl} ppm × {volume_m3}m³ × {fator_cloro}g × 1,2 dem. "
                             f"| {motivo_cloro}.",
            "norma": "APSP / WHO",
        })

    elif crl is not None and crl > FAIXA_CRL_MAX:
        sugestoes.append({
            "prioridade": 2,
            "produto": "Aeração + aguardar",
            "quantidade": 0,
            "unidade": "",
            "acao": f"CRL {crl:.1f} ppm — acima de {FAIXA_CRL_MAX} ppm",
            "justificativa": "Cloro excessivo — não adicionar mais. "
                             "Aeração (agitação da água) e luz solar reduzem naturalmente.",
            "norma": "ABNT NBR 10339",
        })

    # ── 3. ALCALINIDADE ───────────────────────────────────────────────────────
    if alc is not None:
        if alc < FAIXA_ALC_MIN:
            deficit_alc = round(META_ALC - alc, 1)
            dose_g = round((deficit_alc / 10) * DOSE_BICARBONATO * volume_m3)
            qtd_fmt = round(dose_g / 1000, 2) if dose_g >= 1000 else dose_g
            unid = "kg" if dose_g >= 1000 else "g"
            sugestoes.append({
                "prioridade": 3,
                "produto": "Bicarbonato de sódio",
                "quantidade": qtd_fmt,
                "unidade": unid,
                "acao": f"Elevar alcalinidade de {alc:.0f} → {META_ALC:.0f} ppm",
                "justificativa": f"Alcalinidade {alc:.0f} ppm abaixo de {FAIXA_ALC_MIN} ppm. "
                                 f"Déficit: {deficit_alc:.0f} ppm ÷ 10 × {DOSE_BICARBONATO}g × {volume_m3}m³.",
                "norma": "WHO / ABNT NBR 10339",
            })
        elif alc > FAIXA_ALC_MAX:
            excesso_alc = round(alc - META_ALC, 1)
            dose_ml = round((excesso_alc / 0.1) * DOSE_ACIDO_MURIATICO * volume_m3 * 0.5)
            sugestoes.append({
                "prioridade": 3,
                "produto": "Ácido muriático 31%",
                "quantidade": dose_ml,
                "unidade": "mL",
                "acao": f"Reduzir alcalinidade de {alc:.0f} → {META_ALC:.0f} ppm",
                "justificativa": f"Alcalinidade {alc:.0f} ppm acima de {FAIXA_ALC_MAX} ppm. "
                                 "Aplicar ácido muriático com bomba desligada.",
                "norma": "APSP",
            })

    # ── 4. DUREZA ─────────────────────────────────────────────────────────────
    if dc is not None:
        if dc < FAIXA_DC_MIN:
            deficit_dc = round(META_DUREZA - dc, 1)
            dose_g = round((deficit_dc / 10) * DOSE_CLORETO_CALCIO * volume_m3)
            qtd_fmt = round(dose_g / 1000, 2) if dose_g >= 1000 else dose_g
            unid = "kg" if dose_g >= 1000 else "g"
            sugestoes.append({
                "prioridade": 4,
                "produto": "Cloreto de cálcio",
                "quantidade": qtd_fmt,
                "unidade": unid,
                "acao": f"Elevar dureza de {dc:.0f} → {META_DUREZA:.0f} ppm",
                "justificativa": f"Dureza {dc:.0f} ppm abaixo de {FAIXA_DC_MIN} ppm. "
                                 "Água agressiva corrói equipamentos e pisos.",
                "norma": "APSP / WHO",
            })
        elif dc > FAIXA_DC_MAX:
            sugestoes.append({
                "prioridade": 4,
                "produto": "Troca parcial de água",
                "quantidade": round(volume_m3 * 0.2),
                "unidade": "m³",
                "acao": f"Reduzir dureza de {dc:.0f} ppm",
                "justificativa": f"Dureza {dc:.0f} ppm acima de {FAIXA_DC_MAX} ppm. "
                                 "Trocar ~20% da água e reequilibrar.",
                "norma": "APSP",
            })

    # ── 5. CYA — só monitoramento ─────────────────────────────────────────────
    if cya is not None:
        if cya > FAIXA_CYA_MAX:
            sugestoes.append({
                "prioridade": 5,
                "produto": "Troca parcial de água",
                "quantidade": round(volume_m3 * 0.3),
                "unidade": "m³",
                "acao": f"CYA {cya:.0f} ppm acima do limite",
                "justificativa": f"CYA {cya:.0f} ppm acima de {FAIXA_CYA_MAX} ppm. "
                                 "CYA alto inibe o cloro (efeito bloqueio). Trocar ~30% da água.",
                "norma": "WHO",
            })
        elif cya < FAIXA_CYA_MIN and cya > 0:
            sugestoes.append({
                "prioridade": 5,
                "produto": "Monitorar CYA",
                "quantidade": 0,
                "unidade": "",
                "acao": f"CYA {cya:.0f} ppm — abaixo do ideal",
                "justificativa": f"CYA {cya:.0f} ppm abaixo de {FAIXA_CYA_MIN} ppm. "
                                 "Sem ácido cianúrico puro no estoque — monitorar.",
                "norma": "APSP",
            })

    # Ordena por prioridade
    sugestoes.sort(key=lambda x: x["prioridade"])
    return sugestoes


def exibir_sugestoes_dosagem(sugestoes: list[dict]):
    """Exibe as sugestões de dosagem formatadas no Streamlit."""
    if not sugestoes:
        st.success("✅ Todos os parâmetros dentro da faixa ideal. Nenhuma correção necessária.")
        return

    st.markdown("**💊 Sugestões de correção (APSP/WHO):**")
    for s in sugestoes:
        prod = s["produto"]
        qtd  = s["quantidade"]
        unid = s["unidade"]
        acao = s["acao"]
        just = s["justificativa"]
        prio = s["prioridade"]

        if prio == 1:
            icon = "🔴"
        elif prio == 2:
            icon = "🟡"
        else:
            icon = "🔵"

        if qtd and qtd > 0:
            st.markdown(f"{icon} **{prod}** — **{qtd} {unid}** → _{acao}_")
        else:
            st.markdown(f"{icon} **{prod}** → _{acao}_")

        with st.expander("ℹ️ Detalhes técnicos", expanded=False):
            st.caption(f"📐 {just}")
            st.caption(f"📚 Norma: {s.get('norma','')}")

# =========================================
# CONFIGURAÇÃO GERAL
# =========================================

APP_TITLE = "Aqua Gestão – Controle Técnico de Piscinas"
APP_VERSION = "v6_relatorio_premium_aqua_final"
RESPONSAVEL_TÉCNICO = "Thyago Fernando da Silveira"
RESPONSAVEL_TECNICO_ASSINATURA = "Thyago Fernando da Silveira | CRQ 024025748 | Técnico em Química"
CRQ = "CRQ-MG 2ª Região | CRQ 024025748"
CRQ_NUMERO = "024025748"
QUALIFICACAO_RT = "Técnico em Química"
CERTIFICACOES_RT = "NR-26 e NR-6"
EMPRESA_RT = "Aqua Gestão – Controle Técnico de Piscinas"

# ── Dados da Bem Star Piscinas ──────────────────────────────────────────────
EMPRESA_BEM_STAR     = "Bem Star Piscinas"
CNPJ_BEM_STAR        = "26.799.958/0001-88"
ENDERECO_BEM_STAR    = "Avenida Getúlio Vargas, 4411 — CEP 38.412-316 — Uberlândia/MG"

# ── Texto RT para relatório sem RT ──────────────────────────────────────────
TEXTO_RT_SEM_RT = """SOBRE RESPONSABILIDADE TÉCNICA (RT)

Este relatório foi elaborado pela Bem Star Piscinas como registro técnico das análises e dosagens realizadas durante a visita de manutenção.

A Responsabilidade Técnica (RT) é um serviço de supervisão especializada, regulamentado pela Resolução CFQ nº 332/2025 — publicada pelo Conselho Federal de Química em 24 de junho de 2025 — e complementada pela Resolução CFQ nº 345/2026, que tornam obrigatória a Anotação de Responsabilidade Técnica (ART) para o tratamento químico e controle de qualidade da água de piscinas de uso público e coletivo, abrangendo condomínios residenciais, clubes, academias, hotéis e escolas.

A RT consiste na supervisão e assinatura de profissional habilitado e registrado no Conselho Regional de Química (CRQ), garantindo que todos os procedimentos estejam em conformidade com os padrões técnicos e sanitários vigentes, incluindo a ABNT NBR 10339. A ART deve ser emitida anualmente e afixada em local visível, podendo o CRQ realizar fiscalizações preventivas e, em caso de irregularidade, acionar a Vigilância Sanitária municipal.

A Aqua Gestão — empresa parceira especializada em Responsabilidade Técnica — oferece o serviço completo de RT com profissional habilitado pelo CRQ-MG, emissão de laudos mensais, ART anual e total suporte documental, garantindo segurança jurídica e conformidade normativa ao seu condomínio.

Saiba mais sobre o serviço de RT:
Thyago Fernando da Silveira | CRQ-MG 2ª Região | CRQ 024025748
Aqua Gestão – Controle Técnico de Piscinas"""

BASE_DIR = Path(__file__).resolve().parent
GENERATED_DIR = BASE_DIR / "generated"

LOGO_BEM_STAR_CANDIDATOS = [
    BASE_DIR / "bem_star_logo.png",
    BASE_DIR / "bem_star_logo.jpg",
    BASE_DIR / "assets" / "bem_star_logo.png",
]
TEMPLATE_CONTRATO = BASE_DIR / "template_rt_aqua_v2_relatorio.docx"
TEMPLATE_BEM_STAR = BASE_DIR / "template_bem_star.docx"
TEMPLATE_ADITIVO = BASE_DIR / "aditivo.docx"
TEMPLATE_RELATORIO = BASE_DIR / "relatorio_mensal.docx"
DADOS_JSON_NAME = "dados_condominio.json"
MANIFEST_JSON_NAME = "manifest.json"
ANALISES_PADRAO = 12
ANALISES_MAX_SUGERIDO = 40


def calcular_linhas_analises_por_frequencia(verificacoes_semanais: int | str | None = None, mes: str | None = None, ano: str | None = None) -> int:
    """Calcula quantas linhas de análise o relatório deve abrir.

    Regra operacional adotada: 12 linhas como padrão mínimo.
    Para clientes com rotina cadastrada, usa verificações semanais × 4 semanas.
    Se o mês tiver mais lançamentos reais importados, o sistema expande automaticamente.
    """
    try:
        freq = int(float(str(verificacoes_semanais or "").replace(",", ".")))
    except Exception:
        freq = 3
    freq = max(1, min(freq, 7))
    return max(ANALISES_PADRAO, min(freq * 4, ANALISES_MAX_SUGERIDO))


def obter_verificacoes_semanais_cliente(cliente_ou_dados: dict | None) -> int:
    """Obtém a frequência semanal cadastrada para o cliente, com fallback para 3x/semana."""
    dados = cliente_ou_dados or {}
    for chave in ["verificacoes_semanais", "frequencia_verificacoes_semanais", "visitas_semanais", "frequencia_semanal"]:
        valor = dados.get(chave)
        try:
            iv = int(float(str(valor or "").replace(",", ".")))
            if iv > 0:
                return max(1, min(iv, 7))
        except Exception:
            pass
    return 3


ASSINATURA_RT_CANDIDATOS = [
    BASE_DIR / "assinatura_rt.png",
    BASE_DIR / "assinatura_rt.jpg",
    BASE_DIR / "assinatura_rt.jpeg",
    BASE_DIR / "assets" / "assinatura_rt.png",
    BASE_DIR / "images" / "assinatura_rt.png",
]

LOGO_CANDIDATOS = [
    BASE_DIR / "aqua_gestao_logo.png",
    BASE_DIR / "aqua_gestao_logo.jpg",
    BASE_DIR / "aqua_gestao_logo.jpeg",
    BASE_DIR / "logo.png",
    BASE_DIR / "logo.jpg",
    BASE_DIR / "logo.jpeg",
    BASE_DIR / "Aqua Gestão Logo.png",
    BASE_DIR / "Aqua_Gestao_Logo.png",
    BASE_DIR / "assets" / "aqua_gestao_logo.png",
    BASE_DIR / "assets" / "logo.png",
    BASE_DIR / "images" / "aqua_gestao_logo.png",
    BASE_DIR / "images" / "logo.png",
]

GENERATED_DIR.mkdir(exist_ok=True)

# Fallback global — sobrescrito pelo Modo Operador quando necessário
def _autosave_rascunho():
    pass

st.set_page_config(
    page_title="",
    page_icon="📘",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================
# PROTEÇÃO CONTRA TRADUÇÃO AUTOMÁTICA DO NAVEGADOR
# =========================================
# O Google Tradutor/Chrome Translate altera nós internos do React/Streamlit e pode gerar
# erros de front-end como: NotFoundError: Failed to execute 'removeChild' on 'Node'.
# Também distorce siglas técnicas (CT -> TC, ALC -> Álcool, CYA -> Tchau).
def bloquear_traducao_navegador():
    """Compatibilidade com Streamlit 1.56.

    Nesta versão, components.html no fluxo principal dispara rerun completo.
    Mantemos apenas um marcador leve via st.markdown para reduzir riscos sem
    criar iframes/componentes customizados.
    """
    try:
        st.markdown(
            """
            <div class="notranslate" translate="no" lang="pt-BR"></div>
            <style>
                .notranslate { translate: no; }
            </style>
            """,
            unsafe_allow_html=True,
        )
    except Exception:
        pass

bloquear_traducao_navegador()


# =========================================
# ESTILO VISUAL
# =========================================

st.markdown(
    """
    <style>
        .main .block-container {
            padding-top: 1.2rem;
            padding-bottom: 2rem;
            max-width: 1360px;
        }

        .top-card {
            border: 1px solid rgba(20, 85, 160, 0.18);
            border-radius: 18px;
            padding: 18px 22px;
            background: linear-gradient(135deg, #ffffff 0%, #f6fbff 100%);
            box-shadow: 0 10px 25px rgba(10, 50, 100, 0.07);
            margin-bottom: 18px;
        }

        .top-title {
            font-size: 1.7rem;
            font-weight: 700;
            color: #0d3d75;
            margin: 0;
            line-height: 1.15;
        }

        .top-subtitle {
            font-size: 0.95rem;
            color: #3b5d85;
            margin-top: 8px;
        }

        .info-badge {
            display: inline-block;
            padding: 7px 12px;
            border-radius: 999px;
            background: #edf5ff;
            color: #134b8a;
            border: 1px solid #d3e6ff;
            font-size: 0.88rem;
            margin-right: 8px;
            margin-top: 8px;
        }

        .section-card {
            border: 1px solid rgba(20, 85, 160, 0.16);
            border-radius: 18px;
            padding: 18px;
            background: #ffffff;
            box-shadow: 0 8px 20px rgba(10, 50, 100, 0.05);
            margin-bottom: 16px;
        }

        .history-meta {
            font-size: 0.84rem;
            color: #6a7d92;
            margin-top: 2px;
        }

        .confirm-box {
            border: 1px solid rgba(220, 80, 80, 0.35);
            border-radius: 12px;
            padding: 10px;
            margin-top: 8px;
            margin-bottom: 10px;
            background: rgba(220, 80, 80, 0.08);
        }

        .quick-mode-box {
            border: 1px solid rgba(20, 85, 160, 0.18);
            border-radius: 14px;
            padding: 12px 14px;
            background: #f6fbff;
            margin-bottom: 14px;
        }

        .alert-vigente {
            border: 1px solid rgba(40, 140, 80, 0.25);
            border-radius: 14px;
            padding: 12px 14px;
            background: rgba(40, 140, 80, 0.08);
            margin-bottom: 14px;
        }

        .alert-vencendo {
            border: 1px solid rgba(220, 150, 20, 0.35);
            border-radius: 14px;
            padding: 12px 14px;
            background: rgba(255, 190, 40, 0.12);
            margin-bottom: 14px;
        }

        .alert-vencido {
            border: 1px solid rgba(220, 60, 60, 0.35);
            border-radius: 14px;
            padding: 12px 14px;
            background: rgba(220, 60, 60, 0.10);
            margin-bottom: 14px;
        }

        .status-badge {
            display: inline-block;
            padding: 4px 8px;
            border-radius: 999px;
            font-size: 0.75rem;
            margin-top: 6px;
            border: 1px solid rgba(0,0,0,0.08);
        }

        .status-vigente { background: #ecfff3; color: #1d7a43; }
        .status-vencendo { background: #fff8e8; color: #9c6200; }
        .status-vencido { background: #fff0f0; color: #b42318; }
        .status-indefinido { background: #f4f6f8; color: #52606d; }

        .venc-row {
            border: 1px solid rgba(20, 85, 160, 0.12);
            border-radius: 16px;
            padding: 14px;
            background: linear-gradient(135deg, #ffffff 0%, #fbfdff 100%);
            margin-bottom: 8px;
        }

        .venc-nome {
            font-size: 1rem;
            font-weight: 700;
            color: #153f73;
            margin-bottom: 4px;
        }

        .venc-meta {
            font-size: 0.88rem;
            color: #5d7288;
            margin-bottom: 3px;
        }

        .venc-empty {
            border: 1px dashed rgba(20, 85, 160, 0.20);
            border-radius: 14px;
            padding: 14px;
            background: #f9fcff;
            color: #59708b;
        }

        .legacy-note {
            border: 1px solid rgba(120, 120, 120, 0.18);
            border-radius: 12px;
            padding: 10px 12px;
            background: #fafbfd;
            color: #5e6f82;
            font-size: 0.87rem;
            margin-top: 8px;
            margin-bottom: 8px;
        }

        .dash-mini {
            border: 1px solid rgba(20, 85, 160, 0.12);
            border-radius: 16px;
            padding: 14px;
            background: #fbfdff;
            min-height: 120px;
        }

        .dash-title {
            font-size: 0.88rem;
            color: #5f7590;
            margin-bottom: 6px;
        }

        .dash-value {
            font-size: 1.55rem;
            font-weight: 700;
            color: #103f78;
            line-height: 1.1;
        }

        .dash-sub {
            font-size: 0.82rem;
            color: #6d8197;
            margin-top: 8px;
        }

        .docs-note {
            border: 1px solid rgba(20, 85, 160, 0.12);
            border-radius: 12px;
            padding: 10px;
            background: #fbfdff;
            color: #5c7188;
            font-size: 0.84rem;
            margin-top: 8px;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================
# FUNÇÕES AUXILIARES GERAIS
# =========================================

def encontrar_assinatura_rt() -> Path | None:
    for caminho in ASSINATURA_RT_CANDIDATOS:
        if caminho.exists() and caminho.is_file():
            return caminho
    return None


def preparar_assinatura_rt_para_relatorio() -> Path | None:
    assinatura = encontrar_assinatura_rt()
    if not assinatura:
        return None

    destino = GENERATED_DIR / "_assinatura_rt_relatorio.png"
    try:
        with Image.open(assinatura) as img:
            img = img.convert("RGBA")
            fundo = Image.new("RGBA", img.size, (255, 255, 255, 255))
            fundo.alpha_composite(img)
            fundo = fundo.convert("RGB")
            fundo = ImageOps.expand(fundo, border=18, fill="white")
            fundo.save(destino, format="PNG")
        return destino
    except Exception:
        return assinatura


def encontrar_logo() -> Path | None:
    """Retorna somente a logo institucional da Aqua Gestão.

    Proteção importante: nunca aceitar arquivos da Bem Star na busca da Aqua,
    nem na lista fixa nem na varredura genérica.
    """
    for caminho in LOGO_CANDIDATOS:
        nome = caminho.name.lower()
        if "bem_star" in nome or "bemstar" in nome:
            continue
        if caminho.exists() and caminho.is_file():
            return caminho

    for extensao in ("*.png", "*.jpg", "*.jpeg", "*.webp"):
        for pasta in [BASE_DIR, BASE_DIR / "assets", BASE_DIR / "images"]:
            if pasta.exists():
                encontrados = list(pasta.glob(extensao))
                for arq in encontrados:
                    nome = arq.name.lower()
                    if "logo" in nome and "bem_star" not in nome and "bemstar" not in nome:
                        return arq
    return None


def encontrar_logo_bem_star() -> Path | None:
    for caminho in LOGO_BEM_STAR_CANDIDATOS:
        if caminho.exists() and caminho.is_file():
            return caminho
    for extensao in ("*.png", "*.jpg", "*.jpeg", "*.webp"):
        for pasta in [BASE_DIR, BASE_DIR / "assets", BASE_DIR / "images"]:
            if pasta.exists():
                for arq in pasta.glob(extensao):
                    if "bem_star" in arq.name.lower() and "logo" in arq.name.lower():
                        return arq
    return None


def inserir_foto_docx_exif(doc_or_run, foto_path: Path, width_inches: float = 5.5):
    """Insere foto em DOCX com rotação EXIF corrigida automaticamente.
    Salva temporariamente a imagem corrigida e insere no documento.
    """
    import io as _io
    try:
        from PIL import Image as _PIL, ImageOps as _IOS
        _img = _PIL.open(str(foto_path))
        _img = _IOS.exif_transpose(_img)  # Corrige rotação EXIF
        _img.thumbnail((1800, 1800), _PIL.LANCZOS)
        _buf = _io.BytesIO()
        _img.convert("RGB").save(_buf, format="JPEG", quality=88)
        _buf.seek(0)
        if hasattr(doc_or_run, "add_picture"):
            # É um run
            doc_or_run.add_picture(_buf, width=Inches(width_inches))
        else:
            # É um parágrafo ou doc — adicionar run
            doc_or_run.add_run().add_picture(_buf, width=Inches(width_inches))
        return True
    except Exception:
        # Fallback sem correção EXIF
        try:
            if hasattr(doc_or_run, "add_picture"):
                doc_or_run.add_picture(str(foto_path), width=Inches(width_inches))
            else:
                doc_or_run.add_run().add_picture(str(foto_path), width=Inches(width_inches))
            return True
        except Exception:
            return False


def logo_para_base64(path) -> str:
    """Converte imagem para string base64 para exibicao HTML inline."""
    if path is None or not path.exists():
        return ""
    try:
        import base64 as _b64
        ext = path.suffix.lower().lstrip(".")
        mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "webp": "webp"}.get(ext, "png")
        with open(path, "rb") as f:
            dados = _b64.b64encode(f.read()).decode()
        return f"data:image/{mime};base64,{dados}"
    except Exception:
        return ""


def _agora_brasilia() -> str:
    """Retorna horario atual no fuso de Brasilia (UTC-3)."""
    from datetime import timezone, timedelta
    return datetime.now(tz=timezone(timedelta(hours=-3))).strftime("%d/%m/%Y %H:%M:%S")

def salvar_rascunho_operador(nome_cond: str, dados: dict, salvar_sheets: bool = False) -> bool:
    """Salva rascunho. Por padrão salva APENAS localmente (evita reruns pesados durante digitação).
    
    v5: o parâmetro salvar_sheets=True deve ser usado APENAS quando o operador
    clicar explicitamente em "Salvar rascunho". Nunca chamar com sheets=True
    a partir de on_change ou de autosave automático.
    """
    dados["_rascunho_salvo_em"] = _agora_brasilia()
    dados["_rascunho_cond"] = nome_cond.strip()

    # 1. Arquivo local (rápido — sempre executado)
    try:
        pasta = GENERATED_DIR / slugify_nome(nome_cond.strip())
        pasta.mkdir(parents=True, exist_ok=True)
        dados_local = {k: v for k, v in dados.items()
                       if not any(p in str(k).lower() for p in ["b64","base64","assinatura_responsavel_b64"])}
        with open(pasta / "_rascunho_operador.json", "w", encoding="utf-8") as f:
            json.dump(dados_local, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

    # 2. Google Sheets aba _Rascunhos — SOMENTE quando solicitado explicitamente
    # v5: não chamar durante digitação/on_change para evitar reruns e instabilidade
    if salvar_sheets:
        try:
            sh = conectar_sheets()
            if sh:
                try:
                    aba_rasc = obter_aba_sheets("_Rascunhos")
                except Exception:
                    aba_rasc = sh.add_worksheet(title="_Rascunhos", rows=500, cols=4)
                    aba_rasc.update(range_name="A1:D1", values=[["Condomínio", "Operador", "Salvo em", "Dados JSON"]])
                dados_sh = {k: v for k, v in dados.items()
                            if not any(p in str(k).lower() for p in ["b64","base64","assinatura_responsavel_b64"])}
                payload = json.dumps(dados_sh, ensure_ascii=False)
                if len(payload) > 45000:
                    payload = payload[:45000] + "..."
                todos = aba_rasc.get_all_values()
                chave = nome_cond.strip().lower()
                linha_existente = None
                for i, row in enumerate(todos[1:], start=2):
                    if row and str(row[0]).strip().lower() == chave:
                        linha_existente = i
                        break
                nova = [nome_cond.strip(), dados.get("operador",""), dados["_rascunho_salvo_em"], payload]
                if linha_existente:
                    aba_rasc.update(range_name=f"A{linha_existente}:D{linha_existente}", values=[nova], value_input_option="RAW")
                else:
                    proxima = max(len(todos) + 1, 2)
                    aba_rasc.update(range_name=f"A{proxima}:D{proxima}", values=[nova], value_input_option="RAW")
        except Exception:
            pass

    return True


def carregar_rascunho_operador(nome_cond: str) -> dict:
    """Carrega rascunho: arquivo local primeiro, fallback no Sheets apos sleep."""
    # 1. Arquivo local
    try:
        pasta = GENERATED_DIR / slugify_nome(nome_cond.strip())
        rascunho_path = pasta / "_rascunho_operador.json"
        if rascunho_path.exists():
            with open(rascunho_path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass

    # 2. Fallback Google Sheets
    try:
        sh = conectar_sheets()
        if sh:
            try:
                aba_rasc = obter_aba_sheets("_Rascunhos")
            except Exception:
                return {}
            todos = aba_rasc.get_all_values()
            chave = nome_cond.strip().lower()
            for row in todos[1:]:
                if row and str(row[0]).strip().lower() == chave and len(row) >= 4 and row[3].strip():
                    dados = json.loads(row[3])
                    try:
                        pasta = GENERATED_DIR / slugify_nome(nome_cond.strip())
                        pasta.mkdir(parents=True, exist_ok=True)
                        with open(pasta / "_rascunho_operador.json", "w", encoding="utf-8") as f:
                            json.dump(dados, f, ensure_ascii=False, indent=2)
                    except Exception:
                        pass
                    return dados
    except Exception:
        pass

    return {}


def deletar_rascunho_operador(nome_cond: str):
    """Remove rascunho apos salvar lancamento definitivo (local + Sheets)."""
    try:
        pasta = GENERATED_DIR / slugify_nome(nome_cond.strip())
        rascunho_path = pasta / "_rascunho_operador.json"
        if rascunho_path.exists():
            rascunho_path.unlink()
    except Exception:
        pass
    try:
        sh = conectar_sheets()
        if sh:
            try:
                aba_rasc = obter_aba_sheets("_Rascunhos")
                todos = aba_rasc.get_all_values()
                chave = nome_cond.strip().lower()
                for i, row in enumerate(todos[1:], start=2):
                    if row and str(row[0]).strip().lower() == chave:
                        aba_rasc.update(range_name=f"A{i}:D{i}", values=[["","","",""]], value_input_option="RAW")
                        break
            except Exception:
                pass
    except Exception:
        pass


def coletar_rascunho_operador(nome_cond: str, piscinas_ativas: list) -> dict:
    """Coleta estado atual do formulário do operador para salvar como rascunho."""
    dados = {
        "data_visita":  st.session_state.get("op_data_visita", ""),
        "operador":     st.session_state.get("op_operador", ""),
        "obs":          st.session_state.get("op_obs_campo", ""),
        "problemas":    st.session_state.get("op_problemas", ""),
        "resp_local":   st.session_state.get("op_resp_local", ""),
        "parecer":      st.session_state.get("op_parecer_visita", "✅ Satisfatório"),
        "assinatura_responsavel_b64": st.session_state.get("op_assinatura_responsavel_b64", ""),
        "assinatura_responsavel_data": st.session_state.get("op_assinatura_responsavel_data", ""),
        "piscinas": [],
        "dosagens": [],
    }
    # Parâmetros por piscina
    for pisc_nome in piscinas_ativas:
        _slug = slugify_nome(pisc_nome).lower()
        abrev = ("adulto" if "adulto" in _slug else
                 "infantil" if "infantil" in _slug else
                 "family" if "family" in _slug else "outra")
        dados["piscinas"].append({
            "nome":        pisc_nome,
            "ph":          st.session_state.get(f"op_{abrev}_ph", ""),
            "cloro_livre": st.session_state.get(f"op_{abrev}_crl", ""),
            "cloro_total": st.session_state.get(f"op_{abrev}_ct", ""),
            "alcalinidade":st.session_state.get(f"op_{abrev}_alc", ""),
            "dureza":      st.session_state.get(f"op_{abrev}_dc", ""),
            "cianurico":   st.session_state.get(f"op_{abrev}_cya", ""),
        })
    # Dosagens por piscina
    _slug_map_r = {"Piscina Adulto":"adulto","Piscina Infantil":"infantil","Piscina Family":"family"}
    for pisc_nome in piscinas_ativas:
        _slug_r = _slug_map_r.get(pisc_nome, slugify_nome(pisc_nome)[:12])
        _dos_p = []
        for i in range(5):
            prod = st.session_state.get(f"op_dos_{_slug_r}_prod_{i}", "").strip()
            if prod:
                _dos_p.append({
                    "produto":    prod,
                    "quantidade": st.session_state.get(f"op_dos_{_slug_r}_qtd_{i}", ""),
                    "unidade":    st.session_state.get(f"op_dos_{_slug_r}_un_{i}", ""),
                    "finalidade": st.session_state.get(f"op_dos_{_slug_r}_fin_{i}", ""),
                })
        if _dos_p:
            dados["dosagens"].extend(_dos_p)
    # Fallback: campos legados op_dos_prod_i
    if not dados["dosagens"]:
        for i in range(5):
            prod = st.session_state.get(f"op_dos_prod_{i}", "").strip()
            if prod:
                dados["dosagens"].append({
                    "produto":    prod,
                    "quantidade": st.session_state.get(f"op_dos_qtd_{i}", ""),
                    "unidade":    st.session_state.get(f"op_dos_un_{i}", ""),
                    "finalidade": st.session_state.get(f"op_dos_fin_{i}", ""),
                })
    # Fotos já salvas na pasta de rascunho
    _pasta_fotos_rasc = GENERATED_DIR / slugify_nome(nome_cond.strip()) / "fotos_rascunho"
    dados["fotos_rascunho"] = {"antes": [], "depois": [], "cmaq": []}
    if _pasta_fotos_rasc.exists():
        for _fp in sorted(_pasta_fotos_rasc.glob("rasc_*")):
            for _cat in ["antes", "depois", "cmaq", "extras"]:
                if f"rasc_{_cat}_" in _fp.name:
                    dados["fotos_rascunho"][_cat].append(_fp.name)
    return dados


def restaurar_rascunho_operador(rascunho: dict):
    """Restaura rascunho nos campos do formulário do operador."""
    if rascunho.get("data_visita"):
        st.session_state["op_data_visita"] = rascunho["data_visita"]
    if rascunho.get("operador"):
        st.session_state["op_operador"] = rascunho["operador"]
    if rascunho.get("obs"):
        st.session_state["op_obs_campo"] = rascunho["obs"]
    if rascunho.get("problemas"):
        st.session_state["op_problemas"] = rascunho["problemas"]
    if rascunho.get("resp_local"):
        st.session_state["op_resp_local"] = rascunho["resp_local"]
    if rascunho.get("parecer"):
        st.session_state["op_parecer_visita"] = rascunho["parecer"]
    if rascunho.get("assinatura_responsavel_b64"):
        st.session_state["op_assinatura_responsavel_b64"] = rascunho.get("assinatura_responsavel_b64", "")
    if rascunho.get("assinatura_responsavel_data"):
        st.session_state["op_assinatura_responsavel_data"] = rascunho.get("assinatura_responsavel_data", "")
    # Parâmetros por piscina
    for p in rascunho.get("piscinas", []):
        _slug = slugify_nome(p.get("nome","")).lower()
        abrev = ("adulto" if "adulto" in _slug else
                 "infantil" if "infantil" in _slug else
                 "family" if "family" in _slug else "outra")
        for campo, key in [
            ("ph", f"op_{abrev}_ph"), ("cloro_livre", f"op_{abrev}_crl"),
            ("cloro_total", f"op_{abrev}_ct"), ("alcalinidade", f"op_{abrev}_alc"),
            ("dureza", f"op_{abrev}_dc"), ("cianurico", f"op_{abrev}_cya"),
        ]:
            if p.get(campo):
                st.session_state[key] = p[campo]
    # Dosagens
    for i, d in enumerate(rascunho.get("dosagens", [])[:5]):
        st.session_state[f"op_dos_prod_{i}"] = d.get("produto", "")
        st.session_state[f"op_dos_qtd_{i}"]  = d.get("quantidade", "")
        st.session_state[f"op_dos_un_{i}"]   = d.get("unidade", "")
        st.session_state[f"op_dos_fin_{i}"]  = d.get("finalidade", "")


def aplicar_restauracao_pendente_operador():
    """Aplica restauração pendente antes de instanciar os widgets do formulário."""
    rasc = st.session_state.pop("_rascunho_operador_pendente", None)
    if not rasc:
        return False
    restaurar_rascunho_operador(rasc)
    st.session_state["_rascunho_operador_restaurado_msg"] = True
    return True


def buscar_cep(cep: str) -> dict:
    """Consulta ViaCEP via requests com fallback para urllib."""
    cep_limpo = re.sub(r"\D", "", cep or "")
    if len(cep_limpo) != 8:
        return {}
    url = f"https://viacep.com.br/ws/{cep_limpo}/json/"
    # Tentativa 1: requests (mais compatível com Streamlit Cloud)
    try:
        import requests as _req
        resp = _req.get(url, timeout=5)
        if resp.status_code == 200:
            dados = resp.json()
            if not dados.get("erro"):
                return dados
    except Exception:
        pass
    # Tentativa 2: urllib fallback
    try:
        import urllib.request, json as _json
        with urllib.request.urlopen(url, timeout=5) as resp:
            dados = _json.loads(resp.read().decode())
        if not dados.get("erro"):
            return dados
    except Exception:
        pass
    return {}


def filtrar_clientes_por_empresa(clientes: list, empresa_ativa: str) -> list:
    """Filtra clientes por empresa compativel.
    empresa_ativa: 'aqua_gestao' | 'bem_star'
    Critério: servico marcado OU campo empresa do cadastro bate.
    Clientes sem empresa/servico definido aparecem em ambos (legado).
    """
    resultado = []
    for c in clientes:
        servicos = _normalizar_servicos_cliente(c)
        # Campo empresa do cadastro (col M do Sheets)
        _emp = str(c.get("empresa", "") or "").strip().lower()
        _tem_empresa = bool(_emp)
        _is_bem_star = "bem star" in _emp or "bemstar" in _emp
        _is_aqua     = "aqua" in _emp
        _tem_servico_limpeza = servicos.get("limpeza", False)
        _tem_servico_rt      = servicos.get("rt", False)
        _sem_definicao = not _tem_empresa and not _tem_servico_limpeza and not _tem_servico_rt

        if empresa_ativa == "bem_star":
            # Inclui se: servico limpeza marcado, OU empresa = bem star, OU sem definicao (legado)
            if _tem_servico_limpeza or _is_bem_star or _sem_definicao:
                resultado.append(c)
        elif empresa_ativa == "aqua_gestao":
            # Inclui se: servico rt marcado, OU empresa = aqua gestao, OU sem definicao (legado)
            if _tem_servico_rt or _is_aqua or _sem_definicao:
                resultado.append(c)
    return resultado


def _empresa_ativa_codigo() -> str:
    """Retorna o painel administrativo ativo.

    O operador não usa esta escolha; operador é filtrado por PIN/condomínios.
    Fallback: lê admin_empresa_fixa se empresa_ativa não estiver definida corretamente.
    """
    valor = st.session_state.get("empresa_ativa", "")
    if valor not in ("aqua_gestao", "bem_star"):
        # Fallback para empresa fixada no login admin
        valor = st.session_state.get("admin_empresa_fixa", "")
    if valor not in ("aqua_gestao", "bem_star"):
        valor = "aqua_gestao"
    st.session_state["empresa_ativa"] = valor
    return valor

def _empresa_ativa_nome() -> str:
    return "Bem Star Piscinas" if _empresa_ativa_codigo() == "bem_star" else "Aqua Gestão"



def _servicos_padrao_empresa_ativa() -> dict:
    """Serviços gravados automaticamente no cadastro conforme o painel ativo."""
    return {"rt": _empresa_ativa_codigo() == "aqua_gestao", "limpeza": _empresa_ativa_codigo() == "bem_star"}


def _filtrar_clientes_painel_ativo(clientes: list[dict]) -> list[dict]:
    """Mantém os clientes separados visualmente por empresa/painel."""
    return filtrar_clientes_por_empresa(clientes or [], _empresa_ativa_codigo())


def _normalizar_lista_textos_unicos(valores) -> list[str]:
    resultado = []
    vistos = set()
    for item in valores or []:
        valor = re.sub(r"\s+", " ", str(item or "").strip())
        if not valor:
            continue
        chave = _normalizar_chave_acesso(valor)
        if chave in vistos:
            continue
        vistos.add(chave)
        resultado.append(valor)
    return resultado


def _empresa_para_servicos(empresa: str) -> dict:
    emp = _normalizar_chave_acesso(empresa)
    if emp == "ambas":
        return {"rt": True, "limpeza": True}
    if "bem star" in emp:
        return {"rt": False, "limpeza": True}
    return {"rt": True, "limpeza": False}


def _normalizar_servicos_cliente(dados: dict | None) -> dict:
    dados = dados or {}
    servicos = dados.get("servicos") if isinstance(dados.get("servicos"), dict) else {}
    rt = bool(servicos.get("rt"))
    limpeza = bool(servicos.get("limpeza"))
    if not rt and not limpeza:
        legado = _empresa_para_servicos(str(dados.get("empresa", "Aqua Gestão") or "Aqua Gestão"))
        rt = legado["rt"]
        limpeza = legado["limpeza"]
    return {"rt": rt, "limpeza": limpeza}


def _servicos_para_empresa(servicos: dict | None) -> str:
    servicos_norm = _normalizar_servicos_cliente({"servicos": servicos or {}})
    if servicos_norm.get("rt") and servicos_norm.get("limpeza"):
        return "Ambas"
    if servicos_norm.get("limpeza"):
        return "Bem Star Piscinas"
    return "Aqua Gestão"


def _carregar_dados_cliente_local(nome_condominio: str) -> dict:
    nome_limpo = str(nome_condominio or "").strip()
    if not nome_limpo:
        return {}
    pasta = GENERATED_DIR / slugify_nome(nome_limpo)
    if not pasta.exists():
        return {}
    dados = carregar_dados_condominio(pasta) or {}
    if not isinstance(dados, dict):
        return {}
    dados["servicos"] = _normalizar_servicos_cliente(dados)
    dados["operadores_vinculados"] = _normalizar_lista_textos_unicos(dados.get("operadores_vinculados", []))
    dados["empresa"] = _servicos_para_empresa(dados.get("servicos"))
    return dados


def _enriquecer_cliente_com_dados_locais(cliente: dict | None) -> dict:
    cliente_final = dict(cliente or {})
    nome = str(cliente_final.get("nome", "") or "").strip()
    dados_locais = _carregar_dados_cliente_local(nome)

    servicos = _normalizar_servicos_cliente({
        "servicos": dados_locais.get("servicos") if dados_locais else cliente_final.get("servicos"),
        "empresa": dados_locais.get("empresa") if dados_locais else cliente_final.get("empresa", "Aqua Gestão"),
    })
    cliente_final["servicos"] = servicos
    cliente_final["empresa"] = _servicos_para_empresa(servicos)
    cliente_final["operadores_vinculados"] = _normalizar_lista_textos_unicos(
        dados_locais.get("operadores_vinculados", cliente_final.get("operadores_vinculados", []))
    )

    for chave in ["cnpj", "cep", "telefone", "contato", "endereco", "vol_adulto", "vol_infantil", "vol_family", "verificacoes_semanais", "analises_mensais_padrao"]:
        valor_local = dados_locais.get(chave)
        if valor_local not in (None, ""):
            cliente_final[chave] = valor_local

    if dados_locais.get("endereco_condominio"):
        cliente_final["endereco"] = dados_locais.get("endereco_condominio", cliente_final.get("endereco", ""))
    if dados_locais.get("cnpj_condominio"):
        cliente_final["cnpj"] = dados_locais.get("cnpj_condominio", cliente_final.get("cnpj", ""))
    if dados_locais.get("nome_sindico") and not cliente_final.get("contato"):
        cliente_final["contato"] = dados_locais.get("nome_sindico", "")
    if "piscinas_extras" in dados_locais:
        cliente_final["piscinas_extras"] = dados_locais.get("piscinas_extras", [])

    return cliente_final


def slugify_nome(texto: str) -> str:
    texto = (texto or "").strip()
    texto = re.sub(r"[^\w\s-]", "", texto, flags=re.UNICODE)
    texto = re.sub(r"\s+", "_", texto)
    return texto[:120] if texto else "condominio"


def humanizar_nome_pasta(texto: str) -> str:
    texto = (texto or "").strip()
    texto = texto.replace("_", " ").replace("-", " ")
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def limpar_nome_arquivo(texto: str) -> str:
    texto = re.sub(r'[<>:"/\\\\|?*]+', "", texto)
    texto = re.sub(r"\s+", "_", texto.strip())
    return texto[:150]


def carregar_imagem_corrigida_orientacao(origem):
    """Corrige orientação EXIF para preview no Streamlit sem alterar o upload original."""
    try:
        import io as _io
        from PIL import Image as _PILImg, ImageOps as _IOps

        if hasattr(origem, "getbuffer"):
            dados = bytes(origem.getbuffer())
        elif isinstance(origem, (str, Path)):
            dados = Path(origem).read_bytes()
        elif isinstance(origem, bytes):
            dados = origem
        else:
            return origem

        img = _PILImg.open(_io.BytesIO(dados))
        img = _IOps.exif_transpose(img)
        return img
    except Exception:
        return origem


def deduplicar_fotos(lista):
    """Remove fotos duplicadas de uma lista de uploads ou paths.

    Deduplica por hash de conteúdo (uploads) ou por path/string truncada.
    Evita que rascunho + upload atual resultem em foto repetida no PDF.
    """
    resultado = []
    vistos = set()
    import hashlib as _hashlib

    for item in lista or []:
        try:
            if isinstance(item, (str, Path)):
                chave = str(item)[:200]
            elif hasattr(item, "getvalue"):
                pos = item.tell() if hasattr(item, "tell") else None
                dados = item.getvalue()
                if pos is not None:
                    try:
                        item.seek(pos)
                    except Exception:
                        pass
                chave = _hashlib.md5(dados).hexdigest()
            elif hasattr(item, "read"):
                pos = item.tell() if hasattr(item, "tell") else None
                dados = item.read()
                if pos is not None:
                    try:
                        item.seek(pos)
                    except Exception:
                        pass
                chave = _hashlib.md5(dados).hexdigest()
            else:
                chave = str(item)
        except Exception:
            chave = str(item)

        if chave not in vistos:
            vistos.add(chave)
            resultado.append(item)

    return resultado


def hoje_br() -> str:
    return date.today().strftime("%d/%m/%Y")


def apenas_digitos(texto: str) -> str:
    return re.sub(r"\D", "", texto or "")


def formatar_data_hora_arquivo(ts: float) -> str:
    dt = datetime.fromtimestamp(ts)
    return dt.strftime("%d/%m/%Y %H:%M")


def classificar_arquivo(nome_arquivo: str) -> tuple[str, str]:
    nome_lower = nome_arquivo.lower()

    if "contrato" in nome_lower:
        tipo_doc = "Contrato"
    elif "aditivo" in nome_lower:
        tipo_doc = "Aditivo"
    elif "relatorio" in nome_lower:
        tipo_doc = "Relatório"
    else:
        tipo_doc = "Documento"

    if nome_lower.endswith(".pdf"):
        tipo_ext = "PDF"
    elif nome_lower.endswith(".docx"):
        tipo_ext = "DOCX"
    else:
        tipo_ext = "Arquivo"

    return tipo_doc, tipo_ext


def chave_segura(texto: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_]+", "_", texto)


def parse_data_br(texto: str):
    try:
        return datetime.strptime((texto or "").strip(), "%d/%m/%Y").date()
    except Exception:
        return None


def formatar_data_br(dt: date) -> str:
    return dt.strftime("%d/%m/%Y")


def adicionar_um_ano(dt: date) -> date:
    try:
        return dt.replace(year=dt.year + 1)
    except ValueError:
        return dt.replace(month=2, day=28, year=dt.year + 1)


def calcular_renovacao_anual(data_fim_texto: str):
    fim_atual = parse_data_br(data_fim_texto)
    if not fim_atual:
        return None, None

    novo_inicio = fim_atual + timedelta(days=1)
    novo_fim = adicionar_um_ano(novo_inicio) - timedelta(days=1)
    return novo_inicio, novo_fim


def status_vencimento(data_fim_texto: str, alerta_dias: int = 30):
    fim = parse_data_br(data_fim_texto)
    if not fim:
        return {
            "codigo": "indefinido",
            "rotulo": "Sem vigência válida",
            "mensagem": "Data final ausente ou inválida.",
            "dias": None,
            "css": "status-indefinido",
        }

    hoje = date.today()
    dias = (fim - hoje).days

    if dias < 0:
        return {
            "codigo": "vencido",
            "rotulo": "Vencido",
            "mensagem": f"Contrato vencido há {abs(dias)} dia(s).",
            "dias": dias,
            "css": "status-vencido",
        }

    if dias <= alerta_dias:
        return {
            "codigo": "vencendo",
            "rotulo": "Vence em breve",
            "mensagem": f"Contrato vence em {dias} dia(s).",
            "dias": dias,
            "css": "status-vencendo",
        }

    return {
        "codigo": "vigente",
        "rotulo": "Vigente",
        "mensagem": f"Contrato vigente. Restam {dias} dia(s) para o vencimento.",
        "dias": dias,
        "css": "status-vigente",
    }


def texto_dias_restantes(status: dict) -> str:
    dias = status.get("dias")
    if dias is None:
        return "Dias restantes: não disponível"
    if dias < 0:
        return f"Atrasado há {abs(dias)} dia(s)"
    return f"Restam {dias} dia(s)"


def sistema_e_windows() -> bool:
    return platform.system().lower().startswith("win")


def diagnostico_sistema() -> dict:
    return {
        "template_contrato_ok": TEMPLATE_CONTRATO.exists(),
        "template_bem_star_ok": TEMPLATE_BEM_STAR.exists(),
        "template_aditivo_ok": TEMPLATE_ADITIVO.exists(),
        "template_relatorio_ok": TEMPLATE_RELATORIO.exists(),
        "generated_ok": GENERATED_DIR.exists(),
        "logo_ok": encontrar_logo() is not None,
        "windows_ok": sistema_e_windows(),
    }

# =========================================
# MÁSCARAS / FORMATAÇÃO
# =========================================

def formatar_cpf(texto: str) -> str:
    dig = apenas_digitos(texto)[:11]
    if len(dig) <= 3:
        return dig
    if len(dig) <= 6:
        return f"{dig[:3]}.{dig[3:]}"
    if len(dig) <= 9:
        return f"{dig[:3]}.{dig[3:6]}.{dig[6:]}"
    return f"{dig[:3]}.{dig[3:6]}.{dig[6:9]}-{dig[9:]}"


def formatar_cnpj(texto: str) -> str:
    dig = apenas_digitos(texto)[:14]
    if len(dig) <= 2:
        return dig
    if len(dig) <= 5:
        return f"{dig[:2]}.{dig[2:]}"
    if len(dig) <= 8:
        return f"{dig[:2]}.{dig[2:5]}.{dig[5:]}"
    if len(dig) <= 12:
        return f"{dig[:2]}.{dig[2:5]}.{dig[5:8]}/{dig[8:]}"
    return f"{dig[:2]}.{dig[2:5]}.{dig[5:8]}/{dig[8:12]}-{dig[12:]}"


def formatar_telefone(texto: str) -> str:
    dig = apenas_digitos(texto)

    if dig.startswith("55") and len(dig) > 11:
        dig = dig[2:]

    dig = dig[:11]

    if len(dig) <= 2:
        return dig
    if len(dig) <= 6:
        return f"({dig[:2]}) {dig[2:]}"
    if len(dig) <= 10:
        return f"({dig[:2]}) {dig[2:6]}-{dig[6:]}"
    return f"({dig[:2]}) {dig[2:7]}-{dig[7:]}"


def formatar_data_digitada(texto: str) -> str:
    dig = apenas_digitos(texto)[:8]
    if len(dig) <= 2:
        return dig
    if len(dig) <= 4:
        return f"{dig[:2]}/{dig[2:]}"
    return f"{dig[:2]}/{dig[2:4]}/{dig[4:]}"


def moeda_br_sem_prefixo(texto: str) -> str:
    if not texto:
        return ""

    dig = apenas_digitos(str(texto))
    if not dig:
        return ""

    if len(dig) == 1:
        valor = float(f"0.0{dig}")
    elif len(dig) == 2:
        valor = float(f"0.{dig}")
    else:
        valor = float(f"{dig[:-2]}.{dig[-2:]}")

    inteiro = int(valor)
    centavos = int(round((valor - inteiro) * 100))
    inteiro_fmt = f"{inteiro:,}".replace(",", ".")
    return f"{inteiro_fmt},{centavos:02d}"


def moeda_br(texto: str) -> str:
    fmt = moeda_br_sem_prefixo(texto)
    return f"R$ {fmt}" if fmt else ""


def valor_para_template(texto: str) -> str:
    texto = (texto or "").strip()
    if texto.startswith("R$"):
        return texto
    fmt = moeda_br_sem_prefixo(texto)
    return f"R$ {fmt}" if fmt else ""


def on_change_nome_condominio():
    """Salva automaticamente o JSON quando o usuário sai do campo nome_condominio."""
    nome = (st.session_state.get("nome_condominio") or "").strip()
    if nome:
        pasta = GENERATED_DIR / slugify_nome(nome)
        pasta.mkdir(parents=True, exist_ok=True)
        salvar_dados_condominio(pasta, salvar_snapshot_formulario())


def on_change_cpf():
    st.session_state.cpf_sindico = formatar_cpf(st.session_state.get("cpf_sindico", ""))


def _buscar_dados_cnpj(cnpj_digits: str) -> dict | None:
    """Consulta BrasilAPI e retorna dados da empresa, ou None se falhar."""
    try:
        import urllib.request, json as _json
        url = f"https://brasilapi.com.br/api/cnpj/v1/{cnpj_digits}"
        req = urllib.request.Request(url, headers={"User-Agent": "AquaGestaoRT/1.0"})
        with urllib.request.urlopen(req, timeout=6) as resp:
            return _json.loads(resp.read().decode())
    except Exception:
        return None


def on_change_cnpj():
    cnpj_raw = st.session_state.get("cnpj_condominio", "")
    cnpj_fmt = formatar_cnpj(cnpj_raw)
    st.session_state.cnpj_condominio = cnpj_fmt
    digits = re.sub(r"\D", "", cnpj_fmt)
    if len(digits) == 14:
        dados = _buscar_dados_cnpj(digits)
        if dados:
            razao = dados.get("razao_social") or dados.get("nome") or ""
            if razao and not (st.session_state.get("nome_condominio") or "").strip():
                st.session_state.nome_condominio = razao.title()
            logr   = dados.get("logradouro", "")
            num    = dados.get("numero", "")
            compl  = dados.get("complemento", "")
            bairro = dados.get("bairro", "")
            cidade = dados.get("municipio", "")
            uf     = dados.get("uf", "")
            cep_r  = dados.get("cep", "")
            partes = [p for p in [logr, num, compl, bairro, f"{cidade}/{uf}", cep_r] if p and p.strip()]
            end_montado = ", ".join(partes)
            if end_montado and not (st.session_state.get("endereco_condominio") or "").strip():
                st.session_state.endereco_condominio = end_montado
            st.session_state["_cnpj_busca_ok"] = f"Dados encontrados: {razao.title()}"
        else:
            st.session_state["_cnpj_busca_ok"] = ""


def on_change_whatsapp():
    st.session_state.whatsapp_cliente = formatar_telefone(st.session_state.get("whatsapp_cliente", ""))


def on_change_valor_mensal():
    st.session_state.valor_mensal = moeda_br(st.session_state.get("valor_mensal", ""))


def on_change_valor_aditivo():
    st.session_state.valor_aditivo = moeda_br(st.session_state.get("valor_aditivo", ""))


def on_change_data_inicio():
    st.session_state.data_inicio = formatar_data_digitada(st.session_state.get("data_inicio", ""))


def on_change_data_fim():
    st.session_state.data_fim = formatar_data_digitada(st.session_state.get("data_fim", ""))


def on_change_data_assinatura():
    st.session_state.data_assinatura = formatar_data_digitada(st.session_state.get("data_assinatura", ""))

def on_change_bs_cont_data_inicio():
    st.session_state.bs_cont_data_inicio = formatar_data_digitada(st.session_state.get("bs_cont_data_inicio", ""))

def on_change_bs_cont_data_fim():
    st.session_state.bs_cont_data_fim = formatar_data_digitada(st.session_state.get("bs_cont_data_fim", ""))


def on_change_rel_documento_representante():
    atual = st.session_state.get("rel_cpf_cnpj_representante", "")
    dig = apenas_digitos(atual)
    if len(dig) <= 11:
        st.session_state.rel_cpf_cnpj_representante = formatar_cpf(atual)
    else:
        st.session_state.rel_cpf_cnpj_representante = formatar_cnpj(atual)


# =========================================
# VALIDAÇÕES REAIS
# =========================================

def validar_cpf(cpf: str) -> bool:
    cpf = apenas_digitos(cpf)

    if len(cpf) != 11:
        return False
    if cpf == cpf[0] * 11:
        return False

    soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
    dig1 = (soma * 10) % 11
    dig1 = 0 if dig1 == 10 else dig1
    if dig1 != int(cpf[9]):
        return False

    soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
    dig2 = (soma * 10) % 11
    dig2 = 0 if dig2 == 10 else dig2
    return dig2 == int(cpf[10])


def validar_cnpj(cnpj: str) -> bool:
    cnpj = apenas_digitos(cnpj)

    if len(cnpj) != 14:
        return False
    if cnpj == cnpj[0] * 14:
        return False

    pesos1 = [5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma1 = sum(int(cnpj[i]) * pesos1[i] for i in range(12))
    resto1 = soma1 % 11
    dig1 = 0 if resto1 < 2 else 11 - resto1
    if dig1 != int(cnpj[12]):
        return False

    pesos2 = [6] + pesos1
    soma2 = sum(int(cnpj[i]) * pesos2[i] for i in range(13))
    resto2 = soma2 % 11
    dig2 = 0 if resto2 < 2 else 11 - resto2
    return dig2 == int(cnpj[13])


def validar_data_br(texto: str) -> bool:
    try:
        datetime.strptime(texto.strip(), "%d/%m/%Y")
        return True
    except Exception:
        return False


def validar_email(email: str) -> bool:
    email = (email or "").strip()
    if not email:
        return True
    padrao = r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$"
    return re.match(padrao, email) is not None


def validar_campos_formato(dados: dict, email_cliente: str) -> list[str]:
    erros = []

    if not validar_cpf(dados["CPF_SINDICO"]):
        erros.append("CPF do síndico/representante inválido.")
    if not validar_cnpj(dados["CNPJ_CONDOMINIO"]):
        erros.append("CNPJ do condomínio inválido.")
    if not validar_data_br(dados["DATA_ASSINATURA"]):
        erros.append("Data de assinatura inválida. Use o formato dd/mm/aaaa.")
    if not validar_data_br(dados["DATA_INICIO"]):
        erros.append("Data de início inválida. Use o formato dd/mm/aaaa.")
    if not validar_data_br(dados["DATA_FIM"]):
        erros.append("Data de fim inválida. Use o formato dd/mm/aaaa.")
    if email_cliente.strip() and not validar_email(email_cliente):
        erros.append("E-mail inválido.")

    return erros


# =========================================
# SNAPSHOT / MANIFEST / CADASTRO
# =========================================

def salvar_snapshot_formulario() -> dict:
    return {
        "nome_condominio": (st.session_state.get("nome_condominio") or "").strip(),
        "cnpj_condominio": (st.session_state.get("cnpj_condominio") or "").strip(),
        "endereco_condominio": (st.session_state.get("endereco_condominio") or "").strip(),
        "nome_sindico": (st.session_state.get("nome_sindico") or "").strip(),
        "cpf_sindico": (st.session_state.get("cpf_sindico") or "").strip(),
        "valor_mensal": (st.session_state.get("valor_mensal") or "").strip(),
        "valor_aditivo": (st.session_state.get("valor_aditivo") or "").strip(),
        "data_inicio": (st.session_state.get("data_inicio") or "").strip(),
        "data_fim": (st.session_state.get("data_fim") or "").strip(),
        "data_assinatura": (st.session_state.get("data_assinatura") or "").strip(),
        "whatsapp_cliente": (st.session_state.get("whatsapp_cliente") or "").strip(),
        "email_cliente": (st.session_state.get("email_cliente") or "").strip(),
        "observacoes_internas": (st.session_state.get("observacoes_internas") or "").strip(),
        "rel_art_status": (st.session_state.get("rel_art_status") or "Emitida").strip(),
        "rel_art_numero": (st.session_state.get("rel_art_numero") or "").strip(),
        "rel_art_inicio": (st.session_state.get("rel_art_inicio") or "").strip(),
        "rel_art_fim": (st.session_state.get("rel_art_fim") or "").strip(),
        "parametros_ultimos": obter_parametros_ultimos_relatorio(),
        "dosagens_ultimas": obter_dosagens_ultimas_relatorio(),
        "salvo_em": _agora_brasilia(),
    }


def obter_parametros_ultimos_relatorio() -> list[dict]:
    """Coleta as análises físico-químicas atuais para reutilizar como último padrão do condomínio."""
    itens = []
    qtd = int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO)
    qtd = max(ANALISES_PADRAO, min(qtd, ANALISES_MAX_SUGERIDO))
    for i in range(qtd):
        itens.append({
            "data": (st.session_state.get(f"rel_analise_data_{i}") or "").strip(),
            "ph": (st.session_state.get(f"rel_analise_ph_{i}") or "").strip(),
            "cl": (st.session_state.get(f"rel_analise_cl_{i}") or "").strip(),
            "ct": (st.session_state.get(f"rel_analise_ct_{i}") or "").strip(),
            "alc": (st.session_state.get(f"rel_analise_alc_{i}") or "").strip(),
            "dc": (st.session_state.get(f"rel_analise_dc_{i}") or "").strip(),
            "cya": (st.session_state.get(f"rel_analise_cya_{i}") or "").strip(),
            "operador": (st.session_state.get(f"rel_analise_operador_{i}") or "").strip(),
        })
    return itens


def aplicar_parametros_ultimos_no_relatorio(dados_salvos: dict | None):
    """Restaura as últimas análises físico-químicas salvas para o condomínio."""
    parametros = []
    if isinstance(dados_salvos, dict):
        parametros = (
            dados_salvos.get("parametros_ultimos")
            or dados_salvos.get("analises_ultimas")
            or dados_salvos.get("analises")
            or []
        )
    if not isinstance(parametros, list):
        parametros = []

    qtd = max(len(parametros), ANALISES_PADRAO)
    qtd = max(ANALISES_PADRAO, min(qtd, ANALISES_MAX_SUGERIDO))
    garantir_campos_analises(qtd)
    st.session_state["rel_analises_total"] = qtd

    for i in range(qtd):
        item = parametros[i] if i < len(parametros) and isinstance(parametros[i], dict) else {}
        st.session_state[f"rel_analise_data_{i}"] = str(item.get("data") or "").strip()
        st.session_state[f"rel_analise_ph_{i}"] = str(item.get("ph") or "").strip()
        st.session_state[f"rel_analise_cl_{i}"] = str(item.get("cl") or item.get("cloro_livre") or "").strip()
        st.session_state[f"rel_analise_ct_{i}"] = str(item.get("ct") or item.get("cloro_total") or "").strip()
        st.session_state[f"rel_analise_alc_{i}"] = str(item.get("alc") or item.get("alcalinidade") or "").strip()
        st.session_state[f"rel_analise_dc_{i}"] = str(item.get("dc") or item.get("dureza") or "").strip()
        st.session_state[f"rel_analise_cya_{i}"] = str(item.get("cya") or item.get("cianurico") or "").strip()
        st.session_state[f"rel_analise_operador_{i}"] = str(item.get("operador") or "").strip()


def obter_dosagens_ultimas_relatorio() -> list[dict]:
    itens = []
    for i in range(7):
        itens.append({
            "produto": (st.session_state.get(f"rel_dos_produto_{i}") or "").strip(),
            "fabricante_lote": (st.session_state.get(f"rel_dos_lote_{i}") or "").strip(),
            "quantidade": (st.session_state.get(f"rel_dos_qtd_{i}") or "").strip(),
            "unidade": (st.session_state.get(f"rel_dos_un_{i}") or "").strip(),
            "finalidade": (st.session_state.get(f"rel_dos_finalidade_{i}") or "").strip(),
        })
    return itens


def aplicar_dosagens_ultimas_no_relatorio(dados_salvos: dict | None):
    dosagens = []
    if isinstance(dados_salvos, dict):
        dosagens = dados_salvos.get("dosagens_ultimas") or []

    for i in range(7):
        item = dosagens[i] if i < len(dosagens) and isinstance(dosagens[i], dict) else {}
        st.session_state[f"rel_dos_produto_{i}"] = (item.get("produto") or "").strip()
        st.session_state[f"rel_dos_lote_{i}"] = (item.get("fabricante_lote") or "").strip()
        st.session_state[f"rel_dos_qtd_{i}"] = (item.get("quantidade") or "").strip()
        st.session_state[f"rel_dos_un_{i}"] = (item.get("unidade") or "").strip()
        st.session_state[f"rel_dos_finalidade_{i}"] = (item.get("finalidade") or "").strip()


def aplicar_snapshot_relatorio_independente(dados: dict):
    for key, value in dados.items():
        if key not in ("dosagens_ultimas", "parametros_ultimos", "analises_ultimas", "analises"):
            st.session_state[key] = value
    aplicar_parametros_ultimos_no_relatorio(dados)
    aplicar_dosagens_ultimas_no_relatorio(dados)

def obter_snapshot_relatorio_independente() -> dict:
    return {
        "nome_condominio": (st.session_state.get("rel_nome_condominio") or "").strip(),
        "cnpj_condominio": (st.session_state.get("rel_cnpj_condominio") or "").strip(),
        "endereco_condominio": (st.session_state.get("rel_endereco_condominio") or "").strip(),
        "nome_sindico": (st.session_state.get("rel_representante") or "").strip(),
        "cpf_sindico": (st.session_state.get("rel_cpf_cnpj_representante") or "").strip(),
        "rel_art_status": (st.session_state.get("rel_art_status") or "Emitida").strip(),
        "rel_art_numero": (st.session_state.get("rel_art_numero") or "").strip(),
        "rel_art_inicio": (st.session_state.get("rel_art_inicio") or "").strip(),
        "rel_art_fim": (st.session_state.get("rel_art_fim") or "").strip(),
        "ultima_origem_relatorio": (st.session_state.get("rel_tipo_atendimento") or "").strip(),
        "parametros_ultimos": obter_parametros_ultimos_relatorio(),
        "dosagens_ultimas": obter_dosagens_ultimas_relatorio(),
        "salvo_em": _agora_brasilia(),
    }


def sincronizar_relatorio_com_cadastro():
    if not (st.session_state.get("rel_nome_condominio") or "").strip():
        st.session_state.rel_nome_condominio = (st.session_state.get("nome_condominio") or "").strip()
    if not (st.session_state.get("rel_cnpj_condominio") or "").strip():
        st.session_state.rel_cnpj_condominio = (st.session_state.get("cnpj_condominio") or "").strip()
    if not (st.session_state.get("rel_endereco_condominio") or "").strip():
        st.session_state.rel_endereco_condominio = (st.session_state.get("endereco_condominio") or "").strip()
    if not (st.session_state.get("rel_representante") or "").strip():
        st.session_state.rel_representante = (st.session_state.get("nome_sindico") or "").strip()
    if not (st.session_state.get("rel_cpf_cnpj_representante") or "").strip():
        cpf = (st.session_state.get("cpf_sindico") or "").strip()
        cnpj = (st.session_state.get("cnpj_condominio") or "").strip()
        st.session_state.rel_cpf_cnpj_representante = cpf or cnpj


def carregar_dados_cadastro_no_relatorio():
    _nome_base_cad = (st.session_state.get("nome_condominio") or "").strip()
    _dados_locais_cad = _carregar_dados_cliente_local(_nome_base_cad) if _nome_base_cad else {}
    _freq_cad = obter_verificacoes_semanais_cliente(_dados_locais_cad)
    st.session_state["rel_verificacoes_semanais"] = _freq_cad
    st.session_state["rel_analises_total"] = calcular_linhas_analises_por_frequencia(_freq_cad)
    st.session_state.rel_nome_condominio = _nome_base_cad
    st.session_state.rel_cnpj_condominio = (st.session_state.get("cnpj_condominio") or "").strip()
    st.session_state.rel_endereco_condominio = (st.session_state.get("endereco_condominio") or "").strip()
    st.session_state.rel_representante = (st.session_state.get("nome_sindico") or "").strip()
    st.session_state.rel_cpf_cnpj_representante = (st.session_state.get("cpf_sindico") or "").strip() or (st.session_state.get("cnpj_condominio") or "").strip()
    aplicar_dosagens_ultimas_no_relatorio(salvar_snapshot_formulario())


def salvar_relatorio_no_cadastro_principal():
    st.session_state.nome_condominio = (st.session_state.get("rel_nome_condominio") or "").strip()
    st.session_state.cnpj_condominio = (st.session_state.get("rel_cnpj_condominio") or "").strip()
    st.session_state.endereco_condominio = (st.session_state.get("rel_endereco_condominio") or "").strip()
    st.session_state.nome_sindico = (st.session_state.get("rel_representante") or "").strip()
    doc = (st.session_state.get("rel_cpf_cnpj_representante") or "").strip()
    dig = apenas_digitos(doc)
    if len(dig) == 11:
        st.session_state.cpf_sindico = formatar_cpf(doc)
    elif len(dig) == 14:
        st.session_state.cnpj_condominio = formatar_cnpj(doc)


def validar_campos_obrigatorios(dados: dict) -> list[str]:
    campos = {
        "Nome do condomínio": dados.get("NOME_CONDOMINIO", ""),
        "CNPJ do condomínio": dados.get("CNPJ_CONDOMINIO", ""),
        "Endereço do condomínio": dados.get("ENDERECO_CONDOMINIO", ""),
        "Nome do síndico / representante": dados.get("NOME_SINDICO", ""),
        "CPF do síndico / representante": dados.get("CPF_SINDICO", ""),
        "Valor mensal": dados.get("VALOR_MENSAL", ""),
        "Data de início": dados.get("DATA_INICIO", ""),
        "Data de fim": dados.get("DATA_FIM", ""),
        "Data de assinatura": dados.get("DATA_ASSINATURA", ""),
    }
    faltando = [nome for nome, valor in campos.items() if not (valor or "").strip()]
    return faltando


def carregar_manifest_condominio(pasta_condominio: Path) -> dict:
    caminho = pasta_condominio / MANIFEST_JSON_NAME
    if caminho.exists():
        try:
            return json.loads(caminho.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {"condominio": pasta_condominio.name, "documentos": []}


def salvar_manifest_condominio(pasta_condominio: Path, manifest: dict):
    caminho = pasta_condominio / MANIFEST_JSON_NAME
    caminho.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def registrar_documento_manifest(pasta_condominio: Path, nome_condominio: str, tipo: str, arquivo_docx: Path | None, arquivo_pdf: Path | None, pdf_gerado: bool, erro_pdf: str | None, dados_utilizados: dict | None = None, extras: dict | None = None):
    manifest = carregar_manifest_condominio(pasta_condominio)
    documento = {
        "tipo": tipo,
        "nome_condominio": nome_condominio,
        "gerado_em": _agora_brasilia(),
        "arquivo_docx": arquivo_docx.name if arquivo_docx and arquivo_docx.exists() else None,
        "arquivo_pdf": arquivo_pdf.name if arquivo_pdf and arquivo_pdf.exists() else None,
        "pdf_gerado": bool(pdf_gerado),
        "erro_pdf": erro_pdf,
        "dados_utilizados": dados_utilizados or {},
        "extras": extras or {},
    }
    manifest.setdefault("documentos", []).append(documento)
    manifest["ultimo_update"] = _agora_brasilia()
    salvar_manifest_condominio(pasta_condominio, manifest)


def obter_ultimo_documento_manifest(pasta_condominio: Path, tipo: str) -> dict | None:
    manifest = carregar_manifest_condominio(pasta_condominio)
    docs = [d for d in manifest.get("documentos", []) if d.get("tipo") == tipo]
    return docs[-1] if docs else None


# =========================================
# PERSISTÊNCIA DE DADOS DO CONDOMÍNIO
# =========================================

def salvar_dados_condominio(pasta_condominio: Path, dados_para_salvar: dict):
    caminho_json = pasta_condominio / DADOS_JSON_NAME
    with open(caminho_json, "w", encoding="utf-8") as f:
        json.dump(dados_para_salvar, f, ensure_ascii=False, indent=2)


def carregar_dados_condominio(pasta_condominio: Path) -> dict | None:
    caminho_json = pasta_condominio / DADOS_JSON_NAME
    if not caminho_json.exists():
        return None
    try:
        with open(caminho_json, "r", encoding="utf-8") as f:
            dados = json.load(f)
        return dados if isinstance(dados, dict) else None
    except Exception:
        return None


def aplicar_dados_no_formulario(dados_salvos: dict):
    st.session_state.nome_condominio = dados_salvos.get("nome_condominio", "")
    st.session_state.cnpj_condominio = dados_salvos.get("cnpj_condominio", "")
    st.session_state.endereco_condominio = dados_salvos.get("endereco_condominio", "")
    st.session_state.nome_sindico = dados_salvos.get("nome_sindico", "")
    st.session_state.cpf_sindico = dados_salvos.get("cpf_sindico", "")
    st.session_state.valor_mensal = dados_salvos.get("valor_mensal", "")
    st.session_state.valor_aditivo = dados_salvos.get("valor_aditivo", "")
    st.session_state.data_inicio = dados_salvos.get("data_inicio", "")
    st.session_state.data_fim = dados_salvos.get("data_fim", "")
    st.session_state.data_assinatura = dados_salvos.get("data_assinatura", hoje_br())
    st.session_state.whatsapp_cliente = dados_salvos.get("whatsapp_cliente", "")
    st.session_state.email_cliente = dados_salvos.get("email_cliente", "")
    st.session_state.observacoes_internas = dados_salvos.get("observacoes_internas", "")
    st.session_state.cargo_sindico = dados_salvos.get("cargo_sindico", "Síndico")
    st.session_state.frequencia_visitas = dados_salvos.get("frequencia_visitas", "1 (uma)")
    st.session_state.dia_pagamento = dados_salvos.get("dia_pagamento", "")
    st.session_state.forma_pagamento = dados_salvos.get("forma_pagamento", "Pix")
    st.session_state.valor_mensal_extenso = dados_salvos.get("valor_mensal_extenso", "")
    st.session_state.rel_art_status = dados_salvos.get("rel_art_status", "Emitida")
    st.session_state.rel_art_status_widget = st.session_state.rel_art_status
    st.session_state.rel_art_numero = dados_salvos.get("rel_art_numero", "")
    st.session_state.rel_art_inicio = dados_salvos.get("rel_art_inicio", "")
    st.session_state.rel_art_fim = dados_salvos.get("rel_art_fim", "")
    st.session_state.origem_dados_carregados = dados_salvos.get("nome_condominio", "")

    carregar_dados_cadastro_no_relatorio()
    aplicar_parametros_ultimos_no_relatorio(dados_salvos)
    aplicar_dosagens_ultimas_no_relatorio(dados_salvos)


def preparar_cadastro_legado(nome_pasta: str):
    limpar_formulario()
    st.session_state.nome_condominio = humanizar_nome_pasta(nome_pasta)
    st.session_state.origem_dados_carregados = f"{humanizar_nome_pasta(nome_pasta)} (cadastro criado de pasta antiga)"


def preparar_renovacao_no_formulario(dados_salvos: dict) -> tuple[bool, str]:
    data_fim_atual = dados_salvos.get("data_fim", "")
    novo_inicio, novo_fim = calcular_renovacao_anual(data_fim_atual)

    if not novo_inicio or not novo_fim:
        return False, "Não foi possível renovar. A data final salva está ausente ou inválida."

    aplicar_dados_no_formulario(dados_salvos)
    st.session_state.data_inicio = formatar_data_br(novo_inicio)
    st.session_state.data_fim = formatar_data_br(novo_fim)
    st.session_state.data_assinatura = hoje_br()
    st.session_state.origem_dados_carregados = dados_salvos.get("nome_condominio", "")
    return True, "Nova vigência anual preenchida no formulário."


def obter_pasta_atual_do_formulario() -> Path | None:
    nome_condominio = (st.session_state.get("nome_condominio") or "").strip()
    if not nome_condominio:
        return None
    return GENERATED_DIR / slugify_nome(nome_condominio)


def salvar_cadastro_sem_gerar_documentos() -> tuple[bool, str]:
    nome_condominio = (st.session_state.get("nome_condominio") or "").strip()
    if not nome_condominio:
        return False, "Informe ao menos o nome do condomínio para salvar o cadastro."

    dados_base = {
        "DATA_ASSINATURA": (st.session_state.get("data_assinatura") or "").strip(),
        "NOME_CONDOMINIO": nome_condominio,
        "CNPJ_CONDOMINIO": (st.session_state.get("cnpj_condominio") or "").strip(),
        "ENDERECO_CONDOMINIO": (st.session_state.get("endereco_condominio") or "").strip(),
        "NOME_SINDICO": (st.session_state.get("nome_sindico") or "").strip(),
        "CPF_SINDICO": (st.session_state.get("cpf_sindico") or "").strip(),
        "VALOR_MENSAL": valor_para_template((st.session_state.get("valor_mensal") or "").strip()),
        "VALOR_ADITIVO": valor_para_template((st.session_state.get("valor_aditivo") or "").strip()),
        "DATA_INICIO": (st.session_state.get("data_inicio") or "").strip(),
        "DATA_FIM": (st.session_state.get("data_fim") or "").strip(),
    }

    erros = validar_para_geracao(dados_base, (st.session_state.get("email_cliente") or "").strip())
    if erros:
        return False, "Corrija os campos antes de salvar o cadastro: " + " | ".join(erros)

    pasta_condominio = obter_pasta_atual_do_formulario()
    assert pasta_condominio is not None
    pasta_condominio.mkdir(parents=True, exist_ok=True)
    salvar_dados_condominio(pasta_condominio, salvar_snapshot_formulario())
    return True, f"Cadastro salvo/atualizado com sucesso em '{pasta_condominio.name}'."


# =========================================
# HISTÓRICO / PAINEL
# =========================================

def listar_arquivos_pasta(pasta: Path):
    if not pasta.exists():
        return []

    arquivos = []
    for p in pasta.iterdir():
        if p.is_file():
            if p.name in {DADOS_JSON_NAME, MANIFEST_JSON_NAME}:
                continue

            tipo_doc, tipo_ext = classificar_arquivo(p.name)
            arquivos.append(
                {
                    "path": p,
                    "name": p.name,
                    "tipo_doc": tipo_doc,
                    "tipo_ext": tipo_ext,
                    "modificado_em": formatar_data_hora_arquivo(p.stat().st_mtime),
                    "ts": p.stat().st_mtime,
                }
            )

    return sorted(arquivos, key=lambda x: x["ts"], reverse=True)


def localizar_ultimo_documento(arquivos: list, tipo_doc: str, extensao_preferida: str = "PDF"):
    preferidos = [a for a in arquivos if a["tipo_doc"] == tipo_doc and a["tipo_ext"] == extensao_preferida]
    if preferidos:
        return preferidos[0]
    alternativos = [a for a in arquivos if a["tipo_doc"] == tipo_doc]
    return alternativos[0] if alternativos else None


def listar_historico():
    if not GENERATED_DIR.exists():
        return []

    pastas = [p for p in GENERATED_DIR.iterdir() if p.is_dir()]
    pastas = sorted(pastas, key=lambda p: p.stat().st_mtime, reverse=True)

    historico = []
    for pasta in pastas:
        arquivos = listar_arquivos_pasta(pasta)
        dados_salvos = carregar_dados_condominio(pasta)
        data_fim = dados_salvos.get("data_fim", "") if dados_salvos else ""
        status = status_vencimento(data_fim)

        historico.append(
            {
                "nome": pasta.name,
                "pasta": pasta,
                "arquivos": arquivos,
                "total_arquivos": len(arquivos),
                "tem_dados_salvos": (pasta / DADOS_JSON_NAME).exists(),
                "status_vencimento": status,
                "data_fim": data_fim,
            }
        )
    return historico


def listar_painel_vencimentos(alerta_dias: int):
    if not GENERATED_DIR.exists():
        return []

    pastas = [p for p in GENERATED_DIR.iterdir() if p.is_dir()]
    itens = []

    for pasta in pastas:
        dados_salvos = carregar_dados_condominio(pasta)
        arquivos = listar_arquivos_pasta(pasta)

        nome_exibicao = pasta.name
        data_fim = ""
        status = {
            "codigo": "indefinido",
            "rotulo": "Sem vigência válida",
            "mensagem": "Data final ausente ou inválida.",
            "dias": None,
            "css": "status-indefinido",
        }
        origem = "legado_sem_json"

        if dados_salvos:
            nome_exibicao = dados_salvos.get("nome_condominio", pasta.name)
            data_fim = dados_salvos.get("data_fim", "")
            status = status_vencimento(data_fim, alerta_dias)
            origem = "json"

        ultimo_contrato = localizar_ultimo_documento(arquivos, "Contrato")
        ultimo_aditivo = localizar_ultimo_documento(arquivos, "Aditivo")
        ultimo_relatorio = localizar_ultimo_documento(arquivos, "Relatório")

        itens.append(
            {
                "nome_exibicao": nome_exibicao,
                "slug": pasta.name,
                "pasta": pasta,
                "dados": dados_salvos,
                "arquivos": arquivos,
                "data_fim": data_fim,
                "data_fim_dt": parse_data_br(data_fim),
                "status": status,
                "origem": origem,
                "tem_json": dados_salvos is not None,
                "ultimo_contrato": ultimo_contrato,
                "ultimo_aditivo": ultimo_aditivo,
                "ultimo_relatorio": ultimo_relatorio,
            }
        )

    def ordem_item(item):
        status_codigo = item["status"]["codigo"]
        dt_fim = item["data_fim_dt"] or date.max
        dias = item["status"]["dias"]
        dias_sort = dias if dias is not None else 999999

        prioridade = {
            "vencido": 0,
            "vencendo": 1,
            "vigente": 2,
            "indefinido": 3,
        }.get(status_codigo, 9)

        return (prioridade, dt_fim, dias_sort, item["nome_exibicao"].lower())

    return sorted(itens, key=ordem_item)


def filtrar_itens_painel(itens: list, termo: str, filtro_status: str):
    resultado = itens

    termo = (termo or "").strip().lower()
    if termo:
        filtrados = []
        for item in resultado:
            dados = item["dados"] or {}
            campos = [
                item["nome_exibicao"],
                item["slug"],
                dados.get("nome_condominio", ""),
                dados.get("cnpj_condominio", ""),
                dados.get("nome_sindico", ""),
                item["status"]["rotulo"],
            ]
            texto_unico = " ".join(campos).lower()
            if termo in texto_unico:
                filtrados.append(item)
        resultado = filtrados

    mapa_status = {
        "Todos": None,
        "Vigente": "vigente",
        "Vence em breve": "vencendo",
        "Vencido": "vencido",
        "Sem vigência válida": "indefinido",
    }

    codigo = mapa_status.get(filtro_status)
    if codigo:
        resultado = [i for i in resultado if i["status"]["codigo"] == codigo]

    return resultado


def abrir_pasta_windows(caminho: Path):
    try:
        os.startfile(str(caminho))
    except Exception as e:
        st.error(f"Não foi possível abrir a pasta no Windows: {e}")


def abrir_arquivo_windows(caminho: Path):
    try:
        os.startfile(str(caminho))
    except Exception as e:
        st.error(f"Não foi possível abrir o arquivo: {e}")


def deletar_pasta_condominio(pasta: Path):
    if pasta.exists() and pasta.is_dir():
        shutil.rmtree(pasta)


def deletar_arquivo_individual(arquivo: Path):
    pasta = arquivo.parent
    if arquivo.exists() and arquivo.is_file():
        arquivo.unlink()

    if pasta.exists() and pasta.is_dir():
        restantes = list(pasta.iterdir())
        if len(restantes) == 0:
            pasta.rmdir()


# =========================================
# DOCX / PLACEHOLDERS
# =========================================

def substituir_em_paragrafo(paragraph, mapa: dict[str, str]):
    texto_original = paragraph.text
    texto_novo = texto_original
    alterou = False

    for chave, valor in mapa.items():
        if chave in texto_novo:
            texto_novo = texto_novo.replace(chave, valor)
            alterou = True

    if alterou:
        # Clear existing runs and add a new one with the replaced text
        # This handles cases where placeholders might be split across multiple runs
        for i in range(len(paragraph.runs) - 1, -1, -1):
            p = paragraph.runs[i]._element
            p.getparent().remove(p)
        paragraph.add_run(texto_novo)


def substituir_placeholders_doc(doc: Document, mapa: dict[str, str]):
    for paragraph in doc.paragraphs:
        substituir_em_paragrafo(paragraph, mapa)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_em_paragrafo(paragraph, mapa)

    for section in doc.sections:
        header = section.header
        footer = section.footer

        for paragraph in header.paragraphs:
            substituir_em_paragrafo(paragraph, mapa)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        substituir_em_paragrafo(paragraph, mapa)

        for paragraph in footer.paragraphs:
            substituir_em_paragrafo(paragraph, mapa)

        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        substituir_em_paragrafo(paragraph, mapa)


def adicionar_espaco(doc: Document, qtd: int = 1):
    for _ in range(qtd):
        doc.add_paragraph("")


def preencher_celula(cell, linhas, negrito_idx=None):
    negrito_idx = negrito_idx or []
    cell.paragraphs[0].clear()
    primeira = True
    for i, linha in enumerate(linhas):
        p = cell.paragraphs[0] if primeira else cell.add_paragraph()
        primeira = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(linha)
        run.font.size = Pt(9)
        run.bold = (i in negrito_idx)


def converter_docx_para_pdf(docx_path: Path, pdf_path: Path):
    """
    Converte DOCX para PDF.
    - Windows (local): usa docx2pdf + pythoncom
    - Linux (Streamlit Cloud): usa LibreOffice headless ou reportlab fallback
    """
    import platform, shutil, subprocess
    _sistema = platform.system()

    # ── Tentativa 1: Windows — docx2pdf com pythoncom ────────────────────────
    if _sistema == "Windows":
        try:
            import pythoncom
            from docx2pdf import convert as _conv
            pythoncom.CoInitialize()
            _conv(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                return True, None
        except Exception:
            pass

    # ── Tentativa 2: LibreOffice headless (Linux/Cloud) ───────────────────────
    try:
        _lo = shutil.which("libreoffice") or shutil.which("soffice")
        if _lo:
            subprocess.run(
                [_lo, "--headless", "--convert-to", "pdf",
                 "--outdir", str(pdf_path.parent), str(docx_path)],
                capture_output=True, timeout=60
            )
            _pdf_lo = pdf_path.parent / (docx_path.stem + ".pdf")
            if _pdf_lo.exists():
                if _pdf_lo != pdf_path:
                    _pdf_lo.rename(pdf_path)
                return True, None
    except Exception:
        pass

    # ── Tentativa 3: docx2pdf sem pythoncom (Linux) ───────────────────────────
    if _sistema != "Windows":
        try:
            from docx2pdf import convert as _conv2
            _conv2(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                return True, None
        except Exception:
            pass

    # ── Fallback: ReportLab extraindo texto + imagens do DOCX ────────────────
    try:
        from docx import Document as _DocxR
        from docx.oxml.ns import qn as _qn
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.platypus import (SimpleDocTemplate, Paragraph,
            Spacer, Image as RLImage, HRFlowable, Table, TableStyle)
        import io as _io, base64 as _b64

        _doc_r = _DocxR(str(docx_path))
        _styles = getSampleStyleSheet()
        _s_title = ParagraphStyle("t", parent=_styles["Heading1"],
            fontSize=13, alignment=TA_CENTER,
            textColor=colors.HexColor("#0d3d75"), spaceAfter=6)
        _s_h2 = ParagraphStyle("h2", parent=_styles["Heading2"],
            fontSize=11, textColor=colors.HexColor("#0d3d75"),
            spaceBefore=10, spaceAfter=4)
        _s_body = ParagraphStyle("b", parent=_styles["Normal"],
            fontSize=9.5, alignment=TA_JUSTIFY, leading=14, spaceAfter=3)
        _s_center = ParagraphStyle("c", parent=_styles["Normal"],
            fontSize=9, alignment=TA_CENTER, spaceAfter=3)

        _story = []

        # Extrai imagens embutidas no DOCX
        _img_map = {}
        for _rel in _doc_r.part.rels.values():
            if "image" in _rel.reltype:
                try:
                    _img_bytes = _rel.target_part.blob
                    _img_map[_rel.rId] = _img_bytes
                except Exception:
                    pass

        def _safe_para(txt, style):
            try:
                # Escapar caracteres XML problemáticos
                txt = (txt.replace("&","&amp;").replace("<","&lt;")
                          .replace(">","&gt;").replace('"','&quot;'))
                return Paragraph(txt, style)
            except Exception:
                return Paragraph("—", style)

        for _para in _doc_r.paragraphs:
            _txt = _para.text.strip()
            # Verifica se tem imagem no parágrafo
            _has_img = _para._element.findall(
                ".//" + _qn("a:blip"), _para._element.nsmap
                if hasattr(_para._element, "nsmap") else {}
            )
            if not _txt and not _has_img:
                _story.append(Spacer(1, 0.15*cm))
                continue
            _style_name = _para.style.name if _para.style else ""
            if "Heading 1" in _style_name or "Title" in _style_name:
                _story.append(_safe_para(_txt, _s_title))
            elif "Heading" in _style_name:
                _story.append(_safe_para(_txt, _s_h2))
            elif _txt:
                _align = _para.alignment
                _s = _s_center if _align == 1 else _s_body
                _story.append(_safe_para(_txt, _s))

        # Adiciona imagens separadamente das relações
        for _rId, _img_bytes in list(_img_map.items())[:20]:
            try:
                _buf = _io.BytesIO(_img_bytes)
                _img_rl = RLImage(_buf, width=14*cm, kind="proportional")
                _img_rl.hAlign = "CENTER"
                _story.append(Spacer(1, 0.2*cm))
                _story.append(_img_rl)
                _story.append(Spacer(1, 0.2*cm))
            except Exception:
                pass

        _doc_rl = SimpleDocTemplate(
            str(pdf_path), pagesize=A4,
            topMargin=2*cm, bottomMargin=2*cm,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
        )
        _doc_rl.build(_story)
        if pdf_path.exists():
            return True, None
    except Exception as _e:
        return False, str(_e)

    return False, "Nenhum método de conversão disponível"


def gerar_documento(template_path: Path, output_docx: Path, placeholders: dict, incluir_assinaturas: bool = False):
    if not template_path.exists():
        raise FileNotFoundError(f"Template não encontrado: {template_path.name}")

    doc = Document(str(template_path))
    
    substituir_placeholders_doc(doc, placeholders)

    doc.save(str(output_docx))




# =========================================
# ADITIVO RT — PDF PREMIUM REPORTLAB
# =========================================

def salvar_aditivo_rt_pdf_premium_reportlab(placeholders: dict, pdf_path: Path) -> tuple[bool, str | None]:
    """Gera o Termo Aditivo em PDF premium Aqua Gestão, sem depender do Word/LibreOffice."""
    try:
        import html
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, HRFlowable, KeepTogether

        pdf_path = Path(pdf_path)
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        azul_escuro = colors.HexColor("#0B2E59")
        azul_claro = colors.HexColor("#EAF4FF")
        dourado = colors.HexColor("#C9A227")
        cinza = colors.HexColor("#4B5563")
        borda = colors.HexColor("#D7E6F5")

        def ph(chave: str, padrao: str = "") -> str:
            valor = str(placeholders.get(chave, "") or "").strip()
            return valor if valor else padrao

        nome_cond = ph("{{NOME_CONDOMINIO}}", ph("{{NOME_CONTRATANTE}}", "CONTRATANTE"))
        cnpj_cond = ph("{{CNPJ_CONDOMINIO}}", ph("{{CPF_CNPJ_CONTRATANTE}}", ""))
        endereco_cond = ph("{{ENDERECO_CONDOMINIO}}", ph("{{ENDERECO_CONTRATANTE}}", ""))
        sindico = ph("{{NOME_SINDICO}}", "")
        cpf_sindico = ph("{{CPF_SINDICO}}", "")
        cargo_sindico = ph("{{CARGO_SINDICO}}", "Síndico")
        valor_mensal = ph("{{VALOR_MENSAL}}", "R$ —")
        valor_aditivo = ph("{{VALOR_ADITIVO}}", "R$ —")
        data_inicio = ph("{{DATA_INICIO_CONTRATO}}", ph("{{DATA_INICIO}}", ""))
        data_fim = ph("{{DATA_FIM_CONTRATO}}", ph("{{DATA_FIM}}", ""))
        data_ass = ph("{{DATA_ASSINATURA}}", "")
        local_data = ph("{{LOCAL_DATA_ASSINATURA}}", f"Uberlândia/MG, {data_ass}" if data_ass else "Uberlândia/MG")

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="AquaTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=15, leading=19, alignment=TA_CENTER, textColor=azul_escuro, spaceAfter=6))
        styles.add(ParagraphStyle(name="AquaH", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=14, textColor=azul_escuro, spaceBefore=8, spaceAfter=4))
        styles.add(ParagraphStyle(name="AquaBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=13.2, alignment=TA_JUSTIFY, textColor=colors.HexColor("#1F2937"), spaceAfter=5))
        styles.add(ParagraphStyle(name="AquaSmall", parent=styles["Normal"], fontName="Helvetica", fontSize=8, leading=10, textColor=cinza))
        styles.add(ParagraphStyle(name="AquaCenter", parent=styles["Normal"], fontName="Helvetica", fontSize=8.5, leading=11, alignment=TA_CENTER, textColor=colors.HexColor("#1F2937")))

        def esc(txt: str) -> str:
            return html.escape(str(txt or "")).replace("\n", "<br/>")

        try:
            logo_path = encontrar_logo()
        except Exception:
            logo_path = None

        def header_footer(canvas, doc):
            canvas.saveState()
            w, h = A4
            canvas.setFillColor(azul_escuro)
            canvas.rect(0, h - 1.15*cm, w, 1.15*cm, stroke=0, fill=1)
            canvas.setFillColor(dourado)
            canvas.rect(0, h - 1.20*cm, w, 0.05*cm, stroke=0, fill=1)
            canvas.setFont("Helvetica-Bold", 8.5)
            canvas.setFillColor(colors.white)
            canvas.drawString(1.55*cm, h - 0.72*cm, "AQUA GESTÃO – CONTROLE TÉCNICO DE PISCINAS")
            canvas.setFont("Helvetica", 7.5)
            canvas.drawRightString(w - 1.55*cm, h - 0.72*cm, f"Termo Aditivo RT • Página {doc.page}")
            canvas.setStrokeColor(borda)
            canvas.line(1.55*cm, 1.3*cm, w - 1.55*cm, 1.3*cm)
            canvas.setFont("Helvetica", 7.2)
            canvas.setFillColor(cinza)
            canvas.drawString(1.55*cm, 0.92*cm, "Thyago Fernando da Silveira | CRQ-MG 2ª Região | CRQ 024025748 | Técnico em Química")
            canvas.drawRightString(w - 1.55*cm, 0.92*cm, "Documento técnico-comercial vinculado ao contrato base")
            canvas.restoreState()

        doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, topMargin=2.05*cm, bottomMargin=1.75*cm, leftMargin=1.65*cm, rightMargin=1.65*cm, title=f"Aditivo RT - {nome_cond}", author="Aqua Gestão – Controle Técnico de Piscinas")
        story = []

        cabecalho = []
        if logo_path and Path(logo_path).exists():
            try:
                img = RLImage(str(logo_path), width=2.55*cm, height=2.55*cm, kind="proportional")
                cabecalho.append(img)
            except Exception:
                cabecalho.append(Paragraph("<b>AQUA<br/>GESTÃO</b>", styles["AquaCenter"]))
        else:
            cabecalho.append(Paragraph("<b>AQUA<br/>GESTÃO</b>", styles["AquaCenter"]))

        cabecalho.append(Paragraph("<b>1º TERMO ADITIVO AO CONTRATO DE PRESTAÇÃO DE SERVIÇOS</b><br/>RESPONSABILIDADE TÉCNICA – CONTROLE TÉCNICO DE PISCINAS<br/><font color='#4B5563'>Aditivo comercial para concessão de desconto, sem alteração do escopo técnico contratado.</font>", styles["AquaTitle"]))
        tabela_cab = Table([cabecalho], colWidths=[3.0*cm, 14.0*cm])
        tabela_cab.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), azul_claro), ("BOX", (0, 0), (-1, -1), 0.8, borda), ("LINEBELOW", (0, 0), (-1, 0), 1.2, dourado), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9)]))
        story.append(tabela_cab)
        story.append(Spacer(1, 0.35*cm))

        qual = [
            [Paragraph("<b>CONTRATADA</b>", styles["AquaSmall"]), Paragraph("AQUA GESTÃO CONTROLE TÉCNICO LTDA<br/>CNPJ 66.008.795/0001-92", styles["AquaSmall"])],
            [Paragraph("<b>CONTRATANTE</b>", styles["AquaSmall"]), Paragraph(f"{esc(nome_cond)}<br/>CNPJ/CPF: {esc(cnpj_cond)}<br/>{esc(endereco_cond)}", styles["AquaSmall"])],
            [Paragraph("<b>CONTRATO BASE</b>", styles["AquaSmall"]), Paragraph(f"Contrato de Prestação de Serviços de Responsabilidade Técnica firmado em {esc(data_inicio)}" + (f", com vigência até {esc(data_fim)}." if data_fim else "."), styles["AquaSmall"])],
        ]
        tq = Table(qual, colWidths=[3.15*cm, 13.85*cm])
        tq.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.6, borda), ("INNERGRID", (0, 0), (-1, -1), 0.4, borda), ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F8FF")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
        story.append(tq)
        story.append(Spacer(1, 0.32*cm))

        def clausula(titulo, paragrafos):
            story.append(Paragraph(titulo, styles["AquaH"]))
            story.append(HRFlowable(width="100%", thickness=0.7, color=dourado, spaceBefore=1, spaceAfter=5))
            for par in paragrafos:
                story.append(Paragraph(par, styles["AquaBody"]))

        clausula("CLÁUSULA PRIMEIRA — DO DESCONTO COMERCIAL", [
            "A CONTRATADA concede à CONTRATANTE, por mera liberalidade e em caráter comercial específico, desconto especial sobre os honorários mensais previstos no contrato base.",
            f"Em razão do desconto ora concedido, o valor mensal contratual passa de <b>{esc(valor_mensal)}</b> para <b>{esc(valor_aditivo)}</b>.",
            "O desconto previsto neste aditivo possui natureza excepcional, discricionária, revogável e não vinculante, não constituindo direito adquirido, permanência obrigatória, política geral de preços, cláusula de exclusividade ou condição compulsória de contratação de quaisquer outros produtos ou serviços.",
            "O desconto decorre exclusivamente de avaliação comercial global realizada pela CONTRATADA no contexto da contratação, podendo ser revisto, reduzido ou suprimido, total ou parcialmente, mediante notificação prévia por escrito de 30 (trinta) dias, caso deixem de existir as premissas negociais que motivaram sua concessão.",
            "A supressão do desconto não implicará alteração do escopo técnico contratado, passando a vigorar, após o prazo de notificação, o valor contratual ordinário então vigente, já considerado eventual reajuste aplicável.",
        ])
        clausula("CLÁUSULA SEGUNDA — DO REAJUSTE", [
            "O valor com desconto será reajustado anualmente pelo mesmo critério previsto no contrato base.",
            "Na hipótese de cessação do desconto, o valor ordinário restabelecido também observará o mesmo critério de reajuste contratual.",
        ])
        clausula("CLÁUSULA TERCEIRA — DA VIGÊNCIA", [
            "O presente Termo Aditivo entra em vigor na data de sua assinatura e produzirá efeitos financeiros a partir do próximo vencimento contratual.",
            "Permanecem inalteradas e ratificadas todas as demais cláusulas do contrato base não expressamente modificadas por este aditivo.",
        ])

        story.append(Spacer(1, 0.45*cm))
        story.append(Paragraph(esc(local_data), styles["AquaCenter"]))
        story.append(Spacer(1, 0.65*cm))

        assinaturas = [[
            Paragraph("_________________________________________<br/><b>CONTRATADA</b><br/>AQUA GESTÃO CONTROLE TÉCNICO LTDA<br/>Thyago Fernando da Silveira<br/>CRQ 024025748 | Técnico em Química", styles["AquaCenter"]),
            Paragraph(f"_________________________________________<br/><b>CONTRATANTE</b><br/>{esc(nome_cond)}<br/>{esc(sindico)}" + (f" — {esc(cargo_sindico)}" if cargo_sindico else "") + (f"<br/>CPF: {esc(cpf_sindico)}" if cpf_sindico else ""), styles["AquaCenter"]),
        ]]
        ta = Table(assinaturas, colWidths=[8.2*cm, 8.2*cm])
        ta.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5)]))
        story.append(KeepTogether(ta))

        doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
        return True, None if pdf_path.exists() else "PDF não foi criado."
    except Exception as e:
        return False, str(e)

# =========================================
# EXPORTAÇÃO DE CADASTRO
# =========================================

def gerar_html_resumo_cadastro(item: dict) -> str:
    dados = item["dados"] or {}
    status = item["status"]
    nome = item["nome_exibicao"]

    def val(chave):
        return dados.get(chave, "Não informado") or "Não informado"

    html = f"""
    <html>
    <head>
        <meta charset="utf-8" />
        <title>Resumo de Cadastro - {nome}</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                margin: 28px;
                color: #15395f;
            }}
            h1 {{
                font-size: 24px;
                color: #0d3d75;
                margin-bottom: 6px;
            }}
            h2 {{
                font-size: 16px;
                color: #25598d;
                margin-top: 24px;
                margin-bottom: 10px;
            }}
            .sub {{
                color: #5e7691;
                margin-bottom: 20px;
            }}
            .box {{
                border: 1px solid #d7e6f7;
                border-radius: 12px;
                padding: 14px 16px;
                margin-bottom: 8px;
                background: #fbfdff;
            }}
            .line {{
                margin-bottom: 6px;
            }}
            .label {{
                font-weight: bold;
            }}
        </style>
    </head>
    <body>
        <h1>{nome}</h1>
        <div class="sub">Resumo de cadastro exportado pelo sistema Aqua Gestão</div>

        <div class="box">
            <div class="line"><span class="label">Status:</span> {status['rotulo']}</div>
            <div class="line"><span class="label">Mensagem:</span> {status['mensagem']}</div>
            <div class="line"><span class="label">Data final:</span> {item.get('data_fim') or 'Não informada'}</div>
        </div>

        <h2>Dados do condomínio</h2>
        <div class="box">
            <div class="line"><span class="label">Nome do condomínio:</span> {val('nome_condominio')}</div>
            <div class="line"><span class="label">CNPJ:</span> {val('cnpj_condominio')}</div>
            <div class="line"><span class="label">Endereço:</span> {val('endereco_condominio')}</div>
            <div class="line"><span class="label">Síndico / representante:</span> {val('nome_sindico')}</div>
            <div class="line"><span class="label">CPF:</span> {val('cpf_sindico')}</div>
        </div>

        <h2>Dados contratuais</h2>
        <div class="box">
            <div class="line"><span class="label">Valor mensal:</span> {val('valor_mensal')}</div>
            <div class="line"><span class="label">Valor aditivo:</span> {val('valor_aditivo')}</div>
            <div class="line"><span class="label">Data de início:</span> {val('data_inicio')}</div>
            <div class="line"><span class="label">Data de fim:</span> {val('data_fim')}</div>
            <div class="line"><span class="label">Data de assinatura:</span> {val('data_assinatura')}</div>
        </div>

        <h2>Contato</h2>
        <div class="box">
            <div class="line"><span class="label">WhatsApp:</span> {val('whatsapp_cliente')}</div>
            <div class="line"><span class="label">E-mail:</span> {val('email_cliente')}</div>
            <div class="line"><span class="label">Última atualização:</span> {val('salvo_em')}</div>
        </div>
    </body>
    </html>
    """
    return html


def _mockup_dados_relatorio_demo() -> dict:
    """Dados fictícios usados apenas para a pré-visualização do modelo de relatório."""
    lancamento_aqua = {
        "data": "10/04/2026",
        "operador": "João Silva",
        "observacao": "Casa de máquinas organizada, filtro em operação normal e responsável local orientado sobre nova conferência após a recirculação.",
        "problemas": "Piscina infantil com cloro livre abaixo do mínimo operacional e pH levemente abaixo da faixa ideal.",
        "parecer": "Aceitável com correções imediatas na piscina infantil e monitoramento da estabilidade química nas próximas horas.",
        "dosagens": [
            {"produto": "Cloro granulado", "quantidade": "1,2", "unidade": "kg", "finalidade": "Reforço de desinfecção"},
            {"produto": "Barrilha leve", "quantidade": "0,8", "unidade": "kg", "finalidade": "Correção gradual do pH"},
            {"produto": "Algicida manutenção", "quantidade": "250", "unidade": "mL", "finalidade": "Prevenção"},
        ],
        "piscinas": [
            {
                "nome": "Piscina Adulto",
                "ph": "7,4",
                "cloro_livre": "1,8",
                "cloro_total": "2,0",
                "cloraminas": "0,2",
                "alcalinidade": "110",
                "dureza": "180",
                "cianurico": "35",
            },
            {
                "nome": "Piscina Infantil",
                "ph": "7,0",
                "cloro_livre": "0,4",
                "cloro_total": "0,8",
                "cloraminas": "0,4",
                "alcalinidade": "72",
                "dureza": "230",
                "cianurico": "28",
            },
        ],
    }

    lancamentos_periodo = [
        {
            "data": "02/04/2026",
            "operador": "João Silva",
            "observacao": "Início do período com parâmetros estáveis e água visualmente cristalina.",
            "dosagens": [{"produto": "Clarificante", "quantidade": "150", "unidade": "mL", "finalidade": "Auxílio à filtração"}],
            "piscinas": [
                {"nome": "Piscina Adulto", "ph": "7,5", "cloro_livre": "2,0", "cloro_total": "2,2", "cloraminas": "0,2", "alcalinidade": "105", "dureza": "190", "cianurico": "34"},
                {"nome": "Piscina Infantil", "ph": "7,3", "cloro_livre": "1,6", "cloro_total": "1,8", "cloraminas": "0,2", "alcalinidade": "88", "dureza": "220", "cianurico": "30"},
            ],
            "parecer": "Satisfatório.",
        },
        {
            "data": "10/04/2026",
            "operador": "João Silva",
            "observacao": "Casa de máquinas organizada, filtro em operação normal e responsável local orientado sobre nova conferência após a recirculação.",
            "problemas": "Piscina infantil com cloro livre abaixo do mínimo operacional e pH levemente abaixo da faixa ideal.",
            "dosagens": [
                {"produto": "Cloro granulado", "quantidade": "1,2", "unidade": "kg", "finalidade": "Reforço de desinfecção"},
                {"produto": "Barrilha leve", "quantidade": "0,8", "unidade": "kg", "finalidade": "Correção gradual do pH"},
            ],
            "piscinas": [
                {"nome": "Piscina Adulto", "ph": "7,4", "cloro_livre": "1,8", "cloro_total": "2,0", "cloraminas": "0,2", "alcalinidade": "110", "dureza": "180", "cianurico": "35", "dosagens": [{"produto": "Algicida manutenção", "quantidade": "250", "unidade": "mL", "finalidade": "Prevenção"}]},
                {"nome": "Piscina Infantil", "ph": "7,0", "cloro_livre": "0,4", "cloro_total": "0,8", "cloraminas": "0,4", "alcalinidade": "72", "dureza": "230", "cianurico": "28", "dosagens": [{"produto": "Cloro granulado", "quantidade": "1,2", "unidade": "kg", "finalidade": "Reforço de desinfecção"}, {"produto": "Barrilha leve", "quantidade": "0,8", "unidade": "kg", "finalidade": "Correção gradual do pH"}]},
            ],
            "parecer": "Aceitável com correções imediatas na piscina infantil e monitoramento da estabilidade química nas próximas horas.",
        },
        {
            "data": "18/04/2026",
            "operador": "Marcos Lima",
            "observacao": "Rechecagem após tratamento corretivo com melhora dos parâmetros e estabilidade visual da água.",
            "dosagens": [{"produto": "Cloro granulado", "quantidade": "0,4", "unidade": "kg", "finalidade": "Manutenção"}],
            "piscinas": [
                {"nome": "Piscina Adulto", "ph": "7,5", "cloro_livre": "2,1", "cloro_total": "2,3", "cloraminas": "0,2", "alcalinidade": "108", "dureza": "185", "cianurico": "36"},
                {"nome": "Piscina Infantil", "ph": "7,3", "cloro_livre": "1,5", "cloro_total": "1,7", "cloraminas": "0,2", "alcalinidade": "84", "dureza": "228", "cianurico": "29"},
            ],
            "parecer": "Satisfatório.",
        },
    ]

    return {
        "nome_local": "Residencial Águas Claras",
        "cnpj": "12.345.678/0001-90",
        "endereco": "Av. Floriano Peixoto, 1500 - Uberlândia/MG",
        "responsavel": "Carlos Menezes",
        "operador": "João Silva",
        "mes": "04",
        "ano": "2026",
        "lancamento_aqua": lancamento_aqua,
        "lancamentos_periodo": lancamentos_periodo,
        "obs_geral": "Relatório demonstrativo de pré-visualização, espelhando o padrão visual e textual atual do sistema.",
        "fotos_demo": [
            ("Antes do tratamento", "Piscina infantil antes da correção química"),
            ("Após o tratamento", "Área após aplicação dos produtos e recirculação inicial"),
            ("Casa de máquinas", "Conjunto de filtração e circulação em operação"),
        ],
    }


def _coletar_analises_preview_formulario() -> list[dict]:
    """Coleta as linhas atuais do formulário de relatório mensal para a prévia."""
    total = int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO)
    linhas = []
    for i in range(total):
        item = {
            "data": str(st.session_state.get(f"rel_analise_data_{i}", "") or "").strip(),
            "ph": str(st.session_state.get(f"rel_analise_ph_{i}", "") or "").strip(),
            "cloro_livre": str(st.session_state.get(f"rel_analise_cl_{i}", "") or "").strip(),
            "cloro_total": str(st.session_state.get(f"rel_analise_ct_{i}", "") or "").strip(),
            "alcalinidade": str(st.session_state.get(f"rel_analise_alc_{i}", "") or "").strip(),
            "dureza": str(st.session_state.get(f"rel_analise_dc_{i}", "") or "").strip(),
            "cianurico": str(st.session_state.get(f"rel_analise_cya_{i}", "") or "").strip(),
            "operador": str(st.session_state.get(f"rel_analise_operador_{i}", "") or "").strip(),
        }
        if any(item.values()):
            item["cloraminas"] = ""
            try:
                _ct = valor_float(item["cloro_total"])
                _cl = valor_float(item["cloro_livre"])
                if _ct is not None and _cl is not None:
                    item["cloraminas"] = str(round(max(_ct - _cl, 0), 2)).replace(".", ",")
            except Exception:
                item["cloraminas"] = ""
            linhas.append(item)
    return linhas


def _coletar_dosagens_preview_formulario() -> list[dict]:
    """Coleta as dosagens atuais do formulário de relatório mensal para a prévia."""
    dosagens = []
    for i in range(7):
        item = {
            "produto": str(st.session_state.get(f"rel_dos_produto_{i}", "") or "").strip(),
            "fabricante_lote": str(st.session_state.get(f"rel_dos_lote_{i}", "") or "").strip(),
            "quantidade": str(st.session_state.get(f"rel_dos_qtd_{i}", "") or "").strip(),
            "unidade": str(st.session_state.get(f"rel_dos_un_{i}", "") or "").strip(),
            "finalidade": str(st.session_state.get(f"rel_dos_finalidade_{i}", "") or "").strip(),
        }
        if any(item.values()):
            dosagens.append(item)
    return dosagens


def _obter_obs_preview_formulario() -> str:
    obs = []
    for i in range(5):
        txt = str(st.session_state.get(f"rel_obs_{i}", "") or "").strip()
        if txt:
            obs.append(txt)
    obs_geral = str(st.session_state.get("rel_observacoes_gerais", "") or "").strip()
    if obs_geral:
        obs.insert(0, obs_geral)
    if obs:
        return "\n".join(obs)
    return str(st.session_state.get("csr_obs_rel", "") or "").strip()


def _parecer_preview_por_status(status: str, diagnostico: str) -> str:
    diagnostico = str(diagnostico or "").strip()
    status = str(status or "").strip().upper()
    if diagnostico:
        return diagnostico
    if status == "CONFORME":
        return "Satisfatório. Parâmetros em conformidade com o preenchimento atual do formulário."
    if status == "EM CORREÇÃO":
        return "Aceitável com correções em andamento conforme o preenchimento atual do formulário."
    if status == "NÃO CONFORME":
        return "Não satisfatório. O preenchimento atual indica necessidade de intervenção corretiva."
    return "Parecer técnico-operacional não informado no formulário atual."


def _obter_dados_preview_relatorio(empresa: str = "Aqua Gestão", usar_formulario_atual: bool = False) -> tuple[dict, bool, str]:
    dados = _mockup_dados_relatorio_demo()
    empresa = str(empresa or "Aqua Gestão").strip()
    if not usar_formulario_atual:
        return dados, False, "Exibindo dados demonstrativos do modelo."

    nome = ""
    cnpj = ""
    endereco = ""
    responsavel = ""
    operador = ""
    mes = ""
    ano = ""

    if empresa == "Bem Star Piscinas":
        nome = str(st.session_state.get("csr_sel_relatorio", "") or "").strip()
        mes = str(st.session_state.get("csr_mes_rel", "") or "").strip()
        ano = str(st.session_state.get("csr_ano_rel", "") or "").strip()
        operador = str(st.session_state.get("csr_operador_rel", "") or "").strip()
        responsavel = str(st.session_state.get("rel_representante", "") or "").strip()
        cnpj = str(st.session_state.get("rel_cnpj_condominio", "") or "").strip()
        endereco = str(st.session_state.get("rel_endereco_condominio", "") or "").strip()
    else:
        nome = str(st.session_state.get("rel_nome_condominio", "") or "").strip()
        cnpj = str(st.session_state.get("rel_cnpj_condominio", "") or "").strip()
        endereco = str(st.session_state.get("rel_endereco_condominio", "") or "").strip()
        responsavel = str(st.session_state.get("rel_representante", "") or "").strip()
        mes = str(st.session_state.get("rel_mes_referencia", "") or "").strip()
        ano = str(st.session_state.get("rel_ano_referencia", "") or "").strip()

    analises = _coletar_analises_preview_formulario()
    dosagens = _coletar_dosagens_preview_formulario()
    obs_txt = _obter_obs_preview_formulario()
    diagnostico = str(st.session_state.get("rel_diagnostico", "") or "").strip()
    status_agua = str(st.session_state.get("rel_status_agua", "") or "").strip()
    parecer = _parecer_preview_por_status(status_agua, diagnostico)

    if analises and not operador:
        operador = analises[-1].get("operador", "")

    campos_preenchidos = any([nome, cnpj, endereco, responsavel, operador, mes, ano, obs_txt, diagnostico, analises, dosagens])
    if not campos_preenchidos:
        return dados, False, "Nenhum dado atual do formulário foi encontrado; exibindo dados demonstrativos."

    dados["nome_local"] = nome or dados["nome_local"]
    dados["cnpj"] = cnpj or dados["cnpj"]
    dados["endereco"] = endereco or dados["endereco"]
    dados["responsavel"] = responsavel or dados["responsavel"]
    dados["operador"] = operador or dados["operador"]
    dados["mes"] = mes or dados["mes"]
    dados["ano"] = ano or dados["ano"]
    dados["obs_geral"] = obs_txt or dados["obs_geral"]

    if analises:
        lancamentos_periodo = []
        for idx, item in enumerate(analises):
            _obs_item = item["data"] and obs_txt if idx == len(analises) - 1 else ""
            _problemas = ""
            if idx == len(analises) - 1 and status_agua == "NÃO CONFORME":
                _problemas = "Status geral da água marcado como NÃO CONFORME no formulário atual."
            elif idx == len(analises) - 1 and status_agua == "EM CORREÇÃO":
                _problemas = "Status geral da água marcado como EM CORREÇÃO no formulário atual."

            lancamentos_periodo.append({
                "data": item.get("data", ""),
                "operador": item.get("operador", operador),
                "observacao": _obs_item,
                "problemas": _problemas,
                "parecer": parecer if idx == len(analises) - 1 else "",
                "dosagens": dosagens if idx == len(analises) - 1 else [],
                "piscinas": [{
                    "nome": "Piscina",
                    "ph": item.get("ph", ""),
                    "cloro_livre": item.get("cloro_livre", ""),
                    "cloro_total": item.get("cloro_total", ""),
                    "cloraminas": item.get("cloraminas", ""),
                    "alcalinidade": item.get("alcalinidade", ""),
                    "dureza": item.get("dureza", ""),
                    "cianurico": item.get("cianurico", ""),
                }],
            })

        dados["lancamentos_periodo"] = lancamentos_periodo
        ultimo = analises[-1]
        dados["lancamento_aqua"] = {
            "data": ultimo.get("data", ""),
            "operador": ultimo.get("operador", operador),
            "observacao": obs_txt or dados["obs_geral"],
            "problemas": lancamentos_periodo[-1].get("problemas", ""),
            "parecer": parecer,
            "dosagens": dosagens,
            "piscinas": [{
                "nome": "Piscina",
                "ph": ultimo.get("ph", ""),
                "cloro_livre": ultimo.get("cloro_livre", ""),
                "cloro_total": ultimo.get("cloro_total", ""),
                "cloraminas": ultimo.get("cloraminas", ""),
                "alcalinidade": ultimo.get("alcalinidade", ""),
                "dureza": ultimo.get("dureza", ""),
                "cianurico": ultimo.get("cianurico", ""),
            }],
        }
    else:
        dados["lancamento_aqua"]["observacao"] = obs_txt or dados["lancamento_aqua"].get("observacao", "")
        dados["lancamento_aqua"]["operador"] = operador or dados["lancamento_aqua"].get("operador", "")
        dados["lancamento_aqua"]["dosagens"] = dosagens or dados["lancamento_aqua"].get("dosagens", [])
        dados["lancamento_aqua"]["parecer"] = parecer

    return dados, True, "Usando os dados atuais preenchidos no formulário para montar esta prévia."


def _gerar_mockup_relatorio_bem_star_html(dados: dict | None = None) -> str:
    dados = dados or _mockup_dados_relatorio_demo()
    hoje = date.today().strftime("%d/%m/%Y")
    return f"""<!DOCTYPE html>
<html lang=\"pt-BR\">
<head>
<meta charset=\"UTF-8\">
<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">
<title>Relatório técnico-operacional de piscinas — {dados['nome_local']}</title>
<style>
  *{{box-sizing:border-box;margin:0;padding:0;}}
  body{{font-family:Arial,Helvetica,sans-serif;background:#f4f7f9;color:#14313d;}}
  .page{{max-width:680px;margin:0 auto;padding:16px;}}
  .card{{background:#fff;border:1px solid #d5e2e6;border-radius:12px;padding:18px 20px;margin-bottom:12px;}}
  .hdr{{display:flex;justify-content:space-between;gap:12px;align-items:flex-start;}}
  .brand{{display:flex;gap:12px;align-items:center;}}
  .logo{{width:48px;height:48px;border-radius:50%;background:#15707c;color:#fff;display:flex;align-items:center;justify-content:center;font-weight:700;}}
  .tit1{{font-size:16px;font-weight:700;color:#0f3a46;}}
  .tit2{{font-size:10px;color:#6f8790;letter-spacing:.7px;text-transform:uppercase;margin-top:2px;}}
  .doc{{text-align:right;}}
  .doc .ttl{{font-size:13px;font-weight:700;color:#15707c;}}
  .doc .sub{{font-size:10px;color:#7f95a0;margin-top:3px;}}
  .sec{{font-size:10px;font-weight:700;color:#15707c;text-transform:uppercase;letter-spacing:.8px;margin-bottom:10px;padding-bottom:6px;border-bottom:2px solid #15707c;}}
  .grid{{display:grid;grid-template-columns:1fr 1fr;gap:8px 16px;}}
  .lbl{{font-size:10px;color:#7f95a0;text-transform:uppercase;}}
  .val{{font-size:13px;font-weight:600;color:#14313d;margin-top:2px;}}
  .pool{{border:1px solid #dbe5ea;border-radius:10px;padding:12px;margin-top:10px;}}
  .pool h4{{margin:0 0 10px;font-size:13px;color:#0f3a46;}}
  .params{{display:grid;grid-template-columns:repeat(3,1fr);gap:8px;}}
  .box{{border:1px solid #d7e2e8;border-radius:8px;padding:10px 8px;text-align:center;background:#fbfdfe;}}
  .box.ok{{background:#edf7f0;border-color:#cfe6d4;}}
  .box.warn{{background:#fff6ee;border-color:#f1d4be;}}
  .pn{{font-size:9px;color:#7f95a0;text-transform:uppercase;margin-bottom:4px;}}
  .pv{{font-size:18px;font-weight:700;color:#14313d;}}
  .box.ok .pv{{color:#2e7d32;}}
  .box.warn .pv{{color:#c26300;}}
  .ps{{font-size:9px;color:#6f8790;margin-top:3px;}}
  .list{{padding-left:18px;color:#334e58;line-height:1.7;font-size:13px;}}
  .note{{font-size:12px;color:#516a74;line-height:1.7;}}
  .photo{{border:1px dashed #c9d8de;border-radius:10px;padding:12px;background:#f8fbfc;margin-bottom:8px;}}
  .photo strong{{display:block;color:#0f3a46;font-size:12px;margin-bottom:4px;}}
  .rod{{text-align:center;font-size:10px;color:#7f95a0;padding:8px 0 2px;}}
</style>
</head>
<body>
<div class=\"page\">
  <div class=\"card\">
    <div class=\"hdr\">
      <div class=\"brand\">
        <div class=\"logo\">BS</div>
        <div>
          <div class=\"tit1\">BEM STAR PISCINAS</div>
          <div class=\"tit2\">RELATÓRIO TÉCNICO-OPERACIONAL DE PISCINAS</div>
        </div>
      </div>
      <div class=\"doc\">
        <div class=\"ttl\">Relatório de visita</div>
        <div class=\"sub\">Emitido em {hoje}</div>
      </div>
    </div>
    <div style=\"margin-top:12px;height:1px;background:#d7e3e7;\"></div>
    <div class=\"grid\" style=\"margin-top:12px;\">
      <div><div class=\"lbl\">Local / Condomínio</div><div class=\"val\">{dados['nome_local']}</div></div>
      <div><div class=\"lbl\">Período de referência</div><div class=\"val\">{dados['mes']}/{dados['ano']}</div></div>
      <div><div class=\"lbl\">Operador de campo</div><div class=\"val\">{dados['operador']}</div></div>
      <div><div class=\"lbl\">Responsável / Síndico</div><div class=\"val\">{dados['responsavel']}</div></div>
    </div>
  </div>

  <div class=\"card\">
    <div class=\"sec\">1. Identificação do local</div>
    <div class=\"grid\">
      <div><div class=\"lbl\">Local / Condomínio</div><div class=\"val\">{dados['nome_local']}</div></div>
      <div><div class=\"lbl\">CNPJ</div><div class=\"val\">{dados['cnpj']}</div></div>
      <div><div class=\"lbl\">Endereço</div><div class=\"val\">{dados['endereco']}</div></div>
      <div><div class=\"lbl\">Responsável no local</div><div class=\"val\">{dados['responsavel']}</div></div>
    </div>
  </div>

  <div class=\"card\">
    <div class=\"sec\">2. Análises físico-químicas</div>
    <div class=\"pool\">
      <h4>🏊 Piscina Principal</h4>
      <div class=\"params\">
        <div class=\"box ok\"><div class=\"pn\">pH</div><div class=\"pv\">7,5</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box ok\"><div class=\"pn\">CRL mg/L</div><div class=\"pv\">2,2</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box ok\"><div class=\"pn\">CT mg/L</div><div class=\"pv\">2,4</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box ok\"><div class=\"pn\">Clor. mg/L</div><div class=\"pv\">0,2</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box ok\"><div class=\"pn\">Alc. mg/L</div><div class=\"pv\">98</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box ok\"><div class=\"pn\">CYA mg/L</div><div class=\"pv\">32</div><div class=\"ps\">Conforme</div></div>
      </div>
    </div>
    <div class=\"pool\">
      <h4>🏊 Spa / Hidro</h4>
      <div class=\"params\">
        <div class=\"box ok\"><div class=\"pn\">pH</div><div class=\"pv\">7,3</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box warn\"><div class=\"pn\">CRL mg/L</div><div class=\"pv\">0,9</div><div class=\"ps\">Em atenção</div></div>
        <div class=\"box ok\"><div class=\"pn\">CT mg/L</div><div class=\"pv\">1,2</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box ok\"><div class=\"pn\">Clor. mg/L</div><div class=\"pv\">0,3</div><div class=\"ps\">Monitorar</div></div>
        <div class=\"box ok\"><div class=\"pn\">Alc. mg/L</div><div class=\"pv\">90</div><div class=\"ps\">Conforme</div></div>
        <div class=\"box ok\"><div class=\"pn\">CYA mg/L</div><div class=\"pv\">25</div><div class=\"ps\">Conforme</div></div>
      </div>
    </div>
  </div>

  <div class=\"card\">
    <div class=\"sec\">3. Dosagens aplicadas</div>
    <ul class=\"list\">
      <li>Aplicação de 600 g de cloro granulado no spa/hidro.</li>
      <li>Aplicação de 150 mL de clarificante na piscina principal.</li>
      <li>Verificação de cestos, skimmers e pré-filtros.</li>
      <li>Orientação passada ao responsável do local.</li>
    </ul>
  </div>

  <div class=\"card\">
    <div class=\"sec\">4. Observações gerais</div>
    <div class=\"note\">Condição geral satisfatória. Apenas o spa/hidro requer rechecagem breve de cloro residual devido à maior sensibilidade de volume e uso. Estrutura e circulação em conformidade visual no momento da visita.</div>
  </div>

  <div class=\"card\">
    <div class=\"sec\">5. Registro fotográfico</div>
    <div class=\"photo\"><strong>📷 Piscina principal</strong>Vista geral da lâmina d'água.</div>
    <div class=\"photo\"><strong>📷 Spa / Hidro</strong>Condição visual do segundo corpo d'água.</div>
    <div class=\"photo\"><strong>📷 Equipamentos</strong>Painel e equipamentos no momento da visita.</div>
  </div>

  <div class=\"card\">
    <div class=\"sec\">Sobre responsabilidade técnica (RT)</div>
    <div class=\"note\">{TEXTO_RT_SEM_RT.strip().replace(chr(10), '<br><br>')}</div>
  </div>

  <div class=\"rod\">Bem Star Piscinas · Documento de pré-visualização do modelo atual</div>
</div>
</body>
</html>"""


def _gerar_mockup_relatorio_impressao_html(empresa: str = "Aqua Gestão", dados: dict | None = None) -> str:
    dados = dados or _mockup_dados_relatorio_demo()
    incluir_rt = str(empresa or "").strip() != "Bem Star Piscinas"
    titulo_topo = "AQUA GESTÃO – CONTROLE TÉCNICO DE PISCINAS" if incluir_rt else "BEM STAR PISCINAS"
    subtitulo_topo = (
        f"Responsável Técnico: {RESPONSAVEL_TÉCNICO} | {CRQ}<br>{QUALIFICACAO_RT} | Certificações: {CERTIFICACOES_RT}"
        if incluir_rt else
        f"RELATÓRIO TÉCNICO-OPERACIONAL DE PISCINAS<br>CNPJ: {CNPJ_BEM_STAR}  |  Uberlândia/MG"
    )
    blocos_analises = []
    for nome_piscina in ["Piscina Adulto", "Piscina Infantil"]:
        linhas = []
        for lc in dados['lancamentos_periodo']:
            for p in lc.get('piscinas', []):
                if p.get('nome') == nome_piscina:
                    linhas.append(f"""
                    <tr>
                      <td>{lc.get('data','')}</td>
                      <td>{p.get('ph','')}</td>
                      <td>{p.get('cloro_livre','')}</td>
                      <td>{p.get('cloro_total','')}</td>
                      <td>{p.get('cloraminas','')}</td>
                      <td>{p.get('alcalinidade','')}</td>
                      <td>{p.get('dureza','')}</td>
                      <td>{p.get('cianurico','')}</td>
                      <td>{lc.get('operador','')}</td>
                    </tr>
                    """)
        blocos_analises.append(f"""
        <div class=\"subpiscina\">🏊 {nome_piscina}</div>
        <table>
          <thead>
            <tr>
              <th>Data</th><th>pH</th><th>CRL mg/L</th><th>CT mg/L</th><th>Clor. mg/L</th><th>Alc. mg/L</th><th>Dureza mg/L</th><th>CYA mg/L</th><th>Operador</th>
            </tr>
          </thead>
          <tbody>
            {''.join(linhas)}
          </tbody>
        </table>
        """)

    secao_rt_extra = ""
    if not incluir_rt:
        secao_rt_extra = f"""
        <div class=\"sec\">SOBRE RESPONSABILIDADE TÉCNICA (RT)</div>
        <div class=\"texto\">{TEXTO_RT_SEM_RT.strip().replace(chr(10), '<br><br>')}</div>
        """

    secao_assinatura = (
        f"___________________________<br>{RESPONSAVEL_TÉCNICO}<br>{CRQ}<br>{QUALIFICACAO_RT}"
        if incluir_rt else
        "___________________________<br>Bem Star Piscinas"
    )

    return f"""<!DOCTYPE html>
<html lang=\"pt-BR\">
<head>
<meta charset=\"UTF-8\">
<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">
<title>Pré-visualização de impressão — {dados['nome_local']}</title>
<style>
  *{{box-sizing:border-box;}}
  body{{margin:0;background:#e9eef5;font-family:Arial,Helvetica,sans-serif;color:#172b44;padding:24px;}}
  .sheet{{max-width:960px;margin:0 auto;background:#fff;box-shadow:0 10px 40px rgba(20,35,60,.14);padding:38px 42px;border:1px solid #d7e0eb;}}
  .topo{{text-align:center;margin-bottom:18px;}}
  .topo h1{{font-size:18px;margin:0;color:#0d3d75;}}
  .topo .sub{{font-size:11px;color:#4d647d;line-height:1.6;margin-top:6px;}}
  .linha{{height:2px;background:#0d3d75;margin:14px 0 18px;}}
  .sec{{font-size:14px;font-weight:700;margin:18px 0 10px;color:#0d3d75;}}
  table{{width:100%;border-collapse:collapse;margin:8px 0 14px;table-layout:fixed;}}
  th,td{{border:1px solid #c8d3e0;padding:7px 6px;font-size:11px;vertical-align:top;word-wrap:break-word;}}
  th{{background:#0d3d75;color:#fff;font-weight:700;}}
  tbody tr:nth-child(odd) td{{background:#eef3fb;}}
  .info td:first-child{{width:26%;background:#eef3fb;font-weight:700;}}
  .subpiscina{{font-size:12px;font-weight:700;color:#0d3d75;margin:12px 0 6px;}}
  .texto{{font-size:11px;line-height:1.8;color:#2f3f52;}}
  .bloco-fotos{{display:grid;grid-template-columns:repeat(3,1fr);gap:10px;margin-top:8px;}}
  .foto{{border:1px solid #cfd9e5;min-height:120px;padding:10px;background:#f9fbfd;}}
  .foto strong{{display:block;margin-bottom:6px;color:#0d3d75;font-size:11px;}}
  .parecer{{padding:10px 12px;border:1px solid #d9c8a8;background:#fff8ea;font-size:11px;line-height:1.7;}}
  .assinatura{{text-align:center;margin-top:30px;font-size:11px;line-height:1.9;}}
  .rodape{{text-align:center;font-size:10px;color:#748399;margin-top:16px;}}
</style>
</head>
<body>
<div class=\"sheet\">
  <div class=\"topo\">
    <h1>{titulo_topo}</h1>
    <div class=\"sub\">{subtitulo_topo}</div>
  </div>
  <div class=\"linha\"></div>

  <div class=\"sec\">1. IDENTIFICAÇÃO DO LOCAL</div>
  <table class=\"info\">
    <tbody>
      <tr><td>Local / Condomínio</td><td>{dados['nome_local']}</td></tr>
      <tr><td>CNPJ</td><td>{dados['cnpj']}</td></tr>
      <tr><td>Endereço</td><td>{dados['endereco']}</td></tr>
      <tr><td>Responsável / Síndico</td><td>{dados['responsavel']}</td></tr>
      <tr><td>Responsável no local</td><td>{dados['responsavel']}</td></tr>
      <tr><td>Operador de campo</td><td>{dados['operador']}</td></tr>
      <tr><td>Período de referência</td><td>{dados['mes']}/{dados['ano']}</td></tr>
    </tbody>
  </table>

  <div class=\"sec\">2. ANÁLISES FÍSICO-QUÍMICAS</div>
  {''.join(blocos_analises)}

  <div class=\"sec\">3. DOSAGENS APLICADAS</div>
  <table>
    <thead><tr><th>Data</th><th>Piscina</th><th>Produto</th><th>Quantidade</th><th>Finalidade técnica</th></tr></thead>
    <tbody>
      <tr><td>10/04/2026</td><td>Piscina Infantil</td><td>Cloro granulado</td><td>1,2 kg</td><td>Reforço de desinfecção</td></tr>
      <tr><td>10/04/2026</td><td>Piscina Infantil</td><td>Barrilha leve</td><td>0,8 kg</td><td>Correção gradual do pH</td></tr>
      <tr><td>10/04/2026</td><td>Piscina Adulto</td><td>Algicida manutenção</td><td>250 mL</td><td>Prevenção</td></tr>
    </tbody>
  </table>

  <div class=\"sec\">4. PROBLEMAS / OCORRÊNCIAS</div>
  <div class=\"texto\">⚠ 10/04/2026: Piscina infantil com cloro livre abaixo do mínimo operacional e pH levemente abaixo da faixa ideal.</div>

  <div class=\"sec\">5. OBSERVAÇÕES GERAIS</div>
  <div class=\"texto\">{dados['obs_geral']}<br>10/04/2026: Casa de máquinas organizada, filtro em operação normal e responsável local orientado sobre nova conferência após a recirculação.</div>

  <div class=\"sec\">6. REGISTRO FOTOGRÁFICO</div>
  <div class=\"bloco-fotos\">
    <div class=\"foto\"><strong>🔵 Antes do tratamento</strong>Piscina infantil antes da correção química.</div>
    <div class=\"foto\"><strong>🟢 Após o tratamento</strong>Área após aplicação dos produtos e recirculação inicial.</div>
    <div class=\"foto\"><strong>🔧 Casa de máquinas</strong>Conjunto de filtração e circulação em operação.</div>
  </div>

  <div class=\"sec\">7. PARECER TÉCNICO-OPERACIONAL</div>
  <div class=\"parecer\">Parecer da última visita: Aceitável com correções imediatas na piscina infantil e monitoramento da estabilidade química nas próximas horas.</div>

  {secao_rt_extra}

  <div class=\"assinatura\">Uberlândia/MG, {hoje_br()}.<br><br>{secao_assinatura}</div>
  <div class=\"rodape\">Documento demonstrativo de pré-visualização do modelo atual do sistema</div>
</div>
</body>
</html>"""


def gerar_mockup_relatorio_preview_html(empresa: str = "Aqua Gestão", visual: str = "web", dados: dict | None = None) -> str:
    """Retorna o HTML demonstrativo do relatório conforme o modelo atual do sistema."""
    empresa = str(empresa or "Aqua Gestão").strip()
    visual = str(visual or "web").strip().lower()
    dados = dados or _mockup_dados_relatorio_demo()
    if visual == "print":
        return _gerar_mockup_relatorio_impressao_html(empresa, dados=dados)
    if empresa == "Bem Star Piscinas":
        return _gerar_mockup_relatorio_bem_star_html(dados=dados)
    return gerar_html_relatorio_visita(dados["lancamento_aqua"], dados["nome_local"])


# =========================================
# MENSAGENS / LINKS
# =========================================

def montar_mensagem_envio(
    nome_condominio: str,
    nome_sindico: str,
    caminho_contrato_pdf: Path | None = None,
    caminho_aditivo_pdf: Path | None = None,
) -> str:
    partes = [
        f"Prezado(a) {nome_sindico},",
        "",
        f"Segue em anexo a documentação referente ao condomínio {nome_condominio}:",
        "",
    ]

    if caminho_contrato_pdf and caminho_contrato_pdf.exists():
        partes.append("- Contrato de Responsabilidade Técnica (PDF)")
    if caminho_aditivo_pdf and caminho_aditivo_pdf.exists():
        partes.append("- Aditivo contratual (PDF)")

    partes += [
        "",
        "Permaneço à disposição para quaisquer esclarecimentos.",
        "",
        "Atenciosamente,",
        f"{RESPONSAVEL_TÉCNICO}",
        CRQ,
        "Aqua Gestão – Controle Técnico de Piscinas",
    ]

    return "\n".join(partes)


def montar_mensagem_bem_star(
    nome_local: str,
    responsavel: str,
    tipo: str = "contrato",          # "contrato" | "relatorio" | "ambos"
    mes: str = "",
    ano: str = "",
) -> str:
    """Monta mensagem de envio com identidade Bem Star Piscinas."""
    saudacao = f"Prezado(a) {responsavel}," if responsavel else "Prezado(a),"
    partes = [saudacao, ""]

    if tipo == "contrato":
        partes += [
            f"Segue em anexo o contrato de prestação de serviços de limpeza e manutenção "
            f"de piscinas referente a {nome_local}.",
            "",
            "Por favor, verifique as condições acordadas e, havendo qualquer dúvida, "
            "estamos à disposição para esclarecimentos.",
        ]
    elif tipo == "relatorio":
        periodo = f" — {mes}/{ano}" if mes and ano else ""
        partes += [
            f"Segue em anexo o relatório técnico-operacional de piscinas{periodo} "
            f"referente a {nome_local}.",
            "",
            "O documento registra os parâmetros analisados, produtos aplicados, "
            "dosagens e observações das visitas realizadas no período.",
        ]
    else:  # ambos
        partes += [
            f"Segue em anexo a documentação referente a {nome_local}:",
            "",
            "— Contrato de prestação de serviços",
            "— Relatório técnico-operacional de piscinas",
        ]

    partes += [
        "",
        "Permanecemos à disposição para quaisquer esclarecimentos.",
        "",
        "Atenciosamente,",
        "Bem Star Piscinas",
        f"CNPJ: {CNPJ_BEM_STAR}",
        "Av. Getúlio Vargas, 4411 — Uberlândia/MG",
        "(34) 9 9999-9999",
    ]
    return "\n".join(partes)


def exibir_bloco_envio_bem_star(
    nome_local: str,
    pasta: Path,
    telefone: str,
    email: str,
    mensagem: str,
    key_suffix: str = "",
):
    """Bloco de envio com identidade Bem Star (WhatsApp + email + copiar)."""
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader("📤 Enviar para o cliente")

    # Editor da mensagem
    msg_editada = st.text_area(
        "Mensagem",
        value=mensagem,
        height=200,
        key=f"bs_msg_envio_{key_suffix}",
        label_visibility="collapsed",
    )
    componente_copiar(msg_editada)

    _ec1, _ec2, _ec3 = st.columns(3)
    with _ec1:
        _tel = (telefone or "").strip()
        if _tel:
            _url_wa = link_whatsapp(_tel, msg_editada)
            st.link_button("💬 Abrir WhatsApp", _url_wa, use_container_width=True)
        else:
            st.text_input("WhatsApp", placeholder="(34) 99999-9999",
                key=f"bs_tel_envio_{key_suffix}")
            _tel2 = st.session_state.get(f"bs_tel_envio_{key_suffix}", "").strip()
            if _tel2:
                st.link_button("💬 Abrir WhatsApp",
                    link_whatsapp(_tel2, msg_editada), use_container_width=True)

    with _ec2:
        _eml = (email or "").strip()
        if _eml:
            _assunto = f"Documentação Bem Star Piscinas – {nome_local}"
            _url_mail = link_email(_eml, _assunto, msg_editada)
            st.link_button("✉️ Abrir e-mail", _url_mail, use_container_width=True)
        else:
            st.text_input("E-mail", placeholder="email@cliente.com.br",
                key=f"bs_email_envio_{key_suffix}")
            _eml2 = st.session_state.get(f"bs_email_envio_{key_suffix}", "").strip()
            if _eml2:
                _assunto = f"Documentação Bem Star Piscinas – {nome_local}"
                st.link_button("✉️ Abrir e-mail",
                    link_email(_eml2, _assunto, msg_editada), use_container_width=True)

    with _ec3:
        if pasta and pasta.exists():
            if st.button("📁 Abrir pasta", key=f"bs_pasta_{key_suffix}",
                    use_container_width=True):
                abrir_pasta_windows(pasta)

    st.markdown("</div>", unsafe_allow_html=True)


def link_whatsapp(telefone: str, mensagem: str) -> str:
    somente_numeros = apenas_digitos(telefone or "")
    if not somente_numeros.startswith("55") and somente_numeros:
        somente_numeros = "55" + somente_numeros
    return f"https://wa.me/{somente_numeros}?text={quote(mensagem)}"


def link_email(email: str, assunto: str, corpo: str) -> str:
    return f"mailto:{email}?subject={quote(assunto)}&body={quote(corpo)}"


def componente_copiar(texto: str):
    """Fallback nativo sem components.html para Streamlit 1.56."""
    _texto = str(texto or "")
    _linhas = max(4, min(14, _texto.count("\n") + 2))
    st.text_area(
        "Mensagem pronta para copiar",
        value=_texto,
        height=32 + (_linhas * 24),
        key=f"copiar_msg_{chave_segura(_texto[:80] or 'vazio')}",
    )
    st.caption("Copie pelo próprio campo acima (Ctrl+C/Cmd+C).")


# =========================================
# ENVIO DE E-MAIL — AQUA GESTÃO PREMIUM
# =========================================
# _EMAIL_AQUA_TODOS_DOCUMENTOS_V1_

def _email_aqua_configurado() -> tuple[bool, str]:
    try:
        cfg = st.secrets.get("email", {})
    except Exception:
        return False, "Bloco [email] não encontrado em st.secrets."

    obrigatorios = ["smtp_host", "smtp_port", "smtp_user", "smtp_password"]
    faltando = [c for c in obrigatorios if not str(cfg.get(c, "")).strip()]
    if faltando:
        return False, "Configuração de e-mail incompleta em st.secrets: " + ", ".join(faltando)
    return True, ""


def assinatura_email_aqua_gestao() -> str:
    try:
        cfg = st.secrets.get("email", {})
    except Exception:
        cfg = {}

    logo_url = str(cfg.get("logo_url", "") or "").strip()
    email_resp = str(cfg.get("reply_to", "") or cfg.get("smtp_user", "") or "").strip()
    logo_html = ""
    if logo_url:
        logo_html = (
            '<td style="width:118px;vertical-align:top;padding-right:18px;">'
            f'<img src="{logo_url}" alt="Aqua Gestão" style="width:105px;height:auto;display:block;border:0;">'
            '</td>'
        )

    return f"""
    <br><br>
    <table role="presentation" style="width:100%;max-width:760px;border-collapse:collapse;font-family:Arial,Helvetica,sans-serif;color:#1f2937;">
      <tr>
        <td style="border-top:4px solid #0B2E59;padding-top:16px;">
          <table role="presentation" style="width:100%;border-collapse:collapse;">
            <tr>
              {logo_html}
              <td style="vertical-align:top;">
                <div style="font-size:18px;font-weight:700;color:#0B2E59;line-height:1.25;">Thyago Fernando da Silveira</div>
                <div style="font-size:14px;color:#374151;margin-top:3px;">Técnico em Química | Responsável Técnico</div>
                <div style="font-size:13px;color:#374151;margin-top:3px;">CRQ-MG 2ª Região | CRQ 024025748</div>
                <div style="height:1px;background:#D8E2EF;margin:10px 0;"></div>
                <div style="font-size:15px;font-weight:700;color:#145DA0;">Aqua Gestão - Controle Técnico de Piscinas</div>
                <div style="font-size:13px;color:#4b5563;margin-top:4px;line-height:1.5;">
                  Responsabilidade Técnica • ART • Relatórios Técnicos • POPs • Controle de Piscinas Coletivas
                </div>
                <div style="font-size:13px;color:#4b5563;margin-top:8px;line-height:1.5;">
                  Uberlândia/MG<br>
                  E-mail: {email_resp}
                </div>
                <div style="font-size:11px;color:#6b7280;margin-top:12px;line-height:1.45;">
                  Esta mensagem e seus anexos podem conter informações técnicas, contratuais e/ou confidenciais destinadas exclusivamente ao(s) destinatário(s).
                </div>
              </td>
            </tr>
          </table>
        </td>
      </tr>
    </table>
    """


def enviar_email_aqua_smtp(destinatario: str, assunto: str, mensagem: str, anexos: list[Path], cc: str = "", bcc: str = "") -> tuple[bool, str]:
    try:
        import smtplib
        import mimetypes
        from email.message import EmailMessage
        from html import escape

        ok_cfg, erro_cfg = _email_aqua_configurado()
        if not ok_cfg:
            return False, erro_cfg

        cfg = st.secrets.get("email", {})
        smtp_host = str(cfg.get("smtp_host", "smtp.gmail.com"))
        smtp_port = int(cfg.get("smtp_port", 587))
        smtp_user = str(cfg.get("smtp_user", "")).strip()
        smtp_password = str(cfg.get("smtp_password", "")).strip()
        remetente_nome = str(cfg.get("remetente_nome", "Aqua Gestao - Controle Tecnico de Piscinas")).strip()
        reply_to = str(cfg.get("reply_to", smtp_user)).strip()

        destinatario = (destinatario or "").strip()
        if not destinatario:
            return False, "Informe o e-mail do destinatário."
        if not anexos:
            return False, "Selecione pelo menos um anexo."

        anexos_validos = []
        for item in anexos:
            p = Path(item)
            if p.exists() and p.is_file():
                anexos_validos.append(p)
        if not anexos_validos:
            return False, "Nenhum arquivo selecionado foi encontrado no sistema."

        msg = EmailMessage()
        msg["From"] = f"{remetente_nome} <{smtp_user}>"
        msg["To"] = destinatario
        if cc.strip():
            msg["Cc"] = cc.strip()
        if bcc.strip():
            msg["Bcc"] = bcc.strip()
        if reply_to:
            msg["Reply-To"] = reply_to
        msg["Subject"] = assunto or "Documentação técnica Aqua Gestão"

        texto_puro = mensagem or "Segue documentação técnica em anexo."
        msg.set_content(texto_puro + "\n\nAqua Gestão - Controle Técnico de Piscinas")

        corpo_html = "<div style='font-family:Arial,Helvetica,sans-serif;font-size:14px;color:#1f2937;line-height:1.55;'>"
        corpo_html += escape(texto_puro).replace("\n", "<br>")
        corpo_html += assinatura_email_aqua_gestao()
        corpo_html += "</div>"
        msg.add_alternative(corpo_html, subtype="html")

        for p in anexos_validos:
            ctype, encoding = mimetypes.guess_type(str(p))
            if ctype is None:
                ctype = "application/octet-stream"
            maintype, subtype = ctype.split("/", 1)
            msg.add_attachment(p.read_bytes(), maintype=maintype, subtype=subtype, filename=p.name)

        with smtplib.SMTP(smtp_host, smtp_port, timeout=45) as smtp:
            smtp.starttls()
            smtp.login(smtp_user, smtp_password)
            smtp.send_message(msg)

        return True, f"E-mail enviado com sucesso para {destinatario}. Anexos: {len(anexos_validos)}."
    except Exception as e:
        return False, f"Erro ao enviar e-mail: {type(e).__name__}: {e}"


def _coletar_documentos_email_aqua(pasta: Path | None = None, documentos_sugeridos: list | None = None) -> list[Path]:
    docs = []

    def _add(caminho):
        try:
            if not caminho:
                return
            p = Path(caminho)
            if p.exists() and p.is_file() and p.suffix.lower() in (".pdf", ".docx"):
                if p.name not in {DADOS_JSON_NAME, MANIFEST_JSON_NAME}:
                    docs.append(p)
        except Exception:
            pass

    for d in documentos_sugeridos or []:
        _add(d)

    try:
        ultimos = st.session_state.get("ultimos_docs_gerados") or {}
        for v in ultimos.values():
            if isinstance(v, (list, tuple)):
                for item in v:
                    _add(item)
            else:
                _add(v)
    except Exception:
        pass

    try:
        if pasta and Path(pasta).exists():
            for p in sorted(Path(pasta).iterdir(), key=lambda x: x.stat().st_mtime, reverse=True):
                _add(p)
    except Exception:
        pass

    final = []
    vistos = set()
    for p in docs:
        chave = str(p.resolve())
        if chave not in vistos:
            vistos.add(chave)
            final.append(p)
    return final


def exibir_envio_email_documentos_aqua(
    nome_condominio: str,
    pasta_condominio: Path | None = None,
    email_cliente: str = "",
    mensagem_padrao: str = "",
    documentos_sugeridos: list | None = None,
    key_prefix: str = "docs_aqua",
):
    docs = _coletar_documentos_email_aqua(pasta_condominio, documentos_sugeridos)
    # Nao retorna se nao ha docs locais — permite upload direto na tela
    # if not docs: return  # <-- removido para sempre mostrar o modulo

    nome_condominio = (nome_condominio or "cliente").strip()
    email_padrao = (email_cliente or st.session_state.get("email_cliente") or st.session_state.get("termo_email_cliente") or "").strip()
    if not mensagem_padrao:
        mensagem_padrao = (
            f"Prezados,\n\n"
            f"Encaminho em anexo a documentação técnica gerada pela Aqua Gestão referente ao {nome_condominio}.\n\n"
            "Seguem os documentos para conferência, registro e arquivo interno.\n\n"
            "Permaneço à disposição para qualquer esclarecimento."
        )

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader("📧 Enviar documentos por e-mail")
    st.caption("Selecione documentos já gerados ou faça upload de PDFs para enviar ao cliente.")

    status_cfg, erro_cfg = _email_aqua_configurado()
    if not status_cfg:
        st.warning("E-mail SMTP ainda não configurado: " + erro_cfg)

    # Upload direto de arquivos (nao depende de arquivos locais)
    _uploads_email = st.file_uploader(
        "📎 Adicionar PDFs/DOCX para envio",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key=f"email_upload_{key_prefix}",
        help="Faça upload dos documentos que deseja enviar. Eles serão anexados ao e-mail."
    )

    _c1, _c2 = st.columns([1.15, 1.15])
    with _c1:
        destinatario = st.text_input("Destinatário", value=email_padrao, key=f"email_destinatario_{key_prefix}")
    with _c2:
        assunto = st.text_input(
            "Assunto",
            value=f"Documentação técnica Aqua Gestão - {nome_condominio}",
            key=f"email_assunto_{key_prefix}",
        )

    with st.expander("CC / CCO", expanded=False):
        cc = st.text_input("CC", value="", key=f"email_cc_{key_prefix}")
        bcc = st.text_input("CCO", value="", key=f"email_bcc_{key_prefix}")

    mensagem = st.text_area("Mensagem", value=mensagem_padrao, height=180, key=f"email_msg_{key_prefix}")

    opcoes = [p.name for p in docs]
    mapa = {p.name: p for p in docs}
    selecionados = st.multiselect(
        "Anexos (documentos locais)",
        options=opcoes,
        default=opcoes,
        key=f"email_anexos_{key_prefix}",
        help="Documentos gerados na sessão atual. Desmarque o que não quiser enviar.",
    ) if opcoes else []

    _n_uploads = len(st.session_state.get(f"email_upload_{key_prefix}") or [])
    _n_local = len(selecionados)
    st.caption(f"{_n_local} arquivo(s) local(is) + {_n_uploads} upload(s) = {_n_local + _n_uploads} anexo(s) total.")

    if not opcoes and _n_uploads == 0:
        st.info("ℹ️ Nenhum documento encontrado. Faça upload dos PDFs acima para enviar.")

    _b1, _b2 = st.columns([1.2, 1])
    with _b1:
        if st.button("📨 Enviar e-mail agora", type="primary", use_container_width=True, key=f"btn_enviar_email_{key_prefix}"):
            # Anexos locais selecionados
            anexos = [mapa[n] for n in selecionados if n in mapa]
            # Anexos por upload — salva temp e adiciona
            import tempfile as _tmp
            _tmp_dir = Path(_tmp.mkdtemp())
            for _uf in (st.session_state.get(f"email_upload_{key_prefix}") or []):
                _tmp_path = _tmp_dir / _uf.name
                _tmp_path.write_bytes(_uf.getbuffer())
                anexos.append(_tmp_path)
            ok, msg = enviar_email_aqua_smtp(destinatario, assunto, mensagem, anexos, cc=cc, bcc=bcc)
            if ok:
                st.success(msg)
            else:
                st.error(msg)
    with _b2:
        if pasta_condominio and Path(pasta_condominio).exists():
            if st.button("📁 Abrir pasta dos documentos", use_container_width=True, key=f"btn_email_abrir_pasta_{key_prefix}"):
                abrir_pasta_windows(Path(pasta_condominio))

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================
# AÇÕES DO PAINEL
# =========================================

def gerar_aditivo_renovacao_por_painel(pasta: Path, alerta_dias: int) -> tuple[bool, str]:
    dados_salvos = carregar_dados_condominio(pasta)
    if not dados_salvos:
        return False, "Esta pasta ainda não possui dados_condominio.json."

    data_fim_atual = dados_salvos.get("data_fim", "")
    novo_inicio, novo_fim = calcular_renovacao_anual(data_fim_atual)
    if not novo_inicio or not novo_fim:
        return False, "Não foi possível calcular a renovação porque a data final salva está inválida."

    dados_atualizados = dict(dados_salvos)
    dados_atualizados["data_inicio"] = formatar_data_br(novo_inicio)
    dados_atualizados["data_fim"] = formatar_data_br(novo_fim)
    dados_atualizados["data_assinatura"] = hoje_br()
    dados_atualizados["salvo_em"] = _agora_brasilia()

    placeholders = {
        "{{DATA_ASSINATURA}}": dados_atualizados.get("data_assinatura", ""),
        "{{NOME_CONTRATANTE}}": dados_atualizados.get("nome_condominio", ""),
        "{{CPF_CNPJ_CONTRATANTE}}": dados_atualizados.get("cnpj_condominio", ""),
        "{{ENDERECO_CONTRATANTE}}": dados_atualizados.get("endereco_condominio", ""),
        "{{VOLUMES_PISCINAS}}": "", # Placeholder para volumes das piscinas
        "{{VALOR_MENSAL}}": valor_para_template(dados_atualizados.get("valor_mensal", "")),
        "{{VALOR_MENSAL_EXTENSO}}": "", # Necessita de função para converter número em extenso
        "{{DIA_PAGAMENTO}}": "", # Necessita de campo no formulário
        "{{MULTA_ATRASO}}": "", # Necessita de campo no formulário
        "{{JUROS_ATRASO}}": "", # Necessita de campo no formulário
        "{{PRAZO_CONTRATO}}": "", # Necessita de campo no formulário
        "{{DATA_INICIO_CONTRATO}}": dados_atualizados.get("data_inicio", ""),
        "{{DATA_FIM_CONTRATO}}": dados_atualizados.get("data_fim", ""),
        "{{LOCAL_DATA_ASSINATURA}}": f"Uberlândia/MG, {dados_atualizados.get('data_assinatura', '')}",
        "{{NOME_CONDOMINIO}}": dados_atualizados.get("nome_condominio", ""),
        "{{CNPJ_CONDOMINIO}}": dados_atualizados.get("cnpj_condominio", ""),
        "{{ENDERECO_CONDOMINIO}}": dados_atualizados.get("endereco_condominio", ""),
        "{{NOME_SINDICO}}": dados_atualizados.get("nome_sindico", ""),
        "{{CPF_SINDICO}}": dados_atualizados.get("cpf_sindico", ""),
        "{{VALOR_MENSAL}}": valor_para_template(dados_atualizados.get("valor_mensal", "")),
        "{{VALOR_ADITIVO}}": valor_para_template(dados_atualizados.get("valor_aditivo", "")),
        "{{DATA_INICIO}}": dados_atualizados.get("data_inicio", ""),
        "{{DATA_FIM}}": dados_atualizados.get("data_fim", ""),
    }

    nome_condominio = dados_atualizados.get("nome_condominio", pasta.name)
    nome_sindico = dados_atualizados.get("nome_sindico", "")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_nome_aditivo = limpar_nome_arquivo(f"Aditivo_RT_{nome_condominio}_{timestamp}")
    aditivo_docx = pasta / f"{base_nome_aditivo}.docx"
    aditivo_pdf = pasta / f"{base_nome_aditivo}.pdf"

    gerar_documento(
        template_path=TEMPLATE_ADITIVO,
        output_docx=aditivo_docx,
        placeholders=placeholders,
        incluir_assinaturas=True,
        nome_sindico=nome_sindico,
        nome_condominio=nome_condominio,
        cnpj_condominio=dados_atualizados.get("cnpj_condominio", ""),
    )

    ok_pdf, erro_pdf = converter_docx_para_pdf(aditivo_docx, aditivo_pdf)
    salvar_dados_condominio(pasta, dados_atualizados)
    registrar_documento_manifest(
        pasta_condominio=pasta,
        nome_condominio=nome_condominio,
        tipo="Aditivo",
        arquivo_docx=aditivo_docx,
        arquivo_pdf=aditivo_pdf,
        pdf_gerado=ok_pdf,
        erro_pdf=erro_pdf,
        dados_utilizados=dados_atualizados,
    )
    aplicar_dados_no_formulario(dados_atualizados)

    if ok_pdf:
        return True, f"Aditivo de renovação gerado para '{nome_condominio}'. Nova vigência: {dados_atualizados['data_inicio']} até {dados_atualizados['data_fim']}."
    return True, f"Aditivo DOCX de renovação gerado para '{nome_condominio}', mas o PDF falhou: {erro_pdf}"


# =========================================
# RELATÓRIO MENSAL DE RT
# =========================================

def valor_float(texto: str):
    try:
        t = str(texto).replace(",", ".").strip()
        return float(t) if t else None
    except Exception:
        return None


def formatar_data_relatorio_chave(chave: str):
    st.session_state[chave] = formatar_data_digitada(st.session_state.get(chave, ""))


def formatar_art_numero(texto: str) -> str:
    texto = (texto or "").strip()
    return re.sub(r"[^A-Za-z0-9./-]", "", texto)[:40]


def on_change_rel_art_numero():
    st.session_state.rel_art_numero = formatar_art_numero(st.session_state.get("rel_art_numero", ""))


def on_change_rel_art_status():
    status = (st.session_state.get("rel_art_status") or "Emitida").strip()
    if status != "Emitida":
        st.session_state.rel_art_numero = ""
        st.session_state.rel_art_inicio = ""
        st.session_state.rel_art_fim = ""


def obter_status_art_texto(dados_relatorio: dict) -> str:
    status = (dados_relatorio.get("art_status") or "Emitida").strip()
    numero = (dados_relatorio.get("art_numero") or "").strip()
    inicio = (dados_relatorio.get("art_inicio") or "").strip()
    fim = (dados_relatorio.get("art_fim") or "").strip()
    if status == "Emitida":
        numero_final = numero or "N/A"
        if inicio and fim:
            return f"ART nº {numero_final} | Vigência: {inicio} a {fim}"
        return f"ART nº {numero_final}"
    if status == "Em tramitação":
        return "ART em tramitação administrativa"
    return "ART não emitida até a data de emissão deste relatório"


# =========================================
# RELATÓRIO RT PREMIUM — PDF DIRETO REPORTLAB
# =========================================
LOGO_AQUA_OFICIAL_B64 = """/9j/4AAQSkZJRgABAQAAAQABAAD/4gIYSUNDX1BST0ZJTEUAAQEAAAIIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAAGRyWFlaAAABVAAAABRnWFlaAAABaAAAABRiWFlaAAABfAAAABR3dHB0AAABkAAAABRyVFJDAAABpAAAAChnVFJDAAABpAAAAChiVFJDAAABpAAAAChjcHJ0AAABzAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAEYAAAAcAEQAaQBzAHAAbABhAHkAIABQADMAIABHAGEAbQB1AHQAIAB3AGkAdABoACAAcwBSAEcAQgAgAFQAcgBhAG4AcwBmAGUAcgAAWFlaIAAAAAAAAIPdAAA9vv///7tYWVogAAAAAAAASr8AALE3AAAKuVhZWiAAAAAAAAAoOwAAEQsAAMjLWFlaIAAAAAAAAPbWAAEAAAAA0y1wYXJhAAAAAAAEAAAAAmZmAADypwAADVkAABPQAAAKWwAAAAAAAAAAbWx1YwAAAAAAAAABAAAADGVuVVMAAAAgAAAAHABHAG8AbwBnAGwAZQAgAEkAbgBjAC4AIAAyADAAMQA2/9sAQwABAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEB/9sAQwEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEB/8AAEQgBrQJ4AwEiAAIRAQMRAf/EAB8AAQACAQQDAQAAAAAAAAAAAAABAgkDBwgKBQYLBP/EAE8QAAECBAQDBQUHAwQABAMECwECAwQFBhEAByExCBJBCRNRYXEigZGh8AoUFTKxwdFC4fEWFyNSGCQzYhlTcicoRoIaNUNIVleGkqLC8v/EABwBAQABBQEBAAAAAAAAAAAAAAACAQMEBQYHCP/EAD0RAAEDAwMCAwYFAwMEAgMBAAEAAhEDBCESMUEFUQZhcRMigZGh8BQyscHRQuHxByNSFSQzchZiJZLSNf/aAAwDAQACEQMRAD8A79GGGGCJhhhgiYYYYImGGGCJhhhgiYYYYImGGGCJhhhgiYYYYImGGGCJhhhgiYYYYImGGGCJhhhgiYYYYImGGGCJhhhgiYYYjmANgfa6C3Xp5YIp3F/7H4HX5YYEXNzuL/PfDBEwwwwRMMMMETDDDBEwwwwRMMMMETDDDBE+vr6/fFRzXN9um3n5X8P4xbDBFUlQOoHLprfX69384m4vbra/u9dsCLgjC2t/K1vnginEEXFj1xOGCKnIQLBRAGwF/j0HUgabb4sBYWvfzOJwwRMMMMETDDDBEwwIB3APriBboBr5WwRThhhgiYYYYImGGBIGpwRMMV5gSANb392ClAX8QNBrvbQYIrYYgG4v9b2+vDE4ImGGGCJhhhgiYYYYImGGGCJhhhgiYYYYImGGGCJ1tcX8L4YoUEqvcdD42PQ6eGh/bF8ETDDDBEwwwwRMMMMETEFQHr0/v5fQxP7fP61+rY03AARY3FrfD19cEVkqBG4BHiQL+f8AYYt4Hodj4+njjQAJIHj9e/4jGsBYAHW22nwwRThhcbDoBf1OGCJhhhgiYYYYImGGGCJhhiDextvgiE2BPhgDcX+VwbY0Pf528D19/jqcMEX6MMUQSbgkm21z67X93yxfBExo83tFVvdceG/x+tca2NEpI8/TBFrEgbkD1xFx4jptrvtjSN7C4Ol9SPG2IAJ9Op8MEWv7wPU2/XD601xpqTcXuCANNrdAdet7beOLp0AtYje401JN9On98EU4YYYImGH18r/vhgiYYYYImGGIvY6iwOx3ufQXPpff9CKcMMMETDDDBEww8dR8dPeemAtcA72BNtbD18d9N9tNcET+4+GmGG17G+p195wwRMMDqonodx7z/Jvh+nT6+GCJiCQATfQGx666eHr9WOBAPnuLbbjx38PnfTQ1sAki9hfU+F7b+7c7eNhgivv1Hx+fp54qVWIGmu+o0t9H4YW5bAG+t7Wtfe4394JtrvigRcX5ifOx11Ou+l97efngi1FXsQN+nxHXFj5eA/TX5407FI3vqLabXNvHzOLA3uPA299gSfn/AJGCKcRrcaaa3Ph/nE40130sSN9iR4b239+CLU+vr4YY00mwKiT5+7/OLg3Fx1wRTiq/yn3friT4jext77fxjSVcaXJvr+uCKACdhiDcGxwuRsbYYIrJvfS+4v6Y1eayhtYXvcjw21PXbUY0k9bG23v+H8HE3V1Tf1SSfjgi1duvl6/phj89ydLn18PPW/6H341UG+t9rAjf3/2+OCK3iTYWPj06H6/fA309dfgf3t8sVKEm/ta362+Nh77+HniUk6g9Nib3I8dfrxwRW/v8hc/IYYEWJ9q4IAta1t7+Zv54YItMLuQLanw2+v741Pr6+OKIFhfx+WL4ImGGGCJhipUAbG/jjTKiTcbdNB/f9cEWqogA62NtNL640gtV99OtrX/TwviCom1+m2gG/oNffiMEWsFJPW3rYevgNd9MVKjzbjl9x9drm/wHzxp4YIv0YY0edXj8h/GLpWLWUdbncWFtLa2tqb6E+mCK+GHvt8f2BxpqWQVJ8NiL9dz012/fpcispYFwN/0P10+ONPnV4/IfxiuGCK3Orx+Q/jFcMMEUgkbHE86vH5D+MVxNifjb34IrJUeYXO++nw2Hn9WxqC9td8aSUkn/AOki/wAca3v/ALbYImGGGCJhhbU6k38f46YYImGGGCJiiyRa1+t9L/HTTF8NOvu0vr9dcEX5/P8Aj9B9fPAAnYE+7GsQL7JN/Hx3PQ62N/HFrAbe/S3+cI+5H39+RRfnsRrYg6239+P0DYX8MQQDuL+GuF7Acx1O/mep9PE7DrisenzH3z+vYopww/i/u8fTzw+vr4/HTD9Y9c/Dy/cIhAIIO3h44gC1wNBr4/HX+PDbXE4W+W9tbetsUxMc/siqAbkaEWtoD8LkakajcemJCQkaCwNz69L/ACxJ8trn3DxPrh9fXT/GKx97fe6JiBfW/Qm221vL1O+uJwwj0+Y++f17FEAAvbqb+/DDDFPl8/L7/wA4RMMMMETEAAEq6k/sBb5YnDBFHW9tbEX8Be9veddB09MThfW3qfhb+cMETDDDBFUjRVhub+uo10OIK7W0uLe/3+g9dSNtTi50+XzNsV5ATcjQ67nx3Hr1HS3rgikai+tjtf8Am5+rYnAaCwGgvp9Hr/F8Rc6+yRbxKen/AObBFOGISebbz3sNt+unvtfpfE4ImKKtZV97i3qbD39Ri+Ite99rgjxFre74eeCLR5jpfpoLC4+uhJtt0xdHMDY3AtcXHjbr6dPXEqSPzEXvuBuDca9Lj5+WuHNy6KvzC4VbxB03ta46a21ucEU7+Y5rC3UAG3nuL/5xawF/PU+tgP2tjTTqEn/qSPjf9NPj641MEVCsC+huDb6N7aWPn64nRRsQdACCdN79N+nX3YKSSLafmJv5G/1thcaG/wCYC29zufduN7YItJWiiALAWxZG/lY/qLfviVJuSelxfXXp7uvjiyQm9gDcDXxtpvuPPTx+BFbGkdEAdSdQd999fQY1L+1y9bX8uv8AHS/nbFSnmKVAgixBGt9CfDT5/wByKiRcjTTr8MQQOYixsCdOvLf9wN/741gALgA6i++/L6m19dMU0USU9U2PkT439Ol7fqRR7NxZJ0Nzvt8T8+tsXBJUQQeWxsbW16eZ1udenwxTlKQSTYXA06nX08DiOZQ0BNvEgf3wRXUkAEgfM+Wu/n6Y0zYK9m2nXe1wPX6tjWCubXXe2u/640SLG2mwPx9ba+n6YIrgX3QNd1XIvfXxsfdbfFjYW0JttboNL31xUqsABcGyffp+n+PHEBZ6nTrawNuo94wRCskm23hYYYqbX0vbz3wwRaiBYa9enhvi+GGCJhhhgiqUhWpuNPL++KKTYm1ynxNv5v8AEDGrhvgi/Phiyk2Pkdv4xXc2G5wRMMW5FeHzH84ggjcYIoxZJAOoGvjrb4/r0xXDBFrk2UE+N7nw8Nvft/nTUq/oL28Tt8vD34pc67i23n6f3thgiYYYYImGGGCKwTzXtuLW9/rjVA0GngdfHGh6Y/Rp026dNMES3Xx/bDDDBEwwwwRMMMMETDDDBExpuKUmxSAddbkCyQQVa2UR7NyLJOoAJTfmGpiqhcW/a9/Lcb+JNhucVAkgdzCLhq5x+cLia8zNyyh68nU2rXJuqmKHzLlVPZb5k1A3S1VvyiXzxEoiZhJqVjpZEvfh0zhIgrgI2LhwVOsKf+8MPNN+ab42uHt0gIn1aknXXKHNsDYm/N/okjppe1yQBqQMYAsiJOqb9oT2u60IUUM8TFCcotzDncy/eS4q24PK0NSNeXQm4vzwFHOpV+Q9b77m+w5U2sf774vtoMIBMyYJz/6mPorTnEEgH6Dssi3/AI0MgSPZntaG9tBlHmyfjai7b3tc+eB4yMiHPyTetlX0v/tPmmi3Q+yujQQQfG2nXbGPBNMusgqLZSm6fatYAlQABJULEk8qQb30AGNZEiUEm5ToTcqOoIve/tWFtz0Fv6dcV/Dsxvxz20//AM/VU1u7/QfwsjUBxcZGxrqWRUFRwRUTZyZ5cZkS+GBBtZyJi6UZYaB1spbwA66XON3ZDmdl7VCmGqfrWmJnEPW7uDhp1L0zAkg2QqWPPtzAOEjRsw4WfIWGMRYkLyl90myl3VZCbhw8qeawR+cnluqyQTaxCeuP0Gn3lNqQ4lS21WPIpPO2o9DZSFpUCNR7QFiSDqMVFEAQD6SJ8u4n+VUPPInzmP2WRHii4tsj+DqiaYzAz3qWMpmmqyzJpTKanYuAkc2nz0fXFaNzB2n5WqFlELGLhWItmVx77syilNwUI1DOKiHWwlRHpMJx45ARQTyvZlIKgT7eTeaPKkpJBSVMUpEp/KOYagi9ikWJxhK7ULhmzi42+Dp/Iqja+jIKoqKr2nM4KCl9RvuRkJNKjoiT1ZLoCljUUQHpvIm5rDVNGMSyMU/FSqVzRuVOOwMNAtxKm+UvCm+znZkRlrXZhFQNRTSkpT/quSvpSiMktXQEOiW1fJothKldzEyap4ObSmJh7pcYdg1tuIbKC2mDaIdhwdIIOqIwHAx2g+vrnYXkjGD3wf1CyPo44sg7kfesyFKVYAIyWzeUkG2wV/oqx9RoSAQCCFD9SONTIx0DlczKKToT/sxm0kDUW/NRd9dTtsOmOL4y6caIH3ck325NSSqwAF7k3tpa+hGmmP1JoZxIAMM4LkEFTYF7gEHXcFJSpJuQQUqFwoKxX2FMRk8bmNtP/wDP1VNbu/0H8LlGOMjJFQ5kxGYfX2VZPZrA6f8A9HBWo1BCdRYi++KHjIyUuQHcxbjcHJ7NVCbG40U5R6BcdLm9vaANtOMP+ingdGCTtqADfU6Dcm97adb9ca7VEvrUElpWhtzFJsTqLg7EA3AIJFwpN+ZJGIikyQMnIEzv+Xb/APXud01u7/QfwuTSOMXJc695mGdemU2ZhOt7f/hQWGljpp4k6497y+z/AMu8zJwuQ01EVE3MQy9EtNVBRtWUumLaYR3jwhHKgk0uQ860jmdLbTiyWELc3TbHBOeyiS09COxk3jYaDaYSpau8WlLh5bEcoFyOoJNtSDp19c4SMyIDMPimfkUhdKZHR2W9YTyIUi5EbNFTmj5BDl0hSghDcPPo8sNEFKeQknnABOotbEH552AHdSa8zB52Py345+nrGXrDEHpYX28NBtp9dT00wCje4uAQNTuSbAX+Q26eFsWVcU4qpVhceNv1/jEm9lADxsb76eHSxO53ucUAVblI08dL+Pj42Gvn43wRX8/I6etv48cTiNbgkWNjf10t8hf9ddTPlbT9b7j5DBEwxCjYgna2p8PAa66XPS3h5uYaa77ef1fBFPS1gfPr+Yq0+Pyw9+31b53wt+/y3w1FtPPXw1/cWwRMVItY3Jtrqd7G59+tvhi31/HxxU3uLC+hvrbwtgigG+oBO/gN7eYPQb4vh9fziAQb26b4IpxB0B9DicUVzEEW03vcbC/nodNv7YIpFlAa+F7eIsfD4EbHFSPbHnr+unpigVY6W2O/uxcKNwTsdBbb4anr4j34IriwB+enu/bU9d9sTiiVKVpYW6n3bWv+3X4NeYEA2trqPE+J9/Tw01GCK4IN/K494H84jlGm/s7fLf4YhIIKr9Tp88WwRQRe/mQT7rfxjTSojfa2w8fX6/jVJA3x+ca/30/XBFqKv7KSRqdSPC48vjYemIRokqO4B0G2ljtv9bY1D08j+xH74nBFBVflOuqTvrppfpofDex89MaYUkXtza+IHT4YldwABoDf5WtjTGgta1vr6/zgisVEixtvf9fPzxXDEgE7YItRGx9f2GIVbmSDcc1x8NfA+OLJBA18f4xb68P1wRaZuARym17XJB0vpsNv0v02xp/Xn8MfoUeUHy8LfqPoa40RoAoXuL8w020I/Q7fPTBFBBBsRr4eu2GNQqB3uTpyWP5Ttr7wPcPC2GCK+GGGCJhhhgiYYYYIhAO4B9cRYDYAe7E4YIm2NFSua2m18apF/wB/PQjGidCR4E4IowwwwRMMMMETDDDBEwwwwRMagXoBbwG/0Pn78aeBGmuxv6Hptgi/QDfUYY0Qop228OmNRKgba6m+mv8Aj54IrYYYYImGGGCJhhhgiYg2sb+B+Njb5/30xOHQ+l9diAQSNj0B6eul8UOIPMgfMiUXXC4NKdM27Qbtim+VCkw/Etle6sqUAB97y4j1KRqR+U73NrEEnGU9FANk+02BY3O1jqdLH9tPA20xhu4ec2YTK7tAe12bi2wVTziZy4U0tRuQIPKuHUTYouEj7/zD2iOYXAFgrGQKH4rZepwf8SLG4tsbEHpdJBFx5/C4zwKkAgAjSDz/AMQYmI7qy/8AMfh+gW1fHzw957Z05f5WZJ8NmdY4e8x80s4oOTPZloTNS5L6Xp7LfMuvZ5LmvwMtzUuTU0lBQqBCREKOcpL0Q20FE41IjsEe1UUouJ7W2eRD2qkJcjc4IVtSr3Cf/KVM6EC9kgJCyBe4JFjmFkOdcFmFxEcKUlh0hKmc2ayizyalXLw456o5bAm9+dVxYjTUjS+W9N1AEgq5hqRrfmtfX+oqub3ve+uLL3vaQJgQDsJnE8d8ZxkqrA0jJBM4E+nHx5/fPQ44haP7ZLsZFyLPbOHNuW8WXDC1UMpkVZKeqCZVnJEMzmITDw8HOkVXIJRmDl7Opi5ywlO1BLJlOKaE/dlsHNTFqmTErjO0tlvDU/mll5QeZtIv/fqUzFoulq8pmNWhKFxVP1fIYCoZO86hBWlDy5fMWO+Qk8qXQtISi3dp2T7d+vqSiOzuz0yAgmoOsc4uIeX03lhlTltAPQ8ZUEwqScVlT0R/q2JgULciJVTOX8HCRNYz2p41tiVylmUNIiIxl+MhgvaLh2z/AJBkhkPkpktDxImELlHlNl1ls3M1oKFzIUVSMop1cw5FhSm0R7kvci22lqDjTbobWlLiVpTKkarw6dgBpJgcifhk7jflUdEmPLbbYfqsgKMuSpSUhAUCT7KkqAVYE8pKSDYm1wk3I062OPKO4bu0rk2a2bkbwa1FwUS3J1+vomZwknzoOcDNZy2p6ppumqwrlESijZVHSBcBH1ZUM2m8CIZ5DxRMVfeW23ecDfb/AMaMgCSR3JULWGgVcnTVKQQSLhPLZYVblsu1uefCJOZpWGU7+ZUygVy//c+sKlq2VQzqSlxymGXoelKUmQSUoHcT6m6YldQQi0AtvQc2h3mluNuJWavL2AnBMgER3gZ57/L5GAE5yIx9P5WNyFyh7cjuwImtuzGZXdKbmUcRMSlu55bg/hLJURe4AKeflCQpIVce0dl9xD5o8WXDVO8zs84SgW8wJNnrnTlhFLy7k8dJaWfl2WlYRFLQUTL4Oax8xmKxEqg4h5UTGvoiHgtPeMMrStOMz4JGtgbFIIGupOo9T0/fHUD7JziWisuuGmsqdQx3jZ4pOKCZBYNwoTDN2fu+PQAC+m43sMRp6quouGwAA5EgZEHzjfz22PEERsQfgRH6rs4mVwZJSGm7HQWCOuhskgk/DytqAcZXFRkr2iecWd0bLOC7iRyeyfoSmstaJiaspXM2jXp6/F1bUE9r1Jmskmksouex7MEuSyWVQ8RBx0xUERbLrrSGw7Y+RZ4yUqRdcKoLIO9tbjzSddbWvrrqDe3LHg0r7/cyq84KlCSgplGV8sAPtXTCLr9wnmBOhMSetr64qWlrdQgkQcg5OMR89v3xRpEjnPw4/me374eKg7Mbtn6yKhUnGrwrRDKwtLvd0TVaEqCr8xLbWXUMB4E97fS2mMg/Zndnzn5wkVNmPmBxEZ70Xm/VFZSGX0xTsqoKi3qbkdNS1iZJm83joiaRyYaOm8dN4uGl7TbIl8CxBQsCAFRLry+4y9hKrkgga2JtqTYbmx162v092LpBGpCbgctwBexJJFzYgHchO53FtrJqPdEwMbR3hXtI7D7/AMfclT7vr9dPrc4r/X/+U6e8fXvxbEWPNfpy2998W1VCQPU7DxOI5jsBdQsSLjQHY32Pu2664thYb218cES46G/phiNBsNz0H64m1r+vj9af3wRDsdvft78Vtqd9gR5Ek7eG36Yth9fXxwRRaw3J66n5YnDEEXtqRrrYkG2vhgihSuUA2v6EefTfodLftcVgba/EfqMUKSkj2jqTqCfHrqPEkb2t00xAFzYi+hBHhe1ydOlvdc+OCLUKrEC1726+focSBYk+J+v188Vttrze0NRrYeHXbr54kXBUTtfS/v8A7YIpJA3O5sB4k6em9t7b+uKgn2r+KdN7XNuniNfDXW2wvod/I+fkcRobgWvvc6C41GvXp8RgiqpINrAC5t7tf4wCdBf+kn36/wBvmMSVco11VcCwNtCDc+Nhp8QNb4oVEkaFJ8L7+nx/TBFZJCQonYE6/DwxHMCq3MQLbg2F9TY/L3291zYA3tY76ef1rimhWALWtqNLdf4wRag29NNfI2+Jth/f5W/nFNQTYaC5/wDaRvp/7vG/W51viQQq+m3v392CKp5dQVHfrc6i/l9fPFBa+u2Naw8B8BjTULbG9jsAdAfedNP5O2CLU5he3lcHx3/S3zxVShY2Ovv8cadyLHbw9MRgiXJ3v7zhhiQLkDxNsEQC5A8TbF0iyyNdB+tsQUlNiNddrY1NL30ufif5wRTiilcqvEW8dNz/ABiO88vn/bFFK5jtbTx9cEViu4Itv5/2xUGx8uo8RiMMEWpYElQ2AuAAd7fta5+Z3wxAUBe6QdOlh8dNR8MMEWrhhhgiYYYYImGGGCJhhhgiYggEfWmJwwRafd+fy/viCgjbX5fvjVwwRaFjvY6eOIxrL/KfDT5n+2NHBEwwwwRMMMMETE3Og8NvfiMXSm+pGnTf9j+o9MEVMWR+Ye/9Di6kXtygDe+p1/X9sSlACdT7Q8L2Pprppf8AvgithhhgiYYYYImGKqUAbXHifIfX74op1ATfm/qAIGtwd7i2xB0sb7adMEWrijiuVBV4aeO4IFx1F9SMUU+i2h/7am9xYkC9/DQWA8LaHT8rkY1ynmWgC1wCb2tqP+p6ePzxXS4wADkjg9wi6j8rkr854/e1NeaQSGeJujGOZI09jKSnx7ViSdXCbWvYi4PXkAmloxCj7JuToeW3mTcqsRr0sb+OmNDIuQMzzjf7WCKQlLiU8WVKtg2uoBOUFKEpB1uRcE9R5jHL+IokpB5Ic3AJCSgC1gehN9BbcXvfTw2LXENa0B06QMkTON8HPHp5qw9oLiSM458h2K4fR1IRUQYV5aHPvEC+qLgIllbrEVAxS4SKgHImDiWHG3oaIXBRsXBreZWlxUJFxLBV3b7qV+AXQtQRVwud1Y40pIQpJqiow2pOwSpBmndFJTpyJuggWA5bDG4HElljxO1zL8tMtOE+t6Iyzzdr/MOIlMFU2YssZmVKmTyDLfMOuptK45C6SrZ6GXMmqTbahY2FkLzrMQltKn4eGciF44G5XcSnFdwlcXMm4N+1HkFJSiJzdhYGJyEz/o6ElMuy4qWbxMazK4aRPzqWyqn5XEQEzmHeSlbsyk0kqamalclbc9lKqdqeTTOGjqYXAOEknnYTG5MnP6eio1ggkAYkGZ7T9+a5SweUUPL1xL0HK2IZ+MKVxsSxCQrUVHKRqgx0Sw01ExpQQAC+4sgJQGw0pPMfWq2gaZy6kMdUtazyTUlIpe0p6Mm9RTSBlECy0j8zjsXHxEO3c2c05itxTaktJcWCjGTwUGyFE/dhbSxAIvoCdVaXAUCeVSwEqSrnIUMep1vw+ZWZpy+USvMnLujq7gafnssqiTQVZ03LanlkBP5U+HIOMMtmALUWwu33KZQDim4eZyyKiZbFFDMZziQc0QNI7xE9vSAc/tKiQTsSPjHbj5x595WP/g2yvmHHjV7EdljBVBBcMkpj3BXHEOYOMkMjq9EBMWoaZ5f5KPzVhuY1ZPJkG4qXz2uoCDcpWi4ZmIUxMpjVq4CWo7T8klUup2TSiQSWDZl0nkstgJPKZdCpCISXy2WwrUDAQcM1chpiGhGGmWmkkhDbaUEkgk7MZPZgUzNqYhJDCSOW0TM6RgJfK42jpTBJlsllMEYdKIB+mGGmIaGepeJLDrUpdhIdoQrsNHSmNbgp1LJtBI3YM/gCQO+TYKBOlyD4+F9ep8tLYwqjqlR4ke6MNA2E9zgcxyPPAWQ3S0QCI9RPA/ZezJVynxKikixG97EdbG5tqPgMdHHs1KcXN+HyrI1KSq/EjxIN3SdLtZrT4KB11NzY+NtNSLd21ufQS3mWw6kc8Qyi+w9p1IF7L8bHVPuvjqGdlBJWI/hgqx5hvmSOKLihRcC1yM3p8pJvc6BtbYGtrAG1yb3KGprpyB5bzgz22P8AbdQqaXY3McHG65KJo50JP/ECADqRa1r62PTe19Le4jJ72b8KqAiM5IQp5ShvLdXUABxFa6Wva3s7+NraXOOODNKpKU3ZNjr5EX6XtrrY3vvpjmFwWMw8lqnOSGI7tTkJlkoi2lkw9ZqBPlzrVY7m+ulgLj//ABuHMffmoMawFsTvvGd/X743WRPlT4DXU31162vqB5XPriFK5SRa5G+vgLnxx+NEahduVaFa62vfbwJFvh1x+jvEEg2JJ0NumhB66bfI2vjCz6Hzx81kLVBuAfHE4qlSToBa3r/J/bFsETDDDBEwwxCiQCR9a4IpwwuBqRce/wDbEX1IttbXxv8AxbBFPqb/AFoPcNMMQL2F97C/XXrricEVVAqFhpgQSbg291/3/W+w1xKiQCR9a4nBFAAG2GigoWuBoR8PG3iMTip9lKinQn+R64Il7HU3J26aDp8/0GJJsL3IA9oiwO172N9LgD4Ybki19Adhfrsdxt4jE6euhGuu/rf62tgipYKPORdJ5Ra+t9B+xtvhqVFNypRGt/ADpa9rA9Rb95NhyjxNregJ/UXxI1B8ydeuhtuPIC/jYYIihcEajz997dNRscVCLG99ulvXFiba2Py/nFeceB+X84IoJUbjWxNvy9Ph/f06vyed7nwtb4+ONS9x664qq2gN9b29ehPx8+uCKveeXz/ti1gb63URYn18Rf06eOICQLXGvjriqPzE9Tcnw3HmdxgisUXtrsLbeGNOx3scapUQoC+nh5m/9sWIBFj+uCL8+LhWo06BO/nvixRe1tN974coAuRqB4nf5YIpKgNrH3/5xQruQbbX6+PuxTDBExNtL9Ntx5dN+uIxYWIsB7V9D00tv6WN9DgiorY3Nttd/DwxRKSD1A08NfLQkafvj9YSBa4BPUG9r+l/lYWxprAFrefU+X1/jBFUAk2+rYYugbnwt++GCLUwwwwRMMMMETDDDBEwwwwRMMMMETA6AnwwxRew9f2OCKqiTbw3A667X3H10vimGGCJhhhgiYYsBcE8pPhYga+/f5YuEbE/D6+v3IqBJP8AfGqNAB4AYnDBEwwwwRMMMMETGk44EEDXW97DbY9SAfn7sav1rj86klZ1F7HUDTyGpBI22t5eeCL8cTEBIPMoW5gq++oOx00Vtvca6Wvj1mOnghwSSByjUE+0bXsTv0IN76k6gDXHs78PzpKCLBXu30JvYa3N7+Xjv6nM5EqKQsJTclNgRYkAm3tWIF9hrodQRqALjANzmCOfRR96OCe2f1legTivTDXCVaJ5ifaIOpJv8xpe2xFra7Xz3N1TCHUpdFwki4Wo25VGyjZOuttgdNbjfHuk/oKIfCyA5rc6k3TqL+0ABcE20vv4nGxdUZXzRYWWS5zKCgk2UdubVQ5FEi35rIUEJJWRyoJTmM9nuMZHqNsT8c/CDsFbcXY1YziPUf232nyBGHPhSrJcDxQ9qFOn+VRn/FrJVNlRUu/3bJuinFKBPX/ziQrSwsr2hYA86l5l86iSE8vQWSAAB6Wtp5k6X8Bhgzq4Ue1tyo4j+ImtuFihcoq7ykzmzDgcxIeGqioKcamrEyNGU3T8Sh2HmFR0bMIFxtUnUwqH/Eo2EcQ2xFMPIW++2nbx6H7d+FX3cRwmZIxJGnes1FTpTYb25c/Ta9ir2UEdbHrOW7g7RAAOduf3x3zzFZrJhm9DyjOThqmj5QG5VmfV8Tcj/wCbw8Z4QaU6IOinIptJJ0uoX3sXaGZZ5L9onw4VDkZmMxK5XUDBXPsq8xVwLcdN8t8wIWCimJTO4MofhoqIlEYYgyqq5IzGQzE7p+IiYX/gmLUvmMFiw4e8q+0+r/OGj5pxR5S5Y5bZdUUaiqBERSc8gphPJjP5hR9QUpK5WiGgK6rRaWi1UMZGvuOtQLaUwYH3pRJQckUVlfVUMkXESpSb9XBc7G6gk63udFdTqNbSbRpuy5xa4wAN5kDc9sxwc7bKLqjmQBBBOREmMZ2x2ziPkuAPZq8dmZUtmU64EeL+Jdl3EvkuXZJSk/nAKEZt5fShoIk81gJs5ytVDOISUNw8WxNG0iKqemlQk9d+8ThmolozJprdzm5wWh+XQ2KrqIHLZAN7XHMAPZvyqsoKSnDhxn9ng/xTQdN1XTs9jcsM/MuIpmPy5zalLUa1MIEwj70bDyKduy52Gjn5EY5z73BxMMszGQRanYuWc0PFTWXTLh9KK27cfJSGYpWoMhsveIhiVo+6Q9bS+Pl65nOmW1lDT78TI6yo9yNcdaKEByMoeWxzxSHIpp6IW665R9MMJByMEOE6TgbH3hv6ZnyRriXGYAjyxt2HmeMH0XZSi6qj3yxMJXGsymfy5uJMpnLUKlx2HEQlHfwMe2lbBmcnji01+JyovQ6I1DMCthyFmUslkwh/Hu8S0ylbpgahSZVMmUjvGVvB2FiGxzJRFy2KbB+9wTvspbdWERTalFmOaZi230N9dSP4mu2zjme5kvAjR8gfWOYRc+mvdsIXf2VgRuaMlTZuwUol8WNvbAuMfryUyp7XbNDNykcwOKyvaCoHKWQRs2jY3KGmGafi5jPXIuRx8ugoRoyGVzsspbj4iCjYmLmOYkUtAgmVty+PALIrT06my1zml0TBA/pnMb5iQY37yhdvB2yMSIweRMGOY38wuxJD8UrZiYMIig4PvUNqFKNx3zdyL6m4N72Fydr64wu9j1Unc8Ik4iF+3+J8SHEhNkqsdorM6bLsQlKlcxCbgAfl1UATYc0pPlDUximnLRSi3ENLTdtSFKUlxKkm60q5RcA2UVkXHMV/1YV8teDjtoOFiKrWhOHylsj8wcpI+vqsqql4epqjkynoVmpJrER7jzX3ueUVPpYqYc7cZFS6NcmDEJHLiEwcVyKdRi9XZTaG+zcIJEzvuMdp0yZ4jsVBjyS6Y/MAIG40AzjnHA+e67FiK6YaRqUkHrzpCRYm59rlUdtRtpj33IHM9cPmNmg3DqBDsmyycXyEjUGu2ubWwtdsgkEpsN9NMADWW32gmckNM8PvCtCFVk/eYiuYNJBsEkcqM3Ywbm6roFjc7aDLp2anDJxv0tJs0Kp43Wsr5fXdV1BS8LSUjyqmMVNZVA0jTknjFGJm8W+uIb/E4ybTyP7uHhoqIQ3BsNPuqbVEpRjG9yeCNoMnOx2kbT8fLKuZBG0SNsHiIOecHiFmEkNfORiUgqBJG9t9Nid9LkW+Hgd0JbUTjyQAq4JSLFVgPADrqbA30Gm9rDbaSZfqhkIBSQEp1tzE3NtdABcWsfG9wdDjc2XU8Yew5djzaAC5FyLa3AuPW2gOpItPc0GDP6GJHzny5xiVdZ/VknbczxH7L22Hje+QlQ9nmSFkXuNdLXO3jb3HXTHlEO7AC+gJOw13I1Jv4Dbx1vjxcNCFsWF0AAJsADex3J2Bt4gAXNtzfyzbZBBN9AbXt4g28QTbXQC9vDGI8guMbY8+Ari18MMMRRMCL6HD6+vnh1Iv4a7jrawuPA392CJhh8/db5XP64YImGnUX8PXxxUK1sQR4XFr73/b4+4VPtKKegF7i246fM/AeeCLUxAvbU3OCQEjlBvbT9+lh8NMTgiYYYYIoOoNvDBIsAPrU3xOGCJhiqlcova9zb9fr98WB9oDpa5+I9+gv9DBFChcEfW+JAsAPAD9MMQTbQanoPK+uu21/wC+CLT5lAnW4BNwQB1t5fWu2JVdXJbQm519x+v0xATzFV9Dfp7/AFxc2BT5aAeN7A9MEVSom5B0B2tr0126m/wxAV7JTt4E9b3vtci3zPji4SOW2tiT8iPIeHzxo4ItZJ0HW2mgP8X8PLfAqsCRvcj321Hw0xDex9T8rYqofm8BY28SbD16k9NvcSKVL0HLcWvfrfzG/p8MSkHUnrY/rjT6DTW+vmLHpY7eoxYLvdNrC2/gLH3nbywRWX+X0IP6j98SNAkHci4/X064hIBBG40F9Be3kBp78TyjTU6AWsRtqPD1wROUAGwsSCOuCQoCxtby336+vxxJSFDX6/XxxoqFjYX2v0vfXb+/+CLUWCVDUG40ve+mp8dsR3Z8U/E/x+uIBAAIJJA0B+B0sLfH3C2J5zY7dB6Xv/GCKASkkeXztocMTyhSleVtvT+2GCLUwwwwRMMMMETDDDBEwwwwRMMMMETGgSSTfy0/sfDX3+uNfDBFp6ApJI1AuN+nv9PIj1tp417JJvbU731ufHyvviQB0HwGCLRCVHW36YslJB1HTy8sSVgba+/Ed55fP+2CLUwxHMLA9ddL7nwGgN/jgCDb/wB23qN9fLBFOGIJt5nw6nEjYdMETDDEC5Avvrf3G2CKcMMMETDDDBExplpJJva1ha+p8ddB1vt8dMamGCL8jkG0sC6QTubgXOu3hc9dwet+nioiTQLwPOykg6HfqCk7g2FjbS/oTrj2DGnyf+7w/nxFxpv5dLYqCWzkZ23B7fHJ9NuU/wAbjJOwzHP+ML1A0pLFhX/AnQ6H2STbXqBbU+N77bnH5XaLlqrnuWyNVEFIA8TfUnTxFiBre1zjy9V1fSdB07OKtrWpJDSdKU/BOzSfVNU02gJFIZLL2Ld/GzSbTOIhoGAhWeYBcTFvsspUtIU4ASR18eLb7S7wPZEzNyksjoKpeKuqmmnVRs1oZ1ilcrZO62txoQsZmDUjHfzh1fIFoiKMpmqpMWyUrmyHgWDtOndL6n1SoKdjaVbgkgF4Gmmz8pl9UkMZg6veMHjhYd1f2lm3Vc1W0xnBgl0RIDZkkT8eNlngisupS9zc8M0dCdAHDbmN9E3un2NfMb7c3gY/K+nksl2JTCMstpUt1x7labbQgXWp1SlBLfKDdXMRoL9DfqnwHaR9ubx4raj+Ffhmdydy2nPImSVFIqKhkF2Bfd5hFxebuecTKaOmjDTZS4l+lqZgFkquymJK0Mo/bWnZN9oBm3LBUHHl2kNPUXKpiQ+9TFU5q1tW0ngkFQWuFcpyLmmXtAtvQ7gXZmSffYVDouw4pAbKOqo+EXUXsZ1TrHT7RziJt6Ln312DDDp9jb6nat5l0SQJIOOauvFNNod+EsK9YDapVDKFGDEEPqEGAZJgCN5wuxpUk54aaTeWmqs48nqYdSLrRUGYVHyV1IB5Pabmk4hlAEi1yOlhoDjbqIz34F4RwtRHFDw7qevylDOblAxCySdAn7vO3+Y625QFb+eMBtDdj52clGd0MxuPGsK1iEAl1vKymafkEvW6NS2y4JDmmt1HMOZFo4uG3MBa+OSsh4E+yVkBb+515xA1Q8FApi4mYTOHU8sJslSxDZeyGHV0tdlI0HKkC+N8zwh0oODTU8TXTNIOq26BVoCfdB0vu3smZMENIPptzV145uKQhn/SKLsA+26gys4D3Mn2YdpzMj5AZWTmJ4juAZTzjC+Kvh3ZeQuykRea9FQllHp3kTNGG1EHeyyRcEkaY9ypmruE2sltN0nn/kPUrz4BZZkWa1BTeIdSpVvZZgJ6+65e5B5Wyb2ABO2IWpuAvsl5+8+iZVhxG0+HVqUuJhplFvobUbguIDlDTxw+ADjDib6qFgSON9Z9kD2Y9ad+aC46sxKEinEgNIzFpmSziXtKCTyIfXH0ll5FOJCiL3mrYSOYBVgbKvhPpAA0nxNa7e9c9Fq3DTJaM/hZLe8uAEHElQtfHN1Vc1r3dGqnEinfNaTIaY/3cHcnadxhdomS5S0tFttRsvdl8fBPJ52oqFdaioV1B1BREMrW04gAkFSHCkHQkEKGPf4XLKVQ6UhEMyOUXHLykEXP5bKIKlaqItcpIULpIOOoPRXYn8V9APRlU8AnaX0TNY6CbU8hmiq8rrKp51KQFtw8yhsvqgzCk7y4hYSlxiewzUI4bpeR3YKj5wcZP2hLgHiX4riEyVjuInLKRLU9NpvOaJlNeQCpPChSomYQGZGRcSzNKebEKlx0TGupFHdy6sORcvVaIA0Nx4Ua972dN61Y3NRpxbXXtrG7LpHuNp3AEnjUHBsy0wRK6m38TsIabiyq024mpRcy5oge6JlkmMmR8uy7ezFFwDRB7pAPjY6m4OptYa7XNh1sdvPQsjhWPytgm19dL+th0Ft/QeB69vCj9pS4MM7JlCUnnnJqn4YKpiB3Lc3qNbVbZXRce2sNvy9Fa03CJmsmfZX3hiIiqaTkUmhUJBenXMpIV2BaGr2iczaXlNb5c1XTdc0dP4YRckqmkp1LagkE1hieXvYGbyqIioCJRzhTai1ELLbqVNOhtxCkjluodM6n012i+tatGSAHkE0nnEaajZY6RyHZB779Ja31rfN1W9VjzAIZMOGGkAtJkb4xJ2AO69kRCNIFkpCBe1gE7DqT5+83JPnjWDQBJ+A6gHcX1trvY641bg7bdCdCfHTcWOnrrsQSxqjIifl/jbETzyszExzBPntgx27nsoAA2xOGGKIlz0tY7+Plb374fWpvhhgiYqoqB0AN7DXyuf0v64tioN7/APtUR8v4Px+GCK22GG+GCJhhhgiqQRcgXJ0tfT1/bEgg3t03xOICbEm97n+cEU4YYH0vgiYYqFX3sP8A8wOJCknY31I+G/116YIh26+ViB8b9MRY+Z/63I9kncnxB08f2xbEEkGwFz1F7EdNvq2CKBzAagk6W1G3nqRce+/zxC9h6/ziUq5r6W+j/GJUnmFr26/rgiXGvlYHfrtiT023/T99f1xW2pN97H0tiCsk/lJ6C3UDS+xwRQoKJVa4BGlrW89zfpbr79QdPGsVeyCBdR6X946aDbfTXfGieltfHy3+O3zwRat7qFjpY+Qvr0xKhcGw1Nv1xRIsUk9b/pp8b41fr6+GCKiARe/lb53/AG/v0qQQb3sCo6+V73t4df1xZZtbUjf37YqfaCSNdNbanUDe1/1wROYA3Auf+x0vfU+z0101sTa9tcSFKVtvy7Da9/4xARcmx00sfH36DTF0psSb9PD0wRSL3Vfa+n17r+pxpr3Hp+5xqagXJuQCT8/2tjSUb2O2n7kYIrhSQBr0HQ/xgpRHKRsd/MafP/Gl8aWNQe2ANrAeB/xt0OCKwvcm3sqsRtpoPedcMLkEJtfQa3+PTywwRWwwwwRMMMMETDDDBEwwwwRMMMMEUXA3IHrhceI+IxVSSTcdB8d/rpvjTsTsCfdgi1AoXIvppY9Ot9foeGClaaHXyPT3YoEKPS3r9XwKFAX3HW3TBFXDDFk2Bufd5Hx/XBFIQdDex6aaj9MWUDcWGwO19BYfDztvbToBfr5afvf9sMEVU7ajUdSNbep9+LYYYIhvY23sbfDFUnQXOpv79Ti2tx4df2xWxufC4I+RPj1vucEVsMMMES4OxBwxB5r6Wt1ve/u8OvjfTwxOCJh9fDfD9dwNCSRsNdPQnS9rkXxt9mfmrl3ktl9VWaea1WyagsvaIkz9QVXVlRxjcBJ5NKIUAuxEW6rmWsurKIeEg2GnZhGRjkPBQkG9Fvssrkxj6jmsptL3vcGta0EuJJAAAAJJzsJKo5wa1znEBrRJJMADkn0E/wAyYXv61oaSpbikoQhClrWtSUoQgA3WtSiAlIsbqNgLG5G+OuV2i/2iLh+4XZlN8oOF+Xyzih4gYGaKkMzVJpqp3KGh5mOeGfhZ3VEnU9FVfUcFHFmEdpOjVuobjERsFNKmk8yhFwEViH4ru0841O24zkmPBP2alJ1dQHDtFLag68zAiH46lp7WNOffIj73VeaVXQ0GheVmWcU20iFhaDgYiJqar0NxUBNFTSKmiqDlnuNEU92fnYnwIkdFSmnONztD2JQuFqLMKbBlOVuSdQvuOF6Ek8IiImbFPzCURHKmIgJSuMzKnDEM9DzipcvpdNhKWPRuheC31atuy6oVOodRrgPo9ItpDmNJaWVr2sSGW9ISSdTgSAQ3UfdHJ9V6+yjTc6nWp29Bs6rl+HP/ACyKLPzPJiBvjyyvRHeA7tPe1Ai1cRXalcRbnCZwqQ96rakFfuMUe1JZUpKnZZCUZkZGR0pklHhqEdcZaqjNWPVV/K8y5EQNVuRFsb40ZX/ZJcDcHK5XwecMDPFJm7KXSt/iE4gEOx0OqYMFZTO5M1UUufLESly6YWFomhaDlgaah30zt9xQfdxfZ98Vuf3FpWCqxz1zFnVVvIdT+E07DPOyeh6XZF+7hqZo+CWJLK0oSpBiIwtRM5miUNfi83jlp7zHhqVhSe6JF7puR5hJHNe2ttrGwIvvY4986P8A6ZvdSpu69dinRAbHSulPNnY0wAB7OtWYG1rkmRMuaxxwWZK8e6140qBx/wCn0tLpg3l00VKxy0A02ullPiMSO+0ZW697SDjDzojnfxXNGLoennQRD0tlZCJoyXQjRUW0tCZsuRFVxqUNhJUmMqh5q/OUwjSCkq2bgIiZTqZKm85jY6dTWJWFvzWaRT0xmcQdSVvR0W4/FLUVKWQFOqFyq2p12UplCUhu49qyCCQB/Ub6jzNr3JA9LY5AU2yk93p4BKel7XHXS19LXB2Gug76z6B0XpTBSsOm2lANb+ZlBmomBJLtMknu6TuSTkrzPqXWep3jy+5u7iqXYh1V0RgnAxiOMb+S3fppkp5CU2Nkm+pVa4J5djdJKdLkbdRjfunU37vlsNUkhIIOguDpYgk6K06b66bNyFlIQzy6AJAIGp/ouT420OpudOgN96pCk/8AHykDVI6ja/TcgHl8rD0xC6DQz3GtEidozLRxGJM7eULQVajnOy71+EQSDOcSTvt5LTqNslTxsdefUE8pPQDpe400vc6DTGwFSwpcLhsq+vIgEkK0uTa5HMLX0HMAAdwQORFRCxcVYkJWseIF9Ovrcba6742Pn7af+UjUgkg2Ol03Nh4WNgR1GoFjiNqWlkkNJ1CGlswMdzny4IHYK3QfUBJ1FpMbahAEbHVg4jbbjOeNU6RGwEciaSyJjJdM4dRXDR8vffhI+HXb/wBaGjIZTcSysakKbdTchON0cvu0K4xMkZhCPyHN+eVVIYVQD9K5jpFcyWMaAF2/vE4UangE+yghUlqiUWJ9ttzu2Qj1GoWAecWsErKQQOpJIuf+u+x31PidgqngQpDoIvY3vYaXJ009LW9L3Fjidz0LpHVKRp3/AE+1uGuJn2lCk4/lE6fcluJE75MFdD07rHULSo11vdVqJaIGmo/yjUJgjyIJgHGFzQrXiB7NTjdcmkp46OE2AyfzGnBQqG4lOHxD0vnzMwcR/wDrSoIeWw7E8dVDlSy3Dz6WZry2IQ6pcRBQqm27bET/ALPLtBuB6XM8T3Y18Vs34lsko1z/AFLMqVoGJlcyrF9UEoLellUZOPLm9E5rOpgWzCzFuXS6BzBaeQuCgaOgIttMYrg1U0IW1qUNFjZVjrpYXA0IVuRqdr3uQf2ZO8Sud3C5Waa7yRzCn1DzlKiY2Eg31RVNz+FsUKg6lpeOU/IZ/C8tlstzCBdXAvgRcC7CzG0enh+sf6ZuFCs7oF4Q2P8A/I6kTc2FVhYAWU31NVahEwAxxY2SNBxHpfRPG9UVKTOoUtbSQPxduPZ12GWyXaSGvHOQJxBmV2AOzk+0m5UZ4z2DyN466flvDBnoqcQ1NwFWuNzCVZQ1FN1qXCrgKhNQRBnmU9RmZNqgjLardjqbU73aHanlzz33A9otiIYimWYmGfZiIaIaQ9DxEO6h5l9lxIW28y42VNuNOJUFoW2pSFJUFJPKU36OlczXs/8AtpoOW0jxLyWRcHPHkZSuS0VxIUe1BQeW+Zk9Q4yqVyit4OPi4JMyVHRHdswcgqaZNzxtbrkso7MP71HQtPxe3WQvHf2gn2ffOOn+FPjwpqd53cHk4cfNA1NJYx2oDJZE44gqnuR9cTpqDdmsskykpFQZN1ZEy12TCJQ/I/8ASyYuGiKh+fOv+EX0LmrSp2tXpvUmhxd0ysQadYNLQaljWkMqU/6okuAcGkNMBezdI69TuaTagqi4tjpArj8zHENAbWYfeaM5kRIO676gIN7HbQ+INgbEdDYg2PQjE42gyHz2yl4lcqqQzqyQrWU5gZbV1KmpvT1SSdb5ZfbWOV6CjoSKbZmMmm8udSuDnElm8NBTSUTSHjYGZQkPGMPtJ3fx57Up1Kb306jSyox2lzXy04wRB2I2zv5Lqmua4AhwcCAQRkbDntvHMCI5TDDDEFVMQABe3jc+uJ8PrTy99sPr3/VsEVQok25T6nw+FvdfFsDp5+mIvpex2v0v8L/XW2CISNQCLm9gSACR46+dvfgDr1287ab2Ox3G3kdiMRYEhXltt8el7evriQLco3sm1/h+tsEU4i48RpvrthcXIvqPHTpfrilr81iDc3Gvgb6/HBFqdbdfDrhipHtA9APf1/nFsEWmoC1wBoRci2l772xe4BtoD4fX1r54C99QLWPnr00It8j/AC9rW9rEg7n4kW1Phrp+hEUSACBfX4ef7e/wxT2gQBZRIFibA6nY3I23PT9BqYi3tXvpa3z8x+mCKECwNxY31vobeYPobefqMTcD/B+vrzGAIuR1Fv2tc7nr8vPFVaEE7AnbfUfzfBFbcaa3HxxN7De1tulhrt0A+vHEJ/Kn0H6YnBFXXmBFrW1I67i1xv59PHXEJB/qSkeVhe/6W8Ovji5AGg2sPLXrp64YIosNNBp8vTFf6lanQA26bpv6aE3/AL4tcDcge/FNLqubEixB6XA1J2sToOvlgisACVX1/LbS+ljtiLAGx9nX2eUWuNRrb3eG5xYcot7QvYA6iwtf6+Vr4qNTrY6+yNr69LWuDpbBFKrgezpYi9tLaEjbGlc+J+Jxq35hpca2N/DqOuu3hhyJ8P1wRVSSD7V7W6+o8cXsNNB8B9b4hQuNwLHr7/r/ADiSoXtufK3Tpvpa22CJYeA+AxBIGgFz1A6eF7DrrbA+0CBp018bjwJxYbDyAGnkLYIoABsba266kb/R+eGJuRe29iB4agjXy9MMETDDDBEwwwwRMMMMETDDDBEwwwwRMLW6Wvr6364YqFjW+mvnqPHbBFbEFQSPPp9bdRinea7aeP6fWv8AMLINrHx/bBFBAASfG99+n10xX6+vjjVSboI6hKifT6IxRI5jY+F8EVkq3ufC2n8DGpiAAkADp536+OJwRLW636+nl7sQb6WPXX08MThgiYfXxwxB/cfIjBEJtvfAG+1/ficMETDDb68dPr9sVJtcq20AIFzc/PzIF7db6XoSBznf5EA/qPWQE4n7/j68heCqqqKdoimKirKr51LqcpSkpHN6mqaoZvFswEpkVPyGAiJpOpxM46IWiHg5dLJdCxMbHRT60NQ8Mw684tKEKI6JPF/xC57faF+LqH4VuGucRdC8D2T9QNzafVrM2VQUnjoaEjJhBDOWu5UuNhYydxc0TDRkBkzl0pMPHpL0VOpqxLIhdRzGlt3ftAfaE1/xZZ+0/wBkxwfRjtRQjVUy2E4gZxT70S5DVFXcMtiaQ2XkZMoP/wAmzQ+V0v5arzQiHC/ConsI1LI/7saTnLE1xv5/Z3UzwYZQf/D84ZX3zMnZag8UGbUJJGpZN66qqaQ0M9HU4I+Gn0pmUC2iEdVCzCDVFxQk1PuQFHpXCRj1YiYet+CfCtd5pXTqIde3LQ62DxizpOLf+4cOXlplgcInfsOO691VjW1KftNNCnmpGPavAEMEGYkZ+yeX3EJx15G8FeTy+z87KdbdN0nK3G2M9+LGUvwxr3NysW4FEtnkVTlYyswyo6LiUMmAmGYMEIeFgYCGNN5YNy+QQjU2icOtOw35FINuf85NypZUUlSlm5WtatCpxRU4SCeb2jjbyk2JVDJa76TQgShKU8rbD7gTpcDkarVtKQBYWCQNRcDQDktSSJA+EFFOOuqSU2KZDHuouNL8xzD5UpI2UtIb5lJvqRj6b8L+H7boVq5tvbufcV3B91d1S11evUdElznTgEnS1pDWtJAgLxDr/V6t7VLqtUtpU4FKkGltOm0QYjvABmJdPmF5SSw91IT7JF0jQEWufI3Avfa5srdIvjfmmIaxRaxvYadSdSE6WudDa9wdST085T9MyVtqFU5S6nVugqWmGpOZRCWE2Ckocina2hoZ5+2im4Zb3IQEhxZBSjfykKckq3Wg7RCISBSS9EzKYSCYwEHDsMJC3XIiLcriIZSQgK5WLqfiHSmGaZU46lJ6p1arGaRAAxJbkANgZHcnvC4a4u7dxP8AutIxuIImPIHfJgjnGV4WQsEBISnYJsk36E72IN76e/XXXG+NLBzvWkFOwTa521HU2SbcxuRcnU30uf0yNqmG1lbFLyp9vnSptTcumcUjkUolvmdZqB5sqKeVSwVAlalqsUqxyDo2WSiZPMQ0NRspffeW2lpqGlkwXFuFR5glEEZ9DOus2StTsQh5xEO2lbrvI0kqxqrmvWY2TTdAM6ho2AHkACSYAMxC1lSrQqHSHNIOQ4giC7TMHPBHrvIIlfnkUOoJbNjYgXBABGwFuvLtfXexSDbXeaRMLBRZNxcflGvhbwJ1vc20+fucmpuWsOKR/peBcbbNh/5OVKTzcylKstGY5WpRJIWixUhQCXAhZCTvHT8ml5DKG6RgEq50iyJVL3FKAH5StNcxaxzEHTlNjpYDHOXfVHBpb7Nwg5ywwCGxiIzG5jvOystt6b3Q2tTMkCZdiQMEye/6DzXGeo23Qp1JGhKyRzEDUWI0AA38j1ta19l540SFaAaG6b7ixN9em2tyT1tqcZAKxkMvZioxAo+GUyh51KeSBly3PZJSQW11NCuJIWlQAJSeZJSLm5PHWoW5LDd6HqQhm+bmSn7xJoZFlKTrr/r1nTT8/LbYA3uErC9fVa0CmSCASfdgzHYScE7ZGD6wdQpMfArtloAcS12ZLZGHDAJnn13jhFP4ZSVOhSOYE6JKj1B3J0PSxJI0ueltjakh1AL9gJOlxc26m9iVGwudN7G1tLjm9U8LJlwqoqBpVCHmEpbiWm5QuMREKUVcsRC9xmCW0oSkBLzLanlBX/J7INhxZqmKhkLURTYQB7PK5S0x5tCq4UVV2LKBuQlRCgCCAElN+otqtQhp9jG0Q53lwDM/D5czpGmHgGpTyBsHiQCM7keu3M8riBU8JZxY5OpNrq1ub6WPiNjYDm8N9iahhyC57HiAdLDUHwtb4Hfc3xy3qJ2TRyHUtySJho9kuK7pyl5izBxTYUAhtl1+uO7REJPMQHHG0vJ0CidDxtqecyyHUttcgfYVqCl+l4xkkhXKB/y1xdN7G6eXmsQrRKklWxbUeYigYGSA44B0zuecxODjzW7tSxxANRuo/lw6QDpJjOTyecECTK4zz+HCe8Pdp2UCAARYhdvktXsndJUCOUkHLrwt9oVk1xCZVo7PntWIBrM3h6qF2GhMsM/ahiG/9e5D1HCS1+XyCaTOtI15cxg5bKGHnoeQV+lh6a04zMIuVVWJ1l9HzCHp3FBUM9lZUtCpPcKW4kKRIo9xKFWUlJUIWrnVpAI5VcwsNLdUjjzV84bhluASNp0KTp3clj1KVc6kKfqh0uEgaJSoKNyAoEgjmPE3QrHxBaGjd2wp3DAX2t5SMXFvWGktex4yIIkgnS7Ae0gwu66F1C66fXY+jWBY7SH0nEhlVp0gh2oxkEyRkZ5WavLjNziS+zT8ZUsoGrpxNc7ez0z9nBqCAmkjTBREvrGQOw8ngYrMKl2EqbgJHnTQkqXLGKip+DjYKT11Ti4CJS+5LpjSszkPfkyvzQoLOjL6jM1srqsk9b5d5gU7KarpCq5DFNR0on0hnMKiIl8dCPtnmSHEOcrrLyUxMNFIehYpuHiYSJhh84HhC4qaT4lMp5r2a3GFEzKZ5U1zDIgeHmtp3LCY/JzMdoKFMSeVzOOmcS/BwURERES1SYMUxDQUxiYmk4tqJpirHoGWc5+wi458yOzc4uqj7LHi1nf4flHmBWCm8magmiFtSak8yqkiELkMVJZi8tBhKAznu2lUG+HGJFXrsM/FNyyNnlVvJ+XfGPhW4Yazi1v4+0aHVXtZAvaIj/e0jAqgZcBn6R7f0bqrTTpanTQrQA0nUaNQ6TpkzDckDbyOMd9b9DqPr16/LDEJPNoAdL62IGmu50uAQd9QQoDlUCZx5R3HI38tx+q67ByDI3+HE/T5hL2ttuN/Drbz/jQjfAG99tCRob22sD59ffqTvhhgiYorQi39R131+r9NfDF8METoBawF7fX77+PgBIGpwwwRQLbjqB9fPEBIB0HQ6/DTfr+3pi3y1Om+nT44fX18cEVVG3L5qA+OLYYYImK8xB9o6Haw28evp9aC2Hu/tqPlgii4J6Dw8x47nfw8sThtp4n+fLw/vrilxzDzBA9QdcEUDl5rg6m+liN/4/jEubD1/nFgQb2INrfO9/0GJwRLW0Gw29OmITsL72H6YnD08vr3YIqpJKQTqdf1xQueAI13NvPzPW2NTQDoOgH1/n1wUbDQgnmCevW9z7tNDrrgi0vzct9fat7hy/ycXt7RG4t6n+ndXw0ubeXWLAKTY31IPrtbYePnf1xN1FRGm2wJFtrWNjv8+uCKq0gWtpe/ifC3XAG/UkpTpfpbYW6ge7F7EWtr48xP7YAKHMdLkaAeNj4239cERH5fre56e766iqygPW/7edz9dMQSooUTtptsdTvr0tp01wSnUKv0Fvhb3fPBFZQuDpewJ9LC98VSBvbwtv1SCevn1/jE8ySOa+m3XwvtviFG9k9FC9xv1tp018R1wRSfZBI8b9euJBuAfHEEXSQOmnnpb3a4qT+S3p67AjBFqYYi4vbqRf8AX+MMEU4YYYImGGGCJhhhgiYYYYImKKVbQb9fh9fvix2JvbQ/G2n1440MET9Pn5/tb34Y1EpsTcDbY6+/3YsUpJ2Gw0Gnv9+CLRwxqKRtyj11/k4gINxfb1H7HBFT32+tsLX29bD688axSLGwF7G2gwAty6WPX4HQ+Otj6gYIpAAAtqOhxOAFtPM/r+nhhgiYqokctuqgD6a4thb5fwcETDDDBEwwwwRMYwe1648ILs9eCXMnOaXR8sZzVqFv/bnIyWzFpuNRMc0Kmgo0SyYOSpz2ZlLqPlsLM60m8KspYiYGn1QLriFRrQVk8WQEkkgbbmwuSAATcWubC9wBudMdAP7QpxAHjN7TfK3gqkkS69ljwzQcJAV041ExAbiK3rOCl9YZixkO22Q027JKIapalpe+4gvMz52cwqiGnyhzoPDXTW9S6pRpvaTRpf79fcgspkENjb3nlojM5weNd1O6/CWr3gw9/uNkk8tmANoHP03XAfIhqacHPDVUPF5WbD854ouKh2byvLR6pajQqeySlqhf/GJvX8bDPL53JnOolLdYzeMi1PREYmLomWr+7tzqfwbmPaVSyUR80i5pP5jHRsbHRERHzOJen0XM46PmEY6uKiX46NUn/miIyIcdiYl95T6nXnHXPZW4oq3Z4pM4283c33zBuyuEoDLuDRQdCyiVQq5m03J5M6GIuYtuRLbEPDMzSMaWptMPH8qJRCypkf8AJDrcX4Cj4uTNqQpMpLzVgOZ1MGlJIHNzcrkHGLQkAE8iYuyEhWoAJx9f+DulMtKPtqtNpuKo1y9mKVCG+zYzeANoiTJ8l4j4jvnVXGmx7hSZiWnTqcS2SSd5J4mB2XvlMLk8KppECzEI25kQE5hoR9Kb6XfVJ1vkkKSfYilE3IOvsjl/l7ScXPwuLFM19EyyEUyJlNP9Ww7Mtgi/ZLH3qYzOVMwECt24Uy5FvNN+0lJdDnOkerZZzZ9kMRsFlbRsxg4Itd9M6hhXm5ejmVzp72KM0l0tcUtJBQ0W1LWkcwQsEqVzHmOYkmrBmTw8w/FpeZZDsswFM0XUaYGkJe8pCW335PBTOAmhgY6JvyRS4OHLEVfkVznkv6HrqNYBRptfLRNQZDJgGATBO4n4Lyq/uGvLhrDCOXOLifyn3to8szvGy3Ap/LfJ+US6GBZr6dzp1jmiGIWbvmQwkWS2A3FzSMo+RGY9wrmS+qBS5BuOA/cZzEJu9jdCnpAl6GZgIIy6XQbnsPpiZjScs+8NrPttPRswquAnUQg9UTCJimUEBLaAQi23FD0hKp1FQ8NLsvK8nTin20ttQ9TQPfquUp5bQ9EJ7tZUQgFzlUh0pStQKm20ZEcv8oeHKm3ZxAZ9Tqr8q55AQ0J+H0vJK7kOYdXGMWgreg57Jaey8iIWmVMsrhXXIaeTyGm5cfb55fDtFt46q/6lS6bSJcLm4quPu0bema9Z06A5zaLAXlrdUuMaWjJXMexqXNUgPo02tmatV/s6YMCAajiACdmgZJ816rQmR0VUDkOzKY2lo19yDiot+Gaq+g6gmLjMFDvxUS41LJdMZrOI2IQhpRTDstRcQ+8pDDSfbK292ZVS8mlHcw8FL4+DdVCfd412dUW+XH1PLT36oeAiYmDhoCHeShsFhEqYiSElMRFRLhWFaUM/lgqPdhKJjptCSVKwiEcTldAx0yi4cOFSHphOKtzFm0UIkgBazKoSRwhUvlYhEIUUo5S0bUiltwcKasq9UMyhDQg3qekb0KWgkbQj1RvNM6aJ7gw5sbd+EjmxzPUOo3hp+1Jc2kQHeydRfTqNw2dQdqBcQYIAB3Mk4GNRpMNQ0tZNTU0F7Koex2Wg492QQMnyGd1t/S1JS2ILd3JW0VHXv5DNIVSb2F+aGjiAPLUDYACyTyEprL+VqdR3MZRMU6lpTjEIuIiISKiX7KQy2ETeYQbZSXlN96pMSVBBUEgKII3Vo+gcu6kioRiYVlPpZGRjrKEOTinUpglOvuJQlPeQ1QRobaWtXIlx5bEOLgLUlJ9reOUZFRUHOI+UwMxph4QEXyPpFm4wJSUrZeiUGXtOoeUjlUOV5aBcd284hKF48w6t4ssnGpQFepQrBpLm1aVSkyA5oOkmm0OAPY8jiJ7no/hfqVcU67aFKtbuqgSyoxzi4AOOoS4txnOOFxazGyyi2IqbPrlsLEpMbErW4l2LKngXYgl9YTNlnnWQXLBPKgKUBa4I4X1xRSoPvVqkpQhN/baTFOgpIvclybLsB/1NgQRsVHGZaPybrGJed7qqktuOuOLcQxMYpt5RUoqVdTSmykkkqV3ZCVcx5weZQxsRm3lZmDT8uXFxlXVAINX/AAtlE5ixDLWUkoaI50LNwLHmUSTaxJFjc6L4uoh9C2/E2zi7S33tQc4+7sC0yPUiZ3UeseEr2iK9yLW4ZTb7x95mmAW6sB3G36CVg+nkslsBEqiuR+EdbJSVNwDkQsoJPM3yqqSD/wCEoPIppai0UAciB7IxsxUtKZbTKDeXLVTMTAOJMRLXYmbsrc5w6tx+BZg5VU7rzDJSVONLfWuHStAQpTRJZ565jNVDAOxSlzeNcLlypa4uJiQsWVooPrUeXmAuq9/C1tdu6BrGRuTFmQVTk5S9asQhfjY6pIeXSeEnEvgoR/7zETefuzhKpK/KpY0i8W9Eu02hEMR94m7UQphxPstteFts2vTHtsNOii9oMENdMOjJ2HGcCFwL/aU3kh7nOZA0kiYxIdk7bgzOeIWJitqMkLK++g4kO2UohCatm0Gtv2jzFszOQMIbWkgpPdBCErC1IuCFK4+1pJoiaN/d0U0uNimwlLka3V0FHRj/ACGyQ6AmE79YHKC+62uIsE94+tCQkZHOJKKpSa1tP4mgJzIYenkPBUG3A0aKchQ8GmxGfd4iWS96PiWVRJdVDrjHY0pUShETEspbjXuBU9TULcSX4SdQiokJW224uNebUsOFSVNK/EEwYWlST7bSi2lTYUnvGzZaOttabqtqyu+m5rnNafZkAvyAZgxEDeRJ9VuOnXXtdAe8Nc2AC4kQZEgkGMCfXzlcLaroiaf8qoSQ98pJX94glzNb0UyTYrDqW30wwKiTdIcUtFi2tHMCVcaqihZhLVusuS+DgmzdDjL8akkpA5inkcmx5LWuO6Soi17WxkYn3D3njUVPRVdUzIpnPKdhQ7NJpFwcYxHfhTDsSppU0iVmKi4Vcmci7p/GvvSWUvOfdJk4xGlPNw/qen6ygnXG5y7Kn20uXMPNJnIIpsg3Ckf88W44yALBz/lS42lbfLZbjIXr6poVC6nTaHVg4AgOB0ZGSCQQT2gjIzAXedPq1GhjnvpOZ56wTBaZBBn48HK4lxcJAuLBYiUwMY2vvWHfxWPZbQ6lRWhTTrRcUy4FBCklLlkrShfeJKEKTkyzijZnx1cI8Fm2Puf/AIsOFOFQxU8wks8CZ/WuX8sYEbD1M05CxDka7HuQjDs+bfuFtVhJasMCIFNUsNq4STimHotDypexTzsSgqUqAQJbEuqQBcFkQP3lMSkbqbCW3kpA5ApQOPZuHLOmJyGzkp2pItNPppmb97SNdwym30IiaUnymIWOinWR3kM45JItEJPkN9ylD65cqGUhDbjiVefeLOiuuLd1w2kDVoN1MIaAXNIAc1wAJMieT9YXofROpMFRtLVirDHNL5AiCHNmCCMHPdfQ87CftEHe0R4GKPq+s523NM+coYiHynz07xDUPHzapJLLmHpHXj0M0lLfcV9Tr0DOHoiGbbgDUrNTQMIhoy95hnM/Y73Op8ugGm1+t9T1x89DsLc9oHgP7XOs+Ge7UDkrxfQLdOUtCtxqHYGWVE4iMrHKSOg30u93HtwsSapyzaCrJiVzmEjFOL+7pbX9C8EEAg3Fgep31tc+F7EdLWOuPkbxT0sdM6nUaxpZQuGtuKGwGmpGtkbDQ/UAOG7byvYumXIurVpJJfS9x45EBsZEGTAE7TIU4YYY5tbJMQSBv9bfyMQT1Avyk3G3Q7X33xQ6lJOxO3hYi/x9MEV+bUJGoOtvPe5Hlrre+ux6Wwsm9wAPO23yvpc+PvxVIPs3voD4eOny+APuwRWOgudsOuFrgg7eBOh+r/qfPDTp1AOpuddfrw26YIh1AHgb+u2nyxJNz9fVvDwxGGCJhc62O+/nqPXz/YjDDBFBF9Oh38b30/fFFaFA+rDl3/X/ABjUwwRQkBIsP8eWnhicNzbx9w95Og9+IN9PXX0sf3tginDDDBFXlHNzG+4IG23j4j4fvixA1sLX3+vr9bsQb9Ohv6+XvwROUez/AO35/t0+fwpey1GxOg2FzrYe/wATboPTFxpud7m3hpew8bW6b9BidAb9Tp62B3919TvtgioRzam4tsOvmCLHfQi3TfwxI0IT0Cf3tideYEaXuTbTpYX6m3TewxBBKgbWtsbjUX8vK/6YIqX9i3nb98ao0AHhjSAJTYa+1+3ni+vMfC2nrpgigBJBFzYG4+H+cCALGyjYC1gCPK/W+nQfDE2tzWFvC1vAeHnfEe1oL9QTqOgHn/8AV64IpBJVYC3W1rEkgb/HDkGmux+Ivf6Py8CQeZRtc30I12PlfyxbBFpqNlA+CT+hwxaxKgegG+lyT772936jDBFbDDDBEwwwwRMMMMETDDDBForJ5iPC3ptgkAqsfAm3ja3841SkHUjEWSnXb4nwwRWwwwwRNbjT1Phhhta+l9vP0wwRMMMMETC4FvM2HrYn9AcMPDy/gj98ETDXp+l/3GGGCJ9fR/sLeeGGGCJhhh9dP3wReo1/WUly6oWsq/qN8Q0goilqgq2dxBKR3MopyVRc4mTt1kJHdwcG8sqUQkWuSAL4+TXSueOYmbtY8W/F1PTCQ+Zmb9SVBOWn1dwyqRzzNCpHqkmcPLXop1KYZin2I2SMwDanD3ULAQzGiEqI+jn22uaLmTvZUcblasrW3Ev5MzCh4MpPItMVmlOpNliwUr5k8qkOVelwEn8qCSFC+Pl1SOqXpflRIZWwhtaqgqybVFFmIYafK1S5pMug1dw6nuUFLTbHcp7slIbCgLi+PYv9Len0az7i5rAkOr06IEHIbpqOE8gy3yxJmMcd4pr1Wsaynp92m50mN3FoBn5+S3Sp6iqsLCphHPyaGl8IWEuxsTEQMW0yXXEtMgCViPiVOuLsW0tMuuEm45Ty35nZNooimYmAnNUzSY1CxBxLL7khpihIeKMxDK0LMK7O6yTLoaBS6pIQ48iQzlKWVLAh3udSTwNkM3n0c40p2bTJRPJYffYkBATylCUoQtKEBCSkoS2hKEDlKUpAAxyUo2KnjiGkuTuZqQlQsgx0YpKSNCAkvEJB8ACbm++PqrpTKdRsN1taGNY0ACYluJ4B9cb+nhvWhV0xVqUY3cNREk6SfiD5iPouek6qKl60qeYVHK6IjqagY2KXFS+nZhFTip4CWpcIV3UN91XKG4dk2KhCQsvZhYcq7mFZZhm2Wm+VGSFGT3MKNalkhqKjKXUFBDgj5BKpC6hCVJUFIjanjZY++rkBCQxGOOOrKUNhxVm18JKUoetES6DnMxM3pqRRRKmJ9PXo6UwMS22Ct96WBwqjJwWxc9xJoOOiw6Fq+7OtgLxyRyeziq/Kebfi9GzeezSZckRD81RzWYRNOFt9pUMVoo9+MflMe4pDhcbXUTU0hWzyRH4RCRjDa2uhq29yLZ7LNtNtYNwx5MAmIc525PpIyD6ed3f4eo5zXOloIaXNE492cnAkxz+iy0R2XlacPFPSBVJ540m3XE/V9+qOdRVRSUw8LJXFQT0nk1PS+Nfqxhh4RbEa/UU1lkRBKmEOqGlMKDJ2ZouO2spPLyRzKcOOTKvqChI2bRTsVFRMJNp+8198i3lPRDymmqOiJdZ951S0sQEXAw7ZUUtNBBQ2njbG535g5pxEK7mAZLUK2YpyNhWXZDK4WHg4mKYaYfehoSXw0NDMOuQ7SIaIdbbV3jDSGipxtAWrkbk5JZtW9QymmZHRMLPZrM3C3BSyTSxa4x9XduOOqbbah30IQ0y06t199CYdpltx6JeYYQ461zb7O76bZ1bi/r0GV2h9W4uwKbdDSAQGvfBDGjABO2IgLQXLqVS4bQt6DqrTpFKkC4hzyWySGkkk8kTjJyuZGXGVklSlgNnMioV3IJpvLlLcA5qRdiazqppetxDiUgpeMrPscq0t2XynmRSdH/gxZdNB1HLYcFCRHVnMoOXsrUbBtKISFk8Oh1wkc7aW5qSBzXWAL42qoLKmgaebCp7L6YbqCBUtqIksrYl1RTeCiWnVw7yI2ZS1DNPtOMuMGHcRAVDGxLDjbrURCocb5McvpDmJM5PTT9LSSk4KMlMaptTyZ6/GRyElHKQqEh4WKlogFcyUKbcZUt1C0tq78qQkjyTxB1a9ef8Ata1W9aXhrhUd+GaGEta51Jw0Nqw0lzZkPEFrl0fRrG2e+bsC0c1pc0MaK7tQDSAWEOI1OwcjgnZe+0UiOh1NOStukZfEIKO6ikoiI6IaKVIUVtGYxM1ZYcHKf+RtCSknmTykkHfpMLMkIC3Y+HjYqJaDr8a3Dw5AcCRcpeUClZVsSlLSknQpCgMcdKWns0Cm+alZai6jzLQuekpvZRALk3iEpSADa5Bt/Vy3xyjo+bxrzjEMJC3zRC0o9l+NUlCQnmC//MKfDaAkAmy9RYpA0x5H14vpn2/smgAw57nU3Oc0acatZJG8AggQJGZXs3hIUazWW73VWl5GkBr9JI0jIADZ7wDjHMrbicy5wJWXXQVAG6ghHOOg8FG973GpPvTjjpXcXO4Nl9MvqeLlSXBZYTUMTLm3eUKJCm0Rbba7Cw5VJ1uU6BN8c9K2/wBLQKVNxcrhXY9xoqspMSmGb09lTy4VPOoE2uhpPNpYlIN8cM8y4WAiYZx6GpiVx7bLxSI2XvxKmkuFoKDa23AYxhSSoju3+Q835edHK4u74a6jSrVqQq0Whhe0N9qKRaT7uYJBGcgugcglQ8ZdKqW1GqKNd/5JLWGrI1ATiMg7GYGeFj6rqYTiEdeffqaVxZcBKkRoh56SSCSQXpbOEpKhqUhu17WAIAO3kXltG5z0aswdUU829LFOQSmpE69LYFq76ophVS0hCSKCgX4nlW43AzNEHK4t1KQwqLmamC0777mLEQzK4lblMsL9lRW23FxLS1AhQsrYe0dSo3Gpusi3LxERmlPctqpbq2j6Qp8R8F96aKJsmeTWDiG4ppUO/Dx8AZ5CQMS2ptwuIQtlYYfaafbDTzTSsfQ3T23Fa0p1LFzG3VJofRA0hjoDIY4gu0g7ZBA3gwvnh4Y27fSruNMGAXv1mD7oktLcR67jMYWyOa/DtUFJQcVF1FVNCyuWMnkXFmJqJ9zmKiED7pC03Fxq1rAHKlKHnHVKslbqypR4RVJSuXsPBtPRdYVJNI9ReU9KqdpBqEgWwh1QQkVHPp/DRZQ6kJWvvaIP3cuIHdlxtHNkzqLiWezKnT8DVuWdNU03OEpgI5Mrl0yVTD3MVuJciZa9HPxUJDuOKbbW+mOiG2loacal0OlRfb2LrjhYnM7mBVRsJEwrsxUVQ9NttRM57pVrd3AmDejZk7DhRPcJUmYuIbPKqJc/K13/AEvqb6dCgzrVYWldzA8sDqZacNluppIIbtI3PMGTi0bh1vcVKUPqMa4ezrUmOLXN93fBIJ2ye4Cx/SDNmc5STFc0y7ZqSUuoccX3b1f1BDNvlaO7LsQ1SDNIpee7sqCFOuOISV3KFi5c2Wz2mdM5rzqKrSBllCu1HO+5jZ9LpzBijZ2qbqhktxkQZnDPSKmp4h59p5bcY1MXJtMEvNRkzlcLGLPLzWzA4Mc4KbWhFQSsyVcU26/DNziU1NKnIhpCuVTjQmUjhS6nmWkXQFAKKU8xOmOEeYFDMUNHog6ul85ejXEPOohYhgyuURfMpfduNzVbsRFzSGcACo1MFBQESvl5ERrfOXU7U0un3bTd2Rpv9oAG1wQ4PALQBOQ4TIHbcZEHrrDqdN5p0vb1GubEUzqBk6TGl0YwORjHC4XVFTUZAreUrLyIV3K0/wDLK3Kjim0uLH/CURctmsZDEqVYpKXnLaFtJWq42rqWQzSay6NjjRsxhksqKIpURJpmtDhdbV/zh1Qhy646UOsxDjzMU7zKZiHIt5cUSjkVVE/nDDTzMsjBJWHXTEvQ0jbVBIddSChKn3u9iI+MabSAGEx0dF91a4DStuMdU1RVqVLtUtQlKdUJM0jkpB9kJCLPC1uVBSQE2CUkbC2lvabKVFwqtqVJaQ5oaNIaYxO4xzzzsu46Ybmu9jqfs2lrmkOLnasacETGRxkieche95x5oVnTlFcJfEfRUa7Ls0cnangJDAVC2ypiMg6uy+mEDVVEvxjyOW6kRUidmC0uKUi0Y+BZt8oP1isic0ZPnfkplHnHIikSjNTLaiswpc2hXeJZhKwp2XT9hjmJIKmER4ZV7RIU2fQfIqja1nE/yUzJpOfR781alk4pqrJeuPLEW+w8mLYgIlDUTFIL6UqhmHAClwuKEU4grU0Uox9KDsBs0lZs9klwdzmId7+PpSiqlyujUXN2xlTXlWUDLGlqUebnMkp6Uuaf8fK/cAp5bfJ/+qXT6dGnb16YI9ldVKIBAGmnVaKjR3IBaIM7kgr3nwrXe5pbUgOfRpuJEwS2A74xnYgRMZWYv3/Xj9fPFQoa3IFiRvi2g3PxsPD4+4YqUpJ1tcnqSL/oPX/OPGPv+F2Sm3tXv6jx33xHL+XXY+Hnf63xbDEmtyJGD69p3CJhhhi5ob2+p/lEt520Pnr0PuxA9CLAD4X13P8AbE4gADYYtEZMDnz++R80U4dCf+ovvbXp88MPLoRriiKL82p3sD7umniLanr1xPTy09evXp/fCwHwA9w2wwRPTQeFtfLx22sPHrhgb9Le82/Y4a9bfG/7DBEOunp8jf8AbDXqfTyHh+uG+o1HiNR8cMETDDDBEwxUmyvyn2iAVbDy18r+WJIuRr1GnmP0uCN/D1wRTivN7RHUkC/TYnQbDY6Yt9f48fUYiwvfr/Yj98EQmxAte9/da2LAXI8yBiNN/C/u8cB5HXca6+7+2CKiNvmf0/bFid7akAqttoPXEkBIIuBZQG/iTff6HlhYXv4i3XY9On14jBFUKvzX/pv8Bfx66YobKKTtfTfbXfa5+XXGpYa6b774qbBYtoOU728zufEj6O5ESke0LkgqUPDwsQRt5gAdNdTiwuAASDbbS3xPX9sQDqRynU3vrbrrfbXS1tveLkm6QT1v5dSNumCK1wTa+ot49euGNIj/AJB01Hwtr7tPfa3lhgi1cMMMETDDDBEwwwwRMMMMETAi+h64XF7dbX92GCIfL49NxceNyNvT0w21AvfU+egHzG3puNDhhgigW3PxV0v0uPh1/iepF9rehv4dRbbXz1w6W6fxgQCNeu+/TBEsBoNh9dcMAANBiCSAbfDx1wRThiAb38QSD9ehF8TgiYYYYImIJAtfqbD1xOBAO/Q39+CJgfjfDDAEAieC36mEGCD2XX8+05zaKl/Y/wCfcHDKITPK84f5PEJCiEmEVnVRMzUlfKdUqXKWwpJJ9pBBBFln50GVUvlNUSqVQMyHK5JTEwiIVhEQ6tRiXvvLjvKyw62hSkOtpAOqlApSCACfpJ/aQKQcrDsfuKNyHYW8/SUdk1WYCE3WliQ52UAJg8EgFRSxLImYvOrA5Ut86laDHzHaOqeZSeFflcE+YVmMiO/iSzZL7p5EIDf3gf8AIGSE8qkNKT3nKtLl0kg+/f6QVqFNlR1Ya6dOu+WgT7xZSGQO5MwcwcjMLgPGVCrWYWUCWvNNnvA7aXDJjInABHmsndIyrIqm5R3UwVKXaoQlPImOam8zh4dTim7JfZlivuKVI5SktRDbL6lewsOEkjeyl6nkEPy/gNUQMncbKVMPSaiW4Jxk3vZmKKXI1u+3ehfeEm6ShXtYxs0lHd6ppZOirFQBIFybq9kjlTckc1kp5vaBOmOTlKx4bS1vYWPTm1A0B5bkEgWF73sbXx9UdIuqD2N0spsa4xGmMjTuTHrPl3K8D6x0p7S577mpVcf6XHDfyzpwfLsI3hc6oaHgqjdaiphXK4+IdR3a4uby+oH4gJJFkKdEJGRC20keyhSii4uUrAATzDyT4VIfNIKdkmaFOQrLD0OzFLVK6iCmnYghLCFNR8rl6C46pXKgoWQVqCCsX0xvU9P1EMsspcW+5ZLTLYKnXVHQBCEIUta781glpem9jcY5g5bVhW2WERCR02rCu6McdcRHMUdSk8nFI1NMWEpREwURGPMuMpp2VxLTsOGZpFQkXGzBhwREnlkTBrdjobZ39y91u+n06vRZdFpgH3iT7obwcg8nBG8riX2lZj9NQ1BS1AmAIOWneOwOdllEzG4AZNw55byLNDM7NmfR8qqGYLlkklFEZWxE4ej4xMKIxpuPqCKrKXySmoeLaQ4IaKmji3YhTEQISXRzzKYd382W2dsRRUjdkuWMvhaDTNoZyBnVSQ8UiaV7Ope+AXpY7VcTCwSpbLCGmQuDpWV069HpF4+Jinyh8bVo49uKis4FmUROcVRymQNQ8JBw8hkyIGHg2oWGYQw007MI2DjakmTzjaeaJj5zO5lMouIKn34taC2E7qZf8Q+bMsjW4+JqtiePKI7wVNTFHVMl0qtcqM7p6OeSVctlqbebUAolKkqPOOBp9M8UVLK6/wDkLLHqtV1Z1W1tzXNKh7MEGiyoxtrALZGo1fbjW0PAacCfUanSGVqT7Cpc2fs6VNtZ/sw55dpY17mEVNXvGcCMYnlcl8plVXOWUuSaXTaMg2AO/mS0LZkkAhtB9uPnsX93lMvZbbTZC42Ohm+7bTygLDqBzcoiopZAlKZ7VkqfS13QEJTsM5U8S6u4SU/fkRMrp9LSvyhyBns0dUOXkh2lXtwVmOaSsyJu1UVaUtSk2mSmISHU+1F1pJoZpuCaah4VENJ5HXMtkEsShCAVol0pg0RDo53Rzm6ueeV0NllRGXM0rzMDL2TSipnIllrLqn5nO6tWupVlpKFRjlK1BUke5ESmDilNvPxzjIhX4VMQloJUlp5/zrxW65pW7PxdiWVq76dCjaWDqFSqatQsaGipUMmm2NTnCizSwanAAEK/4f8Aw9e6qfhLxgp0WurV7m9bUazRSDXkaWtPv4IDS8lxMBcyMtoylJwUxr8PM5fKWGSszeOcgIGCDyAjkYUpRdWl9wqHLDsRbsQF2SlKmwtY3qk8bEOEsyGaS5tK1ufd0pglsF1KlKCEqdiIZxSne7tcXSOfbQgYx3yDNVufzb8SnUBLXnV2Qy21FVG1CQbAICGJdBmoFwMAxy+yhELCsthSOZQKxzDl/TlayqLlzLEsi1SAKQEOlgsRCXFcvsgxrIZmf5rqVzuRZBIC7kE48Z6/0K/oPNSrRqxUDf8AbdFanTiASSQZcTgw0c47+1+FPEfT7htOjTq0Q+kYNRgFGpUjRjMBoPBJnklexVhU1RyV15uNmBYe7sucjL7aglJNgUhC1BKdrBSQskc3InUninXOb1WwKnVwVTR0GtKSUlKkOgct1Ad082tpQBJP/IhSbKN0GwxvZWFGTt2EMzYikRkJEh1bcYyVxjS1pPJZ086YjvLk3StRH5gCAVHHC3Muk6thUOvGXFTCuezxi4RpChymy1KiYhnkBvqFtIsD+ZViTvPCHT+lVyxlZlvWeC1j2upEe8C0OaAQYIOCJwceS0HjrqPVrY1Hs9vTpOl1MtfqJYQ0tIdkAHG0zPC8BG5/yKqm5pS2Z8cul4yKadek+aNNydt+KgYxCVBMJVkgl7HdTyWOsqW2uNlcGzN4f7vDtlsqfMa1xFzayuzxlUmhqmTL5nWlITgpclVa5fzB6r6bmLagtSImGdkDkRFQTKkJ/wCJuYy2WhtfPDrR3rK1Y9brmUzWIcfS7H03LkBSkrcjazpaHKVKuAVwzc2dmC0IIBAVBuNoUCpOgKjytyumVAOcIcxoum62kdS1ll/mAcwqxRAOzGUMyODn7sylULGspncuk8RUME1Dvw0ui46Eh3GIeYRYcXFqcZZef9YrUanhM2F50gOqWl/fW1pe2dSm6rb2za2lgq0qktfbe8WtLHOdTc90Na0n3vJrWtb+IqV9T6kA27sbKtcW1cP0Vqxo6XeyqjPtXESQYDxGXEbYiY+NqCkqpgZ1FQc5mURJ45mPTJajjp/L4CIiYV3vEtTCDgXpbNfuyXAOYMTqFaeKS04040qyvZs+eObNfNR6GTFMTKgYOClbUoRJMrqqmdE08pqHccWIhcsTCTV0xbgeRDxJ++2MOwwlLLKbpc5Y17WEkTCuNVDMpQ7LyHErbnK4OIYdSCbhKIlSy6vluULaLjyXPaS2ldlYxYZm1jKGKljlUS5EMSZIKQiKbZioN58c3eOwbExbfeahVEcrLb61uXuU8jZSjHqPSrW061c293e9La66tKZZbV3Go6k1lT2ZcGgywE6RDg0ugEDBM83ZdVuBTq21sHNovc1zwQ1v5YGXAAknONhg7LZ2pa3i5g469M4mr491RcUHY6tnot8hwk25oqnlqN1/1pIH5VHnNsbUTOtJRBFao6STuaQSCSqAmFZh+Bc2UFvwjtNqaAUqwDh5SpzkT3hXZJ3Njs35xJ3AtMvpSN7sEluLpiUFK7XAKzDMsEq3us2Ubgk7HHtdMcWfDjMKbqCh+JLhcktYQM1U2/KMxMppt/t9mPSUYhkw5fgfvULGwE4hC2S7+FTGORLA6VOPQywAD0vUK7+l2wfb9Mfd6HU2m2tX0mvLXFjXkNqvptPsxDiNWqAQ2XQ07zplpVv7hjK7BTY9ua5qOkOhsAgAOEnEiY7GSVwAzArbLqYBRRlTK4RTYKAqGqGOgyFhJJJRDQLDYULj/wDZJtzX0Axw5qud0UVuWokw4UVXLdRxjqhcK0CXoAgmwsASBfQ6HmxzrqrKGicxImaxGR+ZDM+ZQX34Okq7lRo2uIaCKnXktPNQ784pmbuQyVhqIjZVOGIdSAl9ULCIJbHBvMvKDMKmH3k1DLGJS2glPfRUxlyml3/+UqHiohT6epDPORbUX0xgXlV91QY6hQcTp1Fr2ZZIaQx0QJb8sZ3ld10YdOtq34Wpdvp1gcMNapqMaY06nSQYJkA+QM42Cqqd0ozI6nhJIw7Brm0pbhHoSMiCh5biI+FeSGktpUy7y8iylYWklPMSgdPoQfZY5i/HdlHTjDiypqU59Z3y+CQonlbhn6ig5q621cGyPvMziVJtyjvXVcwF1KHzi6ygjBPhH3tD7yngkiHae7qxCx7DryGrKKuVCU+0FXJAHLj6Wn2Zij36U7IfI2PiGu7XXNd551kyeVSFOQzublWU7CPkkXUHGZC0ppWymkoINgCPlj/WKqXW7KbqbKVY3DBpZABhjpJERO89/XK9/wDBtJrdJZUfUb7MuBqEkjLDA9Ix85Cz8G56D362/nw1A64qUm41vY3136fx5Yvhj5+0EwZGwGd9gO3dehqpIGtvC+1/Xf8Avvpib6keFr+/Cw5r21tv8dPd5+OnW0gWGmwxcGAB2ARQVAW0Jvtb/PngSBbzNsAkDYfQ9cTuMVRPre/19a4YAAaDFQq5ItoLfvf9ProRWwwxUG5PgLW+d8QeCYjiUVsMMMWyIMIoJA3wBBFxsfd+uJwxRE3udQfAH2dT12uBpYAb+mGIF+vifhbT54E2tpuQPS++CKcMQDcA+QPx+vPE4IosdRe1yDfe9h4dNevx0AxPu+tdfd19dL64dB463+Jt8rf5wwRQOulvff3+XpicMVCrqI9LeZ1v6WwRW10AAN/E216dD+38U3JtuE218bnqNbee/li/19fHEAAXIG+p88EWkQrW5Fk2JGp32sTYnXe/u0F8atx7P/u1HXS19SNNvniE6p115t/cdNsTYaabaDy6YIgIIuMUWLqsN7DTx1Pu+eANgvyuR8/4xZPtAEix+B66b39cETmAFj/SkX2HgPHz/fFTdKU2I/m+vh0HQ9flflSdxcHfU3I08fTAAKA8OnTbT6vgiqRzH0SAb9SQdtBYa7aHDFgAD+vuBt87D6vhginDDDBEwwwwRMMMMES4uR4b+/DEW1J8bfLE4Iosea/Tlt774nDDBEwwwsDuAfXBEwwwwRMMMMETFUqBGpF/DXX0sLYtiD5aHxtfw+v8YIpvvodPn6YYgG/j7xb3/Xhibgbm3vwRByq0uPPr6Xt0JH+cVJKwD56662BIufcP874trr18Lf5/jEDQW6Dbx8Tfx8rfrgimw6aeXT3a9eu2GGGHfz+np990XEPj9yWf4ieCfipyTgYT7/Ncych8y6bkUJ3XfLdqWJpaZO0yllq454hNQMS1cPY3D4bULFOnx7KYiYNh9MRMYZx1h5pCm2XXIhh9pZSkWPIA6HAEKQ6HGysPtupcAULH7ZyhdKkkAgggg2sdCNSbW336b2NrY+S52zHDA9wgdpHxNZYQtOKp2i6grqNzcyyhmodULLHcvs1Ih6rZWmSDk5PwyQzaNntFNpaU4mEi6XiZetYchFso9Q/006u6xva9DDtRZWaxwBB0kB8iM40xPEbrmvEVoyvRa4lw1BzHkTgY059ZkRPzxsZQE9yua7j8Rks7cJSj2YeoSwO8IB2VJom2txYuKFwQq+pxzvyuzL4XKfUxEVFlFPqqWkg9zMq9maYBZ0NnISXQckcV4Hu5gg6EpFwTjEDTUdEFbQQgkAgAhd0lJIAF0gg+BBAV0UMciKPbj5pEiHQlaGmG1xUZEpZcWxBwjIAfiXjoAgFTbLCVcpio16FgocKiIllKvq7pPWG3Vuym+lpLwGgsaWk4adw6R2/5ZHmvD+r9FZ7RzjcVhokwKziMkbDmOREgYI2XYHorjroGkZGlvIvKfK7I+OPNAGu5Dl3Lp9XiXC2FPtwNTVhMapim3W2HYdxbjEu5m1vsc77RdYGNo4d3LmqJ7N6qn2Ydbzyo6gmcZOZ7OZrLWJlNZrNJlEKi46MmEa++27FRD77riy8twPFRJUoJISnHPIp0tJh2WErTDQ6UswrK3EkttklZW4pGjkTEqu/Evcyw484vkUG0sJb5MUK7MJtGQMslsK9GzCPiGYWDgoZsuvPxDzgQ222kKvqVFSyAQhAWpVkg27Hp/T+m0mGvSpinXc3VUq6nPqR7py5+p0GBA2gCAMBed9TsbogarqqKbSYALQIBaf8Aj2BmIJ7rIll3JcqItcOy5XM6g+dbbaFxNNqbSHSUJbbU83MXGEk81yXnQL68wBAOXThx4UeH6vIN2b1TnXOpHKZWmHdmD4bklPSplUUOaGYjqhnMumUpgYiLFlQjD0Z98iW0rXBhRQUjgRkNwfQ6pS/UtdR7s+qKXS+YTaAy/pl6GdEfFy+XREbAyR6ardSqZRk1iW0wKjAtsS6FWsJbiJ02XmFcYmc1agq+atTCeRh7lBeMvk7CVQ0okUM9YGXyaWpIblzLYQlLiUID0Q6lyJjXH4t159zR9Ss7rxHTu7DpHWLnp1SgAKly1mp4Jc1oFNpADgQDMkRxJK5KkD03qFO7r0x1KzIJNvWOKrjHLILYzk4O0crO7C5TVBJKjiIPLSsuHuRS+UTCKRIqhk+aFG1DVMbBJedbgZoueVTM4+YSqbuQqWlvuU3KqSaYdW6ywypCUOH3yUcK+Yc2mD89mFU09Uc3jltvRczfzAkM4jpg6QR3j8e9MomKinBflbLr/KhCiEg2CcYjctK7h4R1glZ1KAoa25bqJBJPKb6HUGxG43xkKoPMqFcg20hSebQXACTZITvyjlubgaEbba3HA9Y8O+J+khjqfUbO5qNaGvuanTHmu8GMuquvHklxMlrQGA7NAV636t0F1SrTubC5t6Tnuf7Nl4GUgCR7sexgYkZMnOcgnkw7kPmhT4StMDBvNtm6XGKip5xFiAdCJsCTbp3YuRqbY1ISbVdTi+4jI6QQa0KB7qJq6lGniGwL2ZcniXVqBBsG21K9q3KTbHF6scw24aJ5VchaWOZtRQ2QpQUOa5JIJSN7WTfw0I9Bg60/HZ3JpLCOlL07nkplHO2lKSgzSYwsvSqyCgGxih3abD2gAVEm+MRnQOrXFq6t1KtQqMYw1XObaENgAEgD2rgHb8/CJCwR13p1K7YzpFtdUnOexjCbqW6i5o/4DAzjBWSSss9JsaNpaiqGqz7/ADxRnkdVapIzM2ysrdaMGxBxkTL4YxMDAQboMzjodaIZLq2XVLdbeDp40/7F55ZvOriZQ1T0Sh5arR00rOSLS6sGy1qRCRkfNAsc3KoGDBGoKSQcbE5n1K1Iqsq2j5Y+6xI6YqCd0+xDXBdmYkM3jpauaTt9KEGYzCKfhTEoLgTCwCH/ALlK4aCgWWYZGzcTXqoAlbSktqIuFtKWhdikKIJQoBSdD+a6tSNEoAFzo3g+vQ6fSr9EfaUK91N024urT29R3tyyq3Uym+jEB7WtGokARmJGT1PxY67vhS6u25uqVs1lv7CjcOpNBohrHBpeHl20zBnBwd+T9ccBuf8AJJeuYRqKHi2kiyjC1jCMqbAFgeaeQ0kZNjtaJXr/AEeHBKf0bmBlfUAmcvnsop2o4JMSGpjJ8yKKg5kyzEIUzGw33mEqhLxYjGSWomFVaHi2roiOZDYbc8VW+ZCo7nS4646NdHHXC2NDskrUlIG1wgECwud8cUawqdpfNypbBHMBy8wTZRH9AUlFySTzBFxrdR5QB6D0TofiBrPZdcven3bXw72dKwNBpAiNQfc1mu+Xn5LUV7rp1e4DulWV5ZkwCK117VxJjltNhAIn3duI773zKsoKPiltVnB5T1F98ccERME1BT9PTiHKlK7xaZpSMwhZYVpcKlIcmsujU8/OXja6T+6a8POXM/lLM6hawqKWoimVPLYhkyiq5OzzG7bENO5eiXvxXMLqQBLAkBJ5X3NObH5UFRKBcKV8qjc35r3srRN1XCQdilNgTqQSTfaqOzBm8hdciZZN5jLXb6vy+NfhHB1JK2XGiR7KdlpQm2o0BHSvsHW7Q+heOoFo/wDHThjCQBiDLRJnjA2jZZ1r0a8uGAUKns3uODpLpJiJII+oJ+JC5X1llNk5K334SbZqxcudYU53iI2Bh4BY5Rqpbcay24D7V7n/ANqRfUDaOpMneFyWSCCqapOI6TfdplFRUMzIZVPKfm9WsmEcW26/MaUp+Gn1SyWDcWi0NHTmUyyEjUrbMI+/3qCr17LjOysqvqmnaerHL9OdtIRc3lkBOJe7Jn4ieMSl+LQxHvSyq5Sw1Ey+LZgnH4mFj53EREvZi2WTGrTCd8637zx5cH+VmWzyq04c55MK0px5l92c0VNI5l6fyB9tx8ffJNHhC353AqaURESwh6aQCkB6EfmsG8lUv1Nfquq4t7H8TV9q/wB8vFP2jHtaBqYXAQ0kEQXYwQACuo6T0K/tR7S+ua7WkBrHU6jGtnEGHN1AYzxniZWy8BV/Bhl/M4eaSaqZ/N5xLH0PwUdMWaoQ22+lDiEvNwzEnl6SENuOJ5XkKPtmwsCMbJ56Z1ZQZlS8S2HqBtDbf/IxFNwM6/EYGIKihKkMPSMwb8MWrCIQZgw8QFJaCSQRwpn0xCS6Pwpttaea6FuxjjqVapUCFuNkBIFrLQ2sK5uZpu3KNmahnsU2hzlZZZSQbqS0F6KuArme71WmnNYm531Fzh3fValixzG1DpJcSAzBEg5mIkAAwMZ9F3HTfC9rWuaV441qtwxrdFSrW1gZaTpA5meD+i9NzVEvg5x3EHO4CbQyQ8+5FwK3lthCdEpcZiGmn23uRBeW1EtFbYClgeygn63HZjZPvZC9nxweZURcH+HzOlsgsvFzyDLfdOMVHP5FDVNUgfQUIIilz6czFyK5kJWYlbpcHPzE/MV7Ljhcc42O0L4asjZpKHp7SU3zBl1WZlwjaVBg5XZd2rCtWI1xopTBy+bSuU/6cS9YIcjpzBQYUXItoD66DTbbTSGmkpQ22gNoS2lKEISn2UpQlI5UJSAAlKdEiw6WHyV/qn1Zl91KlRYYguuHAHAkNazBkjGqZP7z9A+FrP8AD28ulz2sDA/MEENnG3Hbv8NT3/r+4GKK3H/t1PkL/PY7a/LFiQBc6fXlivMg9dxbY+BPhvY48qA8yfVdb9/cqQNb+116gj4Xv/GJPXroRbTW/mdvUYgn2gnYEXuN+vX664flBNybC+tz8LDBFFrFJtyi1rXG9joPE69L7+uL409ynU3OumqQD0tuDpe4Nx4jXGpgiYYqSQSTokDQ+foNfG+ltsASdFWB8B4bfXXBFbDDEcwva+vncfrvgin/AB8d/wC+H+fjv/fDDFl/5j8P0CJhhhiKJhp1+j0ww2wRDsfTyO+3lr5+/FUgi9/LokePh+/u64thgiYYYYIl9/LT5X/fD9fq/wCgw8fP+AP2wufHBEJA1ONNKvaOu+3x09OuL8oPlffz8L9d8AANgB/I29fjpgiqjY+v7DF8URqkgbX2vfXY+Xlti/Ty+umCKqea55hYdNf08rfMH1NsPPx18P773OtzqcMEUFQBsT8j+trYoASEEdCb7eOLFI11Nzv6A30+vDEgACwN7fvr++CKcMMMETDDDBEwwwwRMP3wwwRMDoL/AKb/AAwwwROgPj8feOmGIJItpe9+vWxNvfbAG/xI+BIwRTiFXuVD/wDt95/S/gcThgii5IvbU9Nrft/j0wAIvc7628Pq+Jw+vo/H192CJhh7rfW+GCJha/S/zww+v4+v1wRNPr68/niALgXGth+aw6efX09dsSTb4218/wB9/wCfB439q21+nUW9PfY7bWwRRqVJAvre+nwuenywA0Fxcg6E36G/v3F9fLEk2BPgDiAbgHxwRQdTYEg731t4W38xtp78QQQki5J6HruNNyf84nm32FjbU2vv5b6Ycx3AIt46Eefodf4wRTzaa6X010sT1HU2GumpAIGuOoh9rC4EZxmxkNlvxu5fSxqLqLh1ciKLzYZYbbMa9k9WMyh3JRUKVJKXYmFoCuXWTEwoLbcPKawms4X3MPK4txzt2i+l7H2z79D8DfXr87j1SvqBpDNOiKuy3zBp6V1dQteU3OaQrClp3CiMk9QU3UMviJVOZRMoYqSX4SYQEXEQz6ErQvulqDa0EknYdLvqnTb6heU5BpO9+D+em4jU0zjad4Gd53sXNBlxQqUnGJGHHYHBb5iSvi40ZGVQ7Gw0pk8dGMOXWru0Rv3JhltS0F111anW222kFxCXlvOp9hRKG0glOMm+XOVchdkMDCVFX71SRcQWI+bwFMx0OWnoxtRVCQRmMSmNiHZbLgu6Ety5hcZFreioha2UwbELtz2pPALWvZvcYVeZEThqMjaCjXna1ySrB5C3G6vyqncbEpkbz0SWmkfj9OOtu0pVcMlASieylcZD99LplBPucL6aqF6AebU26tlaVfmQtKNSDa3IlP5dfaSd9bi+Pq/wr1q1ure3eys+azQ5jwfdyGwAM5BjfnHkvKeudLql7yGMb5NA1TLZJkSc535wNlmclFA0JLVBH+jpg+iws/ETOfAgC55rtRMI3cDRSyhCFkFxtASoAb90NTFMQkYxMZMxOpNFsL/43ZfOYxp1rmA5k3iBGXbWjRSVlSVgWKVA2xihoHOOsJKWEyuragl4HKoJhJvHtNkk790l5KNfzWKbFRUre+OcuWOe2Y9RzOXS+Lqp+NZW6hUQ5GSinY91iEaADhERMZRERBUs2bBD5X7ZA1Bv65YtuanszRrseHFoILoJPuTrABxM4nGMbLyfrtuKdvWNQEBrXEkgkHTpxg4MRE87ZCyqZd1PXEtjYJ+U1tNmw0tBaTHwcO+6Ld2eVEVDOy4pKAlsjvGlHm5SU25QPBVTlDTX425NoPMui6Rjpsp+ZRMlrOYNyKHiIx94uxbssiYdL6moV55wrRDPwSGYZZWETF1Ckso4IV9xIT+BnZpmlpqZWmXtJ/FZjBMS+FjYmOfSh77ohxiHbTCswbamwv7t93eciVvIdcU2y02j0KXVPFR8U7MY2Mei42KWXXouKeXERTjpFg45EOqW864BzI5nVrAC1ADUY6OjbexcXU3sovcG+09mA7VGnckDA+J7HZeVt6Fe1Lj8Y2q+hbPaNNEy8OnTBIcSGyDJjO+eFlCkdHT+CW3+H1RldOkm1jKc3stUqWkEglLE2qaURRIuCB3QdJuCi5tjk7QUurvmbbYlbcavmQEJlE7p+e3cJSO7CpLN5mFKJtZKLJPMAFjmF8OcnqFaeUl9eqkr5gUBR5ba/lCCU8xCSpJKbnTUE7xyKp2yG0uKSsEhOvdlQSsi1iQQDc9Ryk3ASLgYsXVncXLCx1ek9rttVIRjTAJa8H14+ZJw7/pFN8h9MkndzCQf6SYkcjHAGSsuNR5fZzz2ALsBlhmBGLhwVNuwtLzqKZKrgKCXGYRxCuYAkhKtSQbkaHXyayizpis3csETHLHMGBgE5hUS/HxkdSU/hYKDgoKp5bGR8VFRT8C2wzDtQ0Mtxbq1oSm3MtQsVDgBReZNR09EsRlMVNPqdjGwQ3EyOdTGUxDQJAVyOwETDODnCbE36Ag3AAzSdm5xZZw1dnlJ8p64rGY1tTdQ0xU70OuoHjFzSWR0kgW5rCrZj1J7+JSlmFi4ZxUWt6IeTGqMTEP/AHWDQx574wqeKOheH+qXNvS6VdWdG0rVKgmvb3NKkWhrnNH+9TqOYJcAXU5Ix2Gd4V6J0a46xYWty+8oVH3NIU3BrKjHODmOaDGlzGlw0zDiCSYK2Gzop6jJZmjmfEVVm5R8nffzDrN8yuTpmVWTVliMqGbRTLEUmTQioOBiy0pKX4Z2MeXCugsRCG3QeXidWEfla2tYgswp88hJ0Wui3VJUkpGoH4xCLCduUKSlVxqARjbjiLqZEPnlnWgO6tZu5mpCeckW/wBazzX2lKIUkpOoVzKU64pZUSOXidUVZr5lJU6d+pFgLC1jYAlWmliRppbTHW+GenPp9G6XUfeVnl1lbOLSxjWtBo0yAAGkwMDLiQAdwtX1DplSv1e+DGljG3VcA84qiJkcjc/3K30qGa0EouhuvpioBR/9aiYpIFr9EVA6oC9iOU6dNbX2TngpCJ7zucwpehRTZJj6fqCDQL811LchYKaEEnqkFCQLLUmwvsjOaqu6v/lNibH29eUg6noTub20uLm9r7XTWrFIDgS5oRZWqhe9wRbrfbXrsNRjqIaxoBrPJAGTA2A89hzjJ+M7iy6G4OaW1HaiQdhuIiTHcYyJ+ONz6zkE1lDUHHuvy+YyKaPmGgqglUT9/lLjjakh5px1pCYmFiodKw7FQkTDMxaWgp1MM8hTalcvqSy44VMtabg5vXZXmVW3IImJjKiYiW6GgnAUFlqWU2plhMwQ22lK3Yqo1zNEW44YluXS0pagV456MzXfp6ZqlsTFIEgnMXBrim4hSlQ0BM4R1CpZOQhQUhtyFdQhmKUEkvy599lxVg2UbzZ4VLEVJTEvmrCVwz0AH0zGBQoKDK2nFsTBkhtag59ziWnQh23toS4pF2yha9fc2LOps9ka9ZugkuFL3Q5vu7x5Ax3niZW3/E9R6XdWtAU6bW1Xx7c9joAkRpBmAT8hlcha94qsu2IZyGk83gJZBMIWiFgZbL34eEYQgFDbcPCwMG3CNICdG0soCUpNgB+QcHa24kpHMnHkonrriFFVk9xF/kIGw+7AAeyL3tYi2gxw4qyonCXvbJGpJ5ipV7e1rYXTfRJtqLGwubcfp1UBStai5Y6gkGx2+BNhc6DoBZJFtDUZY9Jdrpgvfj3nQTDdIBzn1E78bL0G16ZX6iwe3qHSdJ9wcHTsI9BEb4EFciK7rOgqhW8/ENlEaoFRmMLCKailqVfWIVyhMUkbFMUhdgmySCdOH+YqpbBwqXpbM4WP+9uFhDCUOMxrRSgKUX4dxBKUJStPK8hakKdPIkqsMeCndSKHMkLUduv5zY6EAnW+lyRa2mt78tuzV4Fswu0i4u6AyDpZuNhKTXFM1TnJWLRLcPQeU0ljYIVROExJafaM5maX00/ScItC0xdUTWVoiEtS9iPioPz7xb4ss6Nncl1NjHspOJqDAGGmOBJGcDbBMCF3Ph/w9UoVKUVqxaXNApuyDkZAOQJ4IGF2yPsnnAnNqByzzV48a9kyIOY50IXlbkuIxkmYf7YU3OBH1lU7F0pMNAVfWcBLZVBoCe+eaoSIiUpEBM4K3cYTsNLDwO41Itaw91tCNgBpj0XLHLShMmsvaMyqyypqWUZl9l7Tcno+jKWk7bjUtkNOSKBagJXK4Xv3HolxuEhWWmy/FPPRcU4HIqLiIiKeeeX7yVaaEE9La4+Oup39Tqd/XvHyBUcQxp3awHAzPAzk5mCvarS3ZbUGURJLGgTgTIE+eDnOSPmihce//rzePw9fXFLAE3I30JR0sBppoNfn541rfQ1+YxRXidgNuh9d/wBN7a4wVkKdDrp5H1+tumo8cOo9D+oxUeyrl3Ctb+HgNdfrx3XuVH/qCPX6tgisnYeg/TE4qBdIF7aDUYkAjck+vT5n6GCKq7kC2pv+xxAPsqV11167DT+MamKo2ABHtDmBIOn5eg9euCKU3sL3v5774mwO4HvGKJAIFtLKv49B5DF9OugwRAQdsMVSAAQDfW/1bAnodCSAnY331t5aHCB2+/sD5IrXF7XF9rXH15+muGIKbC39QJ9rTmF9xcW6ab7X3viOU9Sfy2N/E7k+Z6/riDhIwAfPy+/vsUm9xbx19LH462/XAki1he5H19aePTE4YtkEYPZEwxHy3t19+n9vDE4oiYYYgCwA8MEU/Vuv0OuA05hvdWh8vL42+OItcg3ta/vvifr6+v2wRMDsb66fWh306frthhgiqm5vzX6AAi2mvl5+734t0J6+Hx9w8P8AGGGCJhhhgim17i9rg6/36E9D+1yKkFViCUiwFtdwNyL7n+bYm9/H3gj9cMEUAWFr38zhicMETDDDof4/fce7BEwwF+up1633JP74X6eP7f56YImGI5UnUgn3kfX+cTgiYYEgakXt9eIw31ta4Bt5HbqfXfBE/bD6+OGGCJhhhgiYYYYImGGIAt+3kBsPd44IpwwwwRQQDa421Hlifq3h5dcMPn49NOuCIR49cV2Ntk2+d/74sL9d/q3wGmKLNgPW/jt6e7TfFdJMY3+/h8UUnlSDcC523PtC5HX18vdoaBaQTbXe2vQ7bg+t/hbGmpROpPKbE67DQA6WOu3h132x+Jx4JNgq56nQgk+Z6+O+/QjF0U8497AxHeO54RftLhCtOUDUkAH8xv11AJ8D69DjTU/1Otvjb5WAGtz01Hn4V2NA5vaB0IFzpYg6i1iDYgjTXrpjw78xCCq6hsNOa1hr5bHfWx6XxcbTJjGn5AbDETjzzwBwoFwB522kdx8z98krHR2r/Zq5W9ptw4TLLKoDL6Zzeo38SqfIrNJUEw/GUXWf4e43+EzOJDaouKoOsCiGldaSZqIT3jTUDOoJhU7puWvwnyzs7sls1uGjN2t8k86KRmdDZl5cz2KkFS09NEWLb0O6r7vMZdFNFcJNpBOIUtTKQT6XvRMrncpfhZhL4t9l1fJ9jV2cISNCLEW0Ve173UrQ28x5DTfGErtg+ycyj7THLpifyuKkWW/FNQstXC5cZrvQREvn0rQ85Et5dZoqgYSKmk4oh2KioqIlUdCtuTqjZtEvTiUMzCDjJ5T057fwn16r0qsyhWLvwlSowNzmkZH5d4a4naeJ5M6TqtkLqn7SmQKow4HctluCMdgRtJJXzdaVnzocaQpQ9lQTqRuCfZsbnQaq3Jub7i/OvK2rGKZpubVK46j7w2y6WUKUSoCHQQ22hPNcF2LcLTihqO7QtBAuTw8zeyczq4TM1p9lBnhRU7y+r6lopLU0p2blpxuJhn0rXBTqQzaDeelM+kM2YaL8pn8liYqUTJvmeaiU93Edx7tKsy3e5g24OMPKlpLrxQVApKjyhpYUbBSeUrUFpIKuXQ4+r/CPW7OtQcTcAvDfckzJOkg+ud4wdzGF5B4o6RXrhtNlEGm54NQ8EN06w6NgSI1Z5lbsS6qIuMjoiOi33HYqNiHol9xSjdx590uOLuDoSon2bgDSwCLY3jkdRv8A/Gkq15QQQo6jpY3sbi+mw3udbbRU/mTM0hP/AJxLgBSOZbUO4s6ABRUtj+oC5ASBdRtdO++lOZoTVvuwowTqDuFwEC4CLXuoqhyCRZJVcacoIAOuPRLRpezWLiSYnkEe6du3pvjC4i/p1KTAxtm2GggAVBwGgY04IjA3jC3Kk0+X7HtHlITqVFN7cwUTpsCRr4Wsb3A3OlNTLQoDvBe6Cfavppa6bnTxTb0t18BTuajpUjvYOTqIBHtSeWqJ0VfT7qTpym9gAn8yrJFxzsy+y9rucUOM26npah6AyYQ9DsRWb2a8ukNG5fhyLP8A5dqXzScwH3+rI2ISh0wknoqUVPOo8tutQcriXkKaxO86nY9Oph93e0aRloY2o4tfUe6AG024c8u2DWgkkYEwFzDrW9un6KXTX1HCJ0ZGkEES6IAMneMbkBbCySr1tLbJVYG1rG50AuCFA6Ab3uCR4DTLh2StRfinGvl7CuOk3pPMVxCSVXLhpWMFwD+YlLagRdVrE73Kcb8y4yOEShoyYyyQ5VyfiFmMC6qFh5vPKVpPKjLmMiWkgLjISAMhnWaNSSwOgOwzk0icrY9YSkRUlYDsM6vQpPtJs2KHqBmrcmcquGvJqdwXfmWTqgMkZVGz+Ww0U2IaKhoepq3jq1nwZj4NT0JHrajmFxDDrjIWlpYSjiPFXUbrxB0DqvSLGxI/H2dW3Ze3FZlGg01GtGot96v/AFEtb7LOnJ2J3PRegmz6nZdQu2spG2r06nsGD2tV0FuJEMG2ZcD6kY/JxIViXs+c8yw6HG1Zv5oqBS5zXSquJ7fdZIHOCPZtqeUW2PFSdVQVFSeYkgm5J1Tpy8hBUTqR1N9T0OvNeVdpNLpzHx7efnCtw/V5Czh+IioypqEoaW5TZgMx8aVLi5m1HQkLP6UmMa7EOuRK0x1J95EPrWYiOHeFxv3FuZZIZ1y5MwyQqCVKmKoWKjJllvUlP0/T+YlNohCn7wXpHAuR0BPJO2HEOtT6lZhNoEQ3IuaNSiMdbgcZfRPEVS2oWfTL+zdauo21vQ9syqytReWMYwljmw4Dj/cpsJ4ACn1PorqDq95RptuWVKr6pa33ajdRDjqGRjyc4DlYpJ3UTvKtQ5gARuSLDUb+Go0ve5Ot7jG18zqN1RNydjqQbnTRV9DcdOo2IAtfIFmMzUFNqWH5SyYYXs+1LoRTakpv7QCIclKAB7V7JT11BA4pTrMR9pTqeSEBuUlIh4VJNidUgMg2BFjc6Eo8cdq1tO4Z7SlcNexw1e6ffA93cT6dxnAnbV2Vy8Fo/BBpwYLwCACwHbb+AZwuOEZPHVBQClEp3F1XvcjQqSddbeP9RvygjeahMzzEytyUzld3FpEIHXFLeD78PCNtwjqmSvlT97gWPub/AHZLa3IBl4r76ZKDfp07zFjSpfI8hA1VyoaZSE2JJASltIF+txoLb7ja6cZkTIJUG4wtlK0rS4FBK0OIWChwKBHKtCuU3SFGwJKSN8T8Q2yqF3tyQT+QDfInnOPPO+2Vualg/qdIMfaNaWua5j5gtJDT2zPeYg74XquZkSmXTWPhmeZxgqLsMtF1JLDguhAN9Sgq7sqFg4GwsJSlYSOL87nDoWr8yTckaKPUp0BFzc/lO52BvY45JVtmLLJhIUTKNj2YaO5ClSFLV3y4ppCUvsMMtnvFpcVZxtKUFLTTiG1rARzK2aylyuzm4sc1KcybyRoyb15XVVRTjUlpyVFCUoZYQh2Mnc6mcU+3J5LJJawsRUzns2jYSUyyGAiYmJY7yGU75z4w6jbWjajxdsOpheGBwhkhuoHsO3liJXoPhizuKlNlN9CPZAU3OfID4DQC08zvPC9fykykzN4h80qMyeyfpKbV7mRX88g6epWmZMz3kRHR8a5yl6JeUtuGl8sgGQ5HzWcRrzMtk8nh4+bTJ5mBgHncfUo7Izsz8uezI4a5dQkvVLKnztr38OqjPjM5mCZafqKqkQaUNUrJohTf35mg6KD0TLqVg4hxCY19+Z1PEw7U5qKZMM8MuyL7JLKbs2aQdq+oIyS5m8UtaSdyW1xmg3BEymkZJFqhXIzLzKxMfCQ8fLqZciIRt2fzqMZhp1WEdCsvR0JKpY1ASWX5y4KoUq5U95qATzcxtfp038raEG+4v8r+K+tVuq1RbUn1PYMcS4gn/cPuxOwjaBkfE59a6XasoD2r2sa8RoAiAJb5QSRzPlC3fRFcyUWI21vofUm+pH/uN7HfGul4EggquDe6QCCNuoULa2JIGuPQoWboXyqJCr6G6vIe4k3NyfC/iMeVajgogpUE3B0JNuu2l9wPLYaaEcKaTsniPiD5/A+XfIW8DxuTG4jf6x6j9V7cHb83UJBN9iqxGgH/ANJv/fFyu97j2SdB/UBuDrbqPXx0x4BqK5iLq1tcWUBfYAXGo8baX6W0v5Ft0LsSryA5ibXFhsPcRpa2m1xaIg7RG/y/c59NlIEEYzwfX5BeQSRewSoE338bet/1328Z5Rrpvvvrv/fGgFKI0NxrqdxrfTrtpqdj7saySOUXOt7eOtzb00t+txgqpcp0NzfRPlbp0Gtx77++2KgG5J11uPL4jTpgkECx8f4wRWxAA08hb3fQxXUk2HKRb2r66g+Hlvr8cSAQVeZuPngikADQfviSAdDiDfpYHz+vG3wxAV/Sb3A9x8SPh64IpAA2GBANiRtqPLE4dP2wRMV15tvZ5d76Xvtb064jUXJ2GyRt+g6+X8Yv4eYB+IuMETDDD601xFzdRmYxG390VdbnTS6bbeOvXp18tr4t7z6aa+v9sMMWUTDD6+OGCJhhb6uMMEUG4FxqfD3/ABxI1sfIaa79d9f28NMQR9XI/TEWPh//AJK/jBFbD3Ha+3TDCw8AfZtv1ud9Dp42wRMMPgfX6/Www91vHf46+VttMES1vH3kn9cMMQBYAeGCKfr4YYkfsfmDhgijDEEgWv12xOCJiBzE6i2/UeIticVVe1huf00v8j9b4IhO4vb2QR49T/Fxvv54t5/Wn14e820oRa6huDax8wBsdfhpti17b339fHw+vHBEIuCPHEgWAHgLYYYIoOote19MB63OxPp+m+JwwRMMMMETDDDBExB2NzbzGIOqtdALFO2p6+/bw8uuIub3AGgIPMQNRc23HyPx2wRXGw6+fjiCL21It4dfXEJOmpFwNbe/wv4Ytgif5w/bFFnQC9jv+uLDYeB1FtNDqLggW03trgiKPKLEana48P28f7Y/OVAbn+9/dvofDzxZR1O+/r1xpLBuD46e/X4ep0xeZ+UfH9UX5H3yOYGw0sLEjxPx01/tfHhn3tCbgjqQSTa+99B6ix11HQ4/e8hQOoO5OhG+hHX5k6Y8NEJUQoAf1aXIG24vtYWJBtqPTS+wdomSPODEfX152VmXjefs/H79V4uJi+W6io2A6n5mw9L2A/S/qcfMCCo8/L4HcWAvqCddTt0vsDoPNzBLhQqxN/a11BtYk2GumvQWud7nHpMzbXyr0Ow08DbX4228euthm02tLwIHBg52IHy2+O+6tvJgkb7E8x/kBeAmM9U0gjmtpcWVqU2/Mq+ug9op0Pp12xnNULQlxSl6DUDm3sq+x6dLEqBCiCMeeniHLLISq1myCNbAEE6C5tpqOvmLjGzdQNPqLhsdEKCSoWPNfXf8p11va2vXfaUqbdQMAEluQO0EDfGBP02WvrFw1OBiAMGYOx/jvmQuCXHbwjcO/HJQK6RzupkvTyUsPpojMiQ/dYTMGg4l1zvyqQTt1l8REtiHf+SZ0vN2Y+n5rzKeiJc1MEwUzh+jxxodm7xA8F1QTWa9wcysmvvzzclzXpOFdXDsQalFcJD19IG1xcXR0zDZSzEPxERHU+/FNPNy2dRTnPDNd+WsGYwNPAcwAB1PKog3CrH2rfmFxfcC46HHB/NGUzSZQ8bCOMmKYfS41EMPIC2olpZKVNrQsLbcaUj2VtLSpC0FTa0qSspPofh29ubZzWtuNIDmGHOIiCJjPlJj04XNX1QQXupipnIwB/TxyeDldDKnKxSyttuNc+7rBCCoq/4SQbG6rktr5gQUuBBKQlwhPMEjkvR8/aie65X0KSrl5HEuJWkn2SAVgkAG/XS3tE7pOU3ia7NDL2vYybVDQUKMratin1RnJJYBLlFx7yiVRCIqmmVMIlTsQ4kuiKp96XwzR71TspjHFlSsQOZPCXxOZERbkYaVm88kraVkVDRKF1HKXUJCuZcZK2odUzl4CQSfxSUNJSLKRGEAKPtPTPFdxbMY2sPbUS3D2OGpp90AnvAEnuuRuum2d+0uoO9jUEe48D83umADnfsTuOwWeLs24TJRqpZrV+YUpputqkkLHeUrSNUustSQzBMMqLg46MQ+3FQsUt+PEPANuRcLEwsrUhyYFlyK+6DHFDi04qePfM/NJVUcQWVL4hqfL0NQlPMZbuVVlNQkoVZtEHQClQM5pVlhbbaEOz1D708nKUNLmEappLcOxhqkOedbyaL5nIlxmIhz3SlwLj8ujIdxsEOFQbdWhpYtZXIhghaSlRJAA5N0dxx5wU2WzLM1sx5OQAUtJqedvw/NYamFVFxTDgPLbu1Mcq9lKQFC172nTep9QHUHXdOpV0tFOldAxRPu6hTJnSTzA29Maun0+/6fTfSbb+2pPIcX0jBdJbIgeWwgRtwuVNPcc/EDT5Zh4V2EkbbQCUswOX1OSLukjQDlhJFCrRygcunKQk2vzbcuaI7Q3P0wjS3Z/Hv8xQgpcgGlIIUDcpSGAgkDoE3FtQq1zw7prtQOIOWpZQM0pbPEgJ5mappWjZ6pYBSClx6bSFyN6anvgocyvaOl985L2tWaob7qLg8r34lKQ33yqMlLSL/9+7hHWGxpsENp2FidMb1rC4Fpt7CsDBaRcw2cQdMEgcxMGNtlqLqnUe5jhQr03Aw6GuBB93P68E+mVkLyI46s0swKkakE6oJ6qYR2I7h9x6l0REIllaihSn33YJ2GbQUn2+8iE2SCoOJONh+1nj8kZVSeWdWUfJacoLPY1i1GQcVRTLNPzuIp6HlEeZlHxqpGYHlRAzdUnRBTJ1pLzsQ68iGinmhHtngdmT2pfEDUjMVByqd0ZSDKytvvqZkMC3GJv3gKhETByPhgsIKS04xBwy0KspI503GNStc0J1WM4jJ9VNSzGo53HuLdiZpOZi7MY94m5CFPxLri1IbBs002Usst2baaQkJTjRNs6VG+Zd1n0aDaZJNGlVLw8iIBkgQfMRnB3C3NrQrVaej2Tnghsy0aicTnBM4nmJ3ghZFMseNiOm8NDUfnNFomALLcPAVw8hAjVkKShuHqz2QmIJbUoNz9KBFpKQuaNRaXXpij2TMWmJFMW3ouWPNMrW2HWXGS2pl9K094lxDjSihaXUkKQ4hSkOI5Sgg6DDxG1IgjmStZIFrpT3YF/wDtzFJsd9xfYm2/sEqz6zFkkjRTcnjkPwvfhMu+/MLjoyDQrmBgYFKnTDltxwqW20+xEpbUsd2kBSb7W38UWdpVPvkt0mGU3SJkECAYEjByN++Rh3fhKs91OvZ6aDnOBe1/5SPdk5+fK5FVjDvyVUQ9FRKGWUe0px11CGgkFQt3ija6jYjlLije3djQq41VNXiWytiXqEUsm3fXWGdLghtCRzvg/wDZPIBcq1tyq3my/wCFHil4hJk1HxVOz6Wyh8pdVUddqjKflMK2ogj8MlsQyY+N52yPZkkmcZWlCVRL6Q2OTNfwo9mHlFl67KqlzIlBzarKEcbjECo4RsUXL4kd2pkQVJKL0JMlMEKPf1G7NWHojljmZfL30pQOe6z4rq3Id+HaLakP6yR7Q4H9P5oP07xK3tl063sg38RUFxWGn3GCQ0gD80SPL479sXnBH2aOfvG9OoKfKH+2uTKZilM7zVqqCfDcXDpSFx0NQkhUuFiKvmJbSIZEQ0/C07CRDgbms3h32RCxHdy4FOD3h44FaE/0lkfTSET+cMsprjM2fph4rMWun2nC82mfTdEOwqElUO4oOy2mpW3ASKVqSl2GgnI8x8xiNr8voKMljUJCsMohYaHYbaZh4VpMLCw7LSEpaah4VlKWodptCQhtlpKUIt7CEpISOWNNKjHUtewvmKQEnTmAJBuAD1uNwOt9dvHut3de69ox9c1Gb7kE7HPLhIA8s7xnq7UgNZpYGSBpDRgYEGAN4AEbgdwuYUmqhx4DlcUFgk2BvsfG+25FtBrYf0jdKTTpSwkc+pAsVqJUdBzKJB5Qs6kJsLE2F7XHHCmoWKAbUq97m5JF7Gw6mxFxoLkai21sb0SRh7/jFjYqNr3B1sT0tpa2otbr48FcNaDu0AzsYOwgfT0yZlb6g9zskZmcgxJ4+ogD12IW90umijZJWbCwub2PmnUW67knTxtb26DjFEpVzG1r25h4npuPInUC51BKcbaStp0Bu97jlSrTwN9wAOvQ72HXHvcCldhdW69BbyG5tc7X1B38zjUV2gEkEDAnzzE8QYO/2dvSJgA/8Qc8HE/XPqvd4R4KA1sSSSAq5uNN+u46AEgdRjz0O6dAkj05gL9Cdba66nSxv1vf1iBaNhfoDob3tpfQa30G4t7RtbTHn4dtVhYEkKBubaD2rga6gXToNfhjArBoEgCeSD5jgbY+O6yac57c+o+yvNNL01O53Gp0uevS9/UD4/qSpJI30Ov16i/ofHH4mgToBsskbbEH+Nsfpb0VfoLX+IP7Yx1dX6xsPQb7+/A3toLn68xiLEqBG1t9vE9beWJUCQQPrXBEA69SBfw26YnEDQD0GBUBa5323wRTiCkG976/K3h69cTiOYa+Vr6Hrt0192CKAkJ26+Pl/nAgXKrnTw8hY+JO3T9cDrYbg62G52IN7jQX1F+uxtpXW6TcE8pubgXVrsNDa5sNAPDyIpKkkEX+R93TyxYbC2o2+HlioJJIIGgvb4W8cSb3JH/XTbcXtv6+mCIEAEEE3F73t7rbW898STqkeN/dYfX1tOH19fK+CINhfqAfiML6geIP7fz8sRc31Glhre+vh7tvd54hN7qv1Nx8/wC2KaW9h9/4+5KKw8/P9dPljTKyNQPZNrG30fj/ADi1jzX6ctvff44ti25ucA/JEN+gHx/scMOpHUfvf+DhiCJhhhgiYYYYIniOhBBwAsBqTbTXw6fv0+OIBBvbpe/la9/0OJBB1GCIdDY7/wCP5GGKgHmJ6EeI6W+vPFvr6+OCJhhhgigpuQb2sf4xOGGCJ/f6+t+mKJ3ULk9NenmNT4W+eL2v4+6/7Ef5xVPX2eXXfx8eg2/fBEKbm/y6H1+XwxbqD4G/yOGGCJ18tP3v+2GKpNxrvrfx3PwxbBE16aaEfHTAaCx1OuuGH18METDDDBEOgudhivNf8ov49LfHFrA7i+FgNgBgiofaIt/SdR4aj47H6tiwFr26knUePoRibDc208fPFCFFR1IHjr5dMEUnre1yABYHU6+Z8R5a4kmwv0wsABc3Ivqfj1+tMQblQBBKSNRY8p336YIoCubTlF/P+bE9dsXvcC21hbp02ttp8+uIIFtQLedrDz+PzxHMb/lJF9/32t88EUFP5iNbm9v1Hv8AKxxokefha/Q3vfzvtbH6DfSxI91/8YryhVgbEk3Fgebe1yLba9PDE2vAAEfc/wCfs4L8LrYUdLa81gelj00APhuN+ml/Gvw1yOVNjYAgjQ3vc2+enTbU3PnFJsCCDsbcoNySRpe3y119+PzlIN7C+pFj+YeNrjQedjtrfW99p0mc/D5/t+sKjhIIXpcTBk6EXBJPNYXG4BN7W+Ow6i9vWY6VLWFEJGt7AgAWSLW0tp1PLe2h/q03QdhkFSr+0T5DmJtfe1yB11FwddbHHjHpek26knUqTzEDfa2nrawvbQaYyG1ZzsQYB5/p84z2jtnOLZpkjJif0+md1sJNZKV89kWIAuk6c2w5U77i4HnppYnG2E4psqKwhm1gSnmB0vpYi9hbW2/jp05XxUmbWQOT821gbJ5SRrpaxFrg3BFuoGPW4qnEOBV2wbg+1a99ToRYAe8a36E2ObTuy0t2IxkGSPyxA5gHbMx5Qsd9vq3GPPjbB4IPaI+IMcDqgoxb3egsj2gQV8pHMdE2sP8ArblsDbTptjZSf5UNRIcKYcKVZQsQrUnoSQeUEXuLEG/kMZLphRrLnOO5SR0HKfzXJNrC25NwL23NzqPUo2g0LCh3IsSRqggBNxtdJA0udyNwTc429DqvsSwh5Bx8dhngbxnb1yNXWsGvB93UD2HeAMTB8o39MLEBUuRaXFOpMChVwQErQtSQTexFwNTe/hoL6gY2SnnDq26pa0wCEuXCgQ2U8q0gBChax5kWTyqJJGwJtjOBGZcsO84MN7HMRZbdkkhCddU7XNgetiADbT1iKyph1m5hkqSQQeVsafmA15Ougvp11PTeW/ieuyAHmIEAH/14zPrHxwtU/olN0kMjMmBEQG/lEDMA5OxJXW5zO7O/J3NOJXG15lXS9RzBQAVOnJSIOoFpSCEByeyhcvnJ7s3KD+IDlKj05r8GswOwyybmy3Y6kKjzHoJ97mWICEjJZUkmSpRPKSxPpcZwEpPKAn8fBKRy94FELHcIicoYO6rwaLEWPK3/AE21Krt28NU2HQ2NiPGP5NQbqDaHCSALEhGu9tCkW3BHketwrGfT8WVQ4F0DmS45EtwPhHP6q0OklgGh7mjESZGNPfjBx25K6MlQdhhXral/6azxlD1k8yW59RcwgSm6rJS89L55MiFWFipmG7tW5AISDtxFdh3xJNqBgszMoolAVcKd/wBcw61D8yTyJpF5IJH9KXVWN9TvjvhvZEwSub/ygUNSbtJvzbjYgaAGxsBYAkjQ40BkZAosVQaOY3uS1tbZRPIrcG4JOvhjYN8avaAWvgbAEknGk5IcP0x81T/ptQHJYRwS1s/0wcwZmR8u66JjPYecRyucRmZOUUMge0nuE11FOr6kgOUnCKBA37zuxe/tWtj3SnewtzAeWn/U+csvYvYONU9RUbF7kEhETMZ5LRqCR7cMbGxKLkgd4b/ZKBIIEEdQdmwnW+5PIRsRfb39f3QuScuQAPuKVK11UhNgQLWPs3ULDe9iLp2IxF/jJ5yS0zE77EtO+r5n+VMdNrgQ3BIEw1o/4ggbGT28hwF1C6B7DXJ+XOsP1nPsx65dQElcuiJhK6ak5KQOZIhpFLhNkIIuhSPx9RSg2S4SAvHP7Kfs3cmsqIpqY0HlNSsgmbbYS1Om5WZlP2mleytLVRTkzKfIQsBHMlEy5VHmuB7KUdg+HydgEFAMAgaCyg2lY1G35RYXvoDbW4Pj7JC5UwTYARCpTexuGbEaeZGhBsRt0IIsMYNbxZViabWtESSDn+nfbjz8gFT/AKO5/wD5aj3RAAJ2GDjSI7jJ3WH6RcMrKFgmVtm6u8WVMgd4oruVKKwoEhZJC0BCubcqOo3tkOQqYMNlMGkKt/WgpTbS40sbi2mnidNsZMIfLGHSCRDFASAQA1yk6nW1iVG3roRr0PsUFl402SS1zA/l9kJ5B+YkkIIB0tbl1Gt7nGsr+JLioCTUMYwXD/6YzGIkYP8ACvUuiUqcFrdRxBLSSB7uxO537cyuCciyl7nkBZGhSQA2okgW5gL2FtxpcA663xvfIcvxC8nKwkEquLAk8tje6rAi1ua2ttDtc45QQlCNoUhRYJ5SkBXJzctzbUBIPKdgbC41uRt7dCUelv2UsbKBUFJH/Q6JBTckmw5bb9L76it1V1Qai/VqmRmP6cYP1jPzWyo9OiDpgiCJETtv8RgYO+FsdJ6SUgA91ygm+xN7ajlBPha9wSN7jXG5Mtp9SAPYIHsK2FhyjUHYWUBva5tYXxudA00hNrpBAKQEpSU8tumg1BG/z1GnsTMjQ2NG0giwNgnUnobD+jwOwGg0xqat4XjBI2JJIk7bDOPKM8rYstg0iYJIGB8ONuBG8du3pUFKj7Psdbnca211PvJBsb6+Z9shJaAlJCdBZQuLnrpv0NiL32TsN/Pw8rCSLpSdAbAAa+G1h06a2Go2x5huDSALgJsNhqd73vbfc2vYX6YwKlYn3jnM7zB3wNjjOZPzxlNpkcR3nc/L+3zleKh4Yg7W2Jvrf36W1Nxa99TuLY8y01awA1te/uvYDYnxG42tbH6Ew6Up0BAH9SrX9OpHw16HcY1kNn2Ta35gDtcco1Gljf8A7dSDriwXSCDvz2H7/NXgANlRKQEje5Nvfa/6C29sa6Eaelirb6PXQD98SlCRoeYK3F7m97i9vHf5HFx4D2QBY+ze48SLaf3xBVVgToLaaa3G3jb9sSDdRTbYA3v4/WvxwBGmo1AtsCR000wIGpNtbXv5bYIoKrC9utt/Xr7sBqQDY2T18iB18LYmwtsLdNBbEAG9yLaWtv53vgihBuDrext8gf3xJIKgL6jcWOvh5ab9fniwAG2mFhvbXxwRVAsrfcG3kNNAPA38tvPFBpy6DU7210Pj/bFh/Xqbi/u32N/Ly6YqNeT1P6/2wRXt7RPiPhti2KcxvcAkW00I8NdsWBuL2IJ3B6fX15kU4Yg7H0PS+ttPicQm/KL/AD9cEVsVSrmF7W1thrzHe1tPDpiwt+Uct9+g/ja31pgiYrzaFVvZAJvfe1+m4vY4E3tY31F7Hp526YtsLaW8NLdf5OCLTJuUgEi4vcX+HT6PljUwsL36+P19fHFSfaSPW/w0xbc0yTIiPjjf7+wVsMMMW0Tw+uh3+t7YD+dvX6v54YYIq8tubX81+nr/AD5YkCwA8MFGw/TS/r4dOtxgk3F/2t/P64Iqj85vrodPhb0sSD52OLAWFr38zio3OmpNr67Drfb+cWUSLWv+YA2F9P48+mCKcMSN9fA+etjbbzwwRRhhhgia9P0v+4wwwwRMMMMEVSk3JCrXt0vtp1xPML21v6H9bW/trticMETEDqPHY+F7G/nrfTz8MTipUAbddLgW67fXxw4n78/vnuitcHQW6nf0+vfhioABuB03+Gmpvr6dNcWw8u/39+qfZ+uyYgm3j69B6+GJw8frTqNPHb9cPv7/AGRVNzre43A01I8xruP2xJsSN7/mG48RY/Hb44joUgacyQBbxKfMG2p2Pl6WBJuok3BA5TrYa7X3GnW+9gcEU3+hYE28/X+97DFbC4NhoLD68+o22tbrP7/32+e37YWw+/085/TZFBFxbbbpfY/X98QEgW0Gnu+Ww+GLYYInh5G/yI/fEW1vci6SnQDQk7332+O3nicMPn95+5RQAepv7gMUUE3AIFjfYAWHnrexBOottjUxBtrfqLbgX3036X8t9/CWt3f6D+EVC1ciytNdCB52v5++3vxplmx29be1t7vPz6+Jx+jfTf19fPTzxBsNL8vwt8xb4Yan9/SQB6cfVF+FcOk7pCibEGx3VrtoLa7ag2NtMflVAtkm49oKBIKVWOl+lgdwOg9SNPMk389B8rW/Qf4xVQFlXA+N9QNDrpofD3a4mKg5mYGYiPkcznjaNk2+P1+8L1tyWoJItuSfy+OutrXtcjX3+GPxvyZsnRBUBYghPKDv/QLi46gXtewAA19vSlJSFWtcdCfDXqfHxxTuumlrlWlxrrrcDW5IsDbqfSftYjMYByJ7CT3+fETskd8jkH1n1+UFejOSFKrkoItYWCTby8NdzbXS3QY/G5TrZBCm03ItoLnbY62BT12N/UY3FLQIN7c2+hO36X8yNbYdyhQI5QDci9twL310v6je3riTbggzJzA2AgS3yxgeW2+yjob28+fL9PvhbVu0s0o2LaelzYe0NTtfU6+gPkRj8iqUbBsGTpyf0De9hrcX2va408QcbtGHRaxSkAbdfWx5kn436ba3gwyDc8l9PG97eVyTfa1zcaYufiSR+bGMSPLj4fQfCns2f8QfX74W0wpJnW7Qubm5RbpoNAd9NNiLG98UXR7CjcNJAKSCCgm+o68viDrpjdr7ogE+xY7bWHpa9j7vniRDJI1T+w0v49d/T54r+J2hw3nBb5fMCMT++BZTwNIOZ2JgyN+OPTv5bTf6NZAJLKABfW3SxtcW9x+RvpiUUjDpAu1cDayPHfTlBPvF+h0ON2fu6LJ9m997kkGwvY3OlwLjXxFrkYGGaubpAuroVa3Hhe5BsTe1zt0NhuSeY2OD6fwTA/QwHs2f8R9Vtg3S0OnlBaHQCyB0AFzcjXXxFtR54/WimGE6JaSAEpAFkjY69db6abae7G4gYaKuUIFwba3toB9a776jXGoIdPRKR7iPjpihuMQSfmOdPbMY+g2EpoZ/xEduPkvQEU60ALtpBsQPZFiTfW4J16EbEdNLY/W3Imv/AJZAJubJIJPuB8dAfAagY92DIFtE2B87ja5BtoT4g309MR3QCiAB5E9f11FvgOmmLJrZIkxPqNx84gfL0VQxvby3++OBG69YalCAfyC2u/5tPy7pvYEnzFhr4/uTLEJtZBFlaGxvsRqNDc67eAvrrjziUACxSDfe25/i2vr+tuRPhfqSb76a+B99/Ta1RULuTxPGPWPKPIhIyIER+4HlnAiZkc438U1BoTrypv5gi/zPrbQdLm2v6RDp/wCum+xI8upvbfUa+eP3WA2AHphihJO5O28/zO2/bKnuf7R9AvzhhIA0Fxubf3HroPHyxqBu1gDoAdLbX9/v8NcamGINcSYJzneN8fXdUWn3Y0vckDQk6A+lv2vpiwTbl12B+J8Njb1senni2GJoot7QVfYWt8f5xATuL6EAa9LC1+pN+uLYqskC4PUDx8cEQpAsb6JHQamw9dr4KsbpvY2uT0sD5+n1pi1tNbG99r7dPDcYiwuD1G3+NsEUgWA62AF/rxw+t9fhv78cQeOzjIy+4CeGLMvidzKlszn0hoGCl7MupaSvwkPOKwquoJjDSWl6WlkTGK+7QcRN5tHQyImNiUOIlktbi5ouGiG4NTER1IYH7X7mrLZtCzysOAaRt5ZRcwch0TCS5v1EzNCwzEd3GMQM/m2WqKXmk0g2yhLkMUQCRE3h4gwZKlt5tr068vGOq29MPptOkmQ0yIJA1ESf7ROyx6tzRpPLHvLSQC3E7keR75yfhsu83ceI+OJ/i/u0/kY6kuen2tjhMlGScjqzhuylzFr3OubzmXQs1yqzYghlxJ6Rkpgoh6cTebVzIF1vJ5p3ES3DwUphJIuLfjVxLkdGIgYSDdac7KPCZmzXme3DXkfnPmZQTWVtcZp5a0zX1QZeNRcfHf6PfqmAYnELJHoiZQEsmBi4SXRsCYxEdAwr7EQ65DuMc7SnDaubK6tGNqXFJ1JriWta4CSQMwJDoO4JAlSp16VUkU3h0NBPcd5+Y2/Rch+XU66E6i24/XAJ1B2t033v/OLfXjit9VDoOW3v+vrXGOMgHurykmwBsT5AXt8/X4YkG/iPXfFNSog7AXHTXTqNevpri400/U3PxOuCJhe5/tYfx0/nfDFTcG+nLsR1v4/p5b4ff36orHbe314YpYBSfE81z46fD4YsNdfHE2FwfDb34ItJB9oj/wBp/Ue7GriAkA3GJwRQSALny9ddMU1GiNb630OvXXbw6dcahAOh2+vTFNEqAGibXO2puep93jh3HdFKTvc3I30tb+cWxVJBKraAnzOmuuuISSVEE6a20t1GLLmgRHPfPb7OfRFfxNiANybfsTiN9j7xb++Kova5N9dN/wCSDi+IoqnQHoLEWA0+X82xVI/KbnQW1FibknobdelxjUtf/JH6YWt/kn9cEUWVffTXS3w1wKgm176kAW8bE6+6/wCguSMTiCAbX6G+CKd9v0P74Ya9Be3Tx+vHDBEwwwwRMMMMETDDDBEwwwwOx5Pbn7KJjhjx38dGSHZ6ZDTDiAz3jJt/ptmo6fpKR09TcNL5hVtX1LUUSpuHktMSyYzGVMR0XBSyGmlQzJKo+H+7SSSzSK/5O7S2vmYpRSCodASfHQE6aix0326HlSStPzdPtOXHo5xL8bDfDRSEw5squEBuYUrGKbiSqGqDOqeBleYk2W2jlR3NJsMSygYJDzZiIeaSesld4lmMaQ3t+h9N/wCqX9Kg4kUQQ+s4A+60aTHInIkDzOMrFu7j8PSLwRqwAILskgZA9ccLPo/9rM7Ptkp5spuKIs8yEuOmlcvVIbbUoBTi2xmTEOlCASpZbSXEoSpSEhQTfsyZc5gUjmtQFF5nZfzuCqWiMwKYkdZUnUEvdS/BTmnaklsNN5TMIVxK1gtxMDFsuJSVqU2SWnOVxCkJ+LbMZBP4GTSOfzGTTSCkVTpmC6cm8VBxLMvnjcnmCpbNlSmJdbQ1MEy6ObcgI1yFW4iEjW+4fBKi3jv1/ZTuO5nNDhzrTgcrONfVW3Dk+/WWWkRGxbTiJvkxWc3cXESaCaLiohtzLytImKZfAH3FuSVjSsDBci4KKbb3/iDw5bWdq26sC92hw9vqIdBMQ4Y90AZIycZzMYVjeValT2daJd+Vu2TEjTvBB/v37bgN/eL4gm3+D+wPv8BdR9kKInbcbeRv6Efta+/u4Z8e3GzlN2f3DVXHEfm7EuvyunGEyylaTgHECe5hV5M0OimKKkSF8yExk1imHHI2PfSYKRSeFmc+mN4GWRDa+Lp0qleoyjSGqrULQ0RkyQJiN54z3PZbYkNBc4gACSSfOPn5LkdmPmnlrk/S0wrfNav6Ny2o6VpKplVVdVLJ6Vp6AHKsp+9zedxkFAsFZbWlsLfu4tJSgLIKcYM85/tM/ZZZVTaNkdNZh5i55TCAeUy+7lJlrPIiQKdaUEuJgqrrVdFyCbMJUbCMksbNYF0e03ELbuvHQR4+e0S4k+0OzYmGaOflaRS6fgphMojL3K6Vx8e1lrlXIY18d3LKXkkREOsGP+5MwkPO6ojmnqiqF2GTER8UlpMLCQmSDgS+zk8enGHTUBmTWcJTvCzlfPIWEmVNTjN6Gmztd1dAR8O2+xM5DlvJ2fxhiXKaebWiMrGZ0gqMZVDxEtaj4VxDx7Wj4csLKiyt1i7cyo5v/jaQHZIJBgST5z3had19cVn6Lak1w1fnjJyBHIBxOZ9OF2NIT7WtwFORoZjMk+KGDhC6UmPNN5ZRXKgEgOiERmcl0p68txfW1zYYyH8Lfb49mDxXTyFpKls/4bLCt43ukQdI56ySYZVPx7zpCEQkrqWfBNBzaPcdIDMvlVVRc0fuO6gylQJ69tSfY/8AMxuTrdpTjZoeYVEGXCmBn2T9RSSTPRIQOVsTSXVvUEdDsF32XHBJomIQkc4hXARjrj8fPZhcYfZz1HDyjiMy8bbo6cTWIk9H5v0XHu1PlXWkYy0YnuZVUKoGWx8rm7sIxExBp+rJNTtTFmFiYtEpVAsuPpnT6T4c6iRSs7t9OsQNIc6Jd7uIcBOeBO+wVDcX9vBq021BifdmB7uR2MgztxOdvrtMPsRLDMRDPNxDD7aHWXmXEutPNOJStt1l1BUl1txKkqbcQVIcSpKkKUlQONXmAFyQBcXPltoRe416A32AJ0x81jsSu3hzX4KswKG4fOJWsZ7mFwbT+OllJwsTPouInVQ8OqZhGmGl9R0nMo6KVEjLGVORDjlWUUtcUxLJGzEzejIaFmEufk0++k1DRkLHwENHQMQxFwMbCMRUJFwriHmImFiW0Ow8Qw82oodZeZcS424gqSpCkqSqx15rqHSrnplwKNX3qbz/ALdQCA5oImCNjgS07SMrY21yy5p6mkAj8zdtOwxjznnnKwRzz7Sv2R1OTmbyGb5615CzSQzeZSOZsJyDzpiUNTCUR0RLY5tqKg6LfhYhtuLhn20vsvONvBvvGlLQUqVlw4Y+JTKbi+yMoHiMyNncxqPKvMuEm0bSM6m1Pzqlo+Oh5JUU2paYKiJFUMFL5vLy1OZJMWECMhGvvDTSIpjvIZ5l1fxy85mUNZu5sJAT7OaOYBuAALisJ5ZYNib8o5AonmSm6EkJA5fqEfZ40j/4OXBajYJpzNYAWtojPnNIpSRYdEgK0ubkm5vja9a6HQ6bY29xQfUdUquY14eQ9oBYHmB5AED4c749ndPrVn03CQJLSNyBAx5zPMdlmlxsNxL8S+TfCHk3VOfmftVPUXlbRrklaqKo2JDUFSuwC6insupqVBEmpiWTedRn3mcTaAhOWCl8QtHfd4pIbQtQ34v7VrDa9+u+3oevoMYNvtHQB7IbiaB15pjkx7JAsoIzsy/WgKChryKF02KeUFRAKwhSefsqLbm7t6LnOa2pVptMGD7zg3t57fArNqvNOlUqAAljSRMRuANwe6/dRf2iPsncw63o/L2j+IWo5vVdc1VT1F01LUZG55wjcfUFUzaDkclhFxsxy9hIKDZiZlHwzDkZFxDMJDpX3r7rTSFODN2CDtqPEEEHroRuLWNxoQdDcED44HBGx3/GlwlNkH2+JrIdJA9k8q80qWSoXHtJHKSbXBvuTcg/Y+RcpSbk+ym5JBuopClK0AF1KJJIABOwFrDc+IukUOk1LZlBznNqsLnayJxoDYEZyfhE+axLG5dcMqOqQNMflHcN3+LpPl81YjQ6HbobePXS2x6ja+MWfaN9rnwtdmM/llLM+mq5qKos126gjKfpTLiWSOfT+DktOKgGY2fzqEnFRSBuXSaIjpi1L5bGF5xExjYaPZhe8TL4xTWTOoZ9KaXkc4qSoJjDSiRU9K5jOp1No11DEDK5RKoV+PmMwjH1kIZhoKDh34l91Z5W2mlrVok4+SD2qPGvMe0B44c4OIhJjoai46aIoTKGVx7q1uSnKWi3Y6DpFtbYLjcO/PXHpnWc0gGA992ndVTJlLr5W0pVvw70ZvVrk+1Lhb0WaqpmAX7BskgB0Z524U7y4NvTYGEOqPIBB23BwPkMDkEZXey4cftLnATxK59ZW5BU/Ted1Fz7NurYCiKdqXMCmqKk9IQlSTlDrMggJtMJbXM5i4dU9mwhJHLlJgH2zNZrANvFLKVOtdiMKCtRfrrYgHXp0IFrgpuOUhV+VaCr4ocZLapoOoYRiZwM7pSqpQafn0IxHQ0ZJZ5L1TGBllUU1N4duJZhouFMVLI2Tz2UR7JQl6FiICYQji21w75+sF2QvHZLO0H4H8rM6YgpZzHkUKctc55Yp5p5yEzQouFhIKdTRruwlTUsrOBdllcShhQKoOW1HDS9xX3iEfU5m+IehUOnspXFoXuou9x7znSRpggg7HIMZnI2VqyunVXFlUtDxENGI/KDxiSZ7ZG5KydkkagpA0Bvv12Hu9b4reyhra4/p2NwdTe2vqD64sQDuL/28x0+WAAAsL7/AA9MckNhiPWZ+q2Mz/f+6G+lre+/7e/fyxQqUNwAL2vY/pqSOugv+mNT47+n19EYxxdpp2kmT3ZpcP8AHZs5jJbqetJ+5ESLKDKSDmKICfZl1a2hpx+GS+YWNVKKXkMPEMzKrKnchIqHlEuUzDsMRk7mklk8wv0aNW4qso0WGpUqFrWsAw6THvR25kfFRe9lNpdUIDQMzyNsd95hc9KsrOlKDkkfVFb1RT1IU3K2u/mVQVPOZdIJHL2Ln/ljJrNoiEgIRsgBRXERDYHOjX20jGKXOLt7OyfySmMRJqk4v6LqqcQved7BZSyOts32krQoIcZVOMtqdqSnkPhQA7tybtrWQQhK1JIHzh+OntEuK7tA69frDiJzOm89kENNouY0blRJ4uKlWVNCpin19zCUtRTLypa5MmYcIhk1JOUTGrpmpATHTdxbrbDfLrhY7ATtPOKeSyuq5DkWzlRRU9goaOllZZ9T9nLaHmMBEpQ6xEQlLOwc3zHfhIhlQimIpVFol8VDrZfho11LrSV9rS8L2NpSa/rF4KJeJFNmhrR+UwHOyYG+ON9wtQ/qNd7i20oNcCQA7Sdjp535zHwld1Wm/tL/AGQNQx6ICJ4h6updLjgaTHVVkLnnAy5JWeQLeioWgI4Q7aSfbdiENIbHtrKUpKhlhyF4teGTijk5n3D1nvlZnJL22ExEUmga0kc/j5eyooQDN5NCRipxJnErcShxmZwMG60s8i0JWCMfPVzb+y5dqVlvJnZzTEjyNzsDCHFvSTK/NKIhZ8AGyStiFzOpXLeBig4NG2YSaGMdUOVLAUoJVjDyq4VO0Kyl4zcs8kMrMts88lOMeIncM/QErgIWoKCrOXNKTEIiqqh51CKYbTQkHL25pFVHUIioujnZFDzFEfFRcF3zLsanQei1qVR1j1Nuqm0u0uc0t90A5wO0cpTvbym5ja1uCCQC4NznSMHG3I7QPNfXuC03tYjQWGntEjRKSFHmJAUfZJNkk6EpCsavEP2wfZw8KmalT5I5+cTlOZfZqUciTLqSj4qj8y51HSpNQSKW1LJy9E01RU6lrn3+STeXR7X3aNeIbiQhwIdbcbTy/wCHSkc2qFyQyxo/PfMqDzgzhkdISuAzHzMl1PQdKQVZVQ02j8Tm8FJpcyxDQkKp49xCuBiGfi2kffomHhYx9+HR81L7R2n7t2wfFQEkcqpbkWtSdgFf7B5agnlSEpAN7pAT7N9AByjGl6H0un1K/qW1Ws4NawnUwATBaBggxOofD5LOvLp1vRbUYxskjcSQDpxnnM8x8F3a3ftEPY7Mmy+MiTqI6oymz7Xc7kp5crze2u3uvvjTT9op7HZR04xJb78ns/kket8riNPifljoe9m92M3Et2oVG5k1zkTXuTFHSnK+r5RRtQwuaM6raUR0XMJ3JkTqEipM1SdB1fDRcGIcqZf+9RMHEIiG+UMLaUHRk3b+yMdoAmxXnzwjN2T/APxLnIT+YdDkwkeNuUiygLjw3Nx0boNrVNKrfPFRhaC3UCRhhkgNJH8lYVO7vaoa5tBrg4e6R5aPLO/HrvC7Rx+0T9jum/8A98eVi19soc/NCNNzlZv4a77Y5l8GvaQcGvH89mKxwl5vf7rnKdNLmuloonMSkW5MqsxP/wDT6Qqu6UpgzFyM/wBNThDgliIwwbkLyRhZUtsK6VqfsinHs4LucQfCQ0VJsbT7OV3TQpJIydRc9LE3A3GuOwh2DnZB5+9lW9xOf735hZR16xnUnKT/AEyvK+Y1jFrlSsvmsw0TT8Ybq2jaVEOl9FWy8QKpY5HECHfMQIYq9rV39p0ajbOfZ3dSrXhoDHTkw3/6j4DO+YWRRq3bntFWloZIEgQDhuDxI3jf6LsIz2oJHTEqmM9qObSyRSKTwj0fNZ3OY+ElcplkFDpKnoqYzKNeYg4GHaSOZb0U822gAlRCUrOMM2e/2hfso8hJ3G0vMeI0ZnVDLnXGIyByQoyrM0Zcw+2QlTSawp+WqoR91JUErRDVQ8puyg6GyLHqP/aMO07zQ4l+LHMvhCpCpZxTfDTw7VMKHmdJyuYvw0uzTzOp4NPVNVdXsQym0TaX07O3XKcpOSR7kVAS8yJ+oWmER81QuG4gdmH2KfE/2oEFUtcURPKUypyUpGbimppmlXgmMxZndUNw8NMIqmqOpaUNmZT+Nk0FHQUdOIuOipFI4BMbBw7M1i5h30M3nWnh+2pWTb7qty+lTeGuYxpAI1aSAXEZPIABPbUrVTqFV1Q0LenqIdBc6SABG0YEk8fyV3CZd9qk7LGNivu8bFcREmhy4lP4jHZMxEXDBKjcPKh5RUExmXIBdRSmXrdsCORSvZOSDhs7X7s3OLJ6EluTXFjllHVNGuNsQ1E1tGTHK2t34hwlKYaBpfMyApSbzdwEEkyaGmLKjcMvObY6vc5+x85lJlt5Fxu0VFzdLKimHm2S9QQMucdAUUpRGQmYM2iWWiRq8Jc+6EJKhDqJCBirz0+zY9qFk9mbSdHSbLWnM6qRrKqZDTcDmtlJUiJrSsjVP5tBSxE1reSVGxTVaUvKpIzE/iE+mztOxcigJXCvvInMUpjkco7pvh24YW0OoPpVBDgHlxyA3uwE+moHYiMlV/E3zS3XSaWkgQBByG5GNhzvuM8n6MnE/wAZPDNwY0bIMwuJ3NqQ5SUbU9RN0lIp7PIGfTCGmNROyqYTpqVstU/KZxFJfclkqmMSlbkOhgiFcb70OcqFcEV9v32QSFFJ42aDURvyUfmy4AfDmay+Wm/W3NfXwIxww7RvsYs689+zJ4UOB7h+zMl1b1vw+VlSNQzyu+IKvqtU7UsLKsvqypyevwM1MmrGOgGVzupYX/TVLtwcPJ6fp5piVQTqGYBtL3XVb+yj9puL81R8KyADc2zVrRSb2SlQ5RlOSCTprZVhe1rHGLY9P6PXpF1z1E0agqOYGhzWBzGuaA8BwnIkwB281crXN4xw9nQD2uAcSWyQTpkbcSR8POF29j2/vZA21416F3F7UXm6fA//AMu/MX+G+I/+P72QQuf/ABsUILkn2qMzdF7gHQf7ejw2tffTx+cxx/8AZ35+9m/mlS2UnEHGUBHVTWVEM5gSaIy5qOa1JJvwKJnk3p9tEZGTan6afYjkx0kjeeHbgnm+6MO93474tI8h2fnZqcQ/aV1jmHQvDtHZcwc/y1peW1dUDeY1TzemIN2UTSbJkrBlcRLKYqX7xEIjCgRDT7cKlppwPJdWElvG9d4Y6S22F06/rC3LGvFWaZpwS0TqAmZjnkeUYjeoXbqns/ZN1kxpjkBpx5HPzGN19Fn/AOP92QB//faoEbi5o7NoWPTfL0Wt1JIsL30vjn7wv8WnD3xnZbv5t8M+ZMvzUy8hqkmdJO1RKpRUcog/9QSdiBipjL0Q9TSeSTBwsQ0zgHu/bhFQrqIpBZfWpLqW/n+J+yi9qA6EpVUHCyykk2S5mzWZVa9iD3OUzoJGqTyKWNgkkcuO372G3AZnZ2dfBpM8gs+YihI2tYnOSuK8bi8up9Majp1cmqOVUnCQQVMJtT1MRomH3mTx33iHVLFJQEpdMU8p/mPPdUsuk2tu19jfm5qlzRpc5hIBAkwAI5G/fErOt611UqBtekKbdJIIECfd3Mefyz68+eKzhB4e+NbLaCyh4laFiMxMvIKrpJXLVNt1hW1HsLqSnWY+GlUXGRdDVFTUyjoWHh5nHBUpi416VxTryHomFVEMQ77GHziX7YPsgOAWezXgBrSio+ZyTJmRSSlZrlnl9kxKq/yyphiYSaHmzNIxbMymDMqjZrDy2ZMPVDBxTMZENR8Y41N3npqIxDOVnj24uaT4GOE/OPiZqxqGmX+31MOuUrTETGJgjWVfTdf4TRNItvX79CZ5UURBQ8XEQzTzkDLBHzDu1ohHCn5GFZTvMbPbNKqqxm6JvXGaOalXVJV07Esgo2aTaoKpqSYx9Qz56Cl0C3Fxz5cjYiOjjDw7LqmIVpYFmWfZyvD3Sj1AVH16lWna0nQzQ8tDqkNLoJ7CBJEe98FbvK4oua1rabqrokkce77syY2ieeIX0OOzbnvYHcfOeFeT/hO4MsqpBnDlZLZPXcXD15kHTVJckJNJjFwTVS0XTDkXOaXS7IpsxAtR8bLJLBRMnjJtI3odwLiG1M9llAASiwTYISkcqQlIAAtYJKuUWGg5lADqrQn5AHZpcZ1QcBfGhkrxHymPmEPTlP1FDSLM+WwRceRVOUtTutSyvZPEQTaj9/EPKFqqORNXIbqKRySYNnvYTlH13aRq2nK8pem62o6dS6pKSq+RSmpqbqGURLcZLJ3Ip7BMTOUzWXxKFcsRBzGWxLEbCvAjvId1twBXMUpseIbGpY3NP/cq1bd7AKL6j3PM4EaiTjMSe0jlTsaralN3usa8E6gwATgeg3HHqvZMVKwFcp3J0ABJOlzcAHl0II5rEggpuLnBRKbG4GwuQVWJIAJSkEkX0JsQm4KrJ5iOrP26Xb4p4KIuc8KPCRESeoOJ6KlQRmBmDFNtzKQZAtzNDL0FBMyuJh35bUWZ0wlb6I6FlUyDsopWGi5RNJ/ATN6Ih5KvU2dpX6hXbb27CXu3JwGiQC5x4AzmcHgmFk1qrKLC9xgDjuZGB3Jn5+i7Fuc3Erw98OsnRUGfOdeV+T8neQtyFjMxa3p2kkRqUCx/Dm51MYOImCwq6e5g2n3uYWKPaQF4rK1+0Z9j/RUa/L3eKl2qIiGfch3V0Nk/nbVkCVNqKFLh5rKcvXZVGsc4IREQUbEMOAczTi0kE/NhZVxV8e+f6g0c3uKHiIzHfiHlOFyeV5W81h4JKlrLjz7sU9K6bkzGosYGl6bgWw28qWwrAUM1mUv2W/tS8x5CzOqplmQmSzkUwmJaprM3NWJjajQOW5ZiITKyksyZPBPA+y6y7OkuMu8yXkNvJW0nqP8A470y0az/AKj1HRV0+81rmtAPukgT2MmSMxkLVm9u6hBo0BpkZImZ05B7fODEScLuSZT9v72SecE1hZHIuMGkKVmUWQGWs1qVzByggitSrIQ5P8x6TpynWVn/AKPTdC0khK0oUtIOWWj64ozMOQS+q6AqynK3pebNB+WVHSU8ldRSGYMkcxcgZxKYqLl8WkJupZh4lzkAsqy1ICvlycVH2f7tP+FyUTSqJ9kdDZv0JKIWJi5hWGQlQN5lQ0ul8MkuPRcZSj0BI8xmIFpoLiX4pmin4OHhm3X4l5ru3GxxL4GO0Z4r+zyrpqseHjMucSORRc4gplW+VM7i42Z5WZgIhHV9/CVVR7kQmXficXCqegYeqJaJdVkrVEufh85he8dbVV/hi0uaXtelX7a72jNNzmkmA3BLdhnIiTjIzE2X9alpbc0NMuA16dwS2D6YAHzkyF9fj6+vr9sMY1+zB7TLJvtOMhEZqZdIbpWvKUiYWRZxZSR80hJhUGXtSRKIn7kvvWu4XMaUqVuDi5jSVSGEhmZrCwsfLlw8PPZROZfLslH+fdc26Dw/uTfHI1adSjUfSqMLKjHFrmuxBifkeCJW0Y9r2h7TLSAZGRmOfjHqq8vtcwOp0sdrW19+gt4Yt6bdL7+/E9D6j9/4xVJ1TfYi5+X8nFsGfv4/oVJCbAnwt1tuQP3wJtp1P11tiFeHTdXkL/2PifjieXUKJ8R8jpp6jU+PXpVFOGINxa1tTb3a36/XTCx8SPh7uh+t8EU4YYYsv/Mfh+gRMMMQRfqR6G2IopwwwwRMMT0PjcfDX97YYIowwwwRMMMMETDDDBEwO2hAPS/iTYfR0wwOgOxtvsRsTrvp6A74qJkYmI25J2G3ltkny5oT9/z5eaxodrZxxy3gC4H83c7mItKMx42VmgMlpaA0t+ZZr1jDxUBTUQiHePI7B0w2iNrGc86FgSeQxSCh1b7DD/yqcpsqc0+KXPmhMo6KVE1jm5npmJLaclUbP5lERbk5qysZwpqKqSopw8IqMXBiIioqfVFOXREPMwENM5g+t/unFDP59pX7Qg8UvGS5w30DVENN8leEpcxpOITLHg7LKgzymRbbzPmT8QhSkzFVGfdYLLqBDgLcrnMmrJUMt2HnMV952U7AbPPgH4VOKqp+JrjYzbgqBmmXVFvyTImRP0RmRWbr9YVqmJlVR1mF0RSNTQkA5T9IMzSn4ATZ6FfiXariIqDSHZW06j0ro9o/pXRKt0yiX3lywPaxol8EAM9BzG8zMQtFXqNubrRqAp03DVmMjTnMjeI4I+S7K3a+9jPRD/ZEZd5fcP8AR0tmGafAHRCKspGYwMEzL59XNJQctMZxAQq+QPKjJrV7bUbmi1LVB9+YVfT0JAS4oiZiCvpE9n1xd1fwLcXuS3EtScVM/utD1S0zX0llzykGr8sp42qW17S8XDKWG4v79IX4iLljUQpLMNPoOTzZlwOy9gn6K0V9og7G+ZQUTAxfFpLoqEjIaIhouGismM/XGXoZ1taIht9lzKotOMrYUtK0qBSpBICS5YH5yHG9JeHSnOKvOeG4S8w4LMzh2mtYzCqcqKhgpRU8h/DaYqd38bhaPjJbWMlkNQNxVDvRkRTH3iNlyDHsSqCmbbyhEq5oeHxdXNG7s+oUaoFYF7TVa4t98jWC44ETIEQJMCAl6GNfSr0Ht1M0yWkZAAg4yT3mSPVfX9oGuqTzPoWj8yaEncHUVF13TMkq+kp9L3CYKc0/P5bDzaUR8KpXtpajICIZdShYDrQWUOoS4haR0F/tavElWNVcW+TPC8mYRcLl5lDlRL8ynJM0sIg5xX+Zk5msGudRyUhSX3ZJStMS+Bkve8xgvx2fch5pi6Bk9+yudoK9m9kbWXAzmNUMHFV1w8pNWZNsRKktTSa5Iz+ZATSSJKllyZnL6t5m42mJCVfdaarCl5ShssyRLicN/wBrDy1ndLdolQGY0VBPpp/NTh3pFMkmCgv7tGTagamqen5/ANPKAQqKlsNMaciYthBC2WZtBrUk96lTmq6LYtsvEL7au1oNJr/Yux7wJboO0ZZqJiYIGRCybyqatlTeyYcW6y38xnTz6xvtGDstmvs2PBTRXF5x+f6vzQgmZ1QXC7R7WcX+mIyDZjJVVleLn8tp7L+XTtmIStpyUSqPjJpWZhy2v71M6UlMA+FwUXGJx9M8AABIACRfQWtrc6Cwte5+JN9bY+dh9lB4gaPyx46MzMnKrmkNKIviKydXK6BdiylpE7rnLmet1TDU2y4pQKY6Y0dGVvNYRlfKHnKffhyVRLsI259FE/IdQbixvY3Gg5rGwNjYEkW1OH4sfWd1N4eXaWtZ7MEw3TpaSQNvzSDicQVd6aGtt2loEzBxJJaG/mPr3jf0igSB1JF7jy2G19TbTzHTGwHFNw1ZWcX2QeZvDtnNJEzqgszaXmFPzNTQYRN5JFPMOKlNUU3GvtupllUUvMvu87p2ZBCxBTWDhnVtvNB1lzkDjw9QTuUU1IZ1UNQTSAkkikcqmE4nU5msUzAyyUymWwrsbMplMY2JW1DQcDAwbL0TFxcQ61DwzDS3nnW20KWnnaLqjKtJ9MkPa9uktkEODmwRGSREgZjbM4z3Q5pDjLSCCCIAG3ygc/pv8XrNXLKZ5R5pZkZUVA43EznLKv60y9msQ00WmYmZ0bU0yp2OfabFnBDRERLFPNgqDiWi0SVlNz9PD7PxxAVHxBdlhw+TCrpnEzqp8rkVJkZHTOLWXYuKlmWU1Eso772+STERENQL9LwL0S5Z6Jehy693jrzrp+arxS5pyrObiY4hM35IpTkjzPzszUzAkqlNBhxcnq6up7PZQpTV0hDi5dHQzrifZKCpSVAKSUj6I32ZOhJnR3ZUUHOZky/DozKzZzcr6Uh5KkB6TrnMJR8JEMFVgWIlyjYl9paQUONuJKCC4L+jeKWtd0qwe8AVtdKZGdTmS75wJjdaPp2LioGA6A17nEzwRgfTJ2+i+bDm8pTmamZy783fZjV0sEkknmqqaqKjfUaXVc6m6jurH1CPs7rnedjlwYqJHsyPN5B8i3n9muggEEi10bgkHQ3tj5f+bsO7DZp5lQ76S2+xmJXbL7SxyuNvN1PN21IdSo3bLbiS2oEaKsCBe5+nl9nQjYGM7HPg/RARbMWYGFzogIxLSrrhY5niCzUVEwjyBzKQ9Dh5JWg68qm1/lcChb8Wj/8AFWJEZqUtUbT7JoO0Ygmczsq9Od/3FV2BLHdzu4R8iM85WbLZar9B8Bp7/HSw+eMGP2jklPZGcSI1u5N8mE9bW/3poFRO1jYDa/zGmc/luSbkXHKRoRfYkWJFrjxB6G9r4wVfaP46EhOyP4iBFvtsGJn2TMHCJcNjERjmcVDutsN3/rU2w64SAeRppxxVkoOOK6T73UrOZzcUp5xraOZ9Mra3J/7et30HHxB3HO2NjkT3+dTwLsKc42uEZAUQpfE5kNa6tSRmnS3X1tr4+AGPsVjQHYG1wCRbT2SUgG9hy9ARbXwx8ejgCZMRxz8HKEgrWrihyFTygaqP+6VLEjSwuLWHQ6DXbH196nqmn6Ipaf1lV06l9PUtSMimdS1LP5xENwMrk8gkUvdmE3nEyinVJZhICXQTERHRkQtYQyw24tRsnmx1HjYF93YU2iT7M4G8lzA04ncgYAnPGy13Sf8AxVi4gMETwcNpmAdsnGdpXWd+1Acf6+HThMlnCfl9UsTKs2+K/wC/y6pFytfJGyLIWSOtN10YiLSs/cDmFHOy2gYVpTaVTSRxlbJh3mlQa79N7sfOBaP7QXjqypydjJSI7KumooZoZ7Rb6nGoWHyroyNgYuZSdZbAceiq1nUTI6HhoVtaHUtT+Pm7ZUmTRKR6N2nnHNUfaEcZObPERM4iIbo+PmQpLKCQL5mm6aylpaJiIKjIMsrUe6mU1aci6qnxHIVVBUU4LYbh1tstdi/7OlxadlzwE8PWYNeZ9cTlGUXxPZ41N3NUSaa07mHGRFH5bUW/FwVF0u1GSmi5hKDETOZRk/rKbuS+ZxRiGptT0FHcrsibSnZNt63R/D/sqFGpUurqDULGOcW1HhoGzZIaw7e7t5wscVW3V4HvcG02OETABgtGJPcHjI3yBOp9q17PNFKTjKrjyylo2WyulZjKpHkfnozT0C1ANS6cydhyHykq+Pl8I2plMFGU8w/l1FTIMhuATIKAk6EExUIcY9fs1XaAHhN412Mia8qRyWZK8WqpVQcYmLeWqTyHOGFU81lXUJSoKTArnkdHRVAzGMb5GFrn0li5otEBJGVtdprjJ7WPsUuLvhnzp4acweM+hzIM2qHnVMCYM0Hm1HO0/PXYUxNM1TAqRl040Y+maiYls8g1cwSuJgUMLUlpbwx81SJ+90lU75kNQNvxdNVI8JHVtPOxUNDxUXIZmpuXVJT8TENQ0a1DxETBsTiSRDrMPEiHdgolbcOV+xHo1K46l0ev06+o1W1KZIp1KzS0OJ95unUMvaRBGwEdyFK4fTo3Ta1FzYdGoNM4bpkwM5n1xxhfbDAuL8x3OiSTa+vLr7R5b8uoB02xIvZN73ub69LHfXXp44xgdkBx3wvaD8C+U2dUzmMtjM15PAIy7zzl8uDcMmCzYpOEgoefR6ZaFqVK4CsISIllayeEWeSHl1RwsH3q3od8pygkgEeJ8BoT5aDQC24Hwx53cUKltWqUag0vpPLXaufegRtuNt9it2x7Xsa9pGktBwZ4z975VVnlTzEGwClHfUJSTrYE8pNkmwvrcEEY+XN9oN4rK24l+0xzyp+bzZ9VA8Oc3eyJyykAPLByiFpQoFcTYsAlpc0qWvHJ3GRcaSXXZXCSOVrIh5XDuD6i7l+VWtxY2HmdBa26jewuQLFRINgMfJK7XmkZnRXaecdMhmrTiI5ziLzDqVlt5JSp2WV1MUV7In7KGjMRI6mlrzKzyhxl1l1A7t1JPV+C6TKl9Ve7T7SnRc5h3AIewyARuRIB5Ewtb1d2mkyJ0lwyIn+knOwAn7Cz1fZbezcy1zqneY/HPnVTUorOGynq+Gy7yQpSoJczMpJAV6zLJfUlT5ix8HHMvQsZN6al82p6XUbdh1EvmkZOZu8mHmMFIomF75SG0gFP5SdVhOgO2tr2N9CQAE8w5glI0HU9+yUZx0xP+D3PjI5uPhE1llnnk5WkXJy82mPNJ5kUhT8LKJwhhxSXnoNycUXP4N55tLgaehkBwnvocL7Yp9sDptca3BIuRY2ItfqAettRjW+I6td/Vbltdznezc1tJjp0tp6QWwJ5wT3J+WVYta23puYBtvEAmBvyTPxG43WiCm+uoJTy83LuTfc7HzJ1O58fHREnksRM4GevSmVvTqWwsdL5fOHYGDcmsBBTJcKuZQUFMFNmMhYSOcgYRUfDMutsxLkHDmIQtUO3yeTUkDlGu19L/m2G26bgEg6EXBFjj573bq9p7x98Ovad8QmUOR3FPmtlnlnS0vyfckFG0zNYKFksqenWTNBTuarhWoiXRTyVTCazCMmEVzPq7yKinlgISsITi9J6dX6nX/D0HspuDXOJJhukFrYJAMEz2nhTuLinbND6jS4a9MtztGM/qDv6mPoQNgWSLjQhRAA6hIBsLaEnkCtRoqxPKq3y4vtGylO9sRxXi/5ITI1sg6CzeQOWRuCdzZd/Mknzx29vs0/FlxC8XvBhnLXPEjmtVOb9aUtxK1DR8pqOqnIJ2YwVMQ2VeVM7hpO2qCgoNKoRmZzqZxyS8lyIMTGPuuPL5+UdQP7RKvve2K4txsWxkqk9dsgMsAFaEWuBe3S25GOm8L2tS065c2zy11SjTOpzRidVMy2Ylvrkx8TgdRqitZsqNEBxbAPAlvqskP2cftPeCbgFyV4j6Q4pM14nL2o8w82aaqSmYBiha/qkTCRSqi4aVvRio2j6ancNDBEyLzCIaLfadUUuENFtYUrsip+0XdkMvX/xQRLabE3Xk3nlzdNOX/bkpHUkhR0A/MCeXot9mz2NPEV2ntHZl1vklXeUdISnK6rZVR0/YzGmdUy+PiY6bSNM+h4mXNSClJ+xEQqIQqQ4XoiHdDqClDKkjmxk4P2STjv5Nc8+F9ZuTy/jOZl7dCVDLfbUEJItrrbbGb1Ww8P1OoVqlzfOp3Jc32jWuA0nQwAAaTEiN+T2hW6FS8bRYKdEOYI96TJkN8xgDGByV2dWftFPZCLWhv8A8VK0qJIu5k3niB6rUcuFAJJJGgFrDYYy15NZv5dcQWVdDZ0ZSVAmqstMzKdgKpoyo0y6aSr8akEzQtcFHCWTuEgpvCd4AVBiZwcO8gf+qyCdeg+z9kp480kf/bdwvJSFBRCZ7mSojboMuClWt7Wsfhju5cAnD/VPCnwX8NfDjW8zkc4rHJzKelqEqGaU05FvyCNm0lgUMxUTKXo6FgI52BdeS8YdcXBQby0XK2GlcqTy/VrXpVBtJ/T7p1eoXQ8F2qB7pEDS2MjGccrPtal08xcUtDfd2JMyGx8Y+HnBkfNK7cbhezK4Ze0l4kWa2p+OgaYzizDqbOvK6plpLspq6j8wZtE1A7EQEaE90Y2np5GzWnJ3AOufe5fFSxD7iTAR0uiYjfrshu3RzW7MKnZ/k/Ncr5VnVkBVlZO1vGU4mdikK4o+fTCXS6UTmaUpP1yibS2ZQk1g5JK3Ium57LmGHY6GRGQE/lbsVM1xn0UOLbgp4Z+OLLV3KriWyskOYtMtu/fJNFxiX4KqKTmiUKbE2o+qZa7CT6m48tqDUQ7LI1luYQyTBzRqPg1uMHq6cSn2RLLycTCInPCXxS1FQ0M4XFIofO+lYSuoND6krIYga6pKIpOaQEElZDcO3H0tUkQ0wkKdjHXUqW5vbXr3TL2yp2HU2FmloaHZLTpDQCIy0457SMb4la0uKVQ1rdwkkw3ckSPdIzPbOcAhZDsnftPnZdZmfhcNWVUZs5Fx8aWW4pGaGWsbFyiBfWhJWYioMt46vYBqESsq/wDOxH3JrkHePoYSSlOcHJniGyI4i6YYrfIjN7LfOClogBQnWXlXSOq4WGU5yKVDzD8Hjox+VxjZWkRMHMm4WLYUlX3lllSFpT83viV+zb9p/wAP0qj6lkOXdI8QtMS5pyJi4rIyrRO6khodoOEuChqmgqSquZupQkuKhaYldQxQSq6QeUk4asrc7eIDhKzXhq3ykrzMTJDNig5y5DRMRIZnNKYncvmcoi1w8fT9SyZammJnB9+y9ATymalgI+XR7CoqWzeXxTS1tCj/AA/0u8Z7Tpd6S9pPuPcHgEgACBDmieTJwdyZRl/Xpui4pQDAmJMy2cxnEkE7k+q+z9cA6WBJ0NxqD7QNx0Ub+1oFEGxJvgQb6EeG/wC+2ugvba9txjC12GnaeTLtOOEmJrPMODkcqz+yeqdOW+c0DT0OuClE6inJTBzeksw5XLFuxJlEvrWUOxSH5WmKcagqkkFSMwjTEtTAIOaO5BsdjrfYD6069Rjjq1vUtqr6FZoFWi4sd6iDPG8yJ7zytox7Hta5uxGrkdjwQDGB8THC+ez9rjCFcdHD6LBJHC3LlLISLn/7UsxCknxPKbDqNb767l/ZAmweI7i/esP+PJagk+1cEd7XkQRYgde6URca8ltrX2x+1yKH/jq4fxf2l8LUtBOu3+6eYIBIsdSFE9Py26jG6X2PwlXEFxjqte2TWWwvp/VXE0Vf4aab76bY9AqCfBzJMEUqY371aQIHlmQP8LSUnE9Sc2fzOETxIYIxBzJG+wC77CRpcAAeliLm42Ph8LbdBXQEAWFibkggBPKRuBZP9KUk6X5QLkhJspKb7i4OlyRcgdTrpYddr38TjHv2n/HBTXZ9cGubfEPNn4dyrJdKRSuU0geSHXanzXqtDkso6XIhrpMRL5fFqeqOolpKUw1MyOcxYWUtJS753QpPr1adGm3U+o9jG4GS6JJgRABE9hncLeufoY5xhrQDnkkQIyCAD9SREbLpwfai+0KRndxJU/wU0BN33MuOGWJE2zKWxEckBUOec6gFJXChtCymLay5pSLRKG3nyh1ioajqiDXCpXL2Hse8fZWOA9nNnOPNHjdzClcLGUPkzAx+VGWUsj4NuKh55mVXUicZrOdJ79taBDUfQszTJkskOCPiq+UpLjDknWF9SSqqoqeu6qqCtKxncxqWr6wnszqSp6insY4/NJ1UM/mD8znM5m8fFcylRcymUY9ExT6wrldedPKoIIH0k+zu7Qzsc+Bfg6yS4bKY41cl0xVC0mw9W83hYesG1VRmVPlqntf1M+67SaH4tMxqWPj/AMOXFd47CyVqVS9hbcLBQ7Q9C6my46Z0eh0+zpVXvqBoq1KLHuIdLXVHamgwXGRvAyBtjR0DSq3RrV6ggRoBIgflAAMcciJELpHdsnwOjgE4+s3smZBCxjGWVRxTWa2TbkQ2vu05bV9GzKLgZG1EkqRFoo2dwc9oYRHeqefZp6FiYtDMRGhtXby+yxdoBHZ8cNFVcGuY87h43MHhcMHG5breeR+KTvIipIhwyyHdbWvvY1eXdUKjKYMWwjuoSnJ1QsvcAdup7hh9om4nOzM4/wDhroysch+LDKOreJHIGpYiY0tIIOEq2HnNdUBWQhpXWlFy+ZxVMQkMZhBxUHIKtk7UymDENzSObwsG6iMmqku9YLs3+MuoeAXjLyb4mZO1MphI6Vn34RmRTktiEsRFV5W1QWZXXUhh0qWmGdjlykibSNiJc+7CqJTJopbzKGW+4m6g/rHQA24p1GXdswwKjHNfrYGhrhqAw4b8Ak5wUbUp0b0OY4Gm4jaNOdJzk4zx5eq+r1xcZ3NcNnC7xCZ/ONMP/wCzmTmY2YsLCxKuVqLmVJ0nNJxJ4BxKuUKRMJpCwcGUkhSkxCeQ2VY/HKr+uatzJrGrsyMwZ/ManrWt6hntZVlU83eVFzOeVFUMdEzeezaOXZsqfjYyJfiHG2u7aQFhmHDbKG0j6x/aWyuB4m+yk4tHMsJmioJXmVwm1xXVCzKVpXENVHKhRbmYNPfh3c3U8ioICBhGoIJulaY9sBJbXyq+SS6j7yy8gFN323glSh/x3eSsJWoAglslXMCFAFFyDsRjeDqYbb39Qtm5YdER7wAjAnYEkzHA34U+ovLqlNp/8YAdHBBLZj57g9l9S7sKOzey14D+DHLepE0/AxnEHnxRNM5jZyV7GMMvz1LlUyyDn8ky8lcaUF2X0pR8tioKDEuh3OSZVA3NZ9HKcio1CGM2xAO6QdtwNwbg+ZHv8umOLfBDnLTvEJwhcN2dFLPQr0ozDyaoGeBMEW/u8DMhT8HBz2UFDP8Axsvyafws0k8RDpt3EVLohpaW1pUhPKX6/n6/XHF3dSrWurh9Yuc/2rmu1EuILTEDyBkeUDgSttR0ilTa1oDQxhxOSRn4Y3nO6otPMkgcoJN7qSVJBv8AmICkKUfEBaFeDiT7Q6BX2qLs66IyMr3K7jayZpOS0fTed0+nOX2dMjp6DZlUq/3WZg3ampyt4aXwrbcMzHV1I4GpIaqVQrMM3FTSm4GZOoiJxUEwi19/gm2tirpYctybgAC6ki5vpdSQdr46lP2uPOWj5Dwf8P8Aka/HQj1e5jZ7t13K5QC27GQ1I5dUZVErnk7Wi5VDw/4xXNOS1l5wD7yqKi0QvO5CRCUbPw5Vq0+r2zaRdpe7TUaDhzCJgjERHxgAzKx79rDbVNcSNIZGeQDkzk98brrifZw+K6oOHDtQMoqLZi3BQHE1DTPI2vJb3n/l1xs0gI6e5dzdDZKG1x8oruVSuWIdKiYeUVHOktoPeJQfqIaH2gQQQLEeGpHQdD77363PyUuxUy3neafarcEdOyGFiX3ZbnXK68mzsM0t0S+n8tJbMcw51HRJaB+7wn3emzCd+stoMRFMw5UtTrTb/wBa0eYsdL7akCwIsbCwATYDpe9yQM7xaykOpMdThr30AXgDHABJ5M784z5WemF/sXNMljHaQfMAH+3EjfeFN+nj/fCw08hYen0MQo2tYXuoD47n4YnHL/CFsksNfMWPpr/OB2tr7tD066eHjhiqVcwva2tsEQA9beXkNbA+gPicW+vhhhtgiYYXHyB9xvb9DhiJaCZJP3ACJhhhiyiYgkDfE4hQ5ha9tb4Ip8NvO9/21wwwwRMMMPcT8Ph6np08cETDEG9jy/m6X2+tvnicES1ut+vp5e7DDDBE+tr4xd9sLx2s9ntwK5tZ3St+DGZ86hUZY5JS6LebCIvNStIGYw8kmJhnLqjYekpXBzau5jBot96gKYiYXnbMWpason7a/Q1v6eGOsN24fY78b3aoZxZaTDLXOHIegMjsoKRi5dSlKV3UGYn43M62qeLRG1hVc0llPUHNpJCqchYCRSOUpTHTCKYgJXFRCoiGM7i4GH2HTKdq+9ofi6opUA7XULpg6S2B2zAG/J5Vmu6o2mTTbqfsAN88nEbnfJEfBfPBoOmq8zozKpuiKag5lWOZGatayun5HBKLkVNamrCs53CyyFC3lklyNm83mDK333AQYh8vuKWlXKvtPRH2QzjRdUFN8T3DIUKBTyuQ+aieRC9ChJFEvhSCf6ykKWACoC/IMm/Y2/ZycyuA/ixRxN8UVe5MZnRNBUtNWcnpJlyuq5l+EV5UjbsomFXTpVV0tTjTTsmpmLnEHIPun359Mynb8xU5BPy6Xujtwi9tbc26gDclR1NyAkX1P9KU2GgGgx1XWPE76danS6XVb7GnTANQNB1EhsAamEAATHaY4zr7ewa9rn3DSHuLSI7Q05AME7zM/qvn1q+yFcZxRdPFFwzB03sfuWagNyb3ChSBCet1BJO1ib3PHXiq+zFcZvCtw7ZscREzzfyRzQlmUdLvVjN6Ky/hcwFVZMqflsZCqqOOlSZ1TUplvJTkjcmNTRbL0Xzvy6URrEKlcWpls/SeGtiNrXHTe1r+6/h08Mfimksl06l0dKJvAwk0lU0g4qXzOWxzDcTBR8BGNKh4uCjIZ5DjMTCxMM46w+w82tpxtxSVIVe6dVQ8U9UZVY51QOYHtc8aWDU0ESJ0yJEiTnIAPKyHdPtyxwZIfETiJ93t5jkjf4r46nAhxfVlwLcWmTXE9RzT8fEZb1MiIqSnRErh01bQE355XXlIriU8yWhOqeiY2HgolSHm4CcogZkpiLchi0r6Nfa18ClH9sXwDUbVmSE4kkyzIkshluffDDVcR3QgKll9U0zCTKMomNmCHeaWynMmnVy+GVGKcXDQFTQFOzOYNrh5a/Do67efH2SPignGdeaU0yBzi4dacyQm1b1DM8qZBW87zPZqqn6Kj45cZI5DO0S3LmcwAfkMJEGUNxDM2jhFw8HDvuPLeddSns99jJwXcU/AJwpucMfEtX2WOY0vousJrNMoZxlvP6vm6JPSFUOvTmc0pNmKupWnXYOHldURM0mUkRAuRsIqCnb8vDEvhYBhl7adav7GqLTqdjcMF7S0l9POpwOkwRAwMtJ3wcmVYtaFVvtKNZjjSOAcYyCCN/v5L5fr0JnPwtZyoaioetsmM88mK5Ydbbi4V6na1oSu6VjmYtla4SLbb+7TGAiGmIkJcbdgIyGU08BEy2KSp3uz8A/2rfJmpaWkFC8ftH1FlrmBK5bCQUZnVlvIYirsuawiYVhtl6a1FSMrJrCjJtHqS25EwlPyaq5I6+7EPMvyJgw0InOvx/8AZG8GnaNSkP520E5IczoGFTCyDOzLuIYpvMyVsISe6l8wmYhImW1fJWF2VDSisJXOoOCKluSluVxJMSrqoZ7/AGQviNkM2jIrhv4o8p8xqdWXHYKWZsSCp8sanhk3s3DPR9Nt5iSKaONNpDaozlp9tZUXEQDHOpoXn9T6H1ukwdRDra5aGj2g5Ia2S18kQc4fjfbiAoXdqZo++0uy3/205M9u4EGDuuxJUP2jHsj5BI3503xNRk/W3CmIaksgyozYiZ1GKSlS0wrELE0XAw7cS4rlaSiMjIFhPsqdiWx7WOqh2v8A9oxr3jmoepeGzhlpKfZLcOFTORMqr2oapdlqs084KdUCyqmpnAyiImMooaipmtwPTKSwE2nM8qOHYhoCazeAlLs4kExxW8e3ZT8ZfZxzCmRxD0HBOUfVsOoyLNHL+ZRdYZbvTNp4svU3MKjErlhkdTNpKIpmUT6Blr81glri5KqZsQcwXB8Q8i4zJGVZr0XMOIyk69rfJSHmoNeU3lhVUpoeuZhJ1odQVU7PJ9KJpLkRUO4pqJVL3zLTGw7L0ExO5FERDU2g9h07oXSabRe2rjfOpj2jGF7XS5oDgAAAJJ74mMzJWPWu7kkUXxSDoa4kQQ2QCTG+J2+PY8n+zZ7ObPntLM/JRlNlVJ46W0NJY+VRmcubkbARiqRywo+KiFOPxUbHIYdZjKtnMJATKFoilW3W42oJskKfMJIIKeT2SfWVyaykoXIHKLLnJPLWUtSCgcqqJp6hKTlTViqFktNyuFlsEYhxQJiY+JbhBFR0W+pT8dHuREXFOOPvOuL6yPD59oF7ELhEyOkGWnDll9m1lvS0jl7S2cs6UyQioGdxkzTDNtvRdRVVM59+EVNVEUGwmbVJNq0msfM3kqXGTGIJQ43zb7Lnt1Mru1C4gM2MkqEyTqvKeHy/oFrMKl5zWlXSWZzytZSxUcLTs3EXTUmgls067L3ZxI4juWqhqBClRi+8fbU3DiL5/r7+r9QJr1bKpb2VoYaHgAAjQJIMcDETjtJWdZi2oe614fUe3iZyGkmcyCYP6ro/duHwQ13wWcfuc0JNqeiIHK7O2tKrzmyWqdtKVyioaYrGdKnk8k8M837ENNKGqabxtKzWUvJbimoWFkk5KXIGfy6KiMgHYQ9upSfZ80tM+Friap6oJlw9z6sJjV9HZgUfBrnVQZUz2om2U1JBzilzEsvT2h5tHQrM4MRT6X6jks5jZq+JRPoKODMu73PF7wW8OvHNlHNMl+JDL+X1xSkdzRMsmAcdllVUdOx/6FSUTUsCWpnTk9hiAkxME6iHmEIXJVOIOYSeIiYB7qC8Tn2RbMWBmkZNuEDiapSopA8tx5iic+pPMqcn8sbQQW2Ga3oiXVFKZ6Sk8qXH6HphLKUpC3IpVrZtr1vpvU+n07Dq00ywNDKwkQWhgDpgwYEHvngqzVtri2qmtbQ8E+80DIBgkEY54Eyuws52/fZIt0+ahVxiUh3XdB9UoRSWZi6kCQkL5P8AT3+jBNS7y8yVoRDFQIISUG6k9Qnt1O3RkPaIyyQcN/DfJp/T/DTSNXsVfUdW1hLmZTU+b9UymEjJfIoqBk7UTExtOUJJGplNo2El80chp7PppEyyYR8BJkySHh3/AA6vsq3alvTAQhd4a24UuKR+JqzdnJYCASUumHZy6ejACm9kqge9AKiUXAGOfnC39kHrBybws84x+KKnpfJGXmX4nLzh9kUzmszmgHItTUTmZXkJKYaUWKVMOiAy7m63WlqXDx8G+lJbla0vDXTKn4sXbrh9N002ag+CNJBAABkRicT2MKj3X9y3RpDGuDZjEg6eQe248zOFid+z28E1XcWvaEZW12xLo5nKPhdqGR5z5m1KmGWZe3OZI5FRWWVHpilNrh1TerK0lTcaqFESl9dM09UUYwlhTLSHOyj9qY4/hkXwxSPgsoaYNt5i8U8DERdeuIiB95kGQ0jmSYebNuQzTgWo5i1JB/6TS2+VQcXT0traEU04otKR2KeF3hTyH4NcoKfyO4dMvpTl5l/TqVOogoAuRk1nk2eZYaj6kqqfzBT84qapJkYdtUfOJvFxMW8hpiHQpmChoWDh+pH2iH2eHtIOPXi9zh4mqiz44Y4OErefqhaGpia1Tm5EO0bltIWxKqFpVow+Vz8GwuDk0O1HzluADcFE1LMZ7M20LcmDjrmKzqlr1XrNO9vajKFratApsfgv0kQTH9RJBJzEARmRd/DVLa2NKiC97zDn8idIzMztAMbLqN8FHBvnTx8cRdIcNORbEmTW9Vw85mz88qmMj4GlKVp+nZbFzSb1FVcxl0onkdCSprumpcy7AyqMjImZx8vlzEMuJiGyjsEQv2SDtBghJic9uEVo2SClqqM4Hikg3HOTk6yk8oJ5CGwq5IBupWOxl2GPYsxnZeSfNuu85Z/QOYPERmfGQFOMVPQ653HSGksrpN3UwYp+SxdRyeRzJuY1LUTjs1qlX4Ww3EtyalmEvLVLFLX2DCn82oubeW3zxLqviu4N0afT3MNs2NLi2dbse8NW+QADAiDuo0On0/Zg1gdZIBAkQTpPciN9/PGYPz1UfZIePZdgviA4S79QZ7nAoEXJIPLlKlJBBNtNyQQdScbPaUdiJxWdmTlxQ+bWbdU5VZjULWtXRlFPTrKuMrGPapSeiA/GZIzUqaro6lHGYepIaHmv4W/LhHNMRUliGo95l2Kgn476p6UlJvoenX18uuh28tNccTuOPhHoLjk4V84eGPMNqHRKcyqVehJNOXGlOxFJ1vKYhie0NWcByKQ4I2l6slsomzSG1BUTDQ0TLXy9BRkQw7i2niu/bcUfxL2mhrHtGtaB7p0iQQAQW7jMwI7q5U6bR0ODGu16SWmc/wBJ9MxzOCTnJXz8/s13H5EcJnG/LciKuiVN5P8AF3ESbLiZrejRCw1L5qwz0R/tVUwaVdhwTSaRkTl5Ft8ra71TKo5a1NyVDSvpfghQCtCTpcWtc3uBb/6VadLK88fPDhPskXaGQEVCTCXcQnClK4+AiYeMgIyBqvOKHjZdGQbyH4SMg4yGyoS41FQjrTb0I+wtt+HfS06l5K2eRffb4fpPm7IslcrpNn3N6YqDOeS0PTspzNn9GxMfE0xUFYS6XMwM5n8ncmssk0wTCzyIhvxNUPFSyGehXYlyHIUhhDi7PiZ/T7mvTu7Ks2oaw/3mDdphhByD8Y3HxUrAV2sdTqtLQ38pJwQIgeUjjnK3i+j/AGuCL/pvvbHTk+02dlPV+bCpd2gWQFHOVFP6KpFqnOJCmKfgFxE+mVF001GxdP5qQ8uh4d1ycOUfLjESCrw0lUazSENI5iG1wNORrzHcb/YDp11vr1/Xx6DFFthYIPKQbpIUkrSUkg+0gq5VG+w0sBuCBjT9Pv6vTbqnc0gZYQHtEw9kjU0jfPxHlsDlV6DLik6m8TIOk9nYz9Oy+QLwL8dedfZ38QVPcQuRsdL3prCQa5DWlHT7vXqTzHoOZRUHGzakahRDOtRjDUW5BQkwk82glCOkU+hYGaQofSiJgo3vncLf2nPs3s8KWlbucFUVPwuZiLQ01OqVzEp2dVBTDEWtJL8TIcwKOlc8k0dI0quIaMqCEpGaK5kmKkcIgFwfn7Qb7NZwccZM/nWZeUc3jeE3N2dl2Mm8woKnIGfZYVLN3LKemc6ysVH09CwExjrETCNo+fUyIx1So+OgJjMOd5zrkV59k67Ryl5nFw9D1/w1ZmSlpxQgpkxWtW0XMY1gaQ6omVT+hH4KBfdTYuQjdRzJDKhYxS9FHsq114d64BVu3us7loaXvyHOnSIJMsIHAxg+WdbTpXtoBTpxVZiBP5QC3gnGJGMRtMLtuV32/fZKUHJ4icO8YNF1cWGlutyzL2nq2raaRCkpChDtMSKmolllbilBptUwioVvvFBT77baVLT88XtZuMDL3js4988OJ/KuR1ZTtAZgKoCX0/AVxDSuBqVcPRGW9JUNETGOgpNNJzAwSZvF087M4aEExeioeFi2WIttiJaeYbyYyD7K72pE3i0tTV/hxpeGU4hJjprm5NZi02g25nO6p6hZ1ELDSTflU0DoClKjpjK3wr/ZDaDksRL6g4y+Jea1462809E5b5FSRyj6fKW1cyoaMzHqdyYVJNIaJQtTb34VSNHxjNiuFjm3FlxN2zf4c6G51zQvX3FUtDPd96RLSY0gCSREkY7gyrVRt9eRTfTFKmMkQAZ90nzEz5AiVyP+yLMJRwB5+vEApieL+qbE6koRktkkmxJ3CiF28SFa6HHWK+0WsljtiOLFSdO9YyPet495kBlje+x3Bvr1tpYX+llw7cNORfCdljJMm+HjLWnMrsuqfuuCp+nmHkmMjXGIWFiJvPZtHvxc5qafxzEHDiZVFUUzmc7mbjEOqPj3jDtFPVJ7WH7PBxg8d3HhnRxR5U5n8PdOUPmNCZcMSaT13UGYEDU0Oqjcs6QoqYLmUNIcup/LmExEzp6Meg+5mj5dg3GFuBp5TrLeB0bq9qOuXV9XeKNGu1wbrJn+jSMcnSZ/ssi6taptKdCmC8sLQcwDlo9SMTscxlcNPs3faX8FnAvklxLUnxR5wsZY1BXOb1MVNScC/SlZ1CqaySCoWEk8XGNv0rTU8ahgxMEOsKbjHmHSVd4y2ptbpHZA/wD0iDsimyf/AL1kKux15crc5V6ag+0jL4AjTS2505rY6rj32SHtE+YlvOXhJWL6D/WmardgLDQqydUSTtvffQ3tjRP2SbtGEjTN3hKtpqK5zTt88mwffYgg6a2xmXtv4cvrqpcv6i5r6ulzwHAAFoYBpEcjv377Rom+pMFJtLU1pGcQD7m8nJaARIx6rtVn7Q92RZHMeK6EtpvlbnGVpFtrGgCeUW2G3wxy54UO044I+NyMzEg+GzPGX125lRT8vqjMFcVTtWUhC0xT80emLUJN4+NraQSKG+6D8ImCop5p137k2wXYtTLC+dPSgH2SftFFpPPm9wlo3/8AxvmjrtaxOUCL9elh1Phnj7B7sW+I7szMweJqouIqq8lK4kGc9AUPR8llmXc4qioUl2n55UkxnCKhhqrommIf7jFwc1h2WkMmP74LdbiIdDdufUdRseh0bZz7O+dWrjTpplwdqaS2YgCCBn4E7rJpVLx7w2rS004EnEgw2DiDjPyxjK4LcVn2sL/avjdiaZ4fMsqXzu4PaDgn6PqibRUzcklU5p1O5HQkXMMxMsqmZEbAyml5Ew1ESGmYWcy+YQ9YQy4ufPJhGI+TRMuzBZBfaPuypzwkMsj53nw/kNUMVDtuTOkc8qVm1KxcmilIQHYZ+p5S3UNCRLXOpXdRUFVTocaSHHhCKP3dG0nHf9ma4FeMCo53mTlnET3hNzTnynIqZzLKuWS+Z5bTmaKvaazfKiYxEvlMLGOX5oo0TN6IajXQYiLbiI1cRFRPXizP+yQ8fNJzCKTlbnRw5ZryNLqzL4qYzOtctJ6/DIUtCVRsmiqYqeTQcQUpQlxhir4xtPMoXcKOZV2hQ8N3dGk01altXawNe550y8BskkgtMHMDT58hW3Ov6L5DRVaSD7sYBLRGYO0jnuCdl3Cqv7avspaRkT0+mPHfw5zeGahlxCoGjq8ldez59CQVGHZpyj0T+fRCyUgtQyJcXVLCEqQpCnLfNa7VLiqyw40+PviG4j8naPXRmXVf1HKW6ahYmBh5XN6gh6WpWRUxGVxPZdDXRAzmuI2TPVLHwZP3qGE2QJk9FTYRsQ7lEkn2VXtSptGoYmi+HKmYYrsY+ZZuzeYMo5rDvQ3IKDnMTYjSxh+axWOUGwOYrgm+yRZW0DOZDXHG/ng5nTFyyNh5g7k5lRLprR+XES7CPh0QFSVtNYsVrVcsiFNobioWVSugz3C4mED70O+sKzbKp0XohqVqV465qPZpaxuZBLCfyiAZaMzA+qs1G3d0GsdTbTaHA6pJdA043mCCc8mY4W5v2R/hnrHLHhLz+4iqphYqWyriRzMpuVUDARTLzAmNI5LymeShyrIYLIZXBzmq6xqeQwym0CwpN1QU5DOQi8dtXe40It1+Atf9PS4x6/S1K05RNOyOkaQkUppilqZlUFI6ep+RwMPLJRJZNLWG4WXyuWS6DbZhIGBgodpDcPDQzTbLSD3aEAI5l+eKgk677g2Gmt76+nXTHJ3ly68uq1y4afavLg07tGAAcnMCT3JJIlbSkz2bGs/4gD5AD9l88/7XKL8duQW4I4WJZy730zTzEtpuAPZNyLWNsbxfY/UJGf3GUo9Mnss+YeKTWc93uBf8oudgQnY64ya9ud2I/E/2m3EjlbnDkbXmS9J09RWS7WXU2g8zJ1Wkqm0ROIetqnqExMCzTlD1NDOS1UDOIZCXXIpl9T6HElgN8jq/f+wh7G3iQ7MHMjiCrHPSuMn6tgM1KGo2maeaywnNXzSLhYqnp9NprHOTNup6MpVtqGcYjWBDqhHoxanW3UOIbAQV9VU6lZnwyyyFww3IptaaWQdQewlsRGAJB+WVqmW1YXxqaPc1CXcYDDnOTgx8guy4SACTsn81jYgajUi1rkEA6G6TY3Bx87j7ULx4nP8A4rZFwjULOoOYZXcKgiEVa5K4tMSxOM96jg4dupYaPeYPcuLy7kKJbSrcNo7L6kmVZQEUlL7CUD6A2bCMx15aV8jJ1NLLzUfpGoEZdmt46Yy+kUVo7LohinYipY2US2bzJiSQ8xVCREzEvlcXHOQrK2IdpS3StPQdqT7Kn2j1W1BPqpqLPDhYm1Q1POZtUM9nEbWeazsXNp9Po2Ims6mkU7/tBzuRk2mcXGR0Y6slT0REOOOLWtay5q/DLrGjdOur6sym6i2KTXZBe4NBcBvgCDjncELKv/bvYKdGmSCRqc10YwCCPTM/2CxV9nP2M/FV2nlMZk1tkVN8sKKpDLOfyml46ps2Y+rpLKJ9U8wlj03jJLTETSdHVcIqZ09J35NMJ4zFNwYhYSfyfkUDFFwZOG/skXaDFpPNnnwipWN0qqjOBRAIsQlScnALaXAAvoADYAHuzdndwc05wG8IOTnDTInoCaTSiqeEVXtTQMGqEZrHMmoHVTiuqmCXEiKUxMJ/FRTEpMWpcRCyCFlUArlTDhCebO31tjKu/FXUPxFQ2ppigHRTlvvFoLQCZPIGrPH1t0um0NDfatdrgTH/ACIBJMeYz+sYPzx2/skPaAhJC89uEgAknu/9T5wrR7SQlRCTk+UhRTdKrWC0+yrmGmMEfHXwOZ3dntxDT7h2z1hpYuppZKJBVEhq2mVzh6i67pWoYRTsLUdHzKdSyRTKPlsPNYWa0xHuxMsgnYOo6fnUAUKMMh2J+wpjA125PZAzTtQMvsqJxlLPqMoniByjn8bByio64fm0DTM/y1qZorqSl5zGSKST6ZNxcBO4OSz6nH2pe8mHdE8gFIQ3PYmLYvdN8U3TrpjeoGmKD8OcGgBhgAOdG4BxB4O4hQuOnN9kTQ1awQQJ3gid+0SB385J4L/ZcuPKV558NFYcCWaE6XNq7yAg4uZ5fy+ePfelVJkFVMV91dkkEt8uORcLlvUke9TzsAvlallNVDSMuhEuQUMTDdYTtoeyuzB7N7iTqeKklLTeJ4V81qpm01yFr1ltUbKZW1M3IyeOZSzuPsgwVW0VDNRUulsPMGkOVHTEDDz6DVErTOzKM5/AR9ns7TvgP4scneJuh85+FeYmg6mbYrKmWqvzWQisstqgZcktf0m8HMpG4VT0ypyMjXpO/ELLMtqOCks4Skvy1oHuZZr5Q5Y565f1HlZnFQtM5kZeVhL3JZUtI1ZKoabSWaQboSB30M+j2IpjR+Dj4ZUPGwcU01Ewb0LENsvtW6vVaHSusPuOnvFe1uGg1mMmNRI1FoP5TPvCd5I7FSFs+4oMZWBbVplo1HkQJAiBGIO8DmN/nV9iN270f2dUviuHbPuR1HmFwtTyoYqfSONp1bUbWeTM5nMQ9ET+Np6TxsbCwNRUXO5hENTqdU4xGS2ZQMzemc7kyouNmUZLI/up5e9tZ2V2ZEnh5vKeOTIOng/DiJMqzGrKDyvn8OhbYIZjJHmKmmZg1EJFgpDTZK3LqQSlSBjA7xf/AGSPLmo4+aVbwR5+RuWD8RFOxULlJnPAzKsKOYS84XBBSXMaUPrrOVQkMP8AihUVDIK7jnUqQh+aJQyVO4h5/wDZbe1OgYtbEtl/D7PmA4ttMbLM4nINhTaTyB4Q8+pSTxiEuBPed2qFSU3sQLEYvXFLw71NxujdusqryDUafdBficOEEk8jJMEBW6b7+2ApikKrREE5Ay2NjJgEiMxtgldsLih+0ZdmTw80lNpjRuc0NxLV0zDO/wCn6ByMh3qkZnMzSFCGh42vIlmCoOTQSXkJejo12fRcSxCBxyXyuPiTDQj/AM/HtD+PrOLtH+Iue5+5tpgpSgQDNM5dZeyRT8TIcuKCl7zsXAU9LHIhCIqaTCMiYuKnNRz2LQ3ETedR0VFNQsslbUslUtzMZY/ZNe0KqybQbGZuavDrlJIHltJmcyZqGrcx5/BwyyQ8qApyU0tIJXMYptP/AKbEVWMsZc3McmwB7LPZ5/Z2OCjgbncmzLq0THiezwkvcxEprfM+US+Fo+lpk2tDwmdFZYQ78xkcsmrL7TK4GfVJH1hUMrW0H5NNpY+t5bl23uuhdEDqts517dOaNLiJ05GWloDRkTyYHMAKjqd5eENqAUaYLSYO5EE+ZP0EYECFxD+zW9krV/ChRs+40OIemo2mc7M56WYpnLbL+dy77lPctsp42Ig5vGTqoWIu8dLKszBjIOAW5JHWYaOp6nJZAMRrbEbPZrASrtcG9iRvY2v4i9vPfx9et8ALDYDQaDppa3nbYaDTE45C6uq17cVbmufeeRpEzoYBho7Af5W1o02UqYpMaQBBJMGT35HrGfNV6J5t7+mtzbYD+LWxa411Gm+u3XX3Y01K1I8NR69N+mvribeH9SSSel7C36kHTw3GosK4r/P06+nTFTcKFrhNug6+fy1+HXEgWAGmgA0208PLAi4ttt+t8EUC9zvbS1xbx2xb6/yLG99sQBYk3Jv4m569fficESwG1vd9e7DC4uR4b+/DFtwdOJjjPp8s/wA8ImGGGLaJhiCCQQPrXACxJ8bfL6GCKcMP8eP6frthgiEA79Df34YYYImGGGCJhhgfq30MEUEA7i9vrywAANwBfyA2GmvqNCdyPPAeh95v+52xJvY23tp64JH39/e/cpe5PjZPwFwOn79MPH68/wB8RbW/lb9/5+OIKSea25tb3Ww+W0ff36bBFa1tPDTD6/v9a+WKptcH+k6nUHzsLamwvti318Bf9ME+/VTc+Og6eOul/EDoCbanbrH1rhiiOo6g6j3f2PwwnEcHjj5Ir4bgjxPTTWxAPgNNCffhiCCRpe/lp5fD+3jgcx2AAjY48x8h5CIRen1/l9Q2adIz6gcyqPpqvaIqqWxUrqSk6uksvqCn55LYpBbfgppKpoxEwUZDuA3Lb7KwlSUrTyrSlQ6wvF19lO4SM2pjF1PwtZm1fwuz2KdfedpGZQCs18rF95oluUSuaziSVnTf/KCpwis5zLQ0O6hpIwTz47VSdhvfW9yCdz1Gn17sWG/1b39NN9dNPDGXa393ZH/t672CZjUYxG42OeNsnsFaqUaVUaXsaR6ZOAMn4Z5810RpT9kBzkem7DNRcZ+V8FIe/s/HSXKSqpnOzClSedcPLphV8qgm30tpWEsuzctN8yQpXIgpX2WOzV7HPhK7MiAms4yigKjrXOOqaehZBWudOYkZBxlVzWVpiWJhFSSQSyVQ0DIqPpuJmjDEW5K5TCLio1cHLUTibztUvhogZX7XUFbpsRfbobX9bjY3362xOo2AHppubknTXW58zbptlXnW+pXzPZXFwXMgS1rQwOwPzRkwRyc777Rp21GmQ5jIcNiTMbbSMTAmPkn7frf9R/GA0Ol0m/MSNLqAABNtVGyQLqvoABa2GGNV8zgb74Cv7beRnc48/P7zlQRc3ub66j2TZRBULjXlKgCU3CSQCRcYn9TqdALk9dAB8sMMMcNA9J/lE+XxH6a4qEgXskelk6nqdtTqdTc38Bri22IBBtbre3nbfBFI02uPeflrt5beA2wwww+/8p/jGP0TFEnmv6mwI0IFtCNLjXUHf0xf6+tsVKeY36ADUHrfz8d/j7n39z9/BPv7hW2201JITdIPh+UjpYXNzYDU6WD33sRqSTYkEjUmwNr6b7HTaAABYX3vr7sTceO2p9PH5H4YoZkREZkfKCPrPqnYdjI+Ij7lMMP8/r/GGKonz1Bt6EHXyuL/AKa4akak3F7H+ocwseVQ1TpYXSQdBqd8MQTb16C+p+up2GniMME+9JHkSD/f6IqqHsjlJNrWJuVctwQL6kW3ve9+t7ESVgaDUnY72JtqTqVbdbnX4W+vDf6267Yi4BA6m9tD53O3kdeuKyI/KB8SSPjOU+XwAH6JbTXXSx+Fjbw3Ou+ASAbgWJ1JGhJsBdRGqrAAC97DQbnE/PED0t8P2J+eKQMTP/7Ecz/P3hPv79E8bGxO53Ot/j16+62Kj2rpV7RHiARYG+gtbQ9bX8TvidzcflKSD43Btt5a9MACCRYlIBsQddfE3Gl9/LEgQBEHj+o8Rnby9c77Ip5E2tYWGwFhqb62FtdTc28tL4np4a3sLi5OmpH5gBsFXtYW1GK28la/+79PaxN73tvqB6gX6+G+KEnET8XeWOOIH2E+n7+vKn4dB9dD+567nDTewBIFyBc3AFva33GmunQdMVSdLHoeX3/WmLbfz00/X54rIxIJMzvv8IPphFBTqkjTlJI0B1OpBuNRqQepB3O5n9AAB7hbe5udNcVsCSbnTUf9SbaAepFjiwJsL/DwwDgCTGT5knjk8QO3A9Un4egA+5/lQo8u/ppb+cAbi4vv+3T6GISCCq/U6fPBYJFh4/ziXtPL6/2VRE5mPLdWsDckXuALKAV0sTc630Gp1seUHlFsRYDQaDTQbEJ2BTsQNLaaWFtsT5e/bDFGmXDtmB2HHyx8lT7+/wC6Ab28LEkj2h+48tdb+OK8qbg7G52ABGhFwRqDa9rFNvUnFj+x2392KpuL3v7zffTptvri5nggfD+/3hPv+37K2ltNtrXJta4trqdt9fXFdefy5d79b+HoN8VUFE6baW1AGvqcanXz/wAn5anyxVFCiRbXdQB0vvfpY/IXHTFQPaIIB5dQSAVJJPMQFkc9rk2F7AaCybDF/h9dfdpiLi9r6j+QN9tzte++mmCT9wFOhuDrc3PN7W21gTYWNiLWNwNbADAaaanzJJJ95JJ95wxW6ua1vZ010H6n49PffCB2+/sBPv7/ALK3j5gixFwfUHQjyxAABvYbg7J6eXLYDyGnhbE7AnoNz0+O2G9/LfESATJ22jzx/bH8pOZ/YR8vqpFhewFyLX/qIv8A1K1Ury5lH34jDbEAg6jFQAJwM+XH36IpwwI0N/TTfwuOmmvr0xTU7DQA2VcXAA1PwF7Cx6YqiFY5rennv9EddfSxsSEgfAfV/LAbDW/meuBIG+CIDfa/vxOIBB2xOCJhhhfUjqBf3Xt+vvwRQABoMQTYpHQ3/a318cWuLkeG/vwwRMP5t7/DDT69/wBfHbDz6nfrr5+J88Y+28/5z/EeSJqBY3JB10sT6Drby3Fz0w+vP4fp46+GGGCKpUQQPE6emlx79fLXfwYthgiYYYYImI6geR+Vv5xOIOxtvY28fd8uowRQg3BPiT+2JJ28zb5E/tiqNj01O+m1sX8ProcEVOZRKgOXQnofE9Lg6/38sVAvZQvfmF726fDTw+WNS2p03+tfPU/zicETHr1RVLT9KwD04qeeSenpTDnlemc9mUJKoBtRSFBCouOdZh0qWCQklz81wkKWOQ+fWSEkjYam2+gUR4W1GqibJFyQRe2FOX0W7xp8dOccjzim81jMqOHeIiJXTGW7Ma/Ay2OjEzJ+n0RMb92dZeSiPfl0zmc3iIRaJhMG3pfLERLcrQtozaGmdRgAE+cgjjzn6FUdIHugEkxBMfE/t5rJ9IOJLh/qaapkdO5z5ZTebuuhhmWQtZyExjz61hCWYdpUalUW6tRslMMlwqVypCSVDG9LsVDsNORMQ60xCsNl5+IfdRDtNMpALjzrj5QhtttN1ulR/wCFNg7bVWOHtU8BPCXVMkiJK5kxSsh75kNpmtKw6pDO4ZwAd24iPglIXEPKKE3bj0xrT7iil1tQUrHBvi9mcbU/ELw38AFOTidUllJNpFT8VVMQzO5nGTaopWhyoGYaRzCaTOJiYyZw8rlVGxDcG1HLjW3pvNWHYuGiEwcvQk1of+Ux/wC0bYzx3xEzCCYGqAYkxtnKydP8TPD1CzEyiIztyramSV90uEVXVNpWhaSEqS4ozHuULCj/AOn3hWPy2JxvBLZnATaDYmcqjoSZS+MbD0LHwEQzGQUUy4eVt1iIh3FtutEm4cSspI1STrbirCcC/CZByISNvI2h4hr7v3Jm8VAvRdSLSlBbVFip3IgT5uMIKimIhYtCmSlCkpHdpaPCXhpbnHDFxzVXwmU3PplPMnatkMwrWn5PMnhFOUrMHZGiqmHWFJt3DobZjpTGvoaQmcs/g8zjW1RKHlutDXZa/LeDycD6SfuYEx3PpustMor6iqgn0+peR1XT03qSmV93UMilk5go+bSNanVsoTNoCHdXEwClusutBMS00rvEpJAC038nP6kkFKyx6dVNOpTT8nhi0mJms7mEJK5dDrfcSywh6MjX4eHbW88tDTSVuguOKS2n2lAjrrRdKZ+xnF/xfZs8PU1eTX+Tdfxk6cpNDUVECuaXnE+j2J1IRAsOj8VSESyEilSSIbLkwDCjL1NziGlhPKrPfiro/iY4DK0qWVMmQ1lJagoiUV/REYpQmFLztyqGENBBeCVxcnmQhVxMmjynncaT92ikMTaHjoViZpGWxyIPYHAExmdjxk58gPqJMCcTtt81lPn2c2UtKtyp+p8y6Dp5iey1ucSRycVZIpeJxKHgkNTOVmLj2jGwLqnE93Fw6HWFK9jnAUhzHqC+J/hyZUUOZ65TBZuAkZg0vdWh9myZktQudLhJIvYJVsfSaCyOyhzSyZyMmmZGXFJVtMIDJ3L2DgI2ppHCTSJgYN+lZPFLhod2KZU4ywp1alFtCj/yXUlCVC5xucOHD/k7UXH1xZ0LUGWdGTmh6QlLLtL0vMJHCxkkkL6pvTjIdlkvfDrEOoNvRCQpAQQl9QCikqBiGCXE/lHaZjGx7ieZzPkqyO6zAt5z5RPUs7XDOZ9Bu0bDTNuTP1WiqpIadYmroaUiWuTcxwl7cwKHW1JgzE98ouNiyevr7fEzw7urQy1nnlM44opQlKcwaVvckAbzYX3udddfHHBDtIcsqKy44N5xT2W1IyGjpPEZl0bMoiU03LoaWQTsa887Dqi3IZhDTK33GYaDZWoXdU3CwyAohhpKNWn5P2YS8u6aRVh4Z256mjqf/wBRqhphTTNTszYySEEyW7+DRSZy3N0xodDqEJVGrjjypSuIWlKgptgHU8hzjsBA2Oe2fPIJxIIUSexiIkdxIP0xtvgbrKPLZxLpzAQs1k8xgJrLI1HfQkxl0XDx0viWFBVnYaMhnlw8Q2OVSlOIWGxyOAK50En02s84MrMuXEM19mJRdHvvI71lipKllMniXmRp3rMJGxLcQtoE8oWhBDluZAKQDjCRwhZk1Xkvw48a+YeXonEXldS9TujJCJqVuJel4msdHzOUqmDKH+RLzcJL5jScwn7DKW0vRyIlUSGI12ISrk/wZcJWVeYuUskzxzxkMvzjzGzXXH1VM5pW16hh4SHVMouFlzKIGNdVAvzBDEKh+Pi4uHW5BxDi4CBTCS+ESziXsmtBJeYEAQJJwJPbHKoHh3u/ljnIOfOIx6rI5RebOWOY4eNAZgUbWf3dIcfTTNRSudOwzZ05oliBinnodJUUpC30NJPOD1APl6orqjqHh4OLrKqKcpSEmMWmXy+IqOcwElZjo9TS3kQUI9MHoZD8SWm3HQy2FOqbbUtttadDio42OFqgMk8v0cSHD1BsZMV/lbOJHMS7SIVL5VNZbN5rByVUOJQ2tcuh4hmPj2IpxMNCswU0l33+VTCFi4V/vYb1TjxzCiMx+FzhJr6MhmoGZVtWFL1DFwbCVJQzFRdHTJyMEOlZ70QqYhx4MFQ/9BTRUsqc5lRFPVBaRoOJO87+kROe+FNZqebmCSghQWgKCgbhQIuCnpYjUeO+PR6wzNy5y9XAt11XVKUe5MG4h6AbqSfyuSrjWoVTLcSqDTMIlhUQGXYmHDqmg4GwuyrEpGPc4Qp+6wtynWHZtci4AZRcEX5rg3FrHW41xhu7TSY09LM6+EmPq6RRVU0hCT6bRdTU7AS/8Yi59IWKmoxc1lMJLQ+wmYPRsH3sOILvG1xHOW0rCjYxYA58ahvtI/b65O+FF0gGJnEfTb1WS9PEhw/OJsnO3Kok6ECvqYG4Ot1TJJO9/wCkHcAAADdqAjYKawUHMpZGwsdL5lDQ0bARsI+3EwsZBRrDb8JFwb7Ki0/DRDTrbrDzayh5pXOggYwxwGa3ALFxLSE8EWYZfU420h5/IKCea5lLslSlqmz/ACICiFKXrypuojQ4zJySXS2USOUymSwLMrk0tlkFASmWw7QYYl8vg4VqGgYOFaa5odlmFhW2mWWUEhptHdJUlCQk1c0AjcScg8bKo2E7wJW38/zzyVpOaRkkqjNagKfnMuWhuYSmb1ZJYGZQLjjTcQ2iKgoiMRFMrWw608kOMp5m3EOJKkKStXiU8SvDwblvO7KtWmnNXtMJ0ChzGxmYNgm6rGxABNjax8XX3DdkPWkRUNUVRlNQs8qaZQ8TFR09mNPQcRNIuJZgCw0+9GqQp91xpqHYS2pTh7tCGkJI5ABjZ7M/IXJzMjJqs59mHlpRdYzmCzQnEpg5nUMihJnFw0uRT1LxbUvYdiUO93DNvxkS42kJT3bjjig6Sq5kGtLS46sAfPnjbcZwoumWxO+Y7SM/ASVlenmc+UdLwUimVSZmUJIZfVMD+LU1GzeqZRL4SoJWW2XRHyd+LimkR0GW4mGWIqDD7IQ82o/n9nj7Q3FjTM5q+qFVpmnw30zQcKtxijoKX5oSuc1pNQiI5W5rOolUdAyCWwsTDoU+1AQaZjGNGMhmop1hxhxL/BXtEP8AQVAZ5cG8FO6TEdldTMDPIWY0PI5KzMW3qWgJtT8ImRS6QtLYai2mYVtpiFl7TrRLDKkNLQ02W1+flOZfZ2zSPhodjhMqVETFPsQrTkbkCpttL0S4hhpT0QYkpbaS4pPeuKX3aG+ZSld2CTQNBbLczETE78Ax8VNZkDGwghBH/eWDBGHRFiMDyPuv3VbIfEUYkEsph1Mnn74r7lDJS8t0oNsbSu8RGRDcd+FrziyyRMe8DRgzXNM98l4q5Q0QJoU84VdPdlYUsgWAChbHdxp1JOMzeInILgnlc6iqLy9rKVwc/rFcqiVQbs7lSXKgchZBytHuHIGXwFFxzUHCOd9CxU0mEI/GsuolzRxzGguCDhUgpE3JG8lqTeaZh0w34nEsRj9RPBLYQYh2okvCbpi1aKLjEY1Zaj3KAnlSaBoEaiQCJECe28xweJRciZtWFKyOnYirp5UUjlNKQkK1HRNSR82gYSQsQbzjbLcW7N33kQSYNbrzTSItTiWHFuoS2taiEq/VIJ/Iqpk0DUFNTqWT6RzJr7zLZzJ42HmUrjocqUA7Bx0Gt6Ei0pUCFLYfUhXKoApIKU47eKDJejMheATiAomgfxpuS/gUVNm0TidRs3Ww9MqnkjzzEIl9wtQcIllLDKYeAhmkOqQYyKL0VE/e38fXCXmpnLwWUzlBWle/iNXcJufEug4x+PaTExjuVVWxkZFw0aWoZtqIdhiRCOzCKhGk/cajlcS+/Lm2p9BRDMXUU5Dy2TsGk4JjcAHGMnnZUkd8/e/bj5/PsFxda0hAVJL6Njaop6Dq2bQwi5XTMROIBqezGFtFK+8wUpW8mOiodKYKLBcYh1tgsO2cshZSXWlICp26HXU0hRWb0IY9qlVziXpqF2BDS3zGNylb6Y5cKlptx5T7TLrZZbWsK5UhWMZ2YlRSWp+0e4XJ3Tczl85ks2yhem0tmctiWYuDj5dHpzRiIaLhnmFLbdYeaWFocSSClY8bHSnzjw7WiigFEsHJmMQsXJ5kmk6nKioG3MErKSbgC5BGtrhTJ7/lLsiNtp9fL1CpqEE7gHj0H7nfkLKFOqkkFMS5ycVJOpTT0pYt30yncwhJdAtXBt3kZFOtQ4J5TyJKkl5Iu0V3AG30kz8yRqeYtyan83cuZxNXnEMsS6XVhIouMfdXflZYh2Y5bsQ4rQ8kOl1aSoJWgG3Ni8kFMs8avG/nbI84ZhMJrljw5R8VJKTy3bjHYOVRceiZvyByZxTTCm3FfeHpZHR0zikusx0U5Gy2XpjFS5l2BVzerDgV4Xavkb0hRlTT1KLWhoMzajYcU3PIJTLqXEOpioDuy+AoJLqY9EWXAOcNiISy6ipYxoGpxkjgSOfiqggiQuXD8TDwcO9FxkRDwsHDNrefiIh1LLTLSCSt155ag20y0kKDjq+VKVA3tykJ2Yc4meHhuZfg7uduVaJlzlr7oa9pzn73m5O6Wr7+G0OBXslorK+b2SEkgYxp8VswnOZHFHkBwHwE9nFN5TPUxJJtVYVOpjETSqpdBwdRRv3GaTGOiHoyZmCk9DqhYH8QiIvvZvM3ZhGMPRcFDtt84GOBrhOgZKmSMZG0QWkw6WhM3oFxVQlKEBP3pdSCKaqD7yCpazEfiLfdqcKeVpBRyVNMNAJJ94SztBgiTHb7HNGvDiQOCR8lyyg46CmMIxHy2Lh4+BimkvQ0ZBvtRcLENqAKXWIlhxxl5pQ1Stpa0EGwWSDjVecah2XYh95thhlDj7z7q0ttMtNoUtx11ayhCG20JUtalrSEoSpSjYa4euF7/UvDTxu17wkSmoplUOT0+pmJrakJVNHnIx+j3zLYSooeHYdPJ3ATCREZJpiWeZqZqh5XM3mWoxESYrLBXZV/oesSjVX+lqgKdhr+ExZB1Wjbc+2nbTWwxFzACADIJbD9okjfgzMZ9FU7H0W3q+I3IBBHPnVlYOblIP8Ar6lhcqKhcWmZKRcagm4OhNxce/UzXFF1sw7GUfVtNVVDMqT3ztOT2WTttnmJA75Utiogo7zcd4ltKNfaXa4xI9mtkBk1mRw7xVQ5k5X0XWM/RX1QS1ma1JI4SZx7MthJTT7jECh6La7xMNDrfeW2hKUhKn1uJJUorPgeOHJemOEFdA8TvDmHstajldXQVOT6mpHGxLVOVDK46Fi5iiEXLHlrQ3BxDMtioKaS5K3JbGQsd95TDQcXDMOuSNNgMF8OgeYmAe85n4duFRpJAJ+8rMLVuZGX9A/dBXFbUpR338PmXip6hlMlMaIXujEmDMyioYxKYfv2EulvmKFOJulIWgH0tPEfkAtRCM68qypNyU/6/pXm03sPxYG4ta299gSDbGH2jtTU9MK94MqjqeQR1QUbHxsRUFRUvCQRmMym0ijJvl/GTKSw8sW9DmNi4yXvxMGmELyFvFfJcALxpwma3AM860E8FNeKWtbaEOO5EwUYgKdc7tKlLVOn1IRzcqlGx5Ug3WSLmIZBg7kbdts7g57j9N5LM9L5hATWChJpLIyHj5dHwkNHQEfBPtxEHGwcU2HoaLhX2Stl6HfaPesutuLS42tC0kbn1CrszsuqBeg4euK7o+kX49kvwTNTVHKZI7Fw7TiWXYiHbmUXDOPMNurQ2t1lDiGlkBdr3T7BT8BLJRJpPKZPL2ZRJZdKoCBlMqh20w0PAS6ChUMQsDDwyVLTDtQUL93ZSyFKS2lIQlRSEgYmOOxVFOcYHCZDZkrkCKAdgI//AFWapfhIanVSxU5ig4ibvxrjUMiFU53N1vOtIS5yEr1PKYzU6BtAxvmO4zv8Jx6RcYBzB4+Y/ZZLZJnnkxUscxLKezZy4nUxilBLEBLK0p2OjXTexSxCw8xcefWo2CUtIUsn2UtrVdJ9sqitqRoiCYmNZ1PIaTgIqMRL4aPqKbQEngomOW29EIg2IqPfh2HYpyHhn3EMNOOOcjSyUgg2xE8UUn7OtvJiuTl+9lEnMD8LcXR5yrmksjZ65UqHG3ZeHWZDHRTCZUp5BXN3JyhLTcu75xhxEa1C42l4wxmPLez14bGc0jNBWcPXcqZjBPS8qeIlzVL1+KeVN/vJMZ+ItyJMAy+Y20Z3jREVzRHfrXIUyTgxHcAHbGN9jJMA4PKks7s6qSR05KYifT+cSqSySEaS/EzabzGElssh2V8pQ+9Mot1uEabUFhKC4tAdUlXdrJsg7ZSDiKyGq2YJk9N5yZZzybLdDLcultaSCIjHnwSC1DsIji5EqNykdwHLqHMkK5bY4E9oFlRmvmFTWQVQUrSM7zTy8omaQkzzHywkEREiOnaXWZI/DRqoGCWY2OCYGCm0o7+EYiY6UmZiIh4V4OuuHZyU5odl9Uj8PR+YWQwyMnzbjDL0LVtARFNPNONvISpqJntNPRLog0cq24t6csw8KtpTv31SEL74staCTq3kDMRtJmczOxwizdKdbab7xbrbbaUlxbq//TQgElayoLsQ2k3UrRKh7QUlJAxs7E8R2QUDMTKIzObK6GmSXjDqgna8ptEQ28LjuVo/EF926n+ptxKHB1bSAojH5x01VNa5zB4ZOFCi6nXSuXudiYV2o5jTrkMYSa0euMhJZLJZBqh1rg4+TtyxuNealkO4uXTN1cpaiW4uAHKrlTKOA/hKlMlYkqclKQmAZh0sKms3YfmVQultIS5EuT514zJqJDg57wkTBssrSO6bZbTy4pkNDnkguAIAkjcYIgDvMGI3IlQ/NjVBBMgdp/iPn5rlDF1bTUDTr1XxdQSOGpeGg/xCIqJ+awTMjZgLhP31ycKfMvTC8y0HvzEd2AtA5jzcw8U5mJQbdHKzFNZ0uKG+7GKNZKnUtFLiFEUYExQnaY1ctVDpjUKhi6mL7kxCVMKebUOY8O84ckqEyK4Ms+KJy8hprCSKKpapp2GprN46cLTFx7zDy/uginnGmIVIbZQ21CQzAdSyh6KcffU685xxiA5/8H9bLaUhZy5iWgk2KVKTmhEtcpsQFJUQW1DmAWLp5gDzYqQXAQYl0SOR7v8A7cn0jPCmsuNP1BJKqk8BUFOTWWzyRzWGRGSybymNhZjLphCulXJEQkbBvPw0Q0opUkOMurQopI9lwLbR+GKrWkIOpISjYqqafh6tj4YxkDTD83gWZ/FwYbinVRMLKVvpjX2EtwUY6p5tlSEoh3CSAkkYH+GLM3NPggpjKSoKxYnFW8KWeMip+frmbEOIl/LCs5+33k7Uyy046ppkxbcRHx0uK0/j8uU7HSttufwkbDRu9tZ1NAVL2rfDvOqcmsPN6cm+T4mUsmsBFNRUvmEvj6MzYiIaJh4tlxcO+y604yvvW1rbOhuORXLIscSRqMRIc3c+RHrjHrHIcgSBPdZklOoaSt1aghttC1rW4QlCUIBWpS3FEJQgJSSpxR5UJuo3A19To/MWhcw4aNi6ErGmKxhZXEIhJlFUxO5bO2IGMWyl0QkW5LIuJahX1IUlxLbrvOW1E8ighwo4RdofnPP6JytlGT+Wq46Kzfz8m7NCUxAyltxUzhZJHvMQdQTNnux3sK8/97hZJBrSC80/NYqJY5BLHFtcUaPo2L7OPiOysYEXMHsjM/KPpHL+sIyKiVuQsmzQkEIxB/jr5dAYhmXJtFuTKFHeciZPPahag1KTJ2UIqAS0uyDmJA2ETg+edyMZKg50ERBmOe8dtscnv5FZq5nOJZJZfFzWdzGBlErg2lPRUwmkXDy6ChWUpJU9ExcU4hhlsEKSVvONIStKkhTiuW+0MBxMcPUymSZPLs7crIqZLdDTUE1XNO/eHXSogIhw5HtJilFRASmGLilEiySbDHETtNMqc280MpKRTljJJtWUDTNZpnFZ0JJYhxEzqGV/hy2YN6FhGOd+cOSiMU4sQLDUVFJMaiZQUJFLgmEq4jSHOPs0qobby/ze4Z5rw+1J3CYWIh60oCPlsc0+tHIt1upadXEz9xaShC3I+bwMIkKCVRfMFK5IMBI1ZOw0gzwBkGPX6lJ94ZxHcQPXOT6ftnPCy6h5tt5paXGnUIcbdQedtaFAKQtCwSFIWlSQCLWUQVaEYuEiytdTcHT5jy18Tj1iixSjVI0w1Qn4QqjESCUN0qqQrYXJVU8mBY/B1SxyHKmVwK4JLCoRaFFLjXIpJUTj2jVQIIsbEDW9zY2H6DEhPP39/FTHz/v94QCwA8MQvUAeKh++A0ABIv4XHu+WLYIqo0BHgo/ti2IFulvd44nBExFtb+VrfPE4fzb36/wfhgii2pPjb5YnDC42vqNx4X2wRMMMMU0jt9iP4+5KJhgQDuAfXCwGwAxaeADjt/b9kS/8fH/OGItvYdQfW1tfgLe70uxFFOJBt8CPiLYjDBEtv5m/yA/bDDDBFBGh6Xv8+uAFut/r6/fE4YImBIAJOwFz6DDA6gjxwRQRcEdLeH79PMG4UnmSQQSRjdzc4Z85aEz6jeJvhamsldqKqYcwGZ2W9WRCoOTVW0pmG5o2XxfPDQ4fioiDg4uIZfjZQqFmUOZhAzMuPxEtXkj+flpr8fo64W387X8TbaxGoIJuDpY3I1N8VBLcgAyIM5x5djjy78oscE1zB7ResoVUgp/IrLDKiOiVIhn61n9dQVRQ8C26tIejoKUMuRDiS22FuMoegJx7fIgN3s6j9nFBwe1nnZD5X5o0rW8hp7iXyihIQy+rIOUxUqpmqnIWJTNG5W7DOR05ipPCwM1+9PyeI55wiHbmU1gY+EXCTFaIPIiUJO6QdQSbC6iAU+0dzYW0Nwd9wMVKBfS9rXsSVXUCTcAnQ66EWIN7HXSReRBaxocOd+xxMxgRvjvhFjTazc7ReFhWqei+F7L6bzxLaWHK1YryChpE+62AgzBckbmySrn/APU+7tTiCuUlKEMcyUI9u4WeFmtqKr6ruIjP2ooKqM8q5aiIJTEqWHpFSMlilsJXAwThYhkPRTsJDQ0A2llhMDKpZCIl8I7GLcjI6I5/cg9nQ2CbHUi2lhb/AKkXP5bEeI0wKNbpsDrqCRrblJ06kfmOhUbKUSoAgXmAA1rdpI742x/E7IuA3DjkhmDl9xR8V2Y1TyNEupPM6dwUZSE0/EZZFqmcMmcTaLcWqBhIuIjYPkbfaIRHw8MFBdkKXy2OxHG5wDzGt4yaZo5AQCIarqldhIbMSgoCIYlEurFQifvLdSw6oiPgpZDzWEjGoSImMC6luEmq0ibNuNTqFUmbZcOS4AsPZ1V0vc29+gAsbDy6m2iQAASNRYC4CSCCLbWIJFthc2GptX2jgZHl5bY2g8bZOfLCoQDE8bfOfr/K27yhkUypnKbLGmpzDmDm9P5eUTI5rCqWy6YaZSmmZVAR7HfQ6lsPBqLh3kB5la23QOdC1AgniDkVkXmLRXGfxMZtz+QpgaFzClsND0rORMZW+qZvIjpFFOIVL4aNemUFZEFEm0dCw1y2ClKisWyBC2wFh6db6nXqbcxOpKjzKuSbzYg2HLyjwFibaAk6FWl7FWoGlvCgefeBJg+hI22xHH8RsqETpiAG/wArhvx1ZVVtnJkJMaKy9kyZ9Ub9UUxNGJcqPl8tSqFl0W89Eu/eppFwcIktDkUkfeA7chLSVlXKfQ657PnI+vslYakoGgKVy7zG/wBMyJbFaU5JYCEm8BVkugIZb6pjFwIQqay+PmCH2Jwwt51EVDxD7zLiYlEM8zkEKUk6gGxBFx/1N0/A7a6Yt6aeY092htbpoNsVLyAA2QOeJ+W88/c00y4kxEbfL+O/cR24QcO1D5l1Tw+z3Ibicy2hJEmAlEVQQj5PMJA/J6xpKJhXIeHmcEqRRrj0snMEoPJW/EwkveiSiXTZCGoyKioaG2GoXLnjQ4RYCPy/yupikOIfKKGmEVMaTYmc/bo6rZE1Hvqi4yCeTEPw8O0VRXNEdxCNTWDjYhTkbCIlDcS9L05VrDXxIKb3NyFG5AN7gXJJtqSSCCFHDlFthubAC1r+HhobadLjQG2KF5me8b98T9/upQOw+SxPVjlNxhcY0VI6Wzvp+l8hMmIGZsTmdyKRVC1UFVVC9BptDQzi4RyKhXQyOb7s5MUSqDl8QoTAwU0jIWAQ37vx38Ndb5i5R5N0RkhSLU4/20q6AiWpL+LSiUtQMhldKx0rgSYmczCAbdHeIgmF92pcUoPreAQRzIyVFIIN7gkEXSSg2OhsUlJFhsQbpsCNbECATcCxve4JBuVcxvYjUnU6WIJGxUMUD3A+7AERpgadxxG8DvwOEIBgGYH3nlY7WM4+P9lDTf8A4SKKcbbQ2gBebUgZX7KUpuFpmb1jpbRVz0N7Y9B4mMt+JfMeq+FbOOmMpJdMKyyyQqqasov/AFhIIeXyuoUTeTTMU8JxGTGEXGwrv4cltcbCtxPslRC0qsFZURoBYAbDTTQXPTztp773vevKLAWGmgAJAA09kDblT0TblFvZABxIPAcHBjQR5fPn5Zwgnnf+P2j7CxwtZ19oSHA25wfUZ3abDvTnFT6EoSNOZSRGqJIABUlSlK2BJO/PCiY6pplSVNzCtJLD05V0dJpdFVLIIOLbmEJJp09CMuTGWw8xa5mo1qDilOwzMSlxYebZQuyblI9v5RYgAAH138dCNfPfzxQIAuLezpbptfoNL6m53J1NyTijnBwG8j5eaqvwzBpT8vjmGhd12FiW20kgcynWVoSLkgDU7kgeeOEfADktmDkdlHU9MZkSRuRTuZ5jzioYSEbmMsmgclcRKKegod5URKYyNh21Kfl8U33S3g8kNhSkJ5gBzt5QNUi1ttdfr3G3zxOt73Pha5CRbayfyjUkmwSfPU4oHEAgHB3+/vYIsanGflDnZVedfDrmtk7QULXr2Urs1mkdBRVQySQQzsauay+OgIKIemcxhIhDUUiFfT30GzEKbHNqCShXmP8AeLtAWwo/+EehHbXUEjN6UpUoJNynmMS4QV2IutCrE3XYa4yI8qemnmNFW8ObexsCU35TYXBxcWGutwb76338DrfW5PzxUPIbpAAERtnjmfJFj64oeFir8/pZlRnFRMyl+WHEnlfByycyl2JWqOk70W4IOaRdIzSZwiYpQg5TOvvv4ZNm4WPhQIiYMxUKuEmjzsJ4KHzg7Q5mHZkEZwuZfxc9LaWV1ixmFLEU8tfLyKjxJvxTvgn8yxDGbw/eBQ/40X5DkgsLAWOm1zciwAuDqUkgWJBuQTc6m82HQAX0sBZOhuCUiw0OoNrjoRgHkANLWujYuEnjzHb48ouBWZGVvEfmPwaZp5d5jR1K1lnBWMmi0yuXUs3DSaXQ4enUFM4KRuzWYvwEvjXYKHbdhUR5Yl8MiHh4VgvTKIS7MYrcfITI1qE4Ssvsjc3aZhXi1QTFO1jTMVEQ8cw08t6IiXGURcA+/DKfhXXWoiGjoN9S2Iptl9l4OMMrb5XdCLXv4+82/wDpJOo2NyCLE3WAvYWub/zbwHgkWAFgLADFRUMRAwZEY7dvT9d5VIBzAWG7JbgezZyQ4v6In8I8/V2R9HwNQsUvUsdOoVUdTsqnUtqNyEp6Lk8XEMRqIiHnk2jVOuSaCipbGCMRNFfdYuIjYKC5PzHJHMCJ4+5Dni3I0Ky7gMt4un3p8qYSxK25o7IZrBCHEt++GaLBiY1lvvWoJxslSiSEp5sc8QAPEWA1ClA3Gxve99r2tfS+o0kAJNxrcBJJtcpSDyp1v7KSTZP5RuACBahqOJJncEbAYIEx2OO6oGgCOJnKxrZtcLucFA59zLie4VpvI1VJVrSoXMvLGrHWoKS1U2+3DIeioCP/AOJjvox+XwEc9BRcZAOQ82YVMoOa3ioqXq/ZFZhdohWbCpDJcisucp4p55mHcrOeVnA1JCQbRcbVEOwsqhYmJKiUBYAMJOC2B3QAU4XRkcsLW2G+hI1sADcEG4sCk7pIuLEm8FCSALAW25bpI1BskpIKUki6kj2ValQJtgHmACGujYkZGQfiPXOd1KANhCx/cTPCNWGa0xywzroKq5JTXErlPBy1LE6ZgYqEpWsjBOGMMpi2S9M4yVwkLHxU3XKH3UTJKoOax8smbLkPFJdhfEO5t9oPDw4kqeF6gY6cJbQyqqGsyJczIi+LNmP/AAf8QailIuS4IZc8h1e0oqYZ5UBvI11vcj0JHvGuh1O1tDbawxBSlQN0g3vc2FzcWJJtcmwsDe4FwDZRw9o6IMHtPG20R2VNIkEY9OczlcCOFjhZrKha6rDiFz3qSDqrPWv4dyCi/wAIcL8gpKTPvQzi5VLHizCfe3+5gZZAIcbYbhZfLZbDQMKuJWuMjYnm7U0E/M6bqCXQqAuLj5JNYGHRzJRzvRcA8w0gKcUhtJU44myluJQk251JSFY80EgEnqd+t7Cwv46AC5uQNBoAMSDY3G+Kau4EggjAiR3+QQiQR3Cw48N0g44+GvL5eXcg4cKZqSEfnsbPlzGa5j0xL3URUdBwEG7DBiDnq0FppMvbUlwhD7nOrvRzJ5z7fP8Ah14o+LWs6SieJ1NH5bZP0dMETtOXVGzFE7mtRTJK7FuYRjURMoVtMS2Pui5lETEGAgX4tuXShEU+mPZyuhI8OnLffQXtrvc3NzuepOmHInwAPiByq6XspNjrYFWvtdb4m6qScNbn8xI97gY3EwOVEMiJPMxxx9/HyWOfjQyjzgqzMnh2zCyioNmtHMoJ5HT+Llbk6k1PQin4OcU1M5bArdm0yhHUsRqJbENKXCNxBh2lDRYFl/tTnVx9c4T/AOEWkSk6FS83JAgWBIsQmLeTcp1BBBHMLHfGQnUEgAgD8trhIPTS1gNToBbU6EWxcCw1J131P677ab7C2KamwBE+ZAmOBvEfDdTXrFGxtRTSlKdmNXSeGp6qY6SS6MqGQwUY3MYOTTmIhG1TGWQ0e1dEY1BRRchkRHOrvUMoWCAo44JcUvDdVOcfEpw51Z/o+WVVlnSJchcwUTaKk5gEQC5suNUzFSiZRSH5mytoXWwzCRSFpUr2C42lOMithpvoLaknoka3JubJTqddD1JJEAm5vfWxBIt53Hx9QDe4BAP05bIdj0wPX73VHNDhBWyNN8OGQlITSFm1NZPZcSWbwDiIiBmUDR8jZjYR5tSFIdh30wXOy82oczbrZDjawlaClSQccZu0byKzIz8yZpSk8sJCioZ5KcxYGoYyDVMpZK+7lUPS9US+IfQ9NY2EZdc+8zCEbDLTheUHAUp5G1HGQYJSDcADQAAdANh4WAGgAHXAhKtDY7jobXBB63FwSm41sSNiRimsl4cSZ8oGII/f/CquK2ccTxU081RsZkRTeWdUSmVSoM1bTNXTaLl8/m0SGIFiGZlUU0iFlkHDy8svlx9U6C3Vv87jLzbLTZ4aZ60vxh8W1FLynqrhmoDLSFjZnKIuKr6eV7KqidlRlke1HKekyYOERMIIulotPLhBMomJg3IiDShCYhbzeXUi4N9TzA7kHQkjUDS2hH/UgW8cQQBc63vvc3PvvfXpfx1F9MVFTSPdDSZ53btzPA+MR2RY6c5+BeJqzJ7I+SZcVy/TucnDbIJHLst63jG1MwkzekkFK21Qc4TDNTB6WQkVHyiDmEBEQcLMkyd9gMmAj4KIjoSK8bAZt9otLWGqfm3DDl3UE3aQhl2tILMaVS2SRSmlBBjXZM7FuO8roHevMIjJYsj/AImYWFbsEZKLAXtcX8CbddtfZI0AsAANrYopIPNcGxuQCSo3GxuCbAf0i+nQDFQ8Ow4E+mQMjgRjfGcx8KaQ3IGXHJPr3/xwFwvm1D8RVb8L2atI5qOUbNs0KxkVRQ0ik1INqg4KBamCu9l0qi5tHxbUFGRDCyGUxDLLEO0wyhpUfNHC7GH0IZAZjo7PlORBkLYzFbpKKlhkImcq7r76uvImfIYEy++CVH/yC0OpX987oFXdlxC7JOQ0A2Nk+0ddgEg31sCLJJAuSkakm5FyTcJAAPUFKrm5PMkWCje+ttD4jfxxUAg8AY5PaCPgPrHmqrjLlFkzCjhcy/yVzXpqBjUQeX8spiqqfjTCx8O3FMMuIeS1Ew7j7SYmEf5YmCmEBEh2Fi22YmGfStsYx65R8DmauSXGnQFUy9EXVuTdKtVEiUVnFzOWmMkcnmlI1VBS6QTCVOxjMa1Gwc2mphVuS2XuS2KRGomDaYcxkTCwOaOwOmhFgQPdy338ALk6nTXQWsUjewvbfrrpobXGm1joCbHUjDWWnSO4yT3jiOP8yixMzLhAzC4qc+8wMzOIhFWZaUZTrMDTuTklpOo6dZn70oh4qN5Jq/MpXET78KCW0OR0UwoMRsVMp68wh9EHKENv/qzO7LPK+a0NULdJVlmvMK5YlEREUeaurVqdycT6EaC4NEbBLlUMEQ8wKXYB+JbeQqFRHOxSrohhyZWghAFgLDXQX0B36jc6q2udTc64FCTqQSfUjY+yd/zJ15TunmXykd4sGWp+No9ePl5D69yqQOw+Xpxtxvuse9NxXHZI8jsroOTUXQL2YNFqiZLXUszCqZqKm1ZSKVwbMPT8ZK5lKImIlbMVGMPtfi78fPYSNdj4Hvm0uNRT5Vs9nXB8afEnQM5yoqHhYy7o9M5MLDuVzO8wZJP0SYMRsM+uYyeDVCiIlsSpTLaRHw5ma4ZouJh2HVOKeGWkhIuSL9QNVWOuqQSbKNyVKHtKOpJIGGi9wLHYGwJOw00JI3TqVA6ix2Alo93DiZP05MnuflgqJY08kenw899/p3M7OcPuVSsksmcvcrHZt+OP0dT0PLomZhostRMe4/ER0wXCMlai1AIiot9iXNKPeMQTcOlR5k8qd5SbAnwBPhtriDoCALWFh7h9b4CxTrY3Gvw1vh67nfMqYEADsAqkiwVyi5PvFr9fdi99vMXHy/nEWSQBpYbC53N/PyO3hriqRa2+3np+Xx8Tc+eCKyRyi176k3/b3f5xFlC3tFXQA36m3nbXwGLYob81teW2pGuuvkbHbBEUogedxfqNQSLG42tqLXFx0wN+YAai+oTsCL+0dT1JOIuCSClRtb2rG506k2JtqB0tbEjSxB/MoHYbG+nXbr69baEV8VKdSfZ1tpYm3xIHwHqTpi2GCJhhhgiqFXUU7Ea+o8fDqNL31xbEWF79bWxOLT9x6fuUUXHw5b+QUSL+ex0H7jDC2vw6gbHQC+mpOGIIpww+vhhgiYYYfX18cETDDDBEwwwwRMMMMETDDDBEwwxRzrbYnT5+/wAMEVkkEqt/1SNj4k4nFeUaeRv/AJxbBEwwxX2//bbrvf3fX6YIrYYYYImAN9temIsLg+F7e/E2Hx3tobeuvn6YImGHx9/0P7YYImGGGCJhhhgifp8/L97+7Aba763tt5efrhhgiYYYYImGGGCJgfQn4fuRhgfUj4fuDgiYYYYImGGGCJhhhgiYYYYImF7frhiFbH9vn/fBFN9xaw0JJtbqB166/DD6/j+cQDe56EAWPW19/ibWOJw+/wBP4z8UUXGnnt8v5GJxpgi5ASTzHW/9JBvceV9D+mpxcXJIsQQToeoHX03+GCKfr6+tsVAsVHxNx899sWwwRQSBvgdfiPkb4jlAN+v84t9fX1+hwRMQCDscTgAfEkD00+AGKj8w9QiYg6gjxBxNwdLj61xUgAlWvy66YvoguCNdLAHS506j0uPXrtiUm4B/e/XxxBPtabAXIO9j4G2/snw+F8E2sLXGnv8A4vf3e7EdIJ1ZnB44++6KQoEgA6nYWOuJuPHEW21OhuNtx7sSBYW1O+/ifMWPn119cSRQCFba7+W2++IUL6aep6dL26/2xB/OPT+dfrxxfBFAuNL38/Tx9d/I312xQWSFJJ1PrbUem/8AbGoRv+xHXwxUJtcgnXfYn9Cf5wRUH5k+h/8A9sauI5dQdSRp0Pj4Dzt/nAm1vMgYIlib26iw+f8AOJ/tr8fo+7FQu1wQNyNQQNPPfr09+LYIqggnb8p6m3j6322OnjpcGp0sDuV30udL9SfUfVzi4Fr+JJv8Tb5YEAqvewFyPXX39B+/XBFOFxcDx292Kr1sPFQxW6QQP+pIuPM9dfC2wHvwRamGKBRIXqBa9t9d7dPIb2xKSSL/AKep/tgithhhi0/cen7lE+Hv+Xj1thiD08z+xwxBEIJtboQfd1wP7j5EYnDBFA/c/MnDqPQ/t/GJwwRMMMMETA3sbb209cMMEVDcEK02sQehPgdh1HTyxIJI1Sfda2/TXEkX9PDx+Fj6a773xI00wRMP8fX87YYYImGGKhV1hFtyRe/hbpbz8cEQ6K2vze61vl18sWIHUX8BYfudMSoWJF/ft0xBFtPDS/jbqcEUAC2w13ulOnodf2xJAO4B9cMMESwGwAwwwwRMMMMETDDDBEwwwwRMMMQBbqT664IpJA3IHrhhhgiYYYYImGGGCJhhhgiYYYYImGGGCJhhhgiYi2txqADqNRuOouPK/jpvicCbeG4GmgFzbYaaD3fsRNLai/gOt+h100Ou+KkqvzG978u+gSTvqonx0vsRp0xbqfIkfA2v77XxB0BPgL4IgIOxB9DibXufDf34qna+mvQAAaX266+ZPlibb20J6/AbegtgiG9tBc33000Ouv7YqfZNzewuVHxJFr2Hw8hvYamxF+pHppiOXxJPkTpgisbgG1r/AK33G/hfr/GIAsCep1IuN7XsPQ3BNyP3JVzC/u8fDBR5Re19bYIhOnXa+munWx20+PhioIA5rqtbrr18Brf3YsNQL+Xh62NgL+HpiCALnoAfZ/p28Pn64IpNyNN9COngevlgVAbn9cQnUD3/AKnFVGxKrXsm9j6/VvPocVbuPUfqissEjTx/nCxunyBv8B+/6Yjmsoje5Gt/G36Yvi+ioNCq/U6dfHwv44tzDQX31H17sQDzC+3T9D4ftiEgGyhcHUWvpbXyv1J364Ir4YYYIhNydNNST4bnbTrb9hfEcw1N9BbXXqBb9bYki4I8RbFEjfXY2t0uABcj5+XjhsiW1uOa25tt62vex66dbYlO6vAkdbnruNx6b9OmLW0t+mlsQAB/PU+uCKoSQLlNyk2BuNjc6a+vvAxfEE2IHjf5YnBEwwwIv4j03wRQCDe3TfE/XwwAt/PU28cVubkeBT8zrgik63AJ0IvY28bb6YhJFgLnW9r6nTxOG3MfrQDBIukHwJHxv/GH380UnTW19QOv7A4m9/8ABH6jEX1t5X+dsTgiYWvv10xRI1JufzEW8fq+mJSrmJ0tYet/kNtfji0/cen7lFa2mltin3HQ2v8A/wDXhhgdbbixvvv5emGIIv/Z"""


def garantir_logo_aqua_oficial() -> Path | None:
    """Garante um arquivo de logo oficial local para uso nos PDFs premium."""
    try:
        candidatos = [
            BASE_DIR / "Logo Aqua.jpeg",
            BASE_DIR / "Logo_Aqua.jpeg",
            BASE_DIR / "logo_aqua.jpeg",
            BASE_DIR / "aqua_gestao_logo.png",
            BASE_DIR / "aqua_gestao_logo.jpg",
            BASE_DIR / "assets" / "Logo Aqua.jpeg",
            BASE_DIR / "assets" / "aqua_gestao_logo.png",
        ]
        for p in candidatos:
            if p.exists():
                return p
        p = BASE_DIR / "Logo Aqua.jpeg"
        import base64 as _b64
        p.write_bytes(_b64.b64decode(LOGO_AQUA_OFICIAL_B64))
        return p if p.exists() else None
    except Exception:
        return None


def _rl_valor(valor, padrao="—"):
    """Valor seguro para células do PDF: evita vazios aparentes e caracteres problemáticos."""
    texto = str(valor or "").strip()
    return texto if texto else padrao


def _rl_f(valor, padrao="—"):
    texto = str(valor or "").strip()
    return texto if texto else padrao


def _limpar_texto_pdf(texto: str) -> str:
    """Remove marcações HTML e caracteres que costumam quebrar na fonte padrão do PDF."""
    texto = str(texto or "")
    texto = re.sub(r"<[^>]+>", "", texto)
    texto = texto.replace("CaCO3", "CaCO3").replace("CaCO3", "CaCO3")
    texto = texto.replace("■", "□")
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def _pegar_alias(item: dict, *chaves, padrao=""):
    for chave in chaves:
        valor = item.get(chave, "") if isinstance(item, dict) else ""
        if str(valor or "").strip():
            return valor
    return padrao


def _rl_linhas_analises(dados_relatorio: dict) -> list[list[str]]:
    """Monta linhas da tabela de análises com aliases e traços para campos não informados."""
    linhas = []
    for item in dados_relatorio.get("analises", []) or []:
        if not isinstance(item, dict):
            continue
        data = _pegar_alias(item, "data", "Data")
        ph = _pegar_alias(item, "ph", "pH")
        crl = _pegar_alias(item, "cloro_livre", "cl", "crl", "CRL")
        ct = _pegar_alias(item, "cloro_total", "ct", "cl_total", "CT")
        alc = _pegar_alias(item, "alcalinidade", "alc", "AT")
        dc = _pegar_alias(item, "dureza", "dc", "dureza_calcica", "DC")
        cya = _pegar_alias(item, "cianurico", "cya", "acido_cianurico", "CYA")
        operador = _pegar_alias(item, "operador", "responsavel", "Operador")
        if any(str(v or "").strip() for v in [data, ph, crl, ct, alc, dc, cya, operador]):
            linhas.append([
                _rl_valor(normalizar_data_visita(data) if data else "", "—"),
                _rl_f(ph),
                _rl_f(crl),
                _rl_f(ct),
                _rl_f(alc),
                _rl_f(dc),
                _rl_f(cya),
                _rl_valor(operador, "—"),
            ])
    return linhas


def _normalizar_dosagem_pdf(item: dict) -> dict:
    """Aceita nomes alternativos de campos de dosagem."""
    if not isinstance(item, dict):
        return {}
    return {
        "produto": _pegar_alias(item, "produto", "produto_quimico", "nome", "Produto"),
        "fabricante_lote": _pegar_alias(item, "fabricante_lote", "fabricante", "lote", "marca", "Fabricante / Lote"),
        "quantidade": _pegar_alias(item, "quantidade", "qtd", "Quantidade"),
        "unidade": _pegar_alias(item, "unidade", "unid", "Unidade"),
        "finalidade": _pegar_alias(item, "finalidade", "finalidade_tecnica", "motivo", "acao", "Finalidade"),
    }


def _rl_linhas_dosagens(dados_relatorio: dict) -> list[list[str]]:
    """Monta tabela de dosagens a partir do formulário, das visitas importadas ou de resumo salvo."""
    linhas = []
    vistos = set()

    def add_dosagem(d: dict):
        d = _normalizar_dosagem_pdf(d)
        if not any(str(v or "").strip() for v in d.values()):
            return
        chave = tuple(str(d.get(k, "")).strip().lower() for k in ["produto", "fabricante_lote", "quantidade", "unidade", "finalidade"])
        if chave in vistos:
            return
        vistos.add(chave)
        linhas.append([
            _rl_valor(d.get("produto"), "—"),
            _rl_valor(d.get("fabricante_lote"), "—"),
            _rl_valor(d.get("quantidade"), "—"),
            _rl_valor(d.get("unidade"), "—"),
            _rl_valor(d.get("finalidade"), "—"),
        ])

    for item in dados_relatorio.get("dosagens", []) or []:
        add_dosagem(item)

    # Dosagens que vieram dentro de lançamentos/visitas/piscinas
    for item in dados_relatorio.get("analises", []) or []:
        if not isinstance(item, dict):
            continue
        for d in item.get("dosagens", []) or []:
            add_dosagem(d)
        resumo = str(item.get("dosagem_resumo", "") or item.get("dosagem", "") or "").strip()
        if resumo:
            add_dosagem({"produto": resumo, "finalidade": "Registro importado da visita"})

    for lc in dados_relatorio.get("lancamentos", []) or []:
        if not isinstance(lc, dict):
            continue
        for d in lc.get("dosagens", []) or []:
            add_dosagem(d)
        for p in lc.get("piscinas", []) or []:
            if isinstance(p, dict):
                for d in p.get("dosagens", []) or []:
                    add_dosagem(d)

    return linhas

def gerar_pdf_relatorio_rt_premium_reportlab(dados_relatorio: dict, fotos: list[Path] | None, pdf_path: Path) -> tuple[bool, str | None]:
    """Gera PDF premium no padrão Aqua Gestão, sem depender do Word/LibreOffice."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            Image as RLImage, PageBreak, HRFlowable
        )
        from PIL import Image as PILImage

        pdf_path = Path(pdf_path)
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        azul = colors.HexColor("#173A5E")
        azul2 = colors.HexColor("#1D78A8")
        azul_claro = colors.HexColor("#D9EAF4")
        cinza = colors.HexColor("#4D5661")
        verde = colors.HexColor("#2E7D32")
        laranja = colors.HexColor("#C96B2C")
        vermelho = colors.HexColor("#B42318")
        branco = colors.white
        preto = colors.HexColor("#222222")

        logo_path = garantir_logo_aqua_oficial()

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle("AqTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=azul, alignment=TA_CENTER, spaceAfter=2))
        styles.add(ParagraphStyle("AqSubtitle", parent=styles["Normal"], fontName="Helvetica", fontSize=10.5, leading=13, textColor=azul2, alignment=TA_CENTER, spaceAfter=8))
        styles.add(ParagraphStyle("AqH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=azul, spaceBefore=10, spaceAfter=4))
        styles.add(ParagraphStyle("AqH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=azul2, spaceBefore=8, spaceAfter=4))
        styles.add(ParagraphStyle("AqBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.7, leading=11.2, textColor=preto, alignment=TA_JUSTIFY, spaceAfter=4))
        styles.add(ParagraphStyle("AqSmall", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.6, leading=9.2, textColor=cinza, spaceAfter=2))
        styles.add(ParagraphStyle("AqCell", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.4, leading=9.2, textColor=preto))
        styles.add(ParagraphStyle("AqCellBold", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=7.5, leading=9.2, textColor=azul))
        styles.add(ParagraphStyle("AqWarn", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.2, leading=10.2, textColor=azul, leftIndent=4, borderColor=azul2, borderWidth=0.6, borderPadding=5, spaceBefore=4, spaceAfter=6))

        def P(texto, style="AqBody"):
            texto = _limpar_texto_pdf(str(texto or ""))
            texto = texto.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return Paragraph(texto, styles[style])

        def table(data, widths=None, header=True):
            conv = []
            for r, row in enumerate(data):
                out = []
                for c in row:
                    if isinstance(c, Paragraph):
                        out.append(c)
                    else:
                        out.append(P(c, "AqCellBold" if (header and r == 0) else "AqCell"))
                conv.append(out)
            t = Table(conv, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
            cmds = [
                ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#AEBFCC")),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                ("LEFTPADDING", (0,0), (-1,-1), 5),
                ("RIGHTPADDING", (0,0), (-1,-1), 5),
                ("TOPPADDING", (0,0), (-1,-1), 5),
                ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ]
            if header:
                cmds += [
                    ("BACKGROUND", (0,0), (-1,0), azul),
                    ("TEXTCOLOR", (0,0), (-1,0), branco),
                    ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                    ("ALIGN", (0,0), (-1,0), "CENTER"),
                ]
            for i in range(1 if header else 0, len(data)):
                if i % 2 == 0:
                    cmds.append(("BACKGROUND", (0,i), (-1,i), azul_claro))
            t.setStyle(TableStyle(cmds))
            return t

        doc = SimpleDocTemplate(
            str(pdf_path), pagesize=A4,
            leftMargin=14*mm, rightMargin=14*mm,
            topMargin=24*mm, bottomMargin=18*mm,
            title="Relatório Mensal de Responsabilidade Técnica",
            author=dados_relatorio.get("responsavel_tecnico", RESPONSAVEL_TÉCNICO),
        )
        w, h = A4

        def header_footer(canvas, doc_obj):
            canvas.saveState()
            if logo_path and Path(logo_path).exists():
                try:
                    canvas.drawImage(str(logo_path), 15*mm, h-19*mm, width=23*mm, height=16*mm, preserveAspectRatio=True, mask="auto")
                except Exception:
                    pass
            canvas.setStrokeColor(azul2)
            canvas.setLineWidth(1.1)
            canvas.line(14*mm, h-22*mm, w-14*mm, h-22*mm)
            canvas.setFont("Helvetica-Bold", 8.7)
            canvas.setFillColor(azul)
            canvas.drawRightString(w-15*mm, h-14*mm, "RELATÓRIO MENSAL DE RESPONSABILIDADE TÉCNICA")
            canvas.setFont("Helvetica", 7.2)
            canvas.setFillColor(cinza)
            canvas.drawRightString(w-15*mm, h-18*mm, f"RT: {dados_relatorio.get('responsavel_tecnico', RESPONSAVEL_TÉCNICO)} | CRQ 024025748 | Técnico em Química")
            canvas.setStrokeColor(azul2)
            canvas.setLineWidth(0.8)
            canvas.line(14*mm, 13*mm, w-14*mm, 13*mm)
            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(cinza)
            canvas.drawString(14*mm, 8.5*mm, "Aqua Gestão – Controle Técnico de Piscinas")
            canvas.drawRightString(w-14*mm, 8.5*mm, f"Página {doc_obj.page}")
            canvas.restoreState()

        story = []
        story.append(Spacer(1, 4*mm))
        if logo_path and Path(logo_path).exists():
            try:
                story.append(RLImage(str(logo_path), width=58*mm, height=40*mm, kind="proportional", hAlign="CENTER"))
                story.append(Spacer(1, 4*mm))
            except Exception:
                pass
        story.append(P("RELATÓRIO MENSAL DE RESPONSABILIDADE TÉCNICA", "AqTitle"))
        story.append(P("Controle Técnico-Operacional de Piscinas", "AqSubtitle"))
        story.append(HRFlowable(width="100%", thickness=1.2, color=azul2, spaceBefore=2, spaceAfter=8))

        ident = [
            ["IDENTIFICAÇÃO DO DOCUMENTO", ""],
            ["Responsável Técnico", dados_relatorio.get("responsavel_tecnico", RESPONSAVEL_TÉCNICO)],
            ["Registro CRQ", "CRQ-MG 2ª Região | CRQ 024025748"],
            ["Qualificação", dados_relatorio.get("qualificacao", QUALIFICACAO_RT)],
            ["Empresa", dados_relatorio.get("empresa_rt", EMPRESA_RT)],
            ["Cliente / Estabelecimento", dados_relatorio.get("nome_condominio", "")],
            ["Endereço do Local", dados_relatorio.get("endereco_condominio", "")],
            ["Período de Referência", f"Mês: {dados_relatorio.get('mes_referencia','')} / Ano: {dados_relatorio.get('ano_referencia','')}"],
            ["Nº ART – CRQ", obter_status_art_texto(dados_relatorio)],
            ["Data de Emissão", dados_relatorio.get("data_emissao", hoje_br())],
        ]
        t_ident = table(ident, widths=[55*mm, 125*mm], header=False)
        t_ident.setStyle(TableStyle([
            ("SPAN", (0,0), (1,0)), ("BACKGROUND", (0,0), (1,0), azul), ("TEXTCOLOR", (0,0), (1,0), branco),
            ("FONTNAME", (0,0), (1,0), "Helvetica-Bold"), ("ALIGN", (0,0), (1,0), "CENTER"),
            ("BACKGROUND", (0,1), (0,-1), colors.HexColor("#EAF3F8")),
            ("BACKGROUND", (0,2), (1,2), azul_claro), ("BACKGROUND", (0,4), (1,4), azul_claro),
            ("BACKGROUND", (0,6), (1,6), azul_claro), ("BACKGROUND", (0,8), (1,8), azul_claro),
            ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#AEBFCC")),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING", (0,0), (-1,-1), 6), ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 6), ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(t_ident)

        story.append(P("1. CONFORMIDADE LEGAL – CFQ / CRQ", "AqH1"))
        story.append(P("Este Relatório de Responsabilidade Técnica é emitido em atendimento às obrigações legais do RT habilitado perante o Conselho Regional de Química de Minas Gerais — CRQ-MG 2ª Região, nos termos:", "AqBody"))
        story.append(table([
            ["Instrumento Normativo", "Atendimento Declarado"],
            ["Lei Federal nº 2.800/1956", "Regulamenta a profissão de Químico e as atribuições do Técnico em Química."],
            ["Decreto nº 85.877/1981 e normas profissionais aplicáveis do Sistema CFQ/CRQs", "Define atividades técnicas relacionadas ao controle de tratamento de água."],
            ["Resolução CFQ nº 332/2025", "Dispõe sobre Responsabilidade Técnica e emissão de ART."],
            ["Portaria GM/MS nº 888/2021", "Referência sanitária complementar para parâmetros de água."],
        ], widths=[58*mm, 122*mm]))

        story.append(P("2. NORMAS TÉCNICAS ABNT APLICÁVEIS", "AqH1"))
        story.append(P("2.1 ABNT NBR 10339 – Qualidade da Água de Piscina", "AqH2"))
        story.append(P("Estabelece os limites físico-químicos aceitáveis para água de piscinas coletivas e individuais. Os parâmetros abaixo são monitorados no local com o aparelho Photometer Color Q.", "AqBody"))
        story.append(table([
            ["Parâmetro", "Mínimo", "Máximo", "Unidade", "Método"],
            ["pH", "7,2", "7,8", "—", "Photometer"],
            ["Cloro Residual Livre (CRL)", "0,5", "3,0", "mg/L", "Photometer"],
            ["Cloro Total", "0,5", "3,0", "mg/L", "Photometer"],
            ["Alcalinidade Total", "80", "120", "mg/L CaCO3", "Photometer"],
            ["Dureza Cálcica", "150", "300", "mg/L CaCO3", "Photometer"],
            ["Ácido Cianúrico", "30", "50", "mg/L", "Photometer"],
        ], widths=[58*mm, 30*mm, 30*mm, 35*mm, 27*mm]))
        story.append(P("NOTA TÉCNICA: análises microbiológicas não são realizadas no local. Tais determinações requerem coleta e envio a laboratório acreditado pela ANVISA/Inmetro, sob responsabilidade de contratação do cliente.", "AqWarn"))

        story.append(P("2.2 NBR 11238 – Segurança e Higiene em Piscinas", "AqH2"))
        conf = dados_relatorio.get("conformidades", {}) or {}
        nbr_status = conf.get("nbr_11238_status", {}) or {}
        nbr_obs = conf.get("nbr_11238", "") or "Sem observações adicionais registradas."
        story.append(table([
            ["Requisito NBR 11238", "Evidência / Observação", "Conforme?"],
            ["Sinalização de profundidade visível", nbr_obs, _texto_sim_nao_marcado(nbr_status.get("profundidade", ""))],
            ["Retrolavagem do sistema de filtragem", "", _texto_sim_nao_marcado(nbr_status.get("retrolavagem", ""))],
            ["Limpeza de skimmers e decantadores", "", _texto_sim_nao_marcado(nbr_status.get("skimmers", ""))],
            ["Área de circulação antiderrapante", "", _texto_sim_nao_marcado(nbr_status.get("circulacao", ""))],
            ["Chuveiro obrigatório antes do acesso", "", _texto_sim_nao_marcado(nbr_status.get("chuveiro", ""))],
        ], widths=[78*mm, 78*mm, 24*mm]))

        story.append(P("3. GESTÃO DE RISCOS E SEGURANÇA DO TRABALHO", "AqH1"))
        story.append(P("3.1 NR-26 – Sinalização / GHS – Produtos Químicos", "AqH2"))
        nr26_obs = conf.get("nr_26", "") or "Checklist de GHS/FDS/FISPQ sem observações adicionais registradas."
        story.append(table([
            ["Item de Verificação – NR-26 / GHS", "Observação", "Status"],
            ["FISPQs disponíveis e atualizadas para os produtos", nr26_obs, "OK"],
            ["Rótulos GHS nos recipientes", "", "OK / Pend."],
            ["Sinalização de perigo no estoque de produtos químicos", "", "OK / Pend."],
            ["Separação de oxidantes e ácidos", "", "OK / Pend."],
            ["Procedimento de emergência / derramamento disponível", "", "OK / Pend."],
        ], widths=[83*mm, 72*mm, 25*mm]))
        story.append(P("3.2 NR-06 – Equipamentos de Proteção Individual (EPI)", "AqH2"))
        epis = dados_relatorio.get("epis", {}) or {}
        def epi_status(base):
            stt = epis.get(f"{base}_status", "Conforme") or "Conforme"
            ca = epis.get(f"{base}_ca", "") or "Conforme"
            fornecido = "Sim" if stt == "Conforme" else "Não"
            fiscalizado = "Sim" if stt == "Conforme" else "Não"
            return ca, fornecido, fiscalizado
        luvas = epi_status("luvas"); oculos = epi_status("oculos"); respirador = epi_status("respirador"); botas = epi_status("botas")
        story.append(table([
            ["EPI", "CA nº", "Fornecido?", "Uso Fiscalizado?"],
            ["Luvas de nitrila resistentes a produtos químicos", luvas[0], luvas[1], luvas[2]],
            ["Óculos de proteção contra respingos", oculos[0], oculos[1], oculos[2]],
            ["Respirador para vapores químicos", respirador[0], respirador[1], respirador[2]],
            ["Botas/calçado fechado resistente a produtos químicos", botas[0], botas[1], botas[2]],
        ], widths=[82*mm, 42*mm, 28*mm, 28*mm]))

        story.append(P("4. CONTROLE OPERACIONAL – MONITORAMENTO IN LOCO", "AqH1"))
        story.append(P("Análises realizadas com Photometer Color Q. Todos os resultados confrontados com os limites da ABNT NBR 10339.", "AqBody"))
        story.append(P("4.1 Registro de Análises Físico-Químicas", "AqH2"))
        linhas_a = _rl_linhas_analises(dados_relatorio)
        if linhas_a:
            story.append(table([["Data", "pH", "CRL (ppm)", "Cl. Total (ppm)", "Alcalinidade (ppm)", "Dureza Cálcica (ppm)", "Ác. Cianúrico (ppm)", "Operador"]] + linhas_a,
                widths=[22*mm, 16*mm, 19*mm, 22*mm, 27*mm, 27*mm, 28*mm, 19*mm]))
        else:
            story.append(P("Nenhum lançamento físico-químico foi encontrado para o período de referência.", "AqWarn"))
        story.append(P("4.2 Registro de Dosagens de Produtos Químicos", "AqH2"))
        linhas_d = _rl_linhas_dosagens(dados_relatorio)
        if linhas_d:
            story.append(table([["Produto Químico", "Fabricante / Lote", "Qtd.", "Unid.", "Finalidade Técnica"]] + linhas_d,
                widths=[48*mm, 38*mm, 22*mm, 22*mm, 50*mm]))
        else:
            story.append(P("Nenhuma dosagem foi informada para o período de referência.", "AqWarn"))

        story.append(P("5. PARECER TÉCNICO", "AqH1"))
        story.append(P("5.1 Diagnóstico da Qualidade da Água", "AqH2"))
        status = dados_relatorio.get("status_agua", "") or "EM CORREÇÃO"
        cor_status = verde if status == "CONFORME" else (laranja if status == "EM CORREÇÃO" else vermelho)
        status_tbl = Table([[P(f"Status geral da água: {status}", "AqCellBold")]], colWidths=[180*mm])
        status_tbl.setStyle(TableStyle([("BOX", (0,0), (-1,-1), 0.8, cor_status), ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F7FBFD")), ("LEFTPADDING", (0,0), (-1,-1), 7), ("TOPPADDING", (0,0), (-1,-1), 7), ("BOTTOMPADDING", (0,0), (-1,-1), 7)]))
        story.append(status_tbl)
        story.append(Spacer(1, 3*mm))
        diagnostico_txt = _limpar_texto_pdf(dados_relatorio.get('diagnostico', ''))
        # Evita repetir o resumo automático dentro do diagnóstico e depois em lista.
        diagnostico_txt = re.sub(r"Resumo automático dos desvios:.*", "", diagnostico_txt).strip()
        if diagnostico_txt:
            story.append(P(f"Diagnóstico: {diagnostico_txt}", "AqBody"))
        detalhes = dados_relatorio.get("avaliacao_automatica", {}).get("detalhes", []) or []
        if detalhes:
            story.append(P("Resumo técnico dos desvios identificados:", "AqH2"))
            for d in detalhes[:10]:
                story.append(P(f"• {d}", "AqSmall"))
        obs = [o for o in (dados_relatorio.get("observacoes", []) or []) if str(o or "").strip()]
        if obs:
            story.append(P("Observações técnicas complementares:", "AqH2"))
            for o in obs[:8]:
                story.append(P(f"• {o}", "AqSmall"))

        story.append(P("5.2 Recomendações Técnicas ao Cliente", "AqH2"))
        recs = []
        for idx, r in enumerate(dados_relatorio.get("recomendacoes", []) or [], start=1):
            if str(r.get("recomendacao", "")).strip():
                recs.append([str(idx), r.get("recomendacao", ""), r.get("prazo", ""), r.get("responsavel", "")])
        if not recs:
            recs = [["1", "Manter rotina de monitoramento e registrar as leituras no sistema.", "Próxima rotina", "Operação / RT"]]
        story.append(table([["Nº", "Recomendação Técnica", "Prazo", "Responsável"]] + recs, widths=[13*mm, 100*mm, 32*mm, 35*mm]))

        story.append(P("6. ASSINATURAS E VALIDAÇÃO", "AqH1"))
        assinatura = [
            ["RESPONSÁVEL TÉCNICO", "REPRESENTANTE DO ESTABELECIMENTO"],
            [f"\n\n{dados_relatorio.get('responsavel_tecnico', RESPONSAVEL_TÉCNICO)}\nCRQ 024025748 – Técnico em Química\nData: ______ / ______ / __________", "\n\nNome: _________________________________\nCPF / CNPJ: ___________________________\nData: ______ / ______ / __________"],
        ]
        story.append(table(assinatura, widths=[90*mm, 90*mm], header=True))
        story.append(P("Documento de uso profissional. Emitido sob responsabilidade técnica do RT identificado neste relatório. Análises microbiológicas não são realizadas no local e dependem de laboratório acreditado, sob responsabilidade de contratação do cliente.", "AqWarn"))

        fotos = [Path(f) for f in (fotos or []) if f and Path(f).exists()]
        if fotos:
            story.append(PageBreak())
            story.append(P("7. REGISTRO FOTOGRÁFICO", "AqH1"))
            story.append(P("Registros visuais do período de referência, inseridos para rastreabilidade técnica e comprovação documental.", "AqBody"))
            for idx, fp in enumerate(fotos[:12], start=1):
                try:
                    legenda = "Vista geral da piscina no período de referência" if idx == 1 else "Registro complementar da área da piscina"
                    story.append(P(f"Foto {idx} — {legenda}", "AqH2"))
                    img = PILImage.open(fp)
                    iw, ih = img.size
                    max_w = 170*mm
                    max_h = 175*mm
                    ratio = min(max_w/iw, max_h/ih)
                    story.append(RLImage(str(fp), width=iw*ratio, height=ih*ratio, hAlign="CENTER"))
                    story.append(Spacer(1, 5*mm))
                except Exception:
                    story.append(P(f"Foto {idx} — arquivo não pôde ser inserido no PDF.", "AqSmall"))

        doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
        return (pdf_path.exists(), None if pdf_path.exists() else "PDF premium não foi criado.")
    except Exception as e:
        _log_sheets_erro("gerar_pdf_relatorio_rt_premium_reportlab", e)
        return False, str(e)


LIMITES_RELATORIO = {
    "ph": (7.2, 7.8, "pH"),
    "cloro_livre": (0.5, 3.0, "cloro residual livre"),
    "alcalinidade": (80.0, 120.0, "alcalinidade total"),
    "dureza": (150.0, 300.0, "dureza cálcica"),
    "cianurico": (30.0, 50.0, "ácido cianúrico"),
}


def avaliar_conformidade_analises(analises: list[dict]) -> dict:
    nao_conformes = []
    houve_leitura = False
    cloraminas_altas = []
    for idx, item in enumerate(analises, start=1):
        linha_tem = any((item.get(k) or "").strip() for k in ["ph", "cloro_livre", "cloro_total", "alcalinidade", "dureza", "cianurico"])
        if not linha_tem:
            continue
        houve_leitura = True
        for campo, (mn, mx, rotulo) in LIMITES_RELATORIO.items():
            v = valor_float(item.get(campo, ""))
            if v is None:
                continue
            if v < mn or v > mx:
                nao_conformes.append(f"Linha {idx}: {rotulo}={v} fora da faixa {mn}–{mx}")
        cl = valor_float(item.get("cloro_livre", ""))
        ct = valor_float(item.get("cloro_total", ""))
        if cl is not None and ct is not None:
            combinado = round(max(ct - cl, 0), 2)
            item["cloro_combinado"] = combinado
            if combinado > 0.2:
                cloraminas_altas.append((idx, combinado))
        else:
            item["cloro_combinado"] = None

    status = "NÃO CONFORME" if nao_conformes else ("CONFORME" if houve_leitura else "EM CORREÇÃO")
    return {
        "status": status,
        "detalhes": nao_conformes,
        "houve_leitura": houve_leitura,
        "cloraminas_altas": cloraminas_altas,
    }


def normalizar_texto(texto: str) -> str:
    texto = (texto or "").lower()
    mapa = str.maketrans("áàâãäéèêëíìîïóòôõöúùûüç", "aaaaaeeeeiiiiooooouuuuc")
    return texto.translate(mapa)


def limpar_paragrafo(paragraph):
    for run in paragraph.runs:
        run.text = ""


def set_cell_text(cell, texto: str):
    texto = texto if texto is not None else ""
    if cell.paragraphs:
        primeira = True
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""
            if primeira:
                if p.runs:
                    p.runs[0].text = str(texto)
                else:
                    p.add_run(str(texto))
                primeira = False
    else:
        cell.text = str(texto)


def encontrar_tabela_por_keywords(doc: Document, keywords: list[str]):
    kws = [normalizar_texto(k) for k in keywords]
    melhor = None
    melhor_score = -1
    for table in doc.tables:
        conteudo = " ".join(normalizar_texto(c.text) for row in table.rows[:6] for c in row.cells)
        score = sum(1 for k in kws if k in conteudo)
        if score > melhor_score:
            melhor = table
            melhor_score = score
    return melhor if melhor_score > 0 else None


def preencher_tabela_generica(table, rows_data: list[list[str]], start_row: int = 1):
    if table is None:
        return False
    total_cols = max((len(r.cells) for r in table.rows), default=0)
    if total_cols == 0:
        return False
    for idx, row_data in enumerate(rows_data, start=start_row):
        if idx >= len(table.rows):
            table.add_row()
        row = table.rows[idx]
        for c in range(total_cols):
            valor = row_data[c] if c < len(row_data) else ""
            set_cell_text(row.cells[c], valor)
    return True


def preencher_tabela_identificacao(doc: Document, dados_relatorio: dict):
    tabela = encontrar_tabela_por_keywords(doc, ["Responsável Técnico", "Registro CRQ", "Cliente", "Período de Referência", "Data de Emissão"])
    if tabela is None:
        return False
    mapa_linhas = {
        "responsavel tecnico": dados_relatorio["responsavel_tecnico"],
        "registro crq": dados_relatorio["crq"],
        "qualificacao": dados_relatorio["qualificacao"],
        "empresa": dados_relatorio["empresa_rt"],
        "cliente / estabelecimento": dados_relatorio["nome_condominio"],
        "endereco do local": dados_relatorio["endereco_condominio"],
        "periodo de referencia": f"Mês: {dados_relatorio['mes_referencia']} / Ano: {dados_relatorio['ano_referencia']}",
        "n° art – crq": obter_status_art_texto(dados_relatorio),
        "nº art – crq": obter_status_art_texto(dados_relatorio),
        "data de emissao": dados_relatorio["data_emissao"],
    }
    for row in tabela.rows:
        primeira = normalizar_texto(row.cells[0].text)
        for chave, valor in mapa_linhas.items():
            if chave in primeira and len(row.cells) > 1:
                set_cell_text(row.cells[1], valor)
    return True


def _texto_sim_nao_marcado(valor) -> str:
    """Converte bool/texto em marcação visual para o relatório."""
    v = str(valor or "").strip().lower()
    if v in ("sim", "true", "1", "ok", "conforme"):
        return "☑ Sim   ☐ Não"
    if v in ("não", "nao", "false", "0", "pendente", "não conforme", "nao conforme"):
        return "☐ Sim   ☑ Não"
    return "☐ Sim   ☐ Não"


def _status_nbr11238_por_texto(texto_requisito: str, dados_relatorio: dict) -> str:
    """Retorna status Sim/Não para cada requisito da NBR 11238."""
    status = (dados_relatorio.get("conformidades", {}) or {}).get("nbr_11238_status", {}) or {}
    t = normalizar_texto(texto_requisito)
    mapa = [
        ("profundidade", "profundidade"),
        ("retrolavagem", "retrolavagem"),
        ("skimmers", "skimmers"),
        ("decantadores", "skimmers"),
        ("circulacao", "circulacao"),
        ("antiderrapante", "circulacao"),
        ("chuveiro", "chuveiro"),
    ]
    for pedaco, chave in mapa:
        if pedaco in t:
            return _texto_sim_nao_marcado(status.get(chave, ""))
    return _texto_sim_nao_marcado("")

def preencher_bloco_conformidades(doc: Document, dados_relatorio: dict):
    tabela_nbr = encontrar_tabela_por_keywords(doc, ["Requisito NBR 11238", "Evidência / Observação"])
    if tabela_nbr is not None and len(tabela_nbr.rows) > 1:
        observacao = dados_relatorio["conformidades"].get("nbr_11238", "") or "Sem observações adicionais registradas."
        for idx in range(1, len(tabela_nbr.rows)):
            row = tabela_nbr.rows[idx]
            if len(row.cells) > 1:
                # Mantém a observação geral preenchida em todas as linhas em branco,
                # evitando relatório com campo vazio quando o modelo tiver várias linhas.
                set_cell_text(row.cells[1], observacao if idx == 1 else "")
            if len(row.cells) > 2:
                set_cell_text(row.cells[2], _status_nbr11238_por_texto(row.cells[0].text, dados_relatorio))

    tabela_nr26 = encontrar_tabela_por_keywords(doc, ["NR-26", "GHS", "FISPQs"])
    if tabela_nr26 is not None and len(tabela_nr26.rows) > 1:
        observacao = dados_relatorio["conformidades"].get("nr_26", "") or "Checklist de GHS/FDS/FISPQ sem observações adicionais registradas."
        for idx in range(1, len(tabela_nr26.rows)):
            row = tabela_nr26.rows[idx]
            if len(row.cells) > 1 and idx == 1:
                set_cell_text(row.cells[1], observacao)
            if len(row.cells) > 2 and idx == 1:
                set_cell_text(row.cells[2], "OK" if observacao else "Pend.")

    tabela_nr6 = encontrar_tabela_por_keywords(doc, ["Equipamentos de Proteção Individual", "CA nº", "Luvas de nitrila"])
    if tabela_nr6 is not None and len(tabela_nr6.rows) > 1:
        ca_map = dados_relatorio["epis"]
        linhas_epi = [
            ("luvas", ca_map.get("luvas_status", "Conforme"), ca_map.get("luvas_ca", "")),
            ("oculos", ca_map.get("oculos_status", "Conforme"), ca_map.get("oculos_ca", "")),
            ("respirador", ca_map.get("respirador_status", "Conforme"), ca_map.get("respirador_ca", "")),
            ("botas", ca_map.get("botas_status", "Conforme"), ca_map.get("botas_ca", "")),
        ]
        mapa_status = {
            "Conforme": ("Sim", "Sim"),
            "Pendente": ("Não", "Não"),
            "Não informado": ("Não", "Não"),
            "N/A": ("Não", "Não"),
        }
        for row in tabela_nr6.rows[1:]:
            texto = normalizar_texto(row.cells[0].text)
            for chave, status_epi, ca in linhas_epi:
                if chave in texto and len(row.cells) > 1:
                    # Coluna 1 = CA nº: apenas o número ou N/A, NUNCA o status
                    ca_valor = ca.strip() if (ca or "").strip() else "N/A"
                    set_cell_text(row.cells[1], ca_valor)
                    fornecido, fiscalizado = mapa_status.get(status_epi, ("Não", "Não"))
                    if len(row.cells) > 2:
                        set_cell_text(row.cells[2], fornecido)
                    if len(row.cells) > 3:
                        set_cell_text(row.cells[3], fiscalizado)


def inserir_assinatura_rt_no_doc(doc: Document):
    assinatura = preparar_assinatura_rt_para_relatorio()
    texto_ass = RESPONSAVEL_TECNICO_ASSINATURA

    # Prioridade máxima: colocar a assinatura no fim do relatório,
    # exatamente acima do campo "Data" do bloco do Responsável Técnico.
    paragrafos = list(doc.paragraphs)
    for i, p in enumerate(paragrafos):
        txt = normalizar_texto(p.text)
        if "responsavel tecnico" in txt:
            for prox in paragrafos[i + 1:i + 8]:
                txt_prox = normalizar_texto(prox.text)
                if "representante do estabelecimento" in txt_prox:
                    break
                if "data:" in txt_prox:
                    try:
                        novo = prox.insert_paragraph_before("")
                    except Exception:
                        novo = doc.add_paragraph()
                    novo.alignment = 1
                    if assinatura:
                        try:
                            novo.add_run().add_picture(str(assinatura), width=Inches(1.45))
                        except Exception:
                            pass
                    return True

    # Fallback: insere ao final se não achar o bloco correto
    if assinatura:
        try:
            novo = doc.add_paragraph()
            novo.alignment = 1
            novo.add_run().add_picture(str(assinatura), width=Inches(1.45))
            novo2 = doc.add_paragraph(texto_ass)
            novo2.alignment = 1
            return True
        except Exception:
            pass

    doc.add_paragraph(texto_ass)
    return False


def salvar_uploads_relatorio(pasta_condominio: Path):
    """Salva fotos localmente E no Google Drive. Retorna caminhos locais para inserir no DOCX."""
    caminhos = []
    arquivos = st.session_state.get("rel_fotos_upload") or []
    pasta_fotos = pasta_condominio / "fotos_relatorio"
    pasta_fotos.mkdir(exist_ok=True)
    nome_cond = (st.session_state.get("rel_nome_condominio") or
                 st.session_state.get("nome_condominio") or
                 pasta_condominio.name)
    mes_ano = datetime.now().strftime("%Y-%m")
    for idx, arquivo in enumerate(arquivos, start=1):
        nome = limpar_nome_arquivo(f"foto_relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{idx}_{arquivo.name}")
        destino = pasta_fotos / nome
        foto_bytes = arquivo.getbuffer()
        with open(destino, "wb") as f:
            f.write(foto_bytes)
        caminhos.append(destino)
        # Upload paralelo para Google Drive
        try:
            drive_upload_foto(
                arquivo_bytes=bytes(foto_bytes),
                nome_arquivo=nome,
                nome_condominio=nome_cond,
                mes_ano=mes_ano,
            )
        except Exception:
            pass  # Falha no Drive não impede gerar o relatório
    return caminhos


def buscar_fotos_drive_para_relatorio(nome_condominio: str, mes_ano: str = None) -> list[Path]:
    """Baixa fotos do Drive para pasta temporária e retorna caminhos locais para inserir no DOCX."""
    import tempfile
    if not mes_ano:
        mes_ano = datetime.now().strftime("%Y-%m")
    try:
        service = conectar_drive()
        if not service:
            return []

        # Busca pasta do condomínio
        q_cond = f"name='{nome_condominio}' and '{DRIVE_FOTOS_FOLDER_ID}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        res_cond = service.files().list(q=q_cond, fields="files(id,name)").execute()
        pastas_cond = res_cond.get("files", [])
        if not pastas_cond:
            return []
        pasta_cond_id = pastas_cond[0]["id"]

        # Busca pasta do mês
        q_mes = f"name='{mes_ano}' and '{pasta_cond_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        res_mes = service.files().list(q=q_mes, fields="files(id,name)").execute()
        pastas_mes = res_mes.get("files", [])
        if not pastas_mes:
            return []
        pasta_mes_id = pastas_mes[0]["id"]

        # Lista fotos
        q_fotos = f"'{pasta_mes_id}' in parents and trashed=false"
        res_fotos = service.files().list(q=q_fotos, fields="files(id,name,mimeType)").execute()
        fotos = res_fotos.get("files", [])

        # Baixa para pasta temp
        caminhos = []
        tmp_dir = Path(tempfile.mkdtemp())
        for foto in fotos:
            try:
                conteudo = service.files().get_media(fileId=foto["id"]).execute()
                dest = tmp_dir / foto["name"]
                with open(dest, "wb") as f:
                    f.write(conteudo)
                caminhos.append(dest)
            except Exception:
                pass
        return caminhos
    except Exception as e:
        _log_sheets_erro("buscar_fotos_drive_para_relatorio", e)
        return []


def garantir_campos_analises(qtd: int):
    qtd = max(ANALISES_PADRAO, int(qtd or ANALISES_PADRAO))
    qtd = min(qtd, ANALISES_MAX_SUGERIDO)
    st.session_state.rel_analises_total = qtd
    for i in range(qtd):
        for sufixo in ["data", "ph", "cl", "ct", "alc", "dc", "cya", "operador"]:
            chave = f"rel_analise_{sufixo}_{i}"
            if chave not in st.session_state:
                st.session_state[chave] = ""


def adicionar_analise_extra():
    atual = int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO)
    garantir_campos_analises(atual + 1)


def coletar_analises_relatorio() -> list[dict]:
    itens = []
    qtd = int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO)
    garantir_campos_analises(qtd)
    for i in range(qtd):
        itens.append({
            "data": (st.session_state.get(f"rel_analise_data_{i}") or "").strip(),
            "ph": (st.session_state.get(f"rel_analise_ph_{i}") or "").strip(),
            "cloro_livre": (st.session_state.get(f"rel_analise_cl_{i}") or "").strip(),
            "cloro_total": (st.session_state.get(f"rel_analise_ct_{i}") or "").strip(),
            "alcalinidade": (st.session_state.get(f"rel_analise_alc_{i}") or "").strip(),
            "dureza": (st.session_state.get(f"rel_analise_dc_{i}") or "").strip(),
            "cianurico": (st.session_state.get(f"rel_analise_cya_{i}") or "").strip(),
            "operador": (st.session_state.get(f"rel_analise_operador_{i}") or "").strip(),
        })
    return itens


def coletar_dosagens_relatorio() -> list[dict]:
    itens = []
    for i in range(7):
        itens.append({
            "produto": (st.session_state.get(f"rel_dos_produto_{i}") or "").strip(),
            "fabricante_lote": (st.session_state.get(f"rel_dos_lote_{i}") or "").strip(),
            "quantidade": (st.session_state.get(f"rel_dos_qtd_{i}") or "").strip(),
            "unidade": (st.session_state.get(f"rel_dos_un_{i}") or "").strip(),
            "finalidade": (st.session_state.get(f"rel_dos_finalidade_{i}") or "").strip(),
        })
    return itens


def coletar_recomendacoes_relatorio() -> list[dict]:
    itens = []
    for i in range(5):
        itens.append({
            "recomendacao": (st.session_state.get(f"rel_rec_texto_{i}") or "").strip(),
            "prazo": (st.session_state.get(f"rel_rec_prazo_{i}") or "").strip(),
            "responsavel": (st.session_state.get(f"rel_rec_resp_{i}") or "").strip(),
        })
    return itens


def coletar_observacoes_relatorio() -> list[str]:
    return [(st.session_state.get(f"rel_obs_{i}") or "").strip() for i in range(5)]


def coletar_conformidades_relatorio() -> dict:
    return {
        "nbr_11238": (st.session_state.get("rel_nbr_11238") or "").strip(),
        "nbr_11238_status": {
            "profundidade": (st.session_state.get("rel_nbr11238_profundidade") or "Sim").strip(),
            "retrolavagem": (st.session_state.get("rel_nbr11238_retrolavagem") or "Sim").strip(),
            "skimmers": (st.session_state.get("rel_nbr11238_skimmers") or "Sim").strip(),
            "circulacao": (st.session_state.get("rel_nbr11238_circulacao") or "Sim").strip(),
            "chuveiro": (st.session_state.get("rel_nbr11238_chuveiro") or "Sim").strip(),
        },
        "nr_26": (st.session_state.get("rel_nr_26") or "").strip(),
        "nr_06": (st.session_state.get("rel_nr_06") or "").strip(),
    }


def coletar_epis_relatorio() -> dict:
    return {
        "luvas_status": (st.session_state.get("rel_epi_luvas_status") or "Conforme").strip(),
        "luvas_ca": (st.session_state.get("rel_epi_luvas_ca") or "").strip(),
        "oculos_status": (st.session_state.get("rel_epi_oculos_status") or "Conforme").strip(),
        "oculos_ca": (st.session_state.get("rel_epi_oculos_ca") or "").strip(),
        "respirador_status": (st.session_state.get("rel_epi_respirador_status") or "Conforme").strip(),
        "respirador_ca": (st.session_state.get("rel_epi_respirador_ca") or "").strip(),
        "botas_status": (st.session_state.get("rel_epi_botas_status") or "Conforme").strip(),
        "botas_ca": (st.session_state.get("rel_epi_botas_ca") or "").strip(),
    }


def gerar_textos_automaticos_relatorio(analises: list[dict], avaliacao: dict) -> dict:
    observacoes = []
    recomendacoes = []
    detalhes = avaliacao.get("detalhes", [])
    if detalhes:
        for detalhe in detalhes[:3]:
            observacoes.append(detalhe.replace("Linha", "Leitura"))
    if avaliacao.get("cloraminas_altas"):
        idx, valor = avaliacao["cloraminas_altas"][0]
        observacoes.append(f"Leitura {idx}: cloro combinado estimado em {valor} mg/L, sugerindo formação de cloraminas e maior carga orgânica.")
        recomendacoes.append("Executar oxidação complementar / supercloração controlada e reavaliar CT, CL e cloro combinado.")
    if any("ph=" in d for d in detalhes):
        recomendacoes.append("Ajustar pH para a faixa operacional de 7,2 a 7,8 e repetir a leitura após estabilização.")
    if any("cloro residual livre" in d for d in detalhes):
        recomendacoes.append("Revisar a dosagem de desinfetante e a demanda oxidante da água.")
    if any("ácido cianúrico" in d for d in detalhes):
        recomendacoes.append("Reavaliar o estabilizante e considerar renovação parcial de água quando tecnicamente indicada.")
    if any("alcalinidade total" in d for d in detalhes):
        recomendacoes.append("Corrigir alcalinidade total para aumentar a estabilidade do pH e reduzir flutuações operacionais.")
    if any("dureza cálcica" in d for d in detalhes):
        recomendacoes.append("Ajustar dureza cálcica para reduzir risco de corrosão ou incrustação.")

    if not observacoes:
        observacoes.append("Parâmetros avaliados sem desvios críticos nas leituras registradas no período.")
    if not recomendacoes:
        recomendacoes.append("Manter rotina de monitoramento, rastreabilidade de dosagens e controle técnico periódico.")

    diagnostico = []
    if avaliacao["status"] == "CONFORME":
        diagnostico.append("No período avaliado, os parâmetros registrados indicam condição satisfatória de controle físico-químico da água, sem desvios relevantes frente às faixas operacionais adotadas pelo sistema.")
    elif avaliacao["status"] == "EM CORREÇÃO":
        diagnostico.append("O sistema registra fase operacional de correção, recomendando continuidade das medidas técnicas, reforço de monitoramento e nova conferência analítica após estabilização.")
    else:
        diagnostico.append("Foram observadas não conformidades físico-químicas que exigem ação corretiva e reavaliação técnica sequencial.")
    if avaliacao.get("cloraminas_altas"):
        diagnostico.append("A presença de cloro combinado acima do desejável sugere formação de cloraminas, condição compatível com carga orgânica elevada, consumo do desinfetante e possível redução da eficiência sanitizante.")
    if detalhes:
        diagnostico.append("Resumo automático dos desvios: " + "; ".join(detalhes[:5]) + ".")

    return {
        "diagnostico": " ".join(diagnostico),
        "observacoes": observacoes[:5],
        "recomendacoes": recomendacoes[:5],
    }


def aplicar_textos_automaticos_relatorio():
    analises = coletar_analises_relatorio()
    avaliacao = avaliar_conformidade_analises(analises)
    textos = gerar_textos_automaticos_relatorio(analises, avaliacao)

    # Nunca apaga parâmetros digitados. Só recalcula os campos textuais do parecer.
    st.session_state.rel_diagnostico = textos["diagnostico"]

    for i in range(5):
        chave_obs = f"rel_obs_{i}"
        texto_obs = textos["observacoes"][i] if i < len(textos["observacoes"]) else ""
        st.session_state[chave_obs] = texto_obs

    for i in range(5):
        chave_rec = f"rel_rec_texto_{i}"
        chave_prazo = f"rel_rec_prazo_{i}"
        chave_resp = f"rel_rec_resp_{i}"
        texto_rec = textos["recomendacoes"][i] if i < len(textos["recomendacoes"]) else ""
        st.session_state[chave_rec] = texto_rec
        st.session_state[chave_prazo] = "Imediato" if texto_rec and i == 0 else ("Próxima rotina" if texto_rec else "")
        st.session_state[chave_resp] = "Operação / RT" if texto_rec else ""

    # Não escrever diretamente na chave do selectbox após o widget existir.
    if st.session_state.get("rel_status_agua") != "EM CORREÇÃO":
        st.session_state["rel_status_agua"] = avaliacao["status"]


def montar_dados_relatorio() -> dict:
    nome_condominio = (st.session_state.get("rel_nome_condominio") or "").strip()
    representante = (st.session_state.get("rel_representante") or "").strip()
    dados_base = obter_snapshot_relatorio_independente()
    analises = coletar_analises_relatorio()
    avaliacao = avaliar_conformidade_analises(analises)
    status_manual = (st.session_state.get("rel_status_agua") or "CONFORME").strip()
    status_final = "EM CORREÇÃO" if status_manual == "EM CORREÇÃO" else avaliacao["status"]
    textos_auto = gerar_textos_automaticos_relatorio(analises, avaliacao)
    diagnostico_base = (st.session_state.get("rel_diagnostico") or "").strip() or textos_auto["diagnostico"]

    recomendacoes = coletar_recomendacoes_relatorio()
    if not any(r["recomendacao"] for r in recomendacoes):
        recomendacoes = [{"recomendacao": t, "prazo": "Imediato" if i == 0 else "Próxima rotina", "responsavel": "Operação / RT"} for i, t in enumerate(textos_auto["recomendacoes"])]

    observacoes = coletar_observacoes_relatorio()
    if not any(observacoes):
        observacoes = textos_auto["observacoes"]

    return {
        "empresa_rt": EMPRESA_RT,
        "responsavel_tecnico": RESPONSAVEL_TÉCNICO,
        "assinatura_rt_texto": RESPONSAVEL_TECNICO_ASSINATURA,
        "crq": CRQ,
        "qualificacao": QUALIFICACAO_RT,
        "certificacoes": CERTIFICACOES_RT,
        "nome_condominio": nome_condominio,
        "cnpj_condominio": dados_base.get("cnpj_condominio", ""),
        "endereco_condominio": dados_base.get("endereco_condominio", ""),
        "representante": representante,
        "cpf_cnpj_representante": dados_base.get("cpf_sindico", ""),
        "tipo_atendimento": (st.session_state.get("rel_tipo_atendimento") or "Contrato ativo").strip(),
        "mes_referencia": (st.session_state.get("rel_mes_referencia") or "").strip(),
        "ano_referencia": (st.session_state.get("rel_ano_referencia") or "").strip(),
        "art_status": (st.session_state.get("rel_art_status") or "Emitida").strip(),
        "art_numero": (st.session_state.get("rel_art_numero") or "").strip() if (st.session_state.get("rel_art_status") or "Emitida").strip() == "Emitida" else "N/A",
        "art_inicio": (st.session_state.get("rel_art_inicio") or "").strip() if (st.session_state.get("rel_art_status") or "Emitida").strip() == "Emitida" else "N/A",
        "art_fim": (st.session_state.get("rel_art_fim") or "").strip() if (st.session_state.get("rel_art_status") or "Emitida").strip() == "Emitida" else "N/A",
        "art_texto": "",
        "data_emissao": (st.session_state.get("rel_data_emissao") or hoje_br()).strip(),
        "status_agua": status_final,
        "diagnostico": diagnostico_base,
        "analises": analises,
        "dosagens": coletar_dosagens_relatorio(),
        "recomendacoes": recomendacoes,
        "observacoes": observacoes,
        "conformidades": coletar_conformidades_relatorio(),
        "epis": coletar_epis_relatorio(),
        "avaliacao_automatica": avaliacao,
    }


def validar_relatorio_mensal(dados_relatorio: dict) -> list[str]:
    erros = []
    if not dados_relatorio.get("nome_condominio"):
        erros.append("Informe o nome do condomínio/estabelecimento no próprio relatório antes de gerar.")
    if not dados_relatorio.get("mes_referencia"):
        erros.append("Informe o mês de referência do relatório.")
    if not dados_relatorio.get("ano_referencia"):
        erros.append("Informe o ano de referência do relatório.")
    if not TEMPLATE_RELATORIO.exists():
        erros.append("O arquivo relatorio_mensal.docx não foi localizado na pasta do projeto.")
    if not validar_data_br(dados_relatorio.get("data_emissao", "")):
        erros.append("Data de emissão do relatório inválida.")
    if dados_relatorio.get("art_status") == "Emitida":
        if not dados_relatorio.get("art_numero"):
            erros.append("Informe o número da ART ou altere o status para Não emitida / Em tramitação.")
        if not validar_data_br(dados_relatorio.get("art_inicio", "")):
            erros.append("Vigência ART - início inválida.")
        if not validar_data_br(dados_relatorio.get("art_fim", "")):
            erros.append("Vigência ART - fim inválida.")
    return erros


def append_relatorio_fallback(doc: Document, dados_relatorio: dict):
    doc.add_page_break()
    doc.add_paragraph("COMPLEMENTO AUTOMÁTICO – DADOS ESTRUTURADOS DO RELATÓRIO MENSAL")
    doc.add_paragraph(f"Condomínio: {dados_relatorio['nome_condominio']}")
    doc.add_paragraph(f"Mês/Ano de referência: {dados_relatorio['mes_referencia']}/{dados_relatorio['ano_referencia']}")
    doc.add_paragraph(f"ART: {obter_status_art_texto(dados_relatorio)}")
    doc.add_paragraph(f"Data de emissão: {dados_relatorio['data_emissao']}")
    doc.add_paragraph(f"Responsável técnico: {dados_relatorio['assinatura_rt_texto']}")
    doc.add_paragraph(f"Certificações relevantes: {dados_relatorio['certificacoes']}")
    doc.add_paragraph(f"Status geral da água: {dados_relatorio['status_agua']}")
    doc.add_paragraph(f"Diagnóstico técnico: {dados_relatorio['diagnostico']}")
    doc.add_paragraph("Base normativa referencial do relatório: Lei nº 2.800/1956; Decreto nº 85.877/1981; Lei nº 6.839/1980; Resolução CFQ nº 332/2025; ABNT NBR 10339; NR-26; NR-06.")
    doc.add_paragraph("Nota técnica: análises microbiológicas não são realizadas no local e dependem de laboratório acreditado, sob responsabilidade de contratação do cliente.")

    doc.add_paragraph("Observações automáticas:")
    for obs in dados_relatorio["observacoes"]:
        if obs:
            doc.add_paragraph(f"• {obs}")

    doc.add_paragraph("Recomendações automáticas:")
    for item in dados_relatorio["recomendacoes"]:
        if item.get("recomendacao"):
            doc.add_paragraph(f"• {item['recomendacao']} | Prazo: {item.get('prazo','')} | Responsável: {item.get('responsavel','')}")


def atualizar_textos_normativos(doc: Document):
    # Mapeia TODAS as variações conhecidas de CRQ errado e normas defasadas
    trocas_ordenadas = [
        # CRQ – variações mais específicas primeiro para evitar substituição parcial
        ("CRQ CRQ 024025748 – 4ª Região", CRQ),
        ("CRQ 024025748 – 4ª Região", CRQ),
        ("CRQ-IV | CRQ 024025748", CRQ),
        ("CRQ-IV – CRQ 024025748", CRQ),
        ("CRQ-IV", "CRQ-MG 2ª Região"),
        # Normas
        ("Resolução CFQ nº 491/2020", "Resolução CFQ nº 332/2025"),
        ("Res. Normativa CFQ nº 01/1982", "Decreto nº 85.877/1981 e normas profissionais aplicáveis do Sistema CFQ/CRQs"),
        ("Portaria MS nº 888/2021", "Portaria GM/MS nº 888/2021 (referência sanitária complementar para água de consumo humano)"),
        ("NBR 10818", "ABNT NBR 10339"),
    ]
    # Aplica na ordem para evitar substituições aninhadas incorretas
    for de, para in trocas_ordenadas:
        substituir_placeholders_doc(doc, {de: para})


def preencher_relatorio_mensal_docx(template_path: Path, output_docx: Path, dados_relatorio: dict, fotos: list[Path] | None = None):
    doc = Document(str(template_path))
    dados_relatorio["art_texto"] = obter_status_art_texto(dados_relatorio)

    # ---- Placeholders primários (substituição em todo o documento) ----
    placeholders = {
        "{{NOME_CONDOMINIO}}": dados_relatorio["nome_condominio"],
        "{{CNPJ_CONDOMINIO}}": dados_relatorio["cnpj_condominio"],
        "{{ENDERECO_CONDOMINIO}}": dados_relatorio["endereco_condominio"],
        "{{NOME_SINDICO}}": dados_relatorio["representante"],
        "{{RESPONSAVEL_TÉCNICO}}": dados_relatorio["responsavel_tecnico"],
        "{{RESPONSAVEL_TECNICO_ASSINATURA}}": dados_relatorio["assinatura_rt_texto"],
        "{{CRQ}}": dados_relatorio["crq"],
        "{{QUALIFICACAO_RT}}": dados_relatorio["qualificacao"],
        "{{CERTIFICACOES_RT}}": dados_relatorio["certificacoes"],
        "{{EMPRESA_RT}}": dados_relatorio["empresa_rt"],
        "{{MES_REFERENCIA}}": dados_relatorio["mes_referencia"],
        "{{ANO_REFERENCIA}}": dados_relatorio["ano_referencia"],
        "{{ART_NUMERO}}": dados_relatorio["art_numero"],
        "{{ART_INICIO}}": dados_relatorio["art_inicio"],
        "{{ART_FIM}}": dados_relatorio["art_fim"],
        "{{ART_TEXTO}}": dados_relatorio["art_texto"],
        "{{DATA_EMISSAO}}": dados_relatorio["data_emissao"],
        "{{STATUS_AGUA}}": dados_relatorio["status_agua"],
        "{{DIAGNOSTICO_TÉCNICO}}": dados_relatorio["diagnostico"],
        "{{DIAGNOSTICO_TECNICO}}": dados_relatorio["diagnostico"],
    }
    substituir_placeholders_doc(doc, placeholders)
    atualizar_textos_normativos(doc)
    preencher_tabela_identificacao(doc, dados_relatorio)
    preencher_bloco_conformidades(doc, dados_relatorio)

    # ---- Tabela de análises ----
    tabela_analises = encontrar_tabela_por_keywords(doc, ["Data", "CRL", "Cl. Total", "Operador"])
    linhas_analises = []
    for item in dados_relatorio["analises"]:
        if any((item.get(k) or "").strip() for k in ["data", "ph", "cloro_livre", "cloro_total", "alcalinidade", "dureza", "cianurico", "operador"]):
            linhas_analises.append([item["data"], item["ph"], item["cloro_livre"], item["cloro_total"], item["alcalinidade"], item["dureza"], item["cianurico"], item["operador"]])
    if linhas_analises:
        preencher_tabela_generica(tabela_analises, linhas_analises, start_row=1)

    # ---- Tabela de dosagens ----
    tabela_dos = encontrar_tabela_por_keywords(doc, ["Produto Químico", "Fabricante / Lote", "Finalidade Técnica"])
    linhas_dos = []
    for item in dados_relatorio["dosagens"]:
        if any(item.values()):
            linhas_dos.append([item["produto"], item["fabricante_lote"], item["quantidade"], item["unidade"], item["finalidade"]])
    if linhas_dos:
        preencher_tabela_generica(tabela_dos, linhas_dos, start_row=1)

    # ---- Tabela de recomendações ----
    tabela_rec = encontrar_tabela_por_keywords(doc, ["Recomendação Técnica", "Prazo", "Responsável"])
    linhas_rec = []
    for idx, item in enumerate(dados_relatorio["recomendacoes"], start=1):
        if item.get("recomendacao"):
            linhas_rec.append([str(idx), item["recomendacao"], item.get("prazo", ""), item.get("responsavel", "")])
    if linhas_rec:
        preencher_tabela_generica(tabela_rec, linhas_rec, start_row=1)

    # ---- Parecer principal e aviso legal em parágrafos soltos ----
    # Também cobre o caso de o template ter esses textos fora de tabelas/placeholders.
    aviso_legal_texto = ""  # removido a pedido do usuário

    # Rastreia se os quadros principais foram encontrados e preenchidos
    quadro_diagnostico_preenchido = False

    for p in doc.paragraphs:
        txt = normalizar_texto(p.text)
        # Quadro de diagnóstico – qualquer parágrafo que contenha o texto modelo do template
        if "diagnostico:" in txt and dados_relatorio["diagnostico"] not in p.text:
            # Preserva formatação dos runs, substituindo só o conteúdo
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = f"Diagnóstico: {dados_relatorio['diagnostico']}"
            else:
                p.add_run(f"Diagnóstico: {dados_relatorio['diagnostico']}")
            quadro_diagnostico_preenchido = True
        elif "conforme" in txt and ("nao conforme" in txt or "em correcao" in txt):
            # Linha de status geral – substitui o texto modelo
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = f"Status geral da água: {dados_relatorio['status_agua']}"
            else:
                p.add_run(f"Status geral da água: {dados_relatorio['status_agua']}")
        elif "aviso legal:" in txt:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = aviso_legal_texto
            else:
                p.add_run(aviso_legal_texto)

    # ---- Também busca diagnóstico dentro de tabelas do template ----
    if not quadro_diagnostico_preenchido:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = normalizar_texto(p.text)
                        if "diagnostico:" in txt and dados_relatorio["diagnostico"] not in p.text:
                            for run in p.runs:
                                run.text = ""
                            if p.runs:
                                p.runs[0].text = f"Diagnóstico: {dados_relatorio['diagnostico']}"
                            else:
                                p.add_run(f"Diagnóstico: {dados_relatorio['diagnostico']}")

    # ---- Assinatura automática REMOVIDA definitivamente. ----
    # O bloco oficial de assinatura existe no template relatorio_mensal.docx
    # e não deve ser duplicado ou complementado aqui.

    # ---- Fotos ----
    if fotos:
        doc.add_page_break()
        doc.add_paragraph("REGISTRO FOTOGRÁFICO")
        for foto in fotos:
            try:
                doc.add_paragraph(foto.name if hasattr(foto, "name") else str(foto))
                _p_foto_m = doc.add_paragraph()
                _p_foto_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _ok_m = inserir_foto_docx_exif(_p_foto_m, foto, width_inches=5.8)
                if not _ok_m:
                    doc.add_paragraph(f"Não foi possível inserir a foto: {foto.name if hasattr(foto, 'name') else foto}")
            except Exception:
                doc.add_paragraph(f"Não foi possível inserir a foto.")

    # ---- Gráfico de tendência mensal (pH e CRL) ----
    try:
        _grafico_tendencia = gerar_grafico_tendencia_ph_crl(
            dados_relatorio.get("analises", []),
            nome_condominio=dados_relatorio.get("nome_condominio", ""),
            mes=dados_relatorio.get("mes_referencia", ""),
            ano=dados_relatorio.get("ano_referencia", ""),
            output_path=output_docx.with_name(output_docx.stem + "_grafico_tendencia.png"),
        )
        if _grafico_tendencia and Path(_grafico_tendencia).exists():
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN

            doc.add_page_break()
            _pt = doc.add_paragraph()
            _pt.alignment = _WD_ALIGN.CENTER
            _rt = _pt.add_run("GRÁFICOS DE TENDÊNCIA — pH e CRL")
            _rt.bold = True
            _rt.font.size = Pt(12)

            _ps = doc.add_paragraph()
            _ps.alignment = _WD_ALIGN.CENTER
            _rs = _ps.add_run("Visualização automática da estabilidade da água ao longo do mês de referência.")
            _rs.font.size = Pt(9)

            _pg = doc.add_paragraph()
            _pg.alignment = _WD_ALIGN.CENTER
            inserir_foto_docx_exif(_pg, Path(_grafico_tendencia), width_inches=6.3)

            _pc = doc.add_paragraph()
            _pc.alignment = _WD_ALIGN.CENTER
            _rc = _pc.add_run("As áreas verdes representam a faixa ideal de operação para cada parâmetro.")
            _rc.font.size = Pt(8)
    except Exception:
        pass

    # ---- Fallback CONDICIONAL: só adiciona se as tabelas principais não foram encontradas ----
    # O fallback NÃO deve ser gerado se o template já tem os quadros de análise, dosagem e recomendação.
    template_tem_analises = encontrar_tabela_por_keywords(doc, ["Data", "CRL", "Cl. Total", "Operador"]) is not None
    template_tem_dosagens = encontrar_tabela_por_keywords(doc, ["Produto Químico", "Fabricante / Lote", "Finalidade Técnica"]) is not None
    if not template_tem_analises and not template_tem_dosagens:
        append_relatorio_fallback(doc, dados_relatorio)

    doc.save(str(output_docx))


def _normalizar_lista_fotos_b64(valor) -> list[str]:
    if not valor:
        return []
    if isinstance(valor, str):
        valor = [valor]
    return [str(v).strip() for v in valor if str(v).strip()]


def _imagem_bytes_para_b64_relatorio(img_bytes: bytes) -> str | None:
    try:
        import io as _io
        import base64 as _b64
        from PIL import Image as _PILR, ImageOps as _IOps

        _img = _PILR.open(_io.BytesIO(img_bytes))
        _img = _IOps.exif_transpose(_img)
        if _img.mode != "RGB":
            _img = _img.convert("RGB")
        _img.thumbnail((1600, 1600))
        _buf = _io.BytesIO()
        _img.save(_buf, format="JPEG", quality=85, optimize=True)
        return _b64.b64encode(_buf.getvalue()).decode("utf-8")
    except Exception:
        return None


def _normalizar_assinatura_b64(valor: str) -> str:
    if not valor:
        return ""
    txt = str(valor).strip()
    if txt.lower().startswith("data:image") and "," in txt:
        txt = txt.split(",", 1)[1].strip()
    try:
        import io as _io
        import base64 as _b64

        bruto = _b64.b64decode(txt)
        with Image.open(_io.BytesIO(bruto)) as _img:
            _img = ImageOps.exif_transpose(_img)
            if _img.mode != "RGBA":
                _img = _img.convert("RGBA")
            bbox = None
            if "A" in _img.getbands():
                _alpha = _img.getchannel("A").point(lambda p: 255 if p > 10 else 0)
                bbox = _alpha.getbbox()
            if not bbox:
                _cinza = _img.convert("L").point(lambda p: 255 if p < 245 else 0)
                bbox = _cinza.getbbox()
            if not bbox:
                return ""
            _img = _img.crop(bbox)
            _canvas = Image.new("RGBA", (_img.width + 40, _img.height + 30), (255, 255, 255, 0))
            _canvas.paste(_img, (20, 15), _img)
            _buf = _io.BytesIO()
            _canvas.save(_buf, format="PNG", optimize=True)
            return _b64.b64encode(_buf.getvalue()).decode("utf-8")
    except Exception:
        return txt


def _assinatura_canvas_para_b64(image_data) -> str:
    if image_data is None:
        return ""
    try:
        import io as _io
        import base64 as _b64

        _img = Image.fromarray(image_data.astype("uint8"), mode="RGBA")
        _buf = _io.BytesIO()
        _img.save(_buf, format="PNG")
        return _normalizar_assinatura_b64(_b64.b64encode(_buf.getvalue()).decode("utf-8"))
    except Exception:
        return ""


def _salvar_assinatura_local(assinatura_b64: str, destino: Path) -> bool:
    try:
        import base64 as _b64
        bruto = _b64.b64decode(_normalizar_assinatura_b64(assinatura_b64))
        destino.parent.mkdir(parents=True, exist_ok=True)
        with open(destino, "wb") as _f:
            _f.write(bruto)
        return True
    except Exception:
        return False


def gerar_grafico_tendencia_ph_crl(analises: list[dict], nome_condominio: str = "", mes: str = "", ano: str = "", output_path: Path | None = None) -> Path | None:
    try:
        from PIL import ImageDraw, ImageFont

        pontos = []
        for idx, item in enumerate(analises or []):
            ph = valor_float(item.get("ph", ""))
            crl = valor_float(item.get("cloro_livre", ""))
            if ph is None and crl is None:
                continue
            data_txt = (item.get("data") or "").strip()
            data_ord = parse_data_br(data_txt) or (date.today() + timedelta(days=idx))
            rotulo = data_txt[:5] if data_txt else f"#{idx+1}"
            pontos.append({"ordem": data_ord, "rotulo": rotulo, "ph": ph, "crl": crl})

        if not pontos:
            return None

        pontos.sort(key=lambda p: p["ordem"])
        output_path = output_path or (GENERATED_DIR / "_grafico_tendencia_relatorio.png")
        output_path.parent.mkdir(parents=True, exist_ok=True)

        largura, altura = 1600, 920
        img = Image.new("RGB", (largura, altura), (250, 252, 255))
        draw = ImageDraw.Draw(img)

        def _fonte(tam: int, bold: bool = False):
            candidatas = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
                "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            ]
            for caminho in candidatas:
                try:
                    return ImageFont.truetype(caminho, tam)
                except Exception:
                    continue
            return ImageFont.load_default()

        f_title = _fonte(34, True)
        f_sub = _fonte(18, False)
        f_axis = _fonte(16, False)
        f_small = _fonte(14, False)
        f_bold = _fonte(18, True)

        azul = (13, 61, 117)
        azul_claro = (74, 144, 226)
        verde_faixa = (233, 246, 236)
        verde_linha = (34, 139, 94)
        grid = (208, 216, 228)
        texto = (26, 42, 74)
        cinza = (122, 136, 160)
        draw.rounded_rectangle((24, 24, largura - 24, altura - 24), radius=28, outline=(210, 220, 234), width=2, fill=(255, 255, 255))
        titulo = "Tendência mensal de pH e cloro livre (CRL)"
        subtitulo = f"{nome_condominio or 'Condomínio'} • referência {mes or '--'}/{ano or '--'}"
        draw.text((60, 54), titulo, fill=texto, font=f_title)
        draw.text((60, 98), subtitulo, fill=cinza, font=f_sub)

        def _plotar(y_top: int, chave: str, titulo_plot: str, cor_linha: tuple[int, int, int], faixa_min: float, faixa_max: float, unidade: str):
            x0, x1 = 120, largura - 70
            y0, y1 = y_top + 46, y_top + 290
            valores = [p[chave] for p in pontos if p[chave] is not None]
            if not valores:
                draw.text((60, y_top + 120), f"{titulo_plot}: sem dados no mês", fill=cinza, font=f_bold)
                return

            y_min = min(valores + [faixa_min])
            y_max = max(valores + [faixa_max])
            pad = max((y_max - y_min) * 0.18, 0.15 if chave == 'ph' else 0.35)
            y_min = max(0.0, y_min - pad)
            y_max = y_max + pad
            if abs(y_max - y_min) < 0.001:
                y_max = y_min + 1.0

            def y_map(v: float) -> float:
                return y1 - ((v - y_min) / (y_max - y_min)) * (y1 - y0)

            faixa_top = y_map(faixa_max)
            faixa_bottom = y_map(faixa_min)
            draw.rounded_rectangle((x0, faixa_top, x1, faixa_bottom), radius=14, fill=verde_faixa)

            for i in range(6):
                y = y0 + (y1 - y0) * i / 5
                valor = y_max - ((y - y0) / (y1 - y0)) * (y_max - y_min)
                draw.line((x0, y, x1, y), fill=grid, width=1)
                draw.text((48, y - 10), f"{valor:.2f}".replace('.', ','), fill=cinza, font=f_small)

            draw.line((x0, y0, x0, y1), fill=grid, width=2)
            draw.line((x0, y1, x1, y1), fill=grid, width=2)
            draw.text((60, y_top + 8), titulo_plot, fill=texto, font=f_bold)
            draw.text((x1 - 240, y_top + 10), f"Faixa ideal: {str(faixa_min).replace('.', ',')} a {str(faixa_max).replace('.', ',')} {unidade}", fill=verde_linha, font=f_small)

            total = max(len(pontos), 1)
            xs = [x0 + ((x1 - x0) * i / max(total - 1, 1)) for i in range(total)]
            pts_linha = []
            for i, ponto in enumerate(pontos):
                x = xs[i]
                draw.line((x, y1, x, y1 + 10), fill=grid, width=1)
                draw.text((x - 18, y1 + 16), ponto['rotulo'], fill=cinza, font=f_small)
                v = ponto[chave]
                if v is None:
                    continue
                y = y_map(v)
                pts_linha.append((x, y, v))

            if len(pts_linha) >= 2:
                draw.line([(x, y) for x, y, _ in pts_linha], fill=cor_linha, width=4)
            for x, y, v in pts_linha:
                draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=cor_linha, outline=(255, 255, 255), width=2)
                draw.text((x - 16, y - 28), str(round(v, 2)).replace('.', ','), fill=cor_linha, font=f_small)

        _plotar(150, 'ph', 'pH', azul, 7.2, 7.8, '')
        _plotar(485, 'crl', 'Cloro livre residual (CRL)', azul_claro, 0.5, 3.0, 'mg/L')
        draw.text((60, altura - 54), 'Áreas em verde representam a faixa operacional ideal. O gráfico é gerado a partir das análises lançadas no mês.', fill=cinza, font=f_sub)

        img.save(output_path, format='PNG', optimize=True)
        return output_path
    except Exception:
        return None


def _resolver_fotos_visita_para_relatorio(lancamento: dict) -> dict[str, list[str]]:
    """Resolve fotos do lançamento priorizando base64, depois Drive e por fim arquivos locais."""
    categorias = {
        "antes": {"b64": [], "ids": [], "nomes": []},
        "depois": {"b64": [], "ids": [], "nomes": []},
        "cmaq": {"b64": [], "ids": [], "nomes": []},
        "extras": {"b64": [], "ids": [], "nomes": []},
    }

    for cat in categorias:
        categorias[cat]["b64"] = _normalizar_lista_fotos_b64(lancamento.get(f"fotos_{cat}_b64", []))
        categorias[cat]["ids"] = list(lancamento.get(f"fotos_{cat}_ids", []) or [])
        categorias[cat]["nomes"] = list(lancamento.get(f"fotos_{cat}", []) or [])

    if not categorias["antes"]["ids"]:
        categorias["antes"]["ids"] = list(lancamento.get("fotos_drive_ids", []) or [])

    nomes_gerais = list(lancamento.get("fotos", []) or [])
    if nomes_gerais:
        for nome in nomes_gerais:
            nome_l = str(nome).lower()
            if ("antes" in nome_l) and nome not in categorias["antes"]["nomes"]:
                categorias["antes"]["nomes"].append(nome)
            elif ("depois" in nome_l) and nome not in categorias["depois"]["nomes"]:
                categorias["depois"]["nomes"].append(nome)
            elif (("cmaq" in nome_l) or ("maq" in nome_l)) and nome not in categorias["cmaq"]["nomes"]:
                categorias["cmaq"]["nomes"].append(nome)

    nome_cond = (lancamento.get("condominio") or "").strip()
    pasta_base = GENERATED_DIR / slugify_nome(nome_cond) if nome_cond else None
    pastas_busca = []
    if pasta_base:
        pastas_busca = [
            pasta_base / "fotos_campo",
            pasta_base / "fotos_relatorio",
            pasta_base / "_previa_exata_relatorio" / "fotos_upload",
            pasta_base / "fotos_rascunho",
        ]

    def _adicionar_b64(cat: str, valor_b64: str | None):
        if valor_b64 and valor_b64 not in categorias[cat]["b64"]:
            categorias[cat]["b64"].append(valor_b64)

    for cat in categorias:
        if not categorias[cat]["b64"]:
            for fid in categorias[cat]["ids"]:
                try:
                    fb = drive_baixar_foto(fid)
                except Exception:
                    fb = None
                if not fb:
                    continue
                _adicionar_b64(cat, _imagem_bytes_para_b64_relatorio(fb))

        if not categorias[cat]["b64"]:
            for nome in categorias[cat]["nomes"]:
                for pasta in pastas_busca:
                    try:
                        caminho = pasta / str(nome)
                    except Exception:
                        caminho = None
                    if caminho and caminho.exists() and caminho.is_file():
                        try:
                            _adicionar_b64(cat, _imagem_bytes_para_b64_relatorio(caminho.read_bytes()))
                        except Exception:
                            pass

        if not categorias[cat]["b64"] and pastas_busca:
            padroes = {
                "antes": ["*antes*", "rasc_antes_*"],
                "depois": ["*depois*", "rasc_depois_*"],
                "cmaq": ["*cmaq*", "*maq*", "rasc_cmaq_*"],
            }
            vistos = set()
            for pasta in pastas_busca:
                if not pasta.exists():
                    continue
                for padrao in padroes.get(cat, []):
                    for caminho in sorted(pasta.glob(padrao)):
                        if not caminho.is_file():
                            continue
                        chave = str(caminho)
                        if chave in vistos:
                            continue
                        vistos.add(chave)
                        try:
                            _adicionar_b64(cat, _imagem_bytes_para_b64_relatorio(caminho.read_bytes()))
                        except Exception:
                            pass

    def _deduplicar_b64_lista(_lista):
        _saida = []
        _vistos = set()
        for _item in _lista or []:
            _texto = str(_item)
            _chave = _texto[:120] + str(len(_texto))
            if _chave in _vistos:
                continue
            _vistos.add(_chave)
            _saida.append(_item)
        return _saida

    resultado = {cat: _deduplicar_b64_lista(categorias[cat]["b64"]) for cat in categorias}

    # Deduplicação global: a mesma foto não deve aparecer em duas seções do PDF.
    _vistos_globais = set()
    for _cat in ["antes", "depois", "cmaq", "extras"]:
        _limpa = []
        for _item in resultado.get(_cat, []) or []:
            _texto = str(_item)
            _chave = _texto[:180] + str(len(_texto))
            if _chave in _vistos_globais:
                continue
            _vistos_globais.add(_chave)
            _limpa.append(_item)
        resultado[_cat] = _limpa

    return resultado


def gerar_html_relatorio_visita(lancamento: dict, nome_condominio: str) -> str:
    """Gera HTML profissional do relatório de visita para download como PDF."""

    def fmt(val, sufixo=""):
        return f"{val}{sufixo}" if val and str(val).strip() else "—"

    def param_box(nome, val, mn, mx, quinzenal=False):
        v = valor_float(val)
        if v is None:
            status_cls = "nd"
            status_txt = "Quinzenal" if quinzenal else "Não medido"
            val_txt = "—"
        elif v < mn or v > mx:
            status_cls = "warn"
            status_txt = "Fora da faixa"
            val_txt = str(val).replace(".", ",")
        else:
            status_cls = "ok"
            status_txt = "Conforme"
            val_txt = str(val).replace(".", ",")
        badge = ' <span class="q15">15d</span>' if quinzenal else ""
        return f"""
        <div class="param-box {status_cls}">
          <div class="pnm">{nome}{badge}</div>
          <div class="pval">{val_txt}</div>
          <div class="pst">{status_txt}</div>
        </div>"""

    # Alertas automáticos
    alertas = []
    checks = [
        (lancamento.get("ph",""), 7.2, 7.8, "pH", "Fora da faixa ideal (7,2–7,8). Corrigir imediatamente."),
        (lancamento.get("cloro_livre",""), 0.5, 3.0, "CRL", "Cloro livre fora da faixa (0,5–3,0 mg/L)."),
        (lancamento.get("alcalinidade",""), 80, 120, "Alcalinidade", "Abaixo do ideal (80–120 mg/L). Aplicar bicarbonato de sódio."),
        (lancamento.get("dureza",""), 150, 300, "Dureza DC", "Fora da faixa (150–300 mg/L)."),
        (lancamento.get("cianurico",""), 30, 50, "CYA", "Ácido cianúrico fora da faixa (30–50 mg/L)."),
    ]
    for val, mn, mx, rot, msg in checks:
        v = valor_float(val)
        if v is not None and (v < mn or v > mx):
            alertas.append(f"{rot}: {str(val).replace('.', ',')} mg/L — {msg}")
    cloraminas = lancamento.get("cloraminas", "")
    if cloraminas and valor_float(cloraminas) is not None and valor_float(cloraminas) > 0.2:
        alertas.append(f"Cloraminas {str(cloraminas).replace('.', ',')} mg/L — acima do limite (0,2 mg/L). Chocar piscina.")

    alertas_html = ""
    if alertas:
        for a in alertas:
            alertas_html += f'<div class="alerta"><div class="alerta-icon">!</div><div class="alerta-txt">{a}</div></div>'
    else:
        alertas_html = '<div class="alerta ok-all"><div class="alerta-txt">Todos os parâmetros medidos dentro da faixa ideal.</div></div>'

    # Dosagens
    dosagens = lancamento.get("dosagens", [])
    dos_html = ""
    for d in dosagens:
        if d.get("produto","").strip():
            qtd = f"{d.get('quantidade','')} {d.get('unidade','')}".strip()
            fin = d.get("finalidade","")
            dos_html += f'<div class="dos-row"><span class="dos-nome">{d["produto"]}</span><span class="dos-detalhe">{qtd}{(" · " + fin) if fin else ""}</span></div>'
    if not dos_html:
        dos_html = '<p style="font-size:12px;color:#8a9ab0;font-style:italic;">Nenhuma dosagem registrada.</p>'

    # Cloraminas box
    clor_box = ""
    if cloraminas and valor_float(cloraminas) is not None:
        v_cl = valor_float(cloraminas)
        cls_cl = "ok" if v_cl <= 0.2 else "warn"
        st_cl = "Conforme" if v_cl <= 0.2 else "Fora da faixa"
        clor_box = f"""
        <div class="param-box {cls_cl}">
          <div class="pnm">Cloraminas</div>
          <div class="pval">{str(cloraminas).replace(".", ",")}</div>
          <div class="pst">{st_cl}</div>
        </div>"""

    obs = lancamento.get("observacao","").strip()
    obs_html = f'<div class="obs-txt">"{obs}"</div>' if obs else '<p style="font-size:12px;color:#8a9ab0;font-style:italic;">Sem observações.</p>'

    data_hoje = date.today().strftime("%d/%m/%Y")
    operador = lancamento.get("operador","") or "—"

    # ── Seção de piscinas (múltiplas ou única) ────────────────────────────────
    piscinas_lista = lancamento.get("piscinas", [])
    if not piscinas_lista:
        # Compatibilidade com lançamentos antigos (sem múltiplas piscinas)
        piscinas_lista = [{
            "nome": "Piscina",
            "ph": lancamento.get("ph",""),
            "cloro_livre": lancamento.get("cloro_livre",""),
            "cloro_total": lancamento.get("cloro_total",""),
            "cloraminas": lancamento.get("cloraminas",""),
            "alcalinidade": lancamento.get("alcalinidade",""),
            "dureza": lancamento.get("dureza",""),
            "cianurico": lancamento.get("cianurico",""),
        }]

    piscinas_html_section = ""
    for pisc in piscinas_lista:
        p_clor_box = ""
        p_clor_val = pisc.get("cloraminas","")
        if p_clor_val and valor_float(p_clor_val) is not None:
            v_cl = valor_float(p_clor_val)
            cls_cl = "ok" if v_cl <= 0.2 else "warn"
            st_cl = "Conforme" if v_cl <= 0.2 else "Fora da faixa"
            p_clor_box = f'''<div class="param-box {cls_cl}"><div class="pnm">Cloraminas</div><div class="pval">{str(p_clor_val).replace(".", ",")}</div><div class="pst">{st_cl}</div></div>'''
        piscinas_html_section += f'''
  <div class="card">
    <div class="sec-title">🏊 {pisc.get("nome","Piscina")} — Parâmetros</div>
    <div class="param-grid">
      {param_box("pH", pisc.get("ph",""), 7.2, 7.8)}
      {param_box("CRL mg/L", pisc.get("cloro_livre",""), 0.5, 3.0)}
      {param_box("CT mg/L", pisc.get("cloro_total",""), 0.5, 5.0)}
      {param_box("Alc. mg/L", pisc.get("alcalinidade",""), 80, 120, quinzenal=True)}
      {param_box("Dureza mg/L", pisc.get("dureza",""), 150, 300, quinzenal=True)}
      {param_box("CYA mg/L", pisc.get("cianurico",""), 30, 50, quinzenal=True)}
      {p_clor_box}
    </div>
  </div>'''

    # Problemas / ocorrências
    problemas = lancamento.get("problemas","").strip()
    problemas_html = f'''
  <div class="card">
    <div class="sec-title">Problemas / Ocorrências</div>
    <div class="obs-txt">"{problemas}"</div>
  </div>''' if problemas else ""

    # ── Fotos do relatório (base64, Drive e fallback local) ──────────────────
    fotos_resolvidas = _resolver_fotos_visita_para_relatorio(lancamento)

    def _b64_to_html(b64_list, titulo):
        if not b64_list:
            return ""
        imgs = ""
        for b64 in b64_list[:6]:
            if not b64:
                continue
            imgs += f'<div style="margin-bottom:6px;"><img src="data:image/jpeg;base64,{b64}" style="width:100%;border-radius:6px;border:1px solid #d0d8e4;" /></div>'
        if not imgs:
            return ""
        return f'<div style="margin-bottom:12px;"><div style="font-size:10px;color:#1e4d8c;font-weight:700;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:4px;">{titulo}</div>{imgs}</div>'

    _fotos_content = (
        _b64_to_html(fotos_resolvidas.get("antes", []),  "Antes do tratamento") +
        _b64_to_html(fotos_resolvidas.get("depois", []), "Depois do tratamento") +
        _b64_to_html(fotos_resolvidas.get("cmaq", []),   "Casa de máquinas")
    )

    fotos_html_section = f'''
  <div class="card">
    <div class="sec-title">Registro fotográfico</div>
    {_fotos_content if _fotos_content else '<p style="font-size:12px;color:#8a9ab0;font-style:italic;">Nenhuma foto registrada nesta visita.</p>'}
  </div>'''

    assinatura_resp_b64 = _normalizar_assinatura_b64(lancamento.get("assinatura_responsavel_b64", ""))
    resp_nome_ass = (lancamento.get("assinatura_responsavel_nome") or lancamento.get("resp_local") or "").strip()
    assinatura_responsavel_html = ""
    if assinatura_resp_b64:
        assinatura_responsavel_html = f'''
  <div class="card">
    <div class="sec-title">Assinatura do responsável no local</div>
    <div class="obs-txt" style="font-style:normal;padding-bottom:8px;">{resp_nome_ass or "Responsável presente na visita"} • {fmt(lancamento.get("assinatura_responsavel_data") or lancamento.get("data", ""))}</div>
    <div style="text-align:center;border:1px solid #d0d8e4;border-radius:10px;padding:14px;background:#fcfdff;">
      <img src="data:image/png;base64,{assinatura_resp_b64}" style="max-width:100%;max-height:150px;object-fit:contain;" />
    </div>
  </div>'''

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Relatório de Visita — {nome_condominio}</title>
<style>
  *{{box-sizing:border-box;margin:0;padding:0;}}
  body{{font-family:Arial,Helvetica,sans-serif;background:#f4f6f9;color:#1a2a4a;}}
  .page{{max-width:640px;margin:0 auto;padding:16px;}}
  .card{{background:#fff;border:1px solid #d0d8e4;border-radius:12px;padding:18px 20px;margin-bottom:12px;}}
  .hdr-top{{display:flex;justify-content:space-between;align-items:flex-start;}}
  .hdr-logo{{display:flex;align-items:center;gap:12px;}}
  .logo-ball{{width:48px;height:48px;border-radius:50%;background:#1e4d8c;display:flex;align-items:center;justify-content:center;font-size:13px;font-weight:700;color:#fff;flex-shrink:0;}}
  .hdr-empresa{{font-size:16px;font-weight:700;color:#1a2a4a;letter-spacing:0.3px;}}
  .hdr-sub{{font-size:10px;color:#8a9ab0;letter-spacing:0.8px;text-transform:uppercase;margin-top:2px;}}
  .hdr-right{{text-align:right;}}
  .doc-titulo{{font-size:14px;font-weight:700;color:#1e4d8c;}}
  .doc-num{{font-size:10px;color:#8a9ab0;margin-top:2px;}}
  hr{{border:none;border-top:1px solid #d0d8e4;margin:12px 0;}}
  .info-grid{{display:grid;grid-template-columns:1fr 1fr;gap:8px 16px;}}
  .info-lbl{{font-size:10px;color:#8a9ab0;text-transform:uppercase;letter-spacing:0.5px;}}
  .info-val{{font-size:13px;color:#1a2a4a;font-weight:600;margin-top:2px;}}
  .sec-title{{font-size:10px;font-weight:700;color:#1e4d8c;text-transform:uppercase;letter-spacing:0.8px;margin-bottom:12px;padding-bottom:6px;border-bottom:2px solid #1e4d8c;}}
  .param-grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:8px;}}
  .param-box{{border:1px solid #d0d8e4;border-radius:8px;padding:10px 8px;text-align:center;}}
  .param-box.ok{{border-color:#2e7d32;background:#f1f8f1;}}
  .param-box.warn{{border-color:#e65100;background:#fff8f0;}}
  .param-box.nd{{border-color:#d0d8e4;background:#f8f9fb;}}
  .pnm{{font-size:9px;color:#8a9ab0;text-transform:uppercase;letter-spacing:0.4px;margin-bottom:4px;}}
  .q15{{background:#e8f0fb;color:#1e4d8c;border-radius:4px;padding:1px 4px;font-size:8px;font-weight:700;}}
  .pval{{font-size:20px;font-weight:700;color:#1a2a4a;margin:2px 0;}}
  .param-box.ok .pval{{color:#2e7d32;}}
  .param-box.warn .pval{{color:#e65100;}}
  .param-box.nd .pval{{color:#b0bec5;}}
  .pst{{font-size:9px;}}
  .param-box.ok .pst{{color:#388e3c;}}
  .param-box.warn .pst{{color:#e65100;}}
  .param-box.nd .pst{{color:#b0bec5;font-style:italic;}}
  .alerta{{display:flex;align-items:flex-start;gap:10px;padding:10px 12px;border-radius:8px;background:#fff8f0;border:1px solid #e65100;margin-bottom:8px;}}
  .alerta.ok-all{{background:#f1f8f1;border-color:#2e7d32;}}
  .alerta-icon{{width:16px;height:16px;border-radius:50%;background:#e65100;display:flex;align-items:center;justify-content:center;font-size:10px;color:#fff;font-weight:700;flex-shrink:0;margin-top:1px;}}
  .alerta-txt{{font-size:12px;color:#b84200;line-height:1.5;}}
  .alerta.ok-all .alerta-txt{{color:#2e7d32;}}
  .dos-row{{display:flex;justify-content:space-between;align-items:center;padding:8px 0;border-bottom:1px solid #eef1f5;}}
  .dos-row:last-child{{border-bottom:none;}}
  .dos-nome{{font-size:13px;color:#1a2a4a;font-weight:600;}}
  .dos-detalhe{{font-size:11px;color:#8a9ab0;text-align:right;}}
  .obs-txt{{font-size:13px;color:#4a5568;line-height:1.7;font-style:italic;padding:4px 0;}}
  .assin-bloco{{display:flex;justify-content:space-between;align-items:flex-end;padding-top:12px;margin-top:6px;border-top:1px solid #d0d8e4;}}
  .assin-esq{{font-size:11px;color:#4a5568;line-height:1.8;}}
  .assin-esq strong{{color:#1a2a4a;font-size:13px;}}
  .crq-badge{{display:inline-block;font-size:9px;color:#1e4d8c;background:#e8f0fb;border:1px solid #b5d0f0;border-radius:99px;padding:2px 8px;margin-top:4px;}}
  .assin-dir{{text-align:center;}}
  .assin-linha{{border-top:1px solid #1a2a4a;width:140px;margin-bottom:4px;}}
  .assin-nome{{font-size:10px;color:#8a9ab0;}}
  .rodape{{text-align:center;font-size:10px;color:#8a9ab0;padding:8px 0 4px;}}
  @media print{{
    body{{background:#fff;}}
    .page{{padding:0;}}
    .card{{border:1px solid #ccc;border-radius:0;box-shadow:none;page-break-inside:avoid;}}
  }}
</style>
</head>
<body>
<div class="page">

  <div class="card">
    <div class="hdr-top">
      <div class="hdr-logo">
        <div class="logo-ball">RT</div>
        <div>
          <div class="hdr-empresa">AQUA GESTÃO</div>
          <div class="hdr-sub">Controle Técnico de Piscinas</div>
        </div>
      </div>
      <div class="hdr-right">
        <div class="doc-titulo">Relatório de Visita</div>
        <div class="doc-num">Emitido em {data_hoje}</div>
      </div>
    </div>
    <hr>
    <div class="info-grid">
      <div><div class="info-lbl">Condomínio / Local</div><div class="info-val">{nome_condominio}</div></div>
      <div><div class="info-lbl">Data da visita</div><div class="info-val">{fmt(lancamento.get("data",""))}</div></div>
      <div><div class="info-lbl">Operador</div><div class="info-val">{operador}</div></div>
      <div><div class="info-lbl">Responsável técnico</div><div class="info-val">Thyago F. Silveira</div></div>
    </div>
  </div>

  {piscinas_html_section}

  <div class="card">
    <div class="sec-title">Alertas técnicos</div>
    {alertas_html}
  </div>

  <div class="card">
    <div class="sec-title">Dosagens aplicadas</div>
    {dos_html}
  </div>

  <div class="card">
    <div class="sec-title">Observações</div>
    {obs_html}
  </div>

  {fotos_html_section}

  {assinatura_responsavel_html}

  {problemas_html}

  <div class="card">
    <div class="sec-title">Responsabilidade técnica</div>
    <div class="assin-bloco">
      <div class="assin-esq">
        <strong>Thyago Fernando da Silveira</strong><br>
        Técnico em Química · NR-26 · NR-6<br>
        <span class="crq-badge">CRQ-MG 2ª Região · CRQ 024025748</span>
      </div>
      <div class="assin-dir">
        <div class="assin-linha"></div>
        <div class="assin-nome">Assinatura / carimbo RT</div>
      </div>
    </div>
  </div>

  <div class="rodape">
    Aqua Gestão – Controle Técnico de Piscinas · Documento de uso operacional
  </div>

</div>
</body>
</html>"""
    return html




def gerar_pdf_relatorio_visita_bem_star(lancamento: dict, nome_condominio: str) -> bytes:
    """Gera PDF premium do relatorio de visita operacional da Bem Star Piscinas."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, HRFlowable, Image, PageBreak)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    try:
        pdfmetrics.registerFont(TTFont("BSans",      "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"))
        pdfmetrics.registerFont(TTFont("BSans-Bold", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"))
        _font_reg  = "BSans"
        _font_bold = "BSans-Bold"
    except Exception:
        _font_reg  = "Helvetica"
        _font_bold = "Helvetica-Bold"

    W, H = A4
    M  = 18*mm
    CW = W - 2*M

    NAVY    = colors.HexColor("#0D2A4A")
    TEAL    = colors.HexColor("#0E7490")
    GOLD    = colors.HexColor("#C8A951")
    CINZA   = colors.HexColor("#4A5568")
    CINZA_L = colors.HexColor("#F7F8FA")
    CINZA_M = colors.HexColor("#E2E8F0")
    VERDE   = colors.HexColor("#1B7A1B")
    LARANJA = colors.HexColor("#C05621")
    BRANCO  = colors.white
    PRETO   = colors.black

    def E(n, **k): return ParagraphStyle(n, fontName=_font_reg, **k)
    def P(t, s):   return Paragraph(str(t or ""), s)
    def Sp(h=3):   return Spacer(1, h*mm)
    def HR():      return HRFlowable(width="100%", thickness=0.5, color=CINZA_M, spaceAfter=3, spaceBefore=3)
    def HR_GOLD(): return HRFlowable(width="100%", thickness=1.5, color=GOLD,    spaceAfter=5, spaceBefore=2)

    E_BODY  = E("BO",  fontSize=9,  textColor=PRETO,  spaceAfter=4, leading=14, alignment=TA_JUSTIFY)
    E_CELL  = E("CE",  fontSize=8,  textColor=CINZA,  alignment=TA_CENTER, leading=11)
    E_CELL_L= E("CL",  fontSize=8,  textColor=CINZA,  alignment=TA_LEFT,   leading=11)
    E_CELL_B= E("CB",  fontSize=8,  textColor=NAVY,   fontName=_font_bold, alignment=TA_LEFT, leading=11)
    E_TH    = E("TH",  fontSize=8,  textColor=BRANCO, fontName=_font_bold, alignment=TA_CENTER, leading=11)
    E_TH_L  = E("TL",  fontSize=8,  textColor=BRANCO, fontName=_font_bold, alignment=TA_LEFT,   leading=11)
    E_SEC   = E("SE",  fontSize=9,  textColor=NAVY,   fontName=_font_bold, spaceBefore=8, spaceAfter=3, leading=12)
    E_OK    = E("OK",  fontSize=8,  textColor=VERDE,  fontName=_font_bold, leading=11)
    E_WARN  = E("WA",  fontSize=8,  textColor=LARANJA,fontName=_font_bold, leading=11)
    E_FOTO  = E("FT",  fontSize=8,  textColor=TEAL,   fontName=_font_bold, leading=11, spaceAfter=2)
    E_CAMPO = E("CA",  fontSize=8,  textColor=colors.HexColor("#718096"), leading=11)
    E_AVISO = E("AV",  fontSize=8,  textColor=colors.HexColor("#7A4000"),
                backColor=colors.HexColor("#FFFBEB"), leading=12, alignment=TA_JUSTIFY)
    E_CAPA_TAG  = E("CT", fontSize=9,  textColor=GOLD,  alignment=TA_CENTER, spaceAfter=6, leading=12)
    E_CAPA_MAIN = E("CM", fontSize=22, textColor=BRANCO, fontName=_font_bold, alignment=TA_CENTER, spaceAfter=4, leading=26)
    E_CAPA_SUB  = E("CS", fontSize=10, textColor=colors.HexColor("#A8D4E0"), alignment=TA_CENTER, spaceAfter=14, leading=15)
    E_CAPA_INFO = E("CI", fontSize=8,  textColor=colors.HexColor("#8FC8CC"), alignment=TA_CENTER, leading=13)
    E_RODAPE    = E("RO", fontSize=7,  textColor=colors.HexColor("#A0AEC0"), alignment=TA_CENTER)

    data_visita = lancamento.get("data", datetime.now().strftime("%d/%m/%Y"))
    operador    = lancamento.get("operador", "")
    observacao  = lancamento.get("observacao", "")
    problemas   = lancamento.get("problemas", "")
    piscinas    = lancamento.get("piscinas", [])
    dosagens    = lancamento.get("dosagens", [])
    servicos_exec = lancamento.get("servicos_executados", [])
    sindico     = lancamento.get("sindico", "")
    endereco    = lancamento.get("endereco", "")

    def capa_fn(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(NAVY)
        canvas.rect(0, 0, W, H, fill=1, stroke=0)
        canvas.setFillColor(GOLD)
        canvas.rect(0, H-3*mm, W, 3*mm, fill=1, stroke=0)
        canvas.rect(0, 0, W, 3*mm, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor("#0A2540"))
        canvas.roundRect(M, H*0.28, W-2*M, H*0.46, 8, fill=1, stroke=0)
        canvas.setFillColor(TEAL)
        canvas.rect(M, H*0.28 + H*0.46 - 2*mm, W-2*M, 2*mm, fill=1, stroke=0)
        canvas.restoreState()

    def hf_fn(canvas, doc):
        canvas.saveState()
        if doc.page > 1:
            canvas.setFillColor(NAVY)
            canvas.rect(0, H-12*mm, W, 12*mm, fill=1, stroke=0)
            canvas.setFillColor(GOLD)
            canvas.rect(0, H-12.8*mm, W, 0.8*mm, fill=1, stroke=0)
            canvas.setFillColor(BRANCO)
            canvas.setFont(_font_bold, 7.5)
            canvas.drawString(M, H-6*mm, "BEM STAR PISCINAS - RELATORIO DE VISITA OPERACIONAL")
            canvas.setFont(_font_reg, 7)
            canvas.drawRightString(W-M, H-6*mm, f"Pagina {doc.page}")
            canvas.drawString(M, H-10*mm, "Documento operacional - nao substitui Relatorio de Responsabilidade Tecnica")
            canvas.setFillColor(CINZA_L)
            canvas.rect(0, 0, W, 10*mm, fill=1, stroke=0)
            canvas.setFillColor(NAVY)
            canvas.rect(0, 9.5*mm, W, 0.5*mm, fill=1, stroke=0)
            canvas.setFillColor(colors.HexColor("#718096"))
            canvas.setFont(_font_reg, 7)
            canvas.drawString(M, 4*mm, "Bem Star Piscinas Ltda  |  CNPJ 26.799.958/0001-88  |  Uberlandia/MG")
            canvas.drawRightString(W-M, 4*mm, f"Pagina {doc.page}")
        canvas.restoreState()

    def tabela(rows, cws, cor=NAVY, stripe=True):
        t = Table(rows, colWidths=cws)
        s = [("BACKGROUND",(0,0),(-1,0),cor),("GRID",(0,0),(-1,-1),0.4,CINZA_M),
             ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
             ("LEFTPADDING",(0,0),(-1,-1),6),("VALIGN",(0,0),(-1,-1),"MIDDLE")]
        if stripe:
            for i in range(1, len(rows)):
                if i % 2 == 0: s.append(("BACKGROUND",(0,i),(-1,i),CINZA_L))
        t.setStyle(TableStyle(s))
        return t

    def sec_hdr(titulo):
        t = Table([[P(titulo, E("SH", fontSize=9, textColor=BRANCO, fontName=_font_bold, leading=12))]],
                  colWidths=[CW])
        t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),NAVY),
            ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
            ("LEFTPADDING",(0,0),(-1,-1),8)]))
        return t

    def bloco(titulo, corpo, cf=colors.HexColor("#EBF8FF"), cb=colors.HexColor("#3182CE")):
        t = Table([[
            P(titulo, E("BT", fontSize=8, textColor=cb, fontName=_font_bold, leading=12)),
            P(corpo,  E("BC2", fontSize=8, textColor=CINZA, leading=13, alignment=TA_JUSTIFY)),
        ]], colWidths=[CW*0.22, CW*0.78])
        t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),cf),("BOX",(0,0),(-1,-1),1,cb),
            ("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),
            ("LEFTPADDING",(0,0),(-1,-1),8),("VALIGN",(0,0),(-1,-1),"TOP")]))
        return t

    def status_param(val, minv, maxv):
        try:
            v = float(str(val).replace(",","."))
            if minv <= v <= maxv: return P("Conforme", E_OK)
            return P("Fora da faixa", E_WARN)
        except: return P("-", E_CELL)

    story = []

    # CAPA
    story += [Sp(30),
        P("RELATORIO DE VISITA TECNICA OPERACIONAL", E_CAPA_TAG),
        P("BEM STAR PISCINAS", E_CAPA_MAIN),
        P("Manutencao e Limpeza Operacional com Inteligencia Quimica Hidrica", E_CAPA_SUB),
        Sp(3), P(f"Condominio: {nome_condominio}  |  Data: {data_visita}", E_CAPA_INFO),
        Sp(2), P("Uberlandia/MG  |  Bem Star Piscinas Ltda", E_CAPA_INFO),
        PageBreak()]

    # APRESENTACAO
    story += [Sp(3), sec_hdr("SOBRE ESTE DOCUMENTO"), Sp(3)]
    story.append(bloco("O que e este relatorio",
        "Este Relatorio de Visita Tecnica Operacional e emitido pela Bem Star Piscinas Ltda para registro "
        "das atividades de limpeza, manutencao e controle basico de qualidade da agua realizadas em cada visita."))
    story.append(Sp(2))
    story.append(bloco("Importante",
        "Este documento NAO constitui Relatorio de Responsabilidade Tecnica (RT) e NAO substitui a documentacao "
        "exigida pela Vigilancia Sanitaria ou pelo CRQ. Para conformidade normativa, emissao de ART e Relatorio "
        "Mensal de RT, e necessario contratar a Aqua Gestao - Controle Tecnico de Piscinas (CRQ 024025748).",
        cf=colors.HexColor("#FFF8F0"), cb=colors.HexColor("#F6AD55")))
    story.append(Sp(2))
    story.append(bloco("Aqua Gestao RT",
        "Para conformidade com ABNT NBR 10339, Resolucoes CFQ e Vigilancia Sanitaria: "
        "Thyago Fernando da Silveira | Tecnico em Quimica | CRQ 024025748 | Uberlandia/MG.",
        cf=colors.HexColor("#F0FFF4"), cb=colors.HexColor("#38A169")))
    story.append(Sp(4))

    # IDENTIFICACAO
    story += [sec_hdr("IDENTIFICACAO DA VISITA"), Sp(2)]
    story.append(tabela([
        [P("Item", E_TH_L), P("Informacao", E_TH_L)],
        [P("Condominio / Local",   E_CELL_B), P(nome_condominio, E_CELL_L)],
        [P("Endereco",             E_CELL_B), P(endereco or "___________________________________", E_CAMPO)],
        [P("Responsavel/Sindico",  E_CELL_B), P(sindico  or "___________________________________", E_CAMPO)],
        [P("Data da Visita",       E_CELL_B), P(data_visita, E_CELL_L)],
        [P("Operador",             E_CELL_B), P(operador, E_CELL_L)],
        [P("Empresa",              E_CELL_B), P("Bem Star Piscinas Ltda", E_CELL_L)],
    ], [CW*0.32, CW*0.68]))
    story.append(Sp(4))

    # PARAMETROS por piscina
    story += [sec_hdr("ANALISE BASICA DA AGUA"), Sp(2),
        P("Verificacao basica de parametros com Photometer Color Q. Estes resultados sao orientativos. "
          "Para parecer normativo completo, consulte o Relatorio de RT da Aqua Gestao.", E_BODY), Sp(2)]

    if piscinas:
        for pisc in piscinas:
            pnome = pisc.get("nome", "Piscina")
            story.append(P(f"{pnome}",
                E("PN", fontSize=8, textColor=TEAL, fontName=_font_bold, spaceAfter=2, spaceBefore=4, leading=11)))
            params = [
                [P("Parametro",E_TH),P("Valor",E_TH),P("Faixa ideal",E_TH),P("Status",E_TH)],
                [P("pH",E_CELL_L),P(str(pisc.get("ph","") or "-"),E_CELL),P("7,2-7,8",E_CELL),status_param(pisc.get("ph"),7.2,7.8)],
                [P("CRL (mg/L)",E_CELL_L),P(str(pisc.get("cloro_livre","") or "-"),E_CELL),P("0,5-3,0",E_CELL),status_param(pisc.get("cloro_livre"),0.5,3.0)],
                [P("CT (mg/L)",E_CELL_L),P(str(pisc.get("cloro_total","") or "-"),E_CELL),P("0,5-5,0",E_CELL),status_param(pisc.get("cloro_total"),0.5,5.0)],
                [P("Alcalinidade",E_CELL_L),P(str(pisc.get("alcalinidade","") or "-"),E_CELL),P("80-120",E_CELL),status_param(pisc.get("alcalinidade"),80,120)],
                [P("CYA (mg/L)",E_CELL_L),P(str(pisc.get("cianurico","") or "-"),E_CELL),P("30-50",E_CELL),status_param(pisc.get("cianurico"),30,50)],
            ]
            story.append(tabela(params,[CW*0.30,CW*0.18,CW*0.22,CW*0.30]))
            story.append(Sp(2))
    else:
        # Dados diretos do lancamento (visita unica sem multiplas piscinas)
        ph  = lancamento.get("ph",""); crl = lancamento.get("cloro_livre","")
        ct  = lancamento.get("cloro_total",""); alc = lancamento.get("alcalinidade","")
        cya = lancamento.get("cianurico","")
        params = [
            [P("Parametro",E_TH),P("Valor",E_TH),P("Faixa ideal",E_TH),P("Status",E_TH)],
            [P("pH",E_CELL_L),P(str(ph or "-"),E_CELL),P("7,2-7,8",E_CELL),status_param(ph,7.2,7.8)],
            [P("CRL (mg/L)",E_CELL_L),P(str(crl or "-"),E_CELL),P("0,5-3,0",E_CELL),status_param(crl,0.5,3.0)],
            [P("CT (mg/L)",E_CELL_L),P(str(ct or "-"),E_CELL),P("0,5-5,0",E_CELL),status_param(ct,0.5,5.0)],
            [P("Alcalinidade",E_CELL_L),P(str(alc or "-"),E_CELL),P("80-120",E_CELL),status_param(alc,80,120)],
            [P("CYA (mg/L)",E_CELL_L),P(str(cya or "-"),E_CELL),P("30-50",E_CELL),status_param(cya,30,50)],
        ]
        story.append(tabela(params,[CW*0.30,CW*0.18,CW*0.22,CW*0.30]))
        story.append(Sp(2))

    story.append(Sp(2))

    # SERVICOS EXECUTADOS
    story += [sec_hdr("SERVICOS EXECUTADOS"), Sp(2)]
    if servicos_exec:
        servs = [[P("Servico",E_TH_L),P("Status",E_TH)]]
        for sv in servicos_exec:
            servs.append([P(sv,E_CELL_L),P("Realizado",E_OK)])
        story.append(tabela(servs,[CW*0.75,CW*0.25]))
    else:
        story.append(P("Servicos nao registrados individualmente nesta visita.", E_CAMPO))
    story.append(Sp(3))

    # DOSAGENS
    story += [sec_hdr("DOSAGENS APLICADAS"), Sp(2)]
    if dosagens:
        dos_rows = [[P("Produto",E_TH_L),P("Quantidade",E_TH),P("Unidade",E_TH),P("Finalidade",E_TH_L)]]
        for d in (dosagens if isinstance(dosagens, list) else []):
            if isinstance(d, dict):
                dos_rows.append([
                    P(d.get("produto",""),E_CELL_L),
                    P(str(d.get("quantidade","")),E_CELL),
                    P(d.get("unidade",""),E_CELL),
                    P(d.get("finalidade",""),E_CELL_L),
                ])
        if len(dos_rows) > 1:
            story.append(tabela(dos_rows,[CW*0.30,CW*0.14,CW*0.10,CW*0.46]))
        else:
            story.append(P(str(dosagens), E_BODY))
    else:
        story.append(P("Nenhuma dosagem registrada.", E_CAMPO))
    story.append(Sp(3))

    # PARECER OPERACIONAL
    story += [sec_hdr("PARECER OPERACIONAL"), Sp(2)]
    t_parecer = Table([[
        P("Condicoes operacionais gerais:",
          E("PT", fontSize=9, textColor=NAVY, fontName=_font_bold, leading=12)),
        P("SATISFATORIAS",
          E("PV", fontSize=11, textColor=VERDE, fontName=_font_bold, alignment=TA_CENTER, leading=14)),
    ]], colWidths=[CW*0.55, CW*0.45])
    t_parecer.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#F0FFF4")),
        ("BOX",(0,0),(-1,-1),1.5,VERDE),
        ("TOPPADDING",(0,0),(-1,-1),8),("BOTTOMPADDING",(0,0),(-1,-1),8),
        ("LEFTPADDING",(0,0),(-1,-1),10),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
    ]))
    story.append(t_parecer)
    story.append(Sp(3))

    if observacao:
        story.append(P(f"Observacoes: {observacao}", E_BODY))
    if problemas:
        prob = Table([[P("Ocorrencia:", E_WARN), P(problemas,
            E("OB", fontSize=8, textColor=PRETO, leading=13, alignment=TA_JUSTIFY))]],
            colWidths=[CW*0.20, CW*0.80])
        prob.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#FFF8F0")),
            ("BOX",(0,0),(-1,-1),1,colors.HexColor("#F6AD55")),
            ("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),
            ("LEFTPADDING",(0,0),(-1,-1),7),("VALIGN",(0,0),(-1,-1),"TOP"),
        ]))
        story.append(prob)
    story.append(Sp(3))

    # ASSINATURAS
    story += [sec_hdr("VALIDACAO E ASSINATURAS"), Sp(3)]
    ass = Table([[
        P(f"{operador or 'Operador'}\nBem Star Piscinas Ltda\n\n\n___________________________\nAssinatura\nData: {data_visita}",
          E("A1", fontSize=8, textColor=CINZA, alignment=TA_CENTER, leading=14)),
        P("", E("SP", fontSize=8)),
        P("Nome: ___________________________\nCargo: ___________________________\n\n\n"
          "___________________________\nRecebido por\nData: ____/____/________",
          E("A2", fontSize=8, textColor=CINZA, alignment=TA_CENTER, leading=14)),
    ]], colWidths=[CW*0.46, CW*0.08, CW*0.46])
    ass.setStyle(TableStyle([
        ("BOX",(0,0),(0,0),0.5,CINZA_M),("BOX",(2,0),(2,0),0.5,CINZA_M),
        ("BACKGROUND",(0,0),(0,0),CINZA_L),("BACKGROUND",(2,0),(2,0),CINZA_L),
        ("TOPPADDING",(0,0),(-1,-1),10),("BOTTOMPADDING",(0,0),(-1,-1),10),
        ("LEFTPADDING",(0,0),(-1,-1),10),("VALIGN",(0,0),(-1,-1),"TOP"),
    ]))
    story += [ass, Sp(4),
        P("Documento de uso operacional emitido pela Bem Star Piscinas Ltda. "
          "Este relatorio NAO constitui Relatorio de Responsabilidade Tecnica (RT) e NAO substitui "
          "a documentacao exigida pela Vigilancia Sanitaria ou pelo CRQ. "
          "Para conformidade normativa, contrate a Aqua Gestao - Controle Tecnico de Piscinas.",
          E_AVISO)]

    # FOTOS
    fotos_b64_antes  = lancamento.get("fotos_antes_b64",  [])
    fotos_b64_depois = lancamento.get("fotos_depois_b64", [])
    fotos_b64_cmaq   = lancamento.get("fotos_cmaq_b64",   [])
    fotos_b64_extras = lancamento.get("fotos_extras_b64", [])
    todas_fotos = (
        [("Antes do tratamento", f) for f in (fotos_b64_antes  or [])] +
        [("Apos o tratamento",   f) for f in (fotos_b64_depois or [])] +
        [("Casa de maquinas",    f) for f in (fotos_b64_cmaq   or [])] +
        [("Outras fotos",        f) for f in (fotos_b64_extras or [])]
    )
    if todas_fotos:
        story.append(PageBreak())
        story += [sec_hdr("REGISTRO FOTOGRAFICO"), HR_GOLD(), Sp(2),
            P("Registros visuais realizados durante a visita operacional.", E_BODY), Sp(3)]
        for legenda, foto_b64 in todas_fotos:
            try:
                import base64
                _img_bytes = base64.b64decode(foto_b64) if isinstance(foto_b64, str) else foto_b64
                import tempfile as _tmp
                _tmp_f = _tmp.NamedTemporaryFile(delete=False, suffix=".jpg")
                _tmp_f.write(_img_bytes); _tmp_f.close()
                story.append(P(legenda, E_FOTO))
                img = Image(_tmp_f.name, width=140*mm, height=90*mm)
                img.hAlign = "CENTER"
                story.append(img)
                story.append(Sp(4))
                import os; os.unlink(_tmp_f.name)
            except Exception:
                pass

    story += [HR(),
        P("Bem Star Piscinas Ltda  |  CNPJ 26.799.958/0001-88  |  Uberlandia/MG", E_RODAPE),
        P("Documento operacional. NAO substitui o Relatorio de Responsabilidade Tecnica.", E_RODAPE)]

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=M, rightMargin=M, topMargin=15*mm, bottomMargin=13*mm)
    doc.build(story, onFirstPage=capa_fn, onLaterPages=hf_fn)
    return buf.getvalue()

def gerar_pdf_relatorio_visita(lancamento: dict, nome_condominio: str) -> bytes:
    """Gera PDF do relatório de visita usando ReportLab. Retorna bytes do PDF."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    import base64 as _b64

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=1.8*cm, rightMargin=1.8*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm,
    )

    # Cores da marca
    AZUL_ESCURO = colors.HexColor("#1a2a4a")
    AZUL_MEDIO  = colors.HexColor("#1e4d8c")
    AZUL_CLARO  = colors.HexColor("#e8f0fb")
    CINZA       = colors.HexColor("#8a9ab0")
    VERDE_OK    = colors.HexColor("#2e7d32")
    VERDE_BG    = colors.HexColor("#f1f8f1")
    LARANJA     = colors.HexColor("#e65100")
    LARANJA_BG  = colors.HexColor("#fff8f0")
    BORDA       = colors.HexColor("#d0d8e4")

    styles = getSampleStyleSheet()

    def estilo(nome, **kw):
        return ParagraphStyle(nome, **kw)

    s_titulo    = estilo("titulo",    fontSize=16, textColor=AZUL_ESCURO, fontName="Helvetica-Bold", spaceAfter=2)
    s_sub       = estilo("sub",       fontSize=8,  textColor=CINZA,       fontName="Helvetica",      spaceAfter=6, leading=10)
    s_sec       = estilo("sec",       fontSize=8,  textColor=AZUL_MEDIO,  fontName="Helvetica-Bold", spaceAfter=6, leading=10)
    s_body      = estilo("body",      fontSize=9,  textColor=AZUL_ESCURO, fontName="Helvetica",      leading=13)
    s_body_sm   = estilo("body_sm",   fontSize=8,  textColor=CINZA,       fontName="Helvetica",      leading=11)
    s_alerta    = estilo("alerta",    fontSize=8,  textColor=LARANJA,     fontName="Helvetica",      leading=11)
    s_ok        = estilo("ok",        fontSize=8,  textColor=VERDE_OK,    fontName="Helvetica",      leading=11)
    s_center    = estilo("center",    fontSize=8,  textColor=CINZA,       fontName="Helvetica",      alignment=TA_CENTER)
    s_bold      = estilo("bold",      fontSize=9,  textColor=AZUL_ESCURO, fontName="Helvetica-Bold", leading=13)

    elems = []

    # ── CABEÇALHO ─────────────────────────────────────────────────────────────
    data_hoje = date.today().strftime("%d/%m/%Y")
    operador  = lancamento.get("operador","") or "—"

    hdr_data = [
        [Paragraph("<b>AQUA GESTÃO</b>", estilo("hdr1", fontSize=14, textColor=AZUL_ESCURO, fontName="Helvetica-Bold")),
         Paragraph(f"<b>Relatório de Visita</b><br/><font size=8 color='#8a9ab0'>Emitido em {data_hoje}</font>", estilo("hdr2", fontSize=11, textColor=AZUL_MEDIO, fontName="Helvetica-Bold", alignment=TA_RIGHT))],
        [Paragraph("Controle Técnico de Piscinas", s_sub), ""],
    ]
    t_hdr = Table(hdr_data, colWidths=["60%","40%"])
    t_hdr.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LINEBELOW", (0,1), (-1,1), 0.5, BORDA),
        ("BOTTOMPADDING", (0,1), (-1,1), 8),
    ]))
    elems.append(t_hdr)
    elems.append(Spacer(1, 6))

    # Info básica
    info_data = [
        [Paragraph("<font size=7 color='#8a9ab0'>CONDOMÍNIO / LOCAL</font><br/>" + f"<b>{nome_condominio}</b>", s_body),
         Paragraph("<font size=7 color='#8a9ab0'>DATA DA VISITA</font><br/>" + f"<b>{lancamento.get('data','—')}</b>", s_body)],
        [Paragraph("<font size=7 color='#8a9ab0'>OPERADOR</font><br/>" + f"<b>{operador}</b>", s_body),
         Paragraph("<font size=7 color='#8a9ab0'>RESP. TÉCNICO</font><br/><b>Thyago F. Silveira</b>", s_body)],
    ]
    t_info = Table(info_data, colWidths=["50%","50%"])
    t_info.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.3, BORDA),
        ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#f8fafd")),
        ("PADDING", (0,0), (-1,-1), 6),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [colors.white, colors.HexColor("#f8fafd")]),
    ]))
    elems.append(t_info)
    elems.append(Spacer(1, 10))

    # ── PISCINAS ──────────────────────────────────────────────────────────────
    piscinas_lista = lancamento.get("piscinas", [])
    if not piscinas_lista:
        piscinas_lista = [{
            "nome": "Piscina", "ph": lancamento.get("ph",""),
            "cloro_livre": lancamento.get("cloro_livre",""), "cloro_total": lancamento.get("cloro_total",""),
            "cloraminas": lancamento.get("cloraminas",""), "alcalinidade": lancamento.get("alcalinidade",""),
            "dureza": lancamento.get("dureza",""), "cianurico": lancamento.get("cianurico",""),
        }]

    PARAMS = [
        ("pH",          "ph",          7.2, 7.8,  False),
        ("CRL mg/L",    "cloro_livre", 0.5, 3.0,  False),
        ("CT mg/L",     "cloro_total", 0.5, 5.0,  False),
        ("Alc. mg/L",   "alcalinidade",80, 120,   True),
        ("Dureza mg/L", "dureza",      150, 300,   True),
        ("CYA mg/L",    "cianurico",   30,  50,    True),
    ]

    for pisc in piscinas_lista:
        elems.append(Paragraph(f"🏊 {pisc.get('nome','Piscina')} — Parâmetros analisados", s_sec))
        elems.append(HRFlowable(width="100%", thickness=1.5, color=AZUL_MEDIO, spaceAfter=4))

        param_rows = []
        header_row = [Paragraph("<b>Parâmetro</b>", s_body_sm),
                      Paragraph("<b>Valor</b>", s_body_sm),
                      Paragraph("<b>Faixa ideal</b>", s_body_sm),
                      Paragraph("<b>Status</b>", s_body_sm),
                      Paragraph("<b>Obs</b>", s_body_sm)]
        param_rows.append(header_row)

        faixas_txt = {"pH":"7,2–7,8","CRL mg/L":"0,5–3,0","CT mg/L":"0,5–5,0","Alc. mg/L":"80–120","Dureza mg/L":"150–300","CYA mg/L":"30–50"}
        row_colors = []

        for label, key, mn, mx, quinzenal in PARAMS:
            val_raw = pisc.get(key, "")
            v = valor_float(val_raw)
            q_txt = " (15d)" if quinzenal else ""
            if v is None:
                status_txt = "Não medido"
                val_fmt = "—"
                bg = colors.white
            elif v < mn or v > mx:
                status_txt = "⚠ Fora da faixa"
                val_fmt = str(val_raw).replace(".", ",")
                bg = LARANJA_BG
            else:
                status_txt = "✓ Conforme"
                val_fmt = str(val_raw).replace(".", ",")
                bg = VERDE_BG
            row_colors.append(bg)
            param_rows.append([
                Paragraph(f"{label}{q_txt}", s_body_sm),
                Paragraph(f"<b>{val_fmt}</b>", s_body),
                Paragraph(faixas_txt.get(label,"—"), s_body_sm),
                Paragraph(status_txt, s_ok if "Conforme" in status_txt else (s_alerta if "Fora" in status_txt else s_body_sm)),
                Paragraph("", s_body_sm),
            ])

        # Cloraminas
        clor_raw = pisc.get("cloraminas","")
        v_cl = valor_float(clor_raw)
        if v_cl is not None:
            bg_cl = VERDE_BG if v_cl <= 0.2 else LARANJA_BG
            st_cl = "✓ Conforme" if v_cl <= 0.2 else "⚠ Fora da faixa"
            row_colors.append(bg_cl)
            param_rows.append([
                Paragraph("Cloraminas", s_body_sm),
                Paragraph(f"<b>{str(clor_raw).replace('.', ',')}</b>", s_body),
                Paragraph("≤ 0,2", s_body_sm),
                Paragraph(st_cl, s_ok if "Conforme" in st_cl else s_alerta),
                Paragraph("", s_body_sm),
            ])

        t_param = Table(param_rows, colWidths=["22%","15%","20%","28%","15%"])
        ts = [
            ("GRID", (0,0), (-1,-1), 0.3, BORDA),
            ("BACKGROUND", (0,0), (-1,0), AZUL_CLARO),
            ("TEXTCOLOR", (0,0), (-1,0), AZUL_MEDIO),
            ("PADDING", (0,0), (-1,-1), 5),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,0), 7),
        ]
        for i, bg in enumerate(row_colors):
            ts.append(("BACKGROUND", (0, i+1), (-1, i+1), bg))
        t_param.setStyle(TableStyle(ts))
        elems.append(t_param)
        elems.append(Spacer(1, 10))

    # ── ALERTAS ───────────────────────────────────────────────────────────────
    alertas_gerais = []
    for pisc in piscinas_lista:
        for val_r, mn, mx, rot in [
            (pisc.get("ph",""), 7.2, 7.8, "pH"),
            (pisc.get("cloro_livre",""), 0.5, 3.0, "CRL"),
            (pisc.get("alcalinidade",""), 80, 120, "Alcalinidade"),
            (pisc.get("dureza",""), 150, 300, "Dureza DC"),
            (pisc.get("cianurico",""), 30, 50, "CYA"),
        ]:
            v = valor_float(val_r)
            if v is not None and (v < mn or v > mx):
                alertas_gerais.append(f"⚠ {pisc.get('nome','Piscina')} — {rot}: {str(val_r).replace('.', ',')} — fora da faixa ideal.")

    if alertas_gerais:
        elems.append(Paragraph("Alertas técnicos", s_sec))
        elems.append(HRFlowable(width="100%", thickness=1.5, color=AZUL_MEDIO, spaceAfter=4))
        for a in alertas_gerais:
            elems.append(Paragraph(a, s_alerta))
            elems.append(Spacer(1, 3))
        elems.append(Spacer(1, 6))

    # ── DOSAGENS ──────────────────────────────────────────────────────────────
    dosagens = lancamento.get("dosagens", [])
    dosagens = [d for d in dosagens if d.get("produto","").strip()]
    if dosagens:
        elems.append(Paragraph("Dosagens aplicadas", s_sec))
        elems.append(HRFlowable(width="100%", thickness=1.5, color=AZUL_MEDIO, spaceAfter=4))
        dos_rows = [[
            Paragraph("<b>Produto</b>", s_body_sm),
            Paragraph("<b>Quantidade</b>", s_body_sm),
            Paragraph("<b>Finalidade</b>", s_body_sm),
        ]]
        for d in dosagens:
            qtd = f"{d.get('quantidade','')} {d.get('unidade','')}".strip()
            dos_rows.append([
                Paragraph(d.get("produto",""), s_body),
                Paragraph(qtd, s_body_sm),
                Paragraph(d.get("finalidade",""), s_body_sm),
            ])
        t_dos = Table(dos_rows, colWidths=["40%","25%","35%"])
        t_dos.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.3, BORDA),
            ("BACKGROUND", (0,0), (-1,0), AZUL_CLARO),
            ("TEXTCOLOR", (0,0), (-1,0), AZUL_MEDIO),
            ("PADDING", (0,0), (-1,-1), 5),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f8fafd")]),
        ]))
        elems.append(t_dos)
        elems.append(Spacer(1, 10))

    # ── PROBLEMAS ─────────────────────────────────────────────────────────────
    problemas = lancamento.get("problemas","").strip()
    if problemas:
        elems.append(Paragraph("Problemas / Ocorrências", s_sec))
        elems.append(HRFlowable(width="100%", thickness=1.5, color=AZUL_MEDIO, spaceAfter=4))
        problemas_formatados = problemas.replace("\n", "<br/>")
        t_prob = Table([[Paragraph(f"⚠ {problemas_formatados}", s_alerta)]], colWidths=["100%"])
        t_prob.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), LARANJA_BG),
            ("BOX", (0,0), (-1,-1), 0.5, LARANJA),
            ("PADDING", (0,0), (-1,-1), 8),
            ("RADIUS", (0,0), (-1,-1), 4),
        ]))
        elems.append(t_prob)
        elems.append(Spacer(1, 10))

    # ── OBSERVAÇÃO ────────────────────────────────────────────────────────────
    obs = lancamento.get("observacao","").strip()
    if obs:
        elems.append(Paragraph("Observações", s_sec))
        elems.append(HRFlowable(width="100%", thickness=1.5, color=AZUL_MEDIO, spaceAfter=4))
        elems.append(Paragraph(f'"{obs}"', s_body_sm))
        elems.append(Spacer(1, 10))

    # ── FOTOS ─────────────────────────────────────────────────────────────────
    def _add_fotos_b64(b64_list, titulo):
        """Adiciona fotos ao PDF a partir de lista de base64 — 1 por linha, tamanho máximo."""
        if not b64_list:
            return
        import io as _io
        from reportlab.platypus import Image as RLImage
        from PIL import Image as _PILR, ImageOps as _IOps
        elems.append(Paragraph(titulo, s_body_sm))
        LARGURA_MAX = 15 * cm   # largura útil da página
        ALTURA_MAX  = 18 * cm   # altura máxima por foto
        for b64_str in b64_list[:6]:
            try:
                fb = _b64.b64decode(b64_str)
                # Aplica rotação EXIF antes de medir
                _pil = _PILR.open(_io.BytesIO(fb))
                _pil = _IOps.exif_transpose(_pil)
                _iw, _ih = _pil.size
                # Salva versão corrigida
                _buf_corr = _io.BytesIO()
                _pil.convert("RGB").save(_buf_corr, format="JPEG", quality=85)
                _buf_corr.seek(0)
                # Calcula dimensões mantendo proporção
                ratio = _iw / _ih
                if ratio >= 1:  # paisagem
                    w = LARGURA_MAX
                    h = min(w / ratio, ALTURA_MAX)
                    w = h * ratio
                else:  # retrato
                    h = ALTURA_MAX
                    w = min(h * ratio, LARGURA_MAX)
                    h = w / ratio
                img = RLImage(_buf_corr, width=w, height=h)
                t_foto = Table([[img]], colWidths=[LARGURA_MAX])
                t_foto.setStyle(TableStyle([
                    ("ALIGN",   (0,0), (-1,-1), "CENTER"),
                    ("VALIGN",  (0,0), (-1,-1), "MIDDLE"),
                    ("PADDING", (0,0), (-1,-1), 2),
                ]))
                elems.append(t_foto)
                elems.append(Spacer(1, 8))
            except Exception:
                pass

    fotos_resolvidas = _resolver_fotos_visita_para_relatorio(lancamento)
    fotos_antes_b64  = fotos_resolvidas.get("antes", [])
    fotos_depois_b64 = fotos_resolvidas.get("depois", [])
    fotos_cmaq_b64   = fotos_resolvidas.get("cmaq", [])
    fotos_extras_b64 = fotos_resolvidas.get("extras", [])

    if fotos_antes_b64 or fotos_depois_b64 or fotos_cmaq_b64 or fotos_extras_b64:
        elems.append(Paragraph("Registro fotográfico", s_sec))
        elems.append(HRFlowable(width="100%", thickness=1.5, color=AZUL_MEDIO, spaceAfter=4))
        _add_fotos_b64(fotos_antes_b64,  "Antes do tratamento:")
        _add_fotos_b64(fotos_depois_b64, "Depois do tratamento:")
        _add_fotos_b64(fotos_cmaq_b64,   "Casa de máquinas:")
        _add_fotos_b64(fotos_extras_b64, "Outras fotos:")
        elems.append(Spacer(1, 6))

    assinatura_resp_b64 = _normalizar_assinatura_b64(lancamento.get("assinatura_responsavel_b64", ""))
    if assinatura_resp_b64:
        elems.append(Paragraph("Assinatura do responsável no local", s_sec))
        elems.append(HRFlowable(width="100%", thickness=1.0, color=AZUL_MEDIO, spaceAfter=4))
        try:
            import io as _io
            from reportlab.platypus import Image as RLImage
            from PIL import Image as _PILS

            _sig_bytes = _b64.b64decode(assinatura_resp_b64)
            _sig_img = _PILS.open(_io.BytesIO(_sig_bytes))
            _sw, _sh = _sig_img.size
            _max_w = 8.5 * cm
            _max_h = 2.8 * cm
            _ratio = (_sw / _sh) if _sh else 3.0
            _w = min(_max_w, _max_h * _ratio)
            _h = _w / _ratio if _ratio else _max_h
            if _h > _max_h:
                _h = _max_h
                _w = _h * _ratio
            _rl_sig = RLImage(_io.BytesIO(_sig_bytes), width=_w, height=_h)
            _resp_nome = (lancamento.get("assinatura_responsavel_nome") or lancamento.get("resp_local") or "Responsável no local").strip()
            _resp_data = lancamento.get("assinatura_responsavel_data") or lancamento.get("data", "")
            _tbl_sig = Table([[
                Paragraph(f"<font size=7 color='#8a9ab0'>RESPONSÁVEL NO LOCAL</font><br/><b>{_resp_nome}</b><br/><font size=7 color='#8a9ab0'>{_resp_data}</font>", s_body),
                _rl_sig,
            ]], colWidths=["48%", "52%"])
            _tbl_sig.setStyle(TableStyle([
                ("BOX", (0,0), (-1,-1), 0.4, BORDA),
                ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#fbfdff")),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                ("PADDING", (0,0), (-1,-1), 8),
                ("ALIGN", (1,0), (1,0), "CENTER"),
            ]))
            elems.append(_tbl_sig)
            elems.append(Spacer(1, 8))
        except Exception:
            pass

    # ── ASSINATURA RT ─────────────────────────────────────────────────────────
    elems.append(HRFlowable(width="100%", thickness=0.5, color=BORDA, spaceAfter=8))
    ass_data = [[
        Paragraph("<b>Thyago Fernando da Silveira</b><br/>Técnico em Química · NR-26 · NR-6<br/><font size=7 color='#1e4d8c'>CRQ-MG 2ª Região · CRQ 024025748</font>", s_body),
        Paragraph("<br/><br/>___________________________<br/><font size=7 color='#8a9ab0'>Assinatura / carimbo RT</font>", s_center),
    ]]
    t_ass = Table(ass_data, colWidths=["60%","40%"])
    t_ass.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "BOTTOM"),
        ("PADDING", (0,0), (-1,-1), 4),
    ]))
    elems.append(t_ass)

    elems.append(Spacer(1, 4))
    elems.append(Paragraph("Aqua Gestão – Controle Técnico de Piscinas · Documento de uso operacional", s_center))

    doc.build(elems)
    buffer.seek(0)
    return buffer.read()


def gerar_relatorio_visita_docx(
    output_path: Path,
    nome_local: str,
    cnpj: str,
    endereco: str,
    responsavel: str,
    operador: str,
    mes: str,
    ano: str,
    lancamentos: list,
    obs_geral: str = "",
    incluir_rt: bool = True,
    fotos: list = None,
) -> tuple[bool, str]:
    """
    Gera relatório técnico de visitas em DOCX — unificado para RT e sem RT.
    Se incluir_rt=True: inclui assinatura e dados do RT.
    Se incluir_rt=False: omite dados do RT (relatório sem RT).
    """
    try:
        from docx import Document as _DocxDoc
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = _DocxDoc()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        def _par(texto, bold=False, size=11, align=None, italic=False):
            p = doc.add_paragraph()
            if align: p.alignment = align
            r = p.add_run(texto)
            r.bold = bold; r.italic = italic
            r.font.size = Pt(size)
            return p

        def _tabela_info(dados: list):
            """Cria tabela de 2 colunas com rótulo → valor, com estilo visual."""
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            t = doc.add_table(rows=len(dados), cols=2)
            t.style = "Table Grid"
            # Largura das colunas
            for row in t.rows:
                row.cells[0].width = __import__("docx.shared", fromlist=["Cm"]).Cm(5)
                row.cells[1].width = __import__("docx.shared", fromlist=["Cm"]).Cm(12)
            for i, (rot, val) in enumerate(dados):
                c0 = t.cell(i, 0)
                c1 = t.cell(i, 1)
                # Fundo azul claro na coluna de rótulos
                tc_pr = c0._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EEF3FB")
                tc_pr.append(shd)
                r0 = c0.paragraphs[0].add_run(rot)
                r0.bold = True
                r0.font.size = Pt(10)
                r1 = c1.paragraphs[0].add_run(str(val or "—"))
                r1.font.size = Pt(10)
                # Padding
                for cell in [c0, c1]:
                    tc = cell._tc.get_or_add_tcPr()
                    tcMar = OxmlElement("w:tcMar")
                    for side in ["top","bottom","left","right"]:
                        m = OxmlElement(f"w:{side}")
                        m.set(qn("w:w"), "80")
                        m.set(qn("w:type"), "dxa")
                        tcMar.append(m)
                    tc.append(tcMar)
            doc.add_paragraph()

        # ── CABEÇALHO ─────────────────────────────────────────────────────────
        if incluir_rt:
            _par("AQUA GESTÃO – CONTROLE TÉCNICO DE PISCINAS", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
            _par(f"Responsável Técnico: {RESPONSAVEL_TÉCNICO} | {CRQ}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            _par(f"{QUALIFICACAO_RT} | Certificações: {CERTIFICACOES_RT}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            # Cabeçalho Bem Star com logo
            _logo_bs_hdr = encontrar_logo_bem_star()
            if _logo_bs_hdr and _logo_bs_hdr.exists():
                try:
                    p_hdr_logo = doc.add_paragraph()
                    p_hdr_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_hdr_logo.add_run().add_picture(str(_logo_bs_hdr), width=Inches(3.2))
                except Exception:
                    _par("BEM STAR PISCINAS", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _par("BEM STAR PISCINAS", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
            _par("RELATÓRIO TÉCNICO-OPERACIONAL DE PISCINAS", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
            _par(f"CNPJ: {CNPJ_BEM_STAR}  |  Uberlândia/MG", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            # Linha separadora visual
            p_linha = doc.add_paragraph()
            p_linha.paragraph_format.space_after = Pt(2)
            p_linha.add_run("─" * 72).font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0x0d, 0x3d, 0x75)
        doc.add_paragraph()

        # ── IDENTIFICAÇÃO ─────────────────────────────────────────────────────
        _par("1. IDENTIFICAÇÃO DO LOCAL", bold=True, size=11)
        # Coleta responsáveis no local das visitas (se disponível)
        _resps_locais = list(dict.fromkeys(
            lc.get("resp_local","") for lc in lancamentos if lc.get("resp_local","").strip()
        ))
        _resp_local_txt = " / ".join(_resps_locais) if _resps_locais else responsavel or "Não informado"
        _tabela_info([
            ("Local / Condomínio", nome_local),
            ("CNPJ", cnpj or "Não informado"),
            ("Endereço", endereco or "Não informado"),
            ("Responsável / Síndico", responsavel or "Não informado"),
            ("Responsável no local", _resp_local_txt),
            ("Operador de campo", operador or "Não informado"),
            ("Período de referência", f"{mes}/{ano}"),
        ])

        # ── ANÁLISES FÍSICO-QUÍMICAS — uma tabela por piscina ────────────────
        _par("2. ANÁLISES FÍSICO-QUÍMICAS", bold=True, size=11)

        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OXE
        from docx.shared import RGBColor as _RGB

        def _shd(cell, fill):
            tc = cell._tc.get_or_add_tcPr()
            s = _OXE("w:shd")
            s.set(_qn("w:val"), "clear"); s.set(_qn("w:color"), "auto")
            s.set(_qn("w:fill"), fill); tc.append(s)

        def _tabela_analises(nome_piscina: str, dados_linhas: list, vol_m3: str = ""):
            """Gera bloco de tabela de análises para uma piscina."""
            # Subtítulo da piscina com volume
            _vol_txt = f" — {vol_m3} m³" if vol_m3 and vol_m3 != "0" else ""
            _p_sub = doc.add_paragraph()
            _r_sub = _p_sub.add_run(f"🏊 {nome_piscina}{_vol_txt}")
            _r_sub.bold = True
            _r_sub.font.size = Pt(10)
            _r_sub.font.color.rgb = _RGB(0x0d, 0x3d, 0x75)

            cabecalho = ["Data", "pH", "CRL mg/L", "CT mg/L", "Clor. mg/L", "Alc. mg/L", "Dureza mg/L", "CYA mg/L", "Operador"]
            t = doc.add_table(rows=1 + len(dados_linhas), cols=len(cabecalho))
            t.style = "Table Grid"

            # Header azul
            for j, cab in enumerate(cabecalho):
                cell = t.cell(0, j)
                r = cell.paragraphs[0].add_run(cab)
                r.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = _RGB(0xFF, 0xFF, 0xFF)
                _shd(cell, "0D3D75")

            # Dados com zebra
            for i, linha in enumerate(dados_linhas):
                _fill = "EEF3FB" if i % 2 == 0 else "FFFFFF"
                for j, val in enumerate(linha):
                    cell = t.cell(i+1, j)
                    cell.paragraphs[0].add_run(str(val or "—"))
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    _shd(cell, _fill)
            doc.add_paragraph()

        # Coleta todas as piscinas distintas presentes nos lançamentos
        _piscinas_ordem = []
        _piscinas_vistas = set()
        for lc in lancamentos:
            _piscs = lc.get("piscinas", [])
            if _piscs:
                for p in _piscs:
                    _pn = p.get("nome", "Piscina").strip() or "Piscina"
                    if _pn not in _piscinas_vistas:
                        _piscinas_vistas.add(_pn)
                        _piscinas_ordem.append(_pn)
            else:
                if "Piscina" not in _piscinas_vistas:
                    _piscinas_vistas.add("Piscina")
                    _piscinas_ordem.append("Piscina")

        # Se só tem uma piscina genérica, usa o nome do local
        if _piscinas_ordem == ["Piscina"] and nome_local:
            _piscinas_ordem = ["Piscina"]

        # Gera uma tabela por piscina
        for _pisc_nome in _piscinas_ordem:
            _linhas_pisc = []
            for lc in lancamentos:
                _piscs = lc.get("piscinas", [])
                if _piscs:
                    for p in _piscs:
                        if (p.get("nome","Piscina").strip() or "Piscina") == _pisc_nome:
                            # Calcula cloraminas se disponível
                            _crl_v = valor_float(p.get("cloro_livre",""))
                            _ct_v  = valor_float(p.get("cloro_total",""))
                            _clor_txt = str(round(max(_ct_v - _crl_v, 0), 2)) if _crl_v is not None and _ct_v is not None else p.get("cloraminas","")
                            _linhas_pisc.append([
                                lc.get("data",""),
                                p.get("ph",""),
                                p.get("cloro_livre",""),
                                p.get("cloro_total",""),
                                _clor_txt,
                                p.get("alcalinidade",""),
                                p.get("dureza",""),
                                p.get("cianurico",""),
                                lc.get("operador",""),
                            ])
                elif _pisc_nome == "Piscina":
                    _crl_fb = valor_float(lc.get("cloro_livre",""))
                    _ct_fb  = valor_float(lc.get("cloro_total",""))
                    _clor_fb = str(round(max(_ct_fb - _crl_fb, 0), 2)) if _crl_fb is not None and _ct_fb is not None else lc.get("cloraminas","")
                    _linhas_pisc.append([
                        lc.get("data",""),
                        lc.get("ph",""),
                        lc.get("cloro_livre",""),
                        lc.get("cloro_total",""),
                        _clor_fb,
                        lc.get("alcalinidade",""),
                        lc.get("dureza",""),
                        lc.get("cianurico",""),
                        lc.get("operador",""),
                    ])
            if _linhas_pisc:
                _tabela_analises(_pisc_nome, _linhas_pisc)

        # ── DOSAGENS ──────────────────────────────────────────────────────────
        _par("3. DOSAGENS APLICADAS", bold=True, size=11)

        # Agrupa dosagens por piscina (se disponível) ou por visita
        # Estrutura: {piscina_nome: [{data, produto, qtd, un, fin}]}
        _dos_por_pisc = {}

        for lc in lancamentos:
            data_lc = lc.get("data","")
            _piscs = lc.get("piscinas", [])

            if _piscs:
                # Dosagens vinculadas a cada piscina
                for p in _piscs:
                    _pn = p.get("nome","Piscina").strip() or "Piscina"
                    _dos_p = p.get("dosagens", [])
                    # Se a piscina não tem dosagens próprias, usa as dosagens gerais da visita
                    if not _dos_p:
                        _dos_p = lc.get("dosagens", [])
                    for d in _dos_p:
                        if d.get("produto","").strip():
                            if _pn not in _dos_por_pisc:
                                _dos_por_pisc[_pn] = []
                            _dos_por_pisc[_pn].append({**d, "data": data_lc})
                # Se nenhuma piscina tem dosagens, coloca nas dosagens gerais
                if not any(_dos_por_pisc.get(p.get("nome","Piscina"),[]) for p in _piscs):
                    for d in lc.get("dosagens",[]):
                        if d.get("produto","").strip():
                            _pn_geral = _piscs[0].get("nome","Geral") if _piscs else "Geral"
                            if _pn_geral not in _dos_por_pisc:
                                _dos_por_pisc[_pn_geral] = []
                            _dos_por_pisc[_pn_geral].append({**d, "data": data_lc})
            else:
                # Sem piscinas — dosagens gerais
                for d in lc.get("dosagens",[]):
                    if d.get("produto","").strip():
                        if "Geral" not in _dos_por_pisc:
                            _dos_por_pisc["Geral"] = []
                        _dos_por_pisc["Geral"].append({**d, "data": data_lc})

        # Se todas as dosagens ficaram na mesma chave (sem distinção por piscina),
        # consolida em tabela única com header "Dosagens do período"
        _chaves_dos = [k for k, v in _dos_por_pisc.items() if v]

        if not _chaves_dos:
            _par("Nenhuma dosagem registrada no período.", size=10, italic=True)
        elif len(_chaves_dos) == 1:
            # Uma piscina ou geral — tabela única sem subtítulo de piscina
            _lista_dos = list(_dos_por_pisc.values())[0]
            t_dos = doc.add_table(rows=1 + len(_lista_dos), cols=5)
            t_dos.style = "Table Grid"
            for j, cab in enumerate(["Data", "Produto", "Qtd.", "Unidade", "Finalidade / Motivo"]):
                _c = t_dos.cell(0,j)
                _r = _c.paragraphs[0].add_run(cab)
                _r.bold = True; _r.font.size = Pt(9)
                _r.font.color.rgb = _RGB(0xFF,0xFF,0xFF)
                _shd(_c, "0D3D75")
            for i, d in enumerate(_lista_dos):
                _fill_d = "EEF3FB" if i % 2 == 0 else "FFFFFF"
                for j, val in enumerate([d.get("data",""), d.get("produto",""),
                        d.get("quantidade",""), d.get("unidade",""), d.get("finalidade","")]):
                    _cd = t_dos.cell(i+1,j)
                    _cd.paragraphs[0].add_run(str(val or "—"))
                    _cd.paragraphs[0].runs[0].font.size = Pt(9)
                    _shd(_cd, _fill_d)
        else:
            # Múltiplas piscinas — uma subseção por piscina
            for _pn_dos in _chaves_dos:
                _lista_dos = _dos_por_pisc[_pn_dos]
                if not _lista_dos:
                    continue
                # Subtítulo da piscina
                _p_sub_dos = doc.add_paragraph()
                _r_sub_dos = _p_sub_dos.add_run(f"🏊 {_pn_dos}")
                _r_sub_dos.bold = True; _r_sub_dos.font.size = Pt(10)
                _r_sub_dos.font.color.rgb = _RGB(0x0d, 0x3d, 0x75)
                # Tabela
                t_dos = doc.add_table(rows=1 + len(_lista_dos), cols=5)
                t_dos.style = "Table Grid"
                for j, cab in enumerate(["Data", "Produto", "Qtd.", "Unidade", "Finalidade / Motivo"]):
                    _c = t_dos.cell(0,j)
                    _r = _c.paragraphs[0].add_run(cab)
                    _r.bold = True; _r.font.size = Pt(9)
                    _r.font.color.rgb = _RGB(0xFF,0xFF,0xFF)
                    _shd(_c, "0D3D75")
                for i, d in enumerate(_lista_dos):
                    _fill_d = "EEF3FB" if i % 2 == 0 else "FFFFFF"
                    for j, val in enumerate([d.get("data",""), d.get("produto",""),
                            d.get("quantidade",""), d.get("unidade",""), d.get("finalidade","")]):
                        _cd = t_dos.cell(i+1,j)
                        _cd.paragraphs[0].add_run(str(val or "—"))
                        _cd.paragraphs[0].runs[0].font.size = Pt(9)
                        _shd(_cd, _fill_d)
                doc.add_paragraph()

        doc.add_paragraph()

        # ── PROBLEMAS / OCORRÊNCIAS ───────────────────────────────────────────
        problemas_lista = [f"{lc.get('data','')}: {lc['problemas']}"
                           for lc in lancamentos if lc.get("problemas","").strip()]
        if problemas_lista:
            _par("4. PROBLEMAS / OCORRÊNCIAS", bold=True, size=11)
            for prob in problemas_lista:
                _par(f"⚠ {prob}", size=10)
            doc.add_paragraph()
            secao_obs = 5
        else:
            secao_obs = 4

        # ── OBSERVAÇÕES ───────────────────────────────────────────────────────
        obs_lista = [f"{lc.get('data','')}: {lc['observacao']}"
                     for lc in lancamentos if lc.get("observacao","").strip()]
        if obs_geral:
            obs_lista.insert(0, obs_geral)
        if obs_lista:
            _par(f"{secao_obs}. OBSERVAÇÕES GERAIS", bold=True, size=11)
            for obs in obs_lista:
                _par(obs, size=10)
            doc.add_paragraph()
            secao_fotos = secao_obs + 1
        else:
            secao_fotos = secao_obs

        # ── REGISTRO FOTOGRÁFICO — organizado por categoria ───────────────────
        # Classifica fotos por categoria baseado no prefixo do nome
        _fotos_antes  = [f for f in (fotos or []) if "antes"  in f.name.lower()]
        _fotos_depois = [f for f in (fotos or []) if "depois" in f.name.lower()]
        _fotos_cmaq   = [f for f in (fotos or []) if "cmaq"   in f.name.lower()
                         or "maquin" in f.name.lower() or "casa" in f.name.lower()]
        _fotos_outras = [f for f in (fotos or [])
                         if f not in _fotos_antes + _fotos_depois + _fotos_cmaq]

        _tem_fotos = any([_fotos_antes, _fotos_depois, _fotos_cmaq, _fotos_outras])

        def _inserir_bloco_fotos(titulo_cat, lista_fotos, icone="📷"):
            if not lista_fotos:
                return
            _p_cat = doc.add_paragraph()
            _r_cat = _p_cat.add_run(f"{icone} {titulo_cat} ({len(lista_fotos)} foto(s))")
            _r_cat.bold = True
            _r_cat.font.size = Pt(10)
            _r_cat.font.color.rgb = _RGB(0x0d, 0x3d, 0x75)
            for foto_path in lista_fotos:
                try:
                    p_foto = doc.add_paragraph()
                    p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _ok_foto = inserir_foto_docx_exif(p_foto, foto_path, width_inches=5.5)
                    if not _ok_foto:
                        _par(f"[Foto não inserida: {foto_path.name}]", size=9, italic=True)
                except Exception:
                    _par(f"[Foto não inserida: {foto_path.name}]", size=9, italic=True)
            doc.add_paragraph()

        if _tem_fotos:
            _par(f"{secao_fotos}. REGISTRO FOTOGRÁFICO", bold=True, size=11)
            _inserir_bloco_fotos("Antes do tratamento",  _fotos_antes,  "🔵")
            _inserir_bloco_fotos("Após o tratamento",    _fotos_depois, "🟢")
            _inserir_bloco_fotos("Casa de máquinas",     _fotos_cmaq,   "🔧")
            _inserir_bloco_fotos("Outras",               _fotos_outras, "📷")
            secao_parecer = secao_fotos + 1
        else:
            secao_parecer = secao_fotos

        # ── PARECER TÉCNICO ───────────────────────────────────────────────────
        _pareceres = [lc.get("parecer","") for lc in lancamentos if lc.get("parecer","").strip()]
        if _pareceres:
            _par(f"{secao_parecer}. PARECER TÉCNICO-OPERACIONAL", bold=True, size=11)
            # Parecer da última visita
            _parecer_final = _pareceres[-1]
            _cor_parecer = "#1a7a1a" if "Satisfatório" in _parecer_final else (
                "#b86800" if "Aceitável" in _parecer_final else "#aa0000")
            doc.add_paragraph()
            _p_parecer = doc.add_paragraph()
            _r_parecer = _p_parecer.add_run(f"Parecer da última visita: {_parecer_final}")
            _r_parecer.font.size = Pt(11)
            _r_parecer.bold = True
            doc.add_paragraph()
            secao_parecer += 1

        # ── TEXTO RT (apenas no relatório sem RT — Bem Star) ──────────────────
        if not incluir_rt:
            doc.add_page_break()
            _par("SOBRE RESPONSABILIDADE TÉCNICA (RT)", bold=True, size=11)
            doc.add_paragraph()
            for linha in TEXTO_RT_SEM_RT.strip().split("\n\n"):
                if linha.startswith("SOBRE"):
                    continue
                _par(linha.strip(), size=10)
                doc.add_paragraph()
            # Logo Bem Star
            _logo_bs_path = None
            for _lp in LOGO_BEM_STAR_CANDIDATOS:
                if _lp.exists():
                    _logo_bs_path = _lp
                    break
            if _logo_bs_path:
                try:
                    p_logo = doc.add_paragraph()
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_logo.add_run().add_picture(str(_logo_bs_path), width=Inches(3.5))
                except Exception:
                    pass

        # ── ASSINATURA ────────────────────────────────────────────────────────
        _par(f"Uberlândia/MG, {hoje_br()}.", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
        t_ass = doc.add_table(rows=1, cols=2)
        t_ass.autofit = True

        if incluir_rt:
            ass_rt = (
                f"___________________________\n"
                f"{RESPONSAVEL_TÉCNICO}\n"
                f"{CRQ}\n"
                f"Aqua Gestão – Controle Técnico de Piscinas"
            )
        else:
            ass_rt = (
                f"___________________________\n"
                f"Bem Star Piscinas\n"
                f"CNPJ: {CNPJ_BEM_STAR}"
            )

        ass_resp = (
            f"___________________________\n"
            f"{responsavel or 'Responsável local'}\n"
            f"{nome_local}"
        )

        for cell_a, texto_a in [(t_ass.cell(0,0), ass_rt), (t_ass.cell(0,1), ass_resp)]:
            cell_a.paragraphs[0].clear()
            cell_a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_a.paragraphs[0].add_run(texto_a).font.size = Pt(9)

        doc.save(str(output_path))
        return True, ""
    except Exception as e:
        return False, str(e)


def _mes_ano_preview_relatorio(mes: str = "", ano: str = "") -> str:
    """Retorna AAAA-MM para buscar fotos do relatório no Drive/local.

    Aceita mês em formatos como "04", "4", "Abril" ou "04 - Abril".
    Se vier vazio ou inválido, usa o mês/ano atual.
    """
    try:
        meses_nome = {
            "janeiro": 1, "fevereiro": 2, "marco": 3, "março": 3,
            "abril": 4, "maio": 5, "junho": 6, "julho": 7,
            "agosto": 8, "setembro": 9, "outubro": 10,
            "novembro": 11, "dezembro": 12,
        }

        mes_txt = str(mes or "").strip().lower()
        ano_txt = str(ano or "").strip()

        mes_num = None
        m = re.search(r"\d{1,2}", mes_txt)
        if m:
            mes_num = int(m.group(0))
        else:
            for nome_mes, numero in meses_nome.items():
                if nome_mes in mes_txt:
                    mes_num = numero
                    break

        ano_num = None
        m_ano = re.search(r"\d{4}", ano_txt)
        if m_ano:
            ano_num = int(m_ano.group(0))

        hoje = datetime.now()
        if not mes_num or mes_num < 1 or mes_num > 12:
            mes_num = hoje.month
        if not ano_num:
            ano_num = hoje.year

        return f"{ano_num:04d}-{mes_num:02d}"
    except Exception:
        return datetime.now().strftime("%Y-%m")


def _salvar_uploads_relatorio_preview(pasta_preview: Path):
    """Salva fotos anexadas apenas na pasta de prévia.

    Diferente de salvar_uploads_relatorio(), esta função não envia os arquivos
    para o Google Drive, evitando duplicidade quando o usuário apenas clica em
    prévia antes de gerar o relatório definitivo.
    """
    caminhos = []
    arquivos = st.session_state.get("rel_fotos_upload") or []
    pasta_preview.mkdir(parents=True, exist_ok=True)
    for idx, arquivo in enumerate(arquivos, start=1):
        try:
            nome_original = getattr(arquivo, "name", f"foto_{idx}.jpg")
            nome = limpar_nome_arquivo(
                f"foto_previa_relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{idx}_{nome_original}"
            )
            destino = pasta_preview / nome
            foto_bytes = arquivo.getbuffer()
            with open(destino, "wb") as f:
                f.write(foto_bytes)
            caminhos.append(destino)
        except Exception:
            continue
    return caminhos

def _resolver_fotos_relatorio_rt(pasta_condominio: Path, nome_condominio: str, mes: str = "", ano: str = "", preview: bool = False) -> tuple[list[Path], str]:
    """Resolve as fotos do relatório RT usando a mesma lógica para prévia e geração final."""
    if preview:
        pasta_preview = pasta_condominio / "_previa_exata_relatorio" / "fotos_upload"
        fotos_upload = _salvar_uploads_relatorio_preview(pasta_preview)
        origem_upload = "anexos atuais do formulário"
    else:
        fotos_upload = salvar_uploads_relatorio(pasta_condominio)
        origem_upload = "anexos atuais do formulário"

    if fotos_upload:
        return fotos_upload, origem_upload

    mes_ano_rel = _mes_ano_preview_relatorio(mes, ano)
    fotos_drive = buscar_fotos_drive_para_relatorio(nome_condominio, mes_ano_rel)
    if fotos_drive:
        return fotos_drive, f"Google Drive ({mes_ano_rel})"

    pasta_fotos_campo = pasta_condominio / "fotos_campo"
    if pasta_fotos_campo.exists():
        fotos_campo = sorted(
            [f for f in pasta_fotos_campo.glob("*") if f.suffix.lower() in (".jpg", ".jpeg", ".png", ".webp")]
        )
        if fotos_campo:
            return fotos_campo, "pasta local fotos_campo"

    return [], "nenhuma foto encontrada"


def _carregar_clientes_bem_star_relatorio() -> dict:
    clientes = {c.get("nome", ""): c for c in (sheets_listar_clientes_completo() or []) if c.get("nome")}
    try:
        caminho_json = GENERATED_DIR / "_clientes_sem_rt.json"
        if caminho_json.exists() and "carregar_clientes_sem_rt" in globals():
            for c in (carregar_clientes_sem_rt() or []):
                nome = c.get("nome")
                if nome and nome not in clientes:
                    clientes[nome] = c
    except Exception:
        pass
    return clientes


def _coletar_contexto_relatorio_bem_star() -> dict:
    csr_sel = str(st.session_state.get("csr_sel_relatorio", "") or "").strip()
    csr_mes = str(st.session_state.get("csr_mes_rel", "") or "").strip()
    csr_ano = str(st.session_state.get("csr_ano_rel", "") or str(datetime.now().year)).strip()
    csr_operador_nome = str(st.session_state.get("csr_operador_rel", "") or "").strip()
    csr_obs_geral = str(st.session_state.get("csr_obs_rel", "") or "").strip()

    erros = []
    if not csr_sel:
        erros.append("Selecione o cliente Bem Star.")
    if not csr_mes:
        erros.append("Informe o mês do relatório Bem Star.")
    if not csr_ano:
        erros.append("Informe o ano do relatório Bem Star.")
    if erros:
        return {"ok": False, "erros": erros, "mensagem": " | ".join(erros)}

    clientes = _carregar_clientes_bem_star_relatorio()
    csr_dados_sel = clientes.get(csr_sel, {})

    pasta_csr = GENERATED_DIR / slugify_nome(csr_sel)
    pasta_csr.mkdir(parents=True, exist_ok=True)
    dados_rel_json = carregar_dados_condominio(pasta_csr) if pasta_csr.exists() else {}
    lancamentos_local = (dados_rel_json or {}).get("lancamentos_campo", [])
    lancamentos_sheets = sheets_listar_lancamentos(csr_sel) if csr_sel else []

    vistos = set()
    lancamentos_todos = []
    for lc in (lancamentos_local or []) + (lancamentos_sheets or []):
        chave = f"{lc.get('data','')}-{lc.get('operador','')}-{lc.get('ph','') or ((lc.get('piscinas') or [{}])[0].get('ph','') if lc.get('piscinas') else '')}"
        if chave not in vistos:
            vistos.add(chave)
            lancamentos_todos.append(lc)

    lancamentos_csr = _filtrar_lancamentos_preview_por_mes(lancamentos_todos, csr_mes, csr_ano)
    if not lancamentos_csr:
        msg = "Nenhum lançamento encontrado para o cliente/período selecionado."
        return {"ok": False, "erros": [msg], "mensagem": msg}

    lanc_para_relatorio = []
    vistos_rel = set()
    for lc in lancamentos_csr:
        chave = f"{lc.get('data','')}-{lc.get('operador','')}-{lc.get('ph','') or ((lc.get('piscinas') or [{}])[0].get('ph','') if lc.get('piscinas') else '')}"
        if chave in vistos_rel:
            continue
        vistos_rel.add(chave)
        piscinas = lc.get("piscinas", [])
        dados = piscinas[0] if piscinas else lc
        lanc_para_relatorio.append({
            "data": lc.get("data", ""),
            "ph": dados.get("ph", lc.get("ph", "")),
            "cloro_livre": dados.get("cloro_livre", lc.get("cloro_livre", "")),
            "cloro_total": dados.get("cloro_total", lc.get("cloro_total", "")),
            "alcalinidade": dados.get("alcalinidade", lc.get("alcalinidade", "")),
            "dureza": dados.get("dureza", lc.get("dureza", "")),
            "cianurico": dados.get("cianurico", lc.get("cianurico", "")),
            "operador": lc.get("operador", csr_operador_nome),
            "observacao": lc.get("observacao", ""),
            "problemas": lc.get("problemas", ""),
            "dosagens": dados.get("dosagens", lc.get("dosagens", [])),
        })

    fotos_paths = _coletar_fotos_bem_star_preview(csr_sel, lancamentos_csr)
    return {
        "ok": True,
        "cliente": csr_sel,
        "mes": csr_mes,
        "ano": csr_ano,
        "operador": csr_operador_nome,
        "obs_geral": csr_obs_geral,
        "dados_cliente": csr_dados_sel,
        "pasta": pasta_csr,
        "lancamentos": lanc_para_relatorio,
        "fotos": fotos_paths,
        "origem_fotos": "fotos_campo das visitas" if fotos_paths else "nenhuma foto encontrada",
    }


def _renderizar_relatorio_rt(preview: bool = False) -> dict:
    dados_relatorio = montar_dados_relatorio()
    erros = validar_relatorio_mensal(dados_relatorio)
    if erros:
        return {
            "ok": False,
            "empresa": "Aqua Gestão",
            "preview": preview,
            "erros": erros,
            "mensagem": " | ".join(erros),
        }

    nome_condominio = dados_relatorio["nome_condominio"]
    pasta_condominio = GENERATED_DIR / slugify_nome(nome_condominio)
    pasta_condominio.mkdir(parents=True, exist_ok=True)

    if not preview:
        if st.session_state.get("rel_salvar_alteracoes_cadastro"):
            salvar_relatorio_no_cadastro_principal()
            salvar_dados_condominio(pasta_condominio, salvar_snapshot_formulario())
        else:
            salvar_dados_condominio(pasta_condominio, obter_snapshot_relatorio_independente())

    fotos_salvas, origem_fotos = _resolver_fotos_relatorio_rt(
        pasta_condominio,
        nome_condominio,
        dados_relatorio.get("mes_referencia", ""),
        dados_relatorio.get("ano_referencia", ""),
        preview=preview,
    )

    pasta_saida = pasta_condominio / "_previa_exata_relatorio" if preview else pasta_condominio
    pasta_saida.mkdir(parents=True, exist_ok=True)
    data_nome = datetime.now().strftime("%Y%m%d_%H%M%S" if preview else "%Y%m%d")
    sufixo = "PREVIA_EXATA_RELATORIO_RT" if preview else "RELATORIO_RT"
    base_nome = limpar_nome_arquivo(f"{data_nome}_{nome_condominio}_{sufixo}")
    relatorio_docx = pasta_saida / f"{base_nome}.docx"
    relatorio_pdf = pasta_saida / f"{base_nome}.pdf"

    preencher_relatorio_mensal_docx(TEMPLATE_RELATORIO, relatorio_docx, dados_relatorio, fotos=fotos_salvas)

    # PDF oficial premium Aqua Gestão: gerado diretamente por ReportLab para manter
    # cores, tabelas, logo e fotos mesmo quando Word/LibreOffice não estiver disponível.
    ok_pdf, erro_pdf = gerar_pdf_relatorio_rt_premium_reportlab(dados_relatorio, fotos_salvas, relatorio_pdf)

    # Fallback técnico: se o PDF premium falhar, tenta converter o DOCX.
    if not ok_pdf:
        ok_pdf, erro_pdf = converter_docx_para_pdf(relatorio_docx, relatorio_pdf)

    if not preview:
        registrar_documento_manifest(
            pasta_condominio=pasta_condominio,
            nome_condominio=nome_condominio,
            tipo="Relatório",
            arquivo_docx=relatorio_docx,
            arquivo_pdf=relatorio_pdf,
            pdf_gerado=ok_pdf,
            erro_pdf=erro_pdf,
            dados_utilizados={
                "TIPO_ATENDIMENTO": dados_relatorio.get("tipo_atendimento"),
                "REPRESENTANTE": dados_relatorio.get("representante"),
                "CPF_CNPJ_REPRESENTANTE": dados_relatorio.get("cpf_cnpj_representante"),
                "ART_STATUS": dados_relatorio.get("art_status"),
                "ART_TEXTO": obter_status_art_texto(dados_relatorio),
            },
            extras={"fotos": [p.name for p in fotos_salvas]},
        )
        st.session_state.ultimos_docs_gerados = st.session_state.get("ultimos_docs_gerados") or {}
        st.session_state.ultimos_docs_gerados.update({
            "relatorio_docx": str(relatorio_docx) if relatorio_docx.exists() else None,
            "relatorio_pdf": str(relatorio_pdf) if ok_pdf and relatorio_pdf.exists() else None,
        })

    return {
        "ok": True,
        "empresa": "Aqua Gestão",
        "preview": preview,
        "mensagem": (
            f"Prévia exata Aqua Gestão atualizada com {len(fotos_salvas)} foto(s), usando o mesmo gerador DOCX/PDF do relatório final."
            if preview else
            f"Relatório mensal registrado com sucesso para {nome_condominio}."
        ),
        "docx": relatorio_docx,
        "pdf": relatorio_pdf,
        "pdf_ok": ok_pdf,
        "erro_pdf": erro_pdf,
        "fotos": fotos_salvas,
        "origem_fotos": origem_fotos,
        "dados": dados_relatorio,
        "pasta": pasta_condominio,
    }


def _renderizar_relatorio_bem_star(preview: bool = False) -> dict:
    ctx = _coletar_contexto_relatorio_bem_star()
    if not ctx.get("ok"):
        return {
            "ok": False,
            "empresa": "Bem Star Piscinas",
            "preview": preview,
            "erros": ctx.get("erros", []),
            "mensagem": ctx.get("mensagem", "Não foi possível montar o relatório Bem Star."),
        }

    pasta_saida = ctx["pasta"] / "_previa_exata_relatorio" if preview else ctx["pasta"]
    pasta_saida.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    sufixo = "PREVIA_EXATA_RELATORIO_BS" if preview else "RELATORIO_BS"
    base_nome = limpar_nome_arquivo(f"{ts}_{ctx['cliente']}_{sufixo}")
    docx_path = pasta_saida / f"{base_nome}.docx"
    pdf_path = pasta_saida / f"{base_nome}.pdf"

    ok_docx, erro_docx = gerar_relatorio_visita_docx(
        output_path=docx_path,
        nome_local=ctx["dados_cliente"].get("nome", ctx["cliente"]),
        cnpj=ctx["dados_cliente"].get("cnpj", ""),
        endereco=ctx["dados_cliente"].get("endereco", ""),
        responsavel=ctx["dados_cliente"].get("contato", ""),
        operador=ctx["operador"],
        mes=ctx["mes"],
        ano=ctx["ano"],
        lancamentos=ctx["lancamentos"],
        obs_geral=ctx["obs_geral"],
        incluir_rt=False,
        fotos=ctx["fotos"],
    )
    if not ok_docx:
        msg = f"Erro ao gerar DOCX do relatório Bem Star: {erro_docx}"
        return {
            "ok": False,
            "empresa": "Bem Star Piscinas",
            "preview": preview,
            "erros": [erro_docx],
            "mensagem": msg,
            "fotos": ctx.get("fotos", []),
            "origem_fotos": ctx.get("origem_fotos", ""),
        }

    ok_pdf, erro_pdf = converter_docx_para_pdf(docx_path, pdf_path)
    if not preview:
        registrar_documento_manifest(ctx["pasta"], ctx["cliente"], "Relatório", docx_path, pdf_path, ok_pdf, erro_pdf)

    return {
        "ok": True,
        "empresa": "Bem Star Piscinas",
        "preview": preview,
        "mensagem": (
            f"Prévia exata Bem Star atualizada com {len(ctx['lancamentos'])} lançamento(s) e {len(ctx['fotos'])} foto(s)."
            if preview else
            f"Relatório Bem Star gerado! {len(ctx['fotos'])} foto(s) incluída(s)."
        ),
        "docx": docx_path,
        "pdf": pdf_path,
        "pdf_ok": ok_pdf,
        "erro_pdf": erro_pdf,
        "fotos": ctx["fotos"],
        "origem_fotos": ctx["origem_fotos"],
        "dados": ctx,
        "pasta": ctx["pasta"],
    }


def renderizar_relatorio_oficial(empresa: str = "Aqua Gestão", preview: bool = False) -> dict:
    empresa = str(empresa or "Aqua Gestão").strip()
    if empresa == "Bem Star Piscinas":
        return _renderizar_relatorio_bem_star(preview=preview)
    return _renderizar_relatorio_rt(preview=preview)


def gerar_relatorio_mensal() -> tuple[bool, str]:
    resultado = renderizar_relatorio_oficial("Aqua Gestão", preview=False)
    if not resultado.get("ok"):
        return False, resultado.get("mensagem", "Não foi possível gerar o relatório mensal.")

    dados_relatorio = resultado["dados"]
    relatorio_docx = Path(resultado["docx"])
    relatorio_pdf = Path(resultado["pdf"])
    ok_pdf = bool(resultado.get("pdf_ok"))
    erro_pdf = resultado.get("erro_pdf")
    pasta_condominio = Path(resultado["pasta"])

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    if dados_relatorio["avaliacao_automatica"]["detalhes"]:
        st.warning("Diagnóstico automático marcou NÃO CONFORME com base nos parâmetros fora de faixa.")
        for item in dados_relatorio["avaliacao_automatica"]["detalhes"]:
            st.write(f"- {item}")
    if dados_relatorio["avaliacao_automatica"].get("cloraminas_altas"):
        for idx, valor in dados_relatorio["avaliacao_automatica"]["cloraminas_altas"]:
            st.write(f"- Linha {idx}: cloro combinado estimado em {valor} mg/L.")
    c1, c2, c3 = st.columns(3)
    with c1:
        if relatorio_docx.exists():
            with open(relatorio_docx, "rb") as f:
                st.download_button("Baixar DOCX do relatório", data=f, file_name=relatorio_docx.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    with c2:
        if ok_pdf and relatorio_pdf.exists():
            with open(relatorio_pdf, "rb") as f:
                st.download_button("Baixar PDF do relatório", data=f, file_name=relatorio_pdf.name, mime="application/pdf", use_container_width=True)
        else:
            st.warning(f"PDF do relatório não gerado. Erro: {erro_pdf}")
    with c3:
        if st.button("Abrir pasta do condomínio", key="abrir_pasta_relatorio", use_container_width=True):
            abrir_pasta_windows(pasta_condominio)
    st.markdown("</div>", unsafe_allow_html=True)
    return True, resultado["mensagem"]


def gerar_previa_exata_relatorio(empresa: str = "Aqua Gestão") -> dict:
    return renderizar_relatorio_oficial(empresa, preview=True)


def exibir_pdf_previa_exata(pdf_path: Path, height: int = 1200):
    """Exibe a prévia do PDF dentro do Streamlit de forma confiável.

    O visual anterior usava iframe/base64. Em algumas versões do Chrome/Streamlit
    isso aparece como uma folha branca, mesmo quando o PDF foi gerado corretamente.

    Esta versão renderiza o PDF como imagem usando PyMuPDF quando disponível.
    """
    pdf_path = Path(pdf_path) if pdf_path else None
    if not pdf_path or not pdf_path.exists():
        st.warning("O PDF da prévia ainda não está disponível.")
        return

    # Botão sempre disponível, mesmo se a visualização interna falhar.
    try:
        with open(pdf_path, "rb") as _f_prev_pdf:
            st.download_button(
                "⬇️ Baixar PDF da prévia para conferir",
                data=_f_prev_pdf,
                file_name=pdf_path.name,
                mime="application/pdf",
                use_container_width=True,
                key=f"btn_download_previa_pdf_inline_{chave_segura(str(pdf_path))}",
            )
    except Exception:
        pass

    # Principal: renderiza páginas do PDF como imagens PNG.
    try:
        import fitz  # PyMuPDF
        import io as _io

        doc_pdf = fitz.open(str(pdf_path))
        total_paginas = len(doc_pdf)
        if total_paginas == 0:
            st.warning("O PDF foi gerado, mas não possui páginas.")
            doc_pdf.close()
            return

        max_paginas = min(total_paginas, 6)
        st.caption(f"Prévia renderizada em imagem: {max_paginas} de {total_paginas} página(s).")

        for idx in range(max_paginas):
            pagina = doc_pdf.load_page(idx)
            matriz = fitz.Matrix(1.45, 1.45)
            pix = pagina.get_pixmap(matrix=matriz, alpha=False)
            img_bytes = pix.tobytes("png")
            st.image(_io.BytesIO(img_bytes), use_container_width=True)
            if idx < max_paginas - 1:
                st.markdown("---")

        if total_paginas > max_paginas:
            st.info(
                f"Prévia limitada às {max_paginas} primeiras páginas para manter o sistema leve. "
                "Baixe o PDF para ver o relatório completo."
            )
        doc_pdf.close()
        return

    except Exception as e:
        st.warning(
            "A prévia visual em imagem não pôde ser carregada. "
            "O PDF foi gerado; use o botão de download acima para conferir. "
            "Para exibir dentro do sistema, inclua `PyMuPDF` no requirements.txt."
        )
        st.caption(f"Detalhe técnico: {type(e).__name__}: {e}")
    # Fallback seguro: evitar components.html no Streamlit 1.56.
    st.info(
        "A visualização embutida do PDF foi desativada nesta versão para evitar rerun infinito. "
        "Use o botão de download acima para abrir o arquivo localmente."
    )


# =========================================
# PROCESSAMENTO DE VALIDAÇÃO
# =========================================

def validar_para_geracao(dados_base: dict, email_cliente: str) -> list[str]:
    faltando = validar_campos_obrigatorios(dados_base)
    if faltando:
        return [f"Preencha o campo obrigatório: {item}" for item in faltando]

    erros_formato = validar_campos_formato(dados_base, email_cliente)
    return erros_formato


# =========================================
# SESSÃO
# =========================================

def inicializar_campos():
    defaults = {
        "nome_condominio": "",
        "cnpj_condominio": "",
        "endereco_condominio": "",
        "nome_sindico": "",
        "cpf_sindico": "",
        "valor_mensal": "",
        "valor_aditivo": "",
        "data_inicio": "",
        "data_fim": "",
        "data_assinatura": hoje_br(),
        "whatsapp_cliente": "",
        "email_cliente": "",
        "observacoes_internas": "",
        "filtro_historico": "",
        "ultima_pasta_gerada": None,
        "confirm_delete_folder": "",
        "confirm_delete_file": "",
        "origem_dados_carregados": "",
        "alerta_vencimento_dias": 30,
        "painel_acao_msg": "",
        "busca_rapida": "",
        "filtro_status_central": "Todos",
        "rel_mes_referencia": datetime.now().strftime("%m"),
        "rel_tipo_atendimento": "Contrato ativo",
        "rel_nome_condominio": "",
        "rel_cnpj_condominio": "",
        "rel_endereco_condominio": "",
        "rel_representante": "",
        "rel_cpf_cnpj_representante": "",
        "rel_salvar_alteracoes_cadastro": False,
        "rel_ano_referencia": str(datetime.now().year),
        "rel_art_status": "Emitida",
        "rel_art_status_widget": "Emitida",
        "rel_art_numero": "",
        "rel_art_inicio": "",
        "rel_art_fim": "",
        "rel_data_emissao": hoje_br(),
        "rel_epi_luvas_ca": "",
        "rel_epi_oculos_ca": "",
        "rel_epi_respirador_ca": "",
        "rel_epi_botas_ca": "",
        "rel_status_agua": "CONFORME",
        "rel_status_agua_select": "CONFORME",
        "rel_diagnostico": "",
        "rel_nbr_11238": "",
        "rel_nr_26": "",
        "rel_nr_06": "",
        "rel_analises_total": ANALISES_PADRAO,
    }
    garantir_campos_analises(st.session_state.get("rel_analises_total", ANALISES_PADRAO) if hasattr(st, "session_state") else ANALISES_PADRAO)
    for i in range(ANALISES_PADRAO):
        defaults[f"rel_analise_data_{i}"] = ""
        defaults[f"rel_analise_ph_{i}"] = ""
        defaults[f"rel_analise_cl_{i}"] = ""
        defaults[f"rel_analise_ct_{i}"] = ""
        defaults[f"rel_analise_alc_{i}"] = ""
        defaults[f"rel_analise_dc_{i}"] = ""
        defaults[f"rel_analise_cya_{i}"] = ""
        defaults[f"rel_analise_operador_{i}"] = ""
    for i in range(7):
        defaults[f"rel_dos_produto_{i}"] = ""
        defaults[f"rel_dos_lote_{i}"] = ""
        defaults[f"rel_dos_qtd_{i}"] = ""
        defaults[f"rel_dos_un_{i}"] = ""
        defaults[f"rel_dos_finalidade_{i}"] = ""
    for i in range(5):
        defaults[f"rel_rec_texto_{i}"] = ""
        defaults[f"rel_rec_prazo_{i}"] = ""
        defaults[f"rel_rec_resp_{i}"] = ""
        defaults[f"rel_obs_{i}"] = ""
    for chave, valor in defaults.items():
        if chave not in st.session_state:
            st.session_state[chave] = valor


def limpar_formulario():
    st.session_state.nome_condominio = ""
    st.session_state.cnpj_condominio = ""
    st.session_state.endereco_condominio = ""
    st.session_state.nome_sindico = ""
    st.session_state.cpf_sindico = ""
    st.session_state.valor_mensal = ""
    st.session_state.valor_aditivo = ""
    st.session_state.data_inicio = ""
    st.session_state.data_fim = ""
    st.session_state.data_assinatura = hoje_br()
    st.session_state.whatsapp_cliente = ""
    st.session_state.email_cliente = ""
    st.session_state.observacoes_internas = ""
    st.session_state.origem_dados_carregados = ""
    st.session_state.rel_tipo_atendimento = "Contrato ativo"
    st.session_state.rel_nome_condominio = ""
    st.session_state.rel_cnpj_condominio = ""
    st.session_state.rel_endereco_condominio = ""
    st.session_state.rel_representante = ""
    st.session_state.rel_cpf_cnpj_representante = ""
    st.session_state.rel_salvar_alteracoes_cadastro = False
    st.session_state.rel_mes_referencia = datetime.now().strftime("%m")
    st.session_state.rel_ano_referencia = str(datetime.now().year)
    st.session_state.rel_art_status = "Emitida"
    st.session_state.rel_art_status_widget = "Emitida"
    st.session_state.rel_art_numero = ""
    st.session_state.rel_art_inicio = ""
    st.session_state.rel_art_fim = ""
    st.session_state.rel_data_emissao = hoje_br()
    st.session_state.rel_epi_luvas_ca = ""
    st.session_state.rel_epi_oculos_ca = ""
    st.session_state.rel_epi_respirador_ca = ""
    st.session_state.rel_epi_botas_ca = ""
    st.session_state.rel_status_agua = "CONFORME"
    st.session_state.rel_status_agua_select = "CONFORME"
    st.session_state.rel_diagnostico = ""
    st.session_state.rel_nbr_11238 = ""
    st.session_state.rel_nr_26 = ""
    st.session_state.rel_nr_06 = ""
    st.session_state.rel_epi_luvas_status = "Conforme"
    st.session_state.rel_epi_oculos_status = "Conforme"
    st.session_state.rel_epi_respirador_status = "Conforme"
    st.session_state.rel_epi_botas_status = "Conforme"
    st.session_state.rel_epi_luvas_ca = ""
    st.session_state.rel_epi_oculos_ca = ""
    st.session_state.rel_epi_respirador_ca = ""
    st.session_state.rel_epi_botas_ca = ""
    total_analises_atual = int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO)
    st.session_state.rel_analises_total = ANALISES_PADRAO
    for i in range(max(total_analises_atual, ANALISES_PADRAO)):
        st.session_state[f"rel_analise_data_{i}"] = ""
        st.session_state[f"rel_analise_ph_{i}"] = ""
        st.session_state[f"rel_analise_cl_{i}"] = ""
        st.session_state[f"rel_analise_ct_{i}"] = ""
        st.session_state[f"rel_analise_alc_{i}"] = ""
        st.session_state[f"rel_analise_dc_{i}"] = ""
        st.session_state[f"rel_analise_cya_{i}"] = ""
        st.session_state[f"rel_analise_operador_{i}"] = ""
    for i in range(7):
        st.session_state[f"rel_dos_produto_{i}"] = ""
        st.session_state[f"rel_dos_lote_{i}"] = ""
        st.session_state[f"rel_dos_qtd_{i}"] = ""
        st.session_state[f"rel_dos_un_{i}"] = ""
        st.session_state[f"rel_dos_finalidade_{i}"] = ""
    for i in range(5):
        st.session_state[f"rel_rec_texto_{i}"] = ""
        st.session_state[f"rel_rec_prazo_{i}"] = ""
        st.session_state[f"rel_rec_resp_{i}"] = ""
        st.session_state[f"rel_obs_{i}"] = ""


RASCUNHO_JSON_NAME = "rascunho_relatorio.json"


def salvar_rascunho_relatorio(pasta_condominio: Path):
    """Salva o estado atual do formulário de relatório como rascunho no JSON."""
    qtd = int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO)
    rascunho = {
        "rel_nome_condominio": (st.session_state.get("rel_nome_condominio") or "").strip(),
        "rel_cnpj_condominio": (st.session_state.get("rel_cnpj_condominio") or "").strip(),
        "rel_endereco_condominio": (st.session_state.get("rel_endereco_condominio") or "").strip(),
        "rel_representante": (st.session_state.get("rel_representante") or "").strip(),
        "rel_cpf_cnpj_representante": (st.session_state.get("rel_cpf_cnpj_representante") or "").strip(),
        "rel_tipo_atendimento": (st.session_state.get("rel_tipo_atendimento") or "Contrato ativo"),
        "rel_mes_referencia": (st.session_state.get("rel_mes_referencia") or ""),
        "rel_ano_referencia": (st.session_state.get("rel_ano_referencia") or ""),
        "rel_data_emissao": (st.session_state.get("rel_data_emissao") or hoje_br()),
        "rel_art_status": (st.session_state.get("rel_art_status") or "Emitida"),
        "rel_art_numero": (st.session_state.get("rel_art_numero") or ""),
        "rel_art_inicio": (st.session_state.get("rel_art_inicio") or ""),
        "rel_art_fim": (st.session_state.get("rel_art_fim") or ""),
        "rel_status_agua": (st.session_state.get("rel_status_agua") or "CONFORME"),
        "rel_diagnostico": (st.session_state.get("rel_diagnostico") or ""),
        "rel_nbr_11238": (st.session_state.get("rel_nbr_11238") or ""),
        "rel_nr_26": (st.session_state.get("rel_nr_26") or ""),
        "rel_nr_06": (st.session_state.get("rel_nr_06") or ""),
        "rel_epi_luvas_status": (st.session_state.get("rel_epi_luvas_status") or "Conforme"),
        "rel_epi_luvas_ca": (st.session_state.get("rel_epi_luvas_ca") or ""),
        "rel_epi_oculos_status": (st.session_state.get("rel_epi_oculos_status") or "Conforme"),
        "rel_epi_oculos_ca": (st.session_state.get("rel_epi_oculos_ca") or ""),
        "rel_epi_respirador_status": (st.session_state.get("rel_epi_respirador_status") or "Conforme"),
        "rel_epi_respirador_ca": (st.session_state.get("rel_epi_respirador_ca") or ""),
        "rel_epi_botas_status": (st.session_state.get("rel_epi_botas_status") or "Conforme"),
        "rel_epi_botas_ca": (st.session_state.get("rel_epi_botas_ca") or ""),
        "rel_analises_total": qtd,
        "analises": [],
        "dosagens": obter_dosagens_ultimas_relatorio(),
        "observacoes": [(st.session_state.get(f"rel_obs_{i}") or "") for i in range(5)],
        "recomendacoes": [
            {
                "texto": (st.session_state.get(f"rel_rec_texto_{i}") or ""),
                "prazo": (st.session_state.get(f"rel_rec_prazo_{i}") or ""),
                "responsavel": (st.session_state.get(f"rel_rec_resp_{i}") or ""),
            }
            for i in range(5)
        ],
        "salvo_em": _agora_brasilia(),
    }
    for i in range(qtd):
        rascunho["analises"].append({
            "data": (st.session_state.get(f"rel_analise_data_{i}") or ""),
            "ph": (st.session_state.get(f"rel_analise_ph_{i}") or ""),
            "cl": (st.session_state.get(f"rel_analise_cl_{i}") or ""),
            "ct": (st.session_state.get(f"rel_analise_ct_{i}") or ""),
            "alc": (st.session_state.get(f"rel_analise_alc_{i}") or ""),
            "dc": (st.session_state.get(f"rel_analise_dc_{i}") or ""),
            "cya": (st.session_state.get(f"rel_analise_cya_{i}") or ""),
            "operador": (st.session_state.get(f"rel_analise_operador_{i}") or ""),
        })
    caminho = pasta_condominio / RASCUNHO_JSON_NAME
    caminho.write_text(json.dumps(rascunho, ensure_ascii=False, indent=2), encoding="utf-8")
    return rascunho


def carregar_rascunho_relatorio(pasta_condominio: Path) -> dict | None:
    caminho = pasta_condominio / RASCUNHO_JSON_NAME
    if not caminho.exists():
        return None
    try:
        return json.loads(caminho.read_text(encoding="utf-8"))
    except Exception:
        return None


def aplicar_rascunho_no_formulario(rascunho: dict):
    """Restaura todos os campos do relatório a partir do rascunho salvo."""
    campos_simples = [
        "rel_nome_condominio", "rel_cnpj_condominio", "rel_endereco_condominio",
        "rel_representante", "rel_cpf_cnpj_representante", "rel_tipo_atendimento",
        "rel_mes_referencia", "rel_ano_referencia", "rel_data_emissao",
        "rel_art_status", "rel_art_numero", "rel_art_inicio", "rel_art_fim",
        "rel_status_agua", "rel_diagnostico", "rel_nbr_11238", "rel_nr_26", "rel_nr_06",
        "rel_nbr11238_profundidade", "rel_nbr11238_retrolavagem", "rel_nbr11238_skimmers", "rel_nbr11238_circulacao", "rel_nbr11238_chuveiro",
        "rel_epi_luvas_status", "rel_epi_luvas_ca", "rel_epi_oculos_status", "rel_epi_oculos_ca",
        "rel_epi_respirador_status", "rel_epi_respirador_ca", "rel_epi_botas_status", "rel_epi_botas_ca",
    ]
    for c in campos_simples:
        if c in rascunho:
            st.session_state[c] = rascunho[c]

    analises = rascunho.get("analises", [])
    qtd = max(len(analises), ANALISES_PADRAO)
    garantir_campos_analises(qtd)
    st.session_state.rel_analises_total = qtd
    for i, a in enumerate(analises):
        st.session_state[f"rel_analise_data_{i}"] = a.get("data", "")
        st.session_state[f"rel_analise_ph_{i}"] = a.get("ph", "")
        st.session_state[f"rel_analise_cl_{i}"] = a.get("cl", "")
        st.session_state[f"rel_analise_ct_{i}"] = a.get("ct", "")
        st.session_state[f"rel_analise_alc_{i}"] = a.get("alc", "")
        st.session_state[f"rel_analise_dc_{i}"] = a.get("dc", "")
        st.session_state[f"rel_analise_cya_{i}"] = a.get("cya", "")
        st.session_state[f"rel_analise_operador_{i}"] = a.get("operador", "")

    dosagens = rascunho.get("dosagens", [])
    for i in range(7):
        d = dosagens[i] if i < len(dosagens) else {}
        st.session_state[f"rel_dos_produto_{i}"] = d.get("produto", "")
        st.session_state[f"rel_dos_lote_{i}"] = d.get("fabricante_lote", "")
        st.session_state[f"rel_dos_qtd_{i}"] = d.get("quantidade", "")
        st.session_state[f"rel_dos_un_{i}"] = d.get("unidade", "")
        st.session_state[f"rel_dos_finalidade_{i}"] = d.get("finalidade", "")

    for i, obs in enumerate(rascunho.get("observacoes", [])[:5]):
        st.session_state[f"rel_obs_{i}"] = obs

    for i, rec in enumerate(rascunho.get("recomendacoes", [])[:5]):
        st.session_state[f"rel_rec_texto_{i}"] = rec.get("texto", "")
        st.session_state[f"rel_rec_prazo_{i}"] = rec.get("prazo", "")
        st.session_state[f"rel_rec_resp_{i}"] = rec.get("responsavel", "")


def excluir_rascunho_relatorio(pasta_condominio: Path):
    caminho = pasta_condominio / RASCUNHO_JSON_NAME
    if caminho.exists():
        caminho.unlink()


inicializar_campos()
sincronizar_relatorio_com_cadastro()

# Auto-restaurar rascunho ao iniciar sessão
if not st.session_state.get("_rascunho_restaurado"):
    nome_atual = (st.session_state.get("rel_nome_condominio") or st.session_state.get("nome_condominio") or "").strip()
    if nome_atual:
        pasta_rasc = GENERATED_DIR / slugify_nome(nome_atual)
        rasc = carregar_rascunho_relatorio(pasta_rasc)
        if rasc:
            aplicar_rascunho_no_formulario(rasc)
    st.session_state["_rascunho_restaurado"] = True

# Quando nome do condomínio do relatório muda, tenta restaurar rascunho automaticamente
_nome_rel_check = (st.session_state.get("rel_nome_condominio") or "").strip()
_ultimo_nome_rasc = st.session_state.get("_ultimo_nome_rasc_check", "")
if _nome_rel_check and _nome_rel_check != _ultimo_nome_rasc:
    st.session_state["_ultimo_nome_rasc_check"] = _nome_rel_check
    _pasta_check = GENERATED_DIR / slugify_nome(_nome_rel_check)
    if _pasta_check.exists():
        _rasc_auto = carregar_rascunho_relatorio(_pasta_check)
        if _rasc_auto:
            aplicar_rascunho_no_formulario(_rasc_auto)

pasta_formulario_atual = obter_pasta_atual_do_formulario()
if pasta_formulario_atual and pasta_formulario_atual.exists():
    dados_auto = carregar_dados_condominio(pasta_formulario_atual)
    if dados_auto:
        dosagens_preenchidas = any((st.session_state.get(f"rel_dos_produto_{i}") or "").strip() or (st.session_state.get(f"rel_dos_lote_{i}") or "").strip() or (st.session_state.get(f"rel_dos_qtd_{i}") or "").strip() or (st.session_state.get(f"rel_dos_un_{i}") or "").strip() or (st.session_state.get(f"rel_dos_finalidade_{i}") or "").strip() for i in range(7))
        if not dosagens_preenchidas:
            aplicar_dosagens_ultimas_no_relatorio(dados_auto)


# _EMERGENCIA_RELATORIO_RT_FUNCOES_V1_
# =========================================
# EMERGÊNCIA — PROTEÇÃO DO RELATÓRIO RT
# =========================================
def _relatorio_rt_rascunho_path() -> Path:
    try:
        GENERATED_DIR.mkdir(exist_ok=True)
    except Exception:
        pass
    return GENERATED_DIR / "_rascunho_relatorio_rt_em_andamento.json"


def _relatorio_rt_tem_dados_em_tela() -> bool:
    chaves_prefixos = ("rel_analise_", "rel_dos_", "rel_obs_")
    chaves_diretas = (
        "rel_nome_condominio", "nome_condominio", "rel_mes_referencia", "rel_ano_referencia",
        "rel_diagnostico", "rel_observacoes_gerais", "csr_obs_rel",
        "rel_status_agua", "rel_verificacoes_semanais",
    )
    for k in chaves_diretas:
        if str(st.session_state.get(k, "") or "").strip():
            return True
    for k, v in st.session_state.items():
        if str(k).startswith(chaves_prefixos) and str(v or "").strip():
            return True
    return False


def _relatorio_rt_coletar_rascunho() -> dict:
    total = int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO)
    dados = {
        "salvo_em": _agora_brasilia() if "_agora_brasilia" in globals() else datetime.now().isoformat(),
        "empresa_ativa": "aqua_gestao",
        "rel_analises_total": max(total, ANALISES_PADRAO),
        "campos": {},
    }
    for k, v in st.session_state.items():
        if str(k).startswith(("rel_", "csr_")) or str(k) in ("nome_condominio", "cnpj_condominio", "endereco_condominio"):
            try:
                if isinstance(v, (str, int, float, bool)) or v is None:
                    dados["campos"][k] = v
                else:
                    dados["campos"][k] = str(v)
            except Exception:
                pass
    return dados


def _relatorio_rt_salvar_rascunho(motivo: str = "autosave") -> bool:
    try:
        if not _relatorio_rt_tem_dados_em_tela():
            return False
        dados = _relatorio_rt_coletar_rascunho()
        dados["motivo"] = motivo
        # 1. Salva localmente
        _relatorio_rt_rascunho_path().write_text(json.dumps(dados, ensure_ascii=False, indent=2), encoding="utf-8")
        st.session_state["_relatorio_rt_ultimo_autosave"] = dados.get("salvo_em", "")
        # 2. Backup no Sheets (persiste apos sleep do servidor)
        try:
            sh = conectar_sheets()
            if sh:
                try:
                    aba_rasc_rt = obter_aba_sheets("_Rascunhos_RT")
                except Exception:
                    aba_rasc_rt = sh.add_worksheet(title="_Rascunhos_RT", rows=100, cols=3)
                    aba_rasc_rt.update("A1:C1", [["Usuario", "Salvo em", "Dados JSON"]])
                payload = json.dumps(dados, ensure_ascii=False)
                if len(payload) > 45000:
                    payload = payload[:45000] + "..."
                todos_rt = aba_rasc_rt.get_all_values()
                linha_ex = None
                for i, row in enumerate(todos_rt[1:], start=2):
                    if row and str(row[0]).strip() == "thyago":
                        linha_ex = i
                        break
                nova = ["thyago", dados.get("salvo_em", ""), payload]
                if linha_ex:
                    aba_rasc_rt.update(f"A{linha_ex}:C{linha_ex}", [nova], value_input_option="RAW")
                else:
                    prox = max(len(todos_rt) + 1, 2)
                    aba_rasc_rt.update(f"A{prox}:C{prox}", [nova], value_input_option="RAW")
        except Exception:
            pass  # Sheets indisponivel — arquivo local e suficiente
        return True
    except Exception as e:
        st.session_state["_relatorio_rt_autosave_erro"] = f"{type(e).__name__}: {e}"
        return False


def _relatorio_rt_carregar_rascunho() -> dict:
    # 1. Tenta arquivo local
    try:
        path = _relatorio_rt_rascunho_path()
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        pass
    # 2. Fallback no Sheets (apos sleep do servidor)
    try:
        sh = conectar_sheets()
        if sh:
            try:
                aba_rasc_rt = obter_aba_sheets("_Rascunhos_RT")
                todos_rt = aba_rasc_rt.get_all_values()
                for row in todos_rt[1:]:
                    if row and str(row[0]).strip() == "thyago" and len(row) >= 3 and row[2].strip():
                        dados = json.loads(row[2])
                        # Reconstitui arquivo local
                        try:
                            _relatorio_rt_rascunho_path().write_text(
                                json.dumps(dados, ensure_ascii=False, indent=2), encoding="utf-8")
                        except Exception:
                            pass
                        return dados
            except Exception:
                pass
    except Exception:
        pass
    return {}


def _relatorio_rt_aplicar_rascunho(dados: dict) -> bool:
    try:
        campos = dados.get("campos", {}) if isinstance(dados, dict) else {}
        if not campos:
            return False
        total = int(dados.get("rel_analises_total", campos.get("rel_analises_total", ANALISES_PADRAO)) or ANALISES_PADRAO)
        st.session_state["rel_analises_total"] = max(total, ANALISES_PADRAO)
        for k, v in campos.items():
            if str(k).startswith("_"):
                continue
            st.session_state[k] = v
        st.session_state["empresa_ativa"] = "aqua_gestao"
        # Correção: não alterar key de widget após renderização.
        # st.session_state["empresa_seletor_admin_sidebar_definitivo"] = ...
        st.session_state["_relatorio_rt_rascunho_restaurado"] = True
        return True
    except Exception as e:
        st.session_state["_relatorio_rt_restore_erro"] = f"{type(e).__name__}: {e}"
        return False


def _relatorio_rt_renderizar_painel_rascunho():
    try:
        rasc = _relatorio_rt_carregar_rascunho()
        if rasc:
            salvo_em = rasc.get("salvo_em", "horário não informado")
            motivo = rasc.get("motivo", "autosave")
            st.warning(f"🛡️ Rascunho de Relatório RT encontrado ({salvo_em} | {motivo}).")
            c_r1, c_r2, c_r3 = st.columns([1, 1, 2])
            with c_r1:
                if st.button("♻️ Restaurar rascunho RT", key="btn_restaurar_rascunho_relatorio_rt", use_container_width=True):
                    if _relatorio_rt_aplicar_rascunho(rasc):
                        st.success("Rascunho restaurado. Recarregando tela...")
                        st.rerun()
                    else:
                        st.error("Não foi possível restaurar o rascunho.")
            with c_r2:
                if st.button("🗑️ Descartar rascunho", key="btn_descartar_rascunho_relatorio_rt", use_container_width=True):
                    try:
                        _relatorio_rt_rascunho_path().unlink(missing_ok=True)
                        st.success("Rascunho descartado.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro ao descartar rascunho: {e}")
        if st.session_state.get("_relatorio_rt_rascunho_restaurado"):
            st.success("✅ Rascunho do relatório RT restaurado nesta sessão.")
            st.session_state["_relatorio_rt_rascunho_restaurado"] = False
        ultimo = st.session_state.get("_relatorio_rt_ultimo_autosave", "")
        if ultimo:
            st.caption(f"🛡️ Autosave do relatório RT ativo. Último salvamento: {ultimo}")
    except Exception as e:
        st.caption(f"Painel de rascunho RT indisponível: {e}")



# =========================================
# ADMIN — LOGIN FIXO POR EMPRESA E SESSÃO ESTÁVEL
# =========================================
ADMIN_SESSION_TIMEOUT_MIN = 12 * 60  # 12 horas. O admin também pode sair manualmente.


def _admin_pin_configurado() -> str:
    """Lê PIN administrativo do Streamlit Secrets, com fallback para o PIN atual."""
    try:
        cfg_admin = st.secrets.get("admin", {})
        pin = str(cfg_admin.get("pin", "") or "").strip()
        if pin:
            return pin
    except Exception:
        pass
    return "@Anajullya10"


def _admin_pin_valido(pin_digitado: str) -> bool:
    return str(pin_digitado or "").strip() == _admin_pin_configurado()


def _admin_nome_empresa(codigo: str) -> str:
    return "Bem Star Piscinas" if codigo == "bem_star" else "Aqua Gestão"


def _admin_icone_empresa(codigo: str) -> str:
    return "⭐" if codigo == "bem_star" else "🔵"


def _admin_entrar_empresa(codigo_empresa: str):
    """Centraliza login administrativo para evitar troca acidental de empresa por widgets antigos."""
    codigo_empresa = "bem_star" if codigo_empresa == "bem_star" else "aqua_gestao"
    agora = datetime.now().isoformat(timespec="seconds")
    st.session_state["modo_atual"] = "escritorio"
    st.session_state["admin_logado"] = True
    st.session_state["admin_empresa_fixa"] = codigo_empresa
    st.session_state["empresa_ativa"] = codigo_empresa
    st.session_state["empresa_login_admin_atual"] = _admin_nome_empresa(codigo_empresa)
    st.session_state["empresa_selecionada_admin"] = f"{_admin_icone_empresa(codigo_empresa)} {_admin_nome_empresa(codigo_empresa)}"
    st.session_state["mostrar_pin_admin"] = False
    st.session_state["admin_login_em"] = agora
    st.session_state["admin_ultimo_ping"] = agora
    st.session_state.pop("pin_admin_input", None)
    st.session_state.pop("empresa_login_admin_escolha", None)


def _admin_sessao_valida() -> bool:
    """Mantém o admin logado durante a sessão e só expira após limite alto de segurança.

    Recuperação robusta: se admin_logado=True mas modo_atual foi corrompido ou
    perdido (ex.: reconexão no Streamlit Cloud), força modo_atual='escritorio'
    para evitar que o painel fique preso na tela de login.
    """
    # Recuperação: admin marcado como logado mas modo_atual não está em escritorio.
    # Isso ocorre quando a sessão é parcialmente restaurada pelo browser após
    # hibernação ou reconexão no Streamlit Cloud.
    # v5: guard da sidebar também protege o operador
    if (st.session_state.get("admin_logado")
            and st.session_state.get("modo_atual") != "escritorio"
            and st.session_state.get("modo_atual") != "operador"):
        st.session_state["modo_atual"] = "escritorio"

    if st.session_state.get("modo_atual") != "escritorio":
        return True
    if not st.session_state.get("admin_logado"):
        # Compatibilidade: se app antigo entrou em escritório, reconstrói o estado admin.
        st.session_state["admin_logado"] = True
        st.session_state["admin_empresa_fixa"] = st.session_state.get("empresa_ativa", "aqua_gestao")
        st.session_state["admin_login_em"] = datetime.now().isoformat(timespec="seconds")
    empresa = st.session_state.get("admin_empresa_fixa") or st.session_state.get("empresa_ativa", "aqua_gestao")
    if empresa not in ("aqua_gestao", "bem_star"):
        empresa = "aqua_gestao"
    # trava empresa ativa pela empresa escolhida no login, impedindo widgets antigos de alternarem painel
    st.session_state["empresa_ativa"] = empresa
    st.session_state["admin_empresa_fixa"] = empresa
    st.session_state["admin_ultimo_ping"] = datetime.now().isoformat(timespec="seconds")
    return True


def _admin_sair_para_entrada(abrir_login: bool = True):
    """Sai do administrativo sem apagar rascunhos/relatórios em andamento."""
    try:
        if st.session_state.get("empresa_ativa") == "aqua_gestao" and "_relatorio_rt_salvar_rascunho" in globals():
            _relatorio_rt_salvar_rascunho("logout_admin")
    except Exception:
        pass
    for chave in [
        "admin_logado", "admin_empresa_fixa", "admin_login_em", "admin_ultimo_ping",
        "empresa_login_admin_atual", "empresa_login_admin_escolha", "pin_admin_input",
    ]:
        st.session_state.pop(chave, None)
    st.session_state["modo_atual"] = "entrada"
    st.session_state["mostrar_pin_admin"] = bool(abrir_login)


def _admin_keepalive_browser():
    """Keepalive desativado — causava loop infinito de rerenders no Streamlit 1.56.
    O components.v1.html com fetch() dispara rerun a cada chamada, impedindo o painel
    de estabilizar. Sessão é mantida pelo próprio Streamlit Cloud sem necessidade de ping.
    """
    pass


def _admin_render_login_empresa():
    """Renderiza login administrativo: primeiro escolhe empresa, depois informa PIN."""
    st.markdown("**Escolha a empresa do acesso administrativo:**")
    c_emp1, c_emp2 = st.columns(2)
    with c_emp1:
        if st.button("🔵 Aqua Gestão", key="btn_admin_login_aqua", use_container_width=True):
            st.session_state["admin_empresa_pendente"] = "aqua_gestao"
    with c_emp2:
        if st.button("⭐ Bem Star Piscinas", key="btn_admin_login_bemstar", use_container_width=True):
            st.session_state["admin_empresa_pendente"] = "bem_star"

    empresa_pendente = st.session_state.get("admin_empresa_pendente", "aqua_gestao")
    st.info(f"Empresa selecionada para login: **{_admin_icone_empresa(empresa_pendente)} {_admin_nome_empresa(empresa_pendente)}**")

    pin_admin = st.text_input(
        "PIN administrativo",
        type="password",
        key="pin_admin_input",
        placeholder="Digite o PIN administrativo",
        label_visibility="collapsed",
    )
    if st.button("Entrar no escritório", key="btn_pin_admin_ok", use_container_width=True):
        if _admin_pin_valido(pin_admin):
            _admin_entrar_empresa(empresa_pendente)
            st.rerun()
        else:
            st.error("PIN incorreto.")


# Garante que a empresa ativa fique travada pela empresa escolhida no login administrativo.
_admin_sessao_valida()
_admin_keepalive_browser()

# =========================================
# TOPO
# =========================================

# Inicializa empresa como Aqua Gestão se sessão nova
if "empresa_selecionada" not in st.session_state:
    st.session_state["empresa_selecionada"] = "🔵 Aqua Gestão"

logo = encontrar_logo()
logo_bs = encontrar_logo_bem_star()

# Topo só aparece no modo escritório (não na entrada nem no operador)
if st.session_state.get("modo_atual", "entrada") == "escritorio":
    _empresa_topo = st.session_state.get("empresa_ativa", "aqua_gestao")
    _eh_bs_topo = _empresa_topo == "bem_star"

    col_top1, col_top2 = st.columns([1, 5])
    with col_top1:
        if _eh_bs_topo:
            if logo_bs:
                st.image(str(logo_bs), width=150)
        else:
            if logo:
                st.image(str(logo), width=150)
    with col_top2:
        if _eh_bs_topo:
            st.markdown(
                f"""
                <div class="top-card">
                    <div class="top-title">Bem Star Piscinas</div>
                    <div class="top-subtitle">
                        Limpeza e Manutenção de Piscinas
                    </div>
                    <div>
                        <span class="info-badge">CNPJ: {CNPJ_BEM_STAR}</span>
                        <span class="info-badge">Uberlândia/MG</span>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f"""
                <div class="top-card">
                    <div class="top-title">{APP_TITLE}</div>
                    <div class="top-subtitle">
                        Responsabilidade Técnica, controle documental, relatórios, POPs e ART
                    </div>
                    <div>
                        <span class="info-badge">{RESPONSAVEL_TÉCNICO}</span>
                        <span class="info-badge">{CRQ}</span>
                        <span class="info-badge">Windows + Word + DOCX/PDF</span>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

# =========================================
# SIDEBAR – CONFIG + HISTÓRICO
# =========================================

with st.sidebar:
    if st.session_state.get("modo_atual", "entrada") == "escritorio":
        # === EMPRESA ADMINISTRATIVA FIXA PELO LOGIN — INÍCIO ===
        # _SIDEBAR_EMPRESA_FIXA_PELO_LOGIN_V1_
        st.markdown("### Painel administrativo")

        _empresa_atual_admin = st.session_state.get("empresa_ativa", "aqua_gestao")
        if _empresa_atual_admin not in ("aqua_gestao", "bem_star"):
            _empresa_atual_admin = "aqua_gestao"
            st.session_state["empresa_ativa"] = _empresa_atual_admin

        _nome_painel_admin = "Bem Star Piscinas" if _empresa_atual_admin == "bem_star" else "Aqua Gestão"
        _icone_painel_admin = "⭐" if _empresa_atual_admin == "bem_star" else "🔵"

        st.info(f"{_icone_painel_admin} Empresa logada: **{_nome_painel_admin}**")

        if st.button("🚪 Sair do admin", key="btn_trocar_empresa_login_admin", use_container_width=True):
            _admin_sair_para_entrada(abrir_login=True)
            st.rerun()

        st.caption(f"Painel ativo: {_nome_painel_admin}")
        st.divider()
        # === EMPRESA ADMINISTRATIVA FIXA PELO LOGIN — FIM ===

    if st.session_state.get("modo_atual", "entrada") != "operador":
        st.header("Histórico recente")

    st.number_input(
        "Lembrete de vencimento (dias)",
        min_value=1,
        max_value=180,
        step=1,
        key="alerta_vencimento_dias",
    )

    st.text_input(
        "Filtrar condomínio",
        key="filtro_historico",
        placeholder="Digite parte do nome...",
    )

    # v4 — no painel Bem Star, o histórico local completo fica sob demanda.
    # Evita varredura pesada de GENERATED_DIR antes dos módulos Bem Star aparecerem.
    _carregar_historico_sidebar = True
    if st.session_state.get("empresa_ativa") == "bem_star":
        _carregar_historico_sidebar = st.checkbox(
            "Carregar histórico local",
            value=False,
            key="sidebar_carregar_historico_bemstar",
            help="Carrega pastas/arquivos locais somente quando necessário.",
        )

    historico = listar_historico() if _carregar_historico_sidebar else []
    filtro = st.session_state.filtro_historico.strip().lower()

    if filtro:
        historico = [h for h in historico if filtro in h["nome"].lower()]

    if not historico:
        st.caption("Histórico local não carregado." if not _carregar_historico_sidebar else "Nenhum histórico encontrado.")
    else:
        for item in historico:
            nome_cond = item["nome"]
            pasta = item["pasta"]
            arquivos = item["arquivos"]
            folder_key = chave_segura(str(pasta))
            status = status_vencimento(item["data_fim"], st.session_state.alerta_vencimento_dias)

            titulo = f"{nome_cond} ({item['total_arquivos']})"

            with st.expander(titulo, expanded=False):
                st.caption(str(pasta))
                st.markdown(
                    f"<span class='status-badge {status['css']}'>{status['rotulo']}</span>",
                    unsafe_allow_html=True,
                )
                st.caption(status["mensagem"])

                col_h1, col_h2 = st.columns(2)
                with col_h1:
                    if st.button(
                        "Carregar dados",
                        key=f"carregar_dados_{folder_key}",
                        use_container_width=True,
                    ):
                        dados_salvos = carregar_dados_condominio(pasta)
                        if dados_salvos:
                            aplicar_dados_no_formulario(dados_salvos)
                            st.success("Dados carregados no formulário.")
                            st.rerun()
                        else:
                            st.warning("Nenhum cadastro salvo encontrado para este condomínio.")

                with col_h2:
                    if st.button(
                        "Abrir pasta",
                        key=f"abrir_pasta_{folder_key}",
                        use_container_width=True,
                    ):
                        abrir_pasta_windows(pasta)

                if st.session_state.confirm_delete_folder == folder_key:
                    if st.button(
                        "Confirmar pasta",
                        key=f"confirmar_del_pasta_{folder_key}",
                        type="primary",
                        use_container_width=True,
                    ):
                        try:
                            deletar_pasta_condominio(pasta)
                            st.session_state.confirm_delete_folder = ""
                            st.success(f"Pasta de '{nome_cond}' excluída.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao excluir pasta: {e}")
                else:
                    if st.button(
                        "Excluir pasta",
                        key=f"pedir_del_pasta_{folder_key}",
                        use_container_width=True,
                    ):
                        st.session_state.confirm_delete_folder = folder_key
                        st.rerun()

                if st.session_state.confirm_delete_folder == folder_key:
                    st.markdown(
                        "<div class='confirm-box'><strong>Confirma excluir toda a pasta deste condomínio?</strong></div>",
                        unsafe_allow_html=True,
                    )
                    if st.button(
                        "Cancelar exclusão da pasta",
                        key=f"cancelar_del_pasta_{folder_key}",
                        use_container_width=True,
                    ):
                        st.session_state.confirm_delete_folder = ""
                        st.rerun()

                st.markdown("**Arquivos:**")
                if not arquivos:
                    st.caption("Sem arquivos.")
                else:
                    for arq in arquivos:
                        caminho = arq["path"]
                        nome = arq["name"]
                        tipo_doc = arq["tipo_doc"]
                        tipo_ext = arq["tipo_ext"]
                        modificado = arq["modificado_em"]

                        file_key = chave_segura(str(caminho))

                        st.markdown(
                            f"**{tipo_doc} • {tipo_ext}**\n\n"
                            f"<div class='history-meta'>{nome}</div>"
                            f"<div class='history-meta'>Atualizado em: {modificado}</div>",
                            unsafe_allow_html=True,
                        )

                        col_a1, col_a2, col_a3 = st.columns([1.15, 1.0, 0.85])

                        with col_a1:
                            if st.button(
                                "Abrir arquivo",
                                key=f"abrir_arq_{file_key}",
                                use_container_width=True,
                            ):
                                abrir_arquivo_windows(caminho)

                        with col_a2:
                            if st.button(
                                "Abrir pasta",
                                key=f"abrir_pasta_arq_{file_key}",
                                use_container_width=True,
                            ):
                                abrir_pasta_windows(caminho.parent)

                        with col_a3:
                            if st.session_state.confirm_delete_file == file_key:
                                if st.button(
                                    "OK",
                                    key=f"confirmar_del_arq_{file_key}",
                                    type="primary",
                                    use_container_width=True,
                                ):
                                    try:
                                        deletar_arquivo_individual(caminho)
                                        st.session_state.confirm_delete_file = ""
                                        st.success(f"Arquivo excluído: {nome}")
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"Erro ao excluir arquivo: {e}")
                            else:
                                if st.button(
                                    "🗑️",
                                    key=f"pedir_del_arq_{file_key}",
                                    help="Excluir somente este arquivo",
                                    use_container_width=True,
                                ):
                                    st.session_state.confirm_delete_file = file_key
                                    st.rerun()

                        if st.session_state.confirm_delete_file == file_key:
                            st.markdown(
                                "<div class='confirm-box'>Excluir este arquivo?</div>",
                                unsafe_allow_html=True,
                            )
                            if st.button(
                                "Cancelar",
                                key=f"cancelar_del_arq_{file_key}",
                                use_container_width=True,
                            ):
                                st.session_state.confirm_delete_file = ""
                                st.rerun()

                        st.markdown("---")

# =========================================
# MODO DE OPERAÇÃO — TELA DE ENTRADA
# =========================================

# PIN padrão — altere aqui para trocar o PIN do operador
PIN_OPERADOR = "2940"

# =========================================
# v5 — BLINDAGEM DO SESSION STATE DO OPERADOR
# Estas chaves são inicializadas uma única vez e NUNCA apagadas em reruns normais.
# Isso impede que o operador perca sessão durante digitação de parâmetros.
# =========================================
st.session_state.setdefault("modo_atual", "entrada")
st.session_state.setdefault("op_pin_ok", False)
st.session_state.setdefault("op_dados_atual", {})
st.session_state.setdefault("op_limpar_campos", False)
st.session_state.setdefault("op_salvo_sucesso", None)
st.session_state.setdefault("_op_ultimo_lancamento", None)

# Inicializa o modo se não estiver definido (mantido para compatibilidade)
if "modo_atual" not in st.session_state:
    st.session_state["modo_atual"] = "entrada"

# GUARD DUPLO v5: protege o modo operador.
# O guard só redireciona para escritório se o operador NÃO estiver logado.
# Antes, esse guard derrubava o operador durante preenchimento caso admin_logado
# estivesse definido em memória de sessão anterior.
if (st.session_state.get("admin_logado")
        and st.session_state.get("modo_atual") != "escritorio"
        and st.session_state.get("modo_atual") != "operador"):
    st.session_state["modo_atual"] = "escritorio"
    _empresa_fix = st.session_state.get("admin_empresa_fixa") or st.session_state.get("empresa_ativa", "aqua_gestao")
    if _empresa_fix not in ("aqua_gestao", "bem_star"):
        _empresa_fix = "aqua_gestao"
    st.session_state["empresa_ativa"] = _empresa_fix
    st.session_state["admin_empresa_fixa"] = _empresa_fix

# Inicializa e preserva a empresa ativa entre reruns.
# Isso evita o Modo Campo voltar para Aqua Gestão quando o operador escolheu Bem Star.
if "empresa_ativa" not in st.session_state or st.session_state.get("empresa_ativa") not in ("aqua_gestao", "bem_star"):
    st.session_state["empresa_ativa"] = "aqua_gestao"

if "empresa_selecionada_admin" not in st.session_state:
    st.session_state["empresa_selecionada_admin"] = (
        "⭐ Bem Star Piscinas" if st.session_state.get("empresa_ativa") == "bem_star" else "🔵 Aqua Gestão"
    )

# Compatibilidade com seletor antigo (Modo Campo ainda usa st.radio internamente)
_modo_interno = st.session_state.get("modo_atual", "entrada")

# ---- TELA DE ENTRADA ----
if _modo_interno == "entrada":
    st.markdown("""
    <style>
    .entrada-card {
        border: 1px solid rgba(20,85,160,0.12);
        border-radius: 20px;
        padding: 18px 22px 16px 22px;
        background: linear-gradient(135deg, #ffffff 0%, #f7fbff 100%);
        box-shadow: 0 6px 20px rgba(10,50,100,0.06);
        margin: 4px 0 10px 0;
        text-align: left;
    }
    .entrada-eyebrow { font-size: 0.76rem; font-weight: 700; letter-spacing: 0.08em; text-transform: uppercase; color: #6f86a2; margin-bottom: 4px; }
    .entrada-title { font-size: 1.28rem; font-weight: 700; color: #0d3d75; margin-bottom: 4px; }
    .entrada-sub { font-size: 0.90rem; color: #5d7288; margin-bottom: 14px; line-height: 1.4; }
    .entrada-admin-note { font-size: 0.76rem; color: #8ea0b5; margin-top: 10px; }
    </style>
    """, unsafe_allow_html=True)

    col_e1, col_e2, col_e3 = st.columns([1.2, 1.6, 1.2])
    with col_e2:
        st.markdown('<div class="entrada-eyebrow"></div>', unsafe_allow_html=True)
        st.markdown('<div class="entrada-title">Acesso do Operador</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="entrada-sub">Entre com seu PIN para acessar os condomínios liberados.<br>'
            'O sistema identifica automaticamente os clientes vinculados.</div>',
            unsafe_allow_html=True
        )

        # Modo Operador: empresa não é escolhida na entrada.
        # O PIN define automaticamente os condomínios liberados ao operador,
        # incluindo clientes Aqua Gestão, Bem Star Piscinas ou Ambas.
        st.caption("O PIN direciona automaticamente para os condomínios vinculados ao operador.")

        if st.button("📱 Acessar como Operador", type="primary", use_container_width=True):
            # Operador entra por PIN; empresa não é escolhida pelo operador.
            st.session_state["modo_atual"] = "operador"
            st.session_state["mostrar_pin_admin"] = False
            st.session_state["op_pin_ok"] = False
            st.rerun()

        if st.button("🔐 Acesso administrativo", use_container_width=True, key="btn_admin_limpo"):
            st.session_state["mostrar_pin_admin"] = not st.session_state.get("mostrar_pin_admin", False)

        st.markdown('<div class="entrada-admin-note">Uso interno do escritório</div>', unsafe_allow_html=True)

        if st.session_state.get("mostrar_pin_admin"):
            # _LOGIN_ADMIN_EMPRESA_DEFINITIVO_V3_
            # Empresa é definida apenas no login administrativo. A sidebar não alterna mais empresa.
            _admin_render_login_empresa()


    st.stop()


# CSS dinâmico: separa visualmente os módulos por empresa ativa.
# Aqua Gestão vê somente RT/controle técnico; Bem Star vê somente limpeza/manutenção.
_empresa_css_flag = st.session_state.get("empresa_ativa", "aqua_gestao")
if _empresa_css_flag == "bem_star":
    st.markdown("""
    <style>
    div.aq-only { display: none !important; }
    </style>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <style>
    div.bs-only { display: none !important; }
    </style>
    """, unsafe_allow_html=True)

# Alias para compatibilidade
modo = "Modo Escritório" if _modo_interno == "escritorio" else (
    "📱 Modo Operador (Campo / Celular)" if _modo_interno == "operador" else "Modo Escritório"
)

# Botão de saída controlada. No admin, só sai quando clicar explicitamente em sair.
if _modo_interno in ("escritorio", "operador"):
    if _modo_interno == "escritorio":
        if st.button("🚪 Sair do admin", key="btn_voltar_inicio"):
            _admin_sair_para_entrada(abrir_login=True)
            st.rerun()
    else:
        if st.button("← Voltar à tela inicial", key="btn_voltar_inicio"):
            st.session_state["modo_atual"] = "entrada"
            st.session_state["op_pin_ok"] = False
            st.session_state.pop("op_dados_atual", None)
            st.rerun()

# =========================================
# MODO OPERADOR — LANÇAMENTO DE CAMPO
# =========================================

if modo == "📱 Modo Operador (Campo / Celular)":

    st.markdown("""
    <style>
    section[data-testid="stSidebar"] { display: none !important; }
    .main .block-container { padding: 0.35rem 0.7rem 1.2rem !important; max-width: 100% !important; }
    .op-card {
        border: 1px solid rgba(20,85,160,0.18);
        border-radius: 16px;
        padding: 12px 14px;
        background: #ffffff;
        margin-bottom: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
    }
    .op-title { font-size: 1.02rem; font-weight: 700; color: #0d3d75; margin-bottom: 2px; line-height: 1.2; }
    /* Campo oculto para captura de assinatura via canvas HTML */
    div[data-testid="stTextInput"]:has(input[aria-label="assinatura_b64_hidden"]) { display: none !important; }
    .op-sub { font-size: 0.78rem; color: #5d7288; margin-bottom: 6px; }
    .op-salvo {
        border: 1px solid rgba(30,140,70,0.3);
        border-radius: 12px;
        padding: 12px 14px;
        background: rgba(30,140,70,0.08);
        color: #1a6e3a;
        font-size: 0.95rem;
        margin-top: 8px;
    }
    .stTextInput input, .stTextArea textarea {
        font-size: 1rem !important;
        min-height: 42px !important;
        border-radius: 10px !important;
    }
    .stButton > button {
        min-height: 46px !important;
        font-size: 0.98rem !important;
        border-radius: 12px !important;
    }
    .stTextInput label, .stSelectbox label, .stTextArea label {
        font-size: 0.95rem !important;
        font-weight: 600 !important;
        color: #1a3a5c !important;
    }
    .element-container { margin-bottom: 4px !important; }
    [data-testid="stExpander"] { border-radius: 12px !important; }
    [data-testid="stExpander"] details summary p { font-size: 0.96rem !important; font-weight: 600 !important; }
    .op-chip { display:inline-block; padding:4px 10px; border-radius:999px; background:#edf5ff; border:1px solid #d3e6ff; color:#134b8a; font-size:0.78rem; margin: 2px 6px 6px 0; }
    .op-note-compact { font-size:0.86rem; color:#4f657c; margin: 2px 0 8px 0; }
    
    .pin-box {
        padding: 20px 0;
        text-align: center;
    }
    </style>
    """, unsafe_allow_html=True)

    # ---- TELA DE PIN ----
    if not st.session_state.get("op_pin_ok"):
        st.markdown('<div class="pin-box">', unsafe_allow_html=True)
        st.markdown("### 🔐 Área do Operador")
        st.markdown("**Acesso simplificado por PIN**")
        st.markdown("Digite o PIN para acessar o lançamento de campo dos condomínios autorizados.")
        pin_digitado = st.text_input("PIN", type="password", key="op_pin_input",
            placeholder="Digite o PIN", label_visibility="collapsed", max_chars=20)
        if st.button("Entrar", type="primary", use_container_width=True):
            op_dados = validar_pin_operador(pin_digitado.strip())
            if op_dados:
                st.session_state["op_pin_ok"] = True
                st.session_state["op_dados_atual"] = op_dados
                st.rerun()
            else:
                st.error("PIN incorreto. Tente novamente.")
        st.stop()

    # Dados do operador logado
    _op_atual = st.session_state.get("op_dados_atual", {"nome": "Operador", "acesso_total": True, "condomínios": []})
    _op_nome_logado = _op_atual.get("nome", "Operador")
    _op_acesso_total = _op_atual.get("acesso_total", False)
    _op_conds_permitidos = _condominios_organizar(_op_atual.get("condomínios", []))

    if st.button("🔒 Sair / Trocar operador", use_container_width=False):
        st.session_state["op_pin_ok"] = False
        st.session_state.pop("op_dados_atual", None)
        st.session_state.pop("op_sel_cond", None)
        st.rerun()

    # v5: diagnóstico discreto da sessão — ajuda a confirmar se a sessão continua viva
    with st.expander("🔧 Diagnóstico da sessão (operador)", expanded=False):
        st.write({
            "modo_atual": st.session_state.get("modo_atual"),
            "op_pin_ok": st.session_state.get("op_pin_ok"),
            "operador": _op_nome_logado if "_op_nome_logado" in dir() else "—",
            "condominio": st.session_state.get("op_sel_cond", "—"),
            "tem_rascunho_local": bool(st.session_state.get("_rascunho_operador_pendente")),
            "limpar_campos_pendente": st.session_state.get("op_limpar_campos"),
        })

    # v4 — operador não escolhe empresa.
    # A empresa administrativa não deve filtrar o modo campo. O PIN mostra os
    # condomínios vinculados ao operador, sejam Aqua Gestão, Bem Star ou ambos.
    _empresa_op_codigo = "operador_multibase"
    _empresa_op_nome = "Aqua Gestão / Bem Star"
    _empresa_op_titulo = "📱 Modo Campo — condomínios vinculados ao PIN"

    st.markdown('<div class="op-card">', unsafe_allow_html=True)
    st.markdown(f'<div class="op-title">📱 {_empresa_op_titulo}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="op-sub">Operador identificado: <strong>{_op_nome_logado}</strong> • Empresa ativa: <strong>{_empresa_op_nome}</strong></div>', unsafe_allow_html=True)
    st.markdown('<span class="op-chip">Condomínios permitidos por PIN</span><span class="op-chip">Aqua/Bem Star conforme vínculo do cliente</span>', unsafe_allow_html=True)

    _salvo = st.session_state.pop("op_salvo_sucesso", None)
    if _salvo:
        st.markdown(f"""
        <div class="op-salvo">
            ✅ <strong>Lançamento salvo!</strong><br>
            Condomínio: {_salvo['nome']}<br>
            Data: {_salvo['data']}<br>
            Operador: {_salvo['operador']}<br>
            Total este mês: {_salvo['total']}
        </div>
        """, unsafe_allow_html=True)
        # Botão de relatório do lançamento recém salvo
        _ult_lanc = st.session_state.get("_op_ultimo_lancamento")
        if _ult_lanc:
            nome_arq = limpar_nome_arquivo(f"Relatorio_Visita_{_salvo['nome']}_{_salvo['data'].replace('/','')}")
            with st.spinner("Gerando PDF..."):
                try:
                    pdf_bytes = gerar_pdf_relatorio_visita(_ult_lanc, _salvo["nome"])
                    st.download_button(
                        "📄 Baixar PDF desta visita",
                        data=pdf_bytes,
                        file_name=f"{nome_arq}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="btn_dl_relatorio_visita",
                    )
                    st.caption("Baixe e compartilhe diretamente pelo WhatsApp.")
                    # Botao PDF Bem Star Premium
                    if st.session_state.get("empresa_ativa") == "bem_star":
                        try:
                            pdf_bs = gerar_pdf_relatorio_visita_bem_star(_ult_lanc, _salvo["nome"])
                            nome_bs = limpar_nome_arquivo(f"Relatorio_BemStar_{_salvo['nome']}_{_salvo['data'].replace('/','')}")
                            st.download_button(
                                "⭐ Baixar PDF Bem Star (com capa premium)",
                                data=pdf_bs,
                                file_name=f"{nome_bs}.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                                key="btn_dl_relatorio_visita_bs",
                            )
                        except Exception as _ebs:
                            st.warning(f"PDF Bem Star nao gerado: {_ebs}")
                except Exception as _e:
                    import traceback
                    _erro_pdf = str(_e)
                    if "reportlab" in _erro_pdf.lower():
                        st.error("Dependência ausente para gerar PDF: reportlab. Instale no ambiente e gere novamente.")
                    st.warning(f"PDF não gerado: {_e}. Baixando versão HTML como alternativa.")
                    st.code(traceback.format_exc(), language="text")
                    html_rel = gerar_html_relatorio_visita(_ult_lanc, _salvo["nome"])
                    st.download_button(
                        "📄 Baixar relatório (HTML)",
                        data=html_rel.encode("utf-8"),
                        file_name=f"{nome_arq}.html",
                        mime="text/html",
                        use_container_width=True,
                        key="btn_dl_relatorio_visita_html",
                    )

    # Busca TODOS os clientes no modo operador; o filtro é por PIN/condomínio, não por empresa.
    @st.cache_data(ttl=60)
    def _buscar_clientes_sheets_completo():
        return sheets_listar_clientes_completo()

    # v5: restauração protegida — falha não derruba o app
    try:
        if aplicar_restauracao_pendente_operador():
            st.rerun()
    except Exception as _e_rest:
        st.warning(f"Aviso ao restaurar rascunho: {_e_rest}")

    if st.session_state.pop("_rascunho_operador_restaurado_msg", False):
        st.success("✅ Rascunho restaurado! Continue de onde parou.")

    _clientes_todos_op_raw = _buscar_clientes_sheets_completo()

    # v4 — modo operador multibase: não filtra por empresa administrativa.
    # O acesso é determinado pelo PIN e pelos vínculos do cliente.
    if _empresa_op_codigo == "operador_multibase":
        _clientes_todos_op = list(_clientes_todos_op_raw or [])
    else:
        _clientes_todos_op = filtrar_clientes_por_empresa(_clientes_todos_op_raw, _empresa_op_codigo)
    clientes_mapa_op = {c["nome"]: c for c in _clientes_todos_op if c.get("nome")}

    # Combina clientes do Sheets com os locais, respeitando a empresa ativa
    pastas_disponiveis = sorted([
        p for p in GENERATED_DIR.iterdir() if p.is_dir()
    ], key=lambda p: p.name) if GENERATED_DIR.exists() else []

    for p in pastas_disponiveis:
        dados_c = carregar_dados_condominio(p) or {}
        nome_ex = dados_c.get("nome_condominio", humanizar_nome_pasta(p.name)) if dados_c else humanizar_nome_pasta(p.name)
        if not nome_ex:
            continue
        cliente_local = _enriquecer_cliente_com_dados_locais({
            "nome": nome_ex,
            "empresa": dados_c.get("empresa", "Aqua Gestão"),
            "operadores_vinculados": dados_c.get("operadores_vinculados", []),
            "servicos": dados_c.get("servicos", {}),
        })
        _serv_local = _normalizar_servicos_cliente(cliente_local)
        if _empresa_op_codigo == "bem_star" and not _serv_local.get("limpeza"):
            continue
        if _empresa_op_codigo == "aqua_gestao" and not _serv_local.get("rt"):
            continue
        # operador_multibase mantém Aqua, Bem Star e clientes com ambos.
        clientes_mapa_op.setdefault(nome_ex, cliente_local)

    opcoes_cond_todas = list(clientes_mapa_op.keys())

    _op_conds_vinculo_direto = []
    _op_nome_norm = _normalizar_chave_acesso(_op_nome_logado)
    if _op_nome_norm and not _op_acesso_total:
        for _nome_cond, _cliente_cond in clientes_mapa_op.items():
            _ops_vinc = _normalizar_lista_textos_unicos(_cliente_cond.get("operadores_vinculados", []))
            if any(_normalizar_chave_acesso(_op) == _op_nome_norm for _op in _ops_vinc):
                _op_conds_vinculo_direto.append(_nome_cond)

    # Operador logado via PIN — nome vem do cadastro, não digitado
    _op_nome_logado_disp = _op_nome_logado if _op_nome_logado != "Operador" else ""
    op_operador = st.text_input(
        "Operador identificado",
        key="op_operador",
        value=_op_nome_logado_disp,
        placeholder="Ex.: João Silva",
        help="Preenchido automaticamente pelo seu PIN de acesso.",
        disabled=True
    )

    # ── FASE 2: Filtra condomínios combinando vínculo direto E permitidos pelo PIN ──
    if _op_acesso_total or any(_normalizar_chave_acesso(c) == "todos" for c in _op_conds_permitidos):
        opcoes_cond = opcoes_cond_todas
    else:
        # Une os dois conjuntos: condomínios vinculados diretamente + liberados pelo PIN
        _todos_permitidos = list(dict.fromkeys(_op_conds_vinculo_direto + list(_op_conds_permitidos)))
        if _todos_permitidos:
            opcoes_cond = _resolver_condominios_permitidos_exatos(_todos_permitidos, opcoes_cond_todas)
            if opcoes_cond:
                st.markdown(
                    f'<div class="op-note-compact">✅ {len(opcoes_cond)} condomínio(s) liberado(s) pelo seu PIN.</div>',
                    unsafe_allow_html=True
                )
            else:
                st.warning("Nenhum condomínio disponível para seu acesso. Contate o administrador.")
        else:
            opcoes_cond = []
            st.warning("Seu PIN ainda não possui condomínios vinculados. Contate o administrador.")

    op_usar_novo = st.checkbox("Lançar em local ainda não cadastrado", key="op_novo_cond")
    if op_usar_novo:
        op_nome_cond = st.text_input("Nome do local", key="op_nome_livre", placeholder="Ex.: Residencial Aquarela")
    else:
        if opcoes_cond:
            # Fase 2: card fixo APENAS quando o operador tem exatamente 1 condomínio
            # registrado no seu cadastro (não conta acesso_total nem lista global)
            _conds_cadastrados = _op_conds_vinculo_direto or (
                [c for c in _op_conds_permitidos
                 if _normalizar_chave_acesso(c) not in ("todos", "")]
                if not _op_acesso_total else []
            )
            _total_permitidos = len(list(dict.fromkeys(
                _op_conds_vinculo_direto + [c for c in _op_conds_permitidos
                if _normalizar_chave_acesso(c) not in ("todos","")]
            )))
            if _total_permitidos == 1 and len(opcoes_cond) == 1 and not _op_acesso_total:
                op_nome_cond = opcoes_cond[0]
                st.markdown(
                    f'<div class="op-card" style="background:#f0f7ff;border-color:#1565A8;">'
                    f'<div class="op-title">📍 {op_nome_cond}</div>'
                    f'<div class="op-sub">Condomínio selecionado automaticamente pelo PIN</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )
            else:
                op_nome_cond = st.selectbox("Condomínio", opcoes_cond, key="op_sel_cond")

            # Badge de tipo de serviço vinculado ao condomínio selecionado
            if op_nome_cond and op_nome_cond in clientes_mapa_op:
                _cliente_sel = clientes_mapa_op[op_nome_cond]
                _servs = _normalizar_servicos_cliente(_cliente_sel)
                _badges = []
                if _servs.get("rt"):
                    _badges.append('<span class="op-chip" style="background:#e8f0fe;border-color:#1565A8;color:#1565A8;">🔬 RT — Aqua Gestão</span>')
                if _servs.get("limpeza"):
                    _badges.append('<span class="op-chip" style="background:#fff3e0;border-color:#e65100;color:#bf360c;">🧹 Limpeza/Manutenção</span>')
                if _badges:
                    st.markdown(" ".join(_badges), unsafe_allow_html=True)
            
            # --- HISTÓRICO DE VISITAS (ÚLTIMAS 3) ---
            if op_nome_cond:
                with st.expander("🕒 Histórico recente", expanded=False):
                    with st.spinner("Buscando histórico..."):
                        historico_v = sheets_listar_lancamentos(op_nome_cond)
                        if not historico_v:
                            st.caption("Nenhuma visita anterior registrada.")
                        else:
                            # Ordena por data decrescente
                            try:
                                historico_v = sorted(
                                    historico_v, 
                                    key=lambda x: datetime.strptime(x["data"], "%d/%m/%Y") if "/" in x["data"] else datetime.strptime(x["data"], "%Y-%m-%d"),
                                    reverse=True
                                )
                            except:
                                pass
                            
                            for v in historico_v[:3]:
                                st.markdown(f"**Data: {v['data']}** | Op: {v.get('operador','–')}")
                                st.markdown(f"pH: `{v.get('ph','–')}` | CRL: `{v.get('cloro_livre','–')}` | Alc: `{v.get('alcalinidade','–')}`")
                                if v.get('problemas'):
                                    st.caption(f"⚠️ {v['problemas']}")
                                st.markdown("---")
        else:
            st.warning("Nenhum condomínio disponível. Verifique seu nome ou contate o administrador.")
            op_nome_cond = ""

    def _fmt_data_op():
        v = apenas_digitos(st.session_state.get("op_data_visita", ""))[:8]
        if len(v) <= 2:
            st.session_state["op_data_visita"] = v
        elif len(v) <= 4:
            st.session_state["op_data_visita"] = f"{v[:2]}/{v[2:]}"
        else:
            st.session_state["op_data_visita"] = f"{v[:2]}/{v[2:4]}/{v[4:]}"
        # v5: autosave na mudança de data — apenas local, nunca Sheets durante digitação
        if op_nome_cond.strip():
            try:
                _d = coletar_rascunho_operador(op_nome_cond, piscinas_ativas if "piscinas_ativas" in dir() else ["Piscina Adulto"])
                salvar_rascunho_operador(op_nome_cond, _d, salvar_sheets=False)
            except Exception:
                pass

    st.text_input("Data da visita",
        key="op_data_visita", placeholder="06/04/2026", on_change=_fmt_data_op)

    st.markdown("</div>", unsafe_allow_html=True)

    if op_nome_cond:

        # ── Piscinas deste condomínio ─────────────────────────────────────────
        # Carrega configuração de piscinas salva ou usa padrão
        _pasta_cond_op = GENERATED_DIR / slugify_nome(op_nome_cond.strip())
        _dados_cond_op = (carregar_dados_condominio(_pasta_cond_op) or {}) if _pasta_cond_op.exists() else {}
        _piscinas_config = _dados_cond_op.get("piscinas", ["Piscina Adulto"])

        # Piscinas vem automaticamente do cadastro da ADM — operador nao configura
        if _piscinas_config and _piscinas_config != ["Piscina Adulto"]:
            # ADM ja cadastrou as piscinas — usa direto
            _piscinas_ativas = _piscinas_config
            st.info(f"🏊 {len(_piscinas_ativas)} piscina(s) deste local: {chr(10).join(f'  • {p}' for p in _piscinas_ativas)}")
        else:
            # Fallback: ADM ainda nao cadastrou — operador pode configurar
            with st.expander("🏊 Piscinas (configure uma vez)", expanded=True):
                st.caption("O administrador ainda nao cadastrou as piscinas deste local. Configure abaixo.")
                _pisc_adulto  = st.checkbox("Piscina Adulto",   value="Piscina Adulto"   in _piscinas_config, key="op_pisc_adulto")
                _pisc_infant  = st.checkbox("Piscina Infantil", value="Piscina Infantil" in _piscinas_config, key="op_pisc_infantil")
                _pisc_family  = st.checkbox("Piscina Family",   value="Piscina Family"   in _piscinas_config, key="op_pisc_family")
                _pisc_outra_check = st.checkbox("Outra piscina", value=any(p not in ["Piscina Adulto","Piscina Infantil","Piscina Family"] for p in _piscinas_config), key="op_pisc_outra_check")
                _pisc_outra_nome = ""
                if _pisc_outra_check:
                    _outra_default = next((p for p in _piscinas_config if p not in ["Piscina Adulto","Piscina Infantil","Piscina Family"]), "")
                    _pisc_outra_nome = st.text_input("Nome da outra piscina", value=_outra_default, key="op_pisc_outra_nome", placeholder="Ex.: Piscina Olímpica")
                _piscinas_ativas = []
                if _pisc_adulto: _piscinas_ativas.append("Piscina Adulto")
                if _pisc_infant: _piscinas_ativas.append("Piscina Infantil")
                if _pisc_family: _piscinas_ativas.append("Piscina Family")
                if _pisc_outra_check and _pisc_outra_nome.strip(): _piscinas_ativas.append(_pisc_outra_nome.strip())
                if not _piscinas_ativas: _piscinas_ativas = ["Piscina Adulto"]
                if st.button("💾 Salvar configuração de piscinas", key="btn_salvar_piscinas"):
                    _pasta_cond_op.mkdir(parents=True, exist_ok=True)
                    _dados_upd = carregar_dados_condominio(_pasta_cond_op) or {}
                    _dados_upd["piscinas"] = _piscinas_ativas
                    _dados_upd["nome_condominio"] = op_nome_cond.strip()
                    salvar_dados_condominio(_pasta_cond_op, _dados_upd)
                    st.success(f"✅ Piscinas salvas: {', '.join(_piscinas_ativas)}")
                    st.rerun()
        piscinas_ativas = _piscinas_ativas

        # ── Indicador de autosave ────────────────────────────────────────────
        _rasc_path = GENERATED_DIR / slugify_nome(op_nome_cond.strip()) / "_rascunho_operador.json"
        if _rasc_path.exists():
            try:
                import json as _json_rasc
                with open(_rasc_path) as _rf:
                    _rasc_info = _json_rasc.load(_rf)
                _salvo_em_auto = _rasc_info.get("_rascunho_salvo_em", "")
                st.markdown(
                    f"<div style='text-align:right;font-size:0.75rem;color:#3a8a3a;padding:2px 8px;'>"
                    f"💾 Salvo automaticamente às {_salvo_em_auto.split()[-1] if _salvo_em_auto else '—'}"
                    f"</div>",
                    unsafe_allow_html=True
                )
            except Exception:
                pass

        # ── Verificar rascunho existente ─────────────────────────────────────
        _rascunho_op = carregar_rascunho_operador(op_nome_cond)
        _key_rascunho_visto = f"_rascunho_visto_{slugify_nome(op_nome_cond)}"
        if _rascunho_op and not st.session_state.get(_key_rascunho_visto):
            _salvo_em = _rascunho_op.get("_rascunho_salvo_em", "")
            # Conta fotos salvas no rascunho
            _fotos_rasc_info = _rascunho_op.get("fotos_rascunho", {})
            _total_fotos_rasc_banner = sum(len(v) for v in _fotos_rasc_info.values())
            _fotos_txt = f" • {_total_fotos_rasc_banner} foto(s) salva(s)" if _total_fotos_rasc_banner > 0 else ""
            st.markdown(f"""
            <div style="background:rgba(255,165,0,0.12);border:1px solid rgba(255,165,0,0.4);
            border-radius:12px;padding:12px 16px;margin-bottom:8px;">
            ⚠️ <strong>Rascunho encontrado</strong> — salvo em {_salvo_em}{_fotos_txt}<br>
            <span style="font-size:0.85rem;color:#aaa;">Você pode restaurar e continuar de onde parou.</span>
            </div>
            """, unsafe_allow_html=True)
            _rc1, _rc2 = st.columns(2)
            with _rc1:
                if st.button("📂 Restaurar rascunho", key="btn_restaurar_rasc",
                        type="primary", use_container_width=True):
                    st.session_state["_rascunho_operador_pendente"] = _rascunho_op
                    st.session_state[_key_rascunho_visto] = True
                    st.rerun()
            with _rc2:
                if st.button("🗑 Descartar rascunho", key="btn_descartar_rasc",
                        use_container_width=True):
                    deletar_rascunho_operador(op_nome_cond)
                    st.session_state[_key_rascunho_visto] = True
                    st.rerun()

        # Limpa campos SE houver limpeza pendente
        if st.session_state.pop("op_limpar_campos", False):
            for pisc in ["adulto","infantil","family","outra"]:
                for k in ["ph","crl","ct","alc","dc","cya"]:
                    st.session_state[f"op_{pisc}_{k}"] = ""
            for i in range(5):
                for s in ["prod","qtd","un","fin"]:
                    st.session_state[f"op_dos_{s}_{i}"] = ""
            st.session_state["op_obs_campo"] = ""
            st.session_state["op_problemas"] = ""
            st.session_state["op_resp_local"] = ""
            st.session_state["op_parecer_visita"] = "✅ Satisfatório"
            st.session_state["op_assinatura_responsavel_b64"] = ""
            st.session_state["op_assinatura_responsavel_data"] = ""
            st.session_state["_op_ass_canvas_nonce"] = st.session_state.get("_op_ass_canvas_nonce", 0) + 1

        # ── Autosave: função chamada a cada mudança de campo ─────────────────
        def _autosave_rascunho():
            """v5: salva rascunho APENAS localmente quando campo muda (on_change).
            Nunca vai ao Google Sheets durante digitação — isso causava reruns pesados
            e era a principal causa de instabilidade do modo operador.
            """
            if op_nome_cond.strip():
                try:
                    _d = coletar_rascunho_operador(op_nome_cond, piscinas_ativas)
                    salvar_rascunho_operador(op_nome_cond, _d, salvar_sheets=False)
                except Exception:
                    pass  # autosave nunca pode derrubar o app

        def _num_op(chave, label, placeholder, quinzenal=False):
            lbl = f"{label} ⏱ 15d" if quinzenal else label
            v = st.text_input(lbl, key=chave, placeholder=placeholder,
                on_change=_autosave_rascunho,
                help="Medição quinzenal — preencha somente nas visitas de medição completa." if quinzenal else None)
            return re.sub(r"[^0-9.,]", "", v).replace(",", ".")

        with st.expander("📋 Faixas ideais", expanded=False):
            st.markdown("""
| Parâmetro | Faixa ideal |
|---|---|
| pH | 7,2 – 7,8 |
| CRL mg/L | 0,5 – 3,0 |
| Alcalinidade mg/L | 80 – 120 |
| Dureza DC mg/L | 150 – 300 |
| CYA mg/L | 30 – 50 |
            """)

        # ── Parâmetros por piscina ────────────────────────────────────────────
        op_piscinas_dados = []
        _slug_map = {"Piscina Adulto":"adulto","Piscina Infantil":"infantil","Piscina Family":"family"}

        for pisc_nome in piscinas_ativas:
            pisc_slug = _slug_map.get(pisc_nome, slugify_nome(pisc_nome)[:12])
            st.markdown('<div class="op-card">', unsafe_allow_html=True)
            st.markdown(f'<div class="op-title">🧪 {pisc_nome}</div>', unsafe_allow_html=True)

            c1, c2 = st.columns(2)
            with c1:
                p_ph  = _num_op(f"op_{pisc_slug}_ph",  "pH", "ex: 7.4")
                p_alc = _num_op(f"op_{pisc_slug}_alc", "Alcalinidade mg/L", "ex: 95",  quinzenal=True)
                p_dc  = _num_op(f"op_{pisc_slug}_dc",  "Dureza DC mg/L",   "ex: 200", quinzenal=True)
            with c2:
                p_crl = _num_op(f"op_{pisc_slug}_crl", "CRL mg/L",            "ex: 1.5")
                p_ct  = _num_op(f"op_{pisc_slug}_ct",  "Cloro Total CT mg/L", "ex: 1.8")
                p_cya = _num_op(f"op_{pisc_slug}_cya", "CYA mg/L",            "ex: 40",  quinzenal=True)

            # Alertas em tempo real
            for val, mn, mx, rot in [
                (p_ph, 7.2, 7.8, "pH"), (p_crl, 0.5, 3.0, "CRL"),
                (p_alc, 80, 120, "Alcalinidade"), (p_dc, 150, 300, "Dureza DC"),
                (p_cya, 30, 50, "CYA"),
            ]:
                v = valor_float(val)
                if v is not None:
                    st.markdown(f"{'⚠️' if v < mn or v > mx else '✅'} **{rot}: {v}** {'— fora da faixa' if v < mn or v > mx else '— conforme'}")

            p_cloraminas = None
            v_crl2 = valor_float(p_crl); v_ct2 = valor_float(p_ct)
            if v_crl2 is not None and v_ct2 is not None:
                p_cloraminas = round(max(v_ct2 - v_crl2, 0), 2)
                st.markdown(f"{'⚠️' if p_cloraminas > 0.2 else '✅'} **Cloraminas: {p_cloraminas} mg/L**")

            op_piscinas_dados.append({
                "nome": pisc_nome,
                "ph": p_ph, "cloro_livre": p_crl, "cloro_total": p_ct,
                "cloraminas": str(p_cloraminas) if p_cloraminas is not None else "",
                "alcalinidade": p_alc, "dureza": p_dc, "cianurico": p_cya,
            })

            # ── Sugestões de dosagem em tempo real ───────────────────────────
            _v_ph  = valor_float(p_ph)
            _v_crl = valor_float(p_crl)
            _v_alc = valor_float(p_alc)
            _v_dc  = valor_float(p_dc)
            _v_cya = valor_float(p_cya)
            _tem_params = any(v is not None for v in [_v_ph, _v_crl, _v_alc, _v_dc, _v_cya])

            if _tem_params:
                # Busca volume m³ do condomínio no Sheets
                _vol_m3 = st.session_state.get(f"_vol_m3_{slugify_nome(op_nome_cond.strip())}", 0.0)
                if not _vol_m3:
                    try:
                        _clientes_vol = sheets_listar_clientes_completo()
                        for _cv in _clientes_vol:
                            if _cv["nome"].lower().strip() == op_nome_cond.strip().lower():
                                # Busca volume da planilha (col D = Volume_m3)
                                _sh_vol = conectar_sheets()
                                if _sh_vol:
                                    _aba_vol = obter_aba_sheets("👥 Clientes")
                                    _rows_vol = _aba_vol.get_all_values()
                                    for _rv in _rows_vol:
                                        if len(_rv) > 3 and _cv["nome"].lower() in str(_rv[2]).lower():
                                            try:
                                                _vol_m3 = float(str(_rv[3]).replace(",",".").strip() or 0)
                                            except Exception:
                                                _vol_m3 = 0.0
                                            break
                                break
                        st.session_state[f"_vol_m3_{slugify_nome(op_nome_cond.strip())}"] = _vol_m3
                    except Exception:
                        _vol_m3 = 0.0

                # Usa volume específico da piscina se disponível
                _vol_pisc = 0.0
                _slug_map2 = {"Piscina Adulto":"vol_adulto","Piscina Infantil":"vol_infantil","Piscina Family":"vol_family"}
                _vol_key = _slug_map2.get(pisc_nome, "")
                try:
                    _clv = sheets_listar_clientes_completo()
                    for _cv2 in _clv:
                        if nomes_condominio_equivalentes(_cv2["nome"], op_nome_cond):
                            if _vol_key:
                                # Piscina padrão (adulto/infantil/family)
                                _vol_pisc = float(_cv2.get(_vol_key, 0) or 0)
                            else:
                                # Piscina extra — busca no JSON local
                                _pasta_extra_vol = GENERATED_DIR / slugify_nome(op_nome_cond.strip())
                                _dados_extra_vol = (carregar_dados_condominio(_pasta_extra_vol) or {}) if _pasta_extra_vol.exists() else {}
                                for _pe in _dados_extra_vol.get("piscinas_extras", []):
                                    if _pe.get("nome","").strip().lower() == pisc_nome.strip().lower():
                                        _vol_pisc = float(_pe.get("vol", 0) or 0)
                                        break
                            break
                except Exception:
                    pass
                _vol_usar = _vol_pisc if _vol_pisc > 0 else _vol_m3

                _sugestoes = []
                if _vol_usar > 0:
                    _sugestoes = calcular_sugestoes_dosagem(
                        ph=_v_ph, crl=_v_crl, alc=_v_alc, dc=_v_dc, cya=_v_cya,
                        volume_m3=_vol_usar
                    )
                if _sugestoes:
                    st.markdown(f"**💊 Sugestões para {pisc_nome} ({_vol_usar:.0f} m³):**")
                    for _s in _sugestoes:
                        _icon = "🔴" if _s["prioridade"] == 1 else ("🟡" if _s["prioridade"] == 2 else "🔵")
                        if _s["quantidade"] and _s["quantidade"] > 0:
                            st.markdown(f"{_icon} **{_s['produto']}** — **{_s['quantidade']} {_s['unidade']}**")
                            st.caption(f"↳ {_s['acao']}")
                        else:
                            st.markdown(f"{_icon} **{_s['produto']}** — {_s['acao']}")
                        with st.expander("ℹ️ Base técnica", expanded=False):
                            st.caption(_s["justificativa"])
                            st.caption(f"📚 {_s.get('norma','')}")
                    # Botão aplicar sugestões nas dosagens desta piscina
                    if st.button(f"✅ Aplicar sugestões de {pisc_nome}",
                            key=f"btn_aplicar_sug_{pisc_slug}",
                            use_container_width=True):
                        # Salva sugestões no session_state para preencher dosagens
                        _key_sug = f"_sug_pisc_{pisc_slug}"
                        st.session_state[_key_sug] = [
                            s for s in _sugestoes
                            if s.get("quantidade") and s["quantidade"] > 0
                        ]
                        st.success(f"✅ Sugestões aplicadas! Verifique as dosagens de {pisc_nome} abaixo.")
                        st.rerun()
                else:
                    st.success("✅ Todos os parâmetros dentro da faixa ideal.")

            st.markdown("</div>", unsafe_allow_html=True)

        # Compatibilidade com código legado (usa dados da primeira piscina)
        op_ph  = op_piscinas_dados[0]["ph"]        if op_piscinas_dados else ""
        op_crl = op_piscinas_dados[0]["cloro_livre"] if op_piscinas_dados else ""
        op_ct  = op_piscinas_dados[0]["cloro_total"] if op_piscinas_dados else ""
        op_alc = op_piscinas_dados[0]["alcalinidade"] if op_piscinas_dados else ""
        op_dc  = op_piscinas_dados[0]["dureza"]      if op_piscinas_dados else ""
        op_cya = op_piscinas_dados[0]["cianurico"]   if op_piscinas_dados else ""
        op_cloraminas = valor_float(op_piscinas_dados[0]["cloraminas"]) if op_piscinas_dados else None

        # ── Dosagens por piscina ──────────────────────────────────────────────
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        st.markdown('<div class="op-title">⚗️ Dosagens aplicadas</div>', unsafe_allow_html=True)
        st.caption("Registre os produtos aplicados em cada piscina. Use as sugestões calculadas acima.")
        op_dosagens = []  # lista global para compatibilidade (primeira piscina)

        for _pisc_d in op_piscinas_dados:
            _pn_d    = _pisc_d["nome"]
            _slug_d  = _slug_map.get(_pn_d, slugify_nome(_pn_d)[:12])
            _key_sug = f"_sug_pisc_{_slug_d}"

            # Se há sugestões aplicadas para esta piscina, pré-preenche os campos
            _sug_aplicadas = st.session_state.get(_key_sug, [])
            if _sug_aplicadas:
                for _idx_s, _s in enumerate(_sug_aplicadas[:5]):
                    st.session_state[f"op_dos_{_slug_d}_prod_{_idx_s}"] = _s.get("produto","")
                    st.session_state[f"op_dos_{_slug_d}_qtd_{_idx_s}"]  = str(round(_s.get("quantidade",0),1))
                    st.session_state[f"op_dos_{_slug_d}_un_{_idx_s}"]   = _s.get("unidade","")
                    st.session_state[f"op_dos_{_slug_d}_fin_{_idx_s}"]  = _s.get("acao","")
                st.session_state.pop(_key_sug, None)  # limpa após aplicar

            st.markdown(f"**🏊 {_pn_d}**")
            _dos_pisc = []
            for i in range(5):
                _k_prod = f"op_dos_{_slug_d}_prod_{i}"
                _k_qtd  = f"op_dos_{_slug_d}_qtd_{i}"
                _k_un   = f"op_dos_{_slug_d}_un_{i}"
                _k_fin  = f"op_dos_{_slug_d}_fin_{i}"
                with st.expander(
                    f"Produto {i+1}" + (f" — {st.session_state.get(_k_prod,'')}" if st.session_state.get(_k_prod) else ""),
                    expanded=(i == 0 or bool(st.session_state.get(_k_prod)))
                ):
                    _dd1, _dd2 = st.columns([2, 1])
                    _prod = _dd1.text_input("Produto", key=_k_prod,
                        label_visibility="collapsed", placeholder="Nome do produto",
                        on_change=_autosave_rascunho)
                    _qtd  = _dd2.text_input("Qtd", key=_k_qtd,
                        label_visibility="collapsed", placeholder="Qtd",
                        on_change=_autosave_rascunho)
                    _dd3, _dd4 = st.columns([1, 2])
                    _un   = _dd3.text_input("Un", key=_k_un,
                        label_visibility="collapsed", placeholder="kg/L/g",
                        on_change=_autosave_rascunho)
                    _fin  = _dd4.text_input("Finalidade", key=_k_fin,
                        label_visibility="collapsed", placeholder="Finalidade / motivo",
                        on_change=_autosave_rascunho)
                    if _prod.strip():
                        _dos_pisc.append({
                            "produto":    _prod.strip(),
                            "quantidade": _qtd.strip(),
                            "unidade":    _un.strip(),
                            "finalidade": _fin.strip(),
                        })

            # Armazena dosagens na piscina correspondente
            _pisc_d["dosagens"] = _dos_pisc
            # Compatibilidade legado (primeira piscina)
            if not op_dosagens and _dos_pisc:
                op_dosagens = _dos_pisc

        # Fallback: se só tem uma piscina, op_dosagens = dosagens dela
        if not op_dosagens and op_piscinas_dados:
            op_dosagens = op_piscinas_dados[0].get("dosagens", [])


        # ── Fotos categorizadas — com salvamento imediato ───────────────────
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        st.markdown('<div class="op-title">📸 Fotos da visita</div>', unsafe_allow_html=True)
        st.caption("Fotos são salvas automaticamente ao fazer upload — não somem se o sinal cair.")

        op_fotos_antes  = st.file_uploader("🔵 Antes do tratamento", type=["jpg","jpeg","png","webp","heic"], accept_multiple_files=True, key="op_fotos_antes")
        op_fotos_depois = st.file_uploader("🟢 Depois do tratamento", type=["jpg","jpeg","png","webp","heic"], accept_multiple_files=True, key="op_fotos_depois")
        op_fotos_cmaq   = st.file_uploader("🔧 Casa de máquinas", type=["jpg","jpeg","png","webp","heic"], accept_multiple_files=True, key="op_fotos_cmaq")
        op_fotos_extras = st.file_uploader("📋 Outras fotos (bordas, ralos, equipamentos...)", type=["jpg","jpeg","png","webp","heic"], accept_multiple_files=True, key="op_fotos_extras")

        # ── Salvamento imediato ao upload ─────────────────────────────────────
        # Pasta temporária de fotos do rascunho
        _pasta_fotos_rasc = (GENERATED_DIR / slugify_nome(op_nome_cond.strip()) / "fotos_rascunho") if op_nome_cond.strip() else None

        def _salvar_foto_imediato(lista_uploads, categoria):
            """v5: Salva fotos na pasta de rascunho. Falha não derruba o app."""
            if not _pasta_fotos_rasc or not lista_uploads:
                return []
            salvos = []
            try:
                _pasta_fotos_rasc.mkdir(parents=True, exist_ok=True)
                for foto in lista_uploads:
                    try:
                        _nome_f = f"rasc_{categoria}_{limpar_nome_arquivo(foto.name)}"
                        _path_f = _pasta_fotos_rasc / _nome_f
                        if not _path_f.exists():
                            with open(_path_f, "wb") as _ff:
                                _ff.write(foto.getbuffer())
                        salvos.append(_nome_f)
                    except Exception:
                        pass  # foto individual falhou — continua as demais
            except Exception:
                pass  # pasta inacessível — retorna lista vazia
            return salvos

        # Salva imediatamente ao fazer upload
        _fotos_rasc_antes  = _salvar_foto_imediato(op_fotos_antes,  "antes")
        _fotos_rasc_depois = _salvar_foto_imediato(op_fotos_depois, "depois")
        _fotos_rasc_cmaq   = _salvar_foto_imediato(op_fotos_cmaq,   "cmaq")
        op_fotos_extras = st.session_state.get("op_fotos_extras") or []
        _fotos_rasc_extras = _salvar_foto_imediato(op_fotos_extras, "extras")

        # Também mostra fotos já salvas do rascunho anterior (sessão anterior)
        _fotos_rasc_existentes = {"antes": [], "depois": [], "cmaq": [], "extras": []}
        if _pasta_fotos_rasc and _pasta_fotos_rasc.exists():
            for _fp in sorted(_pasta_fotos_rasc.glob("rasc_*")):
                for _cat in ["antes", "depois", "cmaq", "extras"]:
                    if f"rasc_{_cat}_" in _fp.name:
                        _fotos_rasc_existentes[_cat].append(_fp)

        # Preview — fotos do upload atual + fotos salvas do rascunho
        _todas_fotos_preview = [
            ("🔵 Antes", "antes", op_fotos_antes, _fotos_rasc_existentes["antes"]),
            ("🟢 Depois", "depois", op_fotos_depois, _fotos_rasc_existentes["depois"]),
            ("🔧 Casa máq.", "cmaq", op_fotos_cmaq, _fotos_rasc_existentes["cmaq"]),
            ("📋 Outras", "extras", op_fotos_extras, _fotos_rasc_existentes["extras"]),
        ]
        _total_fotos_rasc = sum(len(v) for v in _fotos_rasc_existentes.values())
        if _total_fotos_rasc > 0:
            st.caption(f"💾 {_total_fotos_rasc} foto(s) já salvas do rascunho anterior")

        for _cat_label, _cat_cod, _flist_up, _flist_rasc in _todas_fotos_preview:
            _nomes_up = {
                limpar_nome_arquivo(getattr(f, "name", ""))
                for f in (_flist_up or [])
                if getattr(f, "name", "")
            }

            def _nome_original_rascunho(_arquivo_rasc):
                _nome = limpar_nome_arquivo(getattr(_arquivo_rasc, "name", str(_arquivo_rasc)))
                _prefixo = f"rasc_{_cat_cod}_"
                if _nome.startswith(_prefixo):
                    return _nome[len(_prefixo):]
                return _nome

            _rasc_extras = [
                f for f in (_flist_rasc or [])
                if _nome_original_rascunho(f) not in _nomes_up
            ]
            _all_show = deduplicar_fotos(list(_flist_up or []) + _rasc_extras)
            if _all_show:
                st.caption(f"**{_cat_label}:** {len(_all_show)} foto(s)")
                _cols = st.columns(min(len(_all_show), 3))
                for _i, _f in enumerate(_all_show):
                    with _cols[_i % 3]:
                        st.image(carregar_imagem_corrigida_orientacao(_f), use_container_width=True)

        # Botão para limpar fotos do rascunho
        if _total_fotos_rasc > 0:
            if st.button("🗑 Limpar fotos do rascunho", key="btn_limpar_fotos_rasc"):
                import shutil
                if _pasta_fotos_rasc and _pasta_fotos_rasc.exists():
                    shutil.rmtree(str(_pasta_fotos_rasc))
                st.rerun()


        # ── Problemas / Ocorrências ───────────────────────────────────────────
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        st.markdown('<div class="op-title">⚠️ Problemas / Ocorrências</div>', unsafe_allow_html=True)
        op_problemas = st.text_area("Problemas", key="op_problemas", height=80,
            label_visibility="collapsed",
            placeholder="Ex.: Bomba com ruído, vazamento na casa de máquinas, pH instável, equipamento quebrado...",
            on_change=_autosave_rascunho)

        # ── Observação geral ──────────────────────────────────────────────────
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        st.markdown('<div class="op-title">📝 Observação geral</div>', unsafe_allow_html=True)
        op_obs = st.text_area("Obs", key="op_obs_campo", height=80,
            label_visibility="collapsed", placeholder="Ex.: condições gerais da água, recomendações...",
            on_change=_autosave_rascunho)

        # ── Responsável no local ──────────────────────────────────────────────
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        st.markdown('<div class="op-title">👤 Responsável no local</div>', unsafe_allow_html=True)
        op_resp_local = st.text_input("Responsável no local", key="op_resp_local",
            label_visibility="collapsed",
            placeholder="Nome de quem recebeu o técnico (síndico, porteiro, zelador...)",
            on_change=_autosave_rascunho)

        # ── Parecer da visita ─────────────────────────────────────────────────
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        st.markdown('<div class="op-title">✅ Parecer da visita</div>', unsafe_allow_html=True)
        op_parecer = st.radio(
            "Parecer",
            ["✅ Satisfatório", "⚠️ Aceitável com ajustes pontuais", "❌ Insatisfatório"],
            key="op_parecer_visita",
            label_visibility="collapsed",
            horizontal=True,
            on_change=_autosave_rascunho,
        )
        # ── Assinatura do responsável ─────────────────────────────────────────
        # Compatibilidade Streamlit 1.56: removido canvas HTML/iframe, que causava
        # rerun infinito. Mantemos o módulo com captura nativa por upload/câmera.
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        st.markdown('<div class="op-title">✍️ Assinatura do responsável</div>', unsafe_allow_html=True)
        st.caption(
            "Para evitar travamentos nesta versão do Streamlit, a assinatura é anexada por imagem "
            "(foto ou arquivo PNG/JPG) em vez do canvas desenhado na tela."
        )

        _canvas_nonce = st.session_state.get("_op_ass_canvas_nonce", 0)
        _ass_exist_b64 = _normalizar_assinatura_b64(st.session_state.get("op_assinatura_responsavel_b64", ""))

        _ass_col1, _ass_col2 = st.columns(2)
        with _ass_col1:
            _ass_upload = st.file_uploader(
                "Enviar imagem da assinatura",
                type=["png", "jpg", "jpeg"],
                key=f"upload_assinatura_resp_{_canvas_nonce}",
                help="Aceita foto tirada no celular ou imagem escaneada da assinatura.",
            )
        with _ass_col2:
            _ass_camera = st.camera_input(
                "Fotografar assinatura",
                key=f"camera_assinatura_resp_{_canvas_nonce}",
                help="Opcional: fotografe uma assinatura feita em papel.",
            )

        _ass_arquivo = _ass_upload or _ass_camera
        if _ass_arquivo is not None:
            try:
                import base64 as _b64_ass
                import io as _io_ass
                _img_ass = Image.open(_io_ass.BytesIO(_ass_arquivo.getvalue()))
                _img_ass = ImageOps.exif_transpose(_img_ass).convert("RGBA")
                _buf_ass = _io_ass.BytesIO()
                _img_ass.save(_buf_ass, format="PNG")
                _ass_novo_b64 = _normalizar_assinatura_b64(_b64_ass.b64encode(_buf_ass.getvalue()).decode("utf-8"))
                if _ass_novo_b64 and _ass_novo_b64 != _ass_exist_b64:
                    st.session_state["op_assinatura_responsavel_b64"] = _ass_novo_b64
                    st.session_state["op_assinatura_responsavel_data"] = (
                        st.session_state.get("op_data_visita", "") or hoje_br()
                    )
                    _ass_exist_b64 = _ass_novo_b64
                    _autosave_rascunho()
                    st.success("✅ Assinatura anexada com sucesso.")
            except Exception as e:
                st.error(f"Não foi possível processar a assinatura enviada: {e}")

        _acol1, _acol2 = st.columns([1, 1])
        with _acol1:
            if _ass_exist_b64:
                _resp_ass = (st.session_state.get("op_resp_local") or "").strip()
                st.success("✅ Assinatura capturada" + (f" — {_resp_ass}" if _resp_ass else ""))
            else:
                st.caption("Assine no quadro acima e toque em ✅ Confirmar assinatura.")
        with _acol2:
            if st.button("🧹 Limpar assinatura", key="btn_limpar_assinatura", use_container_width=True):
                st.session_state["op_assinatura_responsavel_b64"] = ""
                st.session_state["op_assinatura_responsavel_data"] = ""
                st.session_state["_op_ass_canvas_nonce"] = _canvas_nonce + 1
                _autosave_rascunho()
                st.rerun()

        if _ass_exist_b64:
            try:
                import base64 as _b64_ui
                _resp_leg = (st.session_state.get("op_resp_local") or "").strip()
                st.image(
                    _b64_ui.b64decode(_ass_exist_b64),
                    caption=f"Assinatura do responsável{(' — ' + _resp_leg) if _resp_leg else ''}",
                    use_container_width=True,
                )
            except Exception:
                pass

        # ── Botão salvar rascunho ─────────────────────────────────────────────
        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        _rc_col1, _rc_col2 = st.columns([1, 1])
        with _rc_col1:
            if st.button("📋 Salvar rascunho", key="btn_salvar_rascunho",
                    use_container_width=True,
                    help="Salva o progresso atual. Você pode continuar depois."):
                if op_nome_cond.strip():
                    _dados_rasc = coletar_rascunho_operador(op_nome_cond, piscinas_ativas)
                    # v5: salvar_sheets=True apenas no botão explícito (nunca em on_change)
                    if salvar_rascunho_operador(op_nome_cond, _dados_rasc, salvar_sheets=True):
                        st.success(f"✅ Rascunho salvo! Pode fechar e retomar depois.")
                    else:
                        st.error("Erro ao salvar rascunho.")
                else:
                    st.warning("Selecione o condomínio antes de salvar o rascunho.")
        with _rc_col2:
            _rasc_exists = (GENERATED_DIR / slugify_nome(op_nome_cond.strip()) / "_rascunho_operador.json").exists() if op_nome_cond.strip() else False
            if _rasc_exists:
                st.caption("📋 Rascunho salvo")

        st.markdown('<div class="op-card">', unsafe_allow_html=True)
        if st.button("💾 Salvar lançamento", type="primary", use_container_width=True):
            data_vis = (st.session_state.get("op_data_visita") or "").strip()
            if not op_nome_cond.strip():
                st.error("Informe o condomínio.")
            elif not data_vis:
                st.error("Informe a data da visita.")
            else:
                pasta_op = GENERATED_DIR / slugify_nome(op_nome_cond.strip())
                pasta_op.mkdir(parents=True, exist_ok=True)
                pasta_fotos_op = pasta_op / "fotos_campo"
                pasta_fotos_op.mkdir(exist_ok=True)
                ts_f = datetime.now().strftime("%Y%m%d_%H%M%S")
                mes_ano = datetime.now().strftime("%Y-%m")

                import base64 as _b64

                def _salvar_categoria(lista_uploads, categoria):
                    """Salva fotos localmente, no Drive e retorna bytes para PDF."""
                    nomes = []; ids = []; b64s = []
                    for idx_ff, foto_ff in enumerate(lista_uploads or [], 1):
                        nome_ff = limpar_nome_arquivo(f"{categoria}_{ts_f}_{idx_ff}_{foto_ff.name}")
                        foto_bytes = bytes(foto_ff.getbuffer())
                        # Salva localmente
                        with open(pasta_fotos_op / nome_ff, "wb") as ff:
                            ff.write(foto_bytes)
                        nomes.append(nome_ff)
                        # Guarda base64 para PDF — com rotação EXIF automática
                        try:
                            from PIL import Image as _PILImg, ImageOps as _IOps
                            import io as _io
                            _img = _PILImg.open(_io.BytesIO(foto_bytes))
                            # Corrige rotação EXIF (fotos tiradas na vertical pelo celular)
                            _img = _IOps.exif_transpose(_img)
                            # Reduz para não pesar mas mantém qualidade
                            _img.thumbnail((1200, 1200), _PILImg.LANCZOS)
                            _buf = _io.BytesIO()
                            _img.convert("RGB").save(_buf, format="JPEG", quality=82)
                            b64s.append(_b64.b64encode(_buf.getvalue()).decode("utf-8"))
                        except Exception:
                            b64s.append(_b64.b64encode(foto_bytes).decode("utf-8"))
                        # Upload para o Google Drive (async best-effort)
                        try:
                            drive_id = drive_upload_foto(
                                arquivo_bytes=foto_bytes,
                                nome_arquivo=f"{categoria}_{nome_ff}",
                                nome_condominio=op_nome_cond.strip(),
                                mes_ano=mes_ano,
                            )
                            if drive_id:
                                ids.append(drive_id)
                        except Exception:
                            pass
                    return nomes, ids, b64s

                fotos_antes_nomes,  fotos_antes_ids,  fotos_antes_b64  = _salvar_categoria(deduplicar_fotos(op_fotos_antes),  "antes")
                fotos_depois_nomes, fotos_depois_ids, fotos_depois_b64 = _salvar_categoria(deduplicar_fotos(op_fotos_depois), "depois")
                fotos_cmaq_nomes,   fotos_cmaq_ids,   fotos_cmaq_b64   = _salvar_categoria(deduplicar_fotos(op_fotos_cmaq),   "cmaq")
                op_fotos_extras_final = st.session_state.get("op_fotos_extras") or []
                fotos_extras_nomes, fotos_extras_ids, fotos_extras_b64 = _salvar_categoria(deduplicar_fotos(op_fotos_extras_final), "extras")

                # ── Incorporar fotos salvas do rascunho (sessão anterior) ──────
                # Corrigido: quando a foto estava no upload atual e também no rascunho,
                # ela era adicionada novamente no PDF. Agora comparamos pelo nome original
                # e também pelo conteúdo (hash md5) antes de anexar.
                _pasta_rasc_f = pasta_op / "fotos_rascunho"
                if _pasta_rasc_f.exists():
                    import shutil as _shutil
                    import hashlib as _hashlib

                    def _nome_original_upload(_nome: str, _cat: str) -> str:
                        _n = limpar_nome_arquivo(str(_nome or ""))
                        _prefixo_rasc = f"rasc_{_cat}_"
                        if _n.startswith(_prefixo_rasc):
                            _n = _n[len(_prefixo_rasc):]
                        _n = re.sub(rf"^{_cat}_[0-9]{{8}}_[0-9]{{6}}_[0-9]+_", "", _n)
                        return _n

                    def _hash_arquivo(_path):
                        try:
                            return _hashlib.md5(Path(_path).read_bytes()).hexdigest()
                        except Exception:
                            return ""

                    _uploads_originais_por_cat = {
                        "antes":  {_nome_original_upload(getattr(f, "name", ""), "antes")  for f in (op_fotos_antes or [])},
                        "depois": {_nome_original_upload(getattr(f, "name", ""), "depois") for f in (op_fotos_depois or [])},
                        "cmaq":   {_nome_original_upload(getattr(f, "name", ""), "cmaq")   for f in (op_fotos_cmaq or [])},
                        "extras": {_nome_original_upload(getattr(f, "name", ""), "extras") for f in (op_fotos_extras_final or [])},
                    }

                    _hashes_ja_salvos = set()
                    for _foto_existente in pasta_fotos_op.glob("*"):
                        if _foto_existente.is_file():
                            _h = _hash_arquivo(_foto_existente)
                            if _h:
                                _hashes_ja_salvos.add(_h)

                    for _fp_rasc in sorted(_pasta_rasc_f.glob("rasc_*")):
                        if not _fp_rasc.is_file():
                            continue
                        for _cat, _nomes_cat, _b64s_cat in [
                            ("antes",  fotos_antes_nomes,  fotos_antes_b64),
                            ("depois", fotos_depois_nomes, fotos_depois_b64),
                            ("cmaq",   fotos_cmaq_nomes,   fotos_cmaq_b64),
                            ("extras", fotos_extras_nomes, fotos_extras_b64),
                        ]:
                            if f"rasc_{_cat}_" not in _fp_rasc.name:
                                continue
                            _orig_rasc = _nome_original_upload(_fp_rasc.name, _cat)
                            if _orig_rasc and _orig_rasc in _uploads_originais_por_cat.get(_cat, set()):
                                continue
                            _hash_rasc = _hash_arquivo(_fp_rasc)
                            if _hash_rasc and _hash_rasc in _hashes_ja_salvos:
                                continue
                            _nome_dest = _fp_rasc.name.replace(f"rasc_{_cat}_", f"{_cat}_{ts_f}_")
                            _dest = pasta_fotos_op / _nome_dest
                            if not _dest.exists():
                                _shutil.copy2(str(_fp_rasc), str(_dest))
                            if _hash_rasc:
                                _hashes_ja_salvos.add(_hash_rasc)
                            if _nome_dest not in _nomes_cat:
                                _nomes_cat.append(_nome_dest)
                                try:
                                    from PIL import Image as _PI, ImageOps as _IO
                                    import io as _io2
                                    _img2 = _PI.open(str(_dest))
                                    _img2 = _IO.exif_transpose(_img2)
                                    _img2.thumbnail((1200, 1200), _PI.LANCZOS)
                                    _buf2 = _io2.BytesIO()
                                    _img2.convert("RGB").save(_buf2, format="JPEG", quality=82)
                                    _b64s_cat.append(_b64.b64encode(_buf2.getvalue()).decode())
                                except Exception:
                                    pass
                    _shutil.rmtree(str(_pasta_rasc_f), ignore_errors=True)

                fotos_salvas_op = fotos_antes_nomes + fotos_depois_nomes + fotos_cmaq_nomes + fotos_extras_nomes
                fotos_drive_ids = fotos_antes_ids   + fotos_depois_ids   + fotos_cmaq_ids + fotos_extras_ids

                assinatura_responsavel_b64 = _normalizar_assinatura_b64(st.session_state.get("op_assinatura_responsavel_b64", ""))
                assinatura_responsavel_arquivo = ""
                if assinatura_responsavel_b64:
                    pasta_assinaturas = pasta_op / "assinaturas_visita"
                    pasta_assinaturas.mkdir(exist_ok=True)
                    assinatura_responsavel_arquivo = f"assinatura_responsavel_{ts_f}.png"
                    _salvar_assinatura_local(assinatura_responsavel_b64, pasta_assinaturas / assinatura_responsavel_arquivo)

                dados_ex = carregar_dados_condominio(pasta_op) or {}
                lancamento = {
                    "data": data_vis, "operador": op_operador.strip(),
                    "ph": op_ph, "cloro_livre": op_crl, "cloro_total": op_ct,
                    "cloraminas": str(op_cloraminas) if op_cloraminas is not None else "",
                    "alcalinidade": op_alc, "dureza": op_dc, "cianurico": op_cya,
                    "piscinas": op_piscinas_dados,
                    "problemas": op_problemas.strip(),
                    "servicos_executados": [
                        s for s, k in [
                            ("Aspiração de fundo",             "op_serv_aspiracao"),
                            ("Escovação de paredes e bordas",  "op_serv_escovacao"),
                            ("Peneiração / retirada de resíduos", "op_serv_peneiracao"),
                            ("Limpeza de skimmer e pré-filtro","op_serv_skimmer"),
                            ("Limpeza de borda (azulejo/deck)","op_serv_borda"),
                            ("Retrolavagem do filtro",         "op_serv_retrolavagem"),
                            ("Aplicação de produtos químicos", "op_serv_dosagem"),
                            ("Verificação de equipamentos",    "op_serv_verificacao"),
                            ("Circulação e filtração",         "op_serv_circulacao"),
                        ] if st.session_state.get(k)
                    ] + (
                        [st.session_state.get("op_serv_outro_desc","").strip()]
                        if st.session_state.get("op_serv_outro") and st.session_state.get("op_serv_outro_desc","").strip()
                        else []
                    ),
                    "observacao": op_obs.strip(), "dosagens": op_dosagens,
                    "resp_local": (st.session_state.get("op_resp_local") or "").strip(),
                    "parecer": st.session_state.get("op_parecer_visita", "✅ Satisfatório"),
                    "assinatura_responsavel_b64": assinatura_responsavel_b64,
                    "assinatura_responsavel_nome": (st.session_state.get("op_resp_local") or "").strip(),
                    "assinatura_responsavel_data": data_vis,
                    "assinatura_responsavel_arquivo": assinatura_responsavel_arquivo,
                    "fotos": fotos_salvas_op,
                    "fotos_antes": fotos_antes_nomes,
                    "fotos_depois": fotos_depois_nomes,
                    "fotos_cmaq": fotos_cmaq_nomes,
                    "fotos_drive_ids": fotos_drive_ids,
                    "fotos_antes_ids": fotos_antes_ids,
                    "fotos_depois_ids": fotos_depois_ids,
                    "fotos_cmaq_ids": fotos_cmaq_ids,
                    "fotos_antes_b64": fotos_antes_b64,
                    "fotos_depois_b64": fotos_depois_b64,
                    "fotos_cmaq_b64": fotos_cmaq_b64,
                    "fotos_extras": fotos_extras_nomes,
                    "fotos_extras_ids": fotos_extras_ids,
                    "fotos_extras_b64": fotos_extras_b64,
                    "condominio": op_nome_cond.strip(),
                    "salvo_em": _agora_brasilia(),
                }
                pendentes = dados_ex.get("lancamentos_campo", [])
                pendentes.append(lancamento)
                dados_ex["lancamentos_campo"] = pendentes
                dados_ex["nome_condominio"] = dados_ex.get("nome_condominio", op_nome_cond.strip())
                if op_dosagens:
                    dados_ex["dosagens_ultimas"] = (op_dosagens + [{"produto":"","fabricante_lote":"","quantidade":"","unidade":"","finalidade":""}]*7)[:7]
                salvar_dados_condominio(pasta_op, dados_ex)

                # v5: Sheets com proteção total — falha não derruba o app nem limpa sessão
                try:
                    ok_sheets = sheets_salvar_lancamento_campo(lancamento, op_nome_cond.strip())
                except Exception as _e_sh:
                    ok_sheets = False
                    st.session_state["_sheets_ultimo_erro"] = str(_e_sh)

                if ok_sheets:
                    st.success("✅ Visita salva no Google Sheets e pronta para entrar no relatório mensal.")
                else:
                    erro_sh = st.session_state.get("_sheets_ultimo_erro", "")
                    st.warning(
                        "⚠️ A visita foi salva localmente e o PDF pode ser gerado normalmente. "
                        "Houve falha ao gravar no Google Sheets — o relatório mensal pode não importar automaticamente."
                    )
                    if erro_sh:
                        with st.expander("Detalhes do erro Sheets", expanded=False):
                            st.code(erro_sh[:1500])
                st.session_state["op_salvo_sucesso"] = {
                    "nome": op_nome_cond, "data": data_vis,
                    "operador": op_operador.strip() or "Não informado",
                    "total": len(pendentes),
                }
                # Guarda último lançamento para gerar relatório
                st.session_state["_op_ultimo_lancamento"] = lancamento
                # Sinaliza limpeza para o próximo rerun — não toca nos widgets agora
                st.session_state["op_limpar_campos"] = True
                # Remove rascunho após salvar lançamento definitivo
                deletar_rascunho_operador(op_nome_cond)
                st.rerun()

        pasta_hc = GENERATED_DIR / slugify_nome(op_nome_cond.strip()) if op_nome_cond.strip() else None
        if pasta_hc and pasta_hc.exists():
            dados_hc = carregar_dados_condominio(pasta_hc)
            pend_hc = (dados_hc or {}).get("lancamentos_campo", [])
            if pend_hc:
                st.markdown(f"**{len(pend_hc)} lançamento(s) registrado(s):**")
                for lc in reversed(pend_hc[-3:]):
                    ft = f" | 📸 {len(lc.get('fotos',[]))} foto(s)" if lc.get("fotos") else ""
                    st.caption(f"📅 {lc.get('data','')} | {lc.get('operador','–')} | pH:{lc.get('ph','–')} CRL:{lc.get('cloro_livre','–')}{ft}")

    # Para o restante da página não renderizar no modo operador
    st.stop()

if modo == "Modo Campo":
    st.info("Fluxo compacto habilitado para operação em campo no Windows / tablet Windows.")

if st.session_state.origem_dados_carregados:
    st.markdown(
        f"""
        <div class="quick-mode-box">
            <strong>Dados carregados:</strong> {st.session_state.origem_dados_carregados}<br>
            Agora você pode ajustar valor do aditivo, vigência ou renovar automaticamente.
        </div>
        """,
        unsafe_allow_html=True,
    )

if st.session_state.painel_acao_msg:
    st.success(st.session_state.painel_acao_msg)
    st.session_state.painel_acao_msg = ""

status_formulario = status_vencimento(
    st.session_state.get("data_fim", ""),
    st.session_state.alerta_vencimento_dias,
)

if st.session_state.get("data_fim", "").strip():
    if status_formulario["codigo"] == "vencido":
        st.markdown(
            f"<div class='alert-vencido'><strong>Alerta de vencimento:</strong> {status_formulario['mensagem']}</div>",
            unsafe_allow_html=True,
        )
    elif status_formulario["codigo"] == "vencendo":
        st.markdown(
            f"<div class='alert-vencendo'><strong>Atenção:</strong> {status_formulario['mensagem']}</div>",
            unsafe_allow_html=True,
        )
    elif status_formulario["codigo"] == "vigente":
        st.markdown(
            f"<div class='alert-vigente'><strong>Status:</strong> {status_formulario['mensagem']}</div>",
            unsafe_allow_html=True,
        )

# =========================================
# DADOS DE PAINEL
# =========================================

def obter_metricas_bem_star():
    """Coleta métricas para o Dashboard Bem Star.

    Usa leitura única da aba de visitas (sheets_listar_todas_visitas) para
    evitar loop por cliente que causava travamento e erro 429 de quota.
    """
    try:
        todos_clientes = sheets_listar_clientes_completo()
        clientes_bs = filtrar_clientes_por_empresa(todos_clientes, "bem_star")
        total_ativos = len([c for c in clientes_bs if str(c.get("status", "Ativo")).strip().lower() != "inativo"])
        nomes_bs = {str(c.get("nome", "")).strip().lower() for c in clientes_bs if c.get("nome")}

        # Conta visitas do mês atual com leitura única da aba
        mes_atual = datetime.now().strftime("%m/%Y")
        visitas_mes = 0
        try:
            todas_visitas = sheets_listar_todas_visitas()
            for v in todas_visitas:
                cond = str(v.get("condominio", "")).strip().lower()
                mes_v = str(v.get("mes_ano", "") or "").strip()
                # Aceita tanto "MM/YYYY" no campo mes_ano quanto extrai da data
                if not mes_v:
                    data_v = str(v.get("data", ""))
                    if len(data_v) >= 7:
                        partes = data_v.split("/")
                        mes_v = f"{partes[1]}/{partes[2]}" if len(partes) == 3 else ""
                if mes_v == mes_atual and cond in nomes_bs:
                    visitas_mes += 1
        except Exception:
            visitas_mes = 0

        return {
            "total_ativos": total_ativos,
            "visitas_mes": visitas_mes,
            "ultimos_pareceres": [],
        }
    except Exception as e:
        _log_sheets_erro("obter_metricas_bem_star", e)
        return {"total_ativos": 0, "visitas_mes": 0, "ultimos_pareceres": []}


# v4 — carregamento cirúrgico por empresa:
# O painel de vencimentos é exclusivo da Aqua Gestão/RT e faz varredura de pastas.
# No login Bem Star ele atrasava ou interrompia a primeira pintura antes dos módulos
# operacionais da Bem Star. Para Bem Star, inicializamos listas vazias e deixamos
# os módulos próprios carregarem direto.
if st.session_state.get("empresa_ativa") == "bem_star":
    painel_vencimentos = []
    painel_filtrado = []
    total_monitorado = total_vencidos = total_vencendo = total_vigentes = total_indefinidos = total_com_json = 0
    itens_vencidos = []
    itens_vencendo = []
    itens_indefinidos = []
else:
    painel_vencimentos = listar_painel_vencimentos(st.session_state.alerta_vencimento_dias)
    painel_filtrado = filtrar_itens_painel(
        painel_vencimentos,
        st.session_state.busca_rapida,
        st.session_state.filtro_status_central,
    )

    total_monitorado = len(painel_vencimentos)
    total_vencidos = len([i for i in painel_vencimentos if i["status"]["codigo"] == "vencido"])
    total_vencendo = len([i for i in painel_vencimentos if i["status"]["codigo"] == "vencendo"])
    total_vigentes = len([i for i in painel_vencimentos if i["status"]["codigo"] == "vigente"])
    total_indefinidos = len([i for i in painel_vencimentos if i["status"]["codigo"] == "indefinido"])
    total_com_json = len([i for i in painel_vencimentos if i["tem_json"]])

    itens_vencidos = [i for i in painel_filtrado if i["status"]["codigo"] == "vencido"]
    itens_vencendo = [i for i in painel_filtrado if i["status"]["codigo"] == "vencendo"]
    itens_indefinidos = [i for i in painel_filtrado if i["status"]["codigo"] == "indefinido"]

# =========================================
# DASHBOARD EXECUTIVO / BEM STAR
# =========================================

if st.session_state.get("empresa_ativa") == "bem_star":
    # --- DASHBOARD BEM STAR ---
    metricas_bs = obter_metricas_bem_star()
    
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader("⭐ Painel Bem Star Piscinas")
    
    db1, db2, db3 = st.columns(3)
    with db1:
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Clientes Ativos</div><div class='dash-value'>{metricas_bs['total_ativos']}</div><div class='dash-sub'>Carteira Bem Star</div></div>",
            unsafe_allow_html=True,
        )
    with db2:
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Visitas no Mês</div><div class='dash-value'>{metricas_bs['visitas_mes']}</div><div class='dash-sub'>Total de registros em {datetime.now().strftime('%m/%Y')}</div></div>",
            unsafe_allow_html=True,
        )
    with db3:
        # Média de visitas por cliente ativo (exemplo de métrica extra)
        media = (metricas_bs['visitas_mes'] / metricas_bs['total_ativos']) if metricas_bs['total_ativos'] > 0 else 0
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Média Visitas/Cli</div><div class='dash-value'>{media:.1f}</div><div class='dash-sub'>Frequência mensal média</div></div>",
            unsafe_allow_html=True,
        )
    
    st.markdown("**Últimos pareceres técnicos**")
    if not metricas_bs['ultimos_pareceres']:
        st.info("Nenhum parecer técnico recente encontrado.")
    else:
        for p in metricas_bs['ultimos_pareceres']:
            st.markdown(f"- **{p['cliente']}** ({p['data']}): _{p['parecer']}_")
    
    st.markdown("</div>", unsafe_allow_html=True)

else:
    # --- DASHBOARD AQUA GESTÃO (ORIGINAL) ---
    taxa_estrutura = (total_com_json / total_monitorado * 100) if total_monitorado else 0
    criticos = [i for i in painel_vencimentos if i["status"]["codigo"] in ("vencido", "vencendo")][:5]

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader("Dashboard executivo")

    d1, d2, d3, d4, d5 = st.columns(5)
    with d1:
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Total monitorado</div><div class='dash-value'>{total_monitorado}</div><div class='dash-sub'>Pastas acompanhadas pelo sistema</div></div>",
            unsafe_allow_html=True,
        )
    with d2:
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Vigentes</div><div class='dash-value'>{total_vigentes}</div><div class='dash-sub'>Dentro da vigência regular</div></div>",
            unsafe_allow_html=True,
        )
    with d3:
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Vencem em breve</div><div class='dash-value'>{total_vencendo}</div><div class='dash-sub'>Dentro da faixa de alerta</div></div>",
            unsafe_allow_html=True,
        )
    with d4:
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Vencidos</div><div class='dash-value'>{total_vencidos}</div><div class='dash-sub'>Exigem ação imediata</div></div>",
            unsafe_allow_html=True,
        )
    with d5:
        st.markdown(
            f"<div class='dash-mini'><div class='dash-title'>Cadastros estruturados</div><div class='dash-value'>{taxa_estrutura:.0f}%</div><div class='dash-sub'>{total_com_json} de {total_monitorado} com JSON salvo</div></div>",
            unsafe_allow_html=True,
        )

    cx1, cx2 = st.columns([1.15, 1])
    with cx1:
        st.markdown("**Resumo crítico**")
        if not criticos:
            st.success("Nenhum condomínio em faixa crítica no momento.")
        else:
            for item in criticos:
                st.markdown(
                    f"- **{item['nome_exibicao']}** — {item['status']['rotulo']} — {item['status']['mensagem']}"
                )

    with cx2:
        st.markdown("**Leitura rápida**")
        st.write(f"Sem vigência válida: **{total_indefinidos}**")
        st.write(f"Com ação prioritária agora: **{total_vencidos + total_vencendo}**")
        st.write(f"Faixa de lembrete atual: **{st.session_state.alerta_vencimento_dias} dias**")

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================
# SAÚDE DO SISTEMA
# =========================================

diag = diagnostico_sistema()

st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.subheader("Saúde do sistema")

s1, s2, s3, s4, s5, s6 = st.columns(6)
with s1:
    st.markdown(f"Template do contrato<br><span class='{'health-ok' if diag['template_contrato_ok'] else 'health-no'}'>{'OK' if diag['template_contrato_ok'] else 'Ausente'}</span>", unsafe_allow_html=True)
    st.markdown(f"Template Bem Star<br><span class='{'health-ok' if diag['template_bem_star_ok'] else 'health-no'}'>{'OK' if diag['template_bem_star_ok'] else 'Ausente'}</span>", unsafe_allow_html=True)
with s2:
    st.markdown(f"Template do aditivo<br><span class='{'health-ok' if diag['template_aditivo_ok'] else 'health-no'}'>{'OK' if diag['template_aditivo_ok'] else 'Ausente'}</span>", unsafe_allow_html=True)
with s3:
    st.markdown(f"Template do relatório<br><span class='{'health-ok' if diag['template_relatorio_ok'] else 'health-no'}'>{'OK' if diag['template_relatorio_ok'] else 'Ausente'}</span>", unsafe_allow_html=True)
with s4:
    st.markdown(f"Pasta de documentos<br><span class='{'health-ok' if diag['generated_ok'] else 'health-no'}'>{'OK' if diag['generated_ok'] else 'Ausente'}</span>", unsafe_allow_html=True)
with s5:
    st.markdown(f"Logo institucional<br><span class='{'health-ok' if diag['logo_ok'] else 'health-no'}'>{'OK' if diag['logo_ok'] else 'Não localizada'}</span>", unsafe_allow_html=True)
with s6:
    st.markdown(f"Ambiente Windows<br><span class='{'health-ok' if diag['windows_ok'] else 'health-no'}'>{'OK' if diag['windows_ok'] else 'Fora do padrão'}</span>", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)

# =========================================
# BUSCA RÁPIDA
# =========================================

st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.subheader("Busca rápida profissional")

b1, b2, b3 = st.columns([2.2, 1.2, 1])
with b1:
    st.text_input(
        "Buscar por nome, CNPJ, síndico ou status",
        key="busca_rapida",
        placeholder="Ex.: terra nova, 12.345.678/0001-90, Marcelo, vencido...",
    )
with b2:
    st.selectbox(
        "Filtrar por status",
        options=["Todos", "Vigente", "Vence em breve", "Vencido", "Sem vigência válida"],
        key="filtro_status_central",
    )
with b3:
    st.metric("Resultado da busca", len(painel_filtrado))

st.markdown("</div>", unsafe_allow_html=True)

# =========================================
# GESTÃO DE OPERADORES
# =========================================

st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.subheader("👷 Gestão de Operadores")
st.caption("Tela administrativa segura para cadastrar, editar, ativar/inativar operadores e definir exatamente quais condomínios cada PIN pode acessar.")
st.info("🔐 PIN geral 2940 mantido como acesso mestre do sistema. Ele continua reservado e não pode ser usado no cadastro de operadores comuns.")

_col_op_top1, _col_op_top2 = st.columns([1, 1.35])
with _col_op_top1:
    if st.button("🔧 Inicializar aba de Operadores no Sheets", key="btn_init_op_aba"):
        ok_aba = sheets_criar_aba_operadores()
        if ok_aba:
            st.success("✅ Aba '👷 Operadores' criada/confirmada no Google Sheets!")
        else:
            st.error("❌ Erro ao criar aba. Verifique a conexão com o Sheets.")
with _col_op_top2:
    st.caption("Use este painel para centralizar o controle de PINs, status e permissões por condomínio, sem expor os PINs completos na tela.")

if st.session_state.get("_operadores_erro"):
    st.warning(st.session_state.get("_operadores_erro"))

@st.cache_data(ttl=30)
def _listar_ops():
    return sheets_listar_operadores()

ops_cadastrados = _listar_ops()

@st.cache_data(ttl=60)
def _clientes_para_painel_op():
    return sheets_listar_clientes_completo()

_todos_clientes_painel = _clientes_para_painel_op()
_nomes_todos_clientes = _condominios_organizar([c.get("nome", "") for c in _todos_clientes_painel if c.get("nome")])


def _mascarar_pin_admin(pin: str) -> str:
    _pin = str(pin or "").strip()
    if not _pin:
        return "—"
    if len(_pin) <= 2:
        return "*" * len(_pin)
    return f"{_pin[:2]}{'*' * (len(_pin) - 2)}"


def _op_tem_acesso_total(op: dict) -> bool:
    _conds = _condominios_organizar(op.get("condomínios", []))
    return op.get("acesso_total", False) or any(_normalizar_chave_acesso(c) == "todos" for c in _conds) or not _conds


def _resumo_acesso_admin(op: dict) -> str:
    _conds = _condominios_organizar(op.get("condomínios", []))
    if _op_tem_acesso_total(op):
        return f"Todos os {len(_nomes_todos_clientes)} condomínio(s) cadastrados"
    _exatos = _resolver_condominios_permitidos_exatos(_conds, _nomes_todos_clientes)
    _orfaos = [c for c in _conds if _normalizar_chave_acesso(c) not in {_normalizar_chave_acesso(x) for x in _exatos}]
    _extra = f" | {len(_orfaos)} não localizado(s) exatamente" if _orfaos else ""
    return f"{len(_exatos)} condomínio(s) vinculado(s){_extra}"


def _ordenar_ops_admin(lista: list[dict]) -> list[dict]:
    return sorted(
        lista,
        key=lambda op: (
            0 if op.get("ativo") else 1,
            _normalizar_chave_acesso(op.get("nome", "")),
            _normalizar_chave_acesso(op.get("pin", "")),
        ),
    )


def _filtrar_ops_admin(lista: list[dict], busca: str, status: str) -> list[dict]:
    busca_norm = _normalizar_chave_acesso(busca)
    resultado = []
    for op in lista:
        if status == "Ativos" and not op.get("ativo"):
            continue
        if status == "Inativos" and op.get("ativo"):
            continue
        if busca_norm:
            alvo = " | ".join([
                str(op.get("nome", "")),
                str(op.get("pin", "")),
                " | ".join(_condominios_organizar(op.get("condomínios", []))),
            ])
            if busca_norm not in _normalizar_chave_acesso(alvo):
                continue
        resultado.append(op)
    return _ordenar_ops_admin(resultado)


def _filtrar_condominios_por_busca(opcoes: list[str], busca: str, selecionados: list[str] | None = None) -> list[str]:
    busca_norm = _normalizar_chave_acesso(busca)
    selecionados_limpos = _condominios_organizar(selecionados or [])
    if busca_norm:
        filtrados = [c for c in opcoes if busca_norm in _normalizar_chave_acesso(c)]
    else:
        filtrados = list(opcoes)

    combinados = []
    vistos = set()
    for nome in filtrados + selecionados_limpos:
        chave = _normalizar_chave_acesso(nome)
        if chave and chave not in vistos:
            vistos.add(chave)
            combinados.append(nome)
    return combinados


def _gerar_exportacao_operadores_csv(lista_ops: list[dict]) -> str:
    import csv
    import io

    buf = io.StringIO()
    writer = csv.writer(buf, delimiter=";")
    writer.writerow([
        "Nome",
        "PIN_mascarado",
        "Status",
        "Acesso_total",
        "Qtd_condominios_permitidos",
        "Condominios_permitidos",
        "Condominios_localizados_exatamente",
        "Permissoes_nao_localizadas",
    ])

    for op in _ordenar_ops_admin(lista_ops):
        _conds = _condominios_organizar(op.get("condomínios", []))
        _acesso_total_csv = _op_tem_acesso_total(op)
        _exatos_csv = _resolver_condominios_permitidos_exatos(_conds, _nomes_todos_clientes)
        _exatos_set_csv = {_normalizar_chave_acesso(c) for c in _exatos_csv}
        _orfaos_csv = [c for c in _conds if _normalizar_chave_acesso(c) not in _exatos_set_csv and _normalizar_chave_acesso(c) != "todos"]
        writer.writerow([
            str(op.get("nome", "")).strip(),
            _mascarar_pin_admin(op.get("pin", "")),
            "Ativo" if op.get("ativo") else "Inativo",
            "Sim" if _acesso_total_csv else "Não",
            len(_exatos_csv) if not _acesso_total_csv else len(_nomes_todos_clientes),
            " | ".join(_conds if not _acesso_total_csv else ["TODOS"]),
            " | ".join(_exatos_csv),
            " | ".join(_orfaos_csv),
        ])
    return buf.getvalue()


def _filtrar_clientes_admin_por_empresa(clientes: list[dict], empresa_filtro: str) -> list[str]:
    """Filtra nomes de clientes por painel/serviço sem misturar empresas.

    Aqua Gestão: RT ativo ou empresa marcada como Aqua.
    Bem Star: limpeza/manutenção ativa ou empresa marcada como Bem Star.
    """
    empresa_filtro = str(empresa_filtro or "Todas").strip()
    nomes = []
    for c in clientes or []:
        nome = str(c.get("nome", "")).strip()
        if not nome:
            continue
        serv = _normalizar_servicos_cliente(c)
        emp = _normalizar_chave_acesso(str(c.get("empresa", "") or ""))
        is_bem = "bem star" in emp or "bemstar" in emp
        is_aqua = "aqua" in emp
        tem_rt = bool(serv.get("rt"))
        tem_limpeza = bool(serv.get("limpeza"))

        if empresa_filtro == "Todas":
            nomes.append(nome)
        elif empresa_filtro == "Aqua Gestão" and (tem_rt or is_aqua):
            nomes.append(nome)
        elif empresa_filtro == "Bem Star Piscinas" and (tem_limpeza or is_bem):
            nomes.append(nome)
        elif empresa_filtro == "Ambas" and tem_rt and tem_limpeza:
            nomes.append(nome)
    return _condominios_organizar(nomes)


_empresa_admin_nome = _empresa_ativa_nome()

def _operador_tem_vinculo_no_painel(op: dict, clientes: list[dict], empresa_nome: str | None = None) -> bool:
    """Confere se o operador tem ao menos um condomínio pertencente ao painel ativo."""
    try:
        empresa = str(empresa_nome or _empresa_ativa_nome() or "").strip()
        if empresa not in ("Aqua Gestão", "Bem Star Piscinas"):
            return True

        nomes_painel = _filtrar_clientes_admin_por_empresa(clientes, empresa)
        if not nomes_painel:
            return False

        if _op_tem_acesso_total(op):
            return True

        conds = _condominios_organizar(op.get("condomínios", []))
        for cond in conds:
            if any(nomes_condominio_equivalentes(cond, nome_painel) for nome_painel in nomes_painel):
                return True
        return False
    except Exception:
        return False


_total_ops = len(ops_cadastrados)
_total_ativos = sum(1 for op in ops_cadastrados if op.get("ativo"))
_total_inativos = _total_ops - _total_ativos
_total_restritos = sum(1 for op in ops_cadastrados if not _op_tem_acesso_total(op))

_mop1, _mop2, _mop3, _mop4 = st.columns(4)
with _mop1:
    st.metric("Operadores", _total_ops)
with _mop2:
    st.metric("Ativos", _total_ativos)
with _mop3:
    st.metric("Inativos", _total_inativos)
with _mop4:
    st.metric("Acesso restrito", _total_restritos)

if ops_cadastrados:
    _csv_operadores = _gerar_exportacao_operadores_csv(ops_cadastrados)
    _exp1, _exp2 = st.columns([1.2, 2])
    with _exp1:
        st.download_button(
            "📤 Exportar operadores e permissões",
            data=_csv_operadores.encode("utf-8-sig"),
            file_name=f"operadores_permissoes_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv",
            use_container_width=True,
            key="btn_exportar_operadores_csv",
        )
    with _exp2:
        st.caption("Exportação administrativa em CSV com status, PIN mascarado e condomínios permitidos/localizados para conferência.")

if not ops_cadastrados:
    st.info("Nenhum operador cadastrado ainda. Use a aba 'Cadastrar novo operador'. O PIN 2940 continua funcionando como acesso geral.")

_tab_ops1, _tab_ops2 = st.tabs(["🛡️ Administrar operadores", "➕ Cadastrar novo operador"])

with _tab_ops1:
    if not ops_cadastrados:
        st.caption("Assim que houver operadores cadastrados, esta tela permitirá editar PIN, status e permissões com mais segurança.")
    else:
        _flt1, _flt2 = st.columns([1.1, 1.9])
        with _flt1:
            _status_ops = st.selectbox(
                "Filtrar status",
                ["Todos", "Ativos", "Inativos"],
                key="ops_admin_status",
            )
        with _flt2:
            _busca_ops = st.text_input(
                "Buscar operador, PIN ou condomínio",
                key="ops_admin_busca",
                placeholder="Ex.: João, Terra Nova, 12",
            )

        _ops_visiveis = _filtrar_ops_admin(ops_cadastrados, _busca_ops, _status_ops)
        _empresa_admin_nome = _empresa_ativa_nome()
        _nomes_painel_admin = _filtrar_clientes_admin_por_empresa(_todos_clientes_painel, _empresa_admin_nome)
        _ops_visiveis = [
            op for op in _ops_visiveis
            if _operador_tem_vinculo_no_painel(op, _todos_clientes_painel, _empresa_admin_nome)
        ]
        st.caption(f"Painel ativo: {_empresa_admin_nome}. Exibindo apenas operadores com vínculo neste painel.")

        if not _ops_visiveis:
            st.warning("Nenhum operador encontrado com os filtros atuais.")
        else:
            _rotulos_ops = [
                f"{'🟢' if op.get('ativo') else '🔴'} {op.get('nome', 'Sem nome')} | PIN {_mascarar_pin_admin(op.get('pin', ''))} | {_resumo_acesso_admin(op)}"
                for op in _ops_visiveis
            ]
            _idx_op = st.selectbox(
                "Selecione o operador para administrar",
                options=list(range(len(_ops_visiveis))),
                format_func=lambda i: _rotulos_ops[i],
                key="ops_admin_selector",
            )
            _op_sel = _ops_visiveis[_idx_op]
            _op_nome_sel = _op_sel.get("nome", "")
            _op_pin_sel = str(_op_sel.get("pin", "")).strip()
            _op_conds_sel = _condominios_organizar(_op_sel.get("condomínios", []))
            _op_total_sel = _op_tem_acesso_total(_op_sel)
            _op_exatos_sel = _resolver_condominios_permitidos_exatos(_op_conds_sel, _nomes_painel_admin)
            _op_exatos_set = {_normalizar_chave_acesso(c) for c in _op_exatos_sel}
            _op_orfaos_sel = [c for c in _op_conds_sel if _normalizar_chave_acesso(c) not in _op_exatos_set and _normalizar_chave_acesso(c) != "todos"]

            st.markdown("### Painel do operador selecionado")
            _sum1, _sum2, _sum3 = st.columns([1.1, 1, 1.4])
            with _sum1:
                st.caption("Operador")
                st.markdown(f"**{_op_nome_sel}**")
                st.caption(f"PIN mascarado: {_mascarar_pin_admin(_op_pin_sel)}")
            with _sum2:
                st.caption("Status")
                st.markdown("**🟢 Ativo**" if _op_sel.get("ativo") else "**🔴 Inativo**")
                st.caption("PIN geral 2940 não aparece aqui por segurança.")
            with _sum3:
                st.caption("Escopo de acesso")
                if _op_total_sel:
                    st.markdown(f"**Acesso total** aos {len(_nomes_painel_admin)} condomínio(s) deste painel")
                else:
                    st.markdown(f"**{len(_op_exatos_sel)} condomínio(s)** com correspondência exata")
                    if _op_orfaos_sel:
                        st.caption(f"{len(_op_orfaos_sel)} permissão(ões) salva(s) sem cliente exato no cadastro atual")

            with st.expander("🔎 Ver permissões atuais deste operador", expanded=False):
                if _op_total_sel:
                    st.success("Este operador está com acesso total liberado.")
                else:
                    if _op_exatos_sel:
                        st.markdown("**Condomínios localizados exatamente no cadastro atual:**")
                        for _c in _op_exatos_sel:
                            st.caption(f"✅ {_c}")
                    else:
                        st.caption("Nenhum condomínio localizado exatamente no cadastro atual.")
                    if _op_orfaos_sel:
                        st.markdown("**Permissões salvas que não batem exatamente com o cadastro atual:**")
                        for _c in _op_orfaos_sel:
                            st.caption(f"⚠️ {_c}")

            with st.expander("📋 Duplicar permissões de outro operador", expanded=False):
                _ops_origem_dup = [
                    op for op in _ordenar_ops_admin(ops_cadastrados)
                    if _normalizar_chave_acesso(op.get("nome", "")) != _normalizar_chave_acesso(_op_nome_sel)
                ]
                if not _ops_origem_dup:
                    st.caption("Cadastre pelo menos mais um operador para habilitar a duplicação de permissões.")
                else:
                    _idx_dup = st.selectbox(
                        "Copiar permissões a partir de",
                        options=list(range(len(_ops_origem_dup))),
                        format_func=lambda i: f"{'🟢' if _ops_origem_dup[i].get('ativo') else '🔴'} {_ops_origem_dup[i].get('nome', 'Sem nome')} | {_resumo_acesso_admin(_ops_origem_dup[i])}",
                        key=f"dup_perm_origem_{_normalizar_chave_acesso(_op_nome_sel)}",
                    )
                    _op_origem_dup = _ops_origem_dup[_idx_dup]
                    st.caption("Serão copiadas apenas as permissões de acesso. O PIN e o status do operador selecionado serão preservados.")
                    if st.button(
                        "📥 Duplicar permissões para este operador",
                        key=f"btn_dup_perm_{_normalizar_chave_acesso(_op_nome_sel)}",
                        use_container_width=True,
                    ):
                        _conds_dup = ["TODOS"] if _op_tem_acesso_total(_op_origem_dup) else _condominios_organizar(_op_origem_dup.get("condomínios", []))
                        if sheets_salvar_operador(
                            nome=_op_nome_sel,
                            pin=_op_pin_sel,
                            condomínios=_conds_dup,
                            ativo=_op_sel.get("ativo", True),
                        ):
                            st.session_state.pop("_operadores_erro", None)
                            st.success(f"✅ Permissões de '{_op_origem_dup.get('nome', 'origem')}' copiadas para '{_op_nome_sel}'.")
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error(st.session_state.get("_operadores_erro") or "Erro ao duplicar permissões. Verifique a conexão com o Sheets.")

            _fe1, _fe2 = st.columns([1.1, 1.9])
            with _fe1:
                _empresa_conds_edit = _empresa_ativa_nome()
                st.text_input(
                    "Permissões deste painel",
                    value=_empresa_conds_edit,
                    disabled=True,
                    key=f"empresa_conds_edit_visivel_{_normalizar_chave_acesso(_op_nome_sel)}",
                    help="Para evitar mistura jurídica/visual, este formulário mostra apenas condomínios do painel administrativo ativo.",
                )
            with _fe2:
                _busca_conds_edit = st.text_input(
                    "Buscar condomínio dentro deste formulário",
                    key=f"busca_conds_edit_{_normalizar_chave_acesso(_op_nome_sel)}",
                    placeholder="Digite parte do nome para filtrar a lista de permissões",
                )
            _nomes_empresa_edit = _filtrar_clientes_admin_por_empresa(_todos_clientes_painel, _empresa_conds_edit)
            _opcoes_conds_edit = _filtrar_condominios_por_busca(_nomes_empresa_edit, _busca_conds_edit, _op_exatos_sel)
            _set_painel_edit = {_normalizar_chave_acesso(n) for n in _nomes_empresa_edit}
            _op_default_edit = [c for c in _op_exatos_sel if _normalizar_chave_acesso(c) in _set_painel_edit]
            if not _op_total_sel:
                st.caption(f"Exibindo {len(_opcoes_conds_edit)} de {len(_nomes_empresa_edit)} condomínio(s) para o filtro de empresa atual.")

            _key_conds_edit = f"conds_edit_{_normalizar_chave_acesso(_op_nome_sel)}"
            with st.container():
                st.markdown("#### Editar operador com segurança")
                _ed1, _ed2 = st.columns(2)
                with _ed1:
                    st.text_input("Nome do operador", value=_op_nome_sel, disabled=True, help="Para preservar o comportamento atual do sistema, o nome é tratado como chave de atualização.")
                    _editar_pin = st.checkbox(
                        "Redefinir PIN deste operador",
                        value=False,
                        key=f"edit_pin_toggle_{_normalizar_chave_acesso(_op_nome_sel)}",
                        help="O PIN atual não é exibido em tela. Marque apenas se quiser trocar o PIN deste operador.",
                    )
                    _novo_pin_edit = st.text_input(
                        "Novo PIN",
                        value="",
                        type="password",
                        disabled=not _editar_pin,
                        key=f"novo_pin_{_normalizar_chave_acesso(_op_nome_sel)}",
                        placeholder="Digite um novo PIN exclusivo",
                    )
                with _ed2:
                    _ativo_edit = st.checkbox(
                        "Operador ativo",
                        value=bool(_op_sel.get("ativo", True)),
                        key=f"ativo_edit_{_normalizar_chave_acesso(_op_nome_sel)}",
                    )
                    _acesso_total_edit = st.checkbox(
                        "Acesso total a todos os condomínios",
                        value=_op_total_sel,
                        key=f"total_edit_{_normalizar_chave_acesso(_op_nome_sel)}",
                    )

                if not _acesso_total_edit:
                    _conds_edit = st.multiselect(
                        "Condomínios permitidos para este PIN",
                        options=_opcoes_conds_edit,
                        default=_op_default_edit,
                        key=_key_conds_edit,
                        help="Seleção exata dos condomínios liberados para este operador.",
                    )
                    _selecionados_edit = _condominios_organizar(st.session_state.get(_key_conds_edit, _op_default_edit))
                    _ced1, _ced2 = st.columns([1.1, 1.4])
                    with _ced1:
                        st.caption(f"Selecionados agora: {len(_selecionados_edit)} condomínio(s).")
                    with _ced2:
                        if st.button(
                            "✅ Marcar todos os resultados da busca",
                            key=f"btn_marcar_busca_{_normalizar_chave_acesso(_op_nome_sel)}",
                            use_container_width=True,
                            disabled=not bool(_opcoes_conds_edit),
                        ):
                            st.session_state[_key_conds_edit] = _condominios_organizar(_selecionados_edit + _opcoes_conds_edit)
                            st.rerun()
                else:
                    st.caption("Com acesso total marcado, o operador continuará vendo todos os condomínios disponíveis.")
                    _conds_edit = ["TODOS"]
                    st.caption(f"Selecionados agora: {len(_nomes_todos_clientes)} condomínio(s) via acesso total.")

                _salvar_edit = st.button(
                    "💾 Salvar alterações do operador",
                    type="primary",
                    use_container_width=True,
                    key=f"btn_salvar_edit_{_normalizar_chave_acesso(_op_nome_sel)}",
                )

            if _salvar_edit:
                _pin_final_edit = _novo_pin_edit.strip() if _editar_pin else _op_pin_sel
                if _acesso_total_edit:
                    _conds_final_edit = ["TODOS"]
                else:
                    # Edita somente as permissões do painel ativo e preserva vínculos de outro painel.
                    _set_painel_salvar = {_normalizar_chave_acesso(n) for n in _nomes_empresa_edit}
                    _conds_outros_paineis = [
                        c for c in _op_conds_sel
                        if _normalizar_chave_acesso(c) not in _set_painel_salvar
                        and _normalizar_chave_acesso(c) != "todos"
                    ]
                    _conds_final_edit = _condominios_organizar(_conds_outros_paineis + _condominios_organizar(_conds_edit))

                if not _pin_final_edit or len(_pin_final_edit) < 4:
                    st.error("PIN deve ter pelo menos 4 caracteres.")
                elif _pin_final_edit == "2940":
                    st.error("O PIN 2940 é reservado para acesso geral. Escolha outro para este operador.")
                elif _pin_operador_em_uso(_pin_final_edit, nome_ignorar=_op_nome_sel):
                    st.error(f"O PIN {_pin_final_edit} já está em uso por outro operador.")
                elif not _acesso_total_edit and not _conds_final_edit:
                    st.error("Selecione ao menos um condomínio para este operador ou marque acesso total.")
                else:
                    with st.spinner("Salvando alterações do operador..."):
                        ok_edit = sheets_salvar_operador(
                            nome=_op_nome_sel,
                            pin=_pin_final_edit,
                            condomínios=_conds_final_edit,
                            ativo=_ativo_edit,
                        )
                    if ok_edit:
                        st.session_state.pop("_operadores_erro", None)
                        st.success(f"✅ Operador '{_op_nome_sel}' atualizado com sucesso.")
                        st.cache_data.clear()
                        st.rerun()
                    else:
                        st.error(st.session_state.get("_operadores_erro") or "Erro ao salvar alterações do operador. Verifique a conexão com o Sheets.")

            st.markdown("---")
            st.markdown("**Ações administrativas adicionais**")
            _rm_key = f"confirmar_remocao_{_normalizar_chave_acesso(_op_nome_sel)}"
            _rm1, _rm2, _rm3 = st.columns([1.2, 1.1, 1.1])
            with _rm1:
                if st.button("🗑 Solicitar remoção deste operador", key=f"btn_solicitar_rm_{_normalizar_chave_acesso(_op_nome_sel)}", use_container_width=True):
                    st.session_state[_rm_key] = True
                    st.rerun()
            if st.session_state.get(_rm_key):
                st.warning(f"Confirme a remoção do operador '{_op_nome_sel}'. Esta ação apaga o cadastro da aba 👷 Operadores.")
                with _rm2:
                    if st.button("✅ Confirmar remoção", key=f"btn_confirma_rm_{_normalizar_chave_acesso(_op_nome_sel)}", use_container_width=True):
                        if sheets_deletar_operador(_op_nome_sel):
                            st.session_state.pop(_rm_key, None)
                            st.session_state.pop("_operadores_erro", None)
                            st.success(f"Operador '{_op_nome_sel}' removido.")
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error("Erro ao remover operador. Verifique a conexão com o Sheets.")
                with _rm3:
                    if st.button("Cancelar remoção", key=f"btn_cancela_rm_{_normalizar_chave_acesso(_op_nome_sel)}", use_container_width=True):
                        st.session_state.pop(_rm_key, None)
                        st.rerun()

with _tab_ops2:
    _fn1, _fn2 = st.columns([1.1, 1.9])
    with _fn1:
        _empresa_conds_novo = _empresa_ativa_nome()
        st.text_input(
            "Permissões deste painel",
            value=_empresa_conds_novo,
            disabled=True,
            key=f"empresa_conds_novo_visivel_{_empresa_conds_novo}",
            help="Novo operador será vinculado somente a condomínios do painel administrativo ativo.",
        )
    with _fn2:
        _busca_conds_novo = st.text_input(
            "Buscar condomínio para o novo operador",
            key="busca_conds_novo",
            placeholder="Digite parte do nome para filtrar a lista de permissões",
        )
    _nomes_empresa_novo = _filtrar_clientes_admin_por_empresa(_todos_clientes_painel, _empresa_conds_novo)
    _opcoes_conds_novo = _filtrar_condominios_por_busca(
        _nomes_empresa_novo,
        _busca_conds_novo,
        st.session_state.get("op_novo_conds", []),
    )
    if _nomes_empresa_novo:
        st.caption(f"Exibindo {len(_opcoes_conds_novo)} de {len(_nomes_empresa_novo)} condomínio(s) para o filtro de empresa atual.")
    else:
        st.caption("Nenhum condomínio disponível para o filtro de empresa selecionado.")

    with st.container():
        st.markdown("#### Cadastro seguro de novo operador")
        _novo1, _novo2 = st.columns(2)
        with _novo1:
            op_nome_novo = st.text_input("Nome do operador *", key="op_novo_nome", placeholder="Ex.: João Silva")
            op_pin_novo = st.text_input(
                "PIN exclusivo *",
                key="op_novo_pin",
                placeholder="Ex.: 1234",
                max_chars=10,
                type="password",
                help="Mínimo 4 caracteres. Não use 2940, pois este PIN é reservado para o acesso geral do sistema.",
            )
        with _novo2:
            op_ativo_novo = st.checkbox("Operador ativo", value=True, key="op_novo_ativo")
            op_acesso_total_novo = st.checkbox("Acesso a todos os condomínios", value=False, key="op_novo_acesso_total")

        if not op_acesso_total_novo:
            op_conds_novo = st.multiselect(
                "Condomínios permitidos para este novo PIN",
                options=_opcoes_conds_novo,
                key="op_novo_conds",
                help="Seleção exata dos condomínios liberados para o novo operador.",
            )
            _selecionados_novo = _condominios_organizar(st.session_state.get("op_novo_conds", []))
            _cn1, _cn2 = st.columns([1.1, 1.4])
            with _cn1:
                st.caption(f"Selecionados agora: {len(_selecionados_novo)} condomínio(s).")
            with _cn2:
                if st.button(
                    "✅ Marcar todos os resultados da busca",
                    key="btn_marcar_busca_novo_operador",
                    use_container_width=True,
                    disabled=not bool(_opcoes_conds_novo),
                ):
                    st.session_state["op_novo_conds"] = _condominios_organizar(_selecionados_novo + _opcoes_conds_novo)
                    st.rerun()
        else:
            st.caption("Com acesso total marcado, o novo operador verá todos os condomínios disponíveis no sistema.")
            op_conds_novo = ["TODOS"]
            st.caption(f"Selecionados agora: {len(_nomes_todos_clientes)} condomínio(s) via acesso total.")

        _salvar_novo = st.button(
            "💾 Cadastrar operador",
            type="primary",
            use_container_width=True,
            key="btn_cadastrar_novo_operador_seguro",
        )

    if _salvar_novo:
        _nome_op_limpo = re.sub(r"\s+", " ", op_nome_novo.strip())
        _pin_op_limpo = op_pin_novo.strip()
        _conds_op_final = ["TODOS"] if op_acesso_total_novo else _condominios_organizar(op_conds_novo)

        if not _nome_op_limpo:
            st.error("Informe o nome do operador.")
        elif not _pin_op_limpo or len(_pin_op_limpo) < 4:
            st.error("PIN deve ter pelo menos 4 caracteres.")
        elif _pin_op_limpo == "2940":
            st.error("O PIN 2940 é reservado para acesso geral. Escolha outro.")
        elif _pin_operador_em_uso(_pin_op_limpo, nome_ignorar=_nome_op_limpo):
            st.error(f"O PIN {_pin_op_limpo} já está em uso por outro operador.")
        elif not op_acesso_total_novo and not _conds_op_final:
            st.error("Selecione ao menos um condomínio para este operador ou marque acesso total.")
        else:
            with st.spinner("Salvando operador..."):
                ok_op = sheets_salvar_operador(
                    nome=_nome_op_limpo,
                    pin=_pin_op_limpo,
                    condomínios=_conds_op_final,
                    ativo=op_ativo_novo,
                )
            if ok_op:
                st.session_state.pop("_operadores_erro", None)
                st.success(f"✅ Operador '{_nome_op_limpo}' cadastrado com sucesso.")
                st.cache_data.clear()
                st.rerun()
            else:
                st.error(st.session_state.get("_operadores_erro") or "❌ Erro ao salvar operador. Verifique a conexão com o Sheets.")
                with st.expander("🔍 Ver detalhe técnico do erro", expanded=False):
                    _det = st.session_state.get("_sheets_ultimo_erro", "Sem detalhes.")
                    st.code(_det, language="text")
                # _EXPANDER_ERRO_SHEETS_OK_

st.markdown("</div>", unsafe_allow_html=True)

# =========================================
# CADASTRO DE CLIENTES — GOOGLE SHEETS
# =========================================

st.markdown('<div class="section-card">', unsafe_allow_html=True)
_empresa_cadastro_nome = _empresa_ativa_nome()
st.subheader(f"👥 Cadastro de Clientes — {_empresa_cadastro_nome}")
if _empresa_ativa_codigo() == "aqua_gestao":
    st.caption("Neste painel aparecem somente clientes de RT/controle técnico da Aqua Gestão.")
else:
    st.caption("Neste painel aparecem somente clientes de limpeza/manutenção da Bem Star Piscinas.")

# ── Diagnóstico de conexão (visível para admin) ──────────────────────────────
with st.expander("🔌 Testar conexão com Google Sheets", expanded=False):
    if st.button("▶ Executar teste de conexão agora", key="btn_teste_sheets"):
        with st.spinner("Testando conexão..."):
            sh_teste = conectar_sheets()
        if sh_teste is not None:
            try:
                abas_disponiveis = [w.title for w in sh_teste.worksheets()]
                st.success("✅ Conexão estabelecida com sucesso!")
                st.write(f"**Abas encontradas na planilha:** {abas_disponiveis}")
                abas_esperadas = ["👥 Clientes", "🔬 Visitas"]
                faltando = [a for a in abas_esperadas if a not in abas_disponiveis]
                if faltando:
                    st.warning(
                        f"⚠️ Aba(s) esperada(s) não encontrada(s): {faltando}\n\n"
                        "Verifique se os nomes das abas na planilha estão escritos **exatamente** assim, "
                        "incluindo os emojis."
                    )
                else:
                    st.success("✅ Todas as abas necessárias foram encontradas.")
            except Exception as ex:
                st.warning(f"Conexão ok, mas erro ao listar abas: {ex}")
        else:
            st.error("❌ Falha na conexão com o Google Sheets.")
            erro_detalhado = st.session_state.get("_sheets_ultimo_erro", "Sem detalhes.")
            st.code(erro_detalhado, language="text")
            st.markdown(
                "**Próximos passos para resolver:**\n"
                "1. Confirme que `gspread` e `google-auth` estão no `requirements.txt`\n"
                "2. No Streamlit Cloud → Settings → Secrets → verifique se `[gcp_service_account]` está presente\n"
                "3. Force redeploy: no painel do Streamlit Cloud clique em **Reboot app**\n"
                "4. Verifique se a conta de serviço `aqua-gestao-sheets@aqua-gestao-rt.iam.gserviceaccount.com` tem acesso **Editor** à planilha"
            )
    st.caption("Use este botão sempre que a conexão com o Sheets falhar para identificar a causa exata.")

# Mostra clientes já cadastrados no Sheets
@st.cache_data(ttl=30)
def _clientes_cadastrados():
    return sheets_listar_clientes_completo()

clientes_cadastrados_todos = _clientes_cadastrados()
clientes_cadastrados = _filtrar_clientes_painel_ativo(clientes_cadastrados_todos)

if clientes_cadastrados:
    st.success(f"✅ {len(clientes_cadastrados)} cliente(s) cadastrado(s) para {_empresa_cadastro_nome} no Google Sheets:")
    for c in clientes_cadastrados:
        _serv = _normalizar_servicos_cliente(c)
        _badges = []
        if _serv.get("rt"):
            _badges.append("🔵 RT")
        if _serv.get("limpeza"):
            _badges.append("⭐ Limpeza")
        _ops = _normalizar_lista_textos_unicos(c.get("operadores_vinculados", []))
        _ops_txt = f" • Operadores: {', '.join(_ops)}" if _ops else ""
        st.caption(f"• {c.get('nome', '')} {' '.join(_badges)}{_ops_txt}")
else:
    st.info(f"Nenhum cliente cadastrado para {_empresa_cadastro_nome} ainda. Use o formulário abaixo para adicionar.")

# Processa flag de limpeza ANTES de renderizar os widgets
if st.session_state.pop("_cc_limpar", False):
    for k in ["cc_nome","cc_cnpj","cc_cep","cc_endereco","cc_contato","cc_telefone",
              "cc_vol_adulto","cc_vol_infantil","cc_vol_family",
              "cc_pisc_extra1_nome","cc_pisc_extra1_vol",
              "cc_pisc_extra2_nome","cc_pisc_extra2_vol"]:
        st.session_state[k] = ""
    # Campo numérico: nunca limpar com string vazia, pois isso pode quebrar o widget do Streamlit.
    st.session_state["cc_verificacoes_semanais"] = 3
    st.session_state["cc_srv_rt"] = False
    st.session_state["cc_srv_limpeza"] = False
    st.session_state["cc_operadores_vinculados"] = []

# ── Seletor de edição ────────────────────────────────────────────────────────
_cc_modo = st.radio("Ação", ["➕ Novo cliente", "✏️ Editar cliente existente"],
    key="cc_modo_acao", horizontal=True, label_visibility="collapsed")

_cc_cliente_editar = {}
if _cc_modo == "✏️ Editar cliente existente":
    @st.cache_data(ttl=30)
    def _clientes_completos_edit():
        return sheets_listar_clientes_completo()
    _clientes_edit = _filtrar_clientes_painel_ativo(_clientes_completos_edit())
    if _clientes_edit:
        _nomes_edit = [c["nome"] for c in _clientes_edit]
        _sel_edit = st.selectbox("Selecione o cliente para editar", _nomes_edit, key="cc_sel_editar")
        _cc_cliente_editar = next((c for c in _clientes_edit if c["nome"] == _sel_edit), {})
        if _cc_cliente_editar and st.button("📂 Carregar dados", key="btn_carregar_editar"):
            st.session_state["cc_nome"]         = _cc_cliente_editar.get("nome","")
            st.session_state["cc_cnpj"]         = _cc_cliente_editar.get("cnpj","")
            st.session_state["cc_cep"]          = _cc_cliente_editar.get("cep","")
            st.session_state["cc_endereco"]     = _cc_cliente_editar.get("endereco","")
            _servicos_cc = _servicos_padrao_empresa_ativa()
            # No cadastro separado, o serviço é determinado pelo painel ativo.
            st.session_state["cc_srv_rt"]       = bool(_servicos_cc.get("rt", False))
            st.session_state["cc_srv_limpeza"]  = bool(_servicos_cc.get("limpeza", False))
            st.session_state["cc_operadores_vinculados"] = _normalizar_lista_textos_unicos(_cc_cliente_editar.get("operadores_vinculados", []))
            st.session_state["cc_contato"]      = _cc_cliente_editar.get("contato","")
            st.session_state["cc_telefone"]     = _cc_cliente_editar.get("telefone","")
            st.session_state["cc_vol_adulto"]   = str(_cc_cliente_editar.get("vol_adulto","") or "")
            st.session_state["cc_vol_infantil"] = str(_cc_cliente_editar.get("vol_infantil","") or "")
            st.session_state["cc_vol_family"]   = str(_cc_cliente_editar.get("vol_family","") or "")
            try:
                st.session_state["cc_verificacoes_semanais"] = int(_cc_cliente_editar.get("verificacoes_semanais", 3) or 3)
            except Exception:
                st.session_state["cc_verificacoes_semanais"] = 3
            _piscs_extras = _cc_cliente_editar.get("piscinas_extras", [])
            st.session_state["cc_pisc_extra1_nome"] = _piscs_extras[0]["nome"] if len(_piscs_extras) > 0 else ""
            st.session_state["cc_pisc_extra1_vol"]  = str(_piscs_extras[0].get("vol","") or "") if len(_piscs_extras) > 0 else ""
            st.session_state["cc_pisc_extra2_nome"] = _piscs_extras[1]["nome"] if len(_piscs_extras) > 1 else ""
            st.session_state["cc_pisc_extra2_vol"]  = str(_piscs_extras[1].get("vol","") or "") if len(_piscs_extras) > 1 else ""
            st.rerun()
    else:
        st.info("Nenhum cliente para editar.")

# ── Formulário ───────────────────────────────────────────────────────────────
def _mask_cc_cnpj():
    st.session_state["cc_cnpj"] = formatar_cnpj(st.session_state.get("cc_cnpj",""))

def _mask_cc_telefone():
    st.session_state["cc_telefone"] = formatar_telefone(st.session_state.get("cc_telefone",""))

# Serviço vinculado ao cadastro conforme painel ativo
st.markdown("**🧩 Serviço vinculado ao cadastro**")
_cc_servicos = _servicos_padrao_empresa_ativa()
_cc_empresa_val = _servicos_para_empresa(_cc_servicos)
if _empresa_ativa_codigo() == "aqua_gestao":
    st.info("🔵 Este cadastro será salvo como cliente de RT / Controle Técnico da Aqua Gestão.")
else:
    st.info("⭐ Este cadastro será salvo como cliente de Limpeza / Manutenção da Bem Star Piscinas.")
# Mantém as chaves antigas coerentes para não quebrar edições/sessão.
st.session_state["cc_srv_rt"] = bool(_cc_servicos.get("rt"))
st.session_state["cc_srv_limpeza"] = bool(_cc_servicos.get("limpeza"))

_operadores_disponiveis = []
_ops_raw_cc = (sheets_listar_operadores() or []) + (carregar_operadores() or [])
for _op in _ops_raw_cc:
    _nome_op = re.sub(r"\s+", " ", str(_op.get("nome", "") or "").strip())
    if _nome_op and _nome_op not in _operadores_disponiveis:
        _operadores_disponiveis.append(_nome_op)

# Fallback: se o Sheets falhar em um rerun, mantém a última lista válida.
if _operadores_disponiveis:
    st.session_state["_cache_operadores_disponiveis"] = list(_operadores_disponiveis)
else:
    _operadores_disponiveis = list(st.session_state.get("_cache_operadores_disponiveis", []))

# Evita que o Streamlit apague seleções já carregadas se a lista vier incompleta.
for _nome_sel_op in st.session_state.get("cc_operadores_vinculados", []) or []:
    if _nome_sel_op and _nome_sel_op not in _operadores_disponiveis:
        _operadores_disponiveis.append(_nome_sel_op)

cc_operadores_vinculados = st.multiselect(
    "Operadores vinculados ao condomínio",
    options=_operadores_disponiveis,
    key="cc_operadores_vinculados",
    help="Esses operadores verão este condomínio automaticamente ao entrar com o PIN."
)
if not _operadores_disponiveis:
    st.caption("Cadastre operadores no módulo de operadores para depois vinculá-los aos condomínios.")

cc1, cc2 = st.columns(2)
with cc1:
    cc_nome     = st.text_input("Nome do condomínio / local *", key="cc_nome", placeholder="Ex.: Residencial Bella Vista")
    # CEP com busca automática ViaCEP
    # Aplica CEP formatado se acabou de buscar
    if st.session_state.get("_cc_cep_fmt"):
        st.session_state["cc_cep"] = st.session_state.pop("_cc_cep_fmt")
    _cep_col1, _cep_col2 = st.columns([3, 1])
    with _cep_col1:
        cc_cep = st.text_input("CEP", key="cc_cep", placeholder="00000-000",
            help="Digite o CEP e clique em 🔍 para preencher o endereço automaticamente")
    with _cep_col2:
        st.markdown("<br>", unsafe_allow_html=True)
        _btn_cep = st.button("🔍", key="btn_buscar_cep", help="Buscar CEP")
    if _btn_cep:
        _cep_valor = re.sub(r"\D", "", st.session_state.get("cc_cep", ""))
        if len(_cep_valor) == 8:
            with st.spinner("Buscando CEP..."):
                _dados_cep = buscar_cep(_cep_valor)
            if _dados_cep:
                _logradouro = _dados_cep.get("logradouro", "")
                _bairro     = _dados_cep.get("bairro", "")
                _cidade     = _dados_cep.get("localidade", "")
                _uf         = _dados_cep.get("uf", "")
                _cep_fmt    = f"{_cep_valor[:5]}-{_cep_valor[5:]}"
                _end_auto   = ", ".join(p for p in [_logradouro, _bairro, f"{_cidade}/{_uf}", _cep_fmt] if p)
                st.session_state["cc_endereco"] = _end_auto
                st.session_state["_cc_cep_fmt"] = _cep_fmt
                st.rerun()
            else:
                st.error("CEP não encontrado. Verifique e tente novamente.")
        else:
            st.warning("Digite um CEP válido com 8 dígitos.")
    cc_endereco = st.text_area("Endereço completo", key="cc_endereco", height=70, placeholder="Rua, número, bairro, cidade")
with cc2:
    cc_cnpj     = st.text_input("CNPJ (opcional)", key="cc_cnpj", placeholder="00.000.000/0000-00", on_change=_mask_cc_cnpj)
    cc_contato  = st.text_input("Síndico / responsável", key="cc_contato", placeholder="Nome do responsável")
    cc_telefone = st.text_input("Telefone (opcional)", key="cc_telefone", placeholder="(34) 99999-9999", on_change=_mask_cc_telefone)

# ── Volumes das piscinas ─────────────────────────────────────────────────────
st.markdown("**🏊 Volumes das piscinas (m³)**")
st.caption("Preencha apenas as piscinas que este local possui. O volume é usado para calcular dosagens automaticamente.")

cv1, cv2, cv3 = st.columns(3)
with cv1:
    cc_vol_adulto   = st.text_input("🏊 Adulto (m³)", key="cc_vol_adulto",
        placeholder="ex: 150", help="Volume da piscina adulto em metros cúbicos")
with cv2:
    cc_vol_infantil = st.text_input("🐣 Infantil (m³)", key="cc_vol_infantil",
        placeholder="ex: 30", help="Volume da piscina infantil em metros cúbicos")
with cv3:
    cc_vol_family   = st.text_input("👨‍👩‍👧 Family (m³)", key="cc_vol_family",
        placeholder="ex: 50", help="Volume da piscina family em metros cúbicos")

# Piscinas extras (outra, SPA, coberta, etc.)
st.markdown("**Outras piscinas** (SPA, coberta, olímpica, etc.)")
_cv_extra1, _cv_extra2 = st.columns(2)
with _cv_extra1:
    cc_pisc_extra1_nome = st.text_input("Nome da piscina extra 1", key="cc_pisc_extra1_nome",
        placeholder="ex: SPA, Coberta, Olímpica")
with _cv_extra2:
    cc_pisc_extra1_vol  = st.text_input("Volume (m³)", key="cc_pisc_extra1_vol",
        placeholder="ex: 80")
_cv_extra3, _cv_extra4 = st.columns(2)
with _cv_extra3:
    cc_pisc_extra2_nome = st.text_input("Nome da piscina extra 2", key="cc_pisc_extra2_nome",
        placeholder="ex: Aquecida, Semiolímpica")
with _cv_extra4:
    cc_pisc_extra2_vol  = st.text_input("Volume (m³) ", key="cc_pisc_extra2_vol",
        placeholder="ex: 120")


st.markdown("**🧪 Rotina de verificação técnica**")
# Sanitiza a frequência antes de criar o widget. Isso evita erro visual do Streamlit
# ao alternar entre Novo cliente / Editar cliente existente com estados antigos.
try:
    _cc_freq_atual = int(st.session_state.get("cc_verificacoes_semanais", 3) or 3)
except Exception:
    _cc_freq_atual = 3
_cc_freq_atual = max(1, min(7, _cc_freq_atual))
# v5: não atribui value= quando key já está em session_state (evita conflito Streamlit)
if "cc_verificacoes_semanais" not in st.session_state:
    st.session_state["cc_verificacoes_semanais"] = _cc_freq_atual

_cc_freq_col1, _cc_freq_col2 = st.columns([1, 2.2])
with _cc_freq_col1:
    cc_verificacoes_semanais = st.number_input(
        "Verificações por semana",
        min_value=1, max_value=7,
        step=1, key="cc_verificacoes_semanais",
        help="Ex.: 3 verificações semanais geram 12 linhas padrão no relatório mensal."
    )
with _cc_freq_col2:
    _linhas_previstas = calcular_linhas_analises_por_frequencia(int(cc_verificacoes_semanais or 3))
    st.caption(f"Com {int(cc_verificacoes_semanais)}x por semana, o relatório mensal abrirá pelo menos {_linhas_previstas} linhas de análises. Se houver mais visitas importadas, o sistema expande automaticamente.")

def _parse_vol(v):
    try: return float(str(v).replace(",",".").strip() or 0)
    except: return 0.0

# ── Botão salvar ─────────────────────────────────────────────────────────────
_btn_label = "💾 Salvar cliente no Google Sheets" if _cc_modo == "➕ Novo cliente" else "💾 Salvar alterações"
if st.button(_btn_label, type="primary", use_container_width=True):
    if not cc_nome.strip():
        st.error("Informe o nome do condomínio.")
    else:
        _vol_a = _parse_vol(cc_vol_adulto)
        _vol_i = _parse_vol(cc_vol_infantil)
        _vol_f = _parse_vol(cc_vol_family)
        with st.spinner("Salvando no Google Sheets..."):
            _cc_servicos = _servicos_padrao_empresa_ativa()
            _cc_empresa_val = _servicos_para_empresa(_cc_servicos)
            _cc_servicos_norm = _normalizar_servicos_cliente({"servicos": _cc_servicos, "empresa": _cc_empresa_val})
            _cc_operadores_sel = _normalizar_lista_textos_unicos(cc_operadores_vinculados)
            _piscs_extras_form = []
            for _en, _ev in [
                (st.session_state.get("cc_pisc_extra1_nome","").strip(),
                 st.session_state.get("cc_pisc_extra1_vol","").strip()),
                (st.session_state.get("cc_pisc_extra2_nome","").strip(),
                 st.session_state.get("cc_pisc_extra2_vol","").strip()),
            ]:
                if _en:
                    try:
                        _ev_f = float(_ev.replace(",",".")) if _ev else 0
                    except:
                        _ev_f = 0
                    _piscs_extras_form.append({"nome": _en, "vol": _ev_f})

            if _cc_modo == "✏️ Editar cliente existente" and _cc_cliente_editar.get("id"):
                ok = sheets_editar_cliente(
                    id_cliente=_cc_cliente_editar["id"],
                    nome=cc_nome.strip(), cnpj=cc_cnpj.strip(),
                    endereco=cc_endereco.strip(), contato=cc_contato.strip(),
                    telefone=cc_telefone.strip(),
                    vol_adulto=_vol_a, vol_infantil=_vol_i, vol_family=_vol_f,
                    empresa=_cc_empresa_val,
                )
                msg_ok = f"✅ Cliente '{cc_nome}' atualizado!"
            else:
                ok = sheets_salvar_cliente(
                    nome=cc_nome.strip(), cnpj=cc_cnpj.strip(),
                    endereco=cc_endereco.strip(), contato=cc_contato.strip(),
                    telefone=cc_telefone.strip(),
                    vol_adulto=_vol_a, vol_infantil=_vol_i, vol_family=_vol_f,
                    empresa=_cc_empresa_val,
                )
                msg_ok = f"✅ Cliente '{cc_nome}' salvo! O operador já pode selecioná-lo no celular."

            _pasta_cliente = GENERATED_DIR / slugify_nome(cc_nome.strip())
            _pasta_cliente.mkdir(parents=True, exist_ok=True)
            _dados_cliente_local = carregar_dados_condominio(_pasta_cliente) or {}
            _dados_cliente_local.update({
                "nome_condominio": cc_nome.strip(),
                "cnpj_condominio": cc_cnpj.strip(),
                "cnpj": cc_cnpj.strip(),
                "cep": cc_cep.strip(),
                "endereco_condominio": cc_endereco.strip(),
                "endereco": cc_endereco.strip(),
                "nome_sindico": cc_contato.strip(),
                "contato": cc_contato.strip(),
                "telefone": cc_telefone.strip(),
                "vol_adulto": _vol_a,
                "vol_infantil": _vol_i,
                "vol_family": _vol_f,
                "verificacoes_semanais": int(cc_verificacoes_semanais or 3),
                "analises_mensais_padrao": calcular_linhas_analises_por_frequencia(int(cc_verificacoes_semanais or 3)),
                "empresa": _cc_empresa_val,
                "servicos": _cc_servicos_norm,
                "operadores_vinculados": _cc_operadores_sel,
                "piscinas_extras": _piscs_extras_form,
            })
            # Monta lista de piscinas ativas para o operador (sem precisar configurar)
            _piscinas_adm = []
            if _vol_a: _piscinas_adm.append("Piscina Adulto")
            if _vol_i: _piscinas_adm.append("Piscina Infantil")
            if _vol_f: _piscinas_adm.append("Piscina Family")
            for _pe in _piscs_extras_form:
                if _pe.get("nome"): _piscinas_adm.append(_pe["nome"])
            if _piscinas_adm:
                _dados_cliente_local["piscinas"] = _piscinas_adm
            salvar_dados_condominio(_pasta_cliente, _dados_cliente_local)
        if ok:
            st.success(msg_ok)
            st.cache_data.clear()
            st.session_state["_cc_limpar"] = True
            st.rerun()
        else:
            st.error("❌ Não foi possível salvar no Google Sheets.")
            erro_detalhado = st.session_state.get("_sheets_ultimo_erro","")
            if erro_detalhado:
                with st.expander("🔍 Ver diagnóstico do erro", expanded=True):
                    st.code(erro_detalhado, language="text")

st.markdown("</div>", unsafe_allow_html=True)


if _empresa_ativa_codigo() != "bem_star":
    # =========================================
    # CENTRAL DE ENVIO DE DOCUMENTOS — AQUA GESTÃO
    # =========================================
    # _CENTRAL_EMAIL_DOCUMENTOS_DEFINITIVA_V3_
    st.markdown('<div class="section-card aq-only">', unsafe_allow_html=True)
    st.subheader("📧 Central de Envio de Documentos")
    st.caption("Selecione contrato, aditivo, termo de ciência, POPs, relatórios ou anexos manuais e envie por SMTP Gmail com assinatura premium Aqua Gestão.")

    _pastas_envio = sorted([p for p in GENERATED_DIR.iterdir() if p.is_dir()], key=lambda p: p.name.lower()) if GENERATED_DIR.exists() else []
    _nomes_pastas_envio = [humanizar_nome_pasta(p.name) for p in _pastas_envio]
    _mapa_pastas_envio = dict(zip(_nomes_pastas_envio, _pastas_envio))

    _nome_envio_padrao = (st.session_state.get("rel_nome_condominio") or st.session_state.get("nome_condominio") or "").strip()
    _idx_envio = 0
    if _nome_envio_padrao and _nomes_pastas_envio:
        for _i, _n in enumerate(_nomes_pastas_envio):
            if nomes_condominio_equivalentes(_nome_envio_padrao, _n):
                _idx_envio = _i
                break

    if _nomes_pastas_envio:
        _nome_cond_envio = st.selectbox(
            "Condomínio / cliente para buscar documentos",
            options=_nomes_pastas_envio,
            index=_idx_envio,
            key="central_email_condominio",
        )
        _pasta_cond_envio = _mapa_pastas_envio.get(_nome_cond_envio)
    else:
        _nome_cond_envio = _nome_envio_padrao or "Condomínio"
        _pasta_cond_envio = None
        st.info("Nenhuma pasta local encontrada em generated. Ainda é possível enviar documentos usando upload manual.")

    _email_envio_padrao = (st.session_state.get("email_cliente") or st.session_state.get("termo_email_cliente") or "").strip()
    try:
        _dados_cond_envio = carregar_dados_condominio(_pasta_cond_envio) if _pasta_cond_envio else {}
        _email_envio_padrao = _email_envio_padrao or str(_dados_cond_envio.get("email_cliente", "") or _dados_cond_envio.get("email", "") or "").strip()
    except Exception:
        _dados_cond_envio = {}

    _docs_envio = _coletar_documentos_email_aqua(_pasta_cond_envio, st.session_state.get("ultimos_docs_gerados") or [])
    _docs_por_tipo = {"Contrato": [], "Aditivo": [], "Termo de ciência": [], "POP": [], "Relatório": [], "Outros": []}
    for _doc in _docs_envio:
        _nm = _doc.name.lower()
        if "contrato" in _nm:
            _docs_por_tipo["Contrato"].append(_doc)
        elif "aditivo" in _nm:
            _docs_por_tipo["Aditivo"].append(_doc)
        elif "termo" in _nm or "ciencia" in _nm or "ciência" in _nm:
            _docs_por_tipo["Termo de ciência"].append(_doc)
        elif "pop" in _nm:
            _docs_por_tipo["POP"].append(_doc)
        elif "relatorio" in _nm or "relatório" in _nm:
            _docs_por_tipo["Relatório"].append(_doc)
        else:
            _docs_por_tipo["Outros"].append(_doc)

    with st.expander("📤 Compor e enviar documentação", expanded=False):
        _env_c1, _env_c2 = st.columns(2)
        with _env_c1:
            _dest_env = st.text_input("Destinatário *", value=_email_envio_padrao, key="central_email_destinatario", placeholder="email@condominio.com.br")
            _assunto_env = st.text_input("Assunto *", value=f"Documentação técnica Aqua Gestão - {_nome_cond_envio}", key="central_email_assunto")
        with _env_c2:
            _cc_env = st.text_input("CC", value="", key="central_email_cc", placeholder="administradora@exemplo.com.br")
            _bcc_env = st.text_input("CCO", value="", key="central_email_bcc")

        _msg_env_padrao = (
            f"Prezados,\n\n"
            f"Encaminho em anexo a documentação técnica gerada pela Aqua Gestão referente ao {_nome_cond_envio}.\n\n"
            "Os documentos selecionados seguem para conferência, registro e arquivo interno do condomínio.\n\n"
            "Permaneço à disposição para qualquer esclarecimento."
        )
        _msg_env = st.text_area("Mensagem", value=_msg_env_padrao, height=180, key="central_email_mensagem")

        st.markdown("**Selecionar documentos locais:**")
        _selecionados_envio = []
        for _tipo_doc, _lista_docs in _docs_por_tipo.items():
            if not _lista_docs:
                continue
            _opcoes_tipo = [p.name for p in _lista_docs]
            _default_tipo = _opcoes_tipo[:2] if _tipo_doc in ("Contrato", "Aditivo", "Termo de ciência", "Relatório") else []
            _sel_tipo = st.multiselect(
                _tipo_doc,
                options=_opcoes_tipo,
                default=_default_tipo,
                key=f"central_email_tipo_{slugify_nome(_tipo_doc)}",
            )
            _mapa_tipo = {p.name: p for p in _lista_docs}
            _selecionados_envio.extend([_mapa_tipo[n] for n in _sel_tipo if n in _mapa_tipo])

        _uploads_env = st.file_uploader(
            "📎 Adicionar anexos manuais (PDF/DOCX)",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            key="central_email_uploads",
        )

        _status_cfg_env, _erro_cfg_env = _email_aqua_configurado()
        if not _status_cfg_env:
            st.warning(f"⚠️ SMTP não configurado: {_erro_cfg_env}")

        _qtd_upload_env = len(_uploads_env or [])
        st.caption(f"{len(_selecionados_envio)} documento(s) local(is) + {_qtd_upload_env} upload(s) selecionado(s).")

        if st.button("📨 Enviar documentos selecionados", type="primary", use_container_width=True, key="central_email_btn_enviar", disabled=not _status_cfg_env):
            if not _dest_env.strip():
                st.error("Informe o destinatário.")
            elif not _assunto_env.strip():
                st.error("Informe o assunto.")
            elif not _selecionados_envio and not _uploads_env:
                st.error("Selecione ou envie pelo menos um anexo.")
            else:
                import tempfile as _tmp_env
                _tmp_dir_env = Path(_tmp_env.mkdtemp())
                _anexos_env = list(_selecionados_envio)
                for _uf_env in (_uploads_env or []):
                    _p_env = _tmp_dir_env / _uf_env.name
                    _p_env.write_bytes(_uf_env.getbuffer())
                    _anexos_env.append(_p_env)
                _ok_env, _msg_retorno_env = enviar_email_aqua_smtp(
                    destinatario=_dest_env.strip(),
                    assunto=_assunto_env.strip(),
                    mensagem=_msg_env,
                    anexos=_anexos_env,
                    cc=_cc_env.strip(),
                    bcc=_bcc_env.strip(),
                )
                if _ok_env:
                    st.success(f"✅ {_msg_retorno_env}")
                else:
                    st.error(f"❌ {_msg_retorno_env}")

    st.markdown("</div>", unsafe_allow_html=True)

    # =========================================
    # PAINEL DE VENCIMENTOS
    # =========================================

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader("Painel de vencimentos")

    m1, m2, m3, m4, m5 = st.columns(5)
    with m1:
        st.metric("Total monitorado", len(painel_filtrado))
    with m2:
        st.metric("Vencidos", len([i for i in painel_filtrado if i["status"]["codigo"] == "vencido"]))
    with m3:
        st.metric("Vencem em breve", len([i for i in painel_filtrado if i["status"]["codigo"] == "vencendo"]))
    with m4:
        st.metric("Vigentes", len([i for i in painel_filtrado if i["status"]["codigo"] == "vigente"]))
    with m5:
        st.metric("Sem vigência", len([i for i in painel_filtrado if i["status"]["codigo"] == "indefinido"]))

    st.caption(
        "Painel operacional separado do histórico comum, com filtro central, exportação de cadastro e visualização dos últimos documentos."
    )

    aba1, aba2, aba3 = st.tabs(
        ["Condomínios vencidos", "Condomínios que vencem em breve", "Sem vigência válida"]
    )

    def render_exportacao_e_docs(item: dict, item_key: str):
        dados = item["dados"]
        ultimo_contrato = item["ultimo_contrato"]
        ultimo_aditivo = item["ultimo_aditivo"]
        ultimo_relatorio = item.get("ultimo_relatorio")

        st.markdown("<div class='docs-note'><strong>Documentos mais recentes</strong></div>", unsafe_allow_html=True)

        dc1, dc2, dc3, dc4, dc5 = st.columns(5)
        with dc1:
            if ultimo_contrato:
                if st.button("Abrir último contrato", key=f"abrir_contrato_{item_key}", use_container_width=True):
                    abrir_arquivo_windows(ultimo_contrato["path"])
            else:
                st.button("Sem contrato", key=f"sem_contrato_{item_key}", disabled=True, use_container_width=True)

        with dc2:
            if ultimo_aditivo:
                if st.button("Abrir último aditivo", key=f"abrir_aditivo_{item_key}", use_container_width=True):
                    abrir_arquivo_windows(ultimo_aditivo["path"])
            else:
                st.button("Sem aditivo", key=f"sem_aditivo_{item_key}", disabled=True, use_container_width=True)

        with dc3:
            if ultimo_relatorio:
                if st.button("Abrir último relatório", key=f"abrir_relatorio_{item_key}", use_container_width=True):
                    abrir_arquivo_windows(ultimo_relatorio["path"])
            else:
                st.button("Sem relatório", key=f"sem_relatorio_{item_key}", disabled=True, use_container_width=True)

        with dc4:
            if dados:
                json_bytes = json.dumps(dados, ensure_ascii=False, indent=2).encode("utf-8")
                st.download_button(
                    "Exportar JSON backup",
                    data=json_bytes,
                    file_name=f"{item['slug']}_cadastro.json",
                    mime="application/json",
                    key=f"download_json_{item_key}",
                    use_container_width=True,
                )
            else:
                st.button("Sem JSON", key=f"sem_json_{item_key}", disabled=True, use_container_width=True)

        with dc5:
            if dados:
                html = gerar_html_resumo_cadastro(item).encode("utf-8")
                st.download_button(
                    "Exportar resumo HTML",
                    data=html,
                    file_name=f"{item['slug']}_resumo_cadastro.html",
                    mime="text/html",
                    key=f"download_html_{item_key}",
                    use_container_width=True,
                )
            else:
                st.button("Sem resumo", key=f"sem_resumo_{item_key}", disabled=True, use_container_width=True)


    with aba1:
        if not itens_vencidos:
            st.markdown(
                "<div class='venc-empty'>Nenhum condomínio vencido no momento.</div>",
                unsafe_allow_html=True,
            )
        else:
            for item in itens_vencidos:
                nome_exibicao = item["nome_exibicao"]
                pasta = item["pasta"]
                dados_salvos = item["dados"]
                status = item["status"]
                data_fim = item["data_fim"] or "Não informada"
                item_key = chave_segura(f"painel_vencido_{pasta}")

                st.markdown(
                    f"""
                    <div class="venc-row">
                        <div class="venc-nome">{nome_exibicao}</div>
                        <div class="venc-meta"><strong>Data final:</strong> {data_fim}</div>
                        <div class="venc-meta"><strong>Situação:</strong> {texto_dias_restantes(status)}</div>
                        <span class="status-badge {status['css']}">{status['rotulo']}</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                c1, c2, c3, c4 = st.columns(4)
                with c1:
                    if st.button("Editar cadastro", key=f"editar_vencido_{item_key}", use_container_width=True):
                        aplicar_dados_no_formulario(dados_salvos)
                        st.session_state.painel_acao_msg = f"Cadastro de '{nome_exibicao}' carregado para edição."
                        st.rerun()

                with c2:
                    if st.button("Abrir pasta", key=f"abrir_vencido_{item_key}", use_container_width=True):
                        abrir_pasta_windows(pasta)

                with c3:
                    if st.button("Renovar no formulário", key=f"renovar_vencido_{item_key}", use_container_width=True):
                        ok, msg = preparar_renovacao_no_formulario(dados_salvos)
                        if ok:
                            st.session_state.painel_acao_msg = f"{msg} Condomínio: {nome_exibicao}."
                            st.rerun()
                        else:
                            st.error(msg)

                with c4:
                    if st.button("Gerar aditivo renovação", key=f"aditivo_vencido_{item_key}", use_container_width=True):
                        ok, msg = gerar_aditivo_renovacao_por_painel(pasta, st.session_state.alerta_vencimento_dias)
                        if ok:
                            st.session_state.painel_acao_msg = msg
                            st.rerun()
                        else:
                            st.error(msg)

                render_exportacao_e_docs(item, item_key)

    with aba2:
        if not itens_vencendo:
            st.markdown(
                "<div class='venc-empty'>Nenhum condomínio dentro da faixa de alerta no momento.</div>",
                unsafe_allow_html=True,
            )
        else:
            for item in itens_vencendo:
                nome_exibicao = item["nome_exibicao"]
                pasta = item["pasta"]
                dados_salvos = item["dados"]
                status = item["status"]
                data_fim = item["data_fim"] or "Não informada"
                item_key = chave_segura(f"painel_vencendo_{pasta}")

                st.markdown(
                    f"""
                    <div class="venc-row">
                        <div class="venc-nome">{nome_exibicao}</div>
                        <div class="venc-meta"><strong>Data final:</strong> {data_fim}</div>
                        <div class="venc-meta"><strong>Situação:</strong> {texto_dias_restantes(status)}</div>
                        <span class="status-badge {status['css']}">{status['rotulo']}</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                c1, c2, c3, c4 = st.columns(4)
                with c1:
                    if st.button("Editar cadastro", key=f"editar_vencendo_{item_key}", use_container_width=True):
                        aplicar_dados_no_formulario(dados_salvos)
                        st.session_state.painel_acao_msg = f"Cadastro de '{nome_exibicao}' carregado para edição."
                        st.rerun()

                with c2:
                    if st.button("Abrir pasta", key=f"abrir_vencendo_{item_key}", use_container_width=True):
                        abrir_pasta_windows(pasta)

                with c3:
                    if st.button("Renovar no formulário", key=f"renovar_vencendo_{item_key}", use_container_width=True):
                        ok, msg = preparar_renovacao_no_formulario(dados_salvos)
                        if ok:
                            st.session_state.painel_acao_msg = f"{msg} Condomínio: {nome_exibicao}."
                            st.rerun()
                        else:
                            st.error(msg)

                with c4:
                    if st.button("Gerar aditivo renovação", key=f"aditivo_vencendo_{item_key}", use_container_width=True):
                        ok, msg = gerar_aditivo_renovacao_por_painel(pasta, st.session_state.alerta_vencimento_dias)
                        if ok:
                            st.session_state.painel_acao_msg = msg
                            st.rerun()
                        else:
                            st.error(msg)

                render_exportacao_e_docs(item, item_key)

    with aba3:
        if not itens_indefinidos:
            st.markdown(
                "<div class='venc-empty'>Nenhum condomínio sem vigência válida encontrado.</div>",
                unsafe_allow_html=True,
            )
        else:
            for item in itens_indefinidos:
                nome_exibicao = item["nome_exibicao"]
                pasta = item["pasta"]
                dados_salvos = item["dados"]
                status = item["status"]
                origem = item["origem"]
                total_arquivos = len(item["arquivos"])
                item_key = chave_segura(f"painel_indefinido_{pasta}")

                st.markdown(
                    f"""
                    <div class="venc-row">
                        <div class="venc-nome">{nome_exibicao}</div>
                        <div class="venc-meta"><strong>Status:</strong> {status['rotulo']}</div>
                        <div class="venc-meta"><strong>Arquivos encontrados:</strong> {total_arquivos}</div>
                        <span class="status-badge {status['css']}">{status['rotulo']}</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                if origem == "legado_sem_json":
                    st.markdown(
                        "<div class='legacy-note'>Histórico antigo sem <strong>dados_condominio.json</strong>. "
                        "Agora você pode criar um cadastro inicial diretamente a partir desta pasta.</div>",
                        unsafe_allow_html=True,
                    )

                    c1, c2, c3 = st.columns(3)
                    with c1:
                        if st.button("Criar cadastro desta pasta", key=f"legado_{item_key}", use_container_width=True):
                            preparar_cadastro_legado(pasta.name)
                            st.session_state.painel_acao_msg = f"Cadastro inicial preparado a partir da pasta '{pasta.name}'."
                            st.rerun()

                    with c2:
                        if st.button("Abrir pasta", key=f"abrir_legado_{item_key}", use_container_width=True):
                            abrir_pasta_windows(pasta)

                    with c3:
                        st.button("Gerar aditivo renovação", key=f"aditivo_legado_{item_key}", disabled=True, use_container_width=True)
                else:
                    st.markdown(
                        "<div class='legacy-note'>Existe cadastro salvo, porém sem data final válida. "
                        "Carregue os dados no formulário e corrija a vigência.</div>",
                        unsafe_allow_html=True,
                    )

                    c1, c2, c3 = st.columns(3)
                    with c1:
                        if st.button("Editar cadastro", key=f"editar_indefinido_{item_key}", use_container_width=True):
                            aplicar_dados_no_formulario(dados_salvos)
                            st.session_state.painel_acao_msg = f"Cadastro de '{nome_exibicao}' carregado para edição."
                            st.rerun()
                    with c2:
                        if st.button("Abrir pasta", key=f"abrir_indefinido_{item_key}", use_container_width=True):
                            abrir_pasta_windows(pasta)
                    with c3:
                        st.button("Gerar aditivo renovação", key=f"aditivo_indefinido_{item_key}", disabled=True, use_container_width=True)

                render_exportacao_e_docs(item, item_key)

    st.markdown("</div>", unsafe_allow_html=True)

    # =========================================
    # CLIENTES SEM RT — CADASTRO E RELATÓRIO TÉCNICO
    # =========================================


# =========================================
# MÓDULOS EXCLUSIVOS BEM STAR — só renderizam no painel Bem Star
# =========================================
if st.session_state.get("empresa_ativa", "aqua_gestao") == "bem_star":
    st.markdown('<div class="section-card bs-only">', unsafe_allow_html=True)
    st.subheader("Bem Star — Relatório Técnico Simples / Operacional")
    st.caption("Módulo exclusivo da Bem Star Piscinas para limpeza, manutenção, visitas operacionais, análises básicas, dosagens e relatório sem RT.")
    
    with st.expander("📋 Cadastrar / selecionar cliente sem RT", expanded=False):
    
        # Lista clientes sem RT já cadastrados
        CLIENTES_SEM_RT_JSON = GENERATED_DIR / "_clientes_sem_rt.json"
    
        def carregar_clientes_sem_rt() -> list:
            if CLIENTES_SEM_RT_JSON.exists():
                try:
                    return json.loads(CLIENTES_SEM_RT_JSON.read_text(encoding="utf-8"))
                except Exception:
                    return []
            return []
    
        def salvar_clientes_sem_rt(lista: list):
            GENERATED_DIR.mkdir(exist_ok=True)
            CLIENTES_SEM_RT_JSON.write_text(json.dumps(lista, ensure_ascii=False, indent=2), encoding="utf-8")
    
        clientes_sem_rt = carregar_clientes_sem_rt()
    
        # ── Importar do Sheets ────────────────────────────────────────────────────
        @st.cache_data(ttl=60)
        def _clientes_sheets_csr():
            return sheets_listar_clientes_completo()
    
        _cls_sheets = filtrar_clientes_por_empresa(_clientes_sheets_csr(), "bem_star")
        if _cls_sheets:
            _opcoes_csr_imp = ["— Importar do cadastro principal (Sheets) —"] + [c["nome"] for c in _cls_sheets]
            _sel_csr_imp = st.selectbox(
                "🔗 Importar cliente do Google Sheets",
                _opcoes_csr_imp,
                key="sel_importar_csr",
                help="Importa os dados do cadastro principal para o formulário abaixo."
            )
            if _sel_csr_imp and _sel_csr_imp != "— Importar do cadastro principal (Sheets) —":
                if st.button("⬇️ Importar dados", key="btn_imp_csr", use_container_width=False):
                    _d = next((c for c in _cls_sheets if c["nome"] == _sel_csr_imp), {})
                    if _d:
                        st.session_state["csr_nome"]     = _d.get("nome", "")
                        st.session_state["csr_cnpj"]     = formatar_cnpj(_d.get("cnpj", ""))
                        st.session_state["csr_endereco"] = _d.get("endereco", "")
                        st.session_state["csr_contato"]  = _d.get("contato", "")
                        st.session_state["csr_telefone"] = formatar_telefone(_d.get("telefone", ""))
                        st.success(f"✅ Dados de '{_sel_csr_imp}' importados!")
                        st.rerun()
            st.markdown("---")
    
        st.markdown("**Novo cliente sem RT:**")
    
        def _mask_csr_cnpj():
            st.session_state["csr_cnpj"] = formatar_cnpj(st.session_state.get("csr_cnpj", ""))
    
        def _mask_csr_telefone():
            st.session_state["csr_telefone"] = formatar_telefone(st.session_state.get("csr_telefone", ""))
    
        csr1, csr2 = st.columns(2)
        with csr1:
            csr_nome = st.text_input("Nome do local / condomínio", key="csr_nome", placeholder="Ex.: Residencial Sol Nascente")
            if st.session_state.get("_csr_cep_fmt"):
                st.session_state["csr_cep"] = st.session_state.pop("_csr_cep_fmt")
            _csr_cep_c1, _csr_cep_c2 = st.columns([3, 1])
            with _csr_cep_c1:
                st.text_input("CEP", key="csr_cep", placeholder="00000-000",
                    help="Digite o CEP e clique em 🔍 para preencher o endereço automaticamente")
            with _csr_cep_c2:
                st.markdown("<br>", unsafe_allow_html=True)
                _btn_csr_cep = st.button("🔍", key="btn_buscar_cep_csr", help="Buscar CEP")
            if _btn_csr_cep:
                _cep_v = re.sub(r"\D", "", st.session_state.get("csr_cep", ""))
                if len(_cep_v) == 8:
                    with st.spinner("Buscando CEP..."):
                        _dc = buscar_cep(_cep_v)
                    if _dc:
                        _end = ", ".join(p for p in [_dc.get("logradouro",""), _dc.get("bairro",""), f"{_dc.get('localidade','')}/{_dc.get('uf','')}", f"{_cep_v[:5]}-{_cep_v[5:]}"] if p)
                        st.session_state["csr_endereco"] = _end
                        st.session_state["_csr_cep_fmt"] = f"{_cep_v[:5]}-{_cep_v[5:]}"
                        st.rerun()
                    else:
                        st.error("CEP não encontrado.")
                else:
                    st.warning("Digite um CEP válido com 8 dígitos.")
            csr_endereco = st.text_area("Endereço", key="csr_endereco", height=70, placeholder="Rua, número, bairro, cidade")
        with csr2:
            csr_cnpj = st.text_input("CNPJ (opcional)", key="csr_cnpj", placeholder="00.000.000/0000-00", on_change=_mask_csr_cnpj)
            csr_contato = st.text_input("Responsável / contato", key="csr_contato", placeholder="Nome do responsável")
            csr_telefone = st.text_input("Telefone (opcional)", key="csr_telefone", placeholder="(34) 99999-9999", on_change=_mask_csr_telefone)
    
        if st.button("➕ Salvar cliente (Bem Star)", use_container_width=True):
            if not csr_nome.strip():
                st.error("Informe o nome do local.")
            else:
                novo = {
                    "nome": csr_nome.strip(),
                    "cnpj": formatar_cnpj(csr_cnpj.strip()),
                    "endereco": csr_endereco.strip(),
                    "contato": csr_contato.strip(),
                    "telefone": formatar_telefone(csr_telefone.strip()),
                    "cadastrado_em": _agora_brasilia(),
                }
                # Atualiza se já existe, senão adiciona
                nomes_existentes = [c["nome"].lower() for c in clientes_sem_rt]
                if csr_nome.strip().lower() in nomes_existentes:
                    idx_ex = nomes_existentes.index(csr_nome.strip().lower())
                    clientes_sem_rt[idx_ex] = novo
                    st.success(f"Cliente '{csr_nome}' atualizado.")
                else:
                    clientes_sem_rt.append(novo)
                    st.success(f"Cliente '{csr_nome}' cadastrado com sucesso.")
                salvar_clientes_sem_rt(clientes_sem_rt)
                # Salva também no Google Sheets (col M = Bem Star Piscinas)
                with st.spinner("Sincronizando com Google Sheets..."):
                    _cl_sheets = sheets_listar_clientes_completo()
                    _existe_sheets = next((c for c in _cl_sheets
                        if c["nome"].lower().strip() == csr_nome.strip().lower()), None)
                    if _existe_sheets:
                        sheets_editar_cliente(
                            id_cliente=_existe_sheets["id"],
                            nome=csr_nome.strip(),
                            cnpj=formatar_cnpj(csr_cnpj.strip()),
                            endereco=csr_endereco.strip(),
                            contato=csr_contato.strip(),
                            telefone=formatar_telefone(csr_telefone.strip()),
                            empresa="Bem Star Piscinas",
                        )
                    else:
                        sheets_salvar_cliente(
                            nome=csr_nome.strip(),
                            cnpj=formatar_cnpj(csr_cnpj.strip()),
                            endereco=csr_endereco.strip(),
                            contato=csr_contato.strip(),
                            telefone=formatar_telefone(csr_telefone.strip()),
                            empresa="Bem Star Piscinas",
                        )
                st.cache_data.clear()
                # Cria pasta do cliente no generated
                pasta_csr = GENERATED_DIR / slugify_nome(csr_nome.strip())
                pasta_csr.mkdir(parents=True, exist_ok=True)
                _dados_csr_local = carregar_dados_condominio(pasta_csr) or {}
                _dados_csr_local.update({
                    "nome_condominio": csr_nome.strip(),
                    "cnpj_condominio": csr_cnpj.strip(),
                    "cnpj": csr_cnpj.strip(),
                    "endereco_condominio": csr_endereco.strip(),
                    "endereco": csr_endereco.strip(),
                    "nome_sindico": csr_contato.strip(),
                    "contato": csr_contato.strip(),
                    "telefone": csr_telefone.strip(),
                    "empresa": "Bem Star Piscinas",
                    "servicos": {"rt": False, "limpeza": True},
                    "tipo": "sem_rt",
                    "salvo_em": _agora_brasilia(),
                })
                salvar_dados_condominio(pasta_csr, _dados_csr_local)
                st.rerun()
    
        if clientes_sem_rt:
            st.markdown(f"**{len(clientes_sem_rt)} cliente(s) cadastrado(s) sem RT:**")
            for c in clientes_sem_rt:
                st.caption(f"📍 {c['nome']} | {c.get('contato','–')} | {c.get('endereco','–')[:50]}")
    
    # ---- GERAÇÃO DO RELATÓRIO TÉCNICO SIMPLES ----
    st.markdown("---")
    st.markdown("**📊 Relatório técnico Bem Star Piscinas (sem RT)**")
    
    # Carrega clientes Bem Star do Sheets (fonte principal) + fallback JSON local
    @st.cache_data(ttl=30)
    def _clientes_bem_star_relatorio():
        _todos = sheets_listar_clientes_completo()
        _bs = filtrar_clientes_por_empresa(_todos, "bem_star")
        # Fallback: também inclui clientes do JSON local
        _json_local = carregar_clientes_sem_rt() if CLIENTES_SEM_RT_JSON.exists() else []
        _nomes_sheets = {c["nome"].lower() for c in _bs}
        for _cl in _json_local:
            if _cl["nome"].lower() not in _nomes_sheets:
                _bs.append(_cl)
        return _bs
    
    clientes_sem_rt_reload = _clientes_bem_star_relatorio()
    opcoes_csr = [c["nome"] for c in clientes_sem_rt_reload]
    
    if not opcoes_csr:
        st.info("Cadastre um cliente Bem Star acima para gerar o relatório técnico.")
    else:
        rts1, rts2, rts3 = st.columns([2, 1, 1])
        with rts1:
            csr_sel = st.selectbox("Selecione o cliente", opcoes_csr, key="csr_sel_relatorio")
        with rts2:
            csr_mes = st.text_input("Mês", key="csr_mes_rel", placeholder=datetime.now().strftime("%m"))
        with rts3:
            csr_ano = st.text_input("Ano", key="csr_ano_rel", placeholder=str(datetime.now().year))
    
        csr_dados_sel = next((c for c in clientes_sem_rt_reload if c["nome"] == csr_sel), {})
    
        # ── Busca lançamentos: JSON local + Google Sheets ──────────────────────
        pasta_csr_sel = GENERATED_DIR / slugify_nome(csr_sel) if csr_sel else None
        _lanc_local_csr = []
        if pasta_csr_sel and pasta_csr_sel.exists():
            _dados_json_csr = carregar_dados_condominio(pasta_csr_sel)
            _lanc_local_csr = (_dados_json_csr or {}).get("lancamentos_campo", [])
    
        _lanc_sheets_csr = []
        if csr_sel:
            try:
                _lanc_sheets_csr = sheets_listar_lancamentos(csr_sel)
            except Exception:
                _lanc_sheets_csr = []
    
        # Deduplica local + Sheets
        _vistos_csr = set()
        _lanc_todos_csr = []
        for _lc in _lanc_local_csr + _lanc_sheets_csr:
            _ch = f"{_lc.get('data','')}-{_lc.get('operador','')}-{_lc.get('ph','')}"
            if _ch not in _vistos_csr:
                _vistos_csr.add(_ch)
                _lanc_todos_csr.append(_lc)
    
        # Filtra por mês/ano
        def _filtrar_mes_csr(lancamentos, mes, ano):
            """Filtra lançamentos pelo mês/ano, aceitando vários formatos de data."""
            if not mes or not ano:
                return lancamentos
            return [lc for lc in lancamentos if lancamento_pertence_mes_ano(lc.get("data", ""), mes, ano)]
    
        _mes_csr  = (csr_mes or "").strip()
        _ano_csr  = (csr_ano or str(datetime.now().year)).strip()
        lancamentos_csr = _filtrar_mes_csr(_lanc_todos_csr, _mes_csr, _ano_csr)
    
        # ── Painel de lançamentos disponíveis ────────────────────────────────
        if lancamentos_csr:
            _fonte_csr = "📱 local + Sheets" if _lanc_sheets_csr else "📱 local"
            _periodo_csr = f"{lancamentos_csr[0].get('data','?')} → {lancamentos_csr[-1].get('data','?')}"
            st.markdown(
                f"<div style='border:1px solid rgba(20,120,60,0.3);border-radius:12px;padding:12px 16px;"
                f"background:rgba(20,120,60,0.07);margin-bottom:12px;'>"
                f"<strong>📱 {len(lancamentos_csr)} lançamento(s) encontrado(s) — {_fonte_csr}</strong><br>"
                f"<span style='font-size:0.85rem;color:#3a6a3a;'>Período: {_periodo_csr}</span></div>",
                unsafe_allow_html=True,
            )
            with st.expander("👁 Ver lançamentos"):
                for _lc in lancamentos_csr:
                    st.caption(f"📅 {_lc.get('data','?')} | pH {_lc.get('ph','–')} | CRL {_lc.get('cloro_livre','–')} | op: {_lc.get('operador','–')}")
        else:
            st.info("Nenhum lançamento encontrado para este cliente/período. O operador precisa registrar visitas no modo campo.")
    
        csr_operador_nome = st.text_input("Operador responsável", key="csr_operador_rel", placeholder="Nome do operador")
        csr_obs_geral     = st.text_area("Observações gerais", key="csr_obs_rel", height=80,
            placeholder="Condições gerais da piscina, ocorrências, recomendações...")
    
        # ── Coleta fotos das visitas ─────────────────────────────────────────
        pasta_fotos_csr = (GENERATED_DIR / slugify_nome(csr_sel) / "fotos_campo") if csr_sel else None
        fotos_csr = []
        if pasta_fotos_csr and pasta_fotos_csr.exists():
            for _lc in lancamentos_csr:
                for _nf in _lc.get("fotos", []):
                    _pf = pasta_fotos_csr / _nf
                    if _pf.exists():
                        fotos_csr.append((_lc.get("data",""), _pf))
    
        if fotos_csr:
            st.caption(f"📷 {len(fotos_csr)} foto(s) serão incluídas no relatório.")
    
        
    if st.button("📄 Gerar relatório Bem Star (PDF)", type="primary", use_container_width=True):
        try:
            with st.spinner("Gerando relatório Bem Star..."):
                _resultado_bs = renderizar_relatorio_oficial("Bem Star Piscinas", preview=False)
    
            if not _resultado_bs.get("ok"):
                st.error(_resultado_bs.get("mensagem", "Erro ao gerar relatório Bem Star."))
            else:
                docx_csr = Path(_resultado_bs["docx"])
                pdf_csr = Path(_resultado_bs["pdf"])
                ok_pdf_csr = bool(_resultado_bs.get("pdf_ok"))
                err_pdf_csr = _resultado_bs.get("erro_pdf")
                _ctx_bs = _resultado_bs.get("dados", {})
                st.success(f"✅ {_resultado_bs.get('mensagem', 'Relatório Bem Star gerado com sucesso.')}")
                dl1, dl2 = st.columns(2)
                with dl1:
                    with open(docx_csr, "rb") as _f:
                        st.download_button("⬇️ Baixar DOCX", data=_f, file_name=docx_csr.name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                with dl2:
                    if ok_pdf_csr and pdf_csr.exists():
                        with open(pdf_csr, "rb") as _f:
                            st.download_button("⬇️ Baixar PDF", data=_f, file_name=pdf_csr.name,
                                mime="application/pdf", use_container_width=True)
                    else:
                        st.warning(f"PDF não gerado: {err_pdf_csr}")
    
                _msg_rel = montar_mensagem_bem_star(
                    nome_local=_ctx_bs.get("dados_cliente", {}).get("nome", _ctx_bs.get("cliente", csr_sel)),
                    responsavel=_ctx_bs.get("dados_cliente", {}).get("contato", ""),
                    tipo="relatorio",
                    mes=_ctx_bs.get("mes", csr_mes),
                    ano=_ctx_bs.get("ano", csr_ano),
                )
                exibir_bloco_envio_bem_star(
                    nome_local=_ctx_bs.get("dados_cliente", {}).get("nome", _ctx_bs.get("cliente", csr_sel)),
                    pasta=Path(_resultado_bs["pasta"]),
                    telefone=_ctx_bs.get("dados_cliente", {}).get("telefone", ""),
                    email=_ctx_bs.get("dados_cliente", {}).get("email", ""),
                    mensagem=_msg_rel,
                    key_suffix="relatorio",
                )
    
        except Exception as e:
            st.error(f"Erro ao gerar relatório Bem Star: {e}")
            import traceback
            st.code(traceback.format_exc(), language="text")
    
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    

    # =========================================
    # PROPOSTA COMERCIAL — BEM STAR
    # =========================================

    st.markdown('<div class="section-card bs-only">', unsafe_allow_html=True)
    st.subheader("📄 Proposta Comercial — Bem Star Piscinas")
    st.caption("Gera proposta personalizada em PDF premium para o cliente.")

    with st.expander("✏️ Preencher e gerar proposta", expanded=False):
        _pb_col1, _pb_col2 = st.columns(2)
        with _pb_col1:
            _pb_cliente  = st.text_input("Cliente / Estabelecimento", key="pb_cliente")
            _pb_cnpj     = st.text_input("CNPJ", key="pb_cnpj", placeholder="00.000.000/0001-00")
            _pb_sindico  = st.text_input("Responsavel / Sindico", key="pb_sindico")
            _pb_freq     = st.text_input("Frequencia de visitas", key="pb_freq", placeholder="ex: 2x por semana")
        with _pb_col2:
            _pb_endereco = st.text_input("Endereco", key="pb_endereco")
            _pb_piscinas = st.text_input("Piscina(s) / Volume hidrico", key="pb_piscinas", placeholder="ex: Adulto 150m3, Infantil 40m3")
            _pb_valor    = st.text_input("Valor mensal (R$)", key="pb_valor", placeholder="ex: 850,00")
            _pb_venc     = st.text_input("Dia de vencimento", key="pb_venc", placeholder="ex: 10")
        _pb_pgto  = st.selectbox("Forma de pagamento", ["PIX / Transferencia bancaria", "Boleto", "PIX", "Transferencia bancaria", "Cheque"], key="pb_pgto")
        _pb_prods = st.radio("Produtos quimicos", ["Nao inclusos — por conta do contratante", "Inclusos no valor mensal"], key="pb_prods", horizontal=True)

        if st.button("📄 Gerar Proposta Bem Star", type="primary", use_container_width=True, key="btn_gerar_prop_bs"):
            with st.spinner("Gerando proposta..."):
                _pb_dados = {
                    "cliente": _pb_cliente, "cnpj": _pb_cnpj, "endereco": _pb_endereco,
                    "sindico": _pb_sindico, "piscinas": _pb_piscinas, "frequencia": _pb_freq,
                    "valor": _pb_valor, "vencimento": _pb_venc, "pagamento": _pb_pgto,
                    "produtos_inclusos": _pb_prods,
                }
                try:
                    _pb_bytes = gerar_proposta_pdf(_pb_dados, "Bem Star Piscinas")
                    _pb_nome  = limpar_nome_arquivo(f"Proposta_Bem_Star_{_pb_cliente or 'Cliente'}_{datetime.now().strftime('%Y%m')}.pdf")
                    st.success("✅ Proposta gerada com sucesso!")
                    st.download_button("⬇️ Baixar Proposta PDF", data=_pb_bytes,
                        file_name=_pb_nome, mime="application/pdf", use_container_width=True, key="dl_prop_bs")
                except Exception as _e:
                    st.error(f"Erro ao gerar proposta: {_e}")

    st.markdown("</div>", unsafe_allow_html=True)

    # =========================================
    # CONTRATO BEM STAR PISCINAS
    # =========================================
    
    st.markdown('<div class="section-card bs-only">', unsafe_allow_html=True)
    st.subheader("📝 Contrato Bem Star Piscinas")
    st.caption("Gera o contrato de prestação de serviços de limpeza e manutenção de piscinas em PDF.")
    
    with st.expander("📋 Preencher e gerar contrato", expanded=False):
    
        # ── Seletor de cliente ────────────────────────────────────────────────────
        # Correção definitiva:
        # - evita que a lista suma em reruns;
        # - mantém último carregamento válido em session_state;
        # - permite atualizar manualmente a lista;
        # - filtra clientes da Bem Star sem depender de rádio lateral.
        @st.cache_data(ttl=300, show_spinner=False)
        def _clientes_bs_contrato():
            _todos = sheets_listar_clientes_completo() or []
            _locais = carregar_clientes_sem_rt() if CLIENTES_SEM_RT_JSON.exists() else []

            # Junta clientes locais não encontrados no Sheets
            _nomes_sheets = [str(c.get("nome", "")).strip() for c in _todos]
            for _cl in _locais:
                _nome_local = str(_cl.get("nome", "")).strip()
                if _nome_local and _nome_local not in _nomes_sheets:
                    _todos.append(_cl)

            # Mantém apenas clientes Bem Star, mas preserva cadastros antigos sem empresa definida
            _filtrados = []
            for _c in _todos:
                _nome = str(_c.get("nome", "")).strip()
                if not _nome:
                    continue
                _empresa = str(_c.get("empresa", "")).strip().lower()
                if ("bem star" in _empresa) or (_empresa == ""):
                    _filtrados.append(_c)

            # Deduplica por nome
            _resultado = []
            _vistos = set()
            for _c in _filtrados:
                _chave = str(_c.get("nome", "")).strip().lower()
                if _chave and _chave not in _vistos:
                    _vistos.add(_chave)
                    _resultado.append(_c)

            return sorted(_resultado, key=lambda x: str(x.get("nome", "")).lower())

        col_atualizar_bs, col_info_bs = st.columns([1, 3])
        with col_atualizar_bs:
            if st.button("🔄 Atualizar clientes", key="btn_atualizar_clientes_bs_contrato"):
                _clientes_bs_contrato.clear()
                st.session_state.pop("_bs_clientes_contrato_ultimo_ok", None)
                st.rerun()

        _bs_clientes = _clientes_bs_contrato() or []

        # Fallback: se o Sheets falhar em algum rerun, reaproveita a última lista boa da sessão
        if _bs_clientes:
            st.session_state["_bs_clientes_contrato_ultimo_ok"] = _bs_clientes
        else:
            _bs_clientes = st.session_state.get("_bs_clientes_contrato_ultimo_ok", [])

        _bs_nomes = ["— selecione ou preencha manualmente —"] + [c.get("nome", "") for c in _bs_clientes if c.get("nome")]

        if len(_bs_nomes) == 1:
            st.warning("⚠️ Nenhum cliente Bem Star carregado no momento. Clique em **Atualizar clientes**. Se persistir, teste a conexão com o Google Sheets.")

        _bs_sel = st.selectbox(
            "Carregar dados de cliente cadastrado",
            _bs_nomes,
            key="bs_cont_cliente_sel"
        )
    
        if st.button("📂 Carregar dados do cliente", key="btn_bs_cont_carregar"):
            _bs_dado = next((c for c in _bs_clientes if c["nome"] == _bs_sel), {})
            if _bs_dado:
                st.session_state["bs_cont_nome"]     = _bs_dado.get("nome", "")
                st.session_state["bs_cont_cnpj"]     = _bs_dado.get("cnpj", "")
                st.session_state["bs_cont_endereco"] = _bs_dado.get("endereco", "")
                st.session_state["bs_cont_contato"]  = _bs_dado.get("contato", "")
                st.session_state["bs_cont_telefone"] = _bs_dado.get("telefone", "")
                st.success(f"✅ Dados de '{_bs_dado['nome']}' carregados.")
                st.rerun()
    
        st.markdown("---")
        st.markdown("**Dados do Contratante**")
    
        _bc1, _bc2 = st.columns(2)
        with _bc1:
            bs_nome     = st.text_input("Nome / Razão social *", key="bs_cont_nome",
                placeholder="Ex.: Condomínio Residencial Bella Vista")
            bs_endereco = st.text_area("Endereço completo", key="bs_cont_endereco",
                height=70, placeholder="Rua, número, bairro, cidade/UF, CEP")
        with _bc2:
            bs_cnpj     = st.text_input("CPF / CNPJ", key="bs_cont_cnpj",
                placeholder="00.000.000/0000-00")
            bs_contato  = st.text_input("Representante / síndico", key="bs_cont_contato",
                placeholder="Nome completo do responsável")
            bs_telefone = st.text_input("Telefone / WhatsApp", key="bs_cont_telefone",
                placeholder="(34) 99999-9999")
    
        st.markdown("**Descrição da(s) piscina(s) atendida(s)**")
        bs_piscinas = st.text_area("Piscinas atendidas", key="bs_cont_piscinas",
            height=60,
            placeholder="Ex.: Piscina adulto (150 m³), piscina infantil (30 m³), descobertas")
    
        st.markdown("**Condições do serviço**")
        _bs_c1, _bs_c2, _bs_c3 = st.columns(3)
        with _bs_c1:
            bs_frequencia = st.selectbox("Frequência de visitas", 
                ["1 visita semanal", "2 visitas semanais", "3 visitas semanais", "Outra"],
                key="bs_cont_frequencia")
            if bs_frequencia == "Outra":
                bs_frequencia = st.text_input("Especificar frequência", key="bs_cont_freq_outro",
                    placeholder="Ex.: quinzenal")
        with _bs_c2:
            bs_produtos = st.radio("Produtos químicos", 
                ["Incluídos no valor", "Não incluídos (por conta do contratante)"],
                key="bs_cont_produtos")
        with _bs_c3:
            bs_prazo = st.text_input("Prazo de vigência (meses)", key="bs_cont_prazo",
                placeholder="Ex.: 12")
    
        st.markdown("**Valores e pagamento**")
        _bs_v1, _bs_v2, _bs_v3, _bs_v4 = st.columns(4)
        with _bs_v1:
            bs_valor = st.text_input("Valor mensal (R$) *", key="bs_cont_valor",
                placeholder="Ex.: 350,00")
        with _bs_v2:
            bs_valor_extenso = st.text_input("Valor por extenso", key="bs_cont_valor_extenso",
                placeholder="Ex.: trezentos e cinquenta reais")
        with _bs_v3:
            bs_vencimento = st.text_input("Dia de vencimento", key="bs_cont_vencimento",
                placeholder="Ex.: 10")
        with _bs_v4:
            bs_pagamento = st.selectbox("Forma de pagamento",
                ["Pix", "Boleto", "Transferência bancária", "Dinheiro", "Outro"],
                key="bs_cont_pagamento")
    
        st.markdown("**Duração do Contrato**")
        bs_duracao = st.radio("Tipo de contrato", 
            ["Por tempo indeterminado", "12 meses com prorrogação automática"],
            key="bs_cont_duracao", horizontal=True)
    
        st.markdown("**Datas**")
        _bs_d1, _bs_d2, _bs_d3 = st.columns(3)
        with _bs_d1:
            bs_data_inicio = st.text_input("Data de início", key="bs_cont_data_inicio",
                placeholder="dd/mm/aaaa", value=hoje_br(), on_change=on_change_bs_cont_data_inicio)
        with _bs_d2:
            if "indeterminado" in bs_duracao:
                st.text_input("Data de término", value="Indeterminado", disabled=True)
            else:
                bs_data_fim = st.text_input("Data de término", key="bs_cont_data_fim",
                    placeholder="dd/mm/aaaa", on_change=on_change_bs_cont_data_fim)
        with _bs_d3:
            bs_local_ass = st.text_input("Local de assinatura", key="bs_cont_local",
                placeholder="Ex.: Uberlândia/MG", value="Uberlândia/MG")
        bs_data_ass = st.text_input("Data de assinatura", key="bs_cont_data_ass",
            placeholder="dd/mm/aaaa", value=hoje_br())
    
        st.markdown("---")
    
        if st.button("📄 Gerar Contrato Bem Star (PDF)", type="primary", use_container_width=True,
                key="btn_gerar_contrato_bs"):
            if not (st.session_state.get("bs_cont_nome","")).strip():
                st.error("Informe o nome do contratante.")
            elif not (st.session_state.get("bs_cont_valor","")).strip():
                st.error("Informe o valor mensal.")
            else:
                # ── Tenta gerar via template_bem_star.docx (preferencial) ──────
                if TEMPLATE_BEM_STAR.exists():
                    _duracao_bs  = st.session_state.get("bs_cont_duracao", "12 meses com prorrogação automática")
                    _prazo_bs    = "indeterminado" if "indeterminado" in _duracao_bs else "12 meses"
                    _fim_bs      = "Indeterminado" if "indeterminado" in _duracao_bs else (
                        (st.session_state.get("bs_cont_data_fim","")).strip() or "—"
                    )
                    _freq_bs     = (st.session_state.get("bs_cont_freq_outro","").strip()
                                    or st.session_state.get("bs_cont_frequencia",""))
                    _prods_inc_bs = "incluídos" in st.session_state.get("bs_cont_produtos","").lower()
                    gerar_contrato_bem_star_docx(
                        nome_contratante    = (st.session_state.get("bs_cont_nome","")).strip(),
                        cpf_cnpj            = (st.session_state.get("bs_cont_cnpj","")).strip(),
                        endereco_contratante= (st.session_state.get("bs_cont_endereco","")).strip(),
                        valor_mensal        = valor_para_template((st.session_state.get("bs_cont_valor","")).strip()),
                        valor_extenso       = (st.session_state.get("bs_cont_valor_extenso","")).strip(),
                        dia_pagamento       = (st.session_state.get("bs_cont_vencimento","")).strip() or "10",
                        forma_pagamento     = st.session_state.get("bs_cont_pagamento","Pix"),
                        prazo_contrato      = _prazo_bs,
                        data_inicio         = (st.session_state.get("bs_cont_data_inicio","")).strip() or hoje_br(),
                        data_fim            = _fim_bs,
                        local_data_assinatura = f"{(st.session_state.get('bs_cont_local','Uberlândia/MG')).strip()}, {(st.session_state.get('bs_cont_data_ass','')).strip() or hoje_br()}",
                        piscinas_atendidas  = (st.session_state.get("bs_cont_piscinas","")).strip(),
                        produtos_incluidos  = ("Produtos incluídos no valor mensal" if _prods_inc_bs
                                              else "Produtos não incluídos no valor mensal"),
                    )
                else:
                    # ── Fallback ReportLab (mantido para compatibilidade) ──────
                    try:
                        from reportlab.lib.units import cm
                        from reportlab.lib import colors
                        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                            Table, TableStyle, HRFlowable)
                        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
                        import io as _io
    
                        # ── Coleta valores ─────────────────────────────────────────
                        _nome     = (st.session_state.get("bs_cont_nome","")).strip()
                        _cnpj     = (st.session_state.get("bs_cont_cnpj","")).strip()
                        _end      = (st.session_state.get("bs_cont_endereco","")).strip()
                        _contato  = (st.session_state.get("bs_cont_contato","")).strip()
                        _tel      = (st.session_state.get("bs_cont_telefone","")).strip()
                        _piscinas = (st.session_state.get("bs_cont_piscinas","")).strip()
                        _freq     = st.session_state.get("bs_cont_freq_outro","").strip() or                             st.session_state.get("bs_cont_frequencia","")
                        _prods_inc = "incluídos" in st.session_state.get("bs_cont_produtos","").lower()
                        _prazo    = "indeterminado" if "indeterminado" in _duracao else "12"
                        _valor    = (st.session_state.get("bs_cont_valor","")).strip()
                        _ext      = (st.session_state.get("bs_cont_valor_extenso","")).strip() or _valor
                        _venc     = (st.session_state.get("bs_cont_vencimento","")).strip() or "10"
                        _pgto     = st.session_state.get("bs_cont_pagamento","Pix")
                        _inicio   = (st.session_state.get("bs_cont_data_inicio","")).strip() or hoje_br()
                        _duracao  = st.session_state.get("bs_cont_duracao", "12 meses com prorrogação automática")
                        _fim      = "Indeterminado" if "indeterminado" in _duracao else ((st.session_state.get("bs_cont_data_fim","")).strip() or "—")
                        _local    = (st.session_state.get("bs_cont_local","")).strip() or "Uberlândia/MG"
                        _data_ass = (st.session_state.get("bs_cont_data_ass","")).strip() or hoje_br()
                        _qualif   = f"inscrito(a) no CPF/CNPJ sob nº {_cnpj}," if _cnpj else ""
                        _piscinas_txt = _piscinas or "conforme descrição operacional acordada entre as partes"
                        _prod_txt = "estão incluídos no valor mensal contratado" if _prods_inc                             else "não estão incluídos no valor mensal contratado"
    
                        # ── Estilos ReportLab ──────────────────────────────────────
                        styles = getSampleStyleSheet()
                        s_titulo = ParagraphStyle("titulo", parent=styles["Heading1"],
                            fontSize=14, alignment=TA_CENTER, spaceAfter=4, textColor=colors.HexColor("#0d3d75"))
                        s_sub = ParagraphStyle("sub", parent=styles["Normal"],
                            fontSize=11, alignment=TA_CENTER, spaceAfter=2, textColor=colors.HexColor("#0d3d75"))
                        s_clausula = ParagraphStyle("clausula", parent=styles["Normal"],
                            fontSize=10, spaceBefore=10, spaceAfter=3, fontName="Helvetica-Bold",
                            textColor=colors.HexColor("#0d3d75"),
                            borderPad=4, borderColor=colors.HexColor("#0d3d75"),
                            leftIndent=0)
                        s_body = ParagraphStyle("body", parent=styles["Normal"],
                            fontSize=9.5, alignment=TA_JUSTIFY, spaceBefore=2, spaceAfter=4,
                            leading=14, leftIndent=8)
                        s_center = ParagraphStyle("center", parent=styles["Normal"],
                            fontSize=10, alignment=TA_CENTER, spaceBefore=4)
                        s_small = ParagraphStyle("small", parent=styles["Normal"],
                            fontSize=8, textColor=colors.grey, alignment=TA_CENTER)
    
                        # ── Monta story ────────────────────────────────────────────
                        story = []
    
                        # Logo Bem Star se disponível
                        _logo_bs = encontrar_logo_bem_star()
                        if _logo_bs and _logo_bs.exists():
                            from reportlab.platypus import Image as RLImage
                            _img = RLImage(str(_logo_bs), width=7*cm, height=2.5*cm,
                                kind="proportional")
                            _img.hAlign = "CENTER"
                            story.append(_img)
                            story.append(Spacer(1, 0.4*cm))
    
                        story.append(Paragraph("CONTRATO DE PRESTAÇÃO DE SERVIÇOS", s_sub))
                        story.append(Paragraph("Limpeza e Manutenção de Piscinas", ParagraphStyle(
                            "sub2", parent=styles["Normal"], fontSize=10,
                            alignment=TA_CENTER, textColor=colors.HexColor("#5d7288"), spaceAfter=4)))
                        story.append(Spacer(1, 0.3*cm))
                        story.append(HRFlowable(width="100%", thickness=2,
                            color=colors.HexColor("#0d3d75")))
                        story.append(Spacer(1, 0.3*cm))
    
                        # Tabela identificação
                        id_data = [
                            ["CONTRATADA", "BEM STAR PISCINAS LTDA., CNPJ 26.799.958/0001-88\nAv. Getúlio Vargas, 4411, Jardim das Palmeiras, Uberlândia/MG, CEP 38.412-316"],
                            ["CONTRATANTE", f"{_nome}{', ' + _qualif if _qualif else ''} com endereço em {_end or '—'}."],
                        ]
                        t_id = Table(id_data, colWidths=[3.5*cm, 14*cm])
                        t_id.setStyle(TableStyle([
                            ("FONTNAME",  (0,0), (0,-1), "Helvetica-Bold"),
                            ("FONTSIZE",  (0,0), (-1,-1), 9),
                            ("VALIGN",    (0,0), (-1,-1), "MIDDLE"),
                            ("BOX",       (0,0), (-1,-1), 1, colors.HexColor("#0d3d75")),
                            ("INNERGRID", (0,0), (-1,-1), 0.5, colors.HexColor("#c0c8d8")),
                            ("BACKGROUND",(0,0), (0,-1), colors.HexColor("#0d3d75")),
                            ("TEXTCOLOR", (0,0), (0,-1), colors.white),
                            ("TOPPADDING",(0,0),(-1,-1), 7),
                            ("BOTTOMPADDING",(0,0),(-1,-1), 7),
                            ("LEFTPADDING",(0,0),(-1,-1), 8),
                            ("RIGHTPADDING",(0,0),(-1,-1), 8),
                        ]))
                        story.append(t_id)
                        story.append(Spacer(1, 0.3*cm))
                        story.append(Paragraph(
                            "As partes acima identificadas têm entre si justo e contratado o presente "
                            "instrumento, regido pelas cláusulas e condições seguintes.", s_body))
    
                        # Cláusulas
                        clausulas = [
                            ("CLÁUSULA 1 — DO OBJETO",
                             "O presente contrato tem por objeto a prestação, pela CONTRATADA, de serviços regulares "
                             "de limpeza, conservação e manutenção operacional de piscina(s) localizada(s) no "
                             f"endereço do CONTRATANTE. Piscina(s) atendida(s): {_piscinas_txt}. "
                             "Os serviços abrangem: aspiração, escovação de paredes e bordas, peneiração e retirada "
                             "de resíduos, limpeza de cestos de skimmer e pré-filtro, acompanhamento visual das "
                             "condições da água e operações rotineiras de circulação e filtração. Este contrato "
                             "não inclui obras civis, substituição estrutural de equipamentos, reformas hidráulicas, "
                             "reparos elétricos, laudos, perícias ou outros serviços extraordinários não previstos."),
    
                            ("CLÁUSULA 2 — DA FREQUÊNCIA E EXECUÇÃO",
                             f"Os serviços serão executados com a seguinte frequência: {_freq}. "
                             "As visitas ocorrerão em dias e horários definidos conforme programação operacional "
                             "da CONTRATADA, podendo haver ajustes por necessidade climática, operacional, "
                             "feriados, caso fortuito ou força maior. Serviços extraordinários, emergenciais "
                             "ou fora da rotina poderão ser cobrados à parte, mediante comunicação prévia."),
    
                            ("CLÁUSULA 3 — DOS PRODUTOS E MATERIAIS",
                             f"Os produtos químicos, acessórios, insumos e materiais consumíveis {_prod_txt}. "
                             + ("Quando não incluídos, caberá ao CONTRATANTE providenciar todos os produtos e "
                                "materiais necessários em quantidade e qualidade suficientes para a execução dos serviços. "
                                "A falta de produtos ou condições inadequadas poderá impactar a qualidade do resultado "
                                "operacional, sem que isso caracterize inadimplemento da CONTRATADA."
                                if not _prods_inc else "")),
    
                            ("CLÁUSULA 4 — DO PREÇO E DO PAGAMENTO",
                             f"Pela prestação dos serviços, o CONTRATANTE pagará à CONTRATADA o valor mensal de "
                             f"R$ {_valor} ({_ext}). O vencimento ocorrerá todo dia {_venc} de cada mês, "
                             f"mediante {_pgto}. O atraso sujeitará o CONTRATANTE a multa de 2%, juros de 1% "
                             "ao mês pro rata die e correção monetária. Persistindo a inadimplência, a CONTRATADA "
                             "poderá suspender os serviços após comunicação prévia."),
    
                            ("CLÁUSULA 5 — DO PRAZO DE VIGÊNCIA",
                             f"O presente contrato vigorará pelo prazo de {_prazo} meses, com início em {_inicio} "
                             f"e término em {_fim}. Findo o prazo, poderá ser renovado por acordo entre as partes, "
                             "inclusive de forma tácita, caso a prestação prossiga sem oposição expressa. "
                             "Em contratos superiores a 12 meses, o valor poderá ser reajustado anualmente pelo IPCA/IBGE."),
    
                            ("CLÁUSULA 6 — DAS OBRIGAÇÕES DAS PARTES",
                             "A CONTRATADA executará os serviços com zelo, técnica e boa-fé; informará o "
                             "CONTRATANTE sobre irregularidades que interfiram na conservação da piscina; e "
                             "manterá sigilo sobre informações não públicas. "
                             "O CONTRATANTE garantirá livre acesso ao local; manterá os sistemas básicos em "
                             "funcionamento; comunicará previamente alterações relevantes de uso, eventos ou reformas; "
                             "e efetuará o pagamento na forma e prazo pactuados."),
    
                            ("CLÁUSULA 7 — DAS LIMITAÇÕES DE RESPONSABILIDADE",
                             "A CONTRATADA responde pela execução dos serviços dentro do escopo previsto, "
                             "não se responsabilizando por falhas estruturais preexistentes ou supervenientes; "
                             "defeitos elétricos, hidráulicos ou mecânicos fora do escopo; danos decorrentes de "
                             "mau uso, acesso indevido de terceiros, vandalismo, intempéries ou ausência de insumos."),
    
                            ("CLÁUSULA 8 — DA RESCISÃO",
                             "Este contrato poderá ser rescindido por mútuo acordo; por qualquer das partes, "
                             "mediante aviso prévio por escrito de 30 dias; imediatamente, em caso de "
                             "descumprimento contratual relevante após notificação sem saneamento; ou por "
                             "inadimplência do CONTRATANTE. Permanecerão exigíveis os valores já vencidos e "
                             "serviços efetivamente prestados."),
    
                            ("CLÁUSULA 9 — DAS DISPOSIÇÕES GERAIS",
                             "Os dados fornecidos serão utilizados exclusivamente para execução do contrato, "
                             "comunicações operacionais e rotinas administrativas. Qualquer alteração de escopo, "
                             "frequência, preço ou condição relevante deverá ser formalizada por escrito. "
                             "Fica eleito o foro da Comarca de Uberlândia/MG para dirimir quaisquer controvérsias "
                             "oriundas deste contrato, com renúncia de qualquer outro, por mais privilegiado que seja."),
                        ]
    
                        for titulo_cl, texto_cl in clausulas:
                            story.append(Paragraph(titulo_cl, s_clausula))
                            story.append(Paragraph(texto_cl, s_body))
    
                        story.append(Spacer(1, 0.5*cm))
                        story.append(HRFlowable(width="100%", thickness=0.5,
                            color=colors.HexColor("#c0c8d8")))
                        story.append(Spacer(1, 0.3*cm))
                        story.append(Paragraph(
                            f"E, por estarem justas e contratadas, firmam o presente instrumento em 2 (duas) "
                            f"vias de igual teor e forma.", s_body))
                        story.append(Spacer(1, 0.2*cm))
                        story.append(Paragraph(f"{_local}, {_data_ass}.", s_center))
                        story.append(Spacer(1, 1*cm))
    
                        # Tabela de assinaturas
                        ass_data = [
                            ["___________________________________",
                             "___________________________________"],
                            ["BEM STAR PISCINAS LTDA.\nCONTRATADA",
                             f"{_nome}\nCONTRATANTE"],
                            ["", ""],
                            ["___________________________________",
                             "___________________________________"],
                            ["TESTEMUNHA 1\nNome:\nCPF:",
                             "TESTEMUNHA 2\nNome:\nCPF:"],
                        ]
                        t_ass = Table(ass_data, colWidths=[9*cm, 9*cm])
                        t_ass.setStyle(TableStyle([
                            ("ALIGN",    (0,0), (-1,-1), "CENTER"),
                            ("FONTSIZE", (0,0), (-1,-1), 9),
                            ("VALIGN",   (0,0), (-1,-1), "TOP"),
                            ("TOPPADDING", (0,0),(-1,-1), 4),
                        ]))
                        story.append(t_ass)
                        story.append(Spacer(1, 0.3*cm))
                        story.append(HRFlowable(width="100%", thickness=0.5,
                            color=colors.HexColor("#0d3d75")))
                        story.append(Spacer(1, 0.15*cm))
                        story.append(Paragraph(
                            "Bem Star Piscinas LTDA. | CNPJ 26.799.958/0001-88 | "
                            "Av. Getúlio Vargas, 4411, Uberlândia/MG | Documento de uso operacional",
                            s_small))
    
                        # ── Gera PDF ───────────────────────────────────────────────
                        pasta_bs_cont = GENERATED_DIR / slugify_nome(_nome)
                        pasta_bs_cont.mkdir(parents=True, exist_ok=True)
                        _ts_bs = datetime.now().strftime("%Y%m%d_%H%M%S")
                        pdf_bs_path = pasta_bs_cont / f"{_ts_bs}_{slugify_nome(_nome)}_CONTRATO_BS.pdf"
    
                        _buf = _io.BytesIO()
                        doc_rl = SimpleDocTemplate(
                            _buf,
                            pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.5*cm, rightMargin=2.5*cm,
                            title=f"Contrato Bem Star — {_nome}",
                            author="Bem Star Piscinas",
                        )
                        doc_rl.build(story)
                        _pdf_bytes = _buf.getvalue()
    
                        with open(pdf_bs_path, "wb") as _pf:
                            _pf.write(_pdf_bytes)
    
                        st.success(f"✅ Contrato gerado para {_nome}!")
                        st.download_button(
                            "⬇️ Baixar Contrato PDF",
                            data=_pdf_bytes,
                            file_name=pdf_bs_path.name,
                            mime="application/pdf",
                            use_container_width=True,
                            key="btn_dl_contrato_bs",
                        )
    
                        # ── Bloco de envio ────────────────────────────────────────
                        _msg_cont = montar_mensagem_bem_star(
                            nome_local   = _nome,
                            responsavel  = _contato,
                            tipo         = "contrato",
                        )
                        exibir_bloco_envio_bem_star(
                            nome_local   = _nome,
                            pasta        = pasta_bs_cont,
                            telefone     = _tel,
                            email        = "",
                            mensagem     = _msg_cont,
                            key_suffix   = "contrato",
                        )
    
                    except Exception as _e:
                            st.error(f"Erro ao gerar contrato: {_e}")
                            import traceback as _tb
                            st.code(_tb.format_exc(), language="text")
    
    st.markdown("</div>", unsafe_allow_html=True)
    

    # =========================================
    # PROPOSTA COMERCIAL — AQUA GESTAO
    # =========================================

    st.markdown('<div class="section-card aq-only">', unsafe_allow_html=True)
    st.subheader("📄 Proposta Comercial — Aqua Gestao RT")
    st.caption("Gera proposta personalizada em PDF premium para o cliente.")

    with st.expander("✏️ Preencher e gerar proposta", expanded=False):
        _pa_col1, _pa_col2 = st.columns(2)
        with _pa_col1:
            _pa_cliente  = st.text_input("Cliente / Estabelecimento", key="pa_cliente")
            _pa_cnpj     = st.text_input("CNPJ", key="pa_cnpj", placeholder="00.000.000/0001-00")
            _pa_sindico  = st.text_input("Responsavel / Sindico", key="pa_sindico")
            _pa_freq     = st.text_input("Frequencia de visitas", key="pa_freq", placeholder="ex: 2x por semana")
        with _pa_col2:
            _pa_endereco = st.text_input("Endereco", key="pa_endereco")
            _pa_piscinas = st.text_input("Piscina(s) / Volume hidrico", key="pa_piscinas", placeholder="ex: Adulto 150m3, Infantil 40m3")
            _pa_valor    = st.text_input("Valor mensal (R$)", key="pa_valor", placeholder="ex: 450,00")
            _pa_venc     = st.text_input("Dia de vencimento", key="pa_venc", placeholder="ex: 10")
        _pa_pgto = st.selectbox("Forma de pagamento", ["PIX / Transferencia bancaria", "Boleto", "PIX", "Transferencia bancaria"], key="pa_pgto")

        if st.button("📄 Gerar Proposta Aqua Gestao", type="primary", use_container_width=True, key="btn_gerar_prop_aq"):
            with st.spinner("Gerando proposta..."):
                _pa_dados = {
                    "cliente": _pa_cliente, "cnpj": _pa_cnpj, "endereco": _pa_endereco,
                    "sindico": _pa_sindico, "piscinas": _pa_piscinas, "frequencia": _pa_freq,
                    "valor": _pa_valor, "vencimento": _pa_venc, "pagamento": _pa_pgto,
                }
                try:
                    _pa_bytes = gerar_proposta_pdf(_pa_dados, "Aqua Gestao – Controle Tecnico de Piscinas")
                    _pa_nome  = limpar_nome_arquivo(f"Proposta_Aqua_Gestao_{_pa_cliente or 'Cliente'}_{datetime.now().strftime('%Y%m')}.pdf")
                    st.success("✅ Proposta gerada com sucesso!")
                    st.download_button("⬇️ Baixar Proposta PDF", data=_pa_bytes,
                        file_name=_pa_nome, mime="application/pdf", use_container_width=True, key="dl_prop_aq")
                except Exception as _e:
                    st.error(f"Erro ao gerar proposta: {_e}")

    st.markdown("</div>", unsafe_allow_html=True)

    # =========================================
    # FORMULÁRIO
    # =========================================
    
    # Encerra aqui para impedir que módulos Aqua Gestão apareçam no painel Bem Star.
    st.stop()

st.markdown('<div class="section-card aq-only" id="sec-formulario">', unsafe_allow_html=True)
st.subheader("Dados do contrato de Responsabilidade Técnica")

# ── Seletor de cliente do Sheets ──────────────────────────────────────────────
@st.cache_data(ttl=60)
def _clientes_completos_cache():
    return sheets_listar_clientes_completo()

_clientes_rt = filtrar_clientes_por_empresa(_clientes_completos_cache(), "aqua_gestao")
if _clientes_rt:
    _opcoes_rt = ["— Selecionar cliente cadastrado —"] + [c["nome"] for c in _clientes_rt]
    _sel_rt = st.selectbox(
        "🔗 Carregar dados de cliente cadastrado",
        _opcoes_rt,
        key="sel_cliente_rt",
        help="Selecione um cliente para preencher automaticamente os campos abaixo."
    )
    if _sel_rt and _sel_rt != "— Selecionar cliente cadastrado —":
        if st.button("⬇️ Preencher formulário com dados deste cliente", key="btn_carregar_cliente_rt", use_container_width=True):
            _dados_rt = next((c for c in _clientes_rt if c["nome"] == _sel_rt), {})
            if _dados_rt:
                st.session_state["nome_condominio"]   = _dados_rt.get("nome", "")
                st.session_state["cnpj_condominio"]   = formatar_cnpj(_dados_rt.get("cnpj", ""))
                st.session_state["endereco_condominio"] = _dados_rt.get("endereco", "")
                st.session_state["nome_sindico"]      = _dados_rt.get("contato", "")
                st.session_state["whatsapp_cliente"]  = formatar_telefone(_dados_rt.get("telefone", ""))
                st.session_state["email_cliente"]     = _dados_rt.get("email", "")
                st.success(f"✅ Dados de '{_sel_rt}' carregados! Ajuste o que precisar antes de gerar.")
                st.rerun()
else:
    st.info("💡 Cadastre clientes na seção acima para carregar os dados automaticamente aqui.")

st.markdown("---")
col1, col2 = st.columns(2)

with col1:
    st.text_input("Nome do condomínio", key="nome_condominio", on_change=on_change_nome_condominio)
    st.text_input(
        "CNPJ do condomínio",
        key="cnpj_condominio",
        on_change=on_change_cnpj,
        placeholder="00.000.000/0000-00",
    )
    if st.session_state.get("_rt_cep_fmt"):
        st.session_state["rt_cep"] = st.session_state.pop("_rt_cep_fmt")
    _rt_cep_c1, _rt_cep_c2 = st.columns([3, 1])
    with _rt_cep_c1:
        st.text_input("CEP", key="rt_cep", placeholder="00000-000",
            help="Digite o CEP e clique em 🔍 para preencher o endereço automaticamente")
    with _rt_cep_c2:
        st.markdown("<br>", unsafe_allow_html=True)
        _btn_rt_cep = st.button("🔍", key="btn_buscar_cep_rt", help="Buscar CEP")
    if _btn_rt_cep:
        _cep_v = re.sub(r"\D", "", st.session_state.get("rt_cep", ""))
        if len(_cep_v) == 8:
            with st.spinner("Buscando CEP..."):
                _dc = buscar_cep(_cep_v)
            if _dc:
                _end = ", ".join(p for p in [_dc.get("logradouro",""), _dc.get("bairro",""), f"{_dc.get('localidade','')}/{_dc.get('uf','')}", f"{_cep_v[:5]}-{_cep_v[5:]}"] if p)
                st.session_state["endereco_condominio"] = _end
                st.session_state["_rt_cep_fmt"] = f"{_cep_v[:5]}-{_cep_v[5:]}"
                st.rerun()
            else:
                st.error("CEP não encontrado.")
        else:
            st.warning("Digite um CEP válido com 8 dígitos.")
    st.text_area("Endereço do condomínio", key="endereco_condominio", height=100)
    st.text_input("Nome do síndico / representante", key="nome_sindico")
    st.text_input(
        "CPF do síndico / representante",
        key="cpf_sindico",
        on_change=on_change_cpf,
        placeholder="000.000.000-00",
    )
    st.text_input(
        "Cargo / qualificação do representante",
        key="cargo_sindico",
        placeholder="Ex.: Síndico, Presidente, Administrador",
        value=st.session_state.get("cargo_sindico", "Síndico"),
    )

with col2:
    st.text_input(
        "Valor mensal",
        key="valor_mensal",
        on_change=on_change_valor_mensal,
        placeholder="R$ 1.621,00",
    )
    st.text_input(
        "Valor por extenso",
        key="valor_mensal_extenso",
        placeholder="um mil, seiscentos e vinte e um reais",
    )
    st.selectbox(
        "Frequência de visitas",
        options=["1 (uma)", "2 (duas)", "3 (três)", "4 (quatro)"],
        key="frequencia_visitas",
        help="Número de visitas técnicas semanais contratadas.",
    )
    _col_dia, _col_forma = st.columns(2)
    with _col_dia:
        st.text_input(
            "Dia de vencimento",
            key="dia_pagamento",
            placeholder="10",
            max_chars=2,
        )
    with _col_forma:
        st.selectbox(
            "Forma de pagamento",
            options=["Pix", "Transferência bancária", "Boleto", "Outro"],
            key="forma_pagamento",
        )
    st.text_input(
        "Data de início",
        key="data_inicio",
        placeholder="01/04/2026",
        on_change=on_change_data_inicio,
    )
    st.text_input(
        "Data de fim",
        key="data_fim",
        placeholder="31/03/2027",
        on_change=on_change_data_fim,
    )
    st.text_input(
        "Data de assinatura",
        key="data_assinatura",
        placeholder="dd/mm/aaaa",
        on_change=on_change_data_assinatura,
    )

if modo == "Modo Campo":
    st.markdown("---")
    col_campo1, col_campo2 = st.columns(2)
    with col_campo1:
        st.text_input(
            "WhatsApp do cliente (opcional)",
            key="whatsapp_cliente",
            on_change=on_change_whatsapp,
            placeholder="(34) 99999-9999",
        )
    with col_campo2:
        st.text_input("E-mail do cliente (opcional)", key="email_cliente")
else:
    st.markdown("---")
    col_cont1, col_cont2 = st.columns(2)
    with col_cont1:
        st.text_input(
            "WhatsApp do cliente",
            key="whatsapp_cliente",
            on_change=on_change_whatsapp,
            placeholder="(34) 99999-9999",
        )
    with col_cont2:
        st.text_input("E-mail do cliente", key="email_cliente")

st.text_area(
    "Observações internas (não vai para o contrato)",
    key="observacoes_internas",
    height=100,
    placeholder="Ex.: condição comercial específica, observação operacional, histórico jurídico...",
)

st.markdown('</div>', unsafe_allow_html=True)

# =========================================
# AÇÕES DE CADASTRO / RENOVAÇÃO
# =========================================

# _SECAO_ADITIVO_CNPJ_ADICIONADA_
# =========================================
# GERAR ADITIVO PARA CONTRATO EXISTENTE
# =========================================
st.markdown('<div class="section-card aq-only" id="sec-aditivo-cnpj">', unsafe_allow_html=True)
st.subheader("📄 Gerar aditivo para contrato existente")
st.caption("Localize um contrato já gerado pelo CNPJ e gere o 1º Termo Aditivo de desconto comercial.")

_adt_col1, _adt_col2 = st.columns([1.5, 2])
with _adt_col1:
    _adt_cnpj_input = st.text_input(
        "CNPJ do condomínio",
        key="adt_cnpj_busca",
        placeholder="00.000.000/0000-00",
        help="Digite o CNPJ do condomínio para localizar o contrato gerado.",
    )
    _adt_cnpj_digits = re.sub(r"\D", "", _adt_cnpj_input or "")

# _BUSCA_CNPJ_VIA_SHEETS_
# Busca pelo Sheets (persiste no Cloud) e enriquece com JSON local se disponivel.
# _ADITIVO_CLOUD_CAMPOS_MANUAIS_
_adt_dados_encontrados = None
_adt_pasta_encontrada = None
_adt_cliente_sheets = None
_adt_tem_json_local = False

if len(_adt_cnpj_digits) == 14:
    # 1. Busca no Sheets pelo CNPJ.
    for _cl in (_todos_clientes_painel or []):
        _cl_cnpj = re.sub(r"\D", "", str(_cl.get("cnpj", "") or ""))
        if _cl_cnpj == _adt_cnpj_digits:
            _adt_cliente_sheets = _cl
            break

    # 2. Tenta enriquecer com JSON local.
    # No Streamlit Cloud a pasta generated/ não é persistente, então o Sheets é a fonte principal.
    if _adt_cliente_sheets:
        _adt_nome_slug = slugify_nome(_adt_cliente_sheets.get("nome", ""))
        _adt_pasta_candidata = GENERATED_DIR / _adt_nome_slug
        _adt_json_local = carregar_dados_condominio(_adt_pasta_candidata) if _adt_pasta_candidata.exists() else None

        if _adt_json_local:
            _adt_tem_json_local = True
            _adt_dados_encontrados = _adt_json_local
            _adt_pasta_encontrada = _adt_pasta_candidata
        else:
            # Dados mínimos vindos do Sheets.
            # Os campos contratuais que não existem no Sheets serão preenchidos manualmente na tela.
            _adt_dados_encontrados = {
                "nome_condominio":     _adt_cliente_sheets.get("nome", ""),
                "cnpj_condominio":     _adt_cliente_sheets.get("cnpj", ""),
                "endereco_condominio": _adt_cliente_sheets.get("endereco", ""),
                "nome_sindico":        _adt_cliente_sheets.get("contato", ""),
                "cpf_sindico":         "",
                "cargo_sindico":       "Síndico",
                "valor_mensal":        "",
                "data_inicio":         "",
                "data_fim":            "",
            }
            _adt_pasta_encontrada = _adt_pasta_candidata

    # Segundo passe: enriquece campos contratuais ausentes a partir do JSON local,
    # mesmo quando o primeiro passe (baseado no slug principal) não encontrou o arquivo.
    # Cobre sessões novas no Cloud onde generated/ foi reiniciado mas o JSON ainda existe,
    # ou onde o slug gravado difere levemente do calculado agora.
    if _adt_dados_encontrados and not _adt_tem_json_local:
        _slugs_candidatos = list(dict.fromkeys(filter(None, [
            slugify_nome(_adt_dados_encontrados.get("nome_condominio", "")),
            slugify_nome(_adt_cliente_sheets.get("nome", "")) if _adt_cliente_sheets else None,
        ])))
        for _slug_cand in _slugs_candidatos:
            _pasta_cand2 = GENERATED_DIR / _slug_cand
            _json_cand2 = carregar_dados_condominio(_pasta_cand2) if _pasta_cand2.exists() else None
            if _json_cand2:
                _adt_tem_json_local = True
                # Preenche apenas campos em branco — não sobrescreve o que veio do Sheets
                for _campo in [
                    "valor_mensal", "data_inicio", "data_fim", "cpf_sindico",
                    "cargo_sindico", "nome_sindico", "endereco_condominio",
                    "cnpj_condominio", "dia_pagamento", "forma_pagamento",
                    "prazo_contrato", "valor_mensal_extenso",
                ]:
                    _v = _json_cand2.get(_campo, "")
                    if _v and not _adt_dados_encontrados.get(_campo):
                        _adt_dados_encontrados[_campo] = _v
                _adt_pasta_encontrada = _pasta_cand2
                break

if len(_adt_cnpj_digits) == 14 and _adt_dados_encontrados is None:
    st.warning(
        "Nenhum cliente encontrado com este CNPJ. "
        "Certifique-se de que o cliente está cadastrado na aba Clientes do Sheets com o CNPJ preenchido na coluna N."
    )

if _adt_dados_encontrados:
    _adt_nome = _adt_dados_encontrados.get("nome_condominio") or (_adt_pasta_encontrada.name if _adt_pasta_encontrada else "Condomínio")
    _adt_sindico = _adt_dados_encontrados.get("nome_sindico", "")
    _adt_cargo = _adt_dados_encontrados.get("cargo_sindico", "Síndico")

    with _adt_col2:
        st.success(f"✅ Cliente encontrado no Sheets: **{_adt_nome}**")
        if _adt_tem_json_local:
            st.caption("Cadastro enriquecido com dados locais do contrato.")
        else:
            st.caption("Sem JSON local no Cloud. Preencha os dados contratuais abaixo antes de gerar o aditivo.")

    # Se não houver JSON local, esses campos são obrigatórios para não gerar aditivo em branco.
    _adt_val_atual_base = str(_adt_dados_encontrados.get("valor_mensal", "") or "").strip()
    _adt_inicio_base = str(_adt_dados_encontrados.get("data_inicio", "") or "").strip()
    _adt_fim_base = str(_adt_dados_encontrados.get("data_fim", "") or "").strip()

    # Injeta no session_state ANTES dos widgets — Streamlit só popula o session_state
    # com o parâmetro value= após interação do usuário; sem isso a validação
    # sempre falha na primeira renderização mesmo com os dados visíveis na tela.
    if _adt_val_atual_base and not st.session_state.get("adt_valor_mensal_atual"):
        st.session_state["adt_valor_mensal_atual"] = _adt_val_atual_base
    if _adt_inicio_base and not st.session_state.get("adt_data_inicio_atual"):
        st.session_state["adt_data_inicio_atual"] = _adt_inicio_base
    if _adt_fim_base and not st.session_state.get("adt_data_fim_atual"):
        st.session_state["adt_data_fim_atual"] = _adt_fim_base

    st.markdown("#### Dados do contrato atual")
    _adt_m1, _adt_m2, _adt_m3 = st.columns(3)
    with _adt_m1:
        _adt_valor_mensal_atual = st.text_input(
            "Valor mensal atual do contrato *",
            key="adt_valor_mensal_atual",
            placeholder="R$ 1.621,00",
            help="Valor original/atual do contrato antes do desconto.",
        )
    with _adt_m2:
        _adt_data_inicio_atual = st.text_input(
            "Data de início do contrato *",
            key="adt_data_inicio_atual",
            placeholder="01/04/2026",
        )
    with _adt_m3:
        _adt_data_fim_atual = st.text_input(
            "Data de fim do contrato *",
            key="adt_data_fim_atual",
            placeholder="31/03/2027",
        )

    # Lê do session_state (já injetado acima ou digitado pelo usuário)
    # com fallback para o valor base — garante que a validação sempre encontre o dado
    _adt_val_atual = (st.session_state.get("adt_valor_mensal_atual") or _adt_val_atual_base).strip()
    _adt_inicio = (st.session_state.get("adt_data_inicio_atual") or _adt_inicio_base).strip()
    _adt_fim = (st.session_state.get("adt_data_fim_atual") or _adt_fim_base).strip()

    # Atualiza o dict antes da geração e antes do manifest.
    _adt_dados_encontrados["valor_mensal"] = _adt_val_atual
    _adt_dados_encontrados["data_inicio"] = _adt_inicio
    _adt_dados_encontrados["data_fim"] = _adt_fim

    st.markdown("#### Dados do aditivo")
    _adt_c1, _adt_c2, _adt_c3 = st.columns([1.2, 1.2, 2])
    with _adt_c1:
        _adt_valor_desconto = st.text_input(
            "Valor com desconto *",
            key="adt_valor_desconto",
            placeholder="R$ 810,50",
            help="Valor mensal após o desconto comercial. Será inserido no aditivo.",
        )
    with _adt_c2:
        _adt_data_ass = st.text_input(
            "Data de assinatura",
            key="adt_data_assinatura",
            value=hoje_br(),
            placeholder="dd/mm/aaaa",
        )
    with _adt_c3:
        st.markdown("<br>", unsafe_allow_html=True)
        _adt_gerar = st.button(
            "📄 Gerar aditivo",
            key="btn_gerar_aditivo_cnpj",
            type="primary",
            use_container_width=True,
        )

    if _adt_gerar:
        _adt_val_limpo = (st.session_state.get("adt_valor_desconto") or "").strip()
        _adt_data_limpa = (st.session_state.get("adt_data_assinatura") or "").strip()
        _adt_erros = []

        if not _adt_val_atual:
            _adt_erros.append("Informe o valor mensal atual do contrato.")
        if not _adt_inicio:
            _adt_erros.append("Informe a data de início do contrato.")
        if not _adt_fim:
            _adt_erros.append("Informe a data de fim do contrato.")
        if not _adt_val_limpo:
            _adt_erros.append("Informe o valor com desconto.")

        if _adt_erros:
            for _erro in _adt_erros:
                st.error(_erro)
        else:
            _adt_val_atual_fmt = moeda_br(_adt_val_atual) if not _adt_val_atual.startswith("R$") else _adt_val_atual
            _adt_val_fmt = moeda_br(_adt_val_limpo) if not _adt_val_limpo.startswith("R$") else _adt_val_limpo

            _adt_placeholders = {
                "{{NOME_CONDOMINIO}}": _adt_nome,
                "{{CNPJ_CONDOMINIO}}": _adt_dados_encontrados.get("cnpj_condominio", ""),
                "{{ENDERECO_CONDOMINIO}}": _adt_dados_encontrados.get("endereco_condominio", ""),
                "{{NOME_SINDICO}}": _adt_sindico,
                "{{CPF_SINDICO}}": _adt_dados_encontrados.get("cpf_sindico", ""),
                "{{CARGO_SINDICO}}": _adt_cargo,
                "{{VALOR_MENSAL}}": valor_para_template(_adt_val_atual_fmt),
                "{{VALOR_ADITIVO}}": valor_para_template(_adt_val_fmt),
                "{{DATA_INICIO}}": _adt_inicio,
                "{{DATA_FIM}}": _adt_fim,
                "{{DATA_INICIO_CONTRATO}}": _adt_inicio,
                "{{DATA_FIM_CONTRATO}}": _adt_fim,
                "{{DATA_ASSINATURA}}": _adt_data_limpa,
                "{{LOCAL_DATA_ASSINATURA}}": f"Uberlândia/MG, {_adt_data_limpa}",
                "{{NOME_CONTRATANTE}}": _adt_nome,
                "{{CPF_CNPJ_CONTRATANTE}}": _adt_dados_encontrados.get("cnpj_condominio", ""),
                "{{ENDERECO_CONTRATANTE}}": _adt_dados_encontrados.get("endereco_condominio", ""),
            }

            _adt_ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            _adt_base = limpar_nome_arquivo(f"Aditivo_RT_{_adt_nome}_{_adt_ts}")

            try:
                _adt_pasta_encontrada.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass

            _adt_docx = _adt_pasta_encontrada / f"{_adt_base}.docx"
            _adt_pdf  = _adt_pasta_encontrada / f"{_adt_base}.pdf"

            with st.spinner("Gerando aditivo..."):
                try:
                    gerar_documento(
                        template_path=TEMPLATE_ADITIVO,
                        output_docx=_adt_docx,
                        placeholders=_adt_placeholders,
                        incluir_assinaturas=False,
                    )
                    _adt_ok_pdf, _adt_err_pdf = salvar_aditivo_rt_pdf_premium_reportlab(_adt_placeholders, _adt_pdf)

                    _adt_dados_encontrados["valor_mensal"] = _adt_val_atual_fmt
                    _adt_dados_encontrados["valor_aditivo"] = _adt_val_fmt
                    _adt_dados_encontrados["data_inicio"] = _adt_inicio
                    _adt_dados_encontrados["data_fim"] = _adt_fim
                    _adt_dados_encontrados["data_assinatura"] = _adt_data_limpa
                    salvar_dados_condominio(_adt_pasta_encontrada, _adt_dados_encontrados)

                    registrar_documento_manifest(
                        pasta_condominio=_adt_pasta_encontrada,
                        nome_condominio=_adt_nome,
                        tipo="Aditivo",
                        arquivo_docx=_adt_docx,
                        arquivo_pdf=_adt_pdf,
                        pdf_gerado=_adt_ok_pdf,
                        erro_pdf=_adt_err_pdf,
                        dados_utilizados=_adt_dados_encontrados,
                    )

                    # _DOWNLOAD_ADITIVO_CNPJ_V2_
                    if _adt_ok_pdf:
                        st.success(f"✅ Aditivo gerado para {_adt_nome}. DOCX e PDF disponíveis para download.")
                    else:
                        st.warning(f"⚠️ DOCX gerado. PDF falhou: {_adt_err_pdf}. O DOCX está disponível para download.")

                    st.markdown("### Arquivos do aditivo")
                    _adt_col_docx, _adt_col_pdf, _adt_col_pasta = st.columns(3)

                    with _adt_col_docx:
                        if _adt_docx.exists():
                            with open(_adt_docx, "rb") as _adt_f_docx:
                                st.download_button(
                                    "⬇️ Baixar aditivo DOCX",
                                    data=_adt_f_docx.read(),
                                    file_name=_adt_docx.name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_aditivo_docx_cnpj_{_adt_ts}",
                                    use_container_width=True,
                                )
                        else:
                            st.error("DOCX do aditivo não encontrado.")

                    with _adt_col_pdf:
                        if _adt_pdf.exists():
                            with open(_adt_pdf, "rb") as _adt_f_pdf:
                                st.download_button(
                                    "⬇️ Baixar aditivo PDF",
                                    data=_adt_f_pdf.read(),
                                    file_name=_adt_pdf.name,
                                    mime="application/pdf",
                                    key=f"download_aditivo_pdf_cnpj_{_adt_ts}",
                                    use_container_width=True,
                                )
                        elif not _adt_ok_pdf:
                            st.warning("PDF não disponível. Baixe o DOCX.")

                    with _adt_col_pasta:
                        if st.button("📁 Abrir pasta do aditivo", key="btn_abrir_pasta_aditivo_cnpj"):
                            abrir_pasta_windows(_adt_pasta_encontrada)

                    # _EMAIL_ADITIVO_CNPJ_AQUA_V1_
                    _adt_email_msg = (
                        f"Prezados,\n\n"
                        f"Encaminho em anexo o termo aditivo referente ao {_adt_nome}, gerado pela Aqua Gestão - Controle Técnico de Piscinas.\n\n"
                        "O documento segue para conferência, registro e arquivo interno do condomínio.\n\n"
                        "Permaneço à disposição para qualquer ajuste ou esclarecimento."
                    )
                    exibir_envio_email_documentos_aqua(
                        nome_condominio=_adt_nome,
                        pasta_condominio=_adt_pasta_encontrada,
                        email_cliente=_adt_dados_encontrados.get("email_cliente", ""),
                        mensagem_padrao=_adt_email_msg,
                        documentos_sugeridos=[_adt_docx, _adt_pdf],
                        key_prefix=f"aditivo_cnpj_{_adt_ts}",
                    )

                except FileNotFoundError:
                    st.error("Template aditivo.docx não encontrado. Verifique o diagnóstico de saúde do sistema.")
                except Exception as _adt_ex:
                    st.error(f"Erro ao gerar aditivo: {_adt_ex}")

st.markdown("</div>", unsafe_allow_html=True)

st.markdown('<div class="section-card aq-only" id="sec-cadastro-renovacao">', unsafe_allow_html=True)
st.subheader("Cadastro e renovação")

r1, r2, r3 = st.columns([1.15, 1.15, 2])

with r1:
    if st.button("Salvar cadastro / atualizar JSON", use_container_width=True):
        ok, msg = salvar_cadastro_sem_gerar_documentos()
        if ok:
            st.success(msg)
            st.rerun()
        else:
            st.error(msg)

with r2:
    if st.button("Renovação anual rápida", use_container_width=True):
        novo_inicio, novo_fim = calcular_renovacao_anual(st.session_state.data_fim)
        if not novo_inicio or not novo_fim:
            st.error("Não foi possível renovar. Verifique se a data final atual está válida.")
        else:
            st.session_state.data_inicio = formatar_data_br(novo_inicio)
            st.session_state.data_fim = formatar_data_br(novo_fim)
            st.session_state.data_assinatura = hoje_br()
            st.success("Nova vigência anual preenchida automaticamente.")
            st.rerun()

with r3:
    st.caption(
        "Salvar cadastro atualiza apenas o JSON do condomínio. "
        "Renovação anual rápida ajusta datas no formulário sem gerar documento automaticamente."
    )

st.markdown("---")
st.markdown("**Geração de documento contratual — RT/ART**")
st.info(
    "Fluxo ajustado para CRQ/ART: gera somente contrato único de Responsabilidade Técnica, "
    "sem aditivo de desconto e sem valor reduzido. O contrato preserva remuneração técnica compatível "
    "com as exigências documentais do CRQ-MG."
)

# _BOTAO_ADITIVO_LINHA_CORRETA_
col_btn1, col_btn2, col_btn3, col_btn4 = st.columns([1.6, 1.4, 1, 1])

with col_btn1:
    gerar = st.button(
        "✅ Gerar contrato RT",
        type="primary",
        use_container_width=True,
    )

with col_btn2:
    gerar_aditivo_rapido = st.button(
        "📄 Gerar aditivo (desconto)",
        use_container_width=True,
        help="Gera o 1º Termo Aditivo de desconto comercial. "
             "Documento privado — NÃO é registrado no CRQ. "
             "Preencha o campo 'Valor aditivo' antes de gerar.",
    )

with col_btn3:
    if st.button("🗑️ Limpar", use_container_width=True):
        limpar_formulario()
        st.rerun()

with col_btn4:
    if st.button("📁 Abrir pasta", use_container_width=True):
        abrir_pasta_windows(GENERATED_DIR)

st.markdown("</div>", unsafe_allow_html=True)


# =========================================
# POPs ADAPTATIVOS — RT / ROTINA OPERACIONAL
# =========================================

def _dados_pops_rt_do_formulario() -> dict:
    """Coleta dados atuais do formulário para gerar o Caderno de POPs."""
    return {
        "nome_condominio": (st.session_state.get("nome_condominio") or "").strip(),
        "cnpj_condominio": (st.session_state.get("cnpj_condominio") or "").strip(),
        "endereco_condominio": (st.session_state.get("endereco_condominio") or "").strip(),
        "nome_sindico": (st.session_state.get("nome_sindico") or "").strip(),
        "cpf_sindico": (st.session_state.get("cpf_sindico") or "").strip(),
        "cargo_sindico": (st.session_state.get("cargo_sindico") or "Síndico").strip(),
        "data_assinatura": (st.session_state.get("data_assinatura") or hoje_br()).strip(),
        "volumes_piscinas": st.session_state.get("volumes_piscinas", ""),
        "executor_operacional": st.session_state.get("pops_executor_operacional", "Prestador externo"),
        "frequencia_operacional": st.session_state.get("pops_frequencia_operacional", "3 vezes por semana"),
        "observacao_pops": (st.session_state.get("pops_observacao") or "").strip(),
    }


def _lista_pops_adaptativos(executor: str, frequencia: str) -> list[dict]:
    """Retorna a lista base de POPs adaptada ao executor e à frequência operacional."""
    executor_txt = executor or "equipe operacional designada pela CONTRATANTE"
    frequencia_txt = frequencia or "conforme rotina definida pela CONTRATANTE"

    responsavel_execucao = "Equipe operacional designada pela CONTRATANTE"
    if "prestador" in executor_txt.lower():
        responsavel_execucao = "Prestador externo indicado ou contratado pela CONTRATANTE"
    elif "empresa" in executor_txt.lower():
        responsavel_execucao = "Empresa contratada contratada pela CONTRATANTE"
    elif "zelador" in executor_txt.lower():
        responsavel_execucao = "Zelador ou responsável local designado pela CONTRATANTE"
    elif "funcionário" in executor_txt.lower() or "funcionario" in executor_txt.lower():
        responsavel_execucao = "Funcionário próprio designado pela CONTRATANTE"
    elif "mista" in executor_txt.lower():
        responsavel_execucao = "Equipe operacional mista designada pela CONTRATANTE"

    freq_visita = frequencia_txt
    if "sem rotina" in frequencia_txt.lower():
        freq_visita = "sempre que houver visita operacional ou solicitação formal da CONTRATANTE"

    return [
        {
            "codigo": "POP 01",
            "titulo": "Chegada ao condomínio e inspeção inicial",
            "responsavel": responsavel_execucao,
            "frequencia": freq_visita,
            "objetivo": "Estabelecer uma rotina mínima de conferência antes do início da manutenção operacional da piscina.",
            "procedimento": [
                "Identificar-se na chegada, quando houver controle de acesso.",
                "Verificar visualmente a piscina antes de qualquer intervenção.",
                "Observar transparência da água, presença de sujeira, espuma, odor forte, alteração de cor, reclamações de usuários e condições gerais da área.",
                "Conferir a casa de máquinas, cestos, pré-filtro, pressão do filtro e disponibilidade dos produtos necessários.",
            ],
            "cuidados": [
                "Não liberar uso da piscina quando houver suspeita de risco sanitário, ausência de cloro residual ou condição visual insegura.",
                "Comunicar ao síndico/responsável local e ao RT situações críticas ou fora da rotina.",
            ],
            "registros": "Registrar data, horário, responsável pela visita e observações iniciais."
        },
        {
            "codigo": "POP 02",
            "titulo": "Medição de parâmetros físico-químicos",
            "responsavel": responsavel_execucao,
            "frequencia": freq_visita,
            "objetivo": "Padronizar a medição dos parâmetros mínimos de controle da água.",
            "procedimento": [
                "Coletar amostra em ponto representativo da piscina, longe do retorno imediato de água tratada.",
                "Medir pH e cloro livre em toda visita operacional.",
                "Medir cloro total, alcalinidade, dureza, ácido cianúrico, temperatura e demais parâmetros conforme orientação do RT ou rotina definida.",
                "Utilizar reagentes dentro da validade e equipamento limpo.",
            ],
            "cuidados": [
                "Não registrar valores estimados ou aproximados sem medição real.",
                "Quando houver dúvida no resultado, repetir a análise antes de dosar produto.",
            ],
            "registros": "Registrar todos os parâmetros na planilha, aplicativo ou formulário oficial do condomínio."
        },
        {
            "codigo": "POP 03",
            "titulo": "Registro obrigatório da visita",
            "responsavel": responsavel_execucao,
            "frequencia": freq_visita,
            "objetivo": "Garantir rastreabilidade técnica das ações executadas na rotina operacional.",
            "procedimento": [
                "Registrar data, horário, nome do executante e piscinas atendidas.",
                "Registrar parâmetros medidos antes das correções.",
                "Registrar produtos aplicados, quantidade, unidade e finalidade.",
                "Registrar lavagem de filtro, limpeza de cestos, observações e não conformidades.",
            ],
            "cuidados": [
                "A ausência de registro compromete a rastreabilidade técnica e poderá ser apontada como não conformidade.",
                "Não alterar registros anteriores sem justificativa formal.",
            ],
            "registros": "Planilha diária, ficha de visita, aplicativo ou outro meio definido pela CONTRATANTE e validado pelo RT."
        },
        {
            "codigo": "POP 04",
            "titulo": "Limpeza física da piscina",
            "responsavel": responsavel_execucao,
            "frequencia": freq_visita,
            "objetivo": "Remover sujeiras físicas que prejudicam a qualidade visual e aumentam a demanda química da água.",
            "procedimento": [
                "Realizar peneiração da superfície quando necessário.",
                "Escovar bordas, paredes, escadas, cantos e áreas com acúmulo visível.",
                "Aspirar o fundo conforme necessidade operacional.",
                "Remover sujidades grosseiras antes de realizar correções químicas mais intensas.",
            ],
            "cuidados": [
                "Não realizar aspiração ou manobras hidráulicas sem conhecimento do sistema.",
                "Comunicar excesso recorrente de sujeira, areia, folhas ou material orgânico.",
            ],
            "registros": "Registrar se houve peneiração, escovação, aspiração e qualquer anormalidade observada."
        },
        {
            "codigo": "POP 05",
            "titulo": "Limpeza de cestos e pré-filtro",
            "responsavel": responsavel_execucao,
            "frequencia": freq_visita,
            "objetivo": "Evitar perda de vazão, esforço da motobomba e redução da eficiência de filtração.",
            "procedimento": [
                "Desligar a motobomba antes de abrir pré-filtro ou manusear cestos quando aplicável.",
                "Limpar cestos de skimmer, coadeira, ralo de superfície e pré-filtro da bomba.",
                "Reinstalar corretamente as tampas, vedações e cestos antes de religar o sistema.",
                "Observar ruídos, entrada de ar, vazamentos ou dificuldade de escorva.",
            ],
            "cuidados": [
                "Nunca abrir pré-filtro com sistema pressurizado ou bomba em funcionamento.",
                "Comunicar vazamentos, tampa danificada, anel de vedação comprometido ou ruído anormal.",
            ],
            "registros": "Registrar limpeza realizada e anormalidades no conjunto hidráulico."
        },
        {
            "codigo": "POP 06",
            "titulo": "Retrolavagem e enxágue do filtro",
            "responsavel": responsavel_execucao,
            "frequencia": "Conforme pressão do filtro, perda de vazão, orientação técnica ou rotina operacional definida.",
            "objetivo": "Padronizar a limpeza do meio filtrante sem causar retorno de sujeira para a piscina.",
            "procedimento": [
                "Verificar pressão do filtro e condição visual da água.",
                "Desligar a bomba antes de mudar a posição da válvula seletora.",
                "Executar retrolavagem pelo tempo necessário até melhora visual da água de descarte.",
                "Executar enxágue antes de retornar para a posição filtrar.",
                "Religar o sistema e conferir vazamentos, pressão e funcionamento.",
            ],
            "cuidados": [
                "Nunca mudar a válvula seletora com a bomba ligada.",
                "Comunicar aumento recorrente de pressão, baixa vazão ou necessidade excessiva de retrolavagem.",
            ],
            "registros": "Registrar data, execução de retrolavagem, enxágue e qualquer observação do filtro."
        },
        {
            "codigo": "POP 07",
            "titulo": "Dosagem segura de produtos químicos",
            "responsavel": responsavel_execucao,
            "frequencia": "Somente após medição e conforme orientação técnica, rótulo, FISPQ/FDS e POP aplicável.",
            "objetivo": "Evitar dosagens aleatórias, incompatibilidades químicas e riscos à segurança dos usuários e operadores.",
            "procedimento": [
                "Medir os parâmetros antes da dosagem.",
                "Selecionar o produto conforme orientação técnica e finalidade da correção.",
                "Utilizar EPIs compatíveis antes de manusear produtos.",
                "Aplicar o produto conforme rótulo, orientação técnica e condição operacional da piscina.",
                "Registrar produto, quantidade, unidade, horário e finalidade.",
            ],
            "cuidados": [
                "Não misturar produtos químicos entre si.",
                "Não aplicar produto sem identificação, vencido, contaminado ou armazenado inadequadamente.",
                "Não executar dosagem corretiva sem registro.",
            ],
            "registros": "Registro obrigatório de todos os produtos aplicados ou recomendados."
        },
        {
            "codigo": "POP 08",
            "titulo": "Uso de EPIs e segurança química",
            "responsavel": responsavel_execucao,
            "frequencia": "Sempre que houver manipulação, transporte, diluição, dosagem ou armazenamento de produtos químicos.",
            "objetivo": "Reduzir risco de acidentes durante o manuseio de produtos químicos de piscina.",
            "procedimento": [
                "Utilizar luvas de proteção química, óculos de segurança, calçado fechado antiderrapante e vestimenta compatível.",
                "Utilizar máscara ou respirador compatível quando houver produto volátil, pó, odor intenso ou recomendação em FISPQ/FDS.",
                "Manter acesso à água corrente para lavagem em caso de contato acidental.",
                "Ler rótulos e seguir FISPQ/FDS, POPs e orientações do RT.",
            ],
            "cuidados": [
                "Interromper a atividade se não houver condição segura ou EPI mínimo disponível.",
                "Comunicar vazamentos, derramamentos, produto sem rótulo, odor forte ou condição insegura.",
            ],
            "registros": "Registrar não conformidades de EPI ou segurança química quando identificadas."
        },
        {
            "codigo": "POP 09",
            "titulo": "Organização da casa de máquinas e produtos",
            "responsavel": responsavel_execucao,
            "frequencia": freq_visita,
            "objetivo": "Manter ambiente técnico organizado, seguro e rastreável.",
            "procedimento": [
                "Manter produtos fechados, identificados e separados de acordo com compatibilidade.",
                "Evitar armazenamento direto no piso quando houver risco de umidade ou contaminação.",
                "Manter acesso livre aos equipamentos, registros, painéis e produtos de emergência.",
                "Não deixar recipientes abertos, ferramentas soltas ou resíduos após a visita.",
            ],
            "cuidados": [
                "Não armazenar produtos incompatíveis juntos.",
                "Comunicar ausência de identificação, embalagem comprometida ou condição de risco.",
            ],
            "registros": "Registrar não conformidades na casa de máquinas, quando existentes."
        },
        {
            "codigo": "POP 10",
            "titulo": "Comunicação de não conformidades",
            "responsavel": responsavel_execucao + " e responsável local da CONTRATANTE",
            "frequencia": "Sempre que identificada anormalidade técnica, sanitária, operacional ou de segurança.",
            "objetivo": "Garantir comunicação rápida de situações que exigem correção, registro ou decisão administrativa.",
            "procedimento": [
                "Identificar a não conformidade e registrar evidências quando possível.",
                "Comunicar o síndico/responsável local e o RT quando houver risco relevante.",
                "Informar parâmetros medidos, condição visual, produtos disponíveis e ação executada.",
                "Aguardar orientação técnica em situações críticas ou fora da rotina.",
            ],
            "cuidados": [
                "Não ocultar não conformidade operacional.",
                "Não liberar uso da piscina quando houver recomendação técnica de restrição.",
            ],
            "registros": "Registrar comunicação por meio rastreável: aplicativo, mensagem, e-mail, ficha ou relatório."
        },
        {
            "codigo": "POP 11",
            "titulo": "Verificação visual nos dias sem visita operacional",
            "responsavel": "Responsável local designado pela CONTRATANTE",
            "frequencia": "Diariamente nos dias sem visita do prestador/operador, quando aplicável.",
            "objetivo": "Criar rotina mínima de observação entre as visitas operacionais.",
            "procedimento": [
                "Observar transparência da água, odor forte, espuma, alteração de cor, sujeira excessiva e reclamações de usuários.",
                "Verificar se há sinal evidente de equipamento parado, vazamento ou área insegura.",
                "Comunicar imediatamente o síndico, prestador operacional e/ou RT em caso de anormalidade.",
            ],
            "cuidados": [
                "Esta verificação não substitui análise técnica nem rotina operacional completa.",
                "Em caso de suspeita de risco, evitar o uso até orientação do responsável técnico ou administrativo.",
            ],
            "registros": "Registrar a ocorrência e comunicação quando houver anormalidade."
        },
        {
            "codigo": "POP 12",
            "titulo": "Restrição preventiva de uso",
            "responsavel": "CONTRATANTE, com orientação do RT quando aplicável",
            "frequencia": "Sempre que houver suspeita de risco sanitário, químico, físico ou operacional.",
            "objetivo": "Definir conduta mínima para reduzir risco aos usuários em situações críticas.",
            "procedimento": [
                "Identificar sinais de risco: água turva, ausência de cloro residual, odor forte, suspeita de contaminação, acidente, equipamento crítico parado ou orientação técnica expressa.",
                "Comunicar imediatamente o síndico/responsável local e o RT.",
                "Sinalizar a restrição de uso quando houver decisão administrativa ou recomendação técnica.",
                "Registrar a ocorrência, horário, responsável pela decisão e medidas adotadas.",
            ],
            "cuidados": [
                "Não priorizar liberação da piscina em detrimento da segurança sanitária.",
                "A decisão administrativa cabe à CONTRATANTE, sem prejuízo da recomendação técnica do RT.",
            ],
            "registros": "Notificação, relatório, mensagem rastreável e/ou checklist de ocorrência."
        },
        {
            "codigo": "POP 13",
            "titulo": "Controle de acesso a produtos e casa de máquinas",
            "responsavel": "CONTRATANTE e responsável local",
            "frequencia": "Contínua.",
            "objetivo": "Reduzir acesso indevido de usuários, crianças, visitantes ou pessoas não autorizadas aos produtos e equipamentos.",
            "procedimento": [
                "Manter casa de máquinas e armazenamento de produtos com acesso controlado.",
                "Permitir acesso apenas a pessoas autorizadas e orientadas.",
                "Manter produtos fora do alcance de usuários e devidamente identificados.",
                "Comunicar imediatamente sinais de acesso indevido, violação, vazamento ou produto fora de local.",
            ],
            "cuidados": [
                "Não armazenar produtos em áreas de circulação comum.",
                "Não permitir que pessoa não orientada manipule produto químico.",
            ],
            "registros": "Registrar acesso indevido, ausência de identificação ou condição insegura."
        },
        {
            "codigo": "POP 14",
            "titulo": "Auditoria técnica dos registros pelo RT",
            "responsavel": "Aqua Gestão — Responsável Técnico",
            "frequencia": "Conforme periodicidade contratada e disponibilidade dos registros.",
            "objetivo": "Verificar tecnicamente os registros operacionais e apontar tendências, falhas ou não conformidades.",
            "procedimento": [
                "Avaliar registros de parâmetros, dosagens, ocorrências e comunicações.",
                "Confrontar resultados com histórico, condição da água e recomendações anteriores.",
                "Registrar não conformidades documentais ou operacionais.",
                "Emitir orientação técnica, relatório, notificação ou recomendação de correção quando necessário.",
            ],
            "cuidados": [
                "A ausência de registro ou dado incompleto limita a análise técnica.",
                "A auditoria técnica não substitui a execução operacional diária pela equipe designada pela CONTRATANTE.",
            ],
            "registros": "Relatório técnico, checklist, notificação ou observação em sistema."
        },
        {
            "codigo": "POP 15",
            "titulo": "Revisão técnica dos POPs",
            "responsavel": "Aqua Gestão — Responsável Técnico",
            "frequencia": "Sempre que houver alteração relevante na rotina, equipamento, risco, recorrência de falhas ou necessidade técnica.",
            "objetivo": "Manter os procedimentos atualizados com a realidade operacional do condomínio.",
            "procedimento": [
                "Revisar POPs quando houver mudança de prestador, equipe, produto, equipamento, frequência operacional ou não conformidade recorrente.",
                "Registrar versão, data e motivo da revisão.",
                "Comunicar a versão atualizada à CONTRATANTE e à equipe operacional designada.",
            ],
            "cuidados": [
                "POPs desatualizados podem gerar falhas de execução e rastreabilidade.",
                "A CONTRATANTE deve garantir que a equipe operacional tenha acesso à versão vigente.",
            ],
            "registros": "Controle de versão do Caderno de POPs e termos de ciência relacionados."
        },
    ]


def _gerar_pdf_caderno_pops(dados: dict) -> bytes:
    """Gera PDF Premium com Caderno de POPs adaptativo."""
    import io as _io
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether

    AZUL_ESCURO = colors.HexColor("#0D2A4A")
    AZUL_MEDIO = colors.HexColor("#1565A8")
    AZUL_CLARO = colors.HexColor("#EAF4FF")
    DOURADO = colors.HexColor("#C8960C")
    CINZA_TEXTO = colors.HexColor("#2F3742")
    CINZA_CLARO = colors.HexColor("#F4F7FA")
    BORDA = colors.HexColor("#D9E2EC")
    ALERTA = colors.HexColor("#FFF8E6")

    def val(chave, padrao="Dado não informado"):
        v = str(dados.get(chave, "") or "").strip()
        return v if v else padrao

    nome_cond = val("nome_condominio")
    cnpj = val("cnpj_condominio")
    endereco = val("endereco_condominio")
    sindico = val("nome_sindico")
    cargo_sindico = val("cargo_sindico", "Síndico/Representante legal")
    data_ass = val("data_assinatura", hoje_br())
    volumes = val("volumes_piscinas", "Volumes não informados")
    executor = val("executor_operacional", "Prestador externo")
    frequencia = val("frequencia_operacional", "3 vezes por semana")
    obs = val("observacao_pops", "")

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=2.4 * cm,
        bottomMargin=1.8 * cm,
        title=f"Caderno de POPs — {nome_cond}",
        author="Aqua Gestão Controle Técnico Ltda",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="PopsTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=AZUL_ESCURO, alignment=TA_CENTER, spaceAfter=5))
    styles.add(ParagraphStyle(name="PopsSub", parent=styles["Normal"], fontName="Helvetica", fontSize=9, leading=12, textColor=CINZA_TEXTO, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="PopsH", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13.5, textColor=AZUL_ESCURO, spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name="PopsMiniH", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.4, leading=10.5, textColor=AZUL_MEDIO, spaceBefore=3, spaceAfter=2))
    styles.add(ParagraphStyle(name="PopsBody", parent=styles["Normal"], fontName="Helvetica", fontSize=8.45, leading=11.8, textColor=CINZA_TEXTO, alignment=TA_JUSTIFY, spaceAfter=4))
    styles.add(ParagraphStyle(name="PopsTableHead", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=colors.white, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="PopsTable", parent=styles["Normal"], fontName="Helvetica", fontSize=7.8, leading=9.7, textColor=CINZA_TEXTO, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="PopsBullet", parent=styles["Normal"], fontName="Helvetica", fontSize=8.15, leading=10.6, textColor=CINZA_TEXTO, leftIndent=8, firstLineIndent=-6, spaceAfter=2))

    def _header_footer(canvas, doc_obj):
        canvas.saveState()
        w, h = A4
        canvas.setFillColor(AZUL_ESCURO)
        canvas.rect(0, h - 1.45 * cm, w, 1.45 * cm, fill=1, stroke=0)
        canvas.setFillColor(DOURADO)
        canvas.rect(0, h - 1.50 * cm, w, 0.06 * cm, fill=1, stroke=0)

        logo = encontrar_logo()
        if logo:
            try:
                canvas.drawImage(str(logo), 1.35 * cm, h - 1.27 * cm, width=2.55 * cm, height=0.92 * cm, preserveAspectRatio=True, mask="auto")
            except Exception:
                pass

        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.drawRightString(w - 1.45 * cm, h - 0.65 * cm, "AQUA GESTÃO — CONTROLE TÉCNICO DE PISCINAS")
        canvas.setFont("Helvetica", 7.5)
        canvas.drawRightString(w - 1.45 * cm, h - 1.02 * cm, "Caderno de POPs | Responsabilidade Técnica | Rotina Operacional")

        canvas.setStrokeColor(BORDA)
        canvas.setLineWidth(0.4)
        canvas.line(1.45 * cm, 1.22 * cm, w - 1.45 * cm, 1.22 * cm)
        canvas.setFillColor(CINZA_TEXTO)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(1.45 * cm, 0.82 * cm, "Aqua Gestão Controle Técnico Ltda | CNPJ 66.008.795/0001-92 | Uberlândia/MG")
        canvas.drawRightString(w - 1.45 * cm, 0.82 * cm, f"Página {doc_obj.page}")
        canvas.restoreState()

    story = []
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph("CADERNO DE POPs DO CONDOMÍNIO", styles["PopsTitle"]))
    story.append(Paragraph("Procedimentos Operacionais Padrão adaptados à rotina operacional informada pela CONTRATANTE.", styles["PopsSub"]))

    dados_table = [
        [Paragraph("<b>Condomínio</b>", styles["PopsTableHead"]), Paragraph("<b>Rotina operacional</b>", styles["PopsTableHead"])],
        [
            Paragraph(f"{nome_cond}<br/>CNPJ: {cnpj}<br/>Endereço: {endereco}<br/>Representante: {sindico} — {cargo_sindico}", styles["PopsTable"]),
            Paragraph(f"Executor informado: {executor}<br/>Frequência operacional: {frequencia}<br/>Piscinas/volumes: {volumes}<br/>Data: {data_ass}", styles["PopsTable"]),
        ],
    ]
    t = Table(dados_table, colWidths=[8.4 * cm, 8.6 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), AZUL_ESCURO),
        ("GRID", (0, 0), (-1, -1), 0.4, BORDA),
        ("BACKGROUND", (0, 1), (-1, 1), CINZA_CLARO),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 6 * mm))

    def p(txt):
        story.append(Paragraph(txt, styles["PopsBody"]))

    story.append(Paragraph("1. Finalidade e delimitação técnica", styles["PopsH"]))
    p("Este caderno integra a documentação técnica de Responsabilidade Técnica do condomínio e estabelece procedimentos mínimos recomendados para a rotina operacional de piscinas de uso coletivo.")
    p("A execução dos procedimentos caberá à equipe operacional designada pela CONTRATANTE, seja funcionário próprio, zelador, piscineiro, empresa terceirizada ou prestador de manutenção, competindo à Aqua Gestão a orientação técnica, supervisão documental e registro de não conformidades no âmbito da Responsabilidade Técnica contratada.")
    p("A Aqua Gestão não assume, por este documento, a execução operacional diária da limpeza, dosagem, manutenção física ou operação contínua dos equipamentos, salvo contratação específica em instrumento próprio.")

    if obs and obs != "Dado não informado":
        story.append(Paragraph("Observação específica informada", styles["PopsMiniH"]))
        p(obs)

    box = Table([[Paragraph("A aplicação destes POPs deve respeitar rótulos, FISPQ/FDS, orientações do Responsável Técnico, condições reais da instalação e normas aplicáveis. Na dúvida, a atividade deve ser interrompida e comunicada ao responsável local e/ou ao RT.", styles["PopsBody"])]], colWidths=[17 * cm])
    box.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), ALERTA),
        ("BOX", (0, 0), (-1, -1), 0.6, DOURADO),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(box)
    story.append(PageBreak())

    def bullet_list(items):
        for item in items:
            story.append(Paragraph(f"• {item}", styles["PopsBullet"]))

    pops = _lista_pops_adaptativos(executor, frequencia)
    for pop in pops:
        bloco = []
        bloco.append(Paragraph(f"{pop['codigo']} — {pop['titulo']}", styles["PopsH"]))
        meta = Table([
            [
                Paragraph("<b>Responsável</b>", styles["PopsTableHead"]),
                Paragraph("<b>Frequência</b>", styles["PopsTableHead"]),
            ],
            [
                Paragraph(pop["responsavel"], styles["PopsTable"]),
                Paragraph(pop["frequencia"], styles["PopsTable"]),
            ],
        ], colWidths=[8.4 * cm, 8.4 * cm])
        meta.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), AZUL_MEDIO),
            ("GRID", (0, 0), (-1, -1), 0.35, BORDA),
            ("BACKGROUND", (0, 1), (-1, 1), CINZA_CLARO),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        bloco.append(meta)
        bloco.append(Spacer(1, 3 * mm))
        bloco.append(Paragraph("<b>Objetivo</b>", styles["PopsMiniH"]))
        bloco.append(Paragraph(pop["objetivo"], styles["PopsBody"]))
        bloco.append(Paragraph("<b>Procedimento mínimo</b>", styles["PopsMiniH"]))
        for item in pop["procedimento"]:
            bloco.append(Paragraph(f"• {item}", styles["PopsBullet"]))
        bloco.append(Paragraph("<b>Cuidados críticos</b>", styles["PopsMiniH"]))
        for item in pop["cuidados"]:
            bloco.append(Paragraph(f"• {item}", styles["PopsBullet"]))
        bloco.append(Paragraph("<b>Registro obrigatório</b>", styles["PopsMiniH"]))
        bloco.append(Paragraph(pop["registros"], styles["PopsBody"]))
        bloco.append(Spacer(1, 4 * mm))
        story.append(KeepTogether(bloco))

    story.append(PageBreak())
    story.append(Paragraph("Termo de recebimento dos POPs", styles["PopsTitle"]))
    p("A CONTRATANTE declara ciência de que recebeu este Caderno de Procedimentos Operacionais Padrão — POPs, comprometendo-se a disponibilizar seu conteúdo à equipe operacional própria ou terceirizada responsável pela rotina da piscina.")
    p("A equipe operacional designada pela CONTRATANTE deverá seguir os procedimentos aqui descritos, registrar as ações executadas e comunicar não conformidades ao síndico/responsável local e/ou ao Responsável Técnico quando aplicável.")
    story.append(Spacer(1, 12 * mm))
    story.append(Paragraph(f"Uberlândia/MG, {data_ass}.", styles["PopsBody"]))
    story.append(Spacer(1, 15 * mm))

    ass = [[
        Paragraph("_________________________________________<br/>AQUA GESTÃO CONTROLE TÉCNICO LTDA<br/>Thyago Fernando da Silveira<br/>CRQ-MG 2ª Região | CRQ 024025748", styles["PopsTable"]),
        Paragraph(f"_________________________________________<br/>{sindico}<br/>{cargo_sindico}<br/>{nome_cond}", styles["PopsTable"]),
    ]]
    t_ass = Table(ass, colWidths=[8.3 * cm, 8.3 * cm])
    t_ass.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(t_ass)

    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    pdf = buf.getvalue()
    buf.close()
    return pdf


def gerar_caderno_pops_pdf() -> tuple[bool, str | None, Path | None]:
    """Gera e salva o Caderno de POPs adaptativo do condomínio."""
    try:
        dados = _dados_pops_rt_do_formulario()
        nome_cond = dados.get("nome_condominio") or "condominio"
        if not dados.get("nome_condominio"):
            return False, "Informe o nome do condomínio antes de gerar os POPs.", None

        pasta = GENERATED_DIR / slugify_nome(nome_cond)
        pasta.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saida = pasta / f"{limpar_nome_arquivo('Caderno_POPs_' + nome_cond + '_' + timestamp)}.pdf"
        pdf_bytes = _gerar_pdf_caderno_pops(dados)
        saida.write_bytes(pdf_bytes)

        ultimos = dict(st.session_state.get("ultimos_docs_gerados") or {})
        ultimos["caderno_pops_pdf"] = str(saida)
        st.session_state.ultimos_docs_gerados = ultimos
        st.session_state.ultima_pasta_gerada = str(pasta)

        try:
            registrar_documento_manifest(
                pasta_condominio=pasta,
                nome_condominio=nome_cond,
                tipo="Caderno de POPs — Procedimentos Operacionais Padrão",
                arquivo_docx=None,
                arquivo_pdf=saida,
                pdf_gerado=True,
                erro_pdf=None,
                dados_utilizados=dados,
                extras={
                    "documento": "caderno_pops",
                    "executor_operacional": dados.get("executor_operacional"),
                    "frequencia_operacional": dados.get("frequencia_operacional"),
                    "pops_adaptativos": True,
                },
            )
        except Exception:
            pass

        return True, None, saida
    except Exception as e:
        return False, str(e), None


# =========================================
# TERMOS DE CIÊNCIA — RT / EPIs
# =========================================
# FUNÇÕES — TERMOS DE CIÊNCIA E SEGURANÇA OPERACIONAL
# =========================================

def _dados_termos_rt_do_formulario() -> dict:
    """Coleta dados atuais do formulário para gerar Termos de Ciência.

    Importante: os Termos usam chaves próprias (termo_*).
    Isso evita erro do Streamlit ao tentar alterar st.session_state de widgets
    do contrato RT depois que eles já foram instanciados na página.
    Quando os campos próprios dos termos estiverem vazios, usa os dados do contrato
    como fallback.
    """
    def _termo_ou_base(chave_termo: str, chave_base: str, padrao: str = "") -> str:
        return str(st.session_state.get(chave_termo) or st.session_state.get(chave_base) or padrao or "").strip()

    return {
        "nome_condominio": _termo_ou_base("termo_nome_condominio", "nome_condominio"),
        "cnpj_condominio": _termo_ou_base("termo_cnpj_condominio", "cnpj_condominio"),
        "endereco_condominio": _termo_ou_base("termo_endereco_condominio", "endereco_condominio"),
        "nome_sindico": _termo_ou_base("termo_nome_sindico", "nome_sindico"),
        "cpf_sindico": _termo_ou_base("termo_cpf_sindico", "cpf_sindico"),
        "cargo_sindico": _termo_ou_base("termo_cargo_sindico", "cargo_sindico", "Síndico"),
        "nome_operador": (st.session_state.get("termo_nome_operador") or "").strip(),
        "cpf_operador": (st.session_state.get("termo_cpf_operador") or "").strip(),
        "funcao_operador": (st.session_state.get("termo_funcao_operador") or "Operador/Zelador").strip(),
        "data_assinatura": (st.session_state.get("termo_data_assinatura") or st.session_state.get("data_assinatura") or hoje_br()).strip(),
        "volumes_piscinas": _termo_ou_base("termo_volumes_piscinas", "volumes_piscinas"),
    }


def _gerar_pdf_termo_ciencia_base(dados: dict, tipo: str) -> bytes:
    """Gera PDF Premium para Termo de Ciência do Síndico ou Operador."""
    import io as _io
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    AZUL_ESCURO = colors.HexColor("#0D2A4A")
    AZUL_MEDIO = colors.HexColor("#1565A8")
    AZUL_CLARO = colors.HexColor("#EAF4FF")
    DOURADO = colors.HexColor("#C8960C")
    CINZA_TEXTO = colors.HexColor("#2F3742")
    CINZA_CLARO = colors.HexColor("#F4F7FA")
    BORDA = colors.HexColor("#D9E2EC")
    ALERTA = colors.HexColor("#FFF8E6")

    def val(chave, padrao="Dado não informado"):
        v = str(dados.get(chave, "") or "").strip()
        return v if v else padrao

    nome_cond = val("nome_condominio")
    cnpj = val("cnpj_condominio")
    endereco = val("endereco_condominio")
    sindico = val("nome_sindico")
    cpf_sindico = val("cpf_sindico")
    cargo_sindico = val("cargo_sindico", "Síndico/Representante legal")
    operador = val("nome_operador", "Operador/Zelador indicado pela CONTRATANTE")
    cpf_operador = val("cpf_operador")
    funcao_operador = val("funcao_operador", "Operador/Zelador")
    data_ass = val("data_assinatura", hoje_br())
    volumes = val("volumes_piscinas", "Volumes não informados")

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.7 * cm,
        rightMargin=1.7 * cm,
        topMargin=2.4 * cm,
        bottomMargin=1.8 * cm,
        title=f"Termo de Ciência — {tipo} — Aqua Gestão",
        author="Aqua Gestão Controle Técnico Ltda",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="TermoTitulo", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=15.5, leading=19, textColor=AZUL_ESCURO, alignment=TA_CENTER, spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="TermoSub", parent=styles["Normal"], fontName="Helvetica",
        fontSize=9, leading=12, textColor=CINZA_TEXTO, alignment=TA_CENTER, spaceAfter=9,
    ))
    styles.add(ParagraphStyle(
        name="TermoH", parent=styles["Heading2"], fontName="Helvetica-Bold",
        fontSize=10, leading=13, textColor=AZUL_ESCURO, spaceBefore=8, spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        name="TermoBody", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.7, leading=12.2, textColor=CINZA_TEXTO, alignment=TA_JUSTIFY, spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="TermoTableHead", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=8.2, leading=10, textColor=colors.white, alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        name="TermoTable", parent=styles["Normal"], fontName="Helvetica",
        fontSize=7.8, leading=9.8, textColor=CINZA_TEXTO, alignment=TA_LEFT,
    ))

    def _header_footer(canvas, doc_obj):
        canvas.saveState()
        w, h = A4
        canvas.setFillColor(AZUL_ESCURO)
        canvas.rect(0, h - 1.45 * cm, w, 1.45 * cm, fill=1, stroke=0)
        canvas.setFillColor(DOURADO)
        canvas.rect(0, h - 1.50 * cm, w, 0.06 * cm, fill=1, stroke=0)

        logo = encontrar_logo()
        if logo:
            try:
                canvas.drawImage(str(logo), 1.35 * cm, h - 1.27 * cm, width=2.55 * cm, height=0.92 * cm, preserveAspectRatio=True, mask="auto")
            except Exception:
                pass

        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.drawRightString(w - 1.45 * cm, h - 0.65 * cm, "AQUA GESTÃO — CONTROLE TÉCNICO DE PISCINAS")
        canvas.setFont("Helvetica", 7.3)
        canvas.drawRightString(w - 1.45 * cm, h - 1.02 * cm, "Termo de Ciência | POPs | Segurança Química | EPIs")

        canvas.setStrokeColor(BORDA)
        canvas.setLineWidth(0.35)
        canvas.line(1.5 * cm, 1.23 * cm, w - 1.5 * cm, 1.23 * cm)
        canvas.setFillColor(CINZA_TEXTO)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(1.5 * cm, 0.84 * cm, "Aqua Gestão Controle Técnico Ltda | CNPJ 66.008.795/0001-92 | Uberlândia/MG")
        canvas.drawRightString(w - 1.5 * cm, 0.84 * cm, f"Página {doc_obj.page}")
        canvas.restoreState()

    story = []

    titulo = "TERMO DE CIÊNCIA DO SÍNDICO / REPRESENTANTE" if tipo == "sindico" else "TERMO DE CIÊNCIA DO OPERADOR / ZELADOR"
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph(titulo, styles["TermoTitulo"]))
    story.append(Paragraph("Responsabilidade Técnica — Piscinas de Uso Coletivo", styles["TermoSub"]))

    id_rows = [
        [Paragraph("CONDOMÍNIO / CONTRATANTE", styles["TermoTableHead"]), Paragraph("DADOS DO RESPONSÁVEL", styles["TermoTableHead"])],
        [
            Paragraph(f"<b>{nome_cond}</b><br/>CNPJ/CPF: {cnpj}<br/>Endereço: {endereco}<br/>Piscinas/volumes: {volumes}", styles["TermoTable"]),
            Paragraph(
                (f"<b>{sindico}</b><br/>CPF: {cpf_sindico}<br/>Qualificação: {cargo_sindico}" if tipo == "sindico" else
                 f"<b>{operador}</b><br/>CPF: {cpf_operador}<br/>Função: {funcao_operador}<br/>Responsável administrativo: {sindico}"),
                styles["TermoTable"],
            ),
        ],
    ]
    t = Table(id_rows, colWidths=[8.6 * cm, 8.4 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), AZUL_ESCURO),
        ("GRID", (0, 0), (-1, -1), 0.35, BORDA),
        ("BACKGROUND", (0, 1), (-1, 1), CINZA_CLARO),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 5 * mm))

    def h(txt):
        story.append(Paragraph(txt, styles["TermoH"]))
    def p(txt):
        story.append(Paragraph(txt, styles["TermoBody"]))
    def bullets(items):
        for item in items:
            story.append(Paragraph(f"• {item}", styles["TermoBody"]))

    if tipo == "sindico":
        h("1. Ciência sobre a Responsabilidade Técnica")
        p("A CONTRATANTE declara ciência de que a Aqua Gestão Controle Técnico Ltda atua no escopo de Responsabilidade Técnica, abrangendo supervisão, orientação, análise técnica, registros, relatórios técnicos, recomendações e comunicação de não conformidades relacionadas ao tratamento químico e ao controle da qualidade da água das piscinas de uso coletivo.")
        p("A ciência firmada neste termo não substitui o contrato principal, mas complementa a documentação técnica do condomínio e reforça as obrigações operacionais, administrativas e de segurança necessárias para a execução adequada da RT.")

        h("2. Obrigações operacionais da CONTRATANTE")
        bullets([
            "manter equipe, operador, zelador ou prestador responsável pela rotina operacional diária das piscinas;",
            "garantir acesso do Responsável Técnico às piscinas, casa de máquinas, registros, produtos e informações necessárias;",
            "executar ou fazer executar as orientações técnicas e POPs emitidos pela Aqua Gestão;",
            "manter registros de parâmetros, dosagens, ocorrências, lavagens de filtro e intervenções realizadas;",
            "comunicar imediatamente anormalidades, água turva, odor forte, suspeita de contaminação, acidente, reclamação de usuário ou risco sanitário;",
            "providenciar produtos, reagentes, equipamentos, instrumentos de medição e condições mínimas de segurança compatíveis com o tratamento da água.",
        ])

        h("3. Ciência sobre EPIs e segurança química")
        p("A CONTRATANTE declara ciência de que a manipulação, transporte, diluição, dosagem e armazenamento de produtos químicos para tratamento de piscinas exige o uso de Equipamentos de Proteção Individual — EPIs adequados ao risco da atividade.")
        p("Compete à CONTRATANTE disponibilizar, fiscalizar e exigir o uso dos EPIs necessários pela equipe operacional própria ou terceirizada, incluindo, quando aplicável: luvas de proteção química, óculos de segurança, máscara ou respirador compatível com o produto utilizado, calçado fechado antiderrapante e vestimenta compatível com a atividade.")
        p("A ausência, insuficiência ou não utilização dos EPIs pela equipe operacional será tratada como não conformidade de segurança, podendo ser registrada pela Responsabilidade Técnica em relatório técnico, checklist ou notificação de não conformidade.")

        h("4. Ciência sobre POPs e não conformidades")
        p("A CONTRATANTE declara ciência de que os Procedimentos Operacionais Padrão — POPs, checklists, relatórios técnicos e notificações emitidos pela Aqua Gestão compõem a rotina técnica mínima recomendada para controle da água e segurança operacional.")
        p("O descumprimento de POPs, a ausência de registros, a falta de produtos, o impedimento de acesso ou a execução de rotinas contrárias à orientação técnica poderão ser registrados em relatório técnico e poderão limitar a responsabilidade da Aqua Gestão quanto aos efeitos decorrentes da conduta operacional da CONTRATANTE ou de seus prepostos.")

        h("5. Restrição de uso ou interdição preventiva")
        p("A CONTRATANTE declara ciência de que, diante de risco sanitário, químico, operacional ou de segurança, o Responsável Técnico poderá recomendar restrição de uso ou interdição técnica preventiva da piscina até a regularização das condições identificadas.")

    else:
        h("1. Ciência sobre a função operacional")
        p("O OPERADOR/ZELADOR declara ciência de que sua atuação na rotina da piscina deve seguir as orientações técnicas, POPs, registros e recomendações fornecidas pela Aqua Gestão e/ou pelo responsável administrativo do condomínio.")
        p("O operador reconhece que não deve alterar dosagens, misturar produtos, improvisar procedimentos ou executar intervenção fora de sua capacitação sem comunicação prévia ao responsável local e/ou ao Responsável Técnico, especialmente em situações de risco químico ou sanitário.")

        h("2. Declaração de ciência sobre uso obrigatório de EPIs")
        p("Declaro estar ciente de que é obrigatório o uso de Equipamentos de Proteção Individual — EPIs durante a manipulação, diluição, dosagem, transporte ou armazenamento de produtos químicos utilizados no tratamento da piscina.")
        p("Declaro estar ciente de que não devo manipular produtos químicos sem os EPIs adequados, especialmente luvas de proteção química, óculos de segurança, máscara ou respirador quando aplicável, calçado fechado antiderrapante e vestimenta compatível com a atividade.")
        p("Declaro estar ciente de que a não utilização de EPIs ou o uso inadequado de produtos químicos pode representar risco à saúde, à segurança operacional e à qualidade da água, devendo a atividade ser interrompida sempre que não houver condição segura para sua execução.")

        h("3. Condutas obrigatórias de segurança química")
        bullets([
            "não misturar produtos químicos entre si;",
            "não adicionar água sobre produto químico concentrado sem orientação técnica específica;",
            "não inalar vapores, poeiras ou gases provenientes de produtos químicos;",
            "não manipular produto sem rótulo, identificação ou orientação de uso;",
            "não armazenar produtos incompatíveis juntos;",
            "manter produtos fechados, identificados, organizados e fora do acesso de usuários;",
            "seguir rótulos, FISPQ/FDS, POPs e recomendações do Responsável Técnico;",
            "comunicar imediatamente vazamentos, derramamentos, odor forte, contato acidental, irritação, queimadura química, ausência de EPI ou condição insegura.",
        ])

        h("4. EPIs mínimos recomendados")
        bullets([
            "luvas de proteção química;",
            "óculos de segurança;",
            "máscara ou respirador compatível com o produto utilizado, quando aplicável;",
            "calçado fechado antiderrapante;",
            "vestimenta compatível com a atividade;",
            "acesso à água corrente para lavagem em caso de contato acidental.",
        ])

        h("5. Registros e comunicação de ocorrências")
        p("O operador declara ciência de que deve registrar os parâmetros avaliados, dosagens realizadas, lavagem/retrolavagem de filtro, anormalidades observadas e qualquer ocorrência que possa comprometer a qualidade da água ou a segurança dos usuários.")
        p("Na presença de água turva, odor forte, cloro combinado elevado, ausência de cloro residual, contaminação fecal, acidente, suspeita de produto inadequado ou qualquer risco sanitário, o operador deverá comunicar imediatamente o síndico/responsável local e aguardar orientação técnica antes de liberar o uso da piscina, quando aplicável.")

    story.append(Spacer(1, 4 * mm))
    aviso = Table([[Paragraph("Este termo deve ser arquivado junto aos documentos de Responsabilidade Técnica do condomínio e apresentado quando solicitado em auditorias, fiscalizações, reuniões internas ou comprovações de ciência operacional.", styles["TermoBody"])]], colWidths=[17 * cm])
    aviso.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), ALERTA),
        ("BOX", (0, 0), (-1, -1), 0.6, DOURADO),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(aviso)
    story.append(Spacer(1, 9 * mm))
    story.append(Paragraph(f"Uberlândia/MG, {data_ass}.", styles["TermoBody"]))
    story.append(Spacer(1, 15 * mm))

    if tipo == "sindico":
        ass = [[
            Paragraph("_________________________________________<br/>AQUA GESTÃO CONTROLE TÉCNICO LTDA<br/>Thyago Fernando da Silveira<br/>CRQ-MG 2ª Região | CRQ 024025748", styles["TermoTable"]),
            Paragraph(f"_________________________________________<br/>{sindico}<br/>{cargo_sindico}<br/>{nome_cond}", styles["TermoTable"]),
        ]]
    else:
        ass = [[
            Paragraph("_________________________________________<br/>AQUA GESTÃO CONTROLE TÉCNICO LTDA<br/>Thyago Fernando da Silveira<br/>CRQ-MG 2ª Região | CRQ 024025748", styles["TermoTable"]),
            Paragraph(f"_________________________________________<br/>{operador}<br/>{funcao_operador}<br/>CPF: {cpf_operador}", styles["TermoTable"]),
        ]]

    t_ass = Table(ass, colWidths=[8.3 * cm, 8.3 * cm])
    t_ass.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(t_ass)

    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    pdf = buf.getvalue()
    buf.close()
    return pdf


def gerar_termo_ciencia_pdf(tipo: str) -> tuple[bool, str | None, Path | None]:
    """Gera e salva Termo de Ciência de síndico ou operador."""
    try:
        dados = _dados_termos_rt_do_formulario()
        nome_cond = dados.get("nome_condominio") or "condominio"
        if not dados.get("nome_condominio"):
            return False, "Informe o nome do condomínio antes de gerar o termo.", None
        if tipo == "operador" and not dados.get("nome_operador"):
            return False, "Informe o nome do operador/zelador antes de gerar o termo.", None

        pasta = GENERATED_DIR / slugify_nome(nome_cond)
        pasta.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_tipo = "Termo_Ciencia_Sindico" if tipo == "sindico" else "Termo_Ciencia_Operador_EPI"
        saida = pasta / f"{limpar_nome_arquivo(nome_tipo + '_' + nome_cond + '_' + timestamp)}.pdf"
        pdf_bytes = _gerar_pdf_termo_ciencia_base(dados, tipo)
        saida.write_bytes(pdf_bytes)

        ultimos = dict(st.session_state.get("ultimos_docs_gerados") or {})
        if tipo == "sindico":
            ultimos["termo_sindico_pdf"] = str(saida)
        else:
            ultimos["termo_operador_pdf"] = str(saida)
        st.session_state.ultimos_docs_gerados = ultimos
        st.session_state.ultima_pasta_gerada = str(pasta)

        try:
            registrar_documento_manifest(
                pasta_condominio=pasta,
                nome_condominio=nome_cond,
                tipo="Termo de Ciência — Síndico" if tipo == "sindico" else "Termo de Ciência — Operador/EPIs",
                arquivo_docx=None,
                arquivo_pdf=saida,
                pdf_gerado=True,
                erro_pdf=None,
                dados_utilizados=dados,
                extras={"documento": "termo_ciencia", "tipo": tipo, "epi_sem_avental_padrao": True},
            )
        except Exception:
            pass

        return True, None, saida
    except Exception as e:
        return False, str(e), None

# =========================================

st.markdown('<div class="section-card aq-only" id="sec-termos-ciencia">', unsafe_allow_html=True)
st.subheader("Termos de ciência e segurança operacional")
st.caption(
    "Gere termos complementares para arquivar junto ao contrato de RT. "
    "Os termos reforçam ciência sobre POPs, registros operacionais, comunicação de não conformidades, segurança química e uso de EPIs."
)

# ── Carregamento de dados do condomínio também nos Termos de Ciência ──────────
# Os termos usam os mesmos campos-base do contrato RT. Este bloco evita digitação
# duplicada quando o usuário acessa diretamente a seção de termos.
_clientes_termos = filtrar_clientes_por_empresa((_clientes_completos_cache() if "_clientes_completos_cache" in globals() else sheets_listar_clientes_completo()), "aqua_gestao")
if _clientes_termos:
    _opcoes_termos = ["— Selecionar cliente cadastrado —"] + [c.get("nome", "") for c in _clientes_termos if c.get("nome")]
    _sel_termos = st.selectbox(
        "🔗 Carregar dados do condomínio cadastrado",
        _opcoes_termos,
        key="sel_cliente_termos_ciencia",
        help="Preenche os dados do condomínio, CNPJ, endereço e representante para gerar os termos."
    )
    if _sel_termos and _sel_termos != "— Selecionar cliente cadastrado —":
        if st.button("⬇️ Preencher termos com dados deste condomínio", key="btn_carregar_cliente_termos", use_container_width=True):
            _dados_termos = next((c for c in _clientes_termos if c.get("nome") == _sel_termos), {})
            if _dados_termos:
                st.session_state["termo_nome_condominio"] = _dados_termos.get("nome", "")
                st.session_state["termo_cnpj_condominio"] = formatar_cnpj(_dados_termos.get("cnpj", ""))
                st.session_state["termo_endereco_condominio"] = _dados_termos.get("endereco", "")
                st.session_state["termo_nome_sindico"] = _dados_termos.get("contato", "")
                st.session_state["termo_whatsapp_cliente"] = formatar_telefone(_dados_termos.get("telefone", ""))
                st.session_state["termo_email_cliente"] = _dados_termos.get("email", "")

                _vols = []
                try:
                    if float(_dados_termos.get("vol_adulto") or 0) > 0:
                        _vols.append(f"Piscina adulto: {_dados_termos.get('vol_adulto')} m³")
                    if float(_dados_termos.get("vol_infantil") or 0) > 0:
                        _vols.append(f"Piscina infantil: {_dados_termos.get('vol_infantil')} m³")
                    if float(_dados_termos.get("vol_family") or 0) > 0:
                        _vols.append(f"Piscina family: {_dados_termos.get('vol_family')} m³")
                except Exception:
                    _vols = []
                if _vols:
                    st.session_state["termo_volumes_piscinas"] = " | ".join(_vols)

                st.success(f"✅ Dados de '{_sel_termos}' carregados para os termos.")
                st.rerun()
else:
    st.info("💡 Cadastre o condomínio ou preencha os dados do contrato RT antes de gerar os termos.")

with st.expander("📌 Dados do condomínio usados nos termos", expanded=False):
    _dados_preview_termos = _dados_termos_rt_do_formulario()
    st.write(f"**Condomínio:** {_dados_preview_termos.get('nome_condominio') or 'Não informado'}")
    st.write(f"**CNPJ:** {_dados_preview_termos.get('cnpj_condominio') or 'Não informado'}")
    st.write(f"**Endereço:** {_dados_preview_termos.get('endereco_condominio') or 'Não informado'}")
    st.write(f"**Representante:** {_dados_preview_termos.get('nome_sindico') or 'Não informado'}")
    st.write(f"**Volumes:** {_dados_preview_termos.get('volumes_piscinas') or 'Não informado'}")

_tc1, _tc2, _tc3 = st.columns([1.25, 1.1, 1.1])
with _tc1:
    st.text_input(
        "Nome do operador/zelador",
        key="termo_nome_operador",
        placeholder="Ex.: José da Silva",
    )
with _tc2:
    st.text_input(
        "CPF do operador/zelador",
        key="termo_cpf_operador",
        placeholder="000.000.000-00",
    )
with _tc3:
    st.text_input(
        "Função",
        key="termo_funcao_operador",
        value=st.session_state.get("termo_funcao_operador", "Operador/Zelador"),
    )

st.info(
    "EPIs mínimos nos termos: luvas de proteção química, óculos de segurança, máscara ou respirador quando aplicável, "
    "calçado fechado antiderrapante, vestimenta compatível e acesso à água corrente. O avental não foi incluído como item padrão."
)

_tbtn1, _tbtn2 = st.columns(2)
with _tbtn1:
    if st.button("📄 Gerar Termo de Ciência do Síndico", use_container_width=True):
        with st.spinner("Gerando Termo de Ciência do Síndico..."):
            ok, erro, caminho = gerar_termo_ciencia_pdf("sindico")
        if ok and caminho and caminho.exists():
            st.success("Termo de Ciência do Síndico gerado com sucesso.")
            with open(caminho, "rb") as f:
                st.download_button(
                    "Baixar Termo do Síndico em PDF",
                    data=f,
                    file_name=caminho.name,
                    mime="application/pdf",
                    use_container_width=True,
                    key="dl_termo_sindico_imediato",
                )
        else:
            st.error(f"Não foi possível gerar o termo: {erro}")

with _tbtn2:
    if st.button("👷 Gerar Termo de Ciência do Operador — EPIs", use_container_width=True):
        with st.spinner("Gerando Termo de Ciência do Operador..."):
            ok, erro, caminho = gerar_termo_ciencia_pdf("operador")
        if ok and caminho and caminho.exists():
            st.success("Termo de Ciência do Operador gerado com sucesso.")
            with open(caminho, "rb") as f:
                st.download_button(
                    "Baixar Termo do Operador em PDF",
                    data=f,
                    file_name=caminho.name,
                    mime="application/pdf",
                    use_container_width=True,
                    key="dl_termo_operador_imediato",
                )
        else:
            st.error(f"Não foi possível gerar o termo: {erro}")

st.markdown("</div>", unsafe_allow_html=True)

# =========================================
# POPs — PROCEDIMENTOS OPERACIONAIS PADRÃO
# =========================================

st.markdown('<div class="section-card aq-only" id="sec-pops-adaptativos">', unsafe_allow_html=True)
st.subheader("POPs — Procedimentos Operacionais Padrão")
st.caption(
    "Gere um Caderno de POPs no padrão Aqua Gestão, adaptado ao tipo de operação do condomínio. "
    "Indicado para condomínios com prestador externo, zelador do condomínio, funcionário próprio, empresa contratada ou rotina mista."
)

# Normaliza valores antigos/estranhos salvos no session_state para não manter textos quebrados na tela.
_opcoes_executor_pops = [
    "Prestador externo",
    "Empresa contratada",
    "Zelador do condomínio",
    "Funcionário próprio do condomínio",
    "Rotina mista",
]
_mapa_executor_pops_antigo = {
    "Prestador externo": "Prestador externo",
    "Prestador ado": "Prestador externo",
    "Empresa contratada": "Empresa contratada",
    "Zelador": "Zelador do condomínio",
}
_valor_executor_atual = st.session_state.get("pops_executor_operacional", "Prestador externo")
st.session_state["pops_executor_operacional"] = _mapa_executor_pops_antigo.get(_valor_executor_atual, _valor_executor_atual)
if st.session_state["pops_executor_operacional"] not in _opcoes_executor_pops:
    st.session_state["pops_executor_operacional"] = "Prestador externo"

_popc1, _popc2 = st.columns(2)
with _popc1:
    st.selectbox(
        "Quem executa a rotina operacional da piscina?",
        _opcoes_executor_pops,
        key="pops_executor_operacional",
        index=_opcoes_executor_pops.index(st.session_state["pops_executor_operacional"]),
    )

with _popc2:
    st.selectbox(
        "Frequência operacional informada",
        [
            "3 vezes por semana",
            "Diária",
            "2 vezes por semana",
            "1 vez por semana",
            "Sem rotina definida",
        ],
        key="pops_frequencia_operacional",
        index=[
            "3 vezes por semana",
            "Diária",
            "2 vezes por semana",
            "1 vez por semana",
            "Sem rotina definida",
        ].index(st.session_state.get("pops_frequencia_operacional", "3 vezes por semana"))
        if st.session_state.get("pops_frequencia_operacional", "3 vezes por semana") in [
            "3 vezes por semana",
            "Diária",
            "2 vezes por semana",
            "1 vez por semana",
            "Sem rotina definida",
        ] else 0,
    )

st.text_area(
    "Observação específica para este caderno de POPs (opcional)",
    key="pops_observacao",
    placeholder="Ex.: O condomínio possui prestador externo 3 vezes por semana e responsável local para verificação visual nos demais dias.",
    height=80,
)

st.info(
    "O caderno deixa claro que a execução operacional cabe à equipe designada pela CONTRATANTE, "
    "enquanto a equipe técnica da Aqua Gestão atua na orientação técnica, supervisão documental e registro de não conformidades no âmbito da RT."
)

if st.button("📘 Gerar Caderno de POPs do Condomínio", use_container_width=True):
    with st.spinner("Gerando Caderno de POPs adaptativo..."):
        ok, erro, caminho = gerar_caderno_pops_pdf()
    if ok and caminho and caminho.exists():
        st.success("Caderno de POPs gerado com sucesso.")
        with open(caminho, "rb") as f:
            st.download_button(
                "Baixar Caderno de POPs em PDF",
                data=f,
                file_name=caminho.name,
                mime="application/pdf",
                use_container_width=True,
                key="dl_caderno_pops_imediato",
            )
    else:
        st.error(f"Não foi possível gerar os POPs: {erro}")

st.markdown("</div>", unsafe_allow_html=True)


# =========================================
# DOWNLOADS DOS ÚLTIMOS DOCUMENTOS GERADOS
# =========================================
# Exibe os downloads logo aqui, antes do relatório, independente de onde o botão foi clicado.

_ultimos = st.session_state.get("ultimos_docs_gerados")
if _ultimos:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader("⬇️ Últimos documentos gerados")

    _dc1, _dc2 = st.columns(2)
    with _dc1:
        st.markdown("**Contrato RT — Aqua Gestão**")
        _p = _ultimos.get("contrato_docx")
        if _p and Path(_p).exists():
            with open(_p, "rb") as _f:
                st.download_button(
                    "Baixar DOCX editável do contrato",
                    data=_f,
                    file_name=Path(_p).name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_contrato_docx_top",
                )
    with _dc2:
        st.markdown("**PDF para envio**")
        _p = _ultimos.get("contrato_pdf_premium")
        if _p and Path(_p).exists():
            with open(_p, "rb") as _f:
                st.download_button(
                    "📄 Baixar PDF Premium Aqua Gestão",
                    data=_f,
                    file_name=Path(_p).name,
                    mime="application/pdf",
                    use_container_width=True,
                    key="dl_contrato_pdf_premium_top",
                )

    _p_termo_sindico = _ultimos.get("termo_sindico_pdf")
    _p_termo_operador = _ultimos.get("termo_operador_pdf")
    if (_p_termo_sindico and Path(_p_termo_sindico).exists()) or (_p_termo_operador and Path(_p_termo_operador).exists()):
        st.markdown("**Termos de ciência — RT / EPIs**")
        _dt1, _dt2 = st.columns(2)
        with _dt1:
            if _p_termo_sindico and Path(_p_termo_sindico).exists():
                with open(_p_termo_sindico, "rb") as _f:
                    st.download_button(
                        "Baixar Termo do Síndico",
                        data=_f,
                        file_name=Path(_p_termo_sindico).name,
                        mime="application/pdf",
                        use_container_width=True,
                        key="dl_termo_sindico_top",
                    )
        with _dt2:
            if _p_termo_operador and Path(_p_termo_operador).exists():
                with open(_p_termo_operador, "rb") as _f:
                    st.download_button(
                        "Baixar Termo do Operador — EPIs",
                        data=_f,
                        file_name=Path(_p_termo_operador).name,
                        mime="application/pdf",
                        use_container_width=True,
                        key="dl_termo_operador_top",
                    )

    _p_pops = _ultimos.get("caderno_pops_pdf")
    if _p_pops and Path(_p_pops).exists():
        st.markdown("**POPs — Procedimentos Operacionais Padrão**")
        with open(_p_pops, "rb") as _f:
            st.download_button(
                "📘 Baixar Caderno de POPs",
                data=_f,
                file_name=Path(_p_pops).name,
                mime="application/pdf",
                use_container_width=True,
                key="dl_caderno_pops_top",
            )


    _p_rel = _ultimos.get("relatorio_docx")
    _p_rel_pdf = _ultimos.get("relatorio_pdf")
    if (_p_rel and Path(_p_rel).exists()) or (_p_rel_pdf and Path(_p_rel_pdf).exists()):
        st.markdown("**Relatório mensal**")
        _dc3, _dc4 = st.columns(2)
        with _dc3:
            if _p_rel and Path(_p_rel).exists():
                with open(_p_rel, "rb") as _f:
                    st.download_button("Baixar DOCX do relatório", data=_f, file_name=Path(_p_rel).name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="dl_relatorio_docx_top")
        with _dc4:
            if _p_rel_pdf and Path(_p_rel_pdf).exists():
                with open(_p_rel_pdf, "rb") as _f:
                    st.download_button("Baixar PDF do relatório", data=_f, file_name=Path(_p_rel_pdf).name,
                        mime="application/pdf", use_container_width=True, key="dl_relatorio_pdf_top")

    # _EMAIL_ULTIMOS_DOCUMENTOS_AQUA_V1_
    _email_docs_pasta = None
    try:
        _pasta_tmp = st.session_state.get("ultima_pasta_gerada") or ""
        if _pasta_tmp:
            _email_docs_pasta = Path(_pasta_tmp)
    except Exception:
        _email_docs_pasta = None

    try:
        _dados_email_docs = _dados_termos_rt_do_formulario()
    except Exception:
        _dados_email_docs = {}

    _nome_email_docs = (
        _dados_email_docs.get("nome_condominio")
        or st.session_state.get("nome_condominio")
        or "Condomínio"
    )
    _email_cliente_docs = (
        st.session_state.get("termo_email_cliente")
        or st.session_state.get("email_cliente")
        or ""
    )
    _msg_email_docs = (
        f"Prezados,\n\n"
        f"Encaminho em anexo a documentação técnica gerada pela Aqua Gestão referente ao {_nome_email_docs}.\n\n"
        "Incluí os documentos selecionados no sistema, podendo contemplar contrato, aditivo, termos de ciência, POPs e relatórios técnicos.\n\n"
        "Os arquivos seguem para conferência, registro e arquivo interno do condomínio.\n\n"
        "Permaneço à disposição para qualquer esclarecimento."
    )
    exibir_envio_email_documentos_aqua(
        nome_condominio=_nome_email_docs,
        pasta_condominio=_email_docs_pasta,
        email_cliente=_email_cliente_docs,
        mensagem_padrao=_msg_email_docs,
        documentos_sugeridos=[v for v in (_ultimos or {}).values() if isinstance(v, str)],
        key_prefix="ultimos_docs_aqua",
    )

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================
# RELATÓRIO MENSAL DE RT
# =========================================


st.markdown('<div class="section-card" id="sec-preview-relatorio">', unsafe_allow_html=True)
st.subheader("👁️ Pré-visualizar relatório final")
st.caption("A prévia usa a empresa ativa no acesso administrativo, mantendo Aqua Gestão e Bem Star separadas dentro do sistema.")

# _PREVIEW_RELATORIO_SEMPRE_AQUA_V1_
_prev_empresa_val = "Aqua Gestão"
if st.session_state.get("empresa_ativa", "aqua_gestao") == "bem_star":
    st.warning("🛡️ Prévia de Relatório RT protegida: usando Aqua Gestão, mesmo que o painel lateral esteja em Bem Star.")
st.info(f"Empresa ativa para esta prévia: {_prev_empresa_val}")
_prev_usar_form = st.checkbox(
    "Usar dados reais do formulário e fotos anexadas (prévia exata)",
    value=True,
    key="preview_rel_usar_form",
    help="Quando marcado, a prévia usa o mesmo gerador DOCX/PDF do relatório final, aproveitando os dados atuais do formulário e as fotos anexadas disponíveis.",
)

_prev_tab1, _prev_tab2, _prev_tab3 = st.tabs(["📄 Prévia exata do relatório final", "🧩 Modelo visual de referência", "⬆️ Carregar Relatório"])

with _prev_tab1:
    if _prev_usar_form:
        with st.spinner("Montando prévia exata com os dados atuais..."):
            _prev_result = gerar_previa_exata_relatorio(_prev_empresa_val)

        if _prev_result.get("ok"):
            _pdf_ok = bool(_prev_result.get("pdf_ok")) and _prev_result.get("pdf") and Path(_prev_result["pdf"]).exists()
            _docx_ok = _prev_result.get("docx") and Path(_prev_result["docx"]).exists()
            st.success(f"✅ {_prev_result.get('mensagem', 'Prévia exata atualizada com sucesso.')}")
            st.caption(f"Fotos usadas: {len(_prev_result.get('fotos', []))} · origem: {_prev_result.get('origem_fotos', 'não identificada')}")

            if _pdf_ok:
                exibir_pdf_previa_exata(Path(_prev_result["pdf"]), height=1220 if _prev_empresa_val == "Aqua Gestão" else 1320)
            else:
                st.warning(f"O DOCX da prévia foi gerado, mas o PDF não foi convertido automaticamente. Erro: {_prev_result.get('erro_pdf')}")

            _col_prev_1, _col_prev_2 = st.columns(2)
            with _col_prev_1:
                if _docx_ok:
                    with open(_prev_result["docx"], "rb") as _f:
                        st.download_button(
                            "⬇️ Baixar DOCX da prévia exata",
                            data=_f,
                            file_name=Path(_prev_result["docx"]).name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="btn_dl_previa_exata_docx",
                        )
            with _col_prev_2:
                if _pdf_ok:
                    with open(_prev_result["pdf"], "rb") as _f:
                        st.download_button(
                            "⬇️ Baixar PDF da prévia exata",
                            data=_f,
                            file_name=Path(_prev_result["pdf"]).name,
                            mime="application/pdf",
                            use_container_width=True,
                            key="btn_dl_previa_exata_pdf",
                        )
        else:
            st.warning("Não foi possível montar a prévia exata com os dados atuais.")
            for _erro_prev in _prev_result.get("erros", []) or [_prev_result.get("mensagem", "")]:
                if _erro_prev:
                    st.write(f"- {_erro_prev}")
    else:
        st.info("Ative a opção de prévia exata para usar o mesmo gerador DOCX/PDF do relatório final com os dados reais do formulário.")

with _prev_tab2:
    _prev_dados, _, _ = _obter_dados_preview_relatorio(_prev_empresa_val, False)
    st.info("Exibindo o modelo visual de referência como apoio rápido de layout. A aba ao lado mostra a prévia exata do documento final.")
    _prev_html = gerar_mockup_relatorio_preview_html(_prev_empresa_val, visual="web", dados=_prev_dados)
    _prev_print = gerar_mockup_relatorio_preview_html(_prev_empresa_val, visual="print", dados=_prev_dados)
    _sub_tab1, _sub_tab2 = st.tabs(["🌐 Referência tela / HTML", "🖨️ Referência impressão / PDF"])
    with _sub_tab1:
        st.warning("Visualização HTML embutida desativada para evitar rerun infinito no Streamlit 1.56.")
        st.code(_prev_html[:12000], language="html")
        if len(_prev_html) > 12000:
            st.caption("Prévia textual truncada no app. Use o download abaixo para abrir o HTML completo.")
        st.download_button(
            "⬇️ Baixar HTML de referência (tela)",
            data=_prev_html.encode("utf-8"),
            file_name=f"mockup_relatorio_{slugify_nome(_prev_empresa_val)}_tela.html",
            mime="text/html",
            use_container_width=True,
            key="btn_dl_mockup_rel_tela",
        )
    with _sub_tab2:
        st.warning("Visualização de impressão embutida desativada para evitar rerun infinito no Streamlit 1.56.")
        st.code(_prev_print[:12000], language="html")
        if len(_prev_print) > 12000:
            st.caption("Prévia textual truncada no app. Use o download abaixo para abrir o HTML completo.")
        st.download_button(
            "⬇️ Baixar HTML de referência (impressão)",
            data=_prev_print.encode("utf-8"),
            file_name=f"mockup_relatorio_{slugify_nome(_prev_empresa_val)}_impressao.html",
            mime="text/html",
            use_container_width=True,
            key="btn_dl_mockup_rel_print",
        )

with _prev_tab3:
    st.markdown("### Carregar relatório para edição")
    st.caption("Selecione um arquivo JSON de relatório salvo para preencher o formulário e editar.")
    uploaded_file = st.file_uploader("Escolha um arquivo JSON", type=["json"], key="upload_relatorio_json")
    if uploaded_file is not None:
        try:
            dados_relatorio = json.load(uploaded_file)
            aplicar_snapshot_relatorio_independente(dados_relatorio)
            st.success("✅ Relatório carregado com sucesso! Os campos do formulário foram preenchidos.")
            st.rerun()
        except Exception as e:
            st.error(f"❌ Erro ao carregar o arquivo JSON: {e}")

st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="section-card aq-only" id="sec-relatorio-rt">', unsafe_allow_html=True)
st.subheader("Relatório mensal de responsabilidade técnica")


st.caption(f"Dados fixos automáticos do RT: {RESPONSAVEL_TECNICO_ASSINATURA} | Certificações {CERTIFICACOES_RT}")

# ── Seletor de cliente do Sheets ──────────────────────────────────────────────
@st.cache_data(ttl=60)
def _clientes_completos_rel_cache():
    return sheets_listar_clientes_completo()

_clientes_rel = filtrar_clientes_por_empresa(_clientes_completos_rel_cache(), "aqua_gestao")
if _clientes_rel:
    _opcoes_rel = ["— Selecionar cliente cadastrado —"] + [c["nome"] for c in _clientes_rel]
    _rel_col1, _rel_col2 = st.columns([3, 1])
    with _rel_col1:
        _sel_rel = st.selectbox(
            "🔗 Carregar dados de cliente cadastrado",
            _opcoes_rel,
            key="sel_cliente_rel",
            help="Selecione um cliente para preencher automaticamente os campos do relatório."
        )
    with _rel_col2:
        st.markdown("<div style='margin-top:28px'>", unsafe_allow_html=True)
        if st.button("⬇️ Carregar", key="btn_carregar_cliente_rel", use_container_width=True):
            if _sel_rel and _sel_rel != "— Selecionar cliente cadastrado —":
                _dados_rel = next((c for c in _clientes_rel if c["nome"] == _sel_rel), {})
                if _dados_rel:
                    st.session_state["rel_nome_condominio"]   = _dados_rel.get("nome", "")
                    st.session_state["rel_cnpj_condominio"]   = formatar_cnpj(_dados_rel.get("cnpj", ""))
                    st.session_state["rel_endereco_condominio"] = _dados_rel.get("endereco", "")
                    st.session_state["rel_representante"]     = _dados_rel.get("contato", "")
                    st.session_state["rel_cpf_cnpj_representante"] = ""
                    _freq_rel = obter_verificacoes_semanais_cliente(_dados_rel)
                    st.session_state["rel_verificacoes_semanais"] = _freq_rel
                    st.session_state["rel_analises_total"] = calcular_linhas_analises_por_frequencia(_freq_rel)
                    st.session_state["_rel_auto_importar_cliente"] = True
                    st.success(f"✅ Dados de '{_sel_rel}' carregados no relatório! As visitas do mês serão importadas automaticamente se existirem.")
                    st.rerun()

rr0a, rr0b, rr0c = st.columns([1.1, 1.2, 1.1])
with rr0a:
    st.selectbox("Tipo de atendimento", ["Contrato ativo", "Visita técnica avulsa", "Inspeção técnica", "Acompanhamento sem contrato"], key="rel_tipo_atendimento")
with rr0b:
    if st.button("Carregar dados do cadastro no relatório", use_container_width=True):
        carregar_dados_cadastro_no_relatorio()
        st.rerun()
with rr0c:
    st.checkbox("Salvar alterações deste relatório no cadastro principal", key="rel_salvar_alteracoes_cadastro")

# ---- Importar lançamentos de campo (local + Google Sheets) ----
nome_rel_atual = (st.session_state.get("rel_nome_condominio") or st.session_state.get("nome_condominio") or "").strip()
pasta_rel_atual = GENERATED_DIR / slugify_nome(nome_rel_atual) if nome_rel_atual else None

# Busca lançamentos locais
lancamentos_local = []
if pasta_rel_atual and pasta_rel_atual.exists():
    dados_rel_json = carregar_dados_condominio(pasta_rel_atual)
    lancamentos_local = (dados_rel_json or {}).get("lancamentos_campo", [])

# Busca lançamentos do Google Sheets (aba Visitas)
lancamentos_sheets = []
if nome_rel_atual:
    lancamentos_sheets = sheets_listar_lancamentos(nome_rel_atual)

# Filtra pelo mês de referência
mes_ref = (st.session_state.get("rel_mes_referencia") or "").strip()
ano_ref = (st.session_state.get("rel_ano_referencia") or str(datetime.now().year)).strip()

# Diagnóstico de importação de visitas — ajuda a separar problema de filtro vs. problema de gravação no Sheets.
with st.expander("🧪 Diagnóstico de visitas importadas", expanded=False):
    st.write("**Condomínio usado na busca:**", nome_rel_atual or "Não informado")
    st.write("**Mês/Ano do relatório:**", f"{mes_ref or 'mês não informado'}/{ano_ref or 'ano não informado'}")
    st.write("**Lançamentos locais encontrados:**", len(lancamentos_local))
    st.write("**Lançamentos no Google Sheets encontrados para este condomínio:**", len(lancamentos_sheets))

    if lancamentos_sheets:
        st.write("**Últimas datas encontradas no Sheets:**")
        for lc in lancamentos_sheets[-10:]:
            st.write(
                f"- {lc.get('data', '')} | "
                f"{lc.get('condominio', '')} | "
                f"Operador: {lc.get('operador', '')} | "
                f"ID: {lc.get('id_visita', '')}"
            )
    else:
        st.warning(
            "Nenhuma visita foi encontrada na aba 🔬 Visitas para este condomínio. "
            "Se existe apenas PDF, o relatório mensal não consegue importar automaticamente. "
            "A visita precisa estar salva no Google Sheets."
        )

    erro_sheets = st.session_state.get("_sheets_ultimo_erro", "")
    if erro_sheets:
        st.code(erro_sheets[:1500])

    st.markdown("---")
    st.markdown("**📄 Recuperar visita antiga a partir de PDF**")
    st.caption(
        "Use isto quando o PDF da visita existe, mas a linha não foi gravada na aba 🔬 Visitas. "
        "O sistema extrai os dados do PDF, salva no Google Sheets e libera a importação no relatório mensal."
    )
    pdf_visita_import = st.file_uploader(
        "Enviar PDF do Relatório de Visita para recuperar lançamento",
        type=["pdf"],
        key="rel_pdf_visita_importador",
    )

    if pdf_visita_import is not None:
        if st.button("📥 Importar este PDF para a aba 🔬 Visitas", key="btn_importar_pdf_visita_sheets"):
            try:
                pdf_bytes = pdf_visita_import.getvalue()
                lanc_pdf = extrair_lancamento_de_pdf_visita(pdf_bytes, nome_rel_atual)
                if not lanc_pdf:
                    st.error("Não consegui extrair dados suficientes deste PDF. Confira se é um Relatório de Visita gerado pelo sistema.")
                else:
                    cond_pdf = lanc_pdf.get("condominio") or nome_rel_atual
                    ok_pdf = sheets_salvar_lancamento_campo(lanc_pdf, cond_pdf)

                    try:
                        pasta_pdf = GENERATED_DIR / slugify_nome(cond_pdf)
                        pasta_pdf.mkdir(exist_ok=True)
                        dados_pdf = carregar_dados_condominio(pasta_pdf) or {}
                        pend_pdf = dados_pdf.get("lancamentos_campo", [])
                        pend_pdf.append(lanc_pdf)
                        dados_pdf["lancamentos_campo"] = pend_pdf
                        dados_pdf["nome_condominio"] = dados_pdf.get("nome_condominio", cond_pdf)
                        salvar_dados_condominio(pasta_pdf, dados_pdf)
                    except Exception as e:
                        _log_sheets_erro("salvar_pdf_importado_json_local", e)

                    if ok_pdf:
                        st.success(
                            f"✅ PDF importado e salvo como visita: {cond_pdf} — {lanc_pdf.get('data','data não identificada')}"
                        )
                        st.session_state["_rel_auto_importar_cliente"] = True
                        st.rerun()
                    else:
                        st.warning(
                            "Consegui ler o PDF, mas não consegui salvar na aba 🔬 Visitas. "
                            "Veja o erro técnico abaixo."
                        )
                        erro_pdf = st.session_state.get("_sheets_ultimo_erro", "")
                        if erro_pdf:
                            st.code(erro_pdf[:2000])
            except Exception as e:
                _log_sheets_erro("importar_pdf_visita_ui", e)
                st.error("Erro ao importar o PDF de visita.")
                st.code(st.session_state.get("_sheets_ultimo_erro", "")[:2000])

def _filtrar_mes(lancamentos, mes, ano):
    """Filtra lançamentos pelo mês/ano, aceitando vários formatos de data."""
    if not mes or not ano:
        return lancamentos
    return [lc for lc in lancamentos if lancamento_pertence_mes_ano(lc.get("data", ""), mes, ano)]

# Une local + Sheets sem duplicar (por data+operador)
_vistos = set()
lancamentos_disponiveis = []
for lc in lancamentos_local + lancamentos_sheets:
    _chave = f"{lc.get('data','')}-{lc.get('operador','')}-{lc.get('ph','')}"
    if _chave not in _vistos:
        _vistos.add(_chave)
        lancamentos_disponiveis.append(lc)

# Filtra por mês se informado
lancamentos_disponiveis = _filtrar_mes(lancamentos_disponiveis, mes_ref, ano_ref)

def _importar_lancamentos(lancamentos):
    """Preenche o relatório com os lançamentos de campo."""
    _freq_base = st.session_state.get("rel_verificacoes_semanais", 3)
    _linhas_base = calcular_linhas_analises_por_frequencia(_freq_base, st.session_state.get("rel_mes_referencia"), st.session_state.get("rel_ano_referencia"))
    garantir_campos_analises(max(len(lancamentos), _linhas_base, ANALISES_PADRAO))
    for i, lc in enumerate(lancamentos[:ANALISES_MAX_SUGERIDO]):
        # Suporte a múltiplas piscinas — usa dados da primeira piscina ou direto
        piscinas = lc.get("piscinas", [])
        if piscinas:
            lc_dados = piscinas[0]  # primeira piscina para o relatório principal
        else:
            lc_dados = lc
        st.session_state[f"rel_analise_data_{i}"]     = lc.get("data", "")
        st.session_state[f"rel_analise_ph_{i}"]        = lc_dados.get("ph", lc.get("ph",""))
        st.session_state[f"rel_analise_cl_{i}"]        = lc_dados.get("cloro_livre", lc.get("cloro_livre",""))
        st.session_state[f"rel_analise_ct_{i}"]        = lc_dados.get("cloro_total", lc.get("cloro_total",""))
        st.session_state[f"rel_analise_alc_{i}"]       = lc_dados.get("alcalinidade", lc.get("alcalinidade",""))
        st.session_state[f"rel_analise_dc_{i}"]        = lc_dados.get("dureza", lc.get("dureza",""))
        st.session_state[f"rel_analise_cya_{i}"]       = lc_dados.get("cianurico", lc.get("cianurico",""))
        st.session_state[f"rel_analise_operador_{i}"]  = lc.get("operador", "")
    # Preenche campo operador responsavel com o mais frequente dos lancamentos
    _ops_import = [lc.get("operador","").strip() for lc in lancamentos if lc.get("operador","").strip()]
    if _ops_import:
        _op_mais_freq = max(set(_ops_import), key=_ops_import.count)
        if not st.session_state.get("csr_operador_rel","").strip():
            st.session_state["csr_operador_rel"] = _op_mais_freq

    # Importa dosagens da última visita com dosagem registrada
    for lc in reversed(lancamentos):
        if lc.get("dosagens"):
            aplicar_dosagens_ultimas_no_relatorio({"dosagens_ultimas": lc["dosagens"]})
            break
    # Concatena observações e problemas
    obs_lista = []
    for lc in lancamentos:
        if lc.get("problemas","").strip():
            obs_lista.append(f"[{lc.get('data','')}] ⚠️ {lc['problemas']}")
        if lc.get("observacao","").strip():
            obs_lista.append(f"[{lc.get('data','')}] {lc['observacao']}")
    obs_txt = "\n".join(obs_lista[:10])
    if obs_txt:
        st.session_state["rel_observacoes_gerais"] = obs_txt

# Autoimportação real para o relatório mensal.
# Se houver lançamentos de campo no período, o relatório é alimentado automaticamente
# uma vez por combinação condomínio/mês/ano/quantidade/último lançamento.
if lancamentos_disponiveis:
    try:
        _ultimo_lc = lancamentos_disponiveis[-1] if lancamentos_disponiveis else {}
        _assinatura_auto = "|".join([
            str(nome_rel_atual),
            str(mes_ref),
            str(ano_ref),
            str(len(lancamentos_disponiveis)),
            str(_ultimo_lc.get("id_visita", "")),
            str(_ultimo_lc.get("data", "")),
            str(_ultimo_lc.get("operador", "")),
        ])

        if st.session_state.get("_rel_autoimport_assinatura") != _assinatura_auto:
            _importar_lancamentos(lancamentos_disponiveis)
            st.session_state["_rel_autoimport_assinatura"] = _assinatura_auto
            st.session_state["_rel_autoimport_msg"] = (
                f"✅ {len(lancamentos_disponiveis)} lançamento(s) de campo importado(s) automaticamente "
                f"para o relatório {mes_ref}/{ano_ref}."
            )
            st.rerun()
    except Exception as e:
        _log_sheets_erro("autoimportar_lancamentos_relatorio_mensal", e)

_auto_msg = st.session_state.pop("_rel_autoimport_msg", "")
if _auto_msg:
    st.success(_auto_msg)

if lancamentos_disponiveis:
    _total = len(lancamentos_disponiveis)
    _fonte = "📱 local + Sheets" if lancamentos_sheets else "📱 local"
    _periodo = f"{lancamentos_disponiveis[0].get('data','?')} → {lancamentos_disponiveis[-1].get('data','?')}"

    st.markdown(f"""
    <div style="border:1px solid rgba(20,120,60,0.3);border-radius:12px;padding:12px 16px;
    background:rgba(20,120,60,0.07);margin-bottom:12px;">
    <strong>📱 {_total} lançamento(s) de campo disponível(is) — {_fonte}</strong><br>
    <span style="font-size:0.85rem;color:#3a6a3a;">Período: {_periodo}</span>
    </div>
    """, unsafe_allow_html=True)

    imp1, imp2, imp3 = st.columns([1.5, 1.5, 1])
    with imp1:
        if st.button("📥 Importar lançamentos para o relatório", use_container_width=True, type="primary"):
            _importar_lancamentos(lancamentos_disponiveis)
            st.success(f"✅ {_total} lançamento(s) importado(s) com sucesso!")
            st.rerun()

    with imp2:
        if st.button("🗑️ Limpar lançamentos locais após gerar", use_container_width=True):
            if pasta_rel_atual and pasta_rel_atual.exists():
                dados_limpar = carregar_dados_condominio(pasta_rel_atual) or {}
                dados_limpar["lancamentos_campo"] = []
                salvar_dados_condominio(pasta_rel_atual, dados_limpar)
                st.success("Lançamentos locais limpos.")
                st.rerun()

    with imp3:
        with st.expander("👁 Ver lançamentos"):
            for lc in lancamentos_disponiveis:
                piscinas = lc.get("piscinas",[])
                if piscinas:
                    for p in piscinas:
                        st.caption(f"📅 {lc.get('data','?')} | {p.get('nome','Piscina')} | pH:{p.get('ph','–')} CRL:{p.get('cloro_livre','–')}")
                else:
                    st.caption(f"📅 {lc.get('data','?')} | Op:{lc.get('operador','–')} | pH:{lc.get('ph','–')} CRL:{lc.get('cloro_livre','–')}")
                if lc.get("problemas","").strip():
                    st.caption(f"   ⚠️ {lc['problemas']}")
                if lc.get("observacao","").strip():
                    st.caption(f"   📝 {lc['observacao']}")

else:
    if nome_rel_atual:
        st.info(f"Nenhum lançamento de campo encontrado para **{nome_rel_atual}**{f' no mês {mes_ref}/{ano_ref}' if mes_ref else ''}. O operador precisa registrar as visitas pelo modo campo.")


# Importação automática após carregar cliente cadastrado: evita depender do botão manual
# e alimenta o relatório de RT com as visitas já registradas no modo campo/Sheets.
if st.session_state.pop("_rel_auto_importar_cliente", False):
    if lancamentos_disponiveis:
        # Importa apenas se nao ha dados ja preenchidos na tabela — evita apagar trabalho manual
        _tem_dados_tabela = any(
            str(st.session_state.get(f"rel_analise_ph_{i}") or "").strip()
            for i in range(int(st.session_state.get("rel_analises_total", 12) or 12))
        )
        if not _tem_dados_tabela:
            _importar_lancamentos(lancamentos_disponiveis)
            st.success(f"✅ {len(lancamentos_disponiveis)} lançamento(s) importado(s) automaticamente.")
            st.rerun()
        else:
            st.info(f"ℹ️ {len(lancamentos_disponiveis)} lançamento(s) disponível(is). Clique em Importar para carregar (irá substituir dados atuais).")
    elif nome_rel_atual:
        st.warning(f"Cliente carregado, mas nenhum lançamento de visita foi encontrado para {nome_rel_atual} no período informado.")

st.markdown("**Dados do condomínio / local atendido**")
rd1, rd2 = st.columns(2)
with rd1:
    st.text_input("Condomínio / estabelecimento", key="rel_nome_condominio")
    if st.session_state.get("_rel_cep_fmt"):
        st.session_state["rel_cep"] = st.session_state.pop("_rel_cep_fmt")
    _rel_cep_c1, _rel_cep_c2 = st.columns([3, 1])
    with _rel_cep_c1:
        st.text_input("CEP", key="rel_cep", placeholder="00000-000",
            help="Digite o CEP e clique em 🔍 para preencher o endereço automaticamente")
    with _rel_cep_c2:
        st.markdown("<br>", unsafe_allow_html=True)
        _btn_rel_cep = st.button("🔍", key="btn_buscar_cep_rel", help="Buscar CEP")
    if _btn_rel_cep:
        _cep_v = re.sub(r"\D", "", st.session_state.get("rel_cep", ""))
        if len(_cep_v) == 8:
            with st.spinner("Buscando CEP..."):
                _dc = buscar_cep(_cep_v)
            if _dc:
                _end = ", ".join(p for p in [_dc.get("logradouro",""), _dc.get("bairro",""), f"{_dc.get('localidade','')}/{_dc.get('uf','')}", f"{_cep_v[:5]}-{_cep_v[5:]}"] if p)
                st.session_state["rel_endereco_condominio"] = _end
                st.session_state["_rel_cep_fmt"] = f"{_cep_v[:5]}-{_cep_v[5:]}"
                st.rerun()
            else:
                st.error("CEP não encontrado.")
        else:
            st.warning("Digite um CEP válido com 8 dígitos.")
    st.text_area("Endereço do local", key="rel_endereco_condominio", height=90)
    st.text_input("Representante / síndico / contato local", key="rel_representante")
with rd2:
    st.text_input(
        "CNPJ do condomínio / estabelecimento",
        key="rel_cnpj_condominio",
        on_change=lambda: st.session_state.__setitem__("rel_cnpj_condominio", formatar_cnpj(st.session_state.get("rel_cnpj_condominio", "")))
    )
    st.text_input("CPF/CNPJ do representante", key="rel_cpf_cnpj_representante", on_change=lambda: on_change_rel_documento_representante())
    st.file_uploader(
        "Upload de fotos do relatório",
        type=["png", "jpg", "jpeg", "webp"],
        accept_multiple_files=True,
        key="rel_fotos_upload",
        help="As fotos serão salvas na pasta do condomínio e inseridas no relatório."
    )
    st.caption("Esses dados podem ser preenchidos diretamente aqui, mesmo quando o relatório for avulso e sem contrato.")

r1, r2, r3, r4 = st.columns(4)
with r1:
    st.text_input("Mês de referência", key="rel_mes_referencia", placeholder="04")
with r2:
    st.text_input("Ano de referência", key="rel_ano_referencia", placeholder="2026")
with r3:
    opcoes_art = ["Emitida", "Não emitida", "Em tramitação"]
    status_atual_art = (st.session_state.get("rel_art_status") or "Emitida").strip()
    if status_atual_art not in opcoes_art:
        status_atual_art = "Emitida"
        st.session_state["rel_art_status"] = status_atual_art
    if st.session_state.get("rel_art_status_widget") not in opcoes_art:
        st.session_state["rel_art_status_widget"] = status_atual_art

    status_selecionado_art = st.selectbox(
        "Status da ART",
        opcoes_art,
        key="rel_art_status_widget",
    )

    if status_selecionado_art != st.session_state.get("rel_art_status"):
        st.session_state["rel_art_status"] = status_selecionado_art
        if status_selecionado_art != "Emitida":
            st.session_state["rel_art_numero"] = ""
            st.session_state["rel_art_inicio"] = ""
            st.session_state["rel_art_fim"] = ""
        st.rerun()
with r4:
    st.text_input("Data de emissão", key="rel_data_emissao", placeholder="dd/mm/aaaa", on_change=lambda: formatar_data_relatorio_chave("rel_data_emissao"))

r5, r6, r7, r8 = st.columns(4)
with r5:
    st.text_input("ART nº", key="rel_art_numero", on_change=on_change_rel_art_numero, disabled=st.session_state.get("rel_art_status") != "Emitida")
with r6:
    st.text_input("Vigência da ART - início", key="rel_art_inicio", placeholder="dd/mm/aaaa", on_change=lambda: formatar_data_relatorio_chave("rel_art_inicio"), disabled=st.session_state.get("rel_art_status") != "Emitida")
with r7:
    st.text_input("Vigência da ART - fim", key="rel_art_fim", placeholder="dd/mm/aaaa", on_change=lambda: formatar_data_relatorio_chave("rel_art_fim"), disabled=st.session_state.get("rel_art_status") != "Emitida")
with r8:
    opcoes_status_agua = ["CONFORME", "NÃO CONFORME", "EM CORREÇÃO"]
    status_atual = st.session_state.get("rel_status_agua", "CONFORME")
    if status_atual not in opcoes_status_agua:
        status_atual = "CONFORME"
    opcao_status = st.selectbox(
        "Status geral da água",
        opcoes_status_agua,
        index=opcoes_status_agua.index(status_atual),
        key="rel_status_agua_select",
    )
    st.session_state["rel_status_agua"] = opcao_status

if st.session_state.get("rel_art_status") != "Emitida":
    st.caption("Como a ART não está emitida, os campos ART nº e vigência ficam desabilitados e o relatório preencherá automaticamente como N/A, com observação institucional conforme o status selecionado.")

st.markdown("**Frequência de verificação para dimensionar linhas do relatório**")
_freq_rel_col1, _freq_rel_col2 = st.columns([1, 3])
with _freq_rel_col1:
    st.number_input("Verificações por semana", min_value=1, max_value=7, value=int(st.session_state.get("rel_verificacoes_semanais", 3) or 3), step=1, key="rel_verificacoes_semanais")
with _freq_rel_col2:
    _linhas_freq = calcular_linhas_analises_por_frequencia(st.session_state.get("rel_verificacoes_semanais", 3), st.session_state.get("rel_mes_referencia"), st.session_state.get("rel_ano_referencia"))
    st.caption(f"Base automática: {int(st.session_state.get('rel_verificacoes_semanais', 3) or 3)}x/semana → {_linhas_freq} linhas mínimas. O sistema aumenta se houver mais visitas importadas.")
if int(st.session_state.get("rel_analises_total", ANALISES_PADRAO) or ANALISES_PADRAO) < calcular_linhas_analises_por_frequencia(st.session_state.get("rel_verificacoes_semanais", 3)):
    garantir_campos_analises(calcular_linhas_analises_por_frequencia(st.session_state.get("rel_verificacoes_semanais", 3)))

c_auto1, c_auto2 = st.columns([1,2])
with c_auto1:
    if st.button("Preencher parecer automático", use_container_width=True):
        aplicar_textos_automaticos_relatorio()
with c_auto2:
    st.caption("O sistema calcula não conformidades e cloro combinado (cloraminas = CT - CL), preenche diagnóstico, observações e recomendações, e você ainda pode editar antes de gerar o relatório.")

st.text_area("Diagnóstico técnico", key="rel_diagnostico", height=120, placeholder="Será preenchido automaticamente conforme os parâmetros, mas permanece editável.")

st.markdown("**Análises físico-químicas**")

# _PAINEL_RASCUNHO_RELATORIO_RT_V1_
_relatorio_rt_renderizar_painel_rascunho()
_linhas_minimas_rel = calcular_linhas_analises_por_frequencia(st.session_state.get("rel_verificacoes_semanais", 3), st.session_state.get("rel_mes_referencia"), st.session_state.get("rel_ano_referencia"))
garantir_campos_analises(max(st.session_state.get("rel_analises_total", ANALISES_PADRAO), _linhas_minimas_rel))
ctrl_a1, ctrl_a2, ctrl_a3 = st.columns([1, 1.35, 2.25])
with ctrl_a1:
    if st.button("Adicionar análise extra", use_container_width=True):
        adicionar_analise_extra()
        st.rerun()
with ctrl_a2:
    if st.button("Carregar parâmetros usados pela última vez", use_container_width=True):
        nome_rel = (st.session_state.get("rel_nome_condominio") or st.session_state.get("nome_condominio") or "").strip()
        if nome_rel:
            pasta_rel = GENERATED_DIR / slugify_nome(nome_rel)
            dados_rel_salvos = carregar_dados_condominio(pasta_rel) if pasta_rel.exists() else None
            aplicar_parametros_ultimos_no_relatorio(dados_rel_salvos)
        else:
            aplicar_parametros_ultimos_no_relatorio(obter_snapshot_relatorio_independente())
        st.rerun()
with ctrl_a3:
    st.caption(f"{st.session_state.get('rel_analises_total', ANALISES_PADRAO)} linha(s) disponíveis neste relatório. Ao gerar o relatório, os parâmetros deste condomínio passam a ficar salvos como usados pela última vez.")
# Cabeçalho fixo para evitar que o navegador traduza siglas técnicas como CT, ALC ou CYA.
cab_cols = st.columns([1.05,0.7,0.8,0.95,1.15,0.95,1.25,1.1])
for _col, _label in zip(
    cab_cols,
    ["Data", "pH", "Cloro livre", "Cloro total", "Alcalinidade", "Dureza", "Ácido cianúrico", "Operador"],
):
    _col.caption(f"**{_label}**")

for i in range(int(st.session_state.get('rel_analises_total', ANALISES_PADRAO) or ANALISES_PADRAO)):
    cols = st.columns([1.05,0.7,0.8,0.95,1.15,0.95,1.25,1.1])
    cols[0].text_input(f"Data {i+1}", key=f"rel_analise_data_{i}", label_visibility="collapsed", placeholder="dd/mm/aaaa", on_change=lambda chave=f"rel_analise_data_{i}": formatar_data_relatorio_chave(chave))
    cols[1].text_input(f"pH {i+1}", key=f"rel_analise_ph_{i}", label_visibility="collapsed", placeholder="pH")
    cols[2].text_input(f"Cloro livre {i+1}", key=f"rel_analise_cl_{i}", label_visibility="collapsed", placeholder="Cloro livre")
    cols[3].text_input(f"Cloro total {i+1}", key=f"rel_analise_ct_{i}", label_visibility="collapsed", placeholder="Cloro total")
    cols[4].text_input(f"Alcalinidade {i+1}", key=f"rel_analise_alc_{i}", label_visibility="collapsed", placeholder="Alcalinidade")
    cols[5].text_input(f"Dureza {i+1}", key=f"rel_analise_dc_{i}", label_visibility="collapsed", placeholder="Dureza")
    cols[6].text_input(f"Ácido cianúrico {i+1}", key=f"rel_analise_cya_{i}", label_visibility="collapsed", placeholder="Ácido cianúrico")
    cols[7].text_input(f"Operador {i+1}", key=f"rel_analise_operador_{i}", label_visibility="collapsed", placeholder="Operador")

# _AUTOSAVE_RELATORIO_RT_APOS_LINHAS_V1_
try:
    _relatorio_rt_salvar_rascunho("autosave_apos_renderizar_linhas")
except Exception as _e_autosave_rel:
    st.caption(f"Autosave RT indisponível: {_e_autosave_rel}")

st.markdown("**Dosagens de produtos químicos**")
ctrl_d1, ctrl_d2 = st.columns([1.1, 2.4])
with ctrl_d1:
    if st.button("Carregar usados pela última vez", use_container_width=True):
        nome_rel = (st.session_state.get("rel_nome_condominio") or st.session_state.get("nome_condominio") or "").strip()
        if nome_rel:
            pasta_rel = GENERATED_DIR / slugify_nome(nome_rel)
            dados_rel_salvos = carregar_dados_condominio(pasta_rel) if pasta_rel.exists() else None
            aplicar_dosagens_ultimas_no_relatorio(dados_rel_salvos)
        else:
            aplicar_dosagens_ultimas_no_relatorio(obter_snapshot_relatorio_independente())
        st.rerun()
with ctrl_d2:
    st.caption("Ao gerar o relatório, as dosagens deste condomínio passam a ficar salvas como usadas pela última vez.")
for i in range(7):
    cols = st.columns([1.4,1.1,0.7,0.7,1.3])
    cols[0].text_input(f"Produto {i+1}", key=f"rel_dos_produto_{i}", label_visibility="collapsed", placeholder="Produto químico", on_change=_autosave_rascunho)
    cols[1].text_input(f"Lote {i+1}", key=f"rel_dos_lote_{i}", label_visibility="collapsed", placeholder="Fabricante / Lote", on_change=_autosave_rascunho)
    cols[2].text_input(f"Qtd {i+1}", key=f"rel_dos_qtd_{i}", label_visibility="collapsed", placeholder="Quantidade", on_change=_autosave_rascunho)
    cols[3].text_input(f"Un {i+1}", key=f"rel_dos_un_{i}", label_visibility="collapsed", placeholder="Unidade", on_change=_autosave_rascunho)
    cols[4].text_input(f"Finalidade {i+1}", key=f"rel_dos_finalidade_{i}", label_visibility="collapsed", placeholder="Finalidade técnica", on_change=_autosave_rascunho)
st.markdown("**Observações automáticas / editáveis**")
for i in range(5):
    st.text_area(f"Observação {i+1}", key=f"rel_obs_{i}", height=70)

st.markdown("**Recomendações técnicas**")
for i in range(5):
    cols = st.columns([2.0,0.8,1.0])
    cols[0].text_input(f"Recomendação {i+1}", key=f"rel_rec_texto_{i}", label_visibility="collapsed", placeholder="Recomendação técnica")
    cols[1].text_input(f"Prazo {i+1}", key=f"rel_rec_prazo_{i}", label_visibility="collapsed", placeholder="Prazo")
    cols[2].text_input(f"Responsável {i+1}", key=f"rel_rec_resp_{i}", label_visibility="collapsed", placeholder="Responsável")

st.markdown("**NBR 11238 — Segurança e higiene operacional**")
st.caption("Marque os requisitos verificados. Esses campos alimentam a coluna ‘Conforme?’ do relatório final.")
_nbr_cols = st.columns([1.6, 1.2, 1.6, 1.6, 1.2])
with _nbr_cols[0]:
    st.radio("Sinalização de profundidade visível", ["Sim", "Não"], key="rel_nbr11238_profundidade", horizontal=True)
with _nbr_cols[1]:
    st.radio("Retrolavagem do filtro", ["Sim", "Não"], key="rel_nbr11238_retrolavagem", horizontal=True)
with _nbr_cols[2]:
    st.radio("Limpeza de skimmers/decantadores", ["Sim", "Não"], key="rel_nbr11238_skimmers", horizontal=True)
with _nbr_cols[3]:
    st.radio("Área de circulação antiderrapante", ["Sim", "Não"], key="rel_nbr11238_circulacao", horizontal=True)
with _nbr_cols[4]:
    st.radio("Chuveiro antes do acesso", ["Sim", "Não"], key="rel_nbr11238_chuveiro", horizontal=True)

cx1, cx2, cx3 = st.columns(3)
with cx1:
    st.text_area("NBR 11238 — Evidência / observação", key="rel_nbr_11238", height=90, placeholder="Ex.: sinalização visível, retrolavagem realizada, área antiderrapante preservada...")
with cx2:
    st.text_area("NR-26 / GHS – Checklist", key="rel_nr_26", height=90, placeholder="FISPQs/FDS, rótulos GHS, sinalização, incompatibilidade, emergência...")
with cx3:
    st.text_area("NR-06 / EPI – Checklist", key="rel_nr_06", height=90, placeholder="Fornecimento, fiscalização e uso dos EPIs...")

st.markdown("**EPIs — preenchimento rápido**")
status_epi_opcoes = ["Conforme", "Pendente", "Não informado", "N/A"]
for rotulo, chave_base in [
    ("Luvas", "luvas"),
    ("Óculos", "oculos"),
    ("Respirador", "respirador"),
    ("Botas", "botas"),
]:
    c_status, c_ca = st.columns([1.2, 1])
    with c_status:
        st.selectbox(
            f"{rotulo} — status",
            options=status_epi_opcoes,
            key=f"rel_epi_{chave_base}_status",
        )
    with c_ca:
        st.text_input(f"{rotulo} — CA nº (opcional)", key=f"rel_epi_{chave_base}_ca")

st.caption("O relatório mensal pode ser gerado independentemente de contrato. Basta preencher os dados do local neste próprio módulo.")

# ---- Barra de rascunho ----
_nome_rasc = (st.session_state.get("rel_nome_condominio") or "").strip()
_pasta_rasc = GENERATED_DIR / slugify_nome(_nome_rasc) if _nome_rasc else None
_rasc_existente = carregar_rascunho_relatorio(_pasta_rasc) if _pasta_rasc and _pasta_rasc.exists() else None

col_rasc1, col_rasc2, col_rasc3 = st.columns([1.2, 1.2, 1.6])
with col_rasc1:
    if st.button("💾 Salvar rascunho", use_container_width=True):
        if _nome_rasc:
            _pasta_rasc.mkdir(parents=True, exist_ok=True)
            salvar_rascunho_relatorio(_pasta_rasc)
            st.success("Rascunho salvo! Os dados serão restaurados mesmo após reiniciar o sistema.")
        else:
            st.warning("Informe o nome do condomínio antes de salvar o rascunho.")

with col_rasc2:
    if _rasc_existente:
        if st.button(f"↩️ Restaurar rascunho ({_rasc_existente.get('salvo_em','?')})", use_container_width=True):
            aplicar_rascunho_no_formulario(_rasc_existente)
            st.success("Rascunho restaurado!")
            st.rerun()
    else:
        st.button("↩️ Sem rascunho salvo", disabled=True, use_container_width=True)

with col_rasc3:
    if _rasc_existente:
        st.markdown(
            f"<div style='font-size:0.82rem;color:#3a7a3a;padding:8px 0;'>"
            f"✅ Rascunho disponível — salvo em {_rasc_existente.get('salvo_em','?')}</div>",
            unsafe_allow_html=True
        )
    else:
        st.caption("Salve o rascunho para não perder dados ao reiniciar o sistema.")

st.markdown("---")
btn_col1, btn_col2 = st.columns([2, 1])
with btn_col1:
    rel_gerar = st.button("🚀 Gerar relatório mensal", use_container_width=True, type="primary")
with btn_col2:
    if st.button("🗑️ Limpar rascunho", use_container_width=True):
        if _pasta_rasc and _pasta_rasc.exists():
            excluir_rascunho_relatorio(_pasta_rasc)
            st.success("Rascunho excluído.")
            st.rerun()

st.markdown('</div>', unsafe_allow_html=True)


# =========================================
# =========================================
# FUNÇÕES DE PROCESSAMENTO DE DOCUMENTOS
# =========================================


# =========================================
# E-MAIL SMTP — AQUA GESTÃO
# =========================================

def _email_secrets_aqua() -> dict:
    """Lê configurações de e-mail do st.secrets sem expor senha no app."""
    try:
        cfg = dict(st.secrets.get("email", {}))
    except Exception:
        cfg = {}
    return {
        "smtp_host": str(cfg.get("smtp_host", "")).strip(),
        "smtp_port": int(cfg.get("smtp_port", 587) or 587),
        "smtp_user": str(cfg.get("smtp_user", "")).strip(),
        "smtp_password": str(cfg.get("smtp_password", "")).strip(),
        "remetente_nome": str(cfg.get("remetente_nome", "Aqua Gestão – Controle Técnico de Piscinas")).strip(),
        "logo_url": str(cfg.get("logo_url", "")).strip(),
        "reply_to": str(cfg.get("reply_to", cfg.get("smtp_user", ""))).strip(),
    }


def email_smtp_configurado() -> bool:
    cfg = _email_secrets_aqua()
    return bool(cfg["smtp_host"] and cfg["smtp_user"] and cfg["smtp_password"])


def assinatura_email_aqua_gestao() -> str:
    """Assinatura premium em HTML para todos os e-mails enviados pelo sistema."""
    import html as _html

    cfg = _email_secrets_aqua()
    logo_url = cfg.get("logo_url", "")
    if logo_url:
        logo_html = f"""
            <img src="{_html.escape(logo_url)}"
                 alt="Aqua Gestão"
                 style="width:112px; height:auto; display:block; border:0; outline:none; text-decoration:none;">
        """
    else:
        logo_html = """
            <div style="width:112px; height:86px; border-radius:18px; background:#EAF3FF; border:1px solid #D8E2EF;
                        display:flex; align-items:center; justify-content:center; text-align:center; color:#0B2E59;
                        font-size:15px; font-weight:800; line-height:1.15;">
                Aqua<br>Gestão
            </div>
        """

    return f"""
    <br><br>
    <table role="presentation" cellpadding="0" cellspacing="0" border="0"
           style="width:100%; max-width:760px; border-collapse:collapse; font-family:Arial, Helvetica, sans-serif; color:#1f2937;">
      <tr>
        <td style="border-top:4px solid #0B2E59; padding-top:16px;">
          <table role="presentation" cellpadding="0" cellspacing="0" border="0"
                 style="width:100%; border-collapse:collapse; background:#ffffff;">
            <tr>
              <td style="width:132px; vertical-align:top; padding-right:18px;">
                {logo_html}
              </td>
              <td style="vertical-align:top; border-left:3px solid #C9A227; padding-left:16px;">
                <div style="font-size:18px; font-weight:800; color:#0B2E59; line-height:1.25;">
                  Thyago Fernando da Silveira
                </div>
                <div style="font-size:14px; color:#374151; margin-top:3px;">
                  Técnico em Química | Responsável Técnico
                </div>
                <div style="font-size:13px; color:#374151; margin-top:3px;">
                  CRQ-MG 2ª Região | CRQ 024025748
                </div>
                <div style="height:1px; background:#D8E2EF; margin:10px 0;"></div>
                <div style="font-size:15px; font-weight:800; color:#145DA0;">
                  Aqua Gestão – Controle Técnico de Piscinas
                </div>
                <div style="font-size:13px; color:#4b5563; margin-top:4px; line-height:1.5;">
                  Responsabilidade Técnica • ART • Relatórios Técnicos • Controle de Piscinas Coletivas
                </div>
                <div style="font-size:13px; color:#4b5563; margin-top:8px; line-height:1.5;">
                  Uberlândia/MG<br>
                  WhatsApp: (34) 99291-3171<br>
                  E-mail: thyagosilveira@bemstarpiscinas.com.br
                </div>
                <div style="font-size:11px; color:#6b7280; margin-top:12px; line-height:1.4;">
                  Esta mensagem e seus anexos podem conter informações técnicas, contratuais e/ou confidenciais
                  destinadas exclusivamente ao(s) destinatário(s).
                </div>
              </td>
            </tr>
          </table>
        </td>
      </tr>
    </table>
    """


def montar_email_html_aqua(mensagem_texto: str) -> str:
    """Converte texto digitado em HTML e acrescenta assinatura premium."""
    import html as _html
    corpo = _html.escape(mensagem_texto or "").replace("\n", "<br>")
    return f"""
    <html>
      <body style="margin:0; padding:0; background:#ffffff;">
        <div style="font-family:Arial, Helvetica, sans-serif; font-size:14px; line-height:1.6; color:#1f2937;">
          {corpo}
          {assinatura_email_aqua_gestao()}
        </div>
      </body>
    </html>
    """


def listar_anexos_pasta_condominio(pasta_condominio: Path) -> list[Path]:
    """Lista anexos PDF/DOCX gerados para o condomínio, priorizando os mais recentes."""
    try:
        pasta = Path(pasta_condominio)
        if not pasta.exists():
            return []
        permitidos = {".pdf", ".docx"}
        arquivos = [p for p in pasta.iterdir() if p.is_file() and p.suffix.lower() in permitidos]
        arquivos.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        return arquivos[:30]
    except Exception:
        return []


def enviar_email_smtp_aqua(destinatario: str, assunto: str, mensagem_texto: str, anexos: list[Path] | None = None, cc: str = "", bcc: str = "") -> tuple[bool, str]:
    """Envia e-mail HTML com anexos usando SMTP configurado no Streamlit Secrets."""
    try:
        import smtplib
        import mimetypes
        from email.message import EmailMessage
        from email.utils import formataddr

        destinatario = (destinatario or "").strip()
        assunto = (assunto or "").strip()
        cc = (cc or "").strip()
        bcc = (bcc or "").strip()
        anexos = anexos or []

        if not destinatario:
            return False, "Informe o e-mail do destinatário."
        if not validar_email(destinatario):
            return False, "E-mail do destinatário inválido."
        if not assunto:
            return False, "Informe o assunto do e-mail."

        cfg = _email_secrets_aqua()
        if not email_smtp_configurado():
            return False, "SMTP não configurado. Configure [email] no Streamlit Secrets."

        msg = EmailMessage()
        msg["Subject"] = assunto
        msg["From"] = formataddr((cfg["remetente_nome"], cfg["smtp_user"]))
        msg["To"] = destinatario
        if cc:
            msg["Cc"] = cc
        if cfg.get("reply_to"):
            msg["Reply-To"] = cfg["reply_to"]

        texto_plano = (mensagem_texto or "").strip()
        texto_plano += "\n\n--\nThyago Fernando da Silveira\nTécnico em Química | CRQ-MG 2ª Região | CRQ 024025748\nAqua Gestão – Controle Técnico de Piscinas"
        msg.set_content(texto_plano)
        msg.add_alternative(montar_email_html_aqua(mensagem_texto), subtype="html")

        for anexo in anexos:
            caminho = Path(anexo)
            if not caminho.exists() or not caminho.is_file():
                continue
            ctype, encoding = mimetypes.guess_type(str(caminho))
            if ctype is None or encoding is not None:
                ctype = "application/octet-stream"
            maintype, subtype = ctype.split("/", 1)
            with open(caminho, "rb") as f:
                msg.add_attachment(
                    f.read(),
                    maintype=maintype,
                    subtype=subtype,
                    filename=caminho.name,
                )

        destinatarios_reais = [e.strip() for e in ([destinatario] + cc.split(",") + bcc.split(",")) if e.strip()]
        with smtplib.SMTP(cfg["smtp_host"], cfg["smtp_port"], timeout=30) as smtp:
            smtp.ehlo()
            smtp.starttls()
            smtp.ehlo()
            smtp.login(cfg["smtp_user"], cfg["smtp_password"])
            smtp.send_message(msg, from_addr=cfg["smtp_user"], to_addrs=destinatarios_reais)

        return True, "E-mail enviado com sucesso."
    except Exception as e:
        return False, f"Erro ao enviar e-mail: {type(e).__name__}: {e}"


def exibir_bloco_envio(
    nome_condominio: str,
    pasta_condominio: Path,
    whatsapp_cliente: str,
    email_cliente: str,
    mensagem: str,
):
    """Bloco de envio unificado para Aqua Gestão: WhatsApp, mailto e envio direto SMTP com assinatura premium."""
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader("📤 Enviar para o cliente")

    _slug_envio = slugify_nome(nome_condominio)

    msg_editada = st.text_area(
        "Mensagem",
        value=mensagem,
        height=220,
        key=f"aq_msg_envio_{_slug_envio}",
        label_visibility="collapsed",
    )
    componente_copiar(msg_editada)

    col_env1, col_env2, col_env3 = st.columns(3)

    with col_env1:
        _tel = (whatsapp_cliente or "").strip()
        if _tel:
            url_wa = link_whatsapp(_tel, msg_editada)
            st.link_button("💬 Abrir WhatsApp", url_wa, use_container_width=True)
        else:
            st.text_input("WhatsApp", placeholder="(34) 99999-9999", key=f"aq_tel_envio_{_slug_envio}")
            _tel2 = st.session_state.get(f"aq_tel_envio_{_slug_envio}", "").strip()
            if _tel2:
                st.link_button("💬 Abrir WhatsApp", link_whatsapp(_tel2, msg_editada), use_container_width=True)

    with col_env2:
        _eml = (email_cliente or "").strip()
        if _eml:
            assunto_mailto = f"Documentação Aqua Gestão – {nome_condominio}"
            url_mail = link_email(_eml, assunto_mailto, msg_editada)
            st.link_button("✉️ Abrir e-mail", url_mail, use_container_width=True)
        else:
            st.text_input("E-mail", placeholder="email@cliente.com.br", key=f"aq_email_envio_{_slug_envio}")
            _eml2 = st.session_state.get(f"aq_email_envio_{_slug_envio}", "").strip()
            if _eml2:
                assunto_mailto = f"Documentação Aqua Gestão – {nome_condominio}"
                st.link_button("✉️ Abrir e-mail", link_email(_eml2, assunto_mailto, msg_editada), use_container_width=True)

    with col_env3:
        if st.button("📂 Abrir pasta", use_container_width=True, key=f"aq_abrir_pasta_{_slug_envio}"):
            abrir_pasta_windows(pasta_condominio)

    with st.expander("📧 Enviar direto pelo sistema com assinatura premium Aqua Gestão", expanded=False):
        cfg_ok = email_smtp_configurado()
        if not cfg_ok:
            st.warning("SMTP ainda não está configurado. Configure o bloco [email] no Streamlit Secrets antes de enviar direto pelo sistema.")
            st.code(
                '[email]\n'
                'smtp_host = "smtp.gmail.com"\n'
                'smtp_port = 587\n'
                'smtp_user = "seuemail@gmail.com"\n'
                'smtp_password = "SENHA_DE_APP_DO_GMAIL"\n'
                'remetente_nome = "Aqua Gestão – Controle Técnico de Piscinas"\n'
                'logo_url = "https://link-publico-da-logo.png"\n'
                'reply_to = "seuemail@gmail.com"',
                language="toml",
            )
        else:
            st.success("SMTP configurado. Os e-mails sairão com assinatura premium Aqua Gestão.")

        email_padrao = (email_cliente or st.session_state.get(f"aq_email_envio_{_slug_envio}", "") or "").strip()
        dest = st.text_input("Destinatário", value=email_padrao, placeholder="email@condominio.com.br", key=f"aq_smtp_dest_{_slug_envio}")
        cc = st.text_input("CC opcional", value="", placeholder="administradora@exemplo.com.br", key=f"aq_smtp_cc_{_slug_envio}")
        bcc = st.text_input("BCC opcional", value="", placeholder="", key=f"aq_smtp_bcc_{_slug_envio}")
        assunto = st.text_input(
            "Assunto",
            value=f"Documentação Aqua Gestão – {nome_condominio}",
            key=f"aq_smtp_assunto_{_slug_envio}",
        )

        anexos_disponiveis = listar_anexos_pasta_condominio(pasta_condominio)
        if anexos_disponiveis:
            opcoes = [a.name for a in anexos_disponiveis]
            padrao = [a.name for a in anexos_disponiveis if a.suffix.lower() == ".pdf"][:8]
            selecionados = st.multiselect(
                "Anexos",
                options=opcoes,
                default=padrao,
                key=f"aq_smtp_anexos_{_slug_envio}",
                help="Por padrão o sistema seleciona PDFs. DOCX também pode ser anexado se necessário.",
            )
            anexos_final = [a for a in anexos_disponiveis if a.name in selecionados]
        else:
            st.info("Nenhum PDF/DOCX encontrado na pasta deste condomínio.")
            anexos_final = []

        st.caption("A assinatura premium será adicionada automaticamente no final do e-mail. A senha SMTP não fica no app.py; fica somente no Streamlit Secrets.")

        if st.button("📨 Enviar e-mail agora", type="primary", use_container_width=True, key=f"aq_smtp_enviar_{_slug_envio}", disabled=not cfg_ok):
            ok, retorno = enviar_email_smtp_aqua(dest, assunto, msg_editada, anexos_final, cc=cc, bcc=bcc)
            if ok:
                st.success(f"✅ {retorno}")
            else:
                st.error(retorno)

    st.markdown("</div>", unsafe_allow_html=True)


def gerar_contrato_bem_star_docx(
    nome_contratante: str,
    cpf_cnpj: str,
    endereco_contratante: str,
    valor_mensal: str,
    valor_extenso: str,
    dia_pagamento: str,
    forma_pagamento: str,
    prazo_contrato: str,
    data_inicio: str,
    data_fim: str,
    local_data_assinatura: str,
    piscinas_atendidas: str = "",
    produtos_incluidos: str = "",
) -> Path | None:
    """Gera contrato Bem Star usando template_bem_star.docx.
    Retorna o Path do arquivo gerado, ou None se o template estiver ausente."""
    if not TEMPLATE_BEM_STAR.exists():
        st.warning(
            "Template do contrato Bem Star não encontrado. "
            "Certifique-se de que `template_bem_star.docx` está na pasta do projeto."
        )
        return None

    placeholders = {
        "{{CNPJ_CONTRATADA}}": CNPJ_BEM_STAR,
        "{{ENDERECO_CONTRATADA}}": ENDERECO_BEM_STAR,
        "{{NOME_CONTRATANTE}}": nome_contratante,
        "{{CPF_CNPJ_CONTRATANTE}}": cpf_cnpj,
        "{{ENDERECO_CONTRATANTE}}": endereco_contratante,
        "{{VALOR_MENSAL}}": valor_mensal,
        "{{VALOR_MENSAL_EXTENSO}}": valor_extenso,
        "{{DIA_PAGAMENTO}}": dia_pagamento,
        "{{FORMA_PAGAMENTO}}": forma_pagamento,
        "{{PRAZO_CONTRATO}}": prazo_contrato,
        "{{DATA_INICIO_CONTRATO}}": data_inicio,
        "{{DATA_FIM_CONTRATO}}": data_fim,
        "{{LOCAL_DATA_ASSINATURA}}": local_data_assinatura,
        "{{PISCINAS_ATENDIDAS}}": piscinas_atendidas or "Conforme acordado entre as partes",
        "{{PRODUTOS_INCLUIDOS}}": produtos_incluidos or "Conforme proposta comercial",
    }

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_pasta = slugify_nome(nome_contratante)
    pasta = GENERATED_DIR / nome_pasta
    pasta.mkdir(parents=True, exist_ok=True)
    base_nome = limpar_nome_arquivo(f"Contrato_Limpeza_{nome_contratante}_{timestamp}")
    output_docx = pasta / f"{base_nome}.docx"
    output_pdf  = pasta / f"{base_nome}.pdf"

    gerar_documento(
        template_path=TEMPLATE_BEM_STAR,
        output_docx=output_docx,
        placeholders=placeholders,
        incluir_assinaturas=False,
    )

    ok_pdf, erro_pdf = converter_docx_para_pdf(output_docx, output_pdf)

    st.success("✅ Contrato Bem Star gerado com sucesso!")
    col1, col2 = st.columns(2)
    with col1:
        if output_docx.exists():
            with open(output_docx, "rb") as f:
                st.download_button(
                    "⬇️ Baixar DOCX",
                    data=f,
                    file_name=output_docx.name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
    with col2:
        if ok_pdf and output_pdf.exists():
            with open(output_pdf, "rb") as f:
                st.download_button(
                    "⬇️ Baixar PDF",
                    data=f,
                    file_name=output_pdf.name,
                    mime="application/pdf",
                    use_container_width=True,
                )
        elif erro_pdf:
            st.warning(f"PDF não gerado: {erro_pdf}")

    return output_docx


# =========================================
# CONTRATO RT — PDF PREMIUM REPORTLAB V12
# =========================================

def gerar_contrato_rt_pdf_reportlab(dados: dict) -> bytes:
    """Gera o contrato RT da Aqua Gestão em PDF premium, sem LibreOffice.

    Objetivo:
    - evitar o PDF sem logo/cores gerado pela conversão DOCX -> PDF;
    - manter layout institucional da Aqua Gestão;
    - gerar arquivo pronto para envio ao condomínio/jurídico.
    """
    import io
    import html
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
        KeepTogether,
        HRFlowable,
    )

    buffer = io.BytesIO()

    # Identidade Aqua Gestão
    AZUL_ESCURO = colors.HexColor("#0D2A4A")
    AZUL_MEDIO = colors.HexColor("#1565A8")
    AZUL_CLARO = colors.HexColor("#EAF4FF")
    AZUL_SUAVE = colors.HexColor("#F4FAFF")
    DOURADO = colors.HexColor("#C8960C")
    CINZA_TEXTO = colors.HexColor("#263442")
    CINZA_MEDIO = colors.HexColor("#5E6F82")
    CINZA_CLARO = colors.HexColor("#F5F7FA")
    BORDA = colors.HexColor("#D7E1EA")
    ALERTA = colors.HexColor("#FFF8E6")

    def _valor(*nomes, padrao=""):
        for nome in nomes:
            chaves = [nome, f"{{{{{nome}}}}}"]
            for chave in chaves:
                if chave in dados:
                    valor = dados.get(chave)
                    if valor is not None and str(valor).strip():
                        return str(valor).strip()
        return padrao

    def _esc(valor):
        return html.escape(str(valor or ""), quote=False).replace("\n", "<br/>")

    def _p(texto, estilo):
        return Paragraph(_esc(texto), estilo)

    def _pb(texto, estilo):
        # Texto controlado internamente. Permite <b>, <br/> e <font>.
        return Paragraph(str(texto or ""), estilo)

    nome_contratante = _valor("NOME_CONTRATANTE", "NOME_CONDOMINIO", padrao="Dado não informado")
    cpf_cnpj_contratante = _valor("CPF_CNPJ_CONTRATANTE", "CNPJ_CONDOMINIO", padrao="Dado não informado")
    endereco_contratante = _valor("ENDERECO_CONTRATANTE", "ENDERECO_CONDOMINIO", padrao="Dado não informado")
    representante = _valor("REPRESENTANTE_CONTRATANTE", "NOME_SINDICO", padrao="Dado não informado")
    cpf_sindico = _valor("CPF_SINDICO", padrao="Dado não informado")
    qualificacao_representante = _valor("QUALIFICACAO_REPRESENTANTE", padrao="Síndico")
    valor_mensal = _valor("VALOR_MENSAL", padrao="Dado não informado")
    valor_mensal_extenso = _valor("VALOR_MENSAL_EXTENSO", padrao="")
    dia_pagamento = _valor("DIA_PAGAMENTO", padrao="Dado não informado")
    forma_pagamento = _valor("FORMA_PAGAMENTO", padrao="Pix")
    frequencia_visitas = _valor("FREQUENCIA_VISITAS", padrao="1")
    data_inicio = _valor("DATA_INICIO_CONTRATO", "DATA_INICIO", padrao="Dado não informado")
    data_fim = _valor("DATA_FIM_CONTRATO", "DATA_FIM", padrao="")
    local_data_assinatura = _valor("LOCAL_DATA_ASSINATURA", "DATA_ASSINATURA", padrao="Uberlândia/MG")
    volumes_piscinas = _valor("VOLUMES_PISCINAS", padrao="Dado não informado")

    # Normalização visual de frequência para não repetir "(s)" de forma estranha.
    freq_txt = frequencia_visitas
    if freq_txt.strip() == "1":
        freq_frase = "1 (uma) visita técnica semanal"
    else:
        freq_frase = f"{freq_txt} visitas técnicas semanais"

    # Prepara logo. Se o PNG não existir, o PDF ainda sai bonito com marca textual.
    logo_path = encontrar_logo()
    logo_temp = None
    if logo_path and logo_path.exists():
        try:
            logo_temp = GENERATED_DIR / "_logo_aqua_reportlab.png"
            with Image.open(logo_path) as im:
                im = ImageOps.exif_transpose(im)
                if im.mode not in ("RGB", "RGBA"):
                    im = im.convert("RGBA")
                im.thumbnail((900, 360))
                # Mantém transparência quando possível.
                im.save(logo_temp, format="PNG")
        except Exception:
            logo_temp = logo_path

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.65 * cm,
        leftMargin=1.65 * cm,
        topMargin=2.35 * cm,
        bottomMargin=1.75 * cm,
        title="Contrato de Responsabilidade Técnica - Aqua Gestão",
        author="Aqua Gestão Controle Técnico Ltda",
    )

    styles = getSampleStyleSheet()
    s_title = ParagraphStyle(
        "AquaTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=19,
        textColor=AZUL_ESCURO,
        alignment=TA_CENTER,
        spaceAfter=5,
    )
    s_subtitle = ParagraphStyle(
        "AquaSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.7,
        leading=11,
        textColor=CINZA_MEDIO,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    s_h = ParagraphStyle(
        "AquaClauseTitle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=9.6,
        leading=12,
        textColor=AZUL_ESCURO,
        spaceBefore=7,
        spaceAfter=3,
        keepWithNext=True,
    )
    s_body = ParagraphStyle(
        "AquaBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.15,
        leading=11.2,
        textColor=CINZA_TEXTO,
        alignment=TA_JUSTIFY,
        spaceAfter=4.4,
    )
    s_body_left = ParagraphStyle(
        "AquaBodyLeft",
        parent=s_body,
        alignment=TA_LEFT,
    )
    s_small = ParagraphStyle(
        "AquaSmall",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7.3,
        leading=9.1,
        textColor=CINZA_TEXTO,
        alignment=TA_LEFT,
    )
    s_small_center = ParagraphStyle(
        "AquaSmallCenter",
        parent=s_small,
        alignment=TA_CENTER,
    )
    s_table_header = ParagraphStyle(
        "AquaTableHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=7.5,
        leading=9,
        textColor=colors.white,
        alignment=TA_CENTER,
    )
    s_table_cell = ParagraphStyle(
        "AquaTableCell",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7.25,
        leading=8.9,
        textColor=CINZA_TEXTO,
        alignment=TA_LEFT,
    )
    s_box_title = ParagraphStyle(
        "AquaBoxTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8.1,
        leading=10,
        textColor=AZUL_ESCURO,
        alignment=TA_LEFT,
        spaceAfter=2,
    )
    s_box = ParagraphStyle(
        "AquaBox",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7.55,
        leading=9.8,
        textColor=CINZA_TEXTO,
        alignment=TA_JUSTIFY,
    )

    def _header_footer(canvas, doc_obj):
        canvas.saveState()
        width, height = A4

        # Cabeçalho premium
        canvas.setFillColor(AZUL_ESCURO)
        canvas.rect(0, height - 1.55 * cm, width, 1.55 * cm, fill=1, stroke=0)

        canvas.setFillColor(DOURADO)
        canvas.rect(0, height - 1.61 * cm, width, 0.06 * cm, fill=1, stroke=0)

        # Logo ou marca textual
        if logo_temp and Path(logo_temp).exists():
            try:
                canvas.drawImage(
                    str(logo_temp),
                    1.30 * cm,
                    height - 1.28 * cm,
                    width=2.70 * cm,
                    height=0.92 * cm,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                canvas.setFillColor(colors.white)
                canvas.setFont("Helvetica-Bold", 11)
                canvas.drawString(1.35 * cm, height - 0.92 * cm, "AQUA GESTÃO")
        else:
            # Fallback com aparência de logo textual
            canvas.setFillColor(colors.white)
            canvas.setFont("Helvetica-Bold", 11)
            canvas.drawString(1.35 * cm, height - 0.92 * cm, "AQUA GESTÃO")
            canvas.setFillColor(DOURADO)
            canvas.circle(1.10 * cm, height - 0.90 * cm, 0.12 * cm, fill=1, stroke=0)

        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 8.6)
        canvas.drawRightString(
            width - 1.45 * cm,
            height - 0.58 * cm,
            "CONTRATO DE RESPONSABILIDADE TÉCNICA",
        )
        canvas.setFont("Helvetica", 7.1)
        canvas.drawRightString(
            width - 1.45 * cm,
            height - 0.96 * cm,
            "Aqua Gestão Controle Técnico Ltda | CRQ-MG 2ª Região | CRQ 024025748",
        )

        # Rodapé
        canvas.setStrokeColor(BORDA)
        canvas.setLineWidth(0.4)
        canvas.line(1.45 * cm, 1.13 * cm, width - 1.45 * cm, 1.13 * cm)

        canvas.setFillColor(CINZA_MEDIO)
        canvas.setFont("Helvetica", 6.6)
        canvas.drawString(
            1.45 * cm,
            0.74 * cm,
            "Aqua Gestão Controle Técnico Ltda | CNPJ 66.008.795/0001-92 | Uberlândia/MG | (34) 9777-7227",
        )
        canvas.drawRightString(width - 1.45 * cm, 0.74 * cm, f"Página {doc_obj.page}")

        canvas.restoreState()

    story = []

    # Capa resumida
    story.append(Spacer(1, 2 * mm))
    story.append(_pb("CONTRATO DE PRESTAÇÃO DE SERVIÇOS DE<br/>RESPONSABILIDADE TÉCNICA", s_title))
    story.append(_pb("Controle técnico de piscinas de uso coletivo · Instrumento particular de natureza civil e empresarial · Obrigação de meio", s_subtitle))
    story.append(HRFlowable(width="100%", thickness=0.8, color=DOURADO, spaceBefore=2, spaceAfter=7))

    partes = [
        [_pb("<b>CONTRATADA</b>", s_table_header), _pb("<b>CONTRATANTE</b>", s_table_header)],
        [
            _pb(
                "AQUA GESTÃO CONTROLE TÉCNICO LTDA<br/>"
                "CNPJ: 66.008.795/0001-92<br/>"
                "R. Benito Segatto, 201, Casa 02, Jardim Europa<br/>"
                "Uberlândia/MG — CEP 38.414-680<br/>"
                "Responsável Técnico: Thyago Fernando da Silveira<br/>"
                "CRQ-MG 2ª Região | CRQ 024025748",
                s_table_cell,
            ),
            _pb(
                f"{_esc(nome_contratante)}<br/>"
                f"CPF/CNPJ: {_esc(cpf_cnpj_contratante)}<br/>"
                f"Endereço: {_esc(endereco_contratante)}<br/>"
                f"Representante: {_esc(representante)}<br/>"
                f"CPF: {_esc(cpf_sindico)}<br/>"
                f"Qualificação: {_esc(qualificacao_representante)}",
                s_table_cell,
            ),
        ],
    ]
    t_partes = Table(partes, colWidths=[8.45 * cm, 8.45 * cm], repeatRows=1)
    t_partes.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), AZUL_ESCURO),
                ("GRID", (0, 0), (-1, -1), 0.35, BORDA),
                ("BACKGROUND", (0, 1), (-1, 1), AZUL_SUAVE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(t_partes)
    story.append(Spacer(1, 5 * mm))

    resumo = Table(
        [
            [_pb("<b>Frequência</b>", s_table_header), _pb("<b>Valor mensal</b>", s_table_header), _pb("<b>Vencimento</b>", s_table_header), _pb("<b>Forma</b>", s_table_header)],
            [_p(freq_frase, s_table_cell), _p(f"{valor_mensal} ({valor_mensal_extenso})", s_table_cell), _p(f"Dia {dia_pagamento}", s_table_cell), _p(forma_pagamento, s_table_cell)],
        ],
        colWidths=[4.8 * cm, 5.2 * cm, 3.0 * cm, 3.9 * cm],
    )
    resumo.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), AZUL_MEDIO),
                ("GRID", (0, 0), (-1, -1), 0.35, BORDA),
                ("BACKGROUND", (0, 1), (-1, 1), colors.white),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4.5),
            ]
        )
    )
    story.append(resumo)
    story.append(Spacer(1, 5 * mm))

    def box_info(titulo, texto, cor_fundo=AZUL_CLARO, cor_borda=AZUL_MEDIO):
        t = Table(
            [[_pb(f"<b>{_esc(titulo)}</b>", s_box_title)], [_p(texto, s_box)]],
            colWidths=[16.9 * cm],
        )
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), cor_fundo),
                    ("BOX", (0, 0), (-1, -1), 0.65, cor_borda),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 4 * mm))

    box_info(
        "Natureza técnica do contrato",
        "O presente instrumento trata de Responsabilidade Técnica para supervisão, orientação, controle documental, emissão de relatórios técnicos e conformidade operacional das piscinas de uso coletivo. Não substitui a execução operacional diária, salvo contratação específica em separado.",
    )

    box_info(
        "Finalidade perante o CRQ-MG",
        "Este instrumento também se destina à comprovação formal do vínculo técnico-profissional perante o Conselho Regional de Química de Minas Gerais — CRQ-MG, para fins de obtenção, renovação ou manutenção da Anotação/Certificado de Responsabilidade Técnica, sem prejuízo das obrigações civis e comerciais aqui pactuadas.",
        cor_fundo=colors.HexColor("#FFF8E6"),
        cor_borda=DOURADO,
    )

    def clausula(titulo, paragrafos):
        if isinstance(paragrafos, str):
            paragrafos = [paragrafos]
        bloco = [_pb(f"<b>{_esc(titulo)}</b>", s_h)]
        for ptxt in paragrafos:
            bloco.append(_p(ptxt, s_body))
        story.append(KeepTogether(bloco))

    clausula(
        "1. DAS PARTES",
        f"Pelo presente instrumento particular, de um lado AQUA GESTÃO CONTROLE TÉCNICO LTDA, pessoa jurídica de direito privado inscrita no CNPJ sob nº 66.008.795/0001-92, com atuação técnica vinculada ao Responsável Técnico Thyago Fernando da Silveira, CRQ-MG 2ª Região, CRQ 024025748, doravante denominada CONTRATADA; e, de outro lado, {nome_contratante}, inscrita no CPF/CNPJ sob nº {cpf_cnpj_contratante}, situada em {endereco_contratante}, representada por {representante}, CPF nº {cpf_sindico}, na qualidade de {qualificacao_representante}, doravante denominada CONTRATANTE, resolvem celebrar o presente Contrato de Responsabilidade Técnica.",
    )

    clausula(
        "2. DO OBJETO",
        [
            "O presente contrato tem por objeto a prestação, pela CONTRATADA, de serviços técnicos especializados de Responsabilidade Técnica voltados ao controle técnico da qualidade da água das piscinas de uso coletivo mantidas pela CONTRATANTE, mediante atuação técnica, consultiva, orientativa e fiscalizatória, nos limites deste instrumento.",
            f"Integram o objeto contratual: realização de {freq_frase}; análise técnica das condições gerais das piscinas, casa de máquinas, rotinas operacionais e aspectos relacionados ao tratamento químico; emissão de relatórios técnicos, notificações de não conformidade, orientações técnicas e pareceres; avaliação dos procedimentos adotados pela equipe operacional; elaboração ou revisão de POPs; recomendação de medidas corretivas, preventivas ou emergenciais; e acompanhamento documental para conformidade com normas aplicáveis, incluindo ABNT NBR 10818:2025 e ABNT NBR 10339.",
            "Fica expressamente ajustado que o contrato não inclui execução operacional diária ou extraordinária do tratamento da água; limpeza física; aspiração; escovação; peneiração; retrolavagem; operação rotineira de equipamentos; fornecimento de mão de obra operacional; fornecimento de produtos, insumos, reagentes, materiais, peças ou EPIs; manutenção corretiva ou preventiva de equipamentos; ou garantia de resultado contínuo dos parâmetros entre as visitas técnicas.",
            f"Piscina(s) objeto deste contrato: {volumes_piscinas}.",
        ],
    )

    clausula(
        "3. DA PERIODICIDADE E FORMA DE EXECUÇÃO",
        [
            f"A prestação dos serviços ocorrerá mediante {freq_frase}, sem caráter de plantão permanente, sem disponibilidade irrestrita e sem permanência contínua nas dependências da CONTRATANTE.",
            "As visitas poderão sofrer alteração por necessidade técnica, operacional, caso fortuito, força maior ou impedimento justificado, sem que isso configure inadimplemento, desde que haja remarcação em prazo razoável.",
        ],
    )

    clausula(
        "4. DOS PROCEDIMENTOS OPERACIONAIS PADRÃO — POPs",
        [
            "Os POPs, checklists, relatórios, notificações de não conformidade, orientações técnicas e demais documentos técnicos emitidos pela CONTRATADA passam a integrar o presente contrato para fins de rastreabilidade e comprovação técnica.",
            "A CONTRATANTE reconhece que o cumprimento integral e contínuo dos POPs e das orientações técnicas constitui obrigação contratual essencial. A inobservância, execução incompleta, tardia, inadequada ou recusa de cumprimento afastará a responsabilidade da CONTRATADA pelos efeitos daí decorrentes.",
        ],
    )

    clausula(
        "5. DAS OBRIGAÇÕES DA CONTRATADA",
        [
            "Prestar os serviços com zelo técnico, diligência profissional e observância das normas aplicáveis ao escopo contratado; manter sua regularidade técnica; manter profissional habilitado vinculado ao serviço; emitir documentos técnicos compatíveis; registrar orientações, recomendações e não conformidades; comunicar risco relevante, irregularidade sanitária, operacional ou documental; responder por erros técnicos, relatórios incorretos e omissões relevantes decorrentes de negligência, imprudência ou imperícia própria; e guardar sigilo sobre informações técnicas, operacionais, comerciais e internas da CONTRATANTE.",
        ],
    )

    clausula(
        "6. DAS OBRIGAÇÕES DA CONTRATANTE",
        [
            "Efetuar os pagamentos nos prazos ajustados; garantir acesso às piscinas, casa de máquinas, áreas técnicas, documentos e informações; manter equipe operacional própria ou terceiros habilitados para execução material das rotinas diárias; cumprir e fazer cumprir os POPs, relatórios e orientações técnicas; fornecer informações verdadeiras e atualizadas; comunicar ocorrência anormal, interdição, acidente, falha operacional, suspeita de contaminação ou evento com potencial impacto sanitário; e providenciar produtos, equipamentos, materiais, EPIs, peças, reagentes, acessórios e recursos humanos necessários à operação diária.",
        ],
    )

    clausula(
        "7. DOS HONORÁRIOS, REMUNERAÇÃO TÉCNICA E REGULARIDADE PERANTE O CRQ-MG",
        [
            f"Pelos serviços ordinários objeto deste contrato, a CONTRATANTE pagará à CONTRATADA o valor mensal de {valor_mensal} ({valor_mensal_extenso}).",
            "O valor previsto nesta cláusula corresponde à remuneração técnica mensal pela prestação dos serviços de Responsabilidade Técnica, abrangendo acompanhamento técnico, orientação especializada, registros, relatórios técnicos ordinários e demais obrigações previstas neste instrumento.",
            "As partes reconhecem que o presente contrato será utilizado para comprovação do vínculo técnico-profissional perante o Conselho Regional de Química de Minas Gerais — CRQ-MG, para fins de obtenção, renovação ou manutenção da Anotação/Certificado de Responsabilidade Técnica, conforme exigências documentais aplicáveis.",
            "Em razão da exigência documental do CRQ-MG para contratos de Responsabilidade Técnica firmados com profissional autônomo, a remuneração técnica prevista neste instrumento não poderá ser inferior ao salário mínimo vigente, devendo o contrato observar vigência mínima de 12 meses e conter remuneração expressa pelo serviço técnico prestado.",
            f"O vencimento ocorrerá no dia {dia_pagamento} de cada mês, mediante emissão de nota fiscal, recibo ou documento equivalente pela CONTRATADA. O pagamento será realizado por {forma_pagamento}.",
            "Eventual revisão de valores somente poderá ocorrer por instrumento formal entre as partes, desde que preservadas as exigências mínimas necessárias à regularidade do vínculo técnico perante o CRQ-MG e sem prejuízo das obrigações técnicas assumidas pela CONTRATADA.",
            "O atraso no pagamento sujeitará a CONTRATANTE à incidência de multa moratória de 2% sobre o valor em atraso, juros de mora de 1% ao mês pro rata die e correção monetária pelo IPCA, ou índice oficial que o substitua.",
        ],
    )

    clausula(
        "8. DOS ATENDIMENTOS, RESPOSTAS TÉCNICAS E SERVIÇOS EXTRAORDINÁRIOS",
        [
            "Estão incluídos no valor mensal os relatórios técnicos ordinários decorrentes das visitas contratadas, as orientações técnicas diretamente relacionadas aos parâmetros avaliados e as comunicações de não conformidade identificadas durante a execução regular do serviço.",
            "A CONTRATADA responderá às solicitações técnicas ordinárias da CONTRATANTE em até 2 dias úteis, desde que relacionadas ao escopo contratado e que não caracterizem atendimento emergencial, reunião extraordinária, elaboração de documento adicional ou deslocamento presencial fora da periodicidade contratada.",
            "Situações de risco sanitário imediato, água imprópria para uso, suspeita de contaminação, acidente, interdição ou condição crítica deverão ser comunicadas pela CONTRATANTE imediatamente. Nesses casos, a CONTRATADA deverá emitir orientação inicial ou comunicação técnica em até 24 horas após ciência da ocorrência, sem prejuízo da necessidade de atendimento extraordinário quando aplicável.",
            "Serão considerados serviços extraordinários, sujeitos a orçamento prévio: visitas adicionais, atendimentos emergenciais presenciais, reuniões com síndico, conselho, administradora ou jurídico, acompanhamento de fiscalização, treinamentos, elaboração de parecer técnico complementar, relatórios adicionais fora da rotina contratada e deslocamentos fora do cronograma.",
            "O orçamento dos serviços extraordinários será apresentado em até 48 horas úteis após a solicitação, contendo descrição do serviço, prazo estimado e valor aplicável. A execução dependerá de aprovação prévia da CONTRATANTE, salvo situação emergencial de risco sanitário relevante, devidamente registrada.",
        ],
    )

    clausula(
        "9. DO REAJUSTE",
        "Os honorários contratuais serão reajustados anualmente, no mês de aniversário do contrato, pelo IPCA/IBGE acumulado nos 12 meses anteriores ao reajuste, ou pelo índice que vier a substituí-lo. Na hipótese de índice negativo, os valores permanecerão inalterados até o período seguinte.",
    )

    clausula(
        "10. DA RESPONSABILIDADE TÉCNICA, LIMITES E CORRESPONSABILIDADE OPERACIONAL",
        [
            "A CONTRATADA responderá pelos atos técnicos praticados no âmbito do serviço contratado, incluindo erros técnicos, omissões relevantes, orientações inadequadas, relatórios técnicos incorretos ou falhas diretamente atribuíveis à sua atuação profissional, quando comprovado nexo causal entre a conduta técnica e o dano verificado.",
            "A Responsabilidade Técnica assumida pela CONTRATADA abrange a supervisão, orientação, avaliação, registro e comunicação técnica referentes ao tratamento químico e ao controle da qualidade da água das piscinas objeto deste contrato, nos limites da periodicidade contratada, das informações fornecidas e do acesso efetivamente disponibilizado pela CONTRATANTE.",
            "A CONTRATANTE permanece responsável pela execução operacional diária das rotinas recomendadas, pela disponibilização de equipe, produtos, equipamentos, EPIs, acesso às instalações, registros operacionais e cumprimento das orientações técnicas formalmente emitidas pela CONTRATADA.",
            "A responsabilidade da CONTRATADA não será afastada por cláusula contratual quando houver erro técnico próprio, negligência, imprudência, imperícia ou omissão relevante comprovada. Por outro lado, a CONTRATADA não responderá por danos decorrentes de descumprimento das orientações técnicas, omissão de informações, impedimento de acesso, falta de produtos, falhas de operação por terceiros ou decisões administrativas da CONTRATANTE contrárias às recomendações técnicas registradas.",
            "Havendo divergência técnica ou situação de risco sanitário, a CONTRATADA deverá comunicar formalmente a CONTRATANTE, indicando a não conformidade, o risco identificado e as medidas corretivas recomendadas.",
        ],
    )

    clausula(
        "11. DAS NÃO CONFORMIDADES, RISCO SANITÁRIO E INTERDIÇÃO TÉCNICA",
        [
            "Verificada situação que represente risco sanitário, operacional, normativo ou de segurança, a CONTRATADA poderá emitir notificação formal de não conformidade, com indicação das providências cabíveis e do prazo recomendado para correção.",
            "Quando a situação exigir, poderá recomendar restrição de uso ou interdição técnica parcial ou total da piscina até regularização das não conformidades identificadas. A decisão administrativa e a efetiva execução das medidas caberão à CONTRATANTE.",
        ],
    )

    clausula(
        "12. DA IMPOSSIBILIDADE DE EXECUÇÃO POR FATO ATRIBUÍVEL À CONTRATANTE",
        [
            "Caso a execução do serviço técnico seja impedida ou substancialmente prejudicada por fato atribuível à CONTRATANTE, incluindo ausência de acesso, falta de informações indispensáveis, inexistência de registros operacionais, ausência de produtos mínimos, impedimento de vistoria ou descumprimento reiterado das orientações técnicas, a CONTRATADA deverá registrar formalmente a ocorrência e comunicar a necessidade de regularização.",
            "Enquanto o impedimento for pontual ou parcial, a CONTRATADA continuará responsável pelas atividades que ainda puder executar dentro do escopo contratado, registrando no relatório técnico as limitações encontradas.",
            "Se o impedimento inviabilizar de forma relevante a prestação do serviço por mais de 15 dias, sem regularização pela CONTRATANTE após comunicação formal, a CONTRATADA poderá suspender proporcionalmente as atividades afetadas ou rescindir o contrato por justa causa, permanecendo devidos apenas os valores correspondentes aos serviços já prestados, documentos emitidos, deslocamentos realizados e obrigações técnicas efetivamente executadas até a data da suspensão ou rescisão.",
            "A suspensão integral da cobrança mensal somente será discutida quando houver impossibilidade total de execução do serviço não atribuível à CONTRATANTE ou por acordo formal entre as partes.",
        ],
    )

    clausula(
        "13. DA AUSÊNCIA DE VÍNCULO TRABALHISTA, OPERACIONAL OU SOCIETÁRIO",
        "O presente contrato possui natureza estritamente civil e empresarial, inexistindo vínculo empregatício, societário, associativo, de subordinação hierárquica ou de exclusividade entre a CONTRATANTE e os sócios, empregados, prepostos ou profissionais vinculados à CONTRATADA.",
    )

    clausula(
        "14. DA CONFIDENCIALIDADE",
        "As partes obrigam-se a manter sigilo sobre dados, documentos, informações técnicas, operacionais, comerciais, financeiras e estratégicas a que tiverem acesso em razão deste contrato. A obrigação permanecerá vigente por 2 anos após o término do contrato, ressalvadas exigências legais, judiciais, regulatórias ou administrativas.",
    )

    clausula(
        "15. DO PRAZO E DA RESCISÃO",
        [
            f"O presente contrato terá vigência mínima de 12 meses, iniciando-se em {data_inicio}{f' e com previsão informada até {data_fim}' if data_fim else ''}, em observância às exigências documentais do CRQ-MG para comprovação do vínculo de Responsabilidade Técnica.",
            "Após o prazo mínimo inicial, o contrato poderá ser renovado, prorrogado, ajustado ou rescindido conforme acordo entre as partes, preservadas as obrigações vencidas e os documentos técnicos já emitidos.",
            "O contrato poderá ser rescindido por mútuo acordo; por qualquer das partes mediante aviso prévio escrito de 30 dias; por inadimplemento; por descumprimento contratual relevante não sanado em 15 dias após notificação; por recusa reiterada da CONTRATANTE em cumprir POPs, orientações técnicas ou medidas sanitárias essenciais; por perda de confiança objetiva comprovada, obstrução de acesso ou ocultação de informações.",
        ],
    )

    clausula(
        "16. DA PROVA DOCUMENTAL E DAS COMUNICAÇÕES",
        "Serão considerados meios válidos de prova e comunicação: contratos, aditivos e anexos assinados física ou eletronicamente; relatórios, POPs e notificações técnicas; e-mails; mensagens eletrônicas; registros fotográficos, vídeos e protocolos. A ausência de assinatura da CONTRATANTE em relatório ou notificação técnica não afastará sua validade quando houver prova de envio, entrega ou ciência.",
    )

    clausula(
        "17. DA INTEGRALIDADE E HIERARQUIA DOS DOCUMENTOS",
        "Este contrato representa a integralidade do acordo entre as partes quanto ao seu objeto. Integram este instrumento, para todos os fins, os aditivos contratuais, POPs, checklists, relatórios técnicos, notificações de não conformidade e os Anexos I e II. Em caso de divergência interpretativa, prevalecerá: contrato principal; aditivos; notificações formais; POPs e Anexos; relatórios e checklists.",
    )

    clausula(
        "18. DO FORO",
        "Fica eleito o foro da Comarca de Uberlândia/MG, com renúncia expressa a qualquer outro, por mais privilegiado que seja, para dirimir controvérsias oriundas deste contrato.",
    )

    story.append(Spacer(1, 7 * mm))
    story.append(_p(local_data_assinatura, s_body_left))
    story.append(Spacer(1, 14 * mm))

    assinaturas = Table(
        [
            [
                _pb(
                    "_________________________________________<br/>"
                    "AQUA GESTÃO CONTROLE TÉCNICO LTDA<br/>"
                    "Thyago Fernando da Silveira<br/>"
                    "CRQ-MG 2ª Região | CRQ 024025748",
                    s_small_center,
                ),
                _pb(
                    f"_________________________________________<br/>"
                    f"{_esc(nome_contratante)}<br/>"
                    f"{_esc(representante)}<br/>"
                    f"{_esc(qualificacao_representante)}",
                    s_small_center,
                ),
            ]
        ],
        colWidths=[8.3 * cm, 8.3 * cm],
    )
    assinaturas.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(assinaturas)

    story.append(PageBreak())

    # Anexos
    story.append(_pb("ANEXO I — PARÂMETROS MÍNIMOS DE QUALIDADE DO SERVIÇO TÉCNICO", s_title))
    story.append(_pb("Relatórios técnicos ordinários emitidos no escopo da Responsabilidade Técnica.", s_subtitle))
    story.append(HRFlowable(width="100%", thickness=0.8, color=DOURADO, spaceBefore=2, spaceAfter=8))

    anexo_i_linhas = [
        ["Parâmetro mínimo", "Critério objetivo"],
        ["Visitas técnicas", "Realização das visitas técnicas conforme periodicidade contratada, com registro de data, horário, local avaliado e responsável técnico."],
        ["Parâmetros físico-químicos", "Registro dos parâmetros avaliados, incluindo, quando aplicável: pH, cloro livre, cloro total, alcalinidade total, temperatura, turbidez, ácido cianúrico e demais parâmetros pertinentes."],
        ["Relatório técnico ordinário", "Emissão em até 48 horas após a visita ou após o recebimento completo das informações necessárias à consolidação técnica."],
        ["Não conformidade crítica", "Comunicação formal em até 24 horas após sua constatação, quando houver risco sanitário, operacional ou de segurança relevante."],
        ["Medidas corretivas", "Indicação objetiva das medidas corretivas recomendadas, com prazo sugerido quando aplicável."],
        ["Limitações de avaliação", "Registro de ausência de dados, falta de acesso, ausência de produtos, limitação operacional, impossibilidade de avaliação completa ou impedimentos verificados."],
        ["Respostas técnicas", "Resposta às solicitações técnicas ordinárias em até 2 dias úteis, desde que estejam dentro do escopo contratado."],
        ["POPs", "Revisão ou emissão de POPs quando houver alteração relevante na rotina operacional ou identificação de falha recorrente."],
    ]
    t_anexo_i = Table(
        [[_pb(f"<b>{_esc(c)}</b>" if i == 0 else _esc(c), s_table_header if i == 0 else s_table_cell) for c in row] for i, row in enumerate(anexo_i_linhas)],
        colWidths=[4.2 * cm, 12.7 * cm],
        repeatRows=1,
    )
    t_anexo_i.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), AZUL_ESCURO),
                ("GRID", (0, 0), (-1, -1), 0.35, BORDA),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 1), (-1, 1), colors.white),
                ("BACKGROUND", (0, 2), (-1, 2), CINZA_CLARO),
                ("BACKGROUND", (0, 3), (-1, 3), colors.white),
                ("BACKGROUND", (0, 4), (-1, 4), CINZA_CLARO),
                ("BACKGROUND", (0, 5), (-1, 5), colors.white),
                ("BACKGROUND", (0, 6), (-1, 6), CINZA_CLARO),
                ("BACKGROUND", (0, 7), (-1, 7), colors.white),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4.5),
            ]
        )
    )
    story.append(t_anexo_i)
    story.append(Spacer(1, 8 * mm))

    story.append(_pb("ANEXO II — TABELA DE SERVIÇOS EXTRAORDINÁRIOS", s_title))
    story.append(_pb("Serviços fora do escopo ordinário mensal, mediante orçamento prévio.", s_subtitle))
    story.append(HRFlowable(width="100%", thickness=0.8, color=DOURADO, spaceBefore=2, spaceAfter=8))

    anexo_ii = [
        ["Serviço extraordinário", "Critério", "Observação"],
        ["Visita técnica extraordinária", "Orçamento prévio em até 48 horas úteis", "Aplicável fora da periodicidade regular contratada."],
        ["Atendimento emergencial presencial", "Conforme urgência, deslocamento e complexidade", "Aplicável em risco sanitário, água imprópria, acidente ou demanda crítica."],
        ["Reunião técnica com síndico, conselho, administradora ou jurídico", "Conforme duração e preparação técnica", "Pode envolver histórico de parâmetros, orientação normativa e alinhamento documental."],
        ["Treinamento de operador ou equipe local", "Conforme carga horária", "Pode envolver POPs, GHS, EPIs, medição de parâmetros e segurança química."],
        ["Relatório técnico complementar", "Conforme escopo solicitado", "Documento adicional fora do relatório técnico mensal ordinário."],
        ["Acompanhamento de fiscalização oficial", "Conforme demanda e urgência", "Apoio técnico documental em fiscalização, notificação ou exigência formal."],
    ]
    t_anexo_ii = Table(
        [[_pb(f"<b>{_esc(c)}</b>" if i == 0 else _esc(c), s_table_header if i == 0 else s_table_cell) for c in row] for i, row in enumerate(anexo_ii)],
        colWidths=[5.2 * cm, 5.4 * cm, 6.3 * cm],
        repeatRows=1,
    )
    t_anexo_ii.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), AZUL_ESCURO),
                ("GRID", (0, 0), (-1, -1), 0.35, BORDA),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, CINZA_CLARO]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4.5),
            ]
        )
    )
    story.append(t_anexo_ii)
    story.append(Spacer(1, 5 * mm))

    box_info(
        "Regra de execução",
        "Os serviços extraordinários somente serão executados mediante solicitação da CONTRATANTE e aprovação prévia do orçamento, salvo situações emergenciais em que houver risco sanitário ou operacional relevante e necessidade de comunicação técnica imediata.",
        cor_fundo=ALERTA,
        cor_borda=DOURADO,
    )

    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf


def salvar_contrato_rt_pdf_premium_reportlab(dados: dict, output_pdf: Path) -> tuple[bool, str | None]:
    """Salva o PDF premium do contrato RT e retorna status/erro."""
    try:
        output_pdf.parent.mkdir(parents=True, exist_ok=True)
        pdf_bytes = gerar_contrato_rt_pdf_reportlab(dados)
        output_pdf.write_bytes(pdf_bytes)
        return True, None
    except Exception as e:
        return False, str(e)


# =========================================
# TERMOS DE CIÊNCIA — RT / EPIs
# =========================================

def gerar_contrato_e_aditivo():
    """Gera somente o contrato RT/ART da Aqua Gestão.

    Fluxo v15: removida a geração de aditivo de desconto da tela principal de RT,
    para manter o contrato limpo para CRQ/ART e alinhado ao item 3.2.
    """
    email_cliente = st.session_state.email_cliente.strip()
    dados = {
        "DATA_ASSINATURA": (st.session_state.get("data_assinatura") or "").strip(),
        "NOME_CONDOMINIO": (st.session_state.get("nome_condominio") or "").strip(),
        "CNPJ_CONDOMINIO": (st.session_state.get("cnpj_condominio") or "").strip(),
        "ENDERECO_CONDOMINIO": (st.session_state.get("endereco_condominio") or "").strip(),
        "NOME_SINDICO": (st.session_state.get("nome_sindico") or "").strip(),
        "CPF_SINDICO": (st.session_state.get("cpf_sindico") or "").strip(),
        "VALOR_MENSAL": valor_para_template((st.session_state.get("valor_mensal") or "").strip()),
        "DATA_INICIO": (st.session_state.get("data_inicio") or "").strip(),
        "DATA_FIM": (st.session_state.get("data_fim") or "").strip(),
    }
    erros = validar_para_geracao(dados, email_cliente)

    if erros:
        st.error("Corrija os campos antes de gerar o contrato:")
        for erro in erros:
            st.write(f"- {erro}")
        return

    try:
        nome_condominio = st.session_state.nome_condominio.strip()
        nome_sindico = st.session_state.nome_sindico.strip()
        whatsapp_cliente = st.session_state.whatsapp_cliente.strip()

        nome_pasta = slugify_nome(nome_condominio)
        pasta_condominio = GENERATED_DIR / nome_pasta
        pasta_condominio.mkdir(parents=True, exist_ok=True)

        salvar_dados_condominio(pasta_condominio, salvar_snapshot_formulario())

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_nome_contrato = limpar_nome_arquivo(f"Contrato_RT_{nome_condominio}_{timestamp}")

        contrato_docx = pasta_condominio / f"{base_nome_contrato}.docx"
        contrato_pdf_premium = pasta_condominio / f"{base_nome_contrato}_PDF_PREMIUM_AQUA_GESTAO.pdf"

        placeholders = {
            "{{CNPJ_CONTRATADA}}": "66.008.795/0001-92",
            "{{ENDERECO_CONTRATADA}}": "R. Benito Segatto, nº 201, Casa 02, Jardim Europa, Uberlândia/MG, CEP 38.414-680",
            "{{NOME_CONTRATANTE}}": st.session_state.nome_condominio.strip(),
            "{{CPF_CNPJ_CONTRATANTE}}": st.session_state.cnpj_condominio.strip(),
            "{{ENDERECO_CONTRATANTE}}": st.session_state.endereco_condominio.strip(),
            "{{REPRESENTANTE_CONTRATANTE}}": (st.session_state.get("nome_sindico") or "").strip(),
            "{{CPF_SINDICO}}": (st.session_state.get("cpf_sindico") or "").strip(),
            "{{QUALIFICACAO_REPRESENTANTE}}": (st.session_state.get("cargo_sindico") or "Síndico").strip(),
            "{{VOLUMES_PISCINAS}}": st.session_state.get("volumes_piscinas", ""),
            "{{FREQUENCIA_VISITAS}}": (st.session_state.get("frequencia_visitas") or "1 (uma)").split(" ")[0],
            "{{VALOR_MENSAL}}": valor_para_template(st.session_state.valor_mensal.strip()),
            "{{VALOR_MENSAL_EXTENSO}}": st.session_state.get("valor_mensal_extenso", ""),
            "{{DIA_PAGAMENTO}}": st.session_state.get("dia_pagamento", ""),
            "{{FORMA_PAGAMENTO}}": st.session_state.get("forma_pagamento", "Pix"),
            "{{MULTA_ATRASO}}": st.session_state.get("multa_atraso", ""),
            "{{JUROS_ATRASO}}": st.session_state.get("juros_atraso", ""),
            "{{PRAZO_CONTRATO}}": st.session_state.get("prazo_contrato", ""),
            "{{DATA_INICIO_CONTRATO}}": st.session_state.data_inicio.strip(),
            "{{DATA_FIM_CONTRATO}}": st.session_state.data_fim.strip(),
            "{{LOCAL_DATA_ASSINATURA}}": f"Uberlândia/MG, {st.session_state.data_assinatura.strip()}",
        }

        with st.spinner("Gerando contrato DOCX editável..."):
            gerar_documento(
                template_path=TEMPLATE_CONTRATO,
                output_docx=contrato_docx,
                placeholders=placeholders,
                incluir_assinaturas=False,
            )

        with st.spinner("Gerando PDF Premium Aqua Gestão..."):
            ok_contrato_premium, erro_contrato_premium = salvar_contrato_rt_pdf_premium_reportlab(
                placeholders,
                contrato_pdf_premium,
            )

        registrar_documento_manifest(
            pasta_condominio=pasta_condominio,
            nome_condominio=nome_condominio,
            tipo="Contrato RT — Aqua Gestão",
            arquivo_docx=contrato_docx,
            arquivo_pdf=contrato_pdf_premium,
            pdf_gerado=ok_contrato_premium,
            erro_pdf=erro_contrato_premium,
            dados_utilizados=dados,
            extras={"fluxo": "contrato_unico_sem_aditivo", "crq_item_3_2": True},
        )

        st.session_state.ultima_pasta_gerada = str(pasta_condominio)
        st.session_state.ultimos_docs_gerados = {
            "contrato_docx": str(contrato_docx) if contrato_docx.exists() else None,
            "contrato_pdf_premium": str(contrato_pdf_premium) if ok_contrato_premium and contrato_pdf_premium.exists() else None,
        }

        st.success("Contrato RT gerado com sucesso.")

        if not encontrar_logo():
            st.warning(
                "Atenção: não encontrei o arquivo de logo da Aqua Gestão na pasta do app. "
                "O PDF Premium foi gerado com marca textual. Para sair com a logo oficial, mantenha o arquivo "
                "`aqua_gestao_logo.png` na mesma pasta do app.py ou na pasta assets."
            )

        st.info(
            "Fluxo CRQ/ART: foi gerado somente o contrato único de Responsabilidade Técnica, "
            "sem aditivo de desconto e sem valor reduzido."
        )

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.subheader("Arquivos gerados")

        c1, c2 = st.columns(2)

        with c1:
            st.markdown("**Contrato RT — Aqua Gestão**")
            if contrato_docx.exists():
                with open(contrato_docx, "rb") as f:
                    st.download_button(
                        "Baixar DOCX editável do contrato",
                        data=f,
                        file_name=contrato_docx.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

        with c2:
            st.markdown("**PDF para envio**")
            if ok_contrato_premium and contrato_pdf_premium.exists():
                with open(contrato_pdf_premium, "rb") as f:
                    st.download_button(
                        "📄 Baixar PDF Premium Aqua Gestão",
                        data=f,
                        file_name=contrato_pdf_premium.name,
                        mime="application/pdf",
                        use_container_width=True,
                    )
            else:
                st.error(f"PDF Premium Aqua Gestão não gerado. Erro: {erro_contrato_premium}")

        mensagem = montar_mensagem_envio(
            nome_condominio=nome_condominio,
            nome_sindico=nome_sindico,
            caminho_contrato_pdf=contrato_pdf_premium if contrato_pdf_premium.exists() else None,
            caminho_aditivo_pdf=None,
        )

        exibir_bloco_envio(
            nome_condominio=nome_condominio,
            pasta_condominio=pasta_condominio,
            whatsapp_cliente=whatsapp_cliente,
            email_cliente=email_cliente,
            mensagem=mensagem,
        )

    except Exception as e:
        st.error(f"Erro na geração do contrato: {e}")


def gerar_somente_aditivo_rapido():
    email_cliente = st.session_state.email_cliente.strip()
    dados = {
        "DATA_ASSINATURA": (st.session_state.get("data_assinatura") or "").strip(),
        "NOME_CONDOMINIO": (st.session_state.get("nome_condominio") or "").strip(),
        "CNPJ_CONDOMINIO": (st.session_state.get("cnpj_condominio") or "").strip(),
        "ENDERECO_CONDOMINIO": (st.session_state.get("endereco_condominio") or "").strip(),
        "NOME_SINDICO": (st.session_state.get("nome_sindico") or "").strip(),
        "CPF_SINDICO": (st.session_state.get("cpf_sindico") or "").strip(),
        "VALOR_MENSAL": valor_para_template((st.session_state.get("valor_mensal") or "").strip()),
        "VALOR_ADITIVO": valor_para_template((st.session_state.get("valor_aditivo") or "").strip()),
        "DATA_INICIO": (st.session_state.get("data_inicio") or "").strip(),
        "DATA_FIM": (st.session_state.get("data_fim") or "").strip(),
    }
    erros = validar_para_geracao(dados, email_cliente)

    if erros:
        st.error("Corrija os campos antes de gerar o aditivo rápido:")
        for erro in erros:
            st.write(f"- {erro}")
        return

    try:
        nome_condominio = st.session_state.nome_condominio.strip()
        nome_sindico = st.session_state.nome_sindico.strip()
        whatsapp_cliente = st.session_state.whatsapp_cliente.strip()

        nome_pasta = slugify_nome(nome_condominio)
        pasta_condominio = GENERATED_DIR / nome_pasta
        pasta_condominio.mkdir(parents=True, exist_ok=True)

        salvar_dados_condominio(pasta_condominio, salvar_snapshot_formulario())

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_nome_aditivo = limpar_nome_arquivo(f"Aditivo_RT_{nome_condominio}_{timestamp}")

        aditivo_docx = pasta_condominio / f"{base_nome_aditivo}.docx"
        aditivo_pdf = pasta_condominio / f"{base_nome_aditivo}.pdf"

        placeholders = {
            "{{CNPJ_CONTRATADA}}": "66.008.795/0001-92",
            "{{ENDERECO_CONTRATADA}}": "R. Benito Segatto, nº 201, Casa 02, Jardim Europa, Uberlândia/MG, CEP 38.414-680",
            "{{NOME_CONTRATANTE}}": st.session_state.nome_condominio.strip(),
            "{{CPF_CNPJ_CONTRATANTE}}": st.session_state.cnpj_condominio.strip(),
            "{{ENDERECO_CONTRATANTE}}": st.session_state.endereco_condominio.strip(),
            "{{REPRESENTANTE_CONTRATANTE}}": (st.session_state.get("nome_sindico") or "").strip(),
            "{{CPF_SINDICO}}": (st.session_state.get("cpf_sindico") or "").strip(),
            "{{QUALIFICACAO_REPRESENTANTE}}": (st.session_state.get("cargo_sindico") or "Síndico").strip(),
            "{{VOLUMES_PISCINAS}}": st.session_state.get("volumes_piscinas", ""),
            "{{FREQUENCIA_VISITAS}}": (st.session_state.get("frequencia_visitas") or "1 (uma)").split(" ")[0],
            "{{VALOR_MENSAL}}": valor_para_template(st.session_state.valor_mensal.strip()),
            "{{VALOR_MENSAL_EXTENSO}}": st.session_state.get("valor_mensal_extenso", ""),
            "{{DIA_PAGAMENTO}}": st.session_state.get("dia_pagamento", ""),
            "{{FORMA_PAGAMENTO}}": st.session_state.get("forma_pagamento", "Pix"),
            "{{MULTA_ATRASO}}": st.session_state.get("multa_atraso", ""),
            "{{JUROS_ATRASO}}": st.session_state.get("juros_atraso", ""),
            "{{PRAZO_CONTRATO}}": st.session_state.get("prazo_contrato", ""),
            "{{DATA_INICIO_CONTRATO}}": st.session_state.data_inicio.strip(),
            "{{DATA_FIM_CONTRATO}}": st.session_state.data_fim.strip(),
            "{{LOCAL_DATA_ASSINATURA}}": f"Uberlândia/MG, {st.session_state.data_assinatura.strip()}",
        }

        with st.spinner("Gerando aditivo rápido..."):
            gerar_documento(
                template_path=TEMPLATE_ADITIVO,
                output_docx=aditivo_docx,
                placeholders=placeholders,
                incluir_assinaturas=False,
            )

        # PDF Premium ReportLab — padrão visual Aqua Gestão (sem LibreOffice)
        ok_aditivo, erro_aditivo = salvar_aditivo_rt_pdf_premium_reportlab(placeholders, aditivo_pdf)

        registrar_documento_manifest(
            pasta_condominio=pasta_condominio,
            nome_condominio=nome_condominio,
            tipo="Aditivo",
            arquivo_docx=aditivo_docx,
            arquivo_pdf=aditivo_pdf,
            pdf_gerado=ok_aditivo,
            erro_pdf=erro_aditivo,
            dados_utilizados=dados,
        )

        st.session_state.ultima_pasta_gerada = str(pasta_condominio)
        st.session_state.ultimos_docs_gerados = st.session_state.get("ultimos_docs_gerados") or {}
        st.session_state.ultimos_docs_gerados.update({
            "aditivo_docx": str(aditivo_docx) if aditivo_docx.exists() else None,
            "aditivo_pdf": str(aditivo_pdf) if ok_aditivo and aditivo_pdf.exists() else None,
        })

        st.success("Aditivo rápido gerado com sucesso.")

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.subheader("Arquivo gerado")

        c1, c2 = st.columns(2)

        with c1:
            if aditivo_docx.exists():
                with open(aditivo_docx, "rb") as f:
                    st.download_button(
                        "Baixar DOCX do aditivo",
                        data=f,
                        file_name=aditivo_docx.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

        with c2:
            if ok_aditivo and aditivo_pdf.exists():
                with open(aditivo_pdf, "rb") as f:
                    st.download_button(
                        "📄 Baixar PDF Premium Aqua Gestão",
                        data=f,
                        file_name=aditivo_pdf.name,
                        mime="application/pdf",
                        use_container_width=True,
                    )
            else:
                st.warning(f"PDF do aditivo não gerado. Erro: {erro_aditivo}")


        mensagem = montar_mensagem_envio(
            nome_condominio=nome_condominio,
            nome_sindico=nome_sindico,
            caminho_contrato_pdf=None,
            caminho_aditivo_pdf=aditivo_pdf if aditivo_pdf.exists() else None,
        )

        exibir_bloco_envio(
            nome_condominio=nome_condominio,
            pasta_condominio=pasta_condominio,
            whatsapp_cliente=whatsapp_cliente,
            email_cliente=email_cliente,
            mensagem=mensagem,
        )

    except Exception as e:
        st.error(f"Erro na geração do aditivo rápido: {e}")


# =========================================
# PROCESSAMENTO
# =========================================

if gerar:
    gerar_contrato_e_aditivo()

if gerar_aditivo_rapido:
    gerar_somente_aditivo_rapido()

if rel_gerar:
    ok_rel, msg_rel = gerar_relatorio_mensal()
    if ok_rel:
        st.success(msg_rel)
    else:
        st.error(msg_rel)

# =========================================
# RODAPÉ
# =========================================

st.markdown("---")
st.caption(
    f"{APP_TITLE} • {RESPONSAVEL_TÉCNICO} • {CRQ} • Versão v5_relatorio_premium_aqua"
)



